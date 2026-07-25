#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
大学生申报书 docx 生成脚本

按统一格式标准生成 Word 文档：
- A4 纸张，页边距 2.5cm
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 一级标题：黑体三号，居中
- 二级标题：黑体小三，左对齐
- 三级标题：宋体四号加粗
- 表格：宋体五号，居中

使用方式：
    from build_docx import ApplicationDocBuilder
    builder = ApplicationDocBuilder()
    builder.build_innovation_project(data, output_path)

或命令行：
    python build_docx.py --type innovation --data data.json --out output.docx
"""

import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_TIMES = "Times New Roman"

# 字号
SIZE_ER = Pt(22)       # 二号
SIZE_XIAO_ER = Pt(18)  # 小二
SIZE_SAN = Pt(16)      # 三号
SIZE_XIAO_SAN = Pt(15) # 小三
SIZE_SI = Pt(14)       # 四号
SIZE_XIAO_SI = Pt(12)  # 小四
SIZE_WU = Pt(10.5)     # 五号
SIZE_XIAO_WU = Pt(9)   # 小五


# ============================================================
# 工具函数
# ============================================================

def set_cell_font(cell, font_name=FONT_SONG, font_size=SIZE_WU, bold=False):
    """设置单元格内所有文字字体"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = bold
            # 设置中文字体
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, font_name=FONT_SONG, font_size=SIZE_XIAO_SI, bold=False, color=None):
    """设置 run 的字体（含中英文）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 设置中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: int = 0,
    space_after: int = 0,
):
    """添加带格式的段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)

    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc,
        text,
        font_name=FONT_HEI,
        font_size=SIZE_SAN,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
        space_before=12,
        space_after=12,
    )


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(
        doc,
        text,
        font_name=FONT_HEI,
        font_size=SIZE_XIAO_SAN,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6,
        space_after=6,
    )


def add_heading_level3(doc, text: str):
    """三级标题：宋体四号加粗"""
    return add_paragraph_with_format(
        doc,
        text,
        font_name=FONT_SONG,
        font_size=SIZE_SI,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6,
        space_after=3,
    )


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc,
        text,
        font_name=FONT_SONG,
        font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent,
        line_spacing=1.5,
    )


def add_table_from_data(doc, headers: List[str], rows: List[List[str]]):
    """从数据创建表格，自动应用规范格式"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_font(hdr_cells[i], font_name=FONT_SONG, font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 数据行
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
            set_cell_font(cells[j], font_name=FONT_SONG, font_size=SIZE_WU)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 表格后空一行
    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面、页边距"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_page_number(doc):
    """在页脚添加页码"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 插入页码字段
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


# ============================================================
# 文档构建器
# ============================================================

class ApplicationDocBuilder:
    """大学生申报书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)

        # 设置默认字体
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SONG)

    # ------------------------------------------------------------
    # 通用工具
    # ------------------------------------------------------------

    def add_cover_title(self, title: str, subtitle: str = ""):
        """封面标题：黑体二号，居中"""
        # 顶部空 2 行
        for _ in range(2):
            self.doc.add_paragraph()

        add_paragraph_with_format(
            self.doc,
            title,
            font_name=FONT_HEI,
            font_size=SIZE_ER,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_before=12,
            space_after=12,
        )
        if subtitle:
            add_paragraph_with_format(
                self.doc,
                subtitle,
                font_name=FONT_HEI,
                font_size=SIZE_SAN,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent=False,
                space_after=24,
            )

    def add_cover_info(self, info_items: List[Dict[str, str]]):
        """封面信息：居中、下划线"""
        for item in info_items:
            label = item.get("label", "")
            value = item.get("value", "")
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)

            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI, font_size=SIZE_SI, bold=True)

            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True

    def add_page_break(self):
        """添加分页符"""
        from docx.enum.text import WD_BREAK
        p = self.doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    def add_table(self, headers, rows):
        return add_table_from_data(self.doc, headers, rows)

    def add_h1(self, text):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text):
        return add_heading_level2(self.doc, text)

    def add_h3(self, text):
        return add_heading_level3(self.doc, text)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    # ------------------------------------------------------------
    # 各类型申报书构建
    # ------------------------------------------------------------

    def build_innovation_project(self, data: Dict[str, Any], output_path: str):
        """构建大创项目申报书

        Args:
            data: 包含申报书所有字段的字典
            output_path: 输出 docx 路径
        """
        # 封面
        self.add_cover_title(
            "国家级大学生创新创业训练计划项目申报书",
            f"({data.get('project_type', '创新训练项目')})",
        )
        self.add_cover_info([
            {"label": "项目名称", "value": data.get("project_name", "")},
            {"label": "项目负责人", "value": data.get("leader_name", "")},
            {"label": "指导教师", "value": data.get("advisor_name", "")},
            {"label": "所在学院", "value": data.get("college", "")},
            {"label": "申报日期", "value": data.get("apply_date", "")},
        ])
        self.add_page_break()

        # 一、基本信息表
        self.add_h1("一、基本信息")
        basic_info = data.get("basic_info", {})
        self.add_table(
            ["项目", "内容"],
            [
                ["项目名称", basic_info.get("project_name", "")],
                ["项目类型", basic_info.get("project_type", "创新训练项目")],
                ["项目来源", basic_info.get("project_source", "A 学生自主选题")],
                ["所属学科", basic_info.get("discipline", "")],
                ["起止时间", basic_info.get("duration", "")],
                ["申请经费", basic_info.get("budget", "") + " 元"],
                ["负责人", basic_info.get("leader_info", "")],
                ["团队成员", basic_info.get("team_members", "")],
                ["指导教师", basic_info.get("advisor_info", "")],
            ],
        )

        # 二、项目简介
        self.add_h1("二、项目简介")
        self.add_para(data.get("abstract", ""))

        # 三、立项背景与意义
        self.add_h1("三、立项背景与意义")
        for para in data.get("background", []):
            self.add_para(para)

        # 四、研究内容与目标
        self.add_h1("四、项目研究内容与目标")
        self.add_h2("（一）研究内容")
        for i, content in enumerate(data.get("research_content", []), 1):
            self.add_para(f"{i}. {content}")
        self.add_h2("（二）研究目标")
        self.add_para(data.get("research_goal", ""))
        self.add_h2("（三）拟解决的关键问题")
        for i, q in enumerate(data.get("key_problems", []), 1):
            self.add_para(f"{i}. {q}")

        # 五、创新点
        self.add_h1("五、项目创新点")
        for i, innovation in enumerate(data.get("innovations", []), 1):
            self.add_para(f"创新点 {i}：{innovation}")

        # 六、技术路线
        self.add_h1("六、技术路线与研究方法")
        self.add_h2("（一）总体技术路线")
        self.add_para(data.get("tech_route", ""))
        self.add_h2("（二）研究方法")
        for i, method in enumerate(data.get("methods", []), 1):
            self.add_para(f"{i}. {method}")

        # 七、进度安排
        self.add_h1("七、进度安排")
        schedule = data.get("schedule", [])
        if schedule:
            self.add_table(
                ["阶段", "时间", "主要工作", "阶段成果"],
                [[s.get("phase", ""), s.get("time", ""), s.get("work", ""), s.get("output", "")]
                 for s in schedule],
            )

        # 八、预期成果
        self.add_h1("八、预期成果")
        for outcome in data.get("expected_outcomes", []):
            self.add_para(f"• {outcome}")

        # 九、经费预算
        self.add_h1("九、经费预算")
        budget_items = data.get("budget_items", [])
        if budget_items:
            rows = [[b.get("item", ""), b.get("amount", "") + " 元", b.get("basis", "")]
                    for b in budget_items]
            total = sum(int(b.get("amount", 0)) for b in budget_items)
            rows.append(["合计", f"{total} 元", ""])
            self.add_table(["预算科目", "金额", "计算依据"], rows)

        # 十、前期工作基础
        self.add_h1("十、前期工作基础")
        self.add_h2("（一）团队基础")
        self.add_para(data.get("team_foundation", ""))
        self.add_h2("（二）指导教师基础")
        self.add_para(data.get("advisor_foundation", ""))
        self.add_h2("（三）实验条件")
        self.add_para(data.get("lab_condition", ""))

        # 十一、指导教师意见
        self.add_h1("十一、指导教师意见")
        for _ in range(6):
            self.doc.add_paragraph()
        self.add_para("指导教师签字：____________________    日期：______年____月____日", indent=False)

        # 十二、学院评审意见
        self.add_h1("十二、学院评审意见")
        for _ in range(6):
            self.doc.add_paragraph()
        self.add_para("学院盖章：____________________    日期：______年____月____日", indent=False)

        self.save(output_path)

    def build_scholarship_application(self, data: Dict[str, Any], output_path: str):
        """构建奖学金申请书"""
        # 标题
        add_paragraph_with_format(
            self.doc,
            f"{data.get('scholarship_name', '')}申请书",
            font_name=FONT_HEI,
            font_size=SIZE_ER,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_before=12,
            space_after=24,
        )

        # 称呼
        self.add_para("尊敬的学院评审委员会：", indent=False)

        # 正文开头
        opening = data.get("opening", "")
        if opening:
            self.add_para(opening)

        # 一、思想方面
        self.add_h2("一、思想方面")
        for para in data.get("ideology", []):
            self.add_para(para)

        # 二、学习方面
        self.add_h2("二、学习方面")
        for para in data.get("academic", []):
            self.add_para(para)

        # 三、科研与实践方面
        self.add_h2("三、科研与实践方面")
        for para in data.get("research_practice", []):
            self.add_para(para)

        # 四、生活方面
        self.add_h2("四、生活方面")
        for para in data.get("life", []):
            self.add_para(para)

        # 结尾
        self.add_h2("结尾")
        for para in data.get("ending", []):
            self.add_para(para)

        # 落款
        self.doc.add_paragraph()
        self.add_para("此致", indent=False)
        self.add_para("敬礼！", indent=False)
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # 申请人和日期
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"申请人：{data.get('applicant_name', '')}")
        set_run_font(run, font_name=FONT_SONG, font_size=SIZE_XIAO_SI)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(data.get("apply_date", ""))
        set_run_font(run, font_name=FONT_SONG, font_size=SIZE_XIAO_SI)

        self.save(output_path)

    def build_social_practice(self, data: Dict[str, Any], output_path: str):
        """构建三下乡社会实践立项申报书"""
        # 封面
        self.add_cover_title(
            '大学生暑期"三下乡"社会实践立项申报书',
            data.get("team_name", ""),
        )
        self.add_cover_info([
            {"label": "团队名称", "value": data.get("team_name", "")},
            {"label": "申报单位", "value": data.get("college", "")},
            {"label": "队长姓名", "value": data.get("leader_name", "")},
            {"label": "指导教师", "value": data.get("advisor_name", "")},
            {"label": "申报日期", "value": data.get("apply_date", "")},
        ])
        self.add_page_break()

        # 一、团队基本信息
        self.add_h1("一、团队基本信息")
        basic = data.get("team_info", {})
        self.add_table(
            ["项目", "内容"],
            [
                ["团队名称", basic.get("team_name", "")],
                ["实践主题", basic.get("theme", "")],
                ["实践地点", basic.get("location", "")],
                ["实践时间", basic.get("time", "")],
                ["团队人数", basic.get("headcount", "")],
                ["队长信息", basic.get("leader_info", "")],
                ["指导教师", basic.get("advisor_info", "")],
            ],
        )

        # 二、团队成员信息表
        self.add_h1("二、团队成员信息")
        members = data.get("members", [])
        if members:
            self.add_table(
                ["姓名", "学号", "专业年级", "团队分工", "联系方式"],
                [[m.get("name", ""), m.get("id", ""), m.get("major", ""),
                  m.get("role", ""), m.get("phone", "")]
                 for m in members],
            )

        # 三、实践主题与背景
        self.add_h1("三、实践主题与背景")
        for para in data.get("theme_background", []):
            self.add_para(para)

        # 四、实践目的与意义
        self.add_h1("四、实践目的与意义")
        for para in data.get("purpose", []):
            self.add_para(para)

        # 五、实践内容与实施方案
        self.add_h1("五、实践内容与实施方案")
        for para in data.get("implementation", []):
            self.add_para(para)

        # 六、安全保障预案
        self.add_h1("六、安全保障预案")
        for para in data.get("safety_plan", []):
            self.add_para(para)

        # 七、预期成果
        self.add_h1("七、预期成果")
        for outcome in data.get("expected_outcomes", []):
            self.add_para(f"• {outcome}")

        # 八、经费预算
        self.add_h1("八、经费预算")
        budget_items = data.get("budget_items", [])
        if budget_items:
            rows = [[b.get("item", ""), b.get("amount", "") + " 元", b.get("basis", "")]
                    for b in budget_items]
            total = sum(int(b.get("amount", 0)) for b in budget_items)
            rows.append(["合计", f"{total} 元", ""])
            self.add_table(["预算科目", "金额", "计算依据"], rows)

        # 九、宣传计划
        self.add_h1("九、宣传计划")
        for para in data.get("publicity", []):
            self.add_para(para)

        # 十、指导教师意见
        self.add_h1("十、指导教师意见")
        for _ in range(4):
            self.doc.add_paragraph()
        self.add_para("指导教师签字：____________________    日期：______年____月____日", indent=False)

        # 十一、学院团委意见
        self.add_h1("十一、学院团委意见")
        for _ in range(4):
            self.doc.add_paragraph()
        self.add_para("学院团委盖章：____________________    日期：______年____月____日", indent=False)

        self.save(output_path)

    def build_party_application(self, data: Dict[str, Any], output_path: str):
        """构建入党申请书"""
        # 标题
        add_paragraph_with_format(
            self.doc,
            "入党申请书",
            font_name=FONT_HEI,
            font_size=SIZE_ER,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_before=12,
            space_after=24,
        )

        # 称呼
        self.add_para("敬爱的党组织：", indent=False)

        # 正文 - 4 段结构
        sections = [
            ("一、入党志愿表达", data.get("will_expression", [])),
            ("二、对党的认识", data.get("party_understanding", [])),
            ("三、个人经历与思想变化", data.get("personal_journey", [])),
            ("四、个人不足与今后努力方向", data.get("shortcomings_plan", [])),
        ]
        for title, paragraphs in sections:
            self.add_h2(title)
            for para in paragraphs:
                self.add_para(para)

        # 结尾
        self.add_para("请党组织在实践中考验我！", indent=False)
        self.doc.add_paragraph()
        self.add_para("此致", indent=False)
        self.add_para("敬礼！", indent=False)
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # 落款
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"申请人：{data.get('applicant_name', '')}")
        set_run_font(run, font_name=FONT_SONG, font_size=SIZE_XIAO_SI)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(data.get("apply_date", ""))
        set_run_font(run, font_name=FONT_SONG, font_size=SIZE_XIAO_SI)

        self.save(output_path)

    # ------------------------------------------------------------
    # 保存
    # ------------------------------------------------------------

    def save(self, output_path: str):
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申报书已生成：{output_path}")
        return str(output_path)


# ============================================================
# CLI
# ============================================================

def main():
    import argparse

    parser = argparse.ArgumentParser(description="大学生申报书 docx 生成器")
    parser.add_argument("--type", required=True,
                        choices=["innovation", "scholarship", "social_practice", "party"],
                        help="申报书类型")
    parser.add_argument("--data", required=True, help="数据 JSON 文件路径")
    parser.add_argument("--out", required=True, help="输出 docx 路径")

    args = parser.parse_args()

    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    builder = ApplicationDocBuilder()
    if args.type == "innovation":
        builder.build_innovation_project(data, args.out)
    elif args.type == "scholarship":
        builder.build_scholarship_application(data, args.out)
    elif args.type == "social_practice":
        builder.build_social_practice(data, args.out)
    elif args.type == "party":
        builder.build_party_application(data, args.out)


if __name__ == "__main__":
    main()
