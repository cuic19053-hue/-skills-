#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
大学生暑期"三下乡"社会实践-调研类立项申报书 docx 生成器

按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 一级标题：黑体三号，居中
- 二级标题：黑体小三，左对齐
- 三级标题：宋体四号加粗
- 表格：宋体五号，居中

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第十一章。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)
SIZE_XIAO_ER = Pt(18)
SIZE_SAN = Pt(16)
SIZE_XIAO_SAN = Pt(15)
SIZE_SI = Pt(14)
SIZE_XIAO_SI = Pt(12)
SIZE_WU = Pt(10.5)
SIZE_XIAO_WU = Pt(9)

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
        space_before=12, space_after=12,
    )


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6, space_after=6,
    )


def add_heading_level3(doc, text: str):
    """三级标题：宋体四号加粗，左对齐"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6, space_after=3,
    )


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent,
        line_spacing=1.5,
    )


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 2 and len(headers) == 4 \
                else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False, alignment=align)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def merge_vertical_cells(table, col_idx: int, start_row: int, end_row: int):
    """纵向合并单元格（用于签字栏预留空白）"""
    cells = [table.rows[r].cells[col_idx] for r in range(start_row, end_row + 1)]
    merged = cells[0]
    for c in cells[1:]:
        merged = merged.merge(c)


def compute_sample_size(z: float = 1.96, sigma: float = 0.5, e: float = 0.05) -> int:
    """样本量计算公式 n = Z² × σ² / E²，向上取整"""
    n = (z ** 2) * (sigma ** 2) / (e ** 2)
    return int(n) + (1 if n > int(n) else 0)


def fmt_money(amount) -> str:
    """金额格式化：整数转字符串 + 元"""
    try:
        return f"{int(amount)} 元"
    except (ValueError, TypeError):
        return str(amount)


def safe_int(value, default: int = 0) -> int:
    """安全转整数"""
    try:
        return int(str(value).strip())
    except (ValueError, TypeError):
        return default


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """三下乡社会实践-调研类立项申报书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def add_h1(self, text):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text):
        return add_heading_level2(self.doc, text)

    def add_h3(self, text):
        return add_heading_level3(self.doc, text)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_table(self, headers, rows, col_widths=None):
        return add_table_from_data(self.doc, headers, rows, col_widths)

    def add_page_break(self):
        add_page_break(self.doc)

    # 封面

    def _add_cover(self):
        """封面：黑体二号标题 + 4 行下划线信息"""
        for _ in range(3):
            self.doc.add_paragraph()

        title = "大学生暑期\u201c三下乡\u201d社会实践活动立项申报书"
        add_paragraph_with_format(
            self.doc, title,
            font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_before=12, space_after=12,
        )

        subtitle = f"（{self._get('theme', default='乡村振兴')}主题·调研类）"
        add_paragraph_with_format(
            self.doc, subtitle,
            font_name=FONT_HEI, font_size=SIZE_SAN, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_after=24,
        )

        for _ in range(3):
            self.doc.add_paragraph()

        project_name = self._get("team_name", default="赴 XX 县乡村振兴调研团")
        info_items = [
            ("项目名称", project_name),
            ("团队名称", self._get("team_name")),
            ("申报单位", self._get("college")),
            ("申报日期", self._get("apply_date")),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI,
                         font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True

        self.add_page_break()

    # 一、团队基本信息

    def _add_team_info(self):
        """一、团队基本信息表（7 行 2 列）"""
        self.add_h1("一、团队基本信息")
        team = self._get("team_info", default={})
        if not isinstance(team, dict):
            team = {}
        g = self._get
        leader = team.get("leader", f"{g('leader_name')} / {g('leader_id')} / "
                          f"{g('leader_major')} / {g('leader_grade')} / {g('leader_phone')}")
        advisor = team.get("advisor", f"{g('advisor_name')} / {g('advisor_title')} / "
                           f"{g('advisor_phone')} / {g('advisor_with_team', '随队')}")
        rows = [
            ["团队名称", team.get("team_name", g("team_name", ""))],
            ["实践主题", team.get("theme", g("theme", ""))],
            ["实践地点", team.get("location", g("location", ""))],
            ["实践时间", team.get("practice_time", g("practice_time", ""))],
            ["团队人数", team.get("team_size", g("team_size", ""))],
            ["队长", leader],
            ["指导教师", advisor],
        ]
        self.add_table(["项目", "内容"], rows, col_widths=[4.5, 11.5])

    # 二、团队成员信息表

    def _add_members_table(self):
        """二、团队成员信息表（5 列）"""
        self.add_h1("二、团队成员信息表")
        members = self._get("members", default=[])
        if members and isinstance(members, list):
            rows = []
            for m in members:
                if not isinstance(m, dict):
                    continue
                rows.append([
                    m.get("name", ""),
                    m.get("id", ""),
                    m.get("major", ""),
                    m.get("role", ""),
                    m.get("phone", ""),
                ])
            self.add_table(
                ["姓名", "学号", "专业年级", "团队分工", "联系方式"],
                rows,
                col_widths=[2.5, 2.8, 3.5, 3.5, 3.7],
            )
        else:
            self.add_para("（请填写团队成员信息表，每人一行：姓名 / 学号 / "
                          "专业年级 / 团队分工 / 联系方式。建议跨学院组队，"
                          "专业互补，分工具体到\u201c做什么\u201d。团队人数 6~15 人为宜。）")

    # 三、实践主题与背景

    def _add_theme_background(self):
        """三、实践主题与背景（400~600 字，3 段）"""
        self.add_h1("三、实践主题与背景")
        bg = self._get("theme_background", default=[])
        if isinstance(bg, str):
            bg = [bg]
        if bg:
            for para in bg:
                self.add_para(para)
        else:
            self.add_h2("（一）团中央主题对应")
            self.add_para("（请填写当年团中央发布的实践主题方向，以及本项目对应"
                          "切入的维度，150 字左右。如：2025 年团中央“乡村振兴 青春建功”主题，本项目切入产业/教育/医疗/文化/生态五维度。）")
            self.add_h2("（二）选址理由")
            self.add_para("（请填写选址理由，150~200 字，3 句话讲清“为什么去这里”："
                          "与团中央主题契合度 + 当地实际情况（人口/产业/特色）+ "
                          "团队已有联系。）")
            self.add_h2("（三）实践对象基本情况")
            self.add_para("（请填写实践对象基本情况，100~150 字：调研对象规模、"
                          "人口结构、产业特征。必须有数据，如常住人口 8600 人、"
                          "适龄劳动力 5200 人、农民人均可支配收入 2.1 万元。）")

    # 四、实践目的与意义

    def _add_purpose_significance(self):
        """四、实践目的与意义（300~500 字，2 段）"""
        self.add_h1("四、实践目的与意义")
        ps = self._get("purpose_significance", default=[])
        if isinstance(ps, str):
            ps = [ps]
        if ps:
            for para in ps:
                self.add_para(para)
        else:
            self.add_h2("（一）对当地")
            self.add_para("（请填写对当地的意义，150~250 字。结构：通过 N 份问卷"
                          "+ N 位访谈，系统梳理 X 维度现状与问题，形成约 X 万字"
                          "调研报告，提交 X 部门，重点解决 X 个问题。）")
            self.add_h2("（二）对学生")
            self.add_para("（请填写对学生的意义，150~250 字。结构：掌握问卷设计、"
                          "抽样方法、深度访谈、数据分析等社会调查方法 + 深入田间"
                          "地头深化国情农情认识 + 厚植家国情怀。）")

    # 五、实践内容与实施方案【重点】

    def _add_implementation_plan(self):
        """五、实践内容与实施方案（800~1200 字，3 子节 + 按天表）"""
        self.add_h1("五、实践内容与实施方案")

        self.add_h2("（一）实践形式")
        form = self._get("implementation_plan", default={})
        if not isinstance(form, dict):
            form = {}
        form_text = form.get("form", "")
        if form_text:
            self.add_para(form_text)
        else:
            self.add_para("本项目实践形式为调研类，采用\u201c走访调研 + 问卷调查 + "
                          "深度访谈\u201d三段式实施流程。走访调研覆盖 5 个行政村，"
                          "问卷调查采用分层随机抽样，深度访谈覆盖村干部、致富"
                          "带头人、普通村民三层级。")

        self.add_h2("（二）调研问卷设计")
        questionnaire = form.get("questionnaire", "")
        if questionnaire:
            self.add_para(questionnaire)
        else:
            self.add_para("问卷采用 5 维度 25 题结构：人口学 4 题（性别、年龄、学历、职业）+ 经济 6 题（收入、产业、就业）+ 教育 5 题（义务教育、师资、辍学）+ 医疗 5 题（就医距离、医保、服务满意度）+ 文化 5 题（文化活动、传统习俗、文明建设）。题型分布：单选 8 题 + 多选 6 题 + 李克特 5 级量表 11 题。正式发放前在非调研村做 30 份预调研，Cronbach α ≥ 0.7 后正式发放。")

        self.add_h2("（三）按天实施安排")
        schedule = form.get("schedule", [])
        if schedule and isinstance(schedule, list):
            rows = []
            total_q = 0
            total_int = 0
            for s in schedule:
                if not isinstance(s, dict):
                    continue
                rows.append([
                    s.get("date", ""),
                    s.get("location", ""),
                    s.get("work", ""),
                    s.get("output", ""),
                ])
                output = str(s.get("output", ""))
                if "问卷" in output:
                    for tok in output.replace("问卷", " ").split():
                        if tok.isdigit():
                            total_q += int(tok)
                if "访谈" in output:
                    for tok in output.replace("访谈", " ").split():
                        if tok.isdigit():
                            total_int += int(tok)
            self.add_table(
                ["日期", "地点", "主要工作", "预期产出"],
                rows,
                col_widths=[2.2, 3.0, 6.0, 4.8],
            )
            if total_q or total_int:
                self.add_para(f"合计：问卷 {total_q} 份 + 访谈 {total_int} 人 + "
                              f"宣讲 1 场 + 调研报告初稿 1 份。每天结束当晚召开"
                              f"日例会，整理当日发现并调整次日方案。")
        else:
            self.add_para("（请填写按天实施安排表格，7~10 行覆盖整个实践期。每行 4 列：日期 / 地点（精确到村）/ 主要工作 / 预期产出（问卷数 + 访谈人数）。示例：7.15 抵达 + 预调研 30 份；7.16-17 5 村各 20 户问卷 100 份；7.18 8 位村干部 + 3 位致富带头人访谈；7.19 5 位普通村民访谈 + 1 场宣讲；7.20 数据录入；7.21 报告初稿 + 返程。）")

    # 六、安全保障预案【重点】

    def _add_safety_plan(self):
        """六、安全保障预案（300~500 字，5 段）"""
        self.add_h1("六、安全保障预案")
        plan = self._get("safety_plan", default=[])
        if isinstance(plan, str):
            plan = [plan]
        if plan and isinstance(plan, list):
            for para in plan:
                self.add_para(para)
        else:
            self.add_h2("（一）出行安全")
            self.add_para("（请填写出行安全，80 字：交通方式（统一购买高铁票）、住宿选择（政府招待所 2 人一间男女分区）、每日 18:00 前向指导教师报平安。）")
            self.add_h2("（二）人身安全")
            self.add_para("（请填写人身安全，80 字：防暑（正午 11:00-14:00 不外出）、防疫（每日早晚测温）、防骗（不携带大量现金、单人不离队）。）")
            self.add_h2("（三）应急联系人")
            self.add_para("（请填写应急联系人，60 字：3 名——指导教师、队长、学院团委，附电话。当地对接：XX 县团委刘书记 136XXXXXXXX。）")
            self.add_h2("（四）应急流程")
            self.add_para("（请填写应急流程，120 字：分级响应。一般事件（伤病 < 1000 元）：队长 → 指导教师 → 学院存档。重大事件（≥ 1000 元 / 治安 / 灾害）：队长 → 指导教师 → 学院 → 110/120，30 分钟首报，2 小时书面报告。）")
            self.add_h2("（五）保险购买")
            self.add_para("（请填写保险购买，60 字：全员购买短期意外险，保额 30 万元，保费 15 元/人，保单号 PA20250715001-010，覆盖 7 月 14-22 日。）")

    # 七、预期成果

    def _add_expected_results(self):
        """七、预期成果（必须可量化）"""
        self.add_h1("七、预期成果")
        outcomes = self._get("expected_results", default=[])
        if isinstance(outcomes, str):
            outcomes = [outcomes]
        if outcomes:
            for o in outcomes:
                self.add_para(f"• {o}", indent=False)
        else:
            self.add_para("• 调研报告 1 份（约 1.8 万字，含数据分析与政策建议，"
                          "提交 XX 县团委与 XX 镇政府）", indent=False)
            self.add_para("• 新闻稿 5 篇（中青网 1 篇、校团委公众号 2 篇、"
                          "学院公众号 2 篇）", indent=False)
            self.add_para("• 短视频 3 个（每个 3-5 分钟，发布在 B 站、抖音）",
                          indent=False)
            self.add_para("• 纪录短片 1 个（10 分钟，用于校内汇报）", indent=False)
            self.add_para("• 访谈记录 16 份（村干部 8 + 致富带头人 3 + 普通村民 5）",
                          indent=False)
            self.add_para("• 问卷数据集 1 套（100 份有效问卷，含原始数据与"
                          "分析代码）", indent=False)

    # 八、经费预算

    def _add_budget(self):
        """八、经费预算（3 列表格：科目/金额/计算依据）"""
        self.add_h1("八、经费预算")
        items = self._get("budget_items", default=[])
        if items and isinstance(items, list):
            rows = []
            total = 0
            for b in items:
                if not isinstance(b, dict):
                    continue
                amount_num = safe_int(b.get("amount", 0))
                total += amount_num
                rows.append([
                    b.get("item", ""),
                    fmt_money(amount_num),
                    b.get("basis", ""),
                ])
            rows.append(["合计", fmt_money(total), ""])
            self.add_table(
                ["预算科目", "金额", "计算依据"],
                rows,
                col_widths=[3.5, 3.0, 9.5],
            )
        else:
            self.add_para("（请填写经费预算，5 类标准科目：交通费 / 食宿费 / "
                          "物资费 / 印刷费 / 其他。每项金额非整数，附计算依据。"
                          "示例：交通费 2860 元 = 高铁 130 × 10 × 2 + 包车 "
                          "300 × 4。）")

    # 九、宣传计划

    def _add_publicity_plan(self):
        """九、宣传计划（3 类媒体）"""
        self.add_h1("九、宣传计划")
        plan = self._get("publicity_plan", default=[])
        if isinstance(plan, str):
            plan = [plan]
        if plan and isinstance(plan, list):
            for p in plan:
                self.add_para(p)
        else:
            self.add_h2("（一）校内媒体")
            self.add_para("校团委公众号 2 篇、校报 1 篇、学院公众号 2 篇。"
                          "发布时间节点：7.16 启动报道、7.19 中期进展、7.21 结项报道。")
            self.add_h2("（二）校外媒体")
            self.add_para("中青网 1 篇、XX 县电视台 1 条、XX 日报 1 篇。"
                          "中青网稿件由队长审核后投递，附调研照片 5~8 张。")
            self.add_h2("（三）新媒体")
            self.add_para("B 站 3 个短视频、抖音 3 个短视频、微信视频号 2 个。"
                          "每个视频 3~5 分钟，内容覆盖调研过程、村民故事、"
                          "团队风采三个角度。")

    # 十、十一、指导教师 + 学院团委意见

    def _add_review_section(self):
        """十、指导教师意见 / 十一、学院团委意见

        两栏预留签字空间（空 6 行），后接签字行 + 日期行。
        若 include_school_approval 为 True 则追加"十二、学校审批意见"。
        """
        self.add_h1("十、指导教师意见")
        for _ in range(6):
            self.doc.add_paragraph()
        self.add_para(
            "指导教师签字：____________________    "
            "日期：______年____月____日",
            indent=False,
        )

        self.add_h1("十一、学院团委意见")
        for _ in range(6):
            self.doc.add_paragraph()
        self.add_para(
            "学院盖章：____________________    "
            "日期：______年____月____日",
            indent=False,
        )

        if self._get("include_school_approval", default=False):
            self.add_h1("十二、学校审批意见")
            for _ in range(6):
                self.doc.add_paragraph()
            self.add_para(
                "学校盖章：____________________    "
                "日期：______年____月____日",
                indent=False,
            )

        # 团中央专项报送说明（可选）
        special = self._get("tuancentral_special", default="")
        if special:
            self.add_h1("附：团中央专项报送说明")
            self.add_para(f"对应专项：{special}。本项目与专项主题契合度高，"
                          "预期形成优秀调研报告 1 份、典型案例 1 个，"
                          "报送团中央专项工作组。")

    # 主构建方法

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 11 栏目，生成 docx

        Args:
            data: 申报书字段字典
            output_path: 输出 docx 路径

        Returns:
            实际保存路径
        """
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()

            self._add_cover()
            self._add_team_info()
            self._add_members_table()
            self._add_theme_background()
            self._add_purpose_significance()
            self._add_implementation_plan()
            self._add_safety_plan()
            self._add_expected_results()
            self._add_budget()
            self._add_publicity_plan()
            self._add_review_section()

            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申报书已生成：{output_path}")
        return str(output_path)

    # 数据校验（含调研类专属校验）

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）

        校验分四类：P0 必填字段、调研专属校验、安全预案校验、经费预算校验。
        """
        warnings = []
        p0_fields = [
            ("team_name", "团队名称"), ("theme", "实践主题"),
            ("location", "实践地点"), ("practice_time", "实践时间"),
            ("leader_name", "队长姓名"), ("college", "申报单位"),
        ]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        team = self._get("team_info", default={})
        if not isinstance(team, dict):
            team = {}
        if not team.get("team_size") and not self._get("team_size"):
            warnings.append("缺少 团队人数（team_size）")
        if not team.get("advisor") and not self._get("advisor_name"):
            warnings.append("缺少 指导教师（advisor_name）")
        if not self._get("theme_background"):
            warnings.append("缺少 实践主题与背景（theme_background），将使用占位文本")
        if not self._get("purpose_significance"):
            warnings.append("缺少 实践目的与意义（purpose_significance），将使用占位文本")

        impl = self._get("implementation_plan", default={})
        if not isinstance(impl, dict):
            impl = {}
        if not impl.get("schedule"):
            warnings.append("缺少 按天实施安排（implementation_plan.schedule），将使用占位文本——评审会扣大分")
        else:
            for i, s in enumerate(impl.get("schedule", []), 1):
                if isinstance(s, dict):
                    if not s.get("date") or not s.get("location"):
                        warnings.append(f"按天表第 {i} 行缺少日期或地点")
                    if not s.get("output"):
                        warnings.append(f"按天表第 {i} 行缺少预期产出")

        if not impl.get("form"):
            warnings.append("缺少 实践形式（implementation_plan.form），建议明确为调研类")
        if not impl.get("questionnaire"):
            warnings.append("缺少 调研问卷设计（implementation_plan.questionnaire），评审会扣调研科学性分")

        if not self._get("safety_plan"):
            warnings.append("缺少 安全保障预案（safety_plan），将使用占位文本——安全预案不达标可能一票否决")
        else:
            plan = self._get("safety_plan", default=[])
            plan_text = "\n".join(plan) if isinstance(plan, list) else str(plan)
            if "保险" not in plan_text:
                warnings.append("安全预案未提及保险购买情况——一票否决项")
            if "应急联系" not in plan_text:
                warnings.append("安全预案未列明应急联系人")
            if "应急流程" not in plan_text:
                warnings.append("安全预案未列明应急流程")

        if not self._get("expected_results"):
            warnings.append("缺少 预期成果（expected_results），将使用占位文本")
        else:
            results = self._get("expected_results", default=[])
            results_text = "\n".join(results) if isinstance(results, list) else ""
            if "调研报告" not in results_text:
                warnings.append("预期成果未以调研报告为主——调研类核心产出缺失")
            if "万" not in results_text and "字" not in results_text:
                warnings.append("预期成果调研报告未注明字数（1.5~2 万字）")

        items = self._get("budget_items", default=[])
        if items:
            total = sum(safe_int(b.get("amount", 0)) for b in items if isinstance(b, dict))
            budget_total_num = safe_int(str(self._get("budget_total", default="")).strip(), default=-1)
            if budget_total_num >= 0 and total != budget_total_num:
                warnings.append(f"预算合计 {total} 元 与申请经费 {budget_total_num} 元不一致")
        else:
            warnings.append("缺少 经费预算（budget_items），将使用占位文本")

        members = self._get("members", default=[])
        if members and isinstance(members, list):
            for i, m in enumerate(members, 1):
                if isinstance(m, dict):
                    if not m.get("role"):
                        warnings.append(f"成员 {m.get('name', f'#{i}')} 缺少分工")

        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        self.warnings = warnings
        return warnings


# ============================================================
# 默认示例数据
# ============================================================

DEFAULT_DATA = {
    "team_name": "赴 XX 县乡村振兴调研团",
    "theme": "乡村振兴",
    "location": "XX 省 XX 县 XX 镇 5 个行政村（A/B/C/D/E 村）",
    "practice_time": "2025.07.15-07.21（7 天）",
    "team_size": "10 人",
    "leader_name": "张三",
    "leader_id": "202212345",
    "leader_major": "经济学 2022 级",
    "leader_grade": "大三",
    "leader_phone": "138XXXXXXXX",
    "advisor_name": "李教授",
    "advisor_title": "副教授",
    "advisor_phone": "139XXXXXXXX",
    "advisor_with_team": "随队",
    "college": "经济管理学院",
    "apply_date": "2025 年 5 月 20 日",
    "team_info": {
        "team_name": "赴 XX 县乡村振兴调研团",
        "theme": "乡村振兴（对应 2025 年团中央“乡村振兴 青春建功”主题）",
        "location": "XX 省 XX 县 XX 镇 5 个行政村（A/B/C/D/E 村）",
        "practice_time": "2025.07.15-07.21（7 天）",
        "team_size": "10 人",
        "leader": "张三 / 202212345 / 经济学 2022 级 / 大三 / 138XXXXXXXX",
        "advisor": "李教授 / 副教授 / 139XXXXXXXX / 随队",
    },
    "members": [
        {"name": "张三", "id": "202212345", "major": "经济学 2022 级", "role": "队长 / 总协调", "phone": "138XXXXXXXX"},
        {"name": "李四", "id": "202212346", "major": "社会学 2022 级", "role": "问卷设计与发放", "phone": "138XXXXXXXX"},
        {"name": "王五", "id": "202212347", "major": "统计学 2023 级", "role": "数据分析", "phone": "138XXXXXXXX"},
        {"name": "赵六", "id": "202212348", "major": "新闻学 2022 级", "role": "宣传报道", "phone": "138XXXXXXXX"},
        {"name": "孙七", "id": "202212349", "major": "经济学 2023 级", "role": "访谈记录", "phone": "138XXXXXXXX"},
    ],
    "theme_background": [
        "团中央主题对应：2025 年团中央发布“乡村振兴 青春建功”暑期社会实践主题，鼓励高校学生深入乡村调研乡村振兴战略实施情况。本项目紧扣该主题，聚焦“产业兴旺、生态宜居、乡风文明、治理有效、生活富裕”五大维度开展实证调研。",
        "选址理由：选址 XX 省 XX 县 XX 镇原因有三：一是该镇 2024 年被列入省级乡村振兴示范镇，政策支持力度大；二是该镇下辖 5 个行政村产业结构差异明显（A 村玉米 / B 村果蔬 / C 村养殖 / D 村旅游 / E 村手工业），便于对比分析；三是团队中 2 名成员为该镇籍贯，已与当地团委建立联系。",
        "实践对象基本情况：XX 镇常住人口 8600 人，其中 18~65 岁适龄劳动力 5200 人。5 个行政村共 1850 户，调研对象覆盖 18~65 岁常住村民。该镇 2024 年农民人均可支配收入 2.1 万元，略低于全省平均水平（2.3 万元），具有调研典型性。",
    ],
    "purpose_significance": [
        "对当地：通过 100 份有效问卷与 16 位深度访谈，系统梳理 XX 镇 5 个行政村在产业、教育、医疗、文化、生态五个维度的现状与问题，形成约 1.8 万字的调研报告。报告将提交 XX 县团委与 XX 镇政府，为当地乡村振兴规划提供一线数据支撑，重点解决\u201c产业同质化\u201d\u201c教育资源配置不均\u201d\u201c基层医疗服务可及性差\u201d三个问题。",
        "对学生：团队成员在实践中掌握问卷设计、抽样方法、深度访谈、数据分析等社会调查方法，提升\u201c用数据说话\u201d的研究能力；通过深入田间地头与农户同吃同住同劳动，深化对国情农情的认识，厚植家国情怀与三农感情，培养扎根基层、服务社会的价值观。",
    ],
    "implementation_plan": {
        "form": "本项目实践形式为调研类，采用\u201c走访调研 + 问卷调查 + 深度访谈\u201d三段式实施流程。走访调研覆盖 5 个行政村，问卷调查采用分层随机抽样（按村产业结构分 5 层，村层内按户籍名单简单随机抽户），深度访谈覆盖村干部、致富带头人、普通村民三层级。",
        "questionnaire": "问卷采用 5 维度 25 题结构：人口学 4 题（性别、年龄、学历、职业）+ 经济 6 题（收入、产业、就业）+ 教育 5 题（义务教育、师资、辍学）+ 医疗 5 题（就医距离、医保、服务满意度）+ 文化 5 题（文化活动、传统习俗、文明建设）。题型：单选 8 + 多选 6 + 李克特 5 级量表 11。样本量 n=Z²σ²/E²=1.96²×0.5×0.5/0.05²≈384，考虑 7 天时间限制实际发放 100 份，以描述性统计为主。预调研 30 份，Cronbach α ≥ 0.7 后正式发放。",
        "schedule": [
            {"date": "7.15", "location": "XX 县城", "work": "抵达 XX 县，与县团委对接，召开启动会；下午开展预调研", "output": "启动会 1 次；预调研问卷 30 份"},
            {"date": "7.16-17", "location": "XX 镇 5 村", "work": "5 个村各随机抽取 20 户，发放问卷 100 份，回收率 ≥ 90%", "output": "问卷 100 份"},
            {"date": "7.18", "location": "5 村", "work": "8 位村干部 + 3 位致富带头人深度访谈（每人 45~60 分钟）", "output": "访谈记录 11 份"},
            {"date": "7.19", "location": "A/B 村", "work": "5 位普通村民深度访谈；下午 A 村政策宣讲 1 场", "output": "访谈记录 5 份；宣讲 1 场"},
            {"date": "7.20", "location": "XX 镇政府", "work": "数据录入与初步分析；与镇团委反馈调研发现", "output": "数据集 1 套"},
            {"date": "7.21", "location": "XX 县城", "work": "撰写调研报告初稿；返程", "output": "报告初稿 1 份"},
        ],
    },
    "safety_plan": [
        "一、出行安全：全员统一购买高铁票，不单独行动；7 月 15 日集体乘 GXXX 次列车赴 XX 县。住宿选择 XX 县政府招待所（已与县团委对接预订），2 人一间，男女分区。每日 18:00 前向指导教师报平安。",
        "二、人身安全：携带常用药品（藿香正气水、创可贴、退烧药、止泻药）；正午 11:00-14:00 不外出，避免中暑；每日早晚测温，体温 ≥ 37.3℃ 立即隔离观察；不携带大量现金，单人不离队。",
        "三、应急联系人：指导教师李教授（139XXXXXXXX）、队长张三（138XXXXXXXX）、学院团委王老师（137XXXXXXXX）。当地对接：XX 县团委刘书记（136XXXXXXXX）。",
        "四、应急流程：一般事件（伤病 < 1000 元）：队长处理 → 指导教师报备 → 学院团委存档。重大事件（伤病 ≥ 1000 元 / 治安事件 / 自然灾害）：队长 → 指导教师 → 学院 → 110/120，同步报告当地团委。所有事件 30 分钟内首报，2 小时内书面报告。",
        "五、保险购买：全员购买中国平安短期意外险（保额 30 万元，保费 15 元/人，保单号 PA20250715001-010），覆盖 7 月 14-22 日。",
    ],
    "expected_results": [
        "调研报告 1 份（约 1.8 万字，含数据分析与政策建议，提交 XX 县团委与 XX 镇政府）",
        "新闻稿 5 篇（中青网 1 篇、校团委公众号 2 篇、学院公众号 2 篇）",
        "短视频 3 个（每个 3-5 分钟，发布在 B 站、抖音）",
        "纪录短片 1 个（10 分钟，用于校内汇报）",
        "访谈记录 16 份（村干部 8 + 致富带头人 3 + 普通村民 5）",
        "问卷数据集 1 套（100 份有效问卷，含原始数据与分析代码）",
    ],
    "budget_items": [
        {"item": "交通费", "amount": "2860", "basis": "高铁往返 130 元 × 10 人 × 2 次 + 县内包车 300 元/天 × 4 天"},
        {"item": "食宿费", "amount": "5600", "basis": "招待所 120 元/间 × 5 间 × 7 天 + 餐费 40 元/人/天 × 10 人 × 7 天"},
        {"item": "物资费", "amount": "850", "basis": "问卷印刷 200 份 × 2 元 + 文具 50 套 × 5 元 + 药品 150 元"},
        {"item": "印刷费", "amount": "380", "basis": "调研报告印刷 10 份 × 30 元 + 海报 10 张 × 8 元"},
        {"item": "其他", "amount": "310", "basis": "保险 15 元 × 10 人 + 通讯补贴 16 元 × 10 人"},
    ],
    "budget_total": "10000",
    "publicity_plan": [
        "（一）校内媒体：校团委公众号 2 篇、校报 1 篇、学院公众号 2 篇。发布时间节点：7.16 启动报道、7.19 中期进展、7.21 结项报道。",
        "（二）校外媒体：中青网 1 篇、XX 县电视台 1 条、XX 日报 1 篇。中青网稿件由队长审核后投递，附调研照片 5~8 张。",
        "（三）新媒体：B 站 3 个短视频、抖音 3 个短视频、微信视频号 2 个。每个视频 3~5 分钟，内容覆盖调研过程、村民故事、团队风采三个角度。",
    ],
    "include_school_approval": False,
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="三下乡社会实践-调研类立项申报书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n"
            "JSON 字段定义详见 SKILL.md 第十一章。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
