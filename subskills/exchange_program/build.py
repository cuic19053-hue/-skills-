#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
交流项目申请书 docx 生成器 v1.0.0

格式标准：A4 / 页边距上下 2.54cm 左右 2.5cm / 正文宋体小四 1.5 倍行距首行缩进 2 字符 /
一级标题黑体二号居中 / 二级标题黑体小四左对齐 / 表格宋体五号居中。

支持 3 类交流项目（通过 exchange_type 字段标识）：
- school: 校际交换（学校与境外高校签署合作协议的官方交换项目，学费互免）
- college: 院际交换（院系与境外高校的院级合作项目，简化程序）
- csc: CSC 公派（国家留学基金管理委员会资助的公派留学项目，全额资助 + 回国服务 2 年义务）

5 段结构（按 2000 字标准版默认配比）：
- 一、个人基本情况  15%  约 300 字
- 二、交流动机     25%  约 500 字（4 维度：学术深化/视野拓展/语言提升/职业规划）
- 三、学习计划     30%  约 600 字（3 阶段：行前准备/在外学习/归国总结）
- 四、文化交流意愿  15%  约 300 字（3 维度：融入/传播/学术交流）
- 五、安全保障     15%  约 300 字（4 维度：法律/保险/汇报/预案）

3 档字数版本（通过 word_count_target 字段切换，主要靠字段完整度自然达到）：
- short: 1500 字（院际交换/短期交流/暑期学校）
- standard: 2000 字（校际交换/CSC 公派，默认）
- long: 2500 字（CSC 公派顶尖项目/联合培养博士）

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md §12（含 P0/P1/P2 分级与数组字段子结构）。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号（标题）
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四（正文）
SIZE_WU = Pt(10.5)          # 五号（表格）
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

# 交流项目类型
EXCHANGE_SCHOOL = "school"
EXCHANGE_COLLEGE = "college"
EXCHANGE_CSC = "csc"
VALID_EXCHANGE_TYPES = (EXCHANGE_SCHOOL, EXCHANGE_COLLEGE, EXCHANGE_CSC)

# 字数档位
WORD_COUNT_SHORT = "short"
WORD_COUNT_STANDARD = "standard"
WORD_COUNT_LONG = "long"
VALID_WORD_COUNTS = (WORD_COUNT_SHORT, WORD_COUNT_STANDARD, WORD_COUNT_LONG)

# 禁用句检测列表（交流动机中不得出现，详见 SKILL.md §10.1 共 7 项 + 扩展）
FORBIDDEN_PHRASES = [
    "开阔眼界", "出国镀金", "体验不同文化", "学习先进知识",
    "感受异国风情", "为未来发展奠定基础", "提高英语水平",
    "恳请领导批准", "完全符合条件",
]

# 标题
TITLE_TEXT = "交流项目申请书"


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(doc, text: str, font_name: str = FONT_SONG,
        font_size=SIZE_XIAO_SI, bold: bool = False,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent: bool = True,
        line_spacing: float = 1.5, space_before: float = 0, space_after: float = 0):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_ER,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进），宋体小四"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_section_heading(doc, text: str):
    """正文小节标题（一、二、三…）：黑体小四加粗，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SI,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
        line_spacing=1.5, space_before=6, space_after=3)


def add_cizhi_paragraph(doc, text: str = "此致"):
    """'此致'另起一行，空两格（首行缩进 2 字符）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格（不缩进）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=False, line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None, caption: str = ""):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    if caption:
        add_paragraph_with_format(doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def format_rank_percent(rank: str, rank_total: str) -> str:
    """根据排名与基数计算百分比，返回 '前 X.X%' 字符串，失败返回空串"""
    try:
        r_num = int(str(rank).split("/")[0])
        total_num = int(rank_total) if rank_total else 0
        if total_num > 0:
            return f"前 {round(r_num / total_num * 100, 1)}%"
    except (ValueError, IndexError):
        pass
    return ""


def check_forbidden_phrases(text: str) -> List[str]:
    """检查文本中是否含禁用句，返回命中的禁用句列表"""
    hits = []
    for phrase in FORBIDDEN_PHRASES:
        if phrase in text:
            hits.append(phrase)
    return hits


# ============================================================
# ExchangeProgramDocBuilder 主类
# ============================================================

class ExchangeProgramDocBuilder:
    """交流项目申请书 docx 构建器（按 exchange_type 标识 3 类项目）"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []
        self.exchange_type = EXCHANGE_SCHOOL

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段，非列表返回空列表"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_heading(self, text):
        return add_section_heading(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, caption=""):
        return add_table_from_data(self.doc, headers, rows, col_widths, caption)

    # --------------------------------------------------------
    # 标题（固定"交流项目申请书"）
    # --------------------------------------------------------

    def _add_title(self):
        """标题：黑体二号居中，固定'交流项目申请书'"""
        add_title(self.doc, TITLE_TEXT)

    # --------------------------------------------------------
    # 称呼
    # --------------------------------------------------------

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号"""
        salutation = self._get("salutation", default="尊敬的评审委员会：")
        add_salutation_paragraph(self.doc, salutation)

    # --------------------------------------------------------
    # 开头段落
    # --------------------------------------------------------

    def _add_opening(self):
        """开头段落（约 75~125 字）：身份 + 申请事项 + 汇报句"""
        opening = self._get("opening", default="")
        if opening:
            self.add_para(opening)
            return
        name = self._get("name")
        current_college = self._get("current_college")
        current_major = self._get("current_major")
        grade = self._get("grade")
        target_university = self._get("target_university")
        exchange_duration = self._get("exchange_duration")
        parts = []
        if name and current_college and current_major and grade:
            parts.append(
                f"我是{current_college}{current_major}{grade}学生{name}，"
                f"现申请赴{target_university}参加{exchange_duration}交流项目。"
            )
        else:
            parts.append(
                f"现申请赴{target_university}参加{exchange_duration}交流项目。"
            )
        parts.append("现就交流项目申请事宜汇报如下：")
        self.add_para("".join(parts))

    # --------------------------------------------------------
    # 一、个人基本情况
    # --------------------------------------------------------

    def _add_basic_info(self):
        """个人基本情况（约 225~375 字）：身份 + 学业 + 语言 + 申请事项"""
        self.add_heading("一、个人基本情况")
        basic_info = self._get("basic_info", default="")
        if basic_info:
            if isinstance(basic_info, list):
                for p in basic_info:
                    self.add_para(p)
            else:
                self.add_para(basic_info)
            self._add_current_courses_table()
            return
        gpa = self._get("current_gpa", default="")
        weighted = self._get("current_weighted_avg", default="")
        rank = self._get("current_rank", default="")
        rank_total = self._get("current_rank_total", default="")
        course_count = self._get("current_course_count", default="")
        high_score_count = self._get("current_high_score_count", default="")
        language_score = self._get("language_score", default="")
        language_breakdown = self._get("language_breakdown", default="")
        research = self._get("research_experience", default="")
        competition = self._get("competition_experience", default="")
        target_university = self._get("target_university", default="")
        target_country = self._get("target_country", default="")
        target_qs_rank = self._get("target_qs_rank", default="")
        parts = []
        if gpa:
            gpa_str = f"在校期间 GPA {gpa}"
            if weighted:
                gpa_str += f"，加权平均分 {weighted}"
            if rank:
                gpa_str += f"，专业排名第 {rank}"
                if rank_total:
                    pct = format_rank_percent(rank, rank_total)
                    if pct:
                        gpa_str += f"（{pct}）"
            parts.append(gpa_str + "。")
        if course_count and high_score_count:
            parts.append(f"已修读 {course_count} 门课程，{high_score_count} 门 85 分以上。")
        if language_score:
            parts.append(f"语言成绩：{language_score}")
            if language_breakdown:
                parts.append(f"（{language_breakdown}）。")
            else:
                parts.append("。")
        if research:
            parts.append(research + "。")
        if competition:
            parts.append(competition + "。")
        if target_university and target_country and target_qs_rank:
            parts.append(
                f"本次申请赴{target_country}{target_university}参加交流项目，"
                f"该校{target_qs_rank}，与本人专业方向高度契合。"
            )
        if parts:
            self.add_para("".join(parts))
        self._add_current_courses_table()

    def _add_current_courses_table(self):
        """现专业主干课程表：4~6 门，含课程名/学分/成绩"""
        courses = self._get_list("current_core_courses")
        if not courses:
            return
        rows = []
        for c in courses:
            if not isinstance(c, dict):
                continue
            rows.append([str(c.get("name", "")), str(c.get("credit", "")), str(c.get("score", ""))])
        if rows:
            self.add_table(["课程名称", "学分", "成绩"], rows,
                           col_widths=[8.0, 3.0, 3.0], caption="现专业主干课程成绩：")

    # --------------------------------------------------------
    # 二、交流动机（4 维度）
    # --------------------------------------------------------

    # 4 维度配置：(field, prefix) 元组列表
    MOTIVATION_DIMS = [
        ("motivation_academic", "**学术深化方面**，"),
        ("motivation_vision", "**视野拓展方面**，"),
        ("motivation_language", "**语言提升方面**，"),
        ("motivation_career", "**职业规划方面**，"),
    ]

    def _add_motivation(self):
        """交流动机（约 375~625 字）：4 维度（学术深化/视野拓展/语言提升/职业规划）"""
        self.add_heading("二、交流动机")
        motivation = self._get("motivation", default="")
        if motivation:
            if isinstance(motivation, list):
                for p in motivation:
                    self.add_para(p)
            else:
                self.add_para(motivation)
            self._add_target_courses_table()
            self._add_target_teachers_info()
            return
        # 4 维度拼装
        self._add_dimension_paragraphs(self.MOTIVATION_DIMS)
        self._add_target_courses_table()
        self._add_target_teachers_info()

    def _add_dimension_paragraphs(self, dims: List[tuple]) -> None:
        """通用维度段落拼装：dims 为 (field, prefix) 元组列表"""
        for field, prefix in dims:
            val = self._get(field, default="")
            if val:
                self.add_para(prefix + val)

    def _add_target_courses_table(self):
        """目标院校核心课程表：≥5 门，含课程名/学分/主讲教师"""
        courses = self._get_list("target_courses")
        if not courses:
            return
        rows = []
        for c in courses:
            if not isinstance(c, dict):
                continue
            rows.append([
                str(c.get("name", "")),
                str(c.get("credit", "")),
                str(c.get("teacher", "")),
                str(c.get("semester", "")),
            ])
        if rows:
            self.add_table(
                ["课程名称", "学分", "主讲教师", "学期"],
                rows,
                col_widths=[5.5, 2.0, 4.0, 3.0],
                caption="目标院校核心课程：",
            )

    def _add_target_teachers_info(self):
        """目标院校导师信息：≥2 位，含姓名/职称/研究方向"""
        teachers = self._get_list("target_teachers")
        if not teachers:
            return
        parts = []
        for t in teachers:
            if not isinstance(t, dict):
                continue
            name, title, research = t.get("name", ""), t.get("title", ""), t.get("research", "")
            achievement = t.get("achievement", "")
            seg = name
            if title:
                seg += f"（{title}）"
            if research:
                seg += f"研究方向为{research}"
            if achievement:
                seg += f"，{achievement}"
            if seg:
                parts.append(seg + "；")
        if parts:
            parts[-1] = parts[-1].rstrip("；") + "。"
            self.add_para("目标院校重点关注导师：" + "".join(parts))

    # --------------------------------------------------------
    # 三、学习计划（3 阶段 + 表格）
    # --------------------------------------------------------

    # 3 阶段配置：(field, prefix) 元组列表
    PLAN_STAGES = [
        ("prep_plan", "**行前准备阶段**，"),
        ("abroad_plan", "**在外学习阶段**，"),
        ("return_plan", "**归国总结阶段**，"),
    ]

    def _add_learning_plan(self):
        """学习计划（约 450~750 字）：3 阶段（行前准备/在外学习/归国总结）"""
        self.add_heading("三、学习计划")
        learning_plan = self._get("learning_plan", default="")
        if learning_plan:
            if isinstance(learning_plan, list):
                for p in learning_plan:
                    self.add_para(p)
            else:
                self.add_para(learning_plan)
            self._add_study_plan_table()
            return
        # 3 阶段拼装
        self._add_dimension_paragraphs(self.PLAN_STAGES)
        self._add_study_plan_table()

    def _add_study_plan_table(self):
        """学习计划表：3 阶段各 1 行（4 列：阶段/时间/主要内容/目标）"""
        plans = self._get_list("study_plan_table")
        if not plans:
            return
        rows = []
        for p in plans:
            if not isinstance(p, dict):
                continue
            rows.append([
                str(p.get("stage", "")),
                str(p.get("time", "")),
                str(p.get("content", "")),
                str(p.get("target", "")),
            ])
        if rows:
            self.add_table(
                ["阶段", "时间", "主要内容", "目标"],
                rows,
                col_widths=[2.0, 3.5, 5.5, 3.0],
                caption="学习计划表：",
            )

    # --------------------------------------------------------
    # 四、文化交流意愿（3 维度）
    # --------------------------------------------------------

    # 3 维度配置：(field, prefix) 元组列表
    CULTURAL_DIMS = [
        ("cultural_integration", "**主动融入方面**，"),
        ("cultural_promotion", "**传播中华方面**，"),
        ("academic_exchange", "**学术交流方面**，"),
    ]

    def _add_cultural_exchange(self):
        """文化交流意愿（约 225~375 字）：3 维度（融入/传播/学术交流）"""
        self.add_heading("四、文化交流意愿")
        cultural_exchange = self._get("cultural_exchange", default="")
        if cultural_exchange:
            if isinstance(cultural_exchange, list):
                for p in cultural_exchange:
                    self.add_para(p)
            else:
                self.add_para(cultural_exchange)
            return
        # 3 维度拼装
        self._add_dimension_paragraphs(self.CULTURAL_DIMS)

    # --------------------------------------------------------
    # 五、安全保障（4 维度）
    # --------------------------------------------------------

    # 4 维度配置：(field, prefix) 元组列表
    SAFETY_DIMS = [
        ("safety_legal", "**法律遵守方面**，"),
        ("safety_insurance", "**保险购买方面**，"),
        ("safety_report", "**定期汇报方面**，"),
        ("safety_emergency", "**紧急预案方面**，"),
    ]

    def _add_safety(self):
        """安全保障（约 225~375 字）：4 维度（法律/保险/汇报/预案）"""
        self.add_heading("五、安全保障")
        safety = self._get("safety", default="")
        if safety:
            if isinstance(safety, list):
                for p in safety:
                    self.add_para(p)
            else:
                self.add_para(safety)
            return
        # 4 维度拼装
        self._add_dimension_paragraphs(self.SAFETY_DIMS)

    # --------------------------------------------------------
    # 结尾"此致 敬礼！"
    # --------------------------------------------------------

    def _add_ending(self):
        """结尾（约 75~125 字）：朴素表态 + 此致/敬礼！"""
        ending = self._get("ending", default="")
        if ending:
            self.add_para(ending)
        else:
            self.add_para(
                "以上是我的交流项目申请。无论审批结果如何，我都将继续努力学习，"
                "以更高标准要求自己。恳请评审委员会予以考虑。"
            )
        add_cizhi_paragraph(self.doc, "此致")
        add_jingli_paragraph(self.doc, "敬礼！")

    # --------------------------------------------------------
    # 落款
    # --------------------------------------------------------

    def _add_signature(self):
        """落款：右对齐，含申请人 + 日期"""
        self.doc.add_paragraph()  # 空一行
        name = self._get("name", default="")
        apply_date = self._get("apply_date", default="")
        if name:
            add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if apply_date:
            add_right_aligned_paragraph(self.doc, apply_date)

    # --------------------------------------------------------
    # 主构建方法
    # --------------------------------------------------------

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/开头/个人基本情况/交流动机/
        学习计划/文化交流意愿/安全保障/结尾/落款"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_opening()
            self._add_basic_info()
            self._add_motivation()
            self._add_learning_plan()
            self._add_cultural_exchange()
            self._add_safety()
            self._add_ending()
            self._add_signature()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 交流项目申请书已生成：{output_path}")
        return str(output_path)

    # --------------------------------------------------------
    # 数据校验
    # --------------------------------------------------------

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []
        # 交流项目类型校验
        etype = str(self._get("exchange_type", default=EXCHANGE_SCHOOL)).lower()
        if etype in VALID_EXCHANGE_TYPES:
            self.exchange_type = etype
        else:
            warnings.append(f"交流项目类型 {etype} 非标准值，默认按校际交换（school）处理")
            self.exchange_type = EXCHANGE_SCHOOL
        # P0 必采字段
        for key, label in [
            ("name", "申请人姓名"), ("current_college", "现学院"),
            ("current_major", "现专业"), ("grade", "年级"),
            ("target_university", "目标院校"), ("exchange_duration", "交流时长"),
            ("current_gpa", "现专业 GPA"), ("current_rank", "现专业排名"),
            ("language_score", "语言成绩"),
        ]:
            if not self._get(key):
                warnings.append(f"缺少 {label}（{key}）")
        # 排名校验
        rank_str = str(self._get("current_rank", default=""))
        if rank_str and "/" not in rank_str and not self._get("current_rank_total"):
            warnings.append(
                f"排名 '{rank_str}' 缺基数，应为 'X/N' 格式或补填 current_rank_total"
            )
        # 目标院校 QS 排名校验
        if not self._get("target_qs_rank"):
            warnings.append("缺少 目标院校 QS 排名（target_qs_rank），将无法体现学术深化动机")
        # 课程/师资/学习计划表数量校验
        checks = [
            ("current_core_courses", 4, "现专业主干课程", "4~6 门"),
            ("target_courses", 5, "目标院校核心课程", "至少 5 门"),
            ("target_teachers", 2, "目标院校导师", "至少 2 位"),
            ("study_plan_table", 3, "学习计划表", "3 阶段 3 行"),
        ]
        for key, min_n, label, hint in checks:
            items = self._get_list(key)
            if not items:
                warnings.append(f"缺少 {label}（{key}），将省略对应内容")
            elif len(items) < min_n:
                warnings.append(f"{label}仅 {len(items)} 项，建议{hint}")
        # 4 维度动机校验（至少 3 维度）
        motivation_fields = [
            ("motivation_academic", "学术深化动机"),
            ("motivation_vision", "视野拓展动机"),
            ("motivation_language", "语言提升动机"),
            ("motivation_career", "职业规划动机"),
        ]
        missing_motivation = [
            label for field, label in motivation_fields
            if not self._get(field)
        ]
        if not self._get("motivation") and len(missing_motivation) >= 2:
            warnings.append(
                f"4 维度交流动机缺 {len(missing_motivation)} 维度："
                f"{'、'.join(missing_motivation)}（建议至少展开 3 维度）"
            )
        # 3 阶段学习计划校验
        plan_fields = [
            ("prep_plan", "行前准备计划"),
            ("abroad_plan", "在外学习计划"),
            ("return_plan", "归国总结计划"),
        ]
        missing_plan = [
            label for field, label in plan_fields
            if not self._get(field)
        ]
        if not self._get("learning_plan") and missing_plan:
            warnings.append(
                f"3 阶段学习计划缺 {len(missing_plan)} 阶段："
                f"{'、'.join(missing_plan)}"
            )
        # 安全保障 4 维度校验
        safety_fields = [
            ("safety_legal", "法律遵守"),
            ("safety_insurance", "保险购买"),
            ("safety_report", "定期汇报"),
            ("safety_emergency", "紧急预案"),
        ]
        missing_safety = [
            label for field, label in safety_fields
            if not self._get(field)
        ]
        if not self._get("safety") and missing_safety:
            warnings.append(
                f"安全保障 4 维度缺 {len(missing_safety)} 维度："
                f"{'、'.join(missing_safety)}"
            )
        # CSC 公派特殊校验
        if self.exchange_type == EXCHANGE_CSC:
            safety_report = str(self._get("safety_report", default=""))
            if safety_report and ("回国" not in safety_report and "CSC" not in safety_report):
                warnings.append(
                    "CSC 公派型建议在 safety_report 中明确 CSC 协议回国服务 2 年义务承诺"
                )
        # 禁用句检测
        for key in [
            "motivation_academic", "motivation_vision",
            "motivation_language", "motivation_career",
            "motivation", "prep_plan", "abroad_plan", "return_plan",
        ]:
            val = self._get(key, default="")
            if isinstance(val, str) and val:
                hits = check_forbidden_phrases(val)
                if hits:
                    warnings.append(
                        f"字段 {key} 含禁用句：{'、'.join(hits)}（详见 SKILL.md §10.1）"
                    )
        self.warnings = warnings
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据（校际交换，2000 字版）
# ============================================================

DEFAULT_DATA = {
    "name": "张明", "student_id": "2023123456", "gender": "男",
    "current_college": "计算机科学与技术学院", "current_major": "计算机科学与技术",
    "grade": "2023 级 大二", "class_name": "计科 2301 班",
    "target_university": "伦敦大学学院（UCL）", "target_country": "英国",
    "target_qs_rank": "2025 年 QS 全球第 9，计算机科学学科全球第 8",
    "exchange_type": "school", "exchange_duration": "1 学年（2025.09~2026.06）",
    "phone": "138XXXXXXXX", "email": "zhangming@xxx.edu.cn",
    "apply_date": "2025 年 3 月 15 日",
    "salutation": "尊敬的评审委员会：",
    "current_gpa": "3.85/4.0", "current_weighted_avg": "88.5",
    "current_rank": "5/120", "current_rank_total": "120",
    "current_course_count": "18", "current_high_score_count": "14",
    "current_core_courses": [
        {"name": "数据结构", "credit": "4", "score": "94"},
        {"name": "算法分析与设计", "credit": "4", "score": "92"},
        {"name": "操作系统", "credit": "4", "score": "90"},
        {"name": "计算机网络", "credit": "3", "score": "88"},
        {"name": "机器学习", "credit": "3", "score": "95"},
    ],
    "language_score": "雅思 6.5（2024.09 取得，有效期至 2026.09）",
    "language_breakdown": "听力 7.0 / 阅读 7.0 / 写作 6.0 / 口语 5.5",
    "motivation_academic": (
        "伦敦大学学院（UCL）计算机科学专业在 2025 年 QS 学科排名全球第 8，"
        "其开设的《Machine Learning》（4 学分）由 Alexei Efros 教授"
        "（ACM Fellow，研究方向为计算机视觉）主讲，"
        "团队近 3 年在 CVPR、ICCV 发表论文 25 篇。"
        "同时拥有 Gatsby Computational Neuroscience Unit 等顶级实验室，"
        "定期举办 ICML、NeurIPS 等顶级会议讲座。"
        "我希望通过交流深入学习其计算机视觉与机器学习方向，"
        "弥补本校在该方向的课程空白。"
    ),
    "motivation_vision": (
        "UCL 每年 3 月举办 International Week 跨文化交流活动，"
        "吸引来自 150 多个国家的学生参与。"
        "我计划参与该活动，与来自不同文化背景的学生组成 5 人小组，"
        "共同完成 1 个跨文化主题项目。"
        "此外，我计划加入 UCL Chinese Society，"
        "在传播中华文化的同时，深入了解英国本土文化。"
        "通过这些跨文化活动，我预期能提升跨文化沟通能力，培养全球视野。"
    ),
    "motivation_language": (
        "当前雅思总分 6.5（2024.09 取得，有效期至 2026.09）。"
        "通过 UCL 交流，我目标将雅思总分提升至 7.0 以上，"
        "其中口语提升至 6.5。具体提升方式包括："
        "每周参与 UCL Language Centre 学术英语课程（2 小时/周），"
        "加入 UCL English Speaking Union 每周练习口语（3 小时/周），"
        "与当地学生组成语言交换伙伴（每周 2 次）。"
        "同时通过沉浸式英语环境，提升日常交流能力。"
    ),
    "motivation_career": (
        "我的长期职业目标是成为跨国科技公司（如 Google、Microsoft）的算法工程师，"
        "从事计算机视觉与医学影像 AI 研究。"
        "UCL 计算机视觉方向全球领先，"
        "导师 Alexei Efros 教授的团队近 3 年在 CVPR、ICCV 发表论文 25 篇。"
        "通过交流学习与参与团队研究，我将积累前沿学术经历，"
        "为未来申请海外博士项目（目标：斯坦福、CMU）奠定基础。"
        "归国后，我计划将交流经历转化为科研论文（目标 1 篇 CVPR/ICCV），"
        "并申请清华或北大直博项目，最终实现跨国科技公司算法工程师的职业目标。"
    ),
    "target_courses": [
        {"name": "Machine Learning", "credit": "4", "teacher": "Alexei Efros", "semester": "2025 Fall"},
        {"name": "Computer Vision", "credit": "4", "teacher": "John S.", "semester": "2025 Fall"},
        {"name": "Academic English", "credit": "2", "teacher": "Sarah M.", "semester": "2025 Fall"},
        {"name": "Deep Learning", "credit": "4", "teacher": "Yann L.", "semester": "2026 Spring"},
        {"name": "British Culture", "credit": "2", "teacher": "David W.", "semester": "2026 Spring"},
    ],
    "target_teachers": [
        {
            "name": "Alexei Efros", "title": "教授", "research": "计算机视觉",
            "achievement": "ACM Fellow，近 3 年 CVPR/ICCV 论文 25 篇",
        },
        {
            "name": "John S.", "title": "副教授", "research": "机器学习与推荐系统",
            "achievement": "近 3 年 NeurIPS/ICML 论文 12 篇",
        },
    ],
    "target_labs": "Gatsby Computational Neuroscience Unit",
    "prep_plan": (
        "行前准备阶段（2025.07~2025.09，3 个月）："
        "①语言强化：参加新东方雅思冲刺班（80 课时，2025.07~08），目标雅思 7.0；"
        "②文化预习：阅读《英国文化简史》《UCL 校史》《British Academic Culture》3 本书；"
        "③课程预选：已与 UCL International Office 沟通，预选 5 门课程共 16 学分；"
        "④学分认定：已与本校教务处沟通，5 门课程可转换为 12 学分。"
    ),
    "abroad_plan": (
        "在外学习阶段（2025.09~2026.06，2 学期）："
        "第 1 学期修读《Machine Learning》《Computer Vision》《Academic English》共 10 学分；"
        "第 2 学期修读《Deep Learning》《British Culture》"
        "+ 参与 Alexei Efros 教授课题组研究项目（4 学分研究实习），共 10 学分。"
        "每周时间分配：授课 16 小时、自习 20 小时、课题组研究 8 小时、"
        "学术讲座 4 小时、跨文化交流 6 小时。"
        "学术活动：参加 UCL Computer Science 每周研讨会（约 30 场/学期）"
        "+ Gatsby Unit 每月开放日。"
        "阶段性产出：第 1 学期完成 1 篇课程论文《Deep Learning for Medical Image Analysis》，"
        "第 2 学期目标以第二作者发表 1 篇 CVPR Workshop 论文。"
    ),
    "return_plan": (
        "归国总结阶段（2026.07~2026.09，2 个月）："
        "①学分转换：5 门 UCL 课程共 16 学分已确认为本校 12 学分"
        "（含 8 学分专业核心课 + 4 学分专业选修课），不影响毕业学分要求；"
        "②经验分享：拟在本校计算机学院举办 2 场交流分享会"
        "（面向本科生与研究生），分享 UCL 学习经历与申请经验，"
        "并在本校国际处网站发布 1 篇交流心得；"
        "③后续学业衔接：基于交流期间积累的计算机视觉研究经历，"
        "归国后申请加入本校人工智能研究院李教授课题组，"
        "目标在 2026.12 前完成 CVPR 论文投稿，并申请清华或北大直博项目。"
    ),
    "weekly_schedule": "授课 16h + 自习 20h + 课题组 8h + 讲座 4h + 跨文化 6h",
    "study_plan_table": [
        {
            "stage": "行前准备", "time": "2025.07~09",
            "content": "雅思冲刺 + 文化预习 + 课程预选 + 学分认定",
            "target": "雅思 7.0 + 5 门预选课程 + 12 学分认定",
        },
        {
            "stage": "在外学习", "time": "2025.09~2026.06",
            "content": "5 门课程 + 课题组研究 + 学术讲座",
            "target": "16 学分 + CVPR Workshop 论文 1 篇",
        },
        {
            "stage": "归国总结", "time": "2026.07~09",
            "content": "学分转换 + 经验分享 + 加入本校课题组",
            "target": "12 学分转换 + 2 场分享会 + 加入李教授课题组",
        },
    ],
    "cultural_integration": (
        "UCL 每年 3 月举办 International Week，"
        "我计划参与并与来自 150 多个国家的学生组成 5 人小组完成跨文化项目。"
        "同时申请家庭寄宿（Homestay）3 个月，深入体验英国家庭生活。"
    ),
    "cultural_promotion": (
        "我计划加入 UCL Chinese Society，"
        "参与组织 2026 年春节文化展示活动（含书法、茶艺、汉服展示），"
        "并向英国同学介绍中国传统文化。"
        "同时在学联组织的中秋节、端午节活动中担任志愿者。"
    ),
    "academic_exchange": (
        "我计划参加 UCL Computer Science 每周研讨会（约 30 场/学期），"
        "与其他交换生组成学习小组每周交流 1 次，"
        "并在学期末举办 1 场面向本校师生的线上分享会，"
        "介绍 UCL 的学术研究与文化特色。"
    ),
    "safety_legal": (
        "我承诺严格遵守英国法律与 UCL 校规，"
        "不参与任何非法活动，按时完成学业任务，"
        "尊重英国当地风俗习惯与文化传统，"
        "维护中国留学生良好形象。"
    ),
    "safety_insurance": (
        "我已购买中国平安'留学无忧'海外保险，"
        "含医疗费用（保额 50 万元人民币）、"
        "意外伤害（保额 30 万元人民币）、"
        "紧急救援（保额 100 万元人民币），"
        "保险期限覆盖 2025.09~2026.06 整个交流期。"
    ),
    "safety_report": (
        "我承诺每月 5 日前向本校国际处提交上月学习进度报告，"
        "紧急情况 24 小时内电话汇报，"
        "每学期末提交交流总结报告。"
        "同时与本校导师保持每两周 1 次的视频沟通，确保学业衔接。"
    ),
    "safety_emergency": (
        "我已设立国内紧急联系人（父亲 XX，电话 138XXXXXXXX）"
        "与国外紧急联系人（UCL International Office，电话 +44-20-7679-2000），"
        "记录中国驻英使馆联系方式（+44-20-7299-4049）"
        "与英国报警急救电话（999），"
        "并制定紧急情况应对流程"
        "（先报警 → 再联系使馆 → 再通知家人 → 最后通知本校）。"
    ),
    "gaokao_year": "2023 年", "gaokao_score": "632 分（山东）",
    "admission_mode": "普通批次",
    "research_experience": (
        "2024.09~2025.03 跟随本校李教授参与医疗影像 AI 项目 6 个月，"
        "独立完成数据预处理模块，处理 1200 例 CT 数据。"
    ),
    "competition_experience": (
        "2024.11 全国大学生数学建模竞赛省级二等奖（队长，负责建模与编程）；"
        "2025.04 中国大学生计算机设计大赛省级一等奖。"
    ),
    "word_count_target": "standard",
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="交流项目申请书 docx 生成器（按 exchange_type 标识 3 类项目）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n\n"
            "JSON 字段定义详见 SKILL.md §12 完整字段清单。\n"
            "exchange_type: school（校际交换）/ college（院际交换）/ csc（CSC 公派）\n"
            "word_count_target: short（1500 字）/ standard（2000 字，默认）/ long（2500 字）"
        ),
    )
    parser.add_argument(
        "--data", type=str, default=None,
        help="数据 JSON 文件路径（与 --demo 二选一）",
    )
    parser.add_argument(
        "--out", type=str, required=True,
        help="输出 docx 文件路径",
    )
    parser.add_argument(
        "--demo", action="store_true",
        help="使用内置示例数据生成演示文档",
    )
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        etype_label = {
            EXCHANGE_SCHOOL: "校际交换",
            EXCHANGE_COLLEGE: "院际交换",
            EXCHANGE_CSC: "CSC 公派",
        }.get(data.get("exchange_type", EXCHANGE_SCHOOL), "校际交换")
        print(
            f"ℹ️ 使用内置示例数据生成演示文档"
            f"（{etype_label}，{data.get('target_university', '')}，"
            f"{data.get('exchange_duration', '')}）"
        )
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ExchangeProgramDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
