#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
优秀毕业生申请书 docx 生成器

支持省级 / 校级优秀毕业生申请书生成。评审跨度为大学 4 年（不是 1 学年），
正文 2000~3000 字，书信体格式。必写五维度：思想品德、学业成绩、科研创新、
社会实践、综合素质，并体现 4 年成长叙事与毕业去向。

按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 标题：黑体二号，居中（按规格动态填充"省级/校级优秀毕业生申请书"）
- 称呼：顶格，宋体小四，全角冒号
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 分学年 GPA 表（4 列）、科研创新成果表（5 列）：宋体五号，居中
- "此致"另起一行空两格，"敬礼！"另起一行顶格
- 落款：右对齐

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第三章信息采集清单。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ----- 字体与格式常量 -----

FONT_SONG = "宋体"
FONT_HEI = "黑体"
SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号
SIZE_XIAO_WU = Pt(9)        # 小五
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

def set_run_font(run, font_name: str = FONT_SONG, font_size=SIZE_XIAO_SI,
                 bold: bool = False, color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)

def set_cell_font(cell, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)

def set_cell_text(cell, text: str, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)

def add_paragraph_with_format(doc, text: str, font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI, bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True, line_spacing: float = 1.5,
    space_before: float = 0, space_after: float = 0):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p

def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, space_before=12, space_after=12)

def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=indent, line_spacing=1.5)

def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进），宋体小四"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)

def add_section_heading(doc, text: str):
    """正文小节标题（一、二、三…）：黑体小四加粗，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
        line_spacing=1.5, space_before=6, space_after=3)

def add_cizhi_paragraph(doc, text: str = "此致"):
    """'此致'另起一行，空两格（首行缩进 2 字符）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True, line_spacing=1.5)

def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格（不缩进）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)

def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=False, line_spacing=1.5)

def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None, caption: str = ""):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    if caption:
        add_paragraph_with_format(
            doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG, font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_name=FONT_SONG, font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)

def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)

def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ----- ApplicationDocBuilder 主类 -----
class ApplicationDocBuilder:
    """优秀毕业生申请书 docx 构建器（支持省级 / 校级）"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段，非列表返回空列表"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_heading(self, text):
        return add_section_heading(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, caption=""):
        return add_table_from_data(self.doc, headers, rows, col_widths, caption)

    def _add_title(self):
        """标题：黑体二号居中，按 apply_level 动态填充"省级/校级优秀毕业生申请书" """
        level = str(self._get("apply_level", default="provincial")).lower()
        title = "校级优秀毕业生申请书" if level in ("school", "校级", "校") else "省级优秀毕业生申请书"
        add_title(self.doc, title)

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号"""
        salutation = self._get("salutation", default="尊敬的领导：")
        add_salutation_paragraph(self.doc, salutation)

    def _add_opening(self):
        """开头段落（150~200 字）：身份 + 申报规格 + 4 年综合数据 3 项 + 毕业去向 + 进入正文"""
        opening = self._get("opening", default="")
        if opening:
            self.add_para(opening)
            return
        name, college, major, grade = self._get("name"), self._get("college"), self._get("major"), self._get("grade")
        apply_year = self._get("apply_year", default="2025 届")
        level = str(self._get("apply_level", default="provincial")).lower()
        level_text = "省级" if level not in ("school", "校级", "校") else "校级"
        gpa_4yr, rank_4yr = self._get("gpa_4yr"), self._get("rank_4yr")
        destination = self._get("destination", default="")
        dest_school = self._get("destination_school", default="")
        parts = []
        if name and college and major and grade:
            parts.append(f"我是{college}{major}{grade}学生{name}，特申报{apply_year}{level_text}优秀毕业生。")
        else:
            parts.append(f"特申报{apply_year}{level_text}优秀毕业生。")
        data_parts = []
        if gpa_4yr:
            data_parts.append(f"大学 4 年加权 GPA {gpa_4yr}")
        if rank_4yr:
            data_parts.append(f"专业排名第 {rank_4yr}")
        papers = self._get_list("papers")
        if papers and isinstance(papers[0], dict):
            p = papers[0]
            data_parts.append(f"以{p.get('author_order', '第一作者')}发表{p.get('level', '')}论文 1 篇")
        elif self._get_list("research_projects"):
            data_parts.append("主持国家级大创 1 项")
        if dest_school:
            data_parts.append(f"保研至{dest_school}")
        elif destination:
            data_parts.append(destination)
        if data_parts:
            parts.append("；".join(data_parts) + "。")
        parts.append("现将大学四年情况汇报如下：")
        self.add_para("".join(parts))

    def _add_ideology(self):
        """思想品德（300~400 字）：政治面貌 + 入党轨迹 + 学生干部履职 + 思想觉悟 + 荣誉"""
        self.add_heading("一、思想品德")
        ideology = self._get("ideology", default="")
        if ideology:
            if isinstance(ideology, list):
                for p in ideology:
                    self.add_para(p)
            else:
                self.add_para(ideology)
            return
        political = self._get("political_status", default="共青团员")
        party_history = self._get("party_history", default="")
        party_activities = self._get_list("party_activities")
        parts = []
        if political:
            if "党员" in political and "预备" not in political:
                parts.append(f"作为一名{political}，我始终坚持学习党的理论，认真学习习近平新时代中国特色社会主义思想与党的二十大和二十届三中全会精神。")
            elif "预备党员" in political:
                parts.append("作为一名中共预备党员，我认真学习习近平新时代中国特色社会主义思想，深入学习党的二十大和二十届三中全会精神，时刻以正式党员标准要求自己。")
            elif "积极分子" in political:
                parts.append("作为一名入党积极分子，我认真学习习近平新时代中国特色社会主义思想，时刻以党员标准要求自己。")
            else:
                parts.append(f"作为一名{political}，我拥护中国共产党的领导，认真学习党的创新理论。")
        if party_history:
            parts.append(party_history)
        if parts:
            self.add_para("".join(parts))
        # 第 2 段：学生干部履职
        position_history = self._get_list("position_history")
        position = self._get("position", default="")
        position_work = self._get("position_work", default="")
        pos_parts = []
        if position_history:
            for ph in position_history:
                if not isinstance(ph, dict):
                    continue
                role, term, work = ph.get("role", ""), ph.get("term", ""), ph.get("work", "")
                seg = f"担任{role}" if role else ""
                if term:
                    seg += f"（{term}）"
                if work:
                    seg += f"，{work}"
                if seg:
                    pos_parts.append(seg + "。")
            if pos_parts:
                self.add_para("".join(pos_parts))
        elif position:
            seg = f"担任{position}"
            if position_work:
                seg += f"，主要工作：{position_work}"
            self.add_para(seg + "。")
        # 第 3 段：思想觉悟 + 党课/主题党日
        if party_activities:
            self.add_para("；".join(party_activities) + "。")
        # 第 4 段：荣誉列表（思想品德类）
        honors_ideology = self._get_list("honors_ideology")
        h_parts = []
        for h in honors_ideology:
            if not isinstance(h, dict):
                continue
            time, name, level = h.get("time", ""), h.get("name", ""), h.get("level", "")
            seg = (f"{time} " if time else "") + name
            if level:
                seg += f"（{level}）"
            if seg:
                h_parts.append(seg + "；")
        if h_parts:
            h_parts[-1] = h_parts[-1].rstrip("；") + "。"
            self.add_para("在校期间获得的主要荣誉：" + "".join(h_parts))

    def _add_academics(self):
        """学业成绩（400~600 字）：4 年 GPA + 排名 + 分学年 GPA 表 + 主干课程 + 英语计算机 + 奖学金 + 毕设"""
        self.add_heading("二、学业成绩")
        academics = self._get("academics", default="")
        if academics and isinstance(academics, str):
            self.add_para(academics)
            self._add_yearly_gpa_table()
            self._add_core_courses_table()
            self._add_academics_summary()
            return
        gpa_4yr, weighted_4yr = self._get("gpa_4yr"), self._get("weighted_avg_4yr")
        rank_4yr, rank_total = self._get("rank_4yr"), self._get("rank_total")
        parts = []
        if gpa_4yr:
            gpa_str = f"大学 4 年加权 GPA {gpa_4yr}"
            if weighted_4yr:
                gpa_str += f"，加权平均分 {weighted_4yr}"
            if rank_4yr:
                gpa_str += f"，专业排名第 {rank_4yr}"
                if rank_total:
                    try:
                        r_num = int(str(rank_4yr).split("/")[0]) if "/" in str(rank_4yr) else int(rank_4yr)
                        total_num = int(rank_total)
                        if total_num > 0:
                            gpa_str += f"（前 {round(r_num / total_num * 100, 1)}%）"
                    except (ValueError, IndexError):
                        pass
            parts.append(gpa_str + "。")
        if parts:
            self.add_para("".join(parts))
        self._add_yearly_gpa_table()
        self._add_core_courses_table()
        self._add_academics_summary()

    def _add_yearly_gpa_table(self):
        """分学年 GPA 表（4 列：学年/GPA/加权平均分/排名）"""
        yearly = self._get_list("yearly_gpa")
        rows = []
        for y in yearly:
            if isinstance(y, dict):
                rows.append([str(y.get("year", "")), str(y.get("gpa", "")), str(y.get("weighted", "")), str(y.get("rank", ""))])
        if rows:
            self.add_table(["学年", "GPA", "加权平均分", "专业排名"], rows,
                           col_widths=[3.0, 3.0, 4.0, 3.0], caption="分学年 GPA 情况：")

    def _add_core_courses_table(self):
        """主干课程表（5~8 门）"""
        courses = self._get_list("core_courses")
        rows = []
        for c in courses:
            if isinstance(c, dict):
                rows.append([str(c.get("name", "")), str(c.get("credit", "")), str(c.get("score", ""))])
        if rows:
            self.add_table(["课程名称", "学分", "成绩"], rows,
                           col_widths=[8.0, 3.0, 3.0], caption="主干课程成绩：")

    def _add_academics_summary(self):
        """学业末段：英语计算机 + 奖学金 + 毕设"""
        summary = self._get("academics_summary", default="")
        if summary:
            self.add_para(summary)
            return
        cet4, cet6, toefl = self._get("cet4"), self._get("cet6"), self._get("toefl")
        computer_level = self._get("computer_level")
        scholarship_summary = self._get("scholarship_summary", default="")
        course_highlight = self._get("course_highlight", default="")
        thesis = self._get("thesis", default="")
        parts = []
        if course_highlight:
            parts.append(course_highlight + "。")
        lang_parts = []
        if cet4:
            lang_parts.append(f"CET-4 {cet4} 分")
        if cet6:
            lang_parts.append(f"CET-6 {cet6} 分")
        if toefl:
            lang_parts.append(f"TOEFL {toefl} 分")
        if lang_parts:
            parts.append("、".join(lang_parts) + "，能熟练阅读英文文献；")
        if computer_level:
            parts.append(f"计算机{computer_level}。")
        if scholarship_summary:
            parts.append(scholarship_summary + "。")
        if thesis:
            parts.append(thesis + "。")
        if parts:
            self.add_para("".join(parts))

    def _add_research(self):
        """科研创新（400~600 字）：大创 + 论文 + 专利 + 学科竞赛 + 成果表"""
        self.add_heading("三、科研创新")
        research = self._get("research", default="")
        if research and isinstance(research, str):
            self.add_para(research)
            self._add_research_table()
            return
        self._add_research_projects()
        self._add_papers()
        self._add_patents()
        self._add_competitions()
        self._add_research_table()

    def _add_research_projects(self):
        """大创/科研立项"""
        projects = self._get_list("research_projects")
        proj_parts = []
        for p in projects:
            if not isinstance(p, dict):
                continue
            name, level = p.get("name", ""), p.get("level", "")
            role, duration, output = p.get("role", ""), p.get("duration", ""), p.get("output", "")
            seg = f"{role}《{name}》" if role and name else (f"《{name}》" if name else "")
            if level:
                seg += f"（{level}）"
            if duration:
                seg += f"，{duration}"
            if output:
                seg += f"，{output}"
            if seg:
                proj_parts.append(seg + "。")
        if proj_parts:
            self.add_para("".join(proj_parts))

    def _add_papers(self):
        """论文列表"""
        papers = self._get_list("papers")
        paper_parts = []
        for p in papers:
            if not isinstance(p, dict):
                continue
            title, journal, level = p.get("title", ""), p.get("journal", ""), p.get("level", "")
            author_order, time, if_val = p.get("author_order", ""), p.get("time", ""), p.get("impact_factor", "")
            seg = (f"以{author_order}在" if author_order else "") + (f"《{journal}》" if journal else "")
            if level:
                seg += f"（{level}）"
            if if_val:
                seg += f"，IF {if_val}"
            seg += "发表论文"
            if title:
                seg += f"《{title}》"
            if time:
                seg += f"，{time} 见刊"
            if seg:
                paper_parts.append(seg + "；")
        if paper_parts:
            paper_parts[-1] = paper_parts[-1].rstrip("；") + "。"
            self.add_para("".join(paper_parts))

    def _add_patents(self):
        """专利列表"""
        patents = self._get_list("patents")
        pat_parts = []
        for p in patents:
            if not isinstance(p, dict):
                continue
            name, ptype, status, order = p.get("name", ""), p.get("type", ""), p.get("status", ""), p.get("order", "")
            seg = (f"作为{order}申请{ptype}专利" if order and ptype else (f"申请{ptype}专利" if ptype else ""))
            if name:
                seg += f"《{name}》"
            if status:
                seg += f"（{status}）"
            if seg:
                pat_parts.append(seg + "；")
        if pat_parts:
            pat_parts[-1] = pat_parts[-1].rstrip("；") + "。"
            self.add_para("".join(pat_parts))

    def _add_competitions(self):
        """学科竞赛段"""
        competitions = self._get_list("competitions")
        comp_parts = []
        for c in competitions:
            if not isinstance(c, dict):
                continue
            time, name, award, role = c.get("time", ""), c.get("name", ""), c.get("award", ""), c.get("role", "")
            seg = (f"{time} " if time else "") + name
            if award:
                seg += f" {award}"
            if role:
                seg += f"（{role}）"
            if seg:
                comp_parts.append(seg + "；")
        if comp_parts:
            comp_parts[-1] = comp_parts[-1].rstrip("；") + "。"
            self.add_para("学科竞赛方面，" + "".join(comp_parts))

    def _add_research_table(self):
        """科研创新成果汇总表（5 列：类别/名称/级别/角色/时间）"""
        rows = []
        for p in self._get_list("research_projects"):
            if isinstance(p, dict):
                rows.append(["大创", str(p.get("name", "")), str(p.get("level", "")), str(p.get("role", "")), str(p.get("duration", ""))])
        for p in self._get_list("papers"):
            if isinstance(p, dict):
                rows.append(["论文", str(p.get("title", "")), str(p.get("level", "")), str(p.get("author_order", "")), str(p.get("time", ""))])
        for p in self._get_list("patents"):
            if isinstance(p, dict):
                rows.append(["专利", str(p.get("name", "")), str(p.get("type", "")), str(p.get("order", "")), str(p.get("status", ""))])
        for c in self._get_list("competitions"):
            if isinstance(c, dict):
                rows.append(["竞赛", str(c.get("name", "")), str(c.get("award", "")), str(c.get("role", "")), str(c.get("time", ""))])
        if rows:
            self.add_table(["类别", "名称", "级别", "角色", "时间"], rows,
                           col_widths=[1.5, 5.0, 2.5, 2.0, 3.0], caption="科研创新成果汇总：")

    def _add_practice(self):
        """社会实践（300~500 字）：学生工作 + 三下乡/志愿服务 + 实习"""
        self.add_heading("四、社会实践")
        practice = self._get("practice", default="")
        if practice and isinstance(practice, str):
            self.add_para(practice)
            return
        self._add_student_work()
        self._add_volunteer()
        self._add_internship()

    def _add_student_work(self):
        """学生工作段"""
        student_work = self._get("student_work", default="")
        if student_work:
            self.add_para(student_work)
            return
        position = self._get("position", default="")
        position_work = self._get("position_work", default="")
        if position:
            seg = f"担任{position}"
            if position_work:
                seg += f"，主要工作：{position_work}"
            self.add_para(seg + "。")

    def _add_volunteer(self):
        """三下乡 / 志愿服务段"""
        volunteer_hours = self._get("volunteer_hours", default="")
        volunteer_detail = self._get("volunteer_detail", default="")
        sanxiaxiang = self._get_list("sanxiaxiang")
        parts = []
        if volunteer_hours:
            seg = f"大学 4 年累计志愿服务时长 {volunteer_hours} 小时"
            if volunteer_detail:
                seg += "：" + volunteer_detail
            parts.append(seg + "。")
        if sanxiaxiang:
            sx_parts = []
            for s in sanxiaxiang:
                if not isinstance(s, dict):
                    continue
                time, place, role, output = s.get("time", ""), s.get("place", ""), s.get("role", ""), s.get("output", "")
                seg = (f"{time} " if time else "") + (f"赴{place}" if place else "")
                if role:
                    seg += f"，{role}"
                if output:
                    seg += f"，{output}"
                if seg:
                    sx_parts.append(seg + "；")
            if sx_parts:
                sx_parts[-1] = sx_parts[-1].rstrip("；") + "。"
                parts.append("参与三下乡" + str(len(sanxiaxiang)) + " 次：" + "".join(sx_parts))
        if parts:
            self.add_para("".join(parts))

    def _add_internship(self):
        """实习经历段"""
        internships = self._get_list("internships")
        intern_parts = []
        for it in internships:
            if not isinstance(it, dict):
                continue
            company, position, duration, output = it.get("company", ""), it.get("position", ""), it.get("duration", ""), it.get("output", "")
            seg = f"{duration} 在{company}" if duration and company else (f"在{company}" if company else "")
            if position:
                seg += f"担任{position}"
            if output:
                seg += f"，{output}"
            if seg:
                intern_parts.append(seg + "。")
        if intern_parts:
            self.add_para("".join(intern_parts))

    def _add_quality(self):
        """综合素质（200~300 字）：文体特长 + 班级评价 + 兴趣爱好"""
        self.add_heading("五、综合素质")
        quality = self._get("quality", default="")
        if quality:
            if isinstance(quality, list):
                for p in quality:
                    self.add_para(p)
            else:
                self.add_para(quality)
            return
        talents = self._get_list("talents")
        tal_parts = []
        for t in talents:
            if not isinstance(t, dict):
                continue
            name, level, experience = t.get("name", ""), t.get("level", ""), t.get("experience", "")
            seg = name
            if level:
                seg += f"（{level}）"
            if experience:
                seg += f"，{experience}"
            if seg:
                tal_parts.append(seg + "；")
        if tal_parts:
            tal_parts[-1] = tal_parts[-1].rstrip("；") + "。"
            self.add_para("文体特长方面，" + "".join(tal_parts))
        class_eval = self._get("class_evaluation", default="")
        dorm_honor = self._get("dorm_honor", default="")
        class_honor = self._get("class_honor", default="")
        eval_parts = []
        if class_eval:
            eval_parts.append(class_eval + "。")
        if dorm_honor:
            eval_parts.append(dorm_honor + "。")
        if class_honor:
            eval_parts.append(class_honor + "。")
        if eval_parts:
            self.add_para("".join(eval_parts))

    def _add_ending(self):
        """结尾（80~150 字）：4 年成长总结 + 毕业去向呼应 + 朴素表态 + 此致/敬礼！"""
        ending = self._get("ending", default="")
        if ending:
            self.add_para(ending)
        else:
            dest_school = self._get("destination_school", default="")
            destination = self._get("destination", default="")
            if dest_school:
                dest_seg = f"即将前往{dest_school}攻读研究生，我将带着母校的教诲继续努力。"
            elif destination:
                dest_seg = f"即将{destination}，我将带着母校的教诲继续努力。"
            else:
                dest_seg = "无论未来身在何处，我都将带着母校的教诲继续努力。"
            self.add_para(
                "以上是我大学四年的成长汇报。" + dest_seg +
                "无论评选结果如何，我都将以优秀毕业生的标准严格要求自己。"
            )
        add_cizhi_paragraph(self.doc, "此致")
        add_jingli_paragraph(self.doc, "敬礼！")

    def _add_signature(self):
        """落款：右对齐，含申请人 + 日期"""
        self.doc.add_paragraph()
        name = self._get("name", default="")
        apply_date = self._get("apply_date", default="")
        if name:
            add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if apply_date:
            add_right_aligned_paragraph(self.doc, apply_date)

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/开头/思想/学业/科研/实践/素质/结尾/落款，生成 docx"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_opening()
            self._add_ideology()
            self._add_academics()
            self._add_research()
            self._add_practice()
            self._add_quality()
            self._add_ending()
            self._add_signature()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申请书已生成：{output_path}")
        return str(output_path)

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []
        p0_fields = [("name", "申请人姓名"), ("college", "学院"), ("major", "专业"),
                     ("grade", "年级"), ("gpa_4yr", "4 年加权 GPA"), ("rank_4yr", "4 年专业排名")]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")
        yearly = self._get_list("yearly_gpa")
        if not yearly:
            warnings.append("缺少 分学年 GPA（yearly_gpa），将省略分学年 GPA 表")
        elif len(yearly) < 4:
            warnings.append(f"分学年 GPA 仅 {len(yearly)} 学年，建议 4 学年（大一~大四）")
        rank_str = str(self._get("rank_4yr", default=""))
        if rank_str and "/" not in rank_str:
            if not self._get("rank_total"):
                warnings.append(f"排名 '{rank_str}' 缺基数，应为 'X/N' 格式或补填 rank_total")
        try:
            rank_num = int(rank_str.split("/")[0]) if "/" in rank_str else 0
            rank_total_num = 0
            if "/" in rank_str:
                rank_total_num = int(rank_str.split("/")[1])
            elif self._get("rank_total"):
                rank_total_num = int(self._get("rank_total"))
            if rank_num and rank_total_num:
                pct = rank_num / rank_total_num * 100
                if pct > 10:
                    warnings.append(f"排名前 {pct:.1f}% 不满足优秀毕业生前 10% 门槛，建议补足荣誉后再申")
        except (ValueError, IndexError):
            pass
        if not self._get_list("competitions") and not self._get_list("papers") \
                and not self._get_list("research_projects"):
            warnings.append("缺少 大创/论文/竞赛任一项，建议至少 2 项校级及以上成果")
        if not self._get("destination") and not self._get("destination_school"):
            warnings.append("缺少 毕业去向（destination/destination_school），必填项")
        self.warnings = warnings
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ----- 默认示例数据 -----
DEFAULT_DATA = {
    "name": "张明",
    "student_id": "2021123456",
    "gender": "男",
    "college": "计算机科学与技术学院",
    "major": "计算机科学与技术",
    "grade": "2021 级 大四",
    "class_name": "计科 2101 班",
    "political_status": "中共预备党员",
    "phone": "138XXXXXXXX",
    "apply_level": "provincial",
    "apply_year": "2025 届",
    "apply_date": "2025 年 5 月 20 日",
    "salutation": "尊敬的领导：",
    "gpa_4yr": "3.85/4.0",
    "weighted_avg_4yr": "90.5",
    "rank_4yr": "1/87",
    "rank_total": "87",
    "yearly_gpa": [
        {"year": "大一", "gpa": "3.60", "weighted": "85.2", "rank": "5/87"},
        {"year": "大二", "gpa": "3.80", "weighted": "88.5", "rank": "3/87"},
        {"year": "大三", "gpa": "3.90", "weighted": "91.0", "rank": "1/87"},
        {"year": "大四", "gpa": "3.95", "weighted": "92.5", "rank": "1/87"},
    ],
    "core_courses": [
        {"name": "数据结构", "credit": "4", "score": "92"},
        {"name": "操作系统", "credit": "4", "score": "94"},
        {"name": "计算机网络", "credit": "3", "score": "91"},
        {"name": "数据库原理", "credit": "3", "score": "93"},
        {"name": "计算机组成原理", "credit": "4", "score": "90"},
        {"name": "算法设计与分析", "credit": "3", "score": "95"},
    ],
    "course_highlight": "数据结构 92、操作系统 94、算法设计与分析 95，专业核心课平均 93 分，全部 90+",
    "cet4": "568",
    "cet6": "542",
    "computer_level": "二级 C 语言（优秀）、三级数据库技术",
    "scholarship_summary": "4 年获国家奖学金 1 次（2024 学年）、校级一等奖学金 3 次（2022/2023/2025 学年）",
    "thesis": "毕业设计《基于对比学习的法律问答系统》已通过中期检查，导师为张教授",
    "party_history": "2022.09 递交入党申请书，2023.03 列为入党积极分子，2024.06 转为中共预备党员。",
    "party_activities": [
        "参加学院分党校第 8 期培训班（2023.09-2023.12）结业",
        "2024.10 参与主题党日活动'红色教育基地走访'，撰写调研报告 1 份（约 3000 字）",
        "提交思想汇报 6 篇",
    ],
    "ideology": "",
    "academics": "",
    "research": "",
    "practice": "",
    "quality": "",
    "position": "学院学生会主席（2024.09-2025.06）",
    "position_work": ("组织学院迎新晚会、运动会等大型活动 8 场，累计参与 2000+ 人次；"
                     "牵头建立学院学生服务平台，覆盖学习辅导、心理咨询、就业指导 3 大模块；"
                     "推动学院学生会获评 2024 年度'校级优秀学生组织'"),
    "position_history": [
        {"role": "班级团支书", "term": "2021.09-2023.06（4 学期）",
         "work": "组织主题团日活动 12 场，覆盖同学 600 余人次，所在团支部获评 2024 年度校级五四红旗团支部"},
        {"role": "学院学生会主席", "term": "2024.09-2025.06（1 学年）",
         "work": "组织大型活动 8 场，累计参与 2000+ 人次，推动学院学生会获评'校级优秀学生组织'"},
    ],
    "honors_ideology": [
        {"time": "2024.10", "name": "校级优秀学生干部", "level": "校级"},
        {"time": "2023.05", "name": "校级优秀共青团员", "level": "校级"},
    ],
    "research_projects": [
        {"name": "分布式光伏故障智能诊断系统", "level": "国家级大创项目",
         "role": "主持", "duration": "2024.03-2025.03",
         "output": "发表 SCI 论文 1 篇（第一作者），申请发明专利 1 项，项目结题评估优秀"},
    ],
    "papers": [
        {"title": "基于对比学习的法律问答系统", "journal": "Knowledge-Based Systems",
         "level": "SCI 二区", "author_order": "第一作者",
         "time": "2025.03", "impact_factor": "5.2"},
        {"title": "A Contrastive Learning Approach for Legal QA",
         "journal": "EMNLP 2024", "level": "CCF-B 类会议",
         "author_order": "第二作者", "time": "2024.11"},
    ],
    "patents": [
        {"name": "一种基于对比学习的法律问答方法", "type": "发明专利",
         "status": "实质审查中", "order": "第一发明人"},
    ],
    "competitions": [
        {"name": "全国大学生数学建模竞赛", "award": "国家二等奖",
         "time": "2024.05", "role": "队长，负责整体建模与论文主笔"},
        {"name": "中国大学生计算机设计大赛", "award": "省级一等奖",
         "time": "2024.11", "role": "核心成员，负责前端开发"},
    ],
    "student_work": "",
    "volunteer_hours": "320",
    "volunteer_detail": ("担任图书馆管理员（2023.09-2024.06，每月 8 小时）；"
                        "组织班级'一对一'帮扶活动，服务同学 30 余人次"),
    "sanxiaxiang": [
        {"time": "2023.07", "place": "云南怒江",
         "role": "教育扶贫支教", "output": "授课 32 课时，覆盖学生 80 余人"},
        {"time": "2024.07", "place": "甘肃定西",
         "role": "电商助农调研", "output": "走访 5 个村、访谈 60 户，撰写调研报告 1 份（约 1.5 万字）"},
    ],
    "internships": [
        {"company": "字节跳动", "position": "推荐算法组实习生",
         "duration": "2024.07-2024.09",
         "output": "参与推荐系统优化项目，提出基于对比学习的用户表征方法，CTR 提升 3.2%，获实习优秀评级"},
    ],
    "talents": [
        {"name": "钢琴", "level": "十级（中央音乐学院考级）",
         "experience": "担任校合唱团团长 2 学期，组织校内演出 6 场"},
    ],
    "class_evaluation": "与同学相处融洽，曾帮助室友完成 1 次重要实验调试",
    "dorm_honor": "宿舍连续 4 学期获评'文明宿舍'",
    "class_honor": "所在班级获评 2024 年度校级先进班集体",
    "destination": "",
    "destination_school": "清华大学计算机科学与技术系，人工智能方向",
    "ending": "",
}


# ----- CLI 入口 -----
def main():
    parser = argparse.ArgumentParser(
        description="优秀毕业生申请书 docx 生成器（支持省级/校级）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=("示例：\n"
                "  python build.py --data data.json --out output.docx\n"
                "  python build.py --demo --out demo.docx\n"
                "\n"
                "JSON 字段定义详见 SKILL.md 第三章信息采集清单。"),
    )
    parser.add_argument("--data", type=str, default=None, help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True, help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true", help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)

if __name__ == "__main__":
    main()
