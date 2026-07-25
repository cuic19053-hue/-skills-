#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
大学生暑期"三下乡"社会实践-科技服务类立项申报书 docx 生成器

格式标准：A4，页边距上下 2.54cm 左右 2.5cm；正文宋体小四 1.5 倍行距首行缩进 2 字符；
一级标题黑体三号居中；二级标题黑体小三左对齐；表格宋体五号居中。

科技服务类专属：实施方案按场次 7 列表（场次/日期/地点/对象/内容/人数/专业要求）；
服务方案 4 列结构化表（需求/匹配能力/服务流程/预期解决方案）+ 持续帮扶机制小节。

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx
JSON 字段详见 SKILL.md 第十一章。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)
SIZE_SAN = Pt(16)
SIZE_XIAO_SAN = Pt(15)
SIZE_SI = Pt(14)
SIZE_XIAO_SI = Pt(12)
SIZE_WU = Pt(10.5)

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG, font_size=SIZE_XIAO_SI,
                 bold: bool = False, color: Optional[RGBColor] = None):
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name=FONT_SONG, font_size=SIZE_WU,
                  bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(doc, text: str, font_name: str = FONT_SONG,
        font_size=SIZE_XIAO_SI, bold: bool = False,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent: bool = True,
        line_spacing: float = 1.5, space_before: float = 0, space_after: float = 0):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text):
    """一级标题：黑体三号居中，段前后 12pt"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI,
        font_size=SIZE_SAN, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False, space_before=12, space_after=12)


def add_heading_level2(doc, text):
    """二级标题：黑体小三左对齐，段前后 6pt"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI,
        font_size=SIZE_XIAO_SAN, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, space_before=6, space_after=6)


def add_heading_level3(doc, text):
    """三级标题：宋体四号加粗，左对齐"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG,
        font_size=SIZE_SI, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, space_before=6, space_after=3)


def add_body_paragraph(doc, text, indent=True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent, line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None,
                        left_align_cols: Optional[List[int]] = None):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            align = (WD_ALIGN_PARAGRAPH.LEFT if left_align_cols and j in left_align_cols
                     else WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(cells[j], val, alignment=align)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def fmt_money(amount):
    """金额格式化：整数转字符串 + 元"""
    try:
        return f"{int(amount)} 元"
    except (ValueError, TypeError):
        return str(amount)


def safe_int(value, default=0):
    """安全转整数"""
    try:
        return int(str(value).strip())
    except (ValueError, TypeError):
        return default


def extract_int_from_str(text):
    """从字符串中提取第一个整数（用于场次表合计人次计算）"""
    if not text:
        return 0
    for tok in str(text).split():
        digits = "".join(c for c in tok if c.isdigit())
        if digits:
            try:
                return int(digits)
            except ValueError:
                continue
    return 0


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """三下乡社会实践-科技服务类立项申报书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def add_h1(self, text):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text):
        return add_heading_level2(self.doc, text)

    def add_h3(self, text):
        return add_heading_level3(self.doc, text)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_table(self, headers, rows, col_widths=None, left_align_cols=None):
        return add_table_from_data(self.doc, headers, rows, col_widths,
                                   left_align_cols)

    def add_page_break(self):
        add_page_break(self.doc)

    # 封面

    def _add_cover(self):
        """封面：黑体二号标题 + 4 行下划线信息"""
        for _ in range(3):
            self.doc.add_paragraph()
        title = "大学生暑期\u201c三下乡\u201d社会实践活动立项申报书"
        add_paragraph_with_format(self.doc, title, font_name=FONT_HEI,
            font_size=SIZE_ER, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False, space_before=12, space_after=12)
        theme = self._get("theme", default="乡村振兴")
        add_paragraph_with_format(self.doc, f"（{theme}主题·科技服务类）",
            font_name=FONT_HEI, font_size=SIZE_SAN,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_after=24)
        for _ in range(3):
            self.doc.add_paragraph()
        info_items = [
            ("项目名称", self._get("team_name", default="赴 XX 县数字赋能科技服务团")),
            ("团队名称", self._get("team_name")),
            ("申报单位", self._get("college")),
            ("申报日期", self._get("apply_date")),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI, font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_size=SIZE_SI)
            run_value.underline = True
        self.add_page_break()

    # 一、团队基本信息

    def _add_team_info(self):
        """一、团队基本信息表（7 行 2 列）"""
        self.add_h1("一、团队基本信息")
        team = self._get("team_info", default={})
        if not isinstance(team, dict):
            team = {}
        g = self._get
        leader = team.get("leader", f"{g('leader_name')} / {g('leader_id')} / "
                          f"{g('leader_major')} / {g('leader_grade')} / {g('leader_phone')}")
        advisor = team.get("advisor", f"{g('advisor_name')} / {g('advisor_title')} / "
                           f"{g('advisor_phone')} / {g('advisor_with_team', '随队')} / "
                           f"{g('advisor_major', '软件工程')}")
        rows = [
            ["团队名称", team.get("team_name", g("team_name", ""))],
            ["实践主题", team.get("theme", g("theme", ""))],
            ["实践地点", team.get("location", g("location", ""))],
            ["实践时间", team.get("practice_time", g("practice_time", ""))],
            ["团队人数", team.get("team_size", g("team_size", ""))],
            ["队长", leader],
            ["指导教师", advisor],
        ]
        self.add_table(["项目", "内容"], rows, col_widths=[4.5, 11.5])

    # 二、团队成员信息表

    def _add_members_table(self):
        """二、团队成员信息表（5 列）"""
        self.add_h1("二、团队成员信息表")
        members = self._get("members", default=[])
        if members and isinstance(members, list):
            rows = []
            for m in members:
                if not isinstance(m, dict):
                    continue
                rows.append([m.get("name", ""), m.get("id", ""),
                             m.get("major", ""), m.get("role", ""),
                             m.get("phone", "")])
            self.add_table(["姓名", "学号", "专业年级", "团队分工", "联系方式"],
                           rows, col_widths=[2.5, 2.8, 3.5, 3.5, 3.7])
        else:
            self.add_para("（请填写团队成员信息表，每人一行：姓名 / 学号 / 专业年级 / "
                          "团队分工 / 联系方式。科技服务类建议 5~12 人小团队，"
                          "专业互补覆盖硬件/软件/网络。分工必须具体到\u201c做什么\u201d，"
                          "如\u201c主讲 IT 维修 + 维修操作\u201d。）")

    # 三、实践主题与背景

    def _add_theme_background(self):
        """三、实践主题与背景（400~600 字，3 段）"""
        self.add_h1("三、实践主题与背景")
        bg = self._get("theme_background", default=[])
        if isinstance(bg, str):
            bg = [bg]
        if bg:
            for para in bg:
                self.add_para(para)
        else:
            self.add_h2("（一）团中央主题对应")
            self.add_para("（请填写当年团中央发布的实践主题方向，以及本项目对应切入的"
                          "维度，150 字左右。如：2025 年团中央\u201c乡村振兴 青春建功\u201d主题，"
                          "本项目切入科技服务维度，聚焦\u201c数字赋能\u201d。）")
            self.add_h2("（二）服务子方向选定理由")
            self.add_para("（请填写服务子方向选定理由，200 字左右。结构：当地数字化水平"
                          "数据 + 团队专业匹配度 + 服务可行性。如：XX 镇 60 岁以上老人 "
                          "1240 人，智能手机持有率 68% 但能独立完成视频通话仅 23%；"
                          "电脑家庭拥有率 31%，故障自修率不足 10%。团队 5 人含计算机 3 + "
                          "通信 2，专业匹配度高。）")
            self.add_h2("（三）实践对象基本情况")
            self.add_para("（请填写实践对象基本情况，150 字左右，必须有量化数据。"
                          "如：XX 镇常住人口 8600 人，60 岁以上老人 1240 人；智能手机持有率"
                          " 68%，视频通话独立使用率 23%；数据来源：XX 镇团委 2025 年 5 月摸底调查。）")

    # 四、实践目的与意义

    def _add_purpose_significance(self):
        """四、实践目的与意义（300~500 字，2 段）"""
        self.add_h1("四、实践目的与意义")
        ps = self._get("purpose_significance", default=[])
        if isinstance(ps, str):
            ps = [ps]
        if ps:
            for para in ps:
                self.add_para(para)
        else:
            self.add_h2("（一）对当地")
            self.add_para("（请填写对当地的意义，150~250 字。结构：通过 5 场科技服务覆盖"
                          " 200 人次 + 修复电脑 50 台 + 帮助 20 户村电商开店 + 建立"
                          "\u201cXX 县数字赋能\u201d答疑微信群 + 每月 1 次远程答疑 + 每季度 1 次"
                          "现场回访，形成长效帮扶机制，重点解决\u201c老人不会用智能手机\u201d"
                          "\u201c电脑故障无人修\u201d\u201c村电商不会开店\u201d三类问题。）")
            self.add_h2("（二）对学生")
            self.add_para("（请填写对学生的意义，150~250 字。结构：专业知识应用（把课堂"
                          "所学带到基层）+ 服务意识培养 + 基层国情认识 + 长效机制建设。）")

    # 五、实践内容与实施方案【重点】

    def _add_implementation_plan(self):
        """五、实践内容与实施方案（800~1200 字，3 子节 + 按场次表）"""
        self.add_h1("五、实践内容与实施方案")

        # （一）实践形式
        self.add_h2("（一）实践形式")
        form = self._get("implementation_plan", default={})
        if not isinstance(form, dict):
            form = {}
        form_text = form.get("form", "")
        if form_text:
            self.add_para(form_text)
        else:
            self.add_para("本项目实践形式为科技服务类，采用\u201c按场次集中服务 + 一对一"
                          "上门服务 + 持续远程帮扶\u201d三段式实施流程。按场次集中服务"
                          "覆盖 3 个行政村 5 场 200 人次，一对一上门服务针对行动不便"
                          "老人与村电商上门维修，持续远程帮扶建立答疑群 + 月度远程 + "
                          "季度现场。子方向为科技服务（IT 维修 + 电商培训 + 数字素养），"
                          "与调研/支教/宣讲类的核心区别在于\u201c专业输出\u201d而非数据采集、"
                          "教学方法或一对多传达。")
        self.add_h2("（二）服务对象需求分析")
        needs = form.get("needs_analysis", "")
        # 服务对象需求分析默认按紧迫程度排序，对应场次表
        if needs:
            self.add_para(needs)
        else:
            self.add_para("通过 XX 镇团委 2025 年 5 月摸底调查 + 实地预调研 2 天访谈"
                          "村干部 5 人 + 村民 20 人，梳理出 5 项核心需求并按紧迫程度"
                          "排序：1. 智能手机基础操作（视频通话/扫码支付）— 留守老人 — "
                          "高紧迫 — 高匹配；2. 电脑故障维修与系统重装 — 村民/村电商 — "
                          "高紧迫 — 高匹配；3. 网络反诈识别 — 全体村民 — 高紧迫 — 中匹配"
                          "（需法律指导）；4. 电商开店与直播带货基础 — 村电商 — 中紧迫 — "
                          "中匹配（需电商指导）；5. 数字素养（政务 App/健康码） — 村民 — "
                          "中紧迫 — 高匹配。高优先级 3 项安排主力场次 4 场，中优先级 2 项"
                          "各安排 1 场。")
        self.add_h2("（三）按场次实施安排")
        schedule = form.get("schedule", [])
        if schedule and isinstance(schedule, list):
            rows = []
            total_people = 0
            total_sessions = 0
            for i, s in enumerate(schedule, 1):
                if not isinstance(s, dict):
                    continue
                total_sessions += 1
                people_count = safe_int(s.get("people", 0)) or extract_int_from_str(s.get("people", ""))
                total_people += people_count
                rows.append([
                    s.get("session", f"第 {i} 场"),
                    s.get("date", ""), s.get("location", ""),
                    s.get("audience", ""), s.get("content", ""),
                    s.get("people", ""), s.get("requirement", ""),
                ])
            rows.append([
                "合计", f"{total_sessions} 场", "3 个村", "—",
                "IT 维修 + 电商培训 + 数字素养 + 反诈",
                f"{total_people} 人次", "—",
            ])
            self.add_table(
                ["场次", "日期", "地点", "服务对象", "服务内容", "人数", "专业要求"],
                rows, col_widths=[1.6, 1.8, 2.5, 2.5, 3.2, 1.4, 3.0],
                left_align_cols=[4])
            self.add_para("每天结束当晚召开日例会，整理当日服务记录单（含服务对象"
                          "姓名/问题/解决方案/满意度），调整次日方案。每场服务前 3 天"
                          "与当地对接人确认场地、设备、对象人数；服务后 7 天电话回访，"
                          "30 天现场回访。")
        else:
            self.add_para("（请填写按场次实施安排表，5~8 行覆盖整个实践期。每行 7 列："
                          "场次 / 日期（精确到日，含上下午）/ 地点（精确到村文化礼堂/"
                          "村委会）/ 服务对象（含人数）/ 服务内容（具体到\u201c做什么\u201d）"
                          "/ 人数 / 专业要求（学科背景 + 主讲/助教/维修/记录/后勤分工）。"
                          "底部合计行场次总数 × 单场人次 = 总服务人次，必须与预期成果一致。）")

    # 六、服务方案【科技服务类专属重点】

    def _add_service_plan(self):
        """六、服务方案（4 列结构化表格 + 持续帮扶机制小节）"""
        self.add_h1("六、服务方案")
        # 服务方案为科技服务类专属栏目，含 4 列表 + 持续帮扶机制
        # 4 列：需求 / 匹配专业能力 / 服务流程 / 预期解决方案
        sp = self._get("service_plan", default={})
        if not isinstance(sp, dict):
            sp = {}
        self.add_h2("（一）服务子方向")
        subdirection = sp.get("subdirection", "")
        if subdirection:
            self.add_para(subdirection)
        else:
            self.add_para("本项目服务子方向为科技服务，覆盖 IT 维修、电商培训、数字素养"
                          "三大模块。IT 维修包括电脑故障检测、系统重装、硬件更换、软件"
                          "优化；电商培训包括平台选择、注册开店、商品上架、直播带货基础；"
                          "数字素养包括智能手机操作、政务 App 使用、健康码使用、网络反诈"
                          "识别。")
        self.add_h2("（二）需求 × 能力匹配 × 服务流程 × 解决方案")
        needs_table = sp.get("needs_table", [])
        rows = []
        if needs_table and isinstance(needs_table, list):
            for n in needs_table:
                if isinstance(n, dict):
                    rows.append([n.get("need", ""), n.get("capability", ""),
                                 n.get("process", ""), n.get("solution", "")])
        else:
            rows = [
                ["智能手机基础操作", "计算机/通信，熟悉 Android/iOS",
                 "调研手机型号 → 一对一教学 → 发放手册 → 加答疑群",
                 "60 位老人 80% 独立完成视频通话 + 操作手册 100 份 + 答疑群 1 个"],
                ["电脑故障维修", "计算机，硬件/软件能力",
                 "检测故障 → 维修/重装 → 测试 → 填写维修单 → 加答疑群",
                 "修复 50 台（系统 30+硬件 15+优化 5）+ 维护手册 100 份 + 答疑群 1 个"],
                ["电商开店基础", "计算机/通信，电商知识（指导教师补充）",
                 "平台选择 → 注册开店 → 商品上架 → 直播基础",
                 "帮助 20 户开店 + 录制教学视频 5 个 + 答疑群 1 个"],
                ["网络反诈识别", "计算机/通信 + 法律（指导教师补充）",
                 "案例讲解 → 一对一答疑 → 安装反诈 App → 加群",
                 "200 人次培训 + 反诈单页 200 份 + 反诈 App 安装 150 部"],
            ]
        self.add_table(["需求", "匹配专业能力", "服务流程", "预期解决方案"],
                       rows, col_widths=[3.5, 3.5, 5.5, 3.5],
                       left_align_cols=[2, 3])
        self.add_h2("（三）持续帮扶机制")
        sustain = sp.get("sustain_mechanism", "")
        if sustain:
            self.add_para(sustain)
        else:
            self.add_para("建立\u201c线上 + 线下\u201d长效服务机制：1. 线上答疑群——建立"
                          "\u201cXX 县数字赋能\u201d答疑微信群，5 名团队成员轮值（每人每周"
                          " 1 天），2 小时内响应；2. 月度远程答疑——每月最后一个周末视频"
                          "连线集中答疑，预计全年 12 次；3. 季度现场回访——每季度 1 次"
                          "现场回访（2~3 名队员赴地），重点解决疑难问题，预计全年 4 次；"
                          "4. 学年度接力——下一年度暑期三下乡接力服务，由本届队长对接"
                          "下届队长，9 月前完成交接，形成\u201cXX 县数字赋能\u201d长效品牌。")

    # 七、安全保障预案【重点】

    def _add_safety_plan(self):
        """七、安全保障预案（300~500 字，6 段 + 服务专属安全）"""
        self.add_h1("七、安全保障预案")
        plan = self._get("safety_plan", default=[])
        if isinstance(plan, str):
            plan = [plan]
        if plan and isinstance(plan, list):
            for para in plan:
                self.add_para(para)
        else:
            self.add_h2("（一）出行安全")
            self.add_para("（请填写出行安全，80 字：交通方式（统一购买高铁票）、住宿选择"
                          "（政府招待所 2 人一间男女分区）、每日 18:00 前向指导教师报平安。）")
            self.add_h2("（二）人身安全")
            self.add_para("（请填写人身安全，80 字：防暑（正午 11:00-14:00 不外出）、防疫"
                          "（每日早晚测温）、防骗（不携带大量现金、单人不离队）。）")
            self.add_h2("（三）应急联系人")
            self.add_para("（请填写应急联系人，60 字：3 名——指导教师、队长、学院团委，"
                          "附电话。当地对接：XX 县团委刘书记 136XXXXXXXX。）")
            self.add_h2("（四）应急流程")
            self.add_para("（请填写应急流程，120 字：分级响应。一般事件（< 1000 元）："
                          "队长 → 指导教师 → 学院存档。重大事件（≥ 1000 元 / 治安 / 灾害）："
                          "队长 → 指导教师 → 学院 → 110/120，30 分钟首报，2 小时书面报告。）")
            self.add_h2("（五）保险购买")
            self.add_para("（请填写保险购买，60 字：全员购买短期意外险，保额 30 万元，"
                          "保费 15 元/人，保单号 PA20250715001-005，覆盖 7 月 14-22 日。）")
            self.add_h2("（六）服务类专属安全")
            self.add_para("（请填写服务类专属安全，120 字：维修类带电作业安全（断电操作、"
                          "防静电、工具绝缘、防误触）；电商培训类隐私保护（不收集村民"
                          "身份证号、银行卡号）；反诈咨询类信息核实（与当地公安反诈中心"
                          "对接案例）；医疗/法律类按相关规范执行，复杂案件转介专业机构。）")

    # 八、预期成果

    def _add_expected_results(self):
        """八、预期成果（必须可量化）"""
        self.add_h1("八、预期成果")
        outcomes = self._get("expected_results", default=[])
        if isinstance(outcomes, str):
            outcomes = [outcomes]
        if outcomes:
            for o in outcomes:
                self.add_para(f"• {o}", indent=False)
        else:
            self.add_para("• 服务场次：5 场（A 村 2 + B 村 2 + C 村 1）", indent=False)
            self.add_para("• 服务人次：200 人次（留守老人 40 + 村民 90 + 村电商 20 + 干部 50）", indent=False)
            self.add_para("• 解决方案：3 套（IT 维修 + 电商开店 + 数字素养培训）", indent=False)
            self.add_para("• 服务记录：签到表 5 份 + 维修记录单 50 份 + 满意度反馈表 200 份", indent=False)
            self.add_para("• 持续帮扶机制：答疑微信群 + 月度远程 12 次/年 + 季度回访 4 次/年 + 学年度接力", indent=False)
            self.add_para("• 新闻稿 5 篇（中青网 1 + 校团委公众号 2 + 学院公众号 2）", indent=False)
            self.add_para("• 短视频 3 个 + 纪录短片 1 个（10 分钟，用于校内汇报）", indent=False)

    # 九、经费预算

    def _add_budget(self):
        """九、经费预算（3 列表格：科目/金额/计算依据）"""
        self.add_h1("九、经费预算")
        items = self._get("budget_items", default=[])
        if items and isinstance(items, list):
            rows = []
            total = 0
            for b in items:
                if not isinstance(b, dict):
                    continue
                amount_num = safe_int(b.get("amount", 0))
                total += amount_num
                rows.append([b.get("item", ""), fmt_money(amount_num), b.get("basis", "")])
            rows.append(["合计", fmt_money(total), ""])
            self.add_table(["预算科目", "金额", "计算依据"], rows,
                           col_widths=[3.5, 3.0, 9.5], left_align_cols=[2])
        else:
            self.add_para("（请填写经费预算，5 类标准科目：交通费 / 食宿费 / 物资费 / "
                          "印刷费 / 其他。每项金额非整数，附计算依据。科技服务类物资费"
                          "占比 30~40%。示例：交通费 2860 元 = 高铁 130 × 5 × 2 + 包车 "
                          "300 × 4；物资费 1500 元 = U 盘 50×20 + 维修工具 400 + 培训手册 "
                          "100×6 + 反诈单页 200×0.6 + 小礼品 60。）")

    # 十、宣传计划

    def _add_publicity_plan(self):
        """十、宣传计划（3 类媒体）"""
        self.add_h1("十、宣传计划")
        plan = self._get("publicity_plan", default=[])
        if isinstance(plan, str):
            plan = [plan]
        if plan and isinstance(plan, list):
            for p in plan:
                self.add_para(p)
        else:
            self.add_h2("（一）校内媒体")
            self.add_para("校团委公众号 2 篇、校报 1 篇、学院公众号 2 篇。"
                          "发布时间节点：7.16 启动报道、7.19 中期进展、7.21 结项报道。")
            self.add_h2("（二）校外媒体")
            self.add_para("中青网 1 篇、XX 县电视台 1 条、XX 日报 1 篇。"
                          "中青网稿件由队长审核后投递，附服务照片 5~8 张。")
            self.add_h2("（三）新媒体")
            self.add_para("B 站 3 个短视频、抖音 3 个短视频、微信视频号 2 个。"
                          "每个视频 3~5 分钟，内容覆盖服务过程、村民故事、团队风采。")

    # 十一、指导教师 + 学院团委意见

    def _add_review_section(self):
        """十一/十二/十三 签字栏（双栏 + 可选学校审批）"""
        def _sign_block(title, sign_label):
            self.add_h1(title)
            for _ in range(6):
                self.doc.add_paragraph()
            self.add_para(
                f"{sign_label}：____________________    日期：______年____月____日",
                indent=False)
        _sign_block("十一、指导教师意见", "指导教师签字")
        _sign_block("十二、学院团委意见", "学院盖章")
        if self._get("include_school_approval", default=False):
            _sign_block("十三、学校审批意见", "学校盖章")
        special = self._get("tuancentral_special", default="")
        if special:
            self.add_h1("附：团中央专项报送说明")
            self.add_para(f"对应专项：{special}。本项目与专项主题契合度高，预期形成"
                          "优秀服务案例 1 个、典型解决方案 1 套，报送团中央专项工作组。")

    # 主构建方法

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 11 栏目，生成 docx。返回实际保存路径。"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_cover()
            self._add_team_info()
            self._add_members_table()
            self._add_theme_background()
            self._add_purpose_significance()
            self._add_implementation_plan()
            self._add_service_plan()
            self._add_safety_plan()
            self._add_expected_results()
            self._add_budget()
            self._add_publicity_plan()
            self._add_review_section()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申报书已生成：{output_path}")
        return str(output_path)

    # 数据校验（含科技服务类专属校验）

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）。
        校验分四类：P0 必填字段、服务专属校验、安全预案校验、经费预算校验。"""
        warnings = []
        for key, name in [("team_name", "团队名称"), ("theme", "实践主题"),
                          ("location", "实践地点"), ("practice_time", "实践时间"),
                          ("leader_name", "队长姓名"), ("college", "申报单位")]:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        team = self._get("team_info", default={})
        if not isinstance(team, dict):
            team = {}
        if not team.get("team_size") and not self._get("team_size"):
            warnings.append("缺少 团队人数（team_size）")
        if not team.get("advisor") and not self._get("advisor_name"):
            warnings.append("缺少 指导教师（advisor_name）")
        if not self._get("theme_background"):
            warnings.append("缺少 实践主题与背景（theme_background），将使用占位文本")
        if not self._get("purpose_significance"):
            warnings.append("缺少 实践目的与意义（purpose_significance），将使用占位文本")

        impl = self._get("implementation_plan", default={})
        if not isinstance(impl, dict):
            impl = {}
        if not impl.get("schedule"):
            warnings.append("缺少 按场次实施安排（implementation_plan.schedule），将使用占位文本——评审会扣大分")
        else:
            total_people = 0
            for i, s in enumerate(impl.get("schedule", []), 1):
                if not isinstance(s, dict):
                    continue
                if not s.get("date") or not s.get("location"):
                    warnings.append(f"按场次表第 {i} 行缺少日期或地点")
                if not s.get("content"):
                    warnings.append(f"按场次表第 {i} 行缺少服务内容")
                if not s.get("requirement"):
                    warnings.append(f"按场次表第 {i} 行缺少专业要求——评审无法判断匹配度")
                total_people += safe_int(s.get("people", 0)) or extract_int_from_str(s.get("people", ""))
            if total_people > 0:
                self._service_total_people = total_people
        if not impl.get("form"):
            warnings.append("缺少 实践形式（implementation_plan.form），建议明确为科技服务类")
        if not impl.get("needs_analysis"):
            warnings.append("缺少 服务对象需求分析（implementation_plan.needs_analysis），评审会扣需求匹配度分")

        sp = self._get("service_plan", default={})
        if not isinstance(sp, dict):
            sp = {}
        if not sp.get("subdirection"):
            warnings.append("缺少 服务子方向（service_plan.subdirection），建议明确为科技/医疗/法律/农业之一")
        if not sp.get("needs_table"):
            warnings.append("缺少 需求×能力匹配×流程×方案表（service_plan.needs_table），评审会扣专业性分")
        sustain = sp.get("sustain_mechanism", "")
        if not sustain:
            warnings.append("缺少 持续帮扶机制（service_plan.sustain_mechanism），评审会扣可持续性分——服务类核心加分项")
        elif "长效" in sustain and "线上" not in sustain and "线下" not in sustain:
            warnings.append("持续帮扶机制只有'长效'空话，无线上/线下具体安排——评审会扣大分")

        if not self._get("safety_plan"):
            warnings.append("缺少 安全保障预案（safety_plan），将使用占位文本——安全预案不达标可能一票否决")
        else:
            plan = self._get("safety_plan", default=[])
            plan_text = "\n".join(plan) if isinstance(plan, list) else str(plan)
            if "保险" not in plan_text:
                warnings.append("安全预案未提及保险购买情况——一票否决项")
            if "应急联系" not in plan_text:
                warnings.append("安全预案未列明应急联系人")
            if "应急流程" not in plan_text:
                warnings.append("安全预案未列明应急流程")

        results = self._get("expected_results", default=[])
        if not results:
            warnings.append("缺少 预期成果（expected_results），将使用占位文本")
        else:
            results_text = "\n".join(results) if isinstance(results, list) else ""
            if "服务" not in results_text and "人次" not in results_text:
                warnings.append("预期成果未以服务人次为主——科技服务类核心产出缺失")
            if "持续" not in results_text and "帮扶" not in results_text and "长效" not in results_text:
                warnings.append("预期成果未提及持续帮扶机制——服务类加分项缺失")

        items = self._get("budget_items", default=[])
        if items:
            total = sum(safe_int(b.get("amount", 0)) for b in items if isinstance(b, dict))
            budget_total_num = safe_int(str(self._get("budget_total", default="")).strip(), default=-1)
            if budget_total_num >= 0 and total != budget_total_num:
                warnings.append(f"预算合计 {total} 元 与申请经费 {budget_total_num} 元不一致")
        else:
            warnings.append("缺少 经费预算（budget_items），将使用占位文本")

        members = self._get("members", default=[])
        for i, m in enumerate(members if isinstance(members, list) else [], 1):
            if not isinstance(m, dict):
                continue
            if not m.get("role"):
                warnings.append(f"成员 {m.get('name', f'#{i}')} 缺少分工")
            if not m.get("major"):
                warnings.append(f"成员 {m.get('name', f'#{i}')} 缺少专业年级——评审无法判断专业匹配度")

        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        self.warnings = warnings
        return warnings


# ============================================================
# 默认示例数据
# ============================================================

DEFAULT_DATA = {
    "team_name": "赴 XX 县数字赋能科技服务团",
    "theme": "乡村振兴",
    "location": "XX 省 XX 县 XX 镇 3 个行政村（A/B/C 村）",
    "practice_time": "2025.07.15-07.21（7 天）",
    "team_size": "5 人",
    "leader_name": "张三", "leader_id": "202212345",
    "leader_major": "计算机科学与技术 2022 级", "leader_grade": "大三",
    "leader_phone": "138XXXXXXXX",
    "advisor_name": "李教授", "advisor_title": "副教授",
    "advisor_phone": "139XXXXXXXX", "advisor_with_team": "随队",
    "advisor_major": "软件工程",
    "college": "计算机科学与技术学院",
    "apply_date": "2025 年 5 月 20 日",
    "team_info": {
        "team_name": "赴 XX 县数字赋能科技服务团",
        "theme": "乡村振兴（对应 2025 年团中央“乡村振兴 青春建功”主题）",
        "location": "XX 省 XX 县 XX 镇 3 个行政村（A/B/C 村）",
        "practice_time": "2025.07.15-07.21（7 天）", "team_size": "5 人",
        "leader": "张三 / 202212345 / 计算机科学与技术 2022 级 / 大三 / 138XXXXXXXX",
        "advisor": "李教授 / 副教授 / 139XXXXXXXX / 随队 / 软件工程",
    },
    "members": [
        {"name": "张三", "id": "202212345", "major": "计算机科学与技术 2022 级", "role": "队长 / 总协调 / IT 维修主讲", "phone": "138XXXXXXXX"},
        {"name": "李四", "id": "202212346", "major": "计算机科学与技术 2022 级", "role": "IT 维修操作 / 系统重装", "phone": "138XXXXXXXX"},
        {"name": "王五", "id": "202212347", "major": "软件工程 2023 级", "role": "电商培训主讲 / 培训手册设计", "phone": "138XXXXXXXX"},
        {"name": "赵六", "id": "202212348", "major": "通信工程 2022 级", "role": "数字素养培训 / 网络配置", "phone": "138XXXXXXXX"},
        {"name": "钱七", "id": "202212349", "major": "通信工程 2023 级", "role": "反诈宣讲 / 后勤保障 / 答疑群管理", "phone": "138XXXXXXXX"},
    ],
    "theme_background": [
        "团中央主题对应：2025 年团中央发布“乡村振兴 青春建功”暑期社会实践主题，鼓励高校学生深入乡村开展科技服务。本项目紧扣该主题，聚焦“数字赋能”维度，通过 IT 维修、电商培训、数字素养培训提升村民数字能力，助力乡村数字化转型。",
        "服务子方向选定理由：XX 镇 60 岁以上老人 1240 人，智能手机持有率 68% 但能独立完成视频通话的仅 23%；电脑家庭拥有率 31%，故障自修率不足 10%；村电商 8 户年销售额 5 万元以下 6 户。团队 5 人含计算机 3 + 通信 2，专业匹配度高，指导教师软件工程背景可补强电商培训。",
        "实践对象基本情况：XX 镇常住人口 8600 人，60 岁以上老人 1240 人；智能手机持有率 68%，视频通话独立使用率 23%；电脑家庭拥有率 31%，故障自修率不足 10%；村电商 8 户，年销售额 5 万元以下 6 户。数据来源：XX 镇团委 2025 年 5 月摸底调查。",
    ],
    "purpose_significance": [
        "对当地：通过 5 场科技服务覆盖 200 人次，修复电脑 50 台，帮助 20 户村电商开店，发放操作手册 100 份、维护手册 100 份、反诈单页 200 份，建立“XX 县数字赋能”答疑微信群，每月 1 次远程答疑，每季度 1 次现场回访，形成长效帮扶机制，重点解决“老人不会用智能手机”“电脑故障无人修”“村电商不会开店”三类问题。",
        "对学生：团队成员在实践中将课堂所学的计算机硬件、软件、网络知识应用到基层实际场景，提升“用技术服务他人”的能力；通过一对一解决村民实际问题，培养服务意识与同理心；通过深入田间地头与村民同吃同住同劳动，深化对乡村振兴与基层国情的认识；通过建立持续帮扶机制，培养“长期主义”价值观与社会责任感。",
    ],
    "implementation_plan": {
        "form": "本项目实践形式为科技服务类，采用“按场次集中服务 + 一对一上门服务 + 持续远程帮扶”三段式实施流程。按场次集中服务覆盖 3 个行政村 5 场 200 人次，一对一上门服务针对行动不便老人与村电商上门维修，持续远程帮扶建立答疑群 + 月度远程 + 季度现场。子方向为科技服务（IT 维修 + 电商培训 + 数字素养），与调研/支教/宣讲类的核心区别在于“专业输出”而非数据采集、教学方法或一对多传达。",
        "needs_analysis": "通过 XX 镇团委 2025 年 5 月摸底调查 + 实地预调研 2 天访谈村干部 5 人 + 村民 20 人，梳理出 5 项核心需求并按紧迫程度排序：1. 智能手机基础操作（视频通话/扫码支付）— 留守老人 — 高紧迫 — 高匹配；2. 电脑故障维修与系统重装 — 村民/村电商 — 高紧迫 — 高匹配；3. 网络反诈识别 — 全体村民 — 高紧迫 — 中匹配（需法律指导）；4. 电商开店与直播带货基础 — 村电商 — 中紧迫 — 中匹配（需电商指导）；5. 数字素养（政务 App/健康码） — 村民 — 中紧迫 — 高匹配。高优先级 3 项安排主力场次 4 场，中优先级 2 项各安排 1 场。",
        "schedule": [
            {"session": "第 1 场", "date": "7.15 下午", "location": "XX 镇 A 村文化礼堂", "audience": "A 村留守老人 40 人", "content": "智能手机基础操作（视频通话/扫码支付/拍照定位）", "people": "40", "requirement": "计算机/通信，主讲 1 + 助教 2"},
            {"session": "第 2 场", "date": "7.16 上午", "location": "XX 镇 A 村村委会", "audience": "A 村村民 30 人", "content": "电脑故障维修（系统重装/硬件更换/软件优化）", "people": "30", "requirement": "计算机，维修 2 + 记录 1"},
            {"session": "第 3 场", "date": "7.16 下午", "location": "XX 镇 B 村电商服务站", "audience": "B 村村电商 20 人", "content": "电商开店与直播带货基础（平台选择/注册开店/商品上架/直播基础）", "people": "20", "requirement": "计算机/通信，主讲 1 + 助教 2"},
            {"session": "第 4 场", "date": "7.17 上午", "location": "XX 镇 B 村文化广场", "audience": "B 村村民 60 人", "content": "网络反诈识别 + 反诈 App 安装（案例讲解/一对一答疑/App 安装）", "people": "60", "requirement": "计算机/通信 + 法律指导，主讲 2"},
            {"session": "第 5 场", "date": "7.17 下午", "location": "XX 镇 C 村小学", "audience": "C 村村民 50 人", "content": "数字素养培训（政务 App 使用/健康码使用/网络反诈）", "people": "50", "requirement": "计算机/通信，主讲 1 + 助教 3"},
        ],
    },
    "service_plan": {
        "subdirection": "本项目服务子方向为科技服务，覆盖 IT 维修、电商培训、数字素养三大模块。IT 维修包括电脑故障检测、系统重装、硬件更换、软件优化；电商培训包括平台选择、注册开店、商品上架、直播带货基础；数字素养包括智能手机操作、政务 App 使用、健康码使用、网络反诈识别。",
        "needs_table": [
            {"need": "智能手机基础操作", "capability": "计算机/通信，熟悉 Android/iOS", "process": "调研手机型号 → 一对一教学 → 发放手册 → 加答疑群", "solution": "60 位老人 80% 独立完成视频通话 + 操作手册 100 份 + 答疑群 1 个"},
            {"need": "电脑故障维修", "capability": "计算机，硬件/软件能力", "process": "检测故障 → 维修/重装 → 测试 → 填写维修单 → 加答疑群", "solution": "修复 50 台（系统 30+硬件 15+优化 5）+ 维护手册 100 份 + 答疑群 1 个"},
            {"need": "电商开店基础", "capability": "计算机/通信，电商知识（指导教师补充）", "process": "平台选择 → 注册开店 → 商品上架 → 直播基础", "solution": "帮助 20 户开店 + 录制教学视频 5 个 + 答疑群 1 个"},
            {"need": "网络反诈识别", "capability": "计算机/通信 + 法律（指导教师补充）", "process": "案例讲解 → 一对一答疑 → 安装反诈 App → 加群", "solution": "200 人次培训 + 反诈单页 200 份 + 反诈 App 安装 150 部"},
        ],
        "sustain_mechanism": "建立“线上 + 线下”长效服务机制：1. 线上答疑群——建立“XX 县数字赋能”答疑微信群，5 名团队成员轮值（每人每周 1 天），2 小时内响应；2. 月度远程答疑——每月最后一个周末视频连线集中答疑，预计全年 12 次；3. 季度现场回访——每季度 1 次现场回访（2~3 名队员赴地），重点解决疑难问题，预计全年 4 次；4. 学年度接力——下一年度暑期三下乡接力服务，由本届队长对接下届队长，9 月前完成交接，形成“XX 县数字赋能”长效品牌。",
    },
    "safety_plan": [
        "一、出行安全：全员统一购买高铁票，不单独行动；7 月 15 日集体乘 GXXX 次列车赴 XX 县。住宿选择 XX 县政府招待所（已与县团委对接预订），2 人一间，男女分区。每日 18:00 前向指导教师报平安。",
        "二、人身安全：携带常用药品（藿香正气水、创可贴、退烧药、止泻药）；正午 11:00-14:00 不外出，避免中暑；每日早晚测温，体温 ≥ 37.3℃ 立即隔离观察；不携带大量现金，单人不离队。",
        "三、应急联系人：指导教师李教授（139XXXXXXXX）、队长张三（138XXXXXXXX）、学院团委王老师（137XXXXXXXX）。当地对接：XX 县团委刘书记（136XXXXXXXX）。",
        "四、应急流程：一般事件（伤病 < 1000 元）：队长处理 → 指导教师报备 → 学院团委存档。重大事件（伤病 ≥ 1000 元 / 治安事件 / 自然灾害）：队长 → 指导教师 → 学院 → 110/120，同步报告当地团委。所有事件 30 分钟内首报，2 小时内书面报告。",
        "五、保险购买：全员购买中国平安短期意外险（保额 30 万元，保费 15 元/人，保单号 PA20250715001-005），覆盖 7 月 14-22 日。",
        "六、服务类专属安全：维修类带电作业安全（断电操作、防静电、工具绝缘、防误触）；电商培训类隐私保护（不收集村民身份证号、银行卡号）；反诈咨询类信息核实（与当地公安反诈中心对接案例）；服务过程全程录像留档，复杂法律问题转介律所。",
    ],
    "expected_results": [
        "服务场次：5 场（A 村 2 场 + B 村 2 场 + C 村 1 场）",
        "服务人次：200 人次（留守老人 40 + 村民 90 + 村电商 20 + 干部 50）",
        "解决方案：3 套（IT 维修方案 1 套 + 电商开店方案 1 套 + 数字素养培训方案 1 套）",
        "服务记录：签到表 5 份 + 维修记录单 50 份 + 满意度反馈表 200 份",
        "持续帮扶机制：答疑微信群 1 个 + 月度远程答疑 12 次/年 + 季度现场回访 4 次/年 + 学年度接力",
        "新闻稿 5 篇（中青网 1 篇、校团委公众号 2 篇、学院公众号 2 篇）",
        "短视频 3 个（每个 3-5 分钟，发布在 B 站、抖音）",
        "纪录短片 1 个（10 分钟，用于校内汇报）",
    ],
    "budget_items": [
        {"item": "交通费", "amount": "2860", "basis": "高铁往返 130 元 × 5 人 × 2 次 + 县内包车 300 元/天 × 4 天"},
        {"item": "食宿费", "amount": "3920", "basis": "招待所 120 元/间 × 3 间 × 7 天 + 餐费 40 元/人/天 × 5 人 × 7 天"},
        {"item": "物资费", "amount": "1500", "basis": "U 盘 50×20 + 维修工具 400 + 培训手册 100×6 + 反诈单页 200×0.6 + 小礼品 60"},
        {"item": "印刷费", "amount": "480", "basis": "操作手册 100×3 + 维护手册 100×1.5 + 海报 10×8"},
        {"item": "其他", "amount": "240", "basis": "保险 15 元 × 5 人 + 通讯补贴 16 元 × 5 人 + 应急 65 元"},
    ],
    "budget_total": "9000",
    "publicity_plan": [
        "（一）校内媒体：校团委公众号 2 篇、校报 1 篇、学院公众号 2 篇。发布时间节点：7.16 启动报道、7.19 中期进展、7.21 结项报道。",
        "（二）校外媒体：中青网 1 篇、XX 县电视台 1 条、XX 日报 1 篇。中青网稿件由队长审核后投递，附服务照片 5~8 张。",
        "（三）新媒体：B 站 3 个短视频、抖音 3 个短视频、微信视频号 2 个。每个视频 3~5 分钟，内容覆盖服务过程、村民故事、团队风采三个角度。",
    ],
    "include_school_approval": False,
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="三下乡社会实践-科技服务类立项申报书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="示例：\n  python build.py --data data.json --out output.docx\n"
               "  python build.py --demo --out demo.docx\n\n"
               "JSON 字段定义详见 SKILL.md 第十一章。",
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
