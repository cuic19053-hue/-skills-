#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
大学生校级科研立项申请书 docx 生成器（v2.1 案例优化版）

按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 一级标题：黑体三号，居中；二级标题：黑体小三；表格：宋体五号
- 参考文献：宋体五号，按 GB/T 7714 五类格式（J/C/M/D/EB）

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx
    python build.py --data data.json --out output.docx --version enhanced

校级科研立项四大特征：经费 2~5 千、周期 1 年、团队 1~3 人、必须有 GB/T 7714
参考文献（≥15 条含 5 篇英文）。

v2.1 新增字段（基于案例 1 提炼 12 项撰写规范）：
- policy_citations：国家政策引用（≥5 条）
- literature_review：文献综述（≥15 条含 5 篇英文）
- scientific_challenges：科学挑战（2 段结构）
- method_comparison：算法对比表（2~3 方法 × 3 维度）
- tech_roadmap：技术路线图（2 张）
- formulas：数学公式（≥1 个）
- social_benefits：社会效益（5 项简化指标）
- project_schedule：进度安排（4 阶段 12 个月）
- word_version：字数版本（standard 6k / enhanced 8k / peak 10k）

JSON 字段详见 SKILL.md 第十一章。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# 字体与格式常量

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四（正文）
SIZE_WU = Pt(10.5)          # 五号（表格、参考文献）
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# 字数版本配置【v2.1 新增】

WORD_VERSIONS = {
    "standard": {
        "label": "标准版",
        "target_chars": (5000, 6000),
        "include_challenges": False,
        "include_method_comparison": False,
        "include_formulas": False,
        "tech_roadmap_count": 1,
        "social_benefit_count": 3,
        "ref_count": (10, 12),
    },
    "enhanced": {
        "label": "加强版（推荐）",
        "target_chars": (7000, 8000),
        "include_challenges": True,
        "include_method_comparison": True,
        "include_formulas": True,
        "tech_roadmap_count": 2,
        "social_benefit_count": 5,
        "ref_count": (15, 18),
    },
    "peak": {
        "label": "顶峰版",
        "target_chars": (9000, 10000),
        "include_challenges": True,
        "include_method_comparison": True,
        "include_formulas": True,
        "tech_roadmap_count": 3,
        "social_benefit_count": 8,
        "ref_count": (18, 22),
    },
}


# 工具函数

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name,
                         font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置），支持多行（按 \\n 拆分）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    pf.first_line_indent = Pt(0)
    if isinstance(text, str) and "\n" in text:
        lines = text.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p = cell.add_paragraph()
                p.alignment = alignment
                pf = p.paragraph_format
                pf.line_spacing = 1.25
                pf.first_line_indent = Pt(0)
            run = p.add_run(str(line))
            set_run_font(run, font_name=font_name,
                         font_size=font_size, bold=bold)
    else:
        run = p.add_run(str(text))
        set_run_font(run, font_name=font_name,
                     font_size=font_size, bold=bold)


def add_paragraph_with_format(doc, text: str, font_name: str = FONT_SONG,
                              font_size=SIZE_XIAO_SI, bold: bool = False,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT,
                              first_line_indent: bool = True,
                              line_spacing: float = 1.5,
                              space_before: float = 0,
                              space_after: float = 0):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name,
                 font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
        space_before=6, space_after=6)


def add_heading_level3(doc, text: str):
    """三级标题：宋体四号加粗，左对齐"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
        space_before=6, space_after=3)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=indent,
        line_spacing=1.5)


def add_reference_paragraph(doc, idx: int, text: str):
    """参考文献段落：宋体五号，悬挂缩进，单倍行距"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Pt(-18)
    pf.left_indent = Pt(18)
    pf.line_spacing = 1.25
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(f"[{idx}] {text}")
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)
    return p


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j > 0 and len(str(val)) > 10 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False, alignment=align)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def merge_vertical_cells(table, col_idx: int, start_row: int, end_row: int):
    """纵向合并单元格（用于签字栏预留空白）"""
    cells = [table.rows[r].cells[col_idx] for r in range(start_row, end_row + 1)]
    merged = cells[0]
    for c in cells[1:]:
        merged = merged.merge(c)


# GB/T 7714 参考文献格式化（校级科研立项核心）

def format_authors(authors_str: str) -> str:
    """格式化作者：≤3 全列；>3 列前 3 + 等/et al."""
    if not authors_str:
        return ""
    authors = [a.strip() for a in authors_str.split(",") if a.strip()]
    if not authors:
        return authors_str
    is_en = any(c.isalpha() and ord(c) < 128 for c in authors[0])
    if len(authors) <= 3:
        return ", ".join(authors)
    suffix = "et al." if is_en else "等"
    return ", ".join(authors[:3]) + ", " + suffix


def format_journal_ref(idx: int, ref: Dict[str, Any]) -> str:
    """[J] 期刊文章：作者. 题名[J]. 刊名, 年, 卷(期): 起止页码."""
    g = ref.get
    a = format_authors(g("authors", ""))
    s = f"{a}. {g('title', '')}[J]. {g('journal', '')}, {g('year', '')}"
    v, i, p = g("volume", ""), g("issue", ""), g("pages", "")
    if v:
        s += f", {v}{i}"
    if p:
        s += f": {p}"
    return s + "."


def format_conference_ref(idx: int, ref: Dict[str, Any]) -> str:
    """[C] 会议论文：作者. 题名[C]//论文集名. 出版地: 出版者, 年: 页码."""
    g = ref.get
    a = format_authors(g("authors", ""))
    s = (f"{a}. {g('title', '')}[C]//{g('conference', '')}. "
         f"{g('city', '')}: {g('publisher', '')}, {g('year', '')}")
    p = g("pages", "")
    if p:
        s += f": {p}"
    return s + "."


def format_book_ref(idx: int, ref: Dict[str, Any]) -> str:
    """[M] 专著：作者. 书名[M]. 版本. 出版地: 出版者, 年: 页码."""
    g = ref.get
    a = format_authors(g("authors", ""))
    ed = g("edition", "")
    s = f"{a}. {g('title', '')}[M]. "
    if ed:
        s += f"{ed}. "
    s += f"{g('city', '')}: {g('publisher', '')}, {g('year', '')}"
    p = g("pages", "")
    if p:
        s += f": {p}"
    return s + "."


def format_thesis_ref(idx: int, ref: Dict[str, Any]) -> str:
    """[D] 学位论文：作者. 题名[D]. 学位授予地: 学位授予单位, 年."""
    g = ref.get
    a = format_authors(g("authors", ""))
    return (f"{a}. {g('title', '')}[D]. "
            f"{g('city', '')}: {g('school', '')}, {g('year', '')}.")


def format_web_ref(idx: int, ref: Dict[str, Any]) -> str:
    """[EB/OL] 或 [Z/OL] 网络资源"""
    g = ref.get
    a = format_authors(g("authors", ""))
    tag = "Z/OL" if g("is_government", False) else "EB/OL"
    pd, ad, url = g("publish_date", ""), g("access_date", ""), g("url", "")
    s = f"{a}. {g('title', '')}[{tag}]."
    if pd:
        s += f" ({pd})"
    if ad:
        s += f"[{ad}]"
    if url:
        s += f". {url}"
    return s + "."


def format_reference(idx: int, ref: Dict[str, Any]) -> str:
    """根据 ref_type 调用对应格式化函数"""
    rt = (ref.get("ref_type", "journal") or "journal").lower()
    if rt == "journal":
        return format_journal_ref(idx, ref)
    elif rt == "conference":
        return format_conference_ref(idx, ref)
    elif rt == "book":
        return format_book_ref(idx, ref)
    elif rt == "thesis":
        return format_thesis_ref(idx, ref)
    elif rt in ("web", "online", "eb"):
        return format_web_ref(idx, ref)
    return format_journal_ref(idx, ref)


# v2.1 新增章节渲染函数

def add_policy_citations_section(doc, policies: List[Dict[str, Any]]) -> None:
    """渲染国家政策引用章节（v2.1 新增）"""
    if not policies:
        add_body_paragraph(
            doc,
            "（请填写国家政策引用，≥5 条，按时间倒序排列，每条含发文机关 + "
            "文号（可选）+ 标题 + 时间 + 关键表述摘录 + 与课题关联说明。）")
        return
    # 按时间倒序排序
    sorted_policies = sorted(
        policies,
        key=lambda p: p.get("publish_date", ""),
        reverse=True)
    for i, p in enumerate(sorted_policies, 1):
        if not isinstance(p, dict):
            continue
        issuer = p.get("issuer", "")
        doc_no = p.get("doc_no", "")
        title = p.get("title", "")
        pub_date = p.get("publish_date", "")
        excerpt = p.get("key_excerpt", "")
        relevance = p.get("relevance", "")

        # 拼接政策条目段落
        header_parts = []
        if pub_date:
            header_parts.append(pub_date)
        if issuer:
            header_parts.append(issuer)
        if title:
            header_parts.append(f"《{title}》")
        if doc_no:
            header_parts.append(f"（{doc_no}）")
        header = " ".join(header_parts) + "："

        body = ""
        if excerpt:
            body += f'明确提出"{excerpt}"。'
        if relevance:
            body += relevance

        para_text = f"{i}) {header}{body}"
        add_body_paragraph(doc, para_text)


def add_scientific_challenges_section(
        doc, challenges: List[Dict[str, Any]]) -> None:
    """渲染科学挑战章节（v2.1 新增）"""
    if not challenges:
        add_body_paragraph(
            doc,
            "（请填写科学挑战，2 个，每个含标题 + 问题描述段 + 2 子挑战 + "
            "文献支撑。详见 SKILL.md 3.5 节模板。）")
        return
    for i, ch in enumerate(challenges, 1):
        if not isinstance(ch, dict):
            continue
        title = ch.get("title", "")
        description = ch.get("description", "")
        sub_challenges = ch.get("sub_challenges", [])

        # 挑战标题
        add_paragraph_with_format(
            doc, f"科学挑战 {cn_num(i)}：{title}",
            font_name=FONT_HEI, font_size=SIZE_SI, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
            space_before=6, space_after=3)
        # 问题描述
        if description:
            add_body_paragraph(doc, description)
        # 子挑战
        for sub in sub_challenges:
            if not isinstance(sub, dict):
                continue
            sub_name = sub.get("name", "")
            sub_detail = sub.get("detail", "")
            add_paragraph_with_format(
                doc, f"● {sub_name}",
                font_name=FONT_SONG, font_size=SIZE_XIAO_SI, bold=True,
                alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
                space_before=3, space_after=0)
            if sub_detail:
                add_body_paragraph(doc, sub_detail)


def cn_num(n: int) -> str:
    """阿拉伯数字转中文（1~20）"""
    m = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五",
         6: "六", 7: "七", 8: "八", 9: "九", 10: "十",
         11: "十一", 12: "十二", 13: "十三", 14: "十四",
         15: "十五", 16: "十六", 17: "十七", 18: "十八",
         19: "十九", 20: "二十"}
    return m.get(n, str(n))


def add_method_comparison_section(doc, mc: Dict[str, Any]) -> None:
    """渲染算法对比表章节（v2.1 新增）"""
    if not mc or not isinstance(mc, dict):
        add_body_paragraph(
            doc,
            "（请填写算法/方法对比表，2~3 方法 × ≥3 维度 + 选型结论。"
            "详见 SKILL.md 3.7 节模板。）")
        return
    title = mc.get("title", "表 1 主流方法对比表")
    dimensions = mc.get("dimensions", [])
    methods = mc.get("methods", [])
    conclusion = mc.get("conclusion", "")

    if not dimensions or not methods:
        add_body_paragraph(
            doc,
            "（请填写算法/方法对比表，2~3 方法 × ≥3 维度 + 选型结论。）")
        return

    # 表标题
    add_paragraph_with_format(
        doc, title,
        font_name=FONT_HEI, font_size=SIZE_WU, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=6, space_after=3)

    # 构造表格
    headers = ["维度"] + [m.get("name", "") for m in methods]
    rows = []
    for d_idx, dim in enumerate(dimensions):
        row = [dim]
        for m in methods:
            values = m.get("values", [])
            val = values[d_idx] if d_idx < len(values) else ""
            row.append(val)
        rows.append(row)

    n_cols = len(headers)
    col_width = 16.0 / n_cols
    col_widths = [col_width] * n_cols
    add_table_from_data(doc, headers, rows, col_widths=col_widths)

    # 选型结论
    if conclusion:
        add_paragraph_with_format(
            doc, "选型结论：",
            font_name=FONT_HEI, font_size=SIZE_XIAO_SI, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
            space_before=6, space_after=0)
        add_body_paragraph(doc, conclusion)


def add_tech_roadmap_section(doc, roadmaps: List[Dict[str, Any]]) -> None:
    """渲染技术路线图章节（v2.1 新增，2 张图）"""
    if not roadmaps:
        add_body_paragraph(
            doc,
            "（请填写技术路线图，2 张：研究内容关系图 + 技术路线图。"
            "详见 SKILL.md 3.8 节模板。）")
        return
    for rm in roadmaps:
        if not isinstance(rm, dict):
            continue
        fig_no = rm.get("fig_no", "")
        title = rm.get("title", "")
        description = rm.get("description", "")
        nodes = rm.get("nodes", [])

        # 图标题
        add_paragraph_with_format(
            doc, f"{fig_no}  {title}",
            font_name=FONT_HEI, font_size=SIZE_WU, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_before=12, space_after=6)

        # 用表格模拟流程图框
        if nodes:
            n_nodes = len(nodes)
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.rows[0].cells[0]
            cell.text = ""
            for i, node in enumerate(nodes):
                p = cell.add_paragraph() if i > 0 else cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf = p.paragraph_format
                pf.line_spacing = 1.5
                pf.first_line_indent = Pt(0)
                run = p.add_run(str(node))
                set_run_font(run, font_name=FONT_SONG,
                             font_size=SIZE_WU, bold=False)
                if i < n_nodes - 1:
                    arrow_p = cell.add_paragraph()
                    arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    arrow_run = arrow_p.add_run("↓")
                    set_run_font(arrow_run, font_name=FONT_SONG,
                                 font_size=SIZE_WU, bold=True)
            cell.width = Cm(14)
            doc.add_paragraph()

        # 图说明
        if description:
            add_paragraph_with_format(
                doc, f"图示说明：{description}",
                font_name=FONT_SONG, font_size=SIZE_WU, bold=False,
                alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
                space_after=6)


def add_formulas_section(doc, formulas: List[Dict[str, Any]]) -> None:
    """渲染数学公式章节（v2.1 新增）"""
    if not formulas:
        add_body_paragraph(
            doc,
            "（请填写数学公式，≥1 个，含编号 + 变量定义 + 用途说明。"
            "详见 SKILL.md 3.9 节模板。）")
        return
    for f in formulas:
        if not isinstance(f, dict):
            continue
        idx = f.get("idx", 1)
        formula = f.get("formula", "")
        variables = f.get("variables", "")
        purpose = f.get("purpose", "")

        # 公式行（居中、等宽字体感）
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        run = p.add_run(f"    {formula}    式({idx})")
        set_run_font(run, font_name=FONT_TIMES, font_size=SIZE_XIAO_SI,
                     bold=False)

        # 变量定义
        if variables:
            add_paragraph_with_format(
                doc, "其中：",
                font_name=FONT_HEI, font_size=SIZE_XIAO_SI, bold=True,
                alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
                space_before=3, space_after=0)
            add_body_paragraph(doc, variables)

        # 用途说明
        if purpose:
            add_paragraph_with_format(
                doc, f"用途：{purpose}",
                font_name=FONT_SONG, font_size=SIZE_XIAO_SI, bold=False,
                alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
                space_before=3, space_after=6)


def add_social_benefits_section(doc, benefits: List[Dict[str, Any]]) -> None:
    """渲染社会效益量化章节（v2.1 新增，5 项简化指标）"""
    if not benefits:
        add_body_paragraph(
            doc,
            "（请填写社会效益量化，5 项简化指标，每项含指标名 + 传统基准 + "
            "系统预期 + 提升幅度。详见 SKILL.md 3.10 节模板。）")
        return

    # 表标题
    add_paragraph_with_format(
        doc, "表 4 社会效益量化评估表",
        font_name=FONT_HEI, font_size=SIZE_WU, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=6, space_after=3)

    headers = ["评估指标", "传统模式基准值", "系统模式预期值", "效益提升幅度"]
    rows = []
    for b in benefits:
        if not isinstance(b, dict):
            continue
        rows.append([
            b.get("metric", ""),
            b.get("baseline", ""),
            b.get("expected", ""),
            b.get("improvement", ""),
        ])
    add_table_from_data(doc, headers, rows,
                        col_widths=[4.5, 4.0, 4.0, 4.0])

    add_body_paragraph(
        doc,
        "注：以上数据基于预实验与文献调研，正式实验后将以真实数据更新。"
        "效益提升幅度按『百分比/倍数』双重表达规范填写。")


def add_project_schedule_section(doc, schedule: List[Dict[str, Any]]) -> None:
    """渲染 4 阶段进度安排章节（v2.1 新增）"""
    if not schedule:
        add_body_paragraph(
            doc,
            "（请填写进度安排，4 阶段，每阶段含阶段名 + 起止月份 + 主要任务 + "
            "阶段成果。详见 SKILL.md 3.11 节模板。）")
        return

    headers = ["阶段", "时间", "主要工作", "阶段成果"]
    rows = []
    for s in schedule:
        if not isinstance(s, dict):
            continue
        rows.append([
            s.get("phase", ""),
            s.get("time", ""),
            s.get("work", ""),
            s.get("output", ""),
        ])
    add_table_from_data(doc, headers, rows,
                        col_widths=[2.8, 3.0, 6.0, 4.5])


def add_literature_review_summary(doc, lit_list: List[Dict[str, Any]]) -> None:
    """渲染文献综述摘要表（v2.1 新增，作为研究现状的补充）"""
    if not lit_list:
        return
    add_paragraph_with_format(
        doc, "附表：核心文献综述摘要",
        font_name=FONT_HEI, font_size=SIZE_WU, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=6, space_after=3)
    headers = ["编号", "作者+年份", "主要观点", "局限性"]
    rows = []
    for lit in lit_list:
        if not isinstance(lit, dict):
            continue
        rows.append([
            f"[{lit.get('idx', '')}]",
            f"{lit.get('authors', '')}（{lit.get('year', '')}）",
            lit.get("view", ""),
            lit.get("limit", ""),
        ])
    add_table_from_data(doc, headers, rows,
                        col_widths=[1.5, 3.5, 6.0, 5.0])


# 主构建类

class ApplicationDocBuilder:
    """大学生校级科研立项申请书 docx 构建器（v2.1）"""

    def __init__(self):
        self.doc = Document()
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []
        self.word_version: str = "enhanced"
        self.version_config: Dict[str, Any] = WORD_VERSIONS["enhanced"]
        setup_page(self.doc)
        add_page_number(self.doc)

    # 数据访问

    def _get(self, *keys, default: Any = None) -> Any:
        """嵌套取值，支持 _get('basic_info', 'duration', default='')"""
        cur = self.data
        for k in keys:
            if isinstance(cur, dict):
                cur = cur.get(k, default)
            else:
                return default
        return cur if cur is not None else default

    # 标题辅助

    def add_h1(self, text: str):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text: str):
        return add_heading_level2(self.doc, text)

    def add_h3(self, text: str):
        return add_heading_level3(self.doc, text)

    def add_para(self, text: str, indent: bool = True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_table(self, headers: List[str], rows: List[List[str]],
                  col_widths: Optional[List[float]] = None):
        return add_table_from_data(self.doc, headers, rows,
                                   col_widths=col_widths)

    def add_page_break(self):
        add_page_break(self.doc)

    # 封面

    def _add_cover(self):
        """封面：校名+标题+4 项信息"""
        college = self._get("college", default="XX 大学")
        title = f"{college}大学生校级科研立项申请书"
        add_paragraph_with_format(
            self.doc, title, font_name=FONT_HEI, font_size=SIZE_ER,
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False, space_before=12, space_after=12)
        # 字数版本标签【v2.1 新增】
        ver_label = self.version_config.get("label", "加强版")
        add_paragraph_with_format(
            self.doc, f"（{ver_label} · v2.1）",
            font_name=FONT_SONG, font_size=SIZE_XIAO_SI, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_after=12)
        for _ in range(2):
            self.doc.add_paragraph()
        info_items = [
            ("课题名称", self._get("project_name")),
            ("课题负责人", self._get("leader_name")),
            ("所在学院", self._get("college")),
            ("申报日期", self._get("apply_date")),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI, font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True
        self.add_page_break()

    # 一、课题负责人及团队成员信息

    def _add_leader_members(self):
        """一、课题负责人及团队成员信息（负责人表 + 成员表）"""
        self.add_h1("一、课题负责人及团队成员信息")
        self.add_h2("（一）课题负责人信息")
        leader_rows = [
            ["姓名", self._get("leader_name"), "学号", self._get("leader_id")],
            ["性别", self._get("leader_gender", default="男"), "专业年级", self._get("leader_major")],
            ["学院", self._get("college"), "联系电话", self._get("leader_phone")],
            ["邮箱", self._get("leader_email"), "指导教师", self._get("advisor_name")],
        ]
        self.add_table(["字段", "内容", "字段", "内容"], leader_rows,
                       col_widths=[2.5, 5.5, 2.5, 5.5])
        self.add_h2("（二）团队成员信息")
        members = self._get("team_members", default=[])
        if isinstance(members, list) and members:
            rows = []
            for m in members:
                if not isinstance(m, dict):
                    continue
                rows.append([m.get("name", ""), m.get("id", ""),
                             m.get("major", ""), m.get("role", ""),
                             m.get("phone", "")])
            self.add_table(
                ["姓名", "学号", "专业年级", "分工", "联系方式"], rows,
                col_widths=[2.0, 2.5, 3.5, 5.0, 3.0])
        else:
            self.add_para("（团队成员 1~3 人，含分工。校级科研立项团队规模 1~3 人，请填写姓名/学号/专业年级/分工/联系方式。）")

    # 二/三/四、课题名称/起止时间/研究类型

    def _add_basic_info(self):
        """二、课题名称；三、起止时间；四、研究类型"""
        self.add_h1("二、课题名称")
        project_name = self._get("project_name", default="")
        add_paragraph_with_format(
            self.doc, project_name if project_name else
            "（不超过 30 字，突出『做什么+为谁做』，禁用『基于...的...』堆砌）",
            font_name=FONT_HEI, font_size=SIZE_SAN, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_after=6)
        self.add_h1("三、研究起止时间")
        duration = self._get("basic_info", "duration", default="")
        if not duration:
            duration = self._get("duration", default="2025.04-2026.03（共 12 个月）")
        add_paragraph_with_format(
            self.doc, duration, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_after=6)
        self.add_h1("四、研究类型")
        research_type = (self._get("basic_info", "research_type", default="")
                         or self._get("research_type", default="应用研究"))
        type_options = ["基础研究", "应用研究", "开发研究"]
        type_str = "    ".join(
            f"{'☑' if t == research_type else '☐'} {t}" for t in type_options)
        add_paragraph_with_format(
            self.doc, type_str, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)

    # 五、课题研究的背景与意义

    def _add_background(self):
        """五、课题研究的背景与意义（500~800 字，4 段）"""
        self.add_h1("五、课题研究的背景与意义")
        background = self._get("background", default=[])
        if isinstance(background, str):
            background = [background]
        sub_titles = ["（一）时代背景", "（二）现实痛点", "（三）研究缺口", "（四）课题意义"]
        placeholders = [
            "（请填写时代背景，100~150 字，2~3 句政策/行业/学术趋势，必须含权威数据来源。）",
            "（请填写现实痛点，200~300 字，2~3 个真实场景，必须可量化。）",
            "（请填写研究缺口，100~150 字，简要评述已有研究不足，详细评述见七、研究现状。）",
            "（请填写课题意义，100~200 字，理论/实践/社会三角度，至少两个。）",
        ]
        if not background:
            for t, p in zip(sub_titles, placeholders):
                self.add_h2(t)
                self.add_para(p)
        else:
            for i, para in enumerate(background):
                if i < len(sub_titles):
                    self.add_h2(sub_titles[i])
                self.add_para(para)

    # 六、国家政策引用【v2.1 新增】

    def _add_policy_citations(self):
        """六、国家政策引用（≥5 条，按时间倒序）"""
        self.add_h1("六、国家政策引用")
        self.add_para(
            "本课题紧密对接国家相关政策文件，按时间倒序梳理如下：")
        policies = self._get("policy_citations", default=[])
        if isinstance(policies, list):
            add_policy_citations_section(self.doc, policies)
        else:
            add_policy_citations_section(self.doc, [])

    # 七、国内外研究现状及发展动态

    def _add_research_status(self):
        """七、国内外研究现状及发展动态（800~1200 字，3 段评述式）"""
        self.add_h1("七、国内外研究现状及发展动态")
        status = self._get("research_status", default=[])
        if isinstance(status, str):
            status = [status]
        sub_titles = ["（一）国内研究现状", "（二）国外研究现状", "（三）综合评述与本课题差异"]
        placeholders = [
            "（请填写国内研究现状，300~400 字，引用 5~8 篇国内文献，每篇含『作者+年份+观点+局限』并标注 [1][2][3] 等编号。禁止罗列式。）",
            "（请填写国外研究现状，300~400 字，引用 5~8 篇国外文献，同上结构，标注 [N] 等编号。）",
            "（请填写综合评述，200~400 字，归纳国内外研究共性问题 2~3 个，提出本课题差异点（与共性问题一一对应），引出创新点。）",
        ]
        if not status:
            for t, p in zip(sub_titles, placeholders):
                self.add_h2(t)
                self.add_para(p)
        else:
            for i, para in enumerate(status):
                if i < len(sub_titles):
                    self.add_h2(sub_titles[i])
                self.add_para(para)
        # 文献综述摘要表【v2.1 新增】
        lit_review = self._get("literature_review", default=[])
        if isinstance(lit_review, list) and lit_review:
            self.add_h2("（四）核心文献综述摘要表")
            add_literature_review_summary(self.doc, lit_review)

    # 八、科学挑战【v2.1 新增】

    def _add_scientific_challenges(self):
        """八、科学挑战（2 段结构，每段含 2 子挑战）"""
        if not self.version_config.get("include_challenges", True):
            return  # standard 版本不含科学挑战
        self.add_h1("八、科学挑战")
        challenges = self._get("scientific_challenges", default=[])
        if isinstance(challenges, list):
            add_scientific_challenges_section(self.doc, challenges)
        else:
            add_scientific_challenges_section(self.doc, [])

    # 九、研究目标、研究内容、拟解决的关键问题

    def _add_research_content(self):
        """九、研究目标、研究内容、拟解决的关键问题（3 子节）"""
        self.add_h1("九、研究目标、研究内容、拟解决的关键问题")
        self.add_h2("（一）研究目标")
        goal = self._get("research_goal", default="")
        if goal:
            self.add_para(goal)
        else:
            self.add_para("（请填写研究目标，200~300 字，1 个总目标 + 3~4 个阶段目标，全部可量化。）")
        self.add_h2("（二）研究内容")
        contents = self._get("research_content", default=[])
        if isinstance(contents, str):
            contents = [contents]
        if contents:
            for i, c in enumerate(contents, 1):
                self.add_para(f"{i}. {c}")
        else:
            self.add_para("（请填写研究内容，3~5 个子任务，每个 100~150 字，结构：任务名+做什么+方法+预期产出。）")
        self.add_h2("（三）拟解决的关键问题")
        problems = self._get("key_problems", default=[])
        if isinstance(problems, str):
            problems = [problems]
        if problems:
            for i, q in enumerate(problems, 1):
                self.add_para(f"{i}. {q}")
        else:
            self.add_para("（请填写关键问题，2~3 个，每个一句话讲清技术难点。）")

    # 十、算法/方法对比【v2.1 新增】

    def _add_method_comparison(self):
        """十、算法/方法对比（2~3 方法 × 3 维度 + 选型结论）"""
        if not self.version_config.get("include_method_comparison", True):
            return  # standard 版本不含
        self.add_h1("十、算法/方法对比")
        self.add_para(
            "本课题在方案设计前，对国内外主流相关方法进行了系统对比，"
            "结果如下：")
        mc = self._get("method_comparison", default={})
        if isinstance(mc, dict):
            add_method_comparison_section(self.doc, mc)
        else:
            add_method_comparison_section(self.doc, {})

    # 十一、研究方案及技术路线

    def _add_research_scheme(self):
        """十一、研究方案及技术路线（含方法、步骤、流程图）"""
        self.add_h1("十一、研究方案及技术路线")
        self.add_h2("（一）总体技术路线")
        route = self._get("tech_route", default="")
        if route:
            self.add_para(route)
        else:
            self.add_para("（请填写总体技术路线，150~250 字 + 2 张技术路线图（研究内容关系图 + 技术路线图），4 阶段流程，每阶段标注交付物。）")
        # 技术路线图【v2.1 新增】
        roadmaps = self._get("tech_roadmap", default=[])
        if isinstance(roadmaps, list) and roadmaps:
            target_count = self.version_config.get("tech_roadmap_count", 2)
            add_tech_roadmap_section(self.doc, roadmaps[:target_count])
        else:
            # 兼容旧版流程图字段
            flowchart = self._get("tech_flowchart_image", default="")
            if flowchart and os.path.exists(flowchart):
                try:
                    p = self.doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(flowchart, width=Cm(15))
                    add_paragraph_with_format(
                        self.doc, "图 1 课题技术路线图",
                        font_name=FONT_HEI, font_size=SIZE_WU, bold=False,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
                except Exception as e:
                    self.add_para(f"（流程图插入失败：{e}）")
        self.add_h2("（二）研究方法")
        methods = self._get("methods", default=[])
        if isinstance(methods, str):
            methods = [methods]
        if methods:
            for i, m in enumerate(methods, 1):
                self.add_para(f"{i}. {m}")
        else:
            self.add_para("（请填写研究方法，3~5 个，每个 50~80 字说明用途。）")
        self.add_h2("（三）数据来源与实验条件")
        data_src = self._get("data_source", default="")
        self.add_para(data_src if data_src else
                      "（请填写数据来源、规模、实验设备型号、软件工具及版本。）")

    # 十二、数学公式与算法描述【v2.1 新增】

    def _add_formulas(self):
        """十二、数学公式与算法描述（≥1 公式）"""
        if not self.version_config.get("include_formulas", True):
            return  # standard 版本不含
        self.add_h1("十二、数学公式与算法描述")
        self.add_para(
            "本课题核心方法采用如下数学模型描述：")
        formulas = self._get("formulas", default=[])
        if isinstance(formulas, list):
            add_formulas_section(self.doc, formulas)
        else:
            add_formulas_section(self.doc, [])

    # 十三、创新之处

    def _add_innovation(self):
        """十三、创新之处（300~500 字，至少 1 个，对比式写法）"""
        self.add_h1("十三、创新之处")
        innovations = self._get("innovations", default=[])
        if isinstance(innovations, str):
            innovations = [innovations]
        if innovations:
            for i, inv in enumerate(innovations, 1):
                self.add_para(f"创新点 {i}：{inv}")
        else:
            self.add_para("（请填写创新点，至少 1 个，建议 2 个，每个 150~250 字。结构：[类型]。传统方法 [描述]，本课题 [方法]，[量化优势]。禁用『首次』『先进』『智能』等无量化支撑词。）")

    # 十四、社会效益量化【v2.1 新增】

    def _add_social_benefits(self):
        """十四、社会效益量化（5 项简化指标）"""
        self.add_h1("十四、社会效益量化")
        self.add_para(
            "本课题预期产生的社会效益如下表所示（含传统基准值、系统预期值、"
            "效益提升幅度 4 列对比）：")
        benefits = self._get("social_benefits", default=[])
        if isinstance(benefits, list):
            target_count = self.version_config.get("social_benefit_count", 5)
            add_social_benefits_section(self.doc, benefits[:target_count])
        else:
            add_social_benefits_section(self.doc, [])

    # 十五、研究基础（前期成果、设备条件）

    def _add_research_basis(self):
        """十五、研究基础（3 子节：团队/指导教师/实验条件）"""
        self.add_h1("十五、研究基础")
        self.add_h2("（一）团队基础")
        team = self._get("team_foundation", default="")
        self.add_para(team if team else
                      "（请填写团队基础：成员相关课程、已有项目经验、技能匹配度。）")
        self.add_h2("（二）指导教师基础")
        advisor = self._get("advisor_foundation", default="")
        self.add_para(advisor if advisor else
                      "（请填写指导教师基础：主持项目、发表论文、研究方向匹配度。）")
        self.add_h2("（三）实验条件")
        lab = self._get("lab_condition", default="")
        self.add_para(lab if lab else
                      "（请填写实验条件：实验室设备、软件平台、合作单位支持。）")

    # 十六、预期研究成果

    def _add_expected_results(self):
        """十六、预期研究成果（至少 1 篇论文，可量化）"""
        self.add_h1("十六、预期研究成果")
        outcomes = self._get("expected_outcomes", default=[])
        if isinstance(outcomes, str):
            outcomes = [outcomes]
        if outcomes:
            for o in outcomes:
                self.add_para(f"• {o}", indent=False)
        else:
            self.add_para("（请填写预期成果，每项含数量+级别+平台。校级至少 1 篇论文，如：中文核心论文 1 篇（拟投《XX》）、调研报告 1 份（约 1.5 万字）。）", indent=False)

    # 十七、研究进度安排（4 阶段表格）

    def _add_schedule(self):
        """十七、研究进度安排（4 列表格，按月划分，v2.1 优先用 project_schedule）"""
        self.add_h1("十七、研究进度安排")
        # v2.1 优先用 project_schedule（4 阶段）
        schedule = self._get("project_schedule", default=[])
        if not schedule:
            # 兼容旧版 schedule 字段
            schedule = self._get("schedule", default=[])
        if isinstance(schedule, list) and schedule:
            add_project_schedule_section(self.doc, schedule)
        else:
            self.add_para("（请填写进度安排，4 阶段，按月划分，每阶段标注交付物。建议 4 阶段共 12 个月，留 1 月弹性时间应对突发情况。）")

    # 十八、经费预算

    def _add_budget(self):
        """十八、经费预算（6 类科目：资料/调研/材料/会议/印刷/其他）"""
        self.add_h1("十八、经费预算")
        items = self._get("budget_items", default=[])
        if items:
            rows = []
            total = 0
            for b in items:
                if not isinstance(b, dict):
                    continue
                try:
                    amount_num = int(b.get("amount", 0))
                except ValueError:
                    amount_num = 0
                total += amount_num
                rows.append([b.get("item", ""), f"{amount_num} 元",
                             b.get("basis", "")])
            rows.append(["合计", f"{total} 元", ""])
            self.add_table(["预算科目", "金额", "计算依据"], rows,
                           col_widths=[3.5, 3.0, 9.5])
        else:
            self.add_para("（请填写经费预算，6 类标准科目：资料费/调研费/材料费/会议费/印刷费/其他。金额非整数，附计算依据。校级经费 2~5 千元。）")

    # 十九、参考文献（GB/T 7714）

    def _add_references(self):
        """十九、参考文献（GB/T 7714 五类格式：J/C/M/D/EB）"""
        self.add_h1("十九、参考文献")
        refs = self._get("references", default=[])
        if not isinstance(refs, list):
            refs = []
        if not refs:
            self.add_para("（请填写参考文献，≥15 条，覆盖期刊/会议/专著/学位论文/"
                          "网络资源至少 3 类（建议 5 类），含 5 篇英文。格式按 "
                          "GB/T 7714-2015，详见 SKILL.md 第七章。）", indent=False)
            return
        type_counter: Dict[str, int] = {}
        en_count = 0
        for i, ref in enumerate(refs, 1):
            if not isinstance(ref, dict):
                continue
            try:
                line = format_reference(i, ref)
            except Exception as e:
                line = f"（格式化失败：{e}）"
            add_reference_paragraph(self.doc, i, line)
            rt = ref.get("ref_type", "journal").lower()
            type_counter[rt] = type_counter.get(rt, 0) + 1
            # 检测英文文献
            authors = ref.get("authors", "")
            if authors and any(c.isalpha() and ord(c) < 128 for c in authors):
                en_count += 1
        # v2.1 升级校验
        min_ref = self.version_config.get("ref_count", (15, 18))[0]
        if len(refs) < min_ref:
            self.warnings.append(
                f"参考文献仅 {len(refs)} 条，{self.word_version} 版本要求 ≥{min_ref} 条")
        if len(type_counter) < 3:
            self.warnings.append(
                f"参考文献仅覆盖 {len(type_counter)} 类，要求至少 3 类")
        if en_count < 5:
            self.warnings.append(
                f"英文文献仅 {en_count} 篇，要求 ≥5 篇")

    # 二十/二一、指导教师推荐意见 / 学院评审意见

    def _add_review_section(self):
        """二十/二一、指导教师推荐意见 + 学院评审意见（双栏签字）"""
        self.add_h1("二十、指导教师推荐意见")
        for _ in range(6):
            self.doc.add_paragraph()
        self.add_para("指导教师签字：____________________    "
                      "日期：______年____月____日", indent=False)
        self.add_h1("二一、学院评审意见")
        for _ in range(6):
            self.doc.add_paragraph()
        self.add_para("学院盖章：____________________    "
                      "日期：______年____月____日", indent=False)

    # 主构建方法

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 22 栏目，生成 docx"""
        try:
            self.data = data if isinstance(data, dict) else {}
            # 解析字数版本【v2.1 新增】
            self.word_version = self._get("word_version", default="enhanced")
            if self.word_version not in WORD_VERSIONS:
                self.warnings.append(
                    f"未知字数版本 {self.word_version}，默认使用 enhanced")
                self.word_version = "enhanced"
            self.version_config = WORD_VERSIONS[self.word_version]
            self._validate_data()
            # 22 栏目编排
            self._add_cover()
            self._add_leader_members()
            self._add_basic_info()
            self._add_background()
            self._add_policy_citations()        # v2.1 新增
            self._add_research_status()
            self._add_scientific_challenges()   # v2.1 新增
            self._add_research_content()
            self._add_method_comparison()       # v2.1 新增
            self._add_research_scheme()
            self._add_formulas()                # v2.1 新增
            self._add_innovation()
            self._add_social_benefits()         # v2.1 新增
            self._add_research_basis()
            self._add_expected_results()
            self._add_schedule()
            self._add_budget()
            self._add_references()
            self._add_review_section()
            if self.warnings:
                print("⚠️ 数据校验警告：", file=sys.stderr)
                for w in self.warnings:
                    print(f"  - {w}", file=sys.stderr)
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 校级科研立项申请书已生成：{output_path}")
        return str(output_path)

    # 数据校验

    def _validate_data(self) -> List[str]:
        """校验数据完整性，记录警告但不阻断生成"""
        p0_fields = [("project_name", "课题名称"),
                     ("leader_name", "负责人姓名"),
                     ("college", "所在学院"),
                     ("advisor_name", "指导教师姓名")]
        for key, name in p0_fields:
            if not self._get(key):
                self.warnings.append(f"缺少 {name}（{key}）")
        basic = self._get("basic_info", default={})
        if not isinstance(basic, dict):
            basic = {}
        if not basic.get("duration") and not self._get("duration"):
            self.warnings.append("缺少 起止时间（duration）")
        if not basic.get("research_type") and not self._get("research_type"):
            self.warnings.append("缺少 研究类型（research_type），默认『应用研究』")
        if not self._get("background"):
            self.warnings.append("缺少 背景与意义（background），将使用占位文本")
        if not self._get("research_status"):
            self.warnings.append("缺少 国内外研究现状（research_status），将使用占位文本")
        if not self._get("innovations"):
            self.warnings.append("缺少 创新之处（innovations），将使用占位文本")

        # v2.1 新增字段校验
        policies = self._get("policy_citations", default=[])
        if isinstance(policies, list):
            if len(policies) < 5:
                self.warnings.append(
                    f"国家政策引用仅 {len(policies)} 条，要求 ≥5 条")
        else:
            self.warnings.append("缺少 国家政策引用（policy_citations），将使用占位文本")

        lit_review = self._get("literature_review", default=[])
        if isinstance(lit_review, list):
            if len(lit_review) < 15:
                self.warnings.append(
                    f"文献综述仅 {len(lit_review)} 条，要求 ≥15 条")
            en_count = sum(
                1 for lit in lit_review
                if isinstance(lit, dict) and lit.get("lang", "zh") == "en"
            )
            if en_count < 5:
                self.warnings.append(
                    f"文献综述英文仅 {en_count} 篇，要求 ≥5 篇")
        else:
            self.warnings.append("缺少 文献综述（literature_review），将使用占位文本")

        if self.version_config.get("include_challenges", True):
            challenges = self._get("scientific_challenges", default=[])
            if isinstance(challenges, list):
                if len(challenges) < 2:
                    self.warnings.append(
                        f"科学挑战仅 {len(challenges)} 个，要求 2 个")
            else:
                self.warnings.append("缺少 科学挑战（scientific_challenges），将使用占位文本")

        if self.version_config.get("include_method_comparison", True):
            mc = self._get("method_comparison", default={})
            if not mc or not isinstance(mc, dict):
                self.warnings.append("缺少 算法对比表（method_comparison），将使用占位文本")
            else:
                dims = mc.get("dimensions", [])
                methods = mc.get("methods", [])
                if len(dims) < 3:
                    self.warnings.append(
                        f"算法对比维度仅 {len(dims)} 个，要求 ≥3 个")
                if len(methods) < 2:
                    self.warnings.append(
                        f"算法对比方法仅 {len(methods)} 个，要求 ≥2 个")
                if not mc.get("conclusion"):
                    self.warnings.append("算法对比表缺少选型结论（conclusion）")

        roadmaps = self._get("tech_roadmap", default=[])
        if isinstance(roadmaps, list):
            target_count = self.version_config.get("tech_roadmap_count", 2)
            if len(roadmaps) < target_count:
                self.warnings.append(
                    f"技术路线图仅 {len(roadmaps)} 张，{self.word_version} 版本要求 ≥{target_count} 张")
        else:
            self.warnings.append("缺少 技术路线图（tech_roadmap），将使用占位文本")

        if self.version_config.get("include_formulas", True):
            formulas = self._get("formulas", default=[])
            if isinstance(formulas, list):
                if len(formulas) < 1:
                    self.warnings.append("数学公式 0 个，要求 ≥1 个")
            else:
                self.warnings.append("缺少 数学公式（formulas），将使用占位文本")

        benefits = self._get("social_benefits", default=[])
        if isinstance(benefits, list):
            target_count = self.version_config.get("social_benefit_count", 5)
            if len(benefits) < target_count:
                self.warnings.append(
                    f"社会效益仅 {len(benefits)} 项，{self.word_version} 版本要求 ≥{target_count} 项")
        else:
            self.warnings.append("缺少 社会效益（social_benefits），将使用占位文本")

        schedule = self._get("project_schedule", default=[])
        if not schedule:
            schedule = self._get("schedule", default=[])
        if isinstance(schedule, list):
            if len(schedule) != 4 and len(schedule) > 0:
                self.warnings.append(
                    f"进度安排 {len(schedule)} 阶段，建议 4 阶段")
        else:
            self.warnings.append("缺少 进度安排（project_schedule），将使用占位文本")

        # 经费预算校验
        items = self._get("budget_items", default=[])
        if items:
            total = 0
            for b in items:
                if isinstance(b, dict):
                    try:
                        total += int(b.get("amount", 0))
                    except (ValueError, TypeError):
                        pass
            budget_total_str = str(self._get("budget_total", default="")).strip()
            try:
                budget_total_num = int(budget_total_str)
            except ValueError:
                budget_total_num = -1
            if budget_total_num >= 0 and total != budget_total_num:
                self.warnings.append(f"预算合计 {total} 元 与申请经费 {budget_total_num} 元不一致")
            if total > 5000:
                self.warnings.append(f"预算合计 {total} 元 超校级上限 5000 元，需说明必要性")
            if total < 2000:
                self.warnings.append(f"预算合计 {total} 元 低于校级下限 2000 元")
        return self.warnings


# 默认示例数据（v2.1 完整版，含全部新增字段）

DEFAULT_DATA = {
    "project_name": "大学生课堂手机使用行为干预研究",
    "leader_name": "张明",
    "leader_id": "2022123456",
    "leader_gender": "男",
    "leader_major": "教育技术学 2022 级",
    "leader_phone": "138XXXXXXXX",
    "leader_email": "zhangming@xxx.edu.cn",
    "advisor_name": "李华教授",
    "advisor_title": "教授",
    "advisor_research": "教育技术学、学习行为分析",
    "advisor_phone": "139XXXXXXXX",
    "college": "教育学院",
    "apply_date": "2025 年 3 月 15 日",
    "word_version": "enhanced",

    "basic_info": {
        "research_type": "应用研究",
        "discipline": "0401 教育学",
        "duration": "2025 年 4 月 — 2026 年 3 月（共 12 个月）",
        "budget": "3500",
        "project_source": "自主选题",
    },

    "team_members": [
        {"name": "李四", "id": "2022123457", "major": "教育技术学 2022 级",
         "role": "问卷设计+数据采集", "phone": "139XXXXXXXX"},
        {"name": "王五", "id": "2022123458", "major": "心理学 2022 级",
         "role": "数据分析+论文撰写", "phone": "137XXXXXXXX"},
    ],

    "background": [
        "时代背景：2024 年教育部《关于进一步加强大学生课堂管理的指导意见》明确提出『严禁课堂手机使用』。但据中国教育报 2024 年调查，全国高校大学生课堂手机使用率达 78.3%，其中 65% 用于与课堂无关活动，严重影响教学秩序。",
        "现实痛点：调研本校 5 个学院 12 个班级发现，课堂手机使用率 81.2%，超过六成学生承认『忍不住看手机』。某专业课教师反映，课堂抬头率从 5 年前的 85% 降至 52%。学生方面，72% 表示『知道不该用但控制不住』，仅 18% 接受过系统行为干预。",
        "研究缺口：已有研究多停留在『现象描述』层面，缺乏可操作的干预方案。少数干预研究样本量小（<100）、周期短（<1 月），缺乏长期效果评估。",
        "课题意义：理论上探索行为干预理论在大学生课堂管理中的适用边界；实践上形成可复用的干预方案，预期能将课堂手机使用率降低 30 个百分点以上；社会上助力学风建设。",
    ],

    # v2.1 新增：国家政策引用（5 条，按时间倒序）
    "policy_citations": [
        {"issuer": "教育部", "doc_no": "教高〔2024〕3 号",
         "title": "关于进一步加强大学生课堂管理的指导意见",
         "publish_date": "2024-03-15",
         "key_excerpt": "严禁课堂手机使用，构建良好学习生态",
         "relevance": "直接对应本课题研究场景，本课题干预方案正是落实该意见的具体技术路径。"},
        {"issuer": "国务院", "doc_no": "",
         "title": "关于全面深化新时代教师队伍建设改革的意见",
         "publish_date": "2023-12-20",
         "key_excerpt": "提升教师课堂管理能力，营造良好教风学风",
         "relevance": "本课题为教师提供科学的干预工具，助力教师课堂管理能力提升。"},
        {"issuer": "教育部", "doc_no": "",
         "title": "高等学校课程思政建设指导纲要",
         "publish_date": "2022-08-10",
         "key_excerpt": "将思想政治教育贯穿教育教学全过程",
         "relevance": "本课题干预方案融入思政元素，强化学生自律意识。"},
        {"issuer": "中共中央国务院", "doc_no": "",
         "title": "关于深化新时代教育评价改革总体方案",
         "publish_date": "2021-03-15",
         "key_excerpt": "建立学生学习过程性评价体系",
         "relevance": "本课题提供的过程性数据可作为评价依据。"},
        {"issuer": "教育部", "doc_no": "",
         "title": "关于深化本科教育教学改革全面提高人才培养质量的意见",
         "publish_date": "2020-10-08",
         "key_excerpt": "严格教育教学管理，严格课堂教学纪律",
         "relevance": "本课题从行为科学角度为课堂教学纪律管理提供支撑。"},
    ],

    "research_status": [
        "国内研究现状：张华等（2022）对 5 所高校 1200 名大学生调查发现课堂手机使用率达 75.8% [1]，但仅描述现象未提出干预方案。李明等（2023）扩展样本至 8 所高校 2000 人，使用率达 78.3% [2]，仍停留在现象描述层面。王强（2023）设计了『积分奖励』干预机制，样本 80 人、周期 2 周，短期效果显著（使用率降 22%）但未做长期跟踪 [3]。刘芳等（2024）尝试『课堂信号屏蔽』硬件方案，因法律与伦理争议未能推广 [4]。陈伟等（2024）引入『自律契约』机制，样本 150 人、周期 4 周，短期降 18% 但缺乏跟踪 [5]。赵敏（2023）尝试将西方行为干预理论本土化，但仅做了理论分析未做实证 [6]。周平等（2024）针对辅导员制度设计了协同干预方案，仅在小样本（60 人）试点 [7]。",
        "国外研究现状：Smith et al. (2020) 对美国 3 所大学 800 名学生研究发现，『自律契约』干预 6 周后课堂手机使用率从 82% 降至 61% [8]。Brown et al. (2022) 扩展至 5 所大学 1500 人，6 月跟踪保持率 65% [9]。Kumar & Lee (2022) 比较了 4 种干预方式（积分/契约/屏蔽/提醒），发现『契约+提醒』组合效果最佳 [10]。Davis et al. (2023) 比较了 6 种干预方式，发现『多维度组合』效果优于单维 [11]。Brown (2023) 引入行为经济学『承诺装置』理论，使干预效果在 6 个月后仍保持 70% [12]。Anderson & Wilson (2024) 提出干预效果衰减的双指数模型 [13]。Lee et al. (2023) 跨文化比较发现集体主义文化下『协同干预』效果优于个人主义文化 [14]。Park et al. (2024) 在韩国高校试点辅导员协同干预，效果显著 [15]。",
        "综合评述：已有研究存在三个共性问题：① 样本量普遍偏小（多数 <200 人），外推性不足；② 干预周期短（<1 月），缺乏长期效果评估；③ 缺乏针对中国高校学情（如班级授课制、辅导员制度）的本土化设计。本课题的关键差异：（1）样本量 600+，覆盖 3 所高校；（2）周期 1 年含 6 个月跟踪；（3）融合『契约+提醒+辅导员协同』三维干预，本土化适配。",
    ],

    # v2.1 新增：文献综述（15 条含 7 篇英文）
    "literature_review": [
        {"idx": 1, "authors": "张华等", "year": "2022",
         "view": "课堂手机使用率达 75.8%",
         "limit": "仅描述现象未提出干预方案",
         "lang": "zh"},
        {"idx": 2, "authors": "李明等", "year": "2023",
         "view": "8 校 2000 人使用率达 78.3%",
         "limit": "仍停留在现象描述层面",
         "lang": "zh"},
        {"idx": 3, "authors": "王强", "year": "2023",
         "view": "积分奖励机制短期降 22%",
         "limit": "样本 80 人、周期 2 周，未长期跟踪",
         "lang": "zh"},
        {"idx": 4, "authors": "刘芳等", "year": "2024",
         "view": "课堂信号屏蔽硬件方案",
         "limit": "法律与伦理争议大，未能推广",
         "lang": "zh"},
        {"idx": 5, "authors": "陈伟等", "year": "2024",
         "view": "自律契约机制短期降 18%",
         "limit": "样本 150 人、周期 4 周，缺乏跟踪",
         "lang": "zh"},
        {"idx": 6, "authors": "赵敏", "year": "2023",
         "view": "行为干预理论本土化",
         "limit": "仅理论分析未做实证",
         "lang": "zh"},
        {"idx": 7, "authors": "周平等", "year": "2024",
         "view": "辅导员协同干预方案",
         "limit": "仅小样本 60 人试点",
         "lang": "zh"},
        {"idx": 8, "authors": "Smith et al.", "year": "2020",
         "view": "自律契约干预 6 周后使用率从 82% 降至 61%",
         "limit": "仅美国样本，未跨文化验证",
         "lang": "en"},
        {"idx": 9, "authors": "Brown et al.", "year": "2022",
         "view": "5 校 1500 人 6 月跟踪保持率 65%",
         "limit": "未涉及协同干预机制",
         "lang": "en"},
        {"idx": 10, "authors": "Kumar & Lee", "year": "2022",
         "view": "4 种干预方式比较，契约+提醒组合最佳",
         "limit": "未做跨文化比较",
         "lang": "en"},
        {"idx": 11, "authors": "Davis et al.", "year": "2023",
         "view": "6 种干预方式比较，多维度组合优于单维",
         "limit": "未做长期跟踪",
         "lang": "en"},
        {"idx": 12, "authors": "Brown A", "year": "2023",
         "view": "承诺装置理论使 6 月保持率达 70%",
         "limit": "仅行为经济学视角，未融合教育心理学",
         "lang": "en"},
        {"idx": 13, "authors": "Anderson & Wilson", "year": "2024",
         "view": "干预效果衰减双指数模型",
         "limit": "仅数学建模，未做实证验证",
         "lang": "en"},
        {"idx": 14, "authors": "Lee et al.", "year": "2023",
         "view": "集体主义文化下协同干预优于个人主义",
         "limit": "样本仅 4 国，未含中国大陆",
         "lang": "en"},
        {"idx": 15, "authors": "Park et al.", "year": "2024",
         "view": "韩国高校辅导员协同干预效果显著",
         "limit": "韩国学情与中国大陆有差异",
         "lang": "en"},
    ],

    # v2.1 新增：科学挑战（2 段结构）
    "scientific_challenges": [
        {"title": "三维干预协同机制设计中的混杂变量控制",
         "description": "在课堂手机使用行为干预研究中，『契约+提醒+辅导员协同』三维干预的协同效应识别是核心难题[3][5]。三维干预中各维度可能存在相互抵消效应，且学生自控力、教师授课风格等混杂变量难以完全控制，导致干预效果难以归因。本挑战需要在准实验设计中引入倾向性得分匹配等方法，量化各维度独立贡献与协同贡献。",
         "sub_challenges": [
             {"name": "三维干预的协同效应识别",
              "detail": "传统的单维干预研究将各维度视为独立变量，忽略了维度间的交互作用。本课题需要建立三维干预的协同效应模型，量化各维度独立贡献与协同贡献。具体技术难点包括：① 协同效应的数学定义与度量；② 维度间交互项的识别与建模；③ 协同效应的统计显著性检验方法。已有研究[5]采用方差分解法，但仅能识别主效应，无法捕捉非线性交互。"},
             {"name": "混杂变量的倾向性得分匹配",
              "detail": "学生自控力（个体层）、教师授课风格（班级层）是影响干预效果的关键混杂变量。本课题需要在准实验设计中引入倾向性得分匹配（PSM），通过构建干预组与对照组的可比样本，控制可观测混杂变量[8]。技术难点：① 倾向性得分模型的协变量选择；② 匹配算法（最近邻/卡尺/核匹配）的选型；③ 匹配后平衡性检验；④ 不可观测混杂变量的敏感性分析。"}
         ]},
        {"title": "干预长期效果衰减规律与保持策略",
         "description": "已有干预研究普遍周期 <1 月，缺乏长期效果评估[1][2]。本课题首次引入 6 个月跟踪设计，需要解决长期效果衰减的建模与保持策略设计两大难题。这一挑战对应案例 1 中科学挑战三『深度学习下高性能计算的能耗与实时性关键技术的研究』的思路，即从『短期效果验证』转向『长期效果建模与保持』。",
         "sub_challenges": [
             {"name": "干预效果衰减曲线建模",
              "detail": "行为干预的效果通常随时间衰减，但衰减规律（线性/指数/分段）尚无定论。本课题需要在 6 个月跟踪期内采集 4 个时间点的数据，拟合衰减曲线。技术难点：① 衰减曲线函数族选择（指数衰减/双指数/分段常数）；② 模型参数估计与置信区间；③ 个体差异的随机效应建模；④ 模型选择的信息准则比较。已有研究[13]提出双指数模型，但未做实证验证。"},
             {"name": "效果保持策略设计",
              "detail": "基于衰减曲线，设计效果保持策略是本课题的创新点。技术难点：① 保持策略的触发条件（效果衰减至 X% 触发）；② 策略强度自适应算法；③ 多次干预的疲劳效应控制；④ 策略的伦理审查与隐私保护[12]。已有研究[12]的承诺装置理论提供了一定基础，但缺乏针对课堂场景的适配。"}
         ]}
    ],

    "research_goal": "总目标：开发能将课堂手机使用率降低 ≥30 个百分点、6 个月后效果保持率 ≥70% 的本土化干预方案。阶段目标 1：完成方案设计（2025.06 前）；阶段目标 2：完成 12 周试点实施（2025.10 前）；阶段目标 3：完成效果评估（2025.12 前）；阶段目标 4：完成 6 个月跟踪（2026.03 前）。",

    "research_content": [
        "干预方案设计：基于行为干预理论设计『契约+提醒+协同』三维方案，含 12 项具体干预动作，产出方案文档 1 份。",
        "试点实施：在 3 所高校 6 个班级共 600 名学生中开展 12 周试点，含干预组与对照组，产出实施记录与原始数据。",
        "效果评估：采用前后测+跟踪测设计，分析干预对课堂手机使用率、学习成绩、自我效能感的影响，产出评估报告 1 份。",
        "跟踪研究：干预结束后 6 个月进行跟踪测，评估长期效果，产出跟踪报告 1 份。",
    ],

    "key_problems": [
        "三维干预中各维度的协同机制设计（避免维度间相互抵消）",
        "准实验设计中混杂变量（学生自控力、教师授课风格）的控制",
        "干预长期效果的衰减规律与保持策略",
    ],

    # v2.1 新增：算法对比表
    "method_comparison": {
        "title": "表 1 主流课堂手机干预方法对比表",
        "dimensions": ["技术特点", "优势", "劣势", "适用场景", "短期效率", "长期保持"],
        "methods": [
            {"name": "契约法[3]",
             "values": ["学生签订行为契约，违规扣分",
                        "学生自主性强，伦理风险低",
                        "依赖学生自控力，短期效果",
                        "自控力较强的本科生",
                        "短期降 15~22%",
                        "6 月保持率 40~55%"]},
            {"name": "提醒法[5]",
             "values": ["课前自动提醒，违规震动提醒",
                        "实时性强，操作简单",
                        "容易产生适应性，长期衰减",
                        "中小型课堂",
                        "短期降 18~25%",
                        "6 月保持率 25~35%"]},
            {"name": "屏蔽法[4]",
             "values": ["硬件屏蔽手机信号",
                        "立即生效，效果显著",
                        "法律伦理争议大，不可推广",
                        "涉密/特殊考场",
                        "短期降 80%+",
                        "撤销即失效"]},
        ],
        "conclusion": "基于上表，单一方法均存在局限：契约法依赖自控力、提醒法长期衰减、屏蔽法不可推广。本课题选用『契约+提醒+辅导员协同』三维组合方案，预期短期降 35%+、6 月保持率 70%+，优于任一单维方法的 1.6 倍以上，对应创新点 1。"
    },

    "tech_route": "总体技术路线分 4 阶段（见图 2）：① 准备阶段（2025.04-05），完成文献综述与方案设计，产出综述与方案；② 试点实施（2025.06-09），在 3 校 6 班 600 人中开展 12 周干预，产出实施记录与原始数据；③ 数据分析（2025.10-11），SPSS 27 + Mplus 8 进行统计分析，产出分析报告与论文初稿；④ 跟踪研究（2025.12-2026.03），6 个月后跟踪测 + 论文投稿 + 结题，产出跟踪报告与论文。",

    # v2.1 新增：技术路线图（2 张）
    "tech_roadmap": [
        {"fig_no": "图 1", "title": "项目研究内容及其相互关系",
         "description": "总目标 → 研究内容 1.1（方案设计）→ 研究内容 1.2（试点实施）→ 研究内容 1.3（效果评估与跟踪）→ 研究内容 1.4（长期跟踪研究）。研究内容层层递进，前一阶段的产出作为后一阶段的输入。",
         "nodes": ["总目标：开发本土化课堂手机干预方案",
                   "研究内容 1.1 干预方案设计（产出方案文档）",
                   "研究内容 1.2 试点实施（产出实施记录+原始数据）",
                   "研究内容 1.3 效果评估与跟踪（产出评估报告）",
                   "研究内容 1.4 长期跟踪研究（产出跟踪报告）"]},
        {"fig_no": "图 2", "title": "项目技术路线图",
         "description": "4 阶段流程：准备阶段（2025.04-05）→ 试点实施（2025.06-09）→ 数据分析（2025.10-11）→ 跟踪研究（2025.12-2026.03）。每阶段标注交付物，与十七章进度安排一一对应。",
         "nodes": ["阶段 1：准备阶段（2025.04-05）→ 文献综述 50 篇 + 方案设计",
                   "阶段 2：试点实施（2025.06-09）→ 3 校 6 班 600 人 12 周干预",
                   "阶段 3：数据分析（2025.10-11）→ SPSS + Mplus 统计分析",
                   "阶段 4：跟踪研究（2025.12-2026.03）→ 跟踪测 + 论文 + 结题"]},
    ],

    # v2.1 新增：数学公式
    "formulas": [
        {"idx": 1,
         "formula": "P(y=1|x) = σ(w·x+b) = 1 / (1 + exp(-(w·x+b)))",
         "variables": "y ∈ {0,1} 表示干预后是否达成行为改变（y=1 达成，y=0 未达成）；x = (x₁, x₂, ..., xₙ) 为特征向量，含基线手机使用率、自控力得分、干预维度强度、教师授课风格等 n 个特征；w = (w₁, w₂, ..., wₙ) 为权重向量；b 为偏置项；σ(·) 为 Sigmoid 激活函数。",
         "purpose": "用于干预效果预测。模型训练采用最大似然估计，损失函数为对数损失 L(w,b) = -Σ[yᵢ log(ŷᵢ) + (1-yᵢ) log(1-ŷᵢ)]，其中 ŷᵢ = P(yᵢ=1|xᵢ)，N 为样本量。模型评估采用 AUC、准确率、召回率三项指标，10 折交叉验证避免过拟合。"}
    ],

    "methods": [
        "文献研究法：系统梳理国内外 50+ 篇相关文献，建立理论框架。",
        "准实验法：采用『干预组-对照组』前后测设计，控制混杂变量。",
        "问卷调查法：开发课堂手机使用行为量表（含 5 维度 25 题项），α≥0.85。",
        "深度访谈法：干预后对 30 名学生 + 6 名教师进行半结构化访谈。",
        "数学建模法：采用 Logistic 回归建模干预效果，引入倾向性得分匹配控制混杂变量。",
    ],

    "data_source": "数据来源：3 所合作高校 6 个班级共 600 名学生（干预组 300 + 对照组 300）。实验设备：教育学院教育技术实验室眼动仪 Tobii Pro X3-120、行为观察记录系统 Noldus Observer XT 14。软件工具：SPSS 27、Mplus 8、NVivo 14、Python 3.11（scikit-learn 1.4）。",

    "innovations": [
        "方法创新。已有干预研究多采用单一手段（仅契约/仅提醒/仅屏蔽），效果有限（短期降 15~22%）。本课题融合『契约+提醒+辅导员协同』三维干预，预实验显示三维干预短期降 35%，优于单维方案的 1.6 倍。",
        "数据创新。已有研究样本量普遍 <200、周期 <1 月。本课题样本量 600+、周期 1 年含 6 个月跟踪，是已有研究样本量的 3 倍以上，外推性显著提升。",
    ],

    # v2.1 新增：社会效益量化（5 项简化指标）
    "social_benefits": [
        {"metric": "课堂手机使用率", "baseline": "78.3%", "expected": "≤48%",
         "improvement": "降低 ≥30 个百分点"},
        {"metric": "课堂抬头率", "baseline": "52%", "expected": "≥80%",
         "improvement": "提升 ≥28 个百分点"},
        {"metric": "学习成绩平均分", "baseline": "75.2 分", "expected": "≥82 分",
         "improvement": "提升 ≥7 分"},
        {"metric": "自我效能感得分", "baseline": "3.2/5", "expected": "≥4.0/5",
         "improvement": "提升 ≥0.8 分"},
        {"metric": "教师课堂管理满意度", "baseline": "60%", "expected": "≥85%",
         "improvement": "提升 ≥25 个百分点"},
    ],

    "team_foundation": "团队 3 名成员已修读《教育心理学》《教育统计与测量》《学习行为分析》等核心课程，2 人有校级科研立项参与经验，1 人熟练使用 SPSS、Mplus、Python 等数据分析工具。",
    "advisor_foundation": "指导教师李华教授主持全国教育科学规划课题 1 项，近 3 年发表 CSSCI 论文 5 篇，研究方向匹配度 90%。曾指导 3 项校级科研立项，其中 2 项获优秀结题。",
    "lab_condition": "实验室配备眼动仪 Tobii Pro X3-120、行为观察记录系统 Noldus Observer XT 14、SPSS 27、Mplus 8、NVivo 14、Python 3.11 等软硬件。已与 3 所合作高校签署调研合作协议。",

    "expected_outcomes": [
        "中文核心期刊论文 1 篇（拟投《高等教育研究》）",
        "调研报告 1 份（约 1.5 万字）",
        "干预方案手册 1 套（含 12 项干预动作操作指南）",
    ],

    # v2.1 新增：4 阶段进度安排
    "project_schedule": [
        {"phase": "第一阶段 准备", "time": "2025.04-05（2 月）",
         "work": "① 文献综述 50 篇；② 方案设计；③ 问卷开发与预试；④ 任务分工",
         "output": "综述 1 份 + 方案 1 份 + 问卷 1 套"},
        {"phase": "第二阶段 实施", "time": "2025.06-09（4 月）",
         "work": "① 3 校 6 班 600 人 12 周干预；② 实施记录与数据采集；③ 中期检查与调整",
         "output": "实施记录 + 原始数据"},
        {"phase": "第三阶段 分析", "time": "2025.10-11（2 月）",
         "work": "① SPSS + Mplus 统计分析；② 干预效果评估；③ 论文初稿撰写",
         "output": "分析报告 1 份 + 论文初稿"},
        {"phase": "第四阶段 总结", "time": "2025.12-2026.03（4 月）",
         "work": "① 6 个月后跟踪测；② 论文修改与投稿；③ 结题材料整理",
         "output": "跟踪报告 1 份 + 论文 1 篇 + 结题报告"},
    ],

    "schedule": [
        {"phase": "准备", "time": "2025.04-05", "work": "文献综述 50 篇、方案设计", "output": "综述 1 份 + 方案 1 份"},
    ],

    "budget_items": [
        {"item": "资料费", "amount": "350", "basis": "图书 5 本 × 50 元 + 数据库订阅 100 元"},
        {"item": "调研费", "amount": "1200", "basis": "3 校实地调研 × 400 元（含交通、住宿）"},
        {"item": "材料费", "amount": "600", "basis": "问卷印刷 600 份 × 1 元"},
        {"item": "会议费", "amount": "500", "basis": "参加教育学会议 1 次"},
        {"item": "印刷费", "amount": "450", "basis": "论文版面费 400 + 报告印刷 50"},
        {"item": "其他", "amount": "400", "basis": "礼品感谢 30 份 × 10 元 + 数据备份 U 盘 100 元"},
    ],
    "budget_total": "3500",

    # 参考文献（15 条含 7 篇英文）
    "references": [
        {"ref_type": "journal", "authors": "张华, 李明, 王芳",
         "title": "大学生课堂手机使用行为研究",
         "journal": "高等教育研究", "year": "2022",
         "volume": "43", "issue": "(5)", "pages": "78-86"},
        {"ref_type": "journal", "authors": "李明, 陈伟, 周平, 等",
         "title": "8 所高校大学生课堂手机使用行为调查",
         "journal": "中国高教研究", "year": "2023",
         "volume": "", "issue": "(8)", "pages": "92-98"},
        {"ref_type": "journal", "authors": "王强",
         "title": "课堂积分奖励干预机制研究",
         "journal": "中国高教研究", "year": "2023",
         "volume": "", "issue": "(8)", "pages": "92-98"},
        {"ref_type": "journal", "authors": "刘芳, 陈伟",
         "title": "课堂信号屏蔽方案的伦理思考",
         "journal": "教育研究", "year": "2024",
         "volume": "45", "issue": "(2)", "pages": "112-120"},
        {"ref_type": "journal", "authors": "陈伟, 周平",
         "title": "自律契约机制在课堂手机管理中的应用",
         "journal": "教育发展研究", "year": "2024",
         "volume": "44", "issue": "(3)", "pages": "65-73"},
        {"ref_type": "journal", "authors": "赵敏",
         "title": "行为干预理论本土化研究",
         "journal": "教育学报", "year": "2023",
         "volume": "19", "issue": "(4)", "pages": "88-96"},
        {"ref_type": "journal", "authors": "周平, 吴敏",
         "title": "辅导员协同干预方案设计与试点",
         "journal": "思想教育研究", "year": "2024",
         "volume": "", "issue": "(5)", "pages": "102-108"},
        {"ref_type": "journal", "authors": "Smith J, Brown R, Davis M",
         "title": "Mobile phone use in college classrooms: A multi-campus study",
         "journal": "Journal of Educational Psychology", "year": "2020",
         "volume": "112", "issue": "(8)", "pages": "1542-1558"},
        {"ref_type": "journal", "authors": "Brown A, Wilson B, Lee C",
         "title": "Long-term effects of classroom phone interventions: A 6-month follow-up",
         "journal": "Computers & Education", "year": "2022",
         "volume": "178", "issue": "", "pages": "104392"},
        {"ref_type": "conference", "authors": "Kumar S, Lee H",
         "title": "Comparative study of four classroom phone interventions",
         "conference": "Proceedings of the 2022 International Conference on Educational Technology",
         "city": "New York", "publisher": "ACM", "year": "2022", "pages": "215-222"},
        {"ref_type": "journal", "authors": "Davis M, Anderson K, Wilson R",
         "title": "Six intervention methods for classroom phone management: A comparative study",
         "journal": "Educational Technology Research and Development", "year": "2023",
         "volume": "71", "issue": "(4)", "pages": "1825-1845"},
        {"ref_type": "thesis", "authors": "Brown A",
         "title": "Behavioral economics approaches to classroom management",
         "city": "Stanford", "school": "Stanford University", "year": "2023"},
        {"ref_type": "journal", "authors": "Anderson K, Wilson R",
         "title": "A double-exponential decay model for intervention effects",
         "journal": "Journal of Behavioral Decision Making", "year": "2024",
         "volume": "37", "issue": "(2)", "pages": "145-162"},
        {"ref_type": "journal", "authors": "Lee H, Park J, Kim S",
         "title": "Cross-cultural comparison of classroom phone interventions",
         "journal": "International Journal of Educational Research", "year": "2023",
         "volume": "117", "issue": "", "pages": "102126"},
        {"ref_type": "web", "authors": "教育部",
         "title": "关于进一步加强大学生课堂管理的指导意见",
         "publish_date": "2024-03-15", "access_date": "2025-02-10",
         "url": "http://www.moe.gov.cn/xxx", "is_government": True},
    ],
}


# CLI 入口

def main():
    parser = argparse.ArgumentParser(
        description="大学生校级科研立项申请书 docx 生成器（v2.1）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "  python build.py --data data.json --out output.docx --version enhanced\n"
            "\n"
            "校级科研立项特征：经费 2~5 千、周期 1 年、团队 1~3 人、"
            "必须有 GB/T 7714 参考文献（≥15 条含 5 篇英文）。\n"
            "v2.1 新增：国家政策引用、科学挑战 2 段、算法对比表、技术路线图、"
            "数学公式、社会效益 5 项、进度 4 阶段、3 档字数版本。\n"
            "JSON 字段定义详见 SKILL.md 第十一章。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    parser.add_argument("--version", type=str, default=None,
                        choices=["standard", "enhanced", "peak"],
                        help="字数版本（覆盖 JSON 中 word_version 字段）")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        if args.version:
            data = dict(data)
            data["word_version"] = args.version
        print(f"ℹ️ 使用内置示例数据生成演示文档（版本：{data.get('word_version', 'enhanced')}）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
            if args.version:
                data["word_version"] = args.version
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
