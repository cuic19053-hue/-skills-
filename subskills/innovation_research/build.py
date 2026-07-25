#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
大学生创新创业训练计划-创新训练项目申报书 docx 生成器（v3.0 案例优化版）

按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 一级标题：黑体三号，居中
- 二级标题：黑体小三，左对齐
- 三级标题：宋体四号加粗
- 表格：宋体五号，居中

v3.0 升级要点：
- 基于 28 页真实案例（消防无人机多模态融合研究项目）优化
- 新增 8 个 JSON 字段处理：policy_citations / scientific_challenges /
  literature_review / algorithm_comparison / tech_roadmap / formulas /
  economic_benefits / project_schedule
- DEFAULT_DATA 改为消防无人机多模态融合主题，对齐案例 1
- 行数从 966 行扩展到 ≥1200 行

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段定义详见 SKILL.md 第十一章。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（含中文 eastAsia 字体）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
        space_before=12, space_after=12,
    )


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6, space_after=6,
    )


def add_heading_level3(doc, text: str):
    """三级标题：宋体四号加粗，左对齐"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6, space_after=3,
    )


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent,
        line_spacing=1.5,
    )


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None,
                        header_bold: bool = True):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=header_bold)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def merge_vertical_cells(table, col_idx: int, start_row: int, end_row: int):
    """纵向合并单元格（用于签字栏预留空白）"""
    cells = [table.rows[r].cells[col_idx] for r in range(start_row, end_row + 1)]
    merged = cells[0]
    for c in cells[1:]:
        merged = merged.merge(c)


def add_figure_caption(doc, caption: str):
    """添加图注（黑体五号居中）"""
    add_paragraph_with_format(
        doc, caption,
        font_name=FONT_HEI, font_size=SIZE_WU, bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
        space_before=3, space_after=6,
    )


def add_table_caption(doc, caption: str):
    """添加表题（黑体五号居中，表上方）"""
    add_paragraph_with_format(
        doc, caption,
        font_name=FONT_HEI, font_size=SIZE_WU, bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
        space_before=6, space_after=3,
    )


def add_formula_paragraph(doc, formula_no: str, expression: str,
                          variables: str = "", algorithm: str = ""):
    """添加数学公式段落（公式居中 + 编号右对齐 + 变量解释）"""
    if algorithm:
        add_paragraph_with_format(
            doc, algorithm,
            font_name=FONT_HEI, font_size=SIZE_SI, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=False,
            space_before=6, space_after=3,
        )
    # 公式行：居中公式 + 右侧编号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    run = p.add_run(f"    {expression}        {formula_no}")
    set_run_font(run, font_name=FONT_TIMES, font_size=SIZE_XIAO_SI)
    # 变量解释
    if variables:
        add_paragraph_with_format(
            doc, f"其中，{variables}",
            font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=True,
            line_spacing=1.5,
            space_after=3,
        )


def add_quote_paragraph(doc, text: str):
    """添加引用段落（仿宋小四，左缩进，单倍行距）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.left_indent = Pt(24)
    pf.right_indent = Pt(24)
    pf.line_spacing = 1.5
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, font_name=FONT_FANGSONG, font_size=SIZE_XIAO_SI)


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """大创-创新训练项目申报书 docx 构建器（v3.0 案例优化版）"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def add_h1(self, text):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text):
        return add_heading_level2(self.doc, text)

    def add_h3(self, text):
        return add_heading_level3(self.doc, text)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_table(self, headers, rows, col_widths=None):
        return add_table_from_data(self.doc, headers, rows, col_widths)

    def add_page_break(self):
        add_page_break(self.doc)

    # --------------------------------------------------------
    # 封面
    # --------------------------------------------------------

    def _add_cover(self):
        """封面：黑体二号标题 + 副标题 + 5 行下划线信息"""
        for _ in range(2):
            self.doc.add_paragraph()

        level = self._get("project_level", default="国家级")
        title = f"{level}大学生创新创业训练计划项目申报书"
        add_paragraph_with_format(
            self.doc, title,
            font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_before=12, space_after=12,
        )

        subtitle = f"（{self._get('project_type', default='创新训练项目')}）"
        add_paragraph_with_format(
            self.doc, subtitle,
            font_name=FONT_HEI, font_size=SIZE_SAN, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_after=24,
        )

        for _ in range(2):
            self.doc.add_paragraph()

        info_items = [
            ("项目名称", self._get("project_name")),
            ("项目负责人", self._get("leader_name")),
            ("指导教师", self._get("advisor_name")),
            ("所在学院", self._get("college")),
            ("申报日期", self._get("apply_date")),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI,
                         font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True

        self.add_page_break()

    # --------------------------------------------------------
    # 基本信息
    # --------------------------------------------------------

    def _add_basic_info_table(self):
        """一、基本信息表（9 行 2 列）"""
        self.add_h1("一、基本信息")
        basic = self._get("basic_info", default={})
        if not isinstance(basic, dict):
            basic = {}
        budget_str = str(basic.get("budget", self._get("budget_total", "")))
        if budget_str and not budget_str.endswith("元"):
            budget_str = f"{budget_str} 元"

        rows = [
            ["项目名称", basic.get("project_name", self._get("project_name"))],
            ["项目类型", basic.get("project_type",
                                    self._get("project_type", "创新训练项目"))],
            ["项目来源", basic.get("project_source", "A 学生自主选题")],
            ["所属学科", basic.get("discipline", "")],
            ["起止时间", basic.get("duration", "")],
            ["申请经费", budget_str],
            ["负责人", basic.get("leader_info",
                                  f"{self._get('leader_name')} / "
                                  f"{self._get('leader_id')} / "
                                  f"{self._get('leader_major')} / "
                                  f"{self._get('leader_grade')} / "
                                  f"{self._get('leader_phone')}")],
            ["团队成员", basic.get("team_members", "")],
            ["指导教师", basic.get("advisor_info",
                                    f"{self._get('advisor_name')} / "
                                    f"{self._get('advisor_title')} / "
                                    f"{self._get('advisor_research')}")],
        ]
        self.add_table(["项目", "内容"], rows, col_widths=[4.5, 11.5])

    # --------------------------------------------------------
    # 项目简介
    # --------------------------------------------------------

    def _add_abstract(self):
        """二、项目简介（300~500 字）"""
        self.add_h1("二、项目简介")
        abstract = self._get("abstract", default="")
        if abstract:
            self.add_para(abstract)
        else:
            self.add_para("（请填写项目简介，300~500 字，按 4 句结构撰写："
                          "痛点+做什么 / 怎么做+量化目标 / 产出什么 / 现状。）")

    # --------------------------------------------------------
    # 立项依据与研究内容（v3.0 重点章节）
    # --------------------------------------------------------

    def _add_background(self):
        """三、立项依据与研究内容 - 引言段"""
        self.add_h1("三、立项依据与研究内容")
        intro = self._get("background_intro",
                          default="（一）立项依据")
        if intro:
            self.add_h2(intro)
        background = self._get("background", default=[])
        if isinstance(background, str):
            background = [background]
        if not background:
            self.add_para("（请填写立项依据引言，3~5 段，每段一个小方向："
                          "时代背景/现实痛点/国内外研究现状/项目意义。）")
        else:
            for para in background:
                self.add_para(para)

    # --------------------------------------------------------
    # v3.0 新增：国家政策引用（policy_citations）
    # --------------------------------------------------------

    def _add_policy_citations(self):
        """3.1 国家规划与技术需求 - 国家政策引用规范"""
        self.add_h2("3.1 国家规划与技术需求")
        self.add_para("本项目立项紧扣国家应急救援、智慧消防、无人机产业发展"
                      "战略导向。下文按时间倒序梳理 8+ 项相关政策文件，"
                      "摘录与项目相关的关键表述，论证项目立项的政策合规性"
                      "与战略必要性。")

        policies = self._get("policy_citations", default=[])
        if not policies:
            self.add_para("（请填写国家政策引用，≥8 条（国家级）/ ≥5 条（省级）/ "
                          "≥3 条（校级）。每条含 4 要素：发文机关+文号+标题+时间，"
                          "并摘录与项目相关的关键表述。应包含最高领导人讲话。"
                          "按时间倒序排列。）")
            return

        for i, policy in enumerate(policies, 1):
            if not isinstance(policy, dict):
                continue
            issuer = policy.get("issuer", "")
            doc_no = policy.get("doc_no", "")
            title = policy.get("title", "")
            date = policy.get("date", "")
            key_quote = policy.get("key_quote", "")
            citation_text = (
                f"[{i}] {date} {issuer}《{title}》"
                f"{('（' + doc_no + '）') if doc_no else ''}：{key_quote}"
            )
            self.add_para(citation_text)

    # --------------------------------------------------------
    # v3.0 新增：科学挑战 3 段（scientific_challenges）
    # --------------------------------------------------------

    def _add_scientific_challenges(self):
        """3.2 科学挑战 3 段结构"""
        self.add_h2("3.2 科学挑战")
        self.add_para("本项目针对消防无人机多模态融合目标检测任务，提炼 3 个"
                      "层层递进的科学挑战：数据层（多模态融合）→ 算法层"
                      "（3D 目标检测）→ 系统层（能耗与实时性）。")

        challenges = self._get("scientific_challenges", default=[])
        if not challenges:
            self.add_para("（请填写 3 段科学挑战，层层递进（数据→算法→系统 / "
                          "感知→认知→决策）。每个挑战必须有文献上标 [N] 支撑，"
                          "子挑战用 bullet point 拆解，每个 5-8 句详细说明。"
                          "标题必须含关键技术词。）")
            return

        for i, ch in enumerate(challenges, 1):
            if not isinstance(ch, dict):
                continue
            title = ch.get("title", "")
            description = ch.get("description", "")
            sub_challenges = ch.get("sub_challenges", [])

            self.add_h3(f"科学挑战{'一二三'[i-1] if i <= 3 else str(i)}：{title}")
            if description:
                self.add_para(description)
            if sub_challenges:
                for sub in sub_challenges:
                    if isinstance(sub, dict):
                        sub_name = sub.get("name", "")
                        sub_detail = sub.get("detail", "")
                        self.add_para(f"● {sub_name}：{sub_detail}")
                    elif isinstance(sub, str):
                        self.add_para(f"● {sub}")

    # --------------------------------------------------------
    # v3.0 新增：国内外研究现状综述（literature_review）
    # --------------------------------------------------------

    def _add_literature_review(self):
        """3.3 国内外研究现状综述"""
        self.add_h2("3.3 研究意义与国内外研究现状")

        # 研究意义（从 background 段中提取或单独字段）
        significance = self._get("research_significance", default="")
        if significance:
            self.add_h3("3.3.1 研究意义")
            if isinstance(significance, list):
                for s in significance:
                    self.add_para(s)
            else:
                self.add_para(significance)
        else:
            self.add_h3("3.3.1 研究意义")
            self.add_para("（请填写研究意义，3-5 段，每段一个小方向："
                          "无人机目标跟踪与智慧消防的重要意义 / 多模态融合技术"
                          "为 3D 目标跟踪与检测开辟新途径 / 构建并行运算构架对"
                          "多模态深度学习的重要意义。）")

        # 国内外研究现状
        self.add_h3("3.3.2 国内外研究现状")
        review_text = self._get("literature_review_text", default="")
        if review_text:
            if isinstance(review_text, list):
                for para in review_text:
                    self.add_para(para)
            else:
                self.add_para(review_text)
        else:
            self.add_para("（请填写国内外研究现状，按技术方向分段，每段引用"
                          "3-5 篇文献。国家级 ≥30 篇，英文 ≥60%，含 SCI 一区/二区"
                          "≥5 篇，近 3 年 ≥50%。GB/T 7714-2015 格式。）")

        # 研究现状总结
        self.add_h3("3.3.3 对现有研究现状的总结与分析")
        summary = self._get("literature_summary", default="")
        if summary:
            self.add_para(summary)
        else:
            self.add_para("（请填写研究现状总结与分析，指出已有研究的 gap 与"
                          "本项目的差异点。）")

        # 文献列表
        self.add_h3("3.3.4 参考文献")
        refs = self._get("literature_review", default=[])
        if refs:
            for i, ref in enumerate(refs, 1):
                if not isinstance(ref, dict):
                    continue
                author = ref.get("author", "")
                title = ref.get("title", "")
                journal = ref.get("journal", "")
                year = ref.get("year", "")
                volume = ref.get("volume", "")
                pages = ref.get("pages", "")
                lang = ref.get("language", "zh")
                if lang == "en":
                    ref_str = (f"[{i}] {author} {title}[J]. {journal}, "
                               f"{year}{(', ' + volume) if volume else ''}"
                               f"{(': ' + pages) if pages else ''}.")
                else:
                    ref_str = (f"[{i}] {author}. {title}[J]. {journal}, "
                               f"{year}{(', ' + volume) if volume else ''}"
                               f"{(': ' + pages) if pages else ''}.")
                self.add_para(ref_str, indent=False)

    # --------------------------------------------------------
    # v3.0 新增：算法对比表（algorithm_comparison）
    # --------------------------------------------------------

    def _add_algorithm_comparison(self):
        """3.4 算法对比表"""
        self.add_h2("3.4 算法对比分析")
        comp = self._get("algorithm_comparison", default={})
        if not isinstance(comp, dict):
            comp = {}

        if comp:
            title = comp.get("title", "表 1 算法对比表")
            add_table_caption(self.doc, title)
            headers = comp.get("headers", ["模型", "文献", "技术特点",
                                            "优势", "劣势", "适用场景"])
            rows = comp.get("rows", [])
            self.add_table(headers, rows)
            conclusion = comp.get("conclusion", "")
            if conclusion:
                self.add_h3("选型结论")
                self.add_para(conclusion)
        else:
            self.add_para("（请填写算法对比表，≥3 算法 × ≥4 维度"
                          "（技术特点/优势/劣势/适用场景）。每算法有文献引用，"
                          "最后给选型结论，如『基于上表，本项目选用 YOLOv8 算法』。）")

    # --------------------------------------------------------
    # 研究内容、目标与关键科学问题
    # --------------------------------------------------------

    def _add_research_content(self):
        """四、项目研究内容、目标与关键科学问题"""
        self.add_h1("四、项目研究内容、目标与关键科学问题")

        self.add_h2("（一）研究内容")
        contents = self._get("research_content", default=[])
        if isinstance(contents, str):
            contents = [contents]
        if contents:
            for i, c in enumerate(contents, 1):
                self.add_para(f"{i}. {c}")
        else:
            self.add_para("（请填写研究内容，3~5 个子任务，每个 200~400 字，"
                          "结构：任务名+做什么+方法+产出。每个子任务对应一个科学挑战。）")

        self.add_h2("（二）研究目标")
        goal = self._get("research_goal", default="")
        if goal:
            self.add_para(goal)
        else:
            self.add_para("（请填写研究目标，1 个总目标 + 3~4 个阶段目标，全部可量化。）")

        self.add_h2("（三）拟解决的关键科学问题")
        problems = self._get("key_problems", default=[])
        if isinstance(problems, str):
            problems = [problems]
        if problems:
            for i, q in enumerate(problems, 1):
                self.add_para(f"{i}. {q}")
        else:
            self.add_para("（请填写关键问题，2~3 个，每个 100~150 字讲清技术难点。）")

    # --------------------------------------------------------
    # 创新点
    # --------------------------------------------------------

    def _add_innovation(self):
        """五、项目创新点（600~900 字，至少 3 个，国家级）"""
        self.add_h1("五、项目创新点")
        innovations = self._get("innovations", default=[])
        if isinstance(innovations, str):
            innovations = [innovations]
        if innovations:
            for i, inv in enumerate(innovations, 1):
                self.add_para(f"创新点 {i}：{inv}")
        else:
            self.add_para("（请填写创新点，至少 3 个（国家级）/ 2 个（省级）/ "
                          "1 个（校级），每个 150~250 字。结构：[类型]。"
                          "传统方法[描述]，本项目[方法]，[量化优势]。"
                          "禁止使用『首次』『先进』『实现』等无支撑词。）")

    # --------------------------------------------------------
    # 研究方案及可行性分析（含技术路线图、数学公式）
    # --------------------------------------------------------

    def _add_research_scheme(self):
        """六、研究方案及可行性分析"""
        self.add_h1("六、研究方案及可行性分析")

        # 6.1 研究方法和思路
        self.add_h2("（一）研究方法和思路")
        route = self._get("tech_route", default="")
        if route:
            self.add_para(route)
        else:
            self.add_para("（请填写研究方法和思路，1 段文字 + 图 1 项目研究方法图。"
                          "总体技术路线分阶段展开，含大模型串联小模型、轻量化"
                          "卷积神经网络、CPU/GPU 异构并行运算三大核心技术。）")

        # 6.2 技术路线图（v3.0 新增 tech_roadmap 字段）
        self.add_h2("（二）技术路线图")
        self._add_tech_roadmap()

        # 6.3 实验方法和关键技术（含数学公式）
        self.add_h2("（三）实验方法和关键技术")
        methods = self._get("methods", default=[])
        if isinstance(methods, str):
            methods = [methods]
        if methods:
            for i, m in enumerate(methods, 1):
                self.add_para(f"{i}. {m}")
        else:
            self.add_para("（请填写实验方法和关键技术，含数学公式与算法流程图。"
                          "每个核心算法至少 1 个数学公式，公式后必须解释变量含义。）")

        # 6.4 数学公式（v3.0 新增 formulas 字段）
        self.add_h2("（四）核心算法与数学公式")
        self._add_formulas()

        # 6.5 数据来源
        self.add_h2("（五）数据来源与实验条件")
        data_src = self._get("data_source", default="")
        if data_src:
            self.add_para(data_src)
        else:
            self.add_para("（请填写数据来源、实验设备型号、软件工具及版本。）")

        # 6.6 可行性分析
        self.add_h2("（六）可行性分析")
        feasibility = self._get("feasibility_analysis", default="")
        if feasibility:
            if isinstance(feasibility, list):
                for f in feasibility:
                    self.add_para(f)
            else:
                self.add_para(feasibility)
        else:
            self.add_para("（请填写可行性分析，3 段：研究方案可行性 + 工作基础"
                          "可行性 + 研究条件可行性。）")

    # --------------------------------------------------------
    # v3.0 新增：技术路线图（tech_roadmap，≥3 张图）
    # --------------------------------------------------------

    def _add_tech_roadmap(self):
        """技术路线图：≥3 张图（研究内容关系图 + 研究方法图 + 技术路线图）"""
        roadmaps = self._get("tech_roadmap", default=[])
        if not roadmaps:
            self.add_para("（请填写技术路线图，≥3 张图：研究内容关系图 + "
                          "研究方法图 + 技术路线图。每图含图号+标题+自含说明，"
                          "节点对应章节编号。）")
            return

        for rm in roadmaps:
            if not isinstance(rm, dict):
                continue
            fig_no = rm.get("figure_no", "")
            title = rm.get("title", "")
            description = rm.get("description", "")
            image_path = rm.get("image_path", "")

            # 尝试插入图片
            if image_path and os.path.exists(image_path):
                try:
                    p = self.doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(image_path, width=Cm(15))
                except Exception as e:
                    self.add_para(f"（图片插入失败：{e}）")
            else:
                # 用文字框代替图片
                add_quote_paragraph(self.doc, f"【{fig_no} {title}】\n{description}")

            # 图注
            add_figure_caption(self.doc, f"{fig_no} {title}")

            # 图说明
            if description:
                self.add_para(description)

    # --------------------------------------------------------
    # v3.0 新增：数学公式（formulas）
    # --------------------------------------------------------

    def _add_formulas(self):
        """数学公式：每个核心算法 ≥1 个公式，公式编号 + 变量解释"""
        formulas = self._get("formulas", default=[])
        if not formulas:
            self.add_para("（请填写数学公式，每个核心算法至少 1 个数学公式。"
                          "公式必须编号（如式(1)、式(2)），公式后必须解释每个"
                          "变量的含义。必须配算法流程图（用文字描述流程）。"
                          "算法名称必须有缩写。）")
            return

        for f in formulas:
            if not isinstance(f, dict):
                continue
            no = f.get("no", "")
            expression = f.get("expression", "")
            variables = f.get("variables", "")
            algorithm = f.get("algorithm", "")
            flowchart = f.get("flowchart", "")
            add_formula_paragraph(self.doc, no, expression, variables, algorithm)
            if flowchart:
                self.add_para(f"算法流程：{flowchart}")

    # --------------------------------------------------------
    # v3.0 新增：社会经济效益量化（economic_benefits）
    # --------------------------------------------------------

    def _add_economic_benefits(self):
        """七、社会经济效益分析（≥10 项指标 × 4 列）"""
        self.add_h1("七、社会经济效益分析")
        self.add_para("本项目通过对比传统模式与系统模式的关键指标，量化评估"
                      "系统社会经济效益。指标涵盖经济损失、识别精度、重构精度、"
                      "响应时间、调度准确率、装备效率等多个维度。")

        benefits = self._get("economic_benefits", default=[])
        if not benefits:
            self.add_para("（请填写社会经济效益，≥10 项指标 × 4 列"
                          "（指标名/传统基准/系统预期/提升幅度）。"
                          "提升幅度必须用百分比+倍数双重表达。）")
            return

        add_table_caption(self.doc, "表 1 系统社会经济效益量化评估表")
        headers = ["评估指标", "传统模式基准值", "系统模式预期值", "效益提升幅度"]
        rows = []
        for b in benefits:
            if not isinstance(b, dict):
                continue
            rows.append([
                b.get("indicator", ""),
                b.get("baseline", ""),
                b.get("expected", ""),
                b.get("improvement", ""),
            ])
        self.add_table(headers, rows, col_widths=[5.0, 3.5, 3.5, 4.0])

    # --------------------------------------------------------
    # 项目实施方案与进度安排（v3.0 升级为 project_schedule 4 阶段）
    # --------------------------------------------------------

    def _add_implementation_plan(self):
        """八、项目实施方案与进度安排（4 阶段甘特图式表格）"""
        self.add_h1("八、项目实施方案与进度安排")

        # 优先使用 v3.0 的 project_schedule 字段（4 阶段）
        schedule = self._get("project_schedule", default=[])
        if schedule:
            self.add_para("本项目按 4 阶段实施，每阶段明确起止月份、主要任务"
                          "与阶段成果。国家级周期 24 个月，省级 18 个月，"
                          "校级 12 个月。")
            add_table_caption(self.doc, "表 2 项目实施进度安排表")
            rows = []
            for s in schedule:
                if not isinstance(s, dict):
                    continue
                phase = s.get("phase", "")
                time = s.get("time", "")
                tasks = s.get("tasks", [])
                if isinstance(tasks, list):
                    tasks_str = "；".join(f"①{t}" if i == 0 else f"{i+1}{t}"
                                          for i, t in enumerate(tasks))
                else:
                    tasks_str = str(tasks)
                output = s.get("output", "")
                rows.append([phase, time, tasks_str, output])
            self.add_table(
                ["阶段", "起止时间", "主要研究内容", "阶段成果"],
                rows,
                col_widths=[3.5, 3.5, 6.0, 3.0],
            )
            return

        # 兼容 v2.0 的 schedule 字段
        schedule_v2 = self._get("schedule", default=[])
        if schedule_v2:
            rows = []
            for s in schedule_v2:
                if not isinstance(s, dict):
                    continue
                rows.append([
                    s.get("phase", ""),
                    s.get("time", ""),
                    s.get("work", ""),
                    s.get("output", ""),
                ])
            self.add_table(
                ["阶段", "时间", "主要工作", "阶段成果"],
                rows,
                col_widths=[2.5, 3.0, 6.5, 4.0],
            )
        else:
            self.add_para("（请填写进度安排，4 阶段按月划分，每阶段含 3-5 任务"
                          "与交付物。国家级 24 个月 / 省级 18 个月 / 校级 12 个月。）")

    # --------------------------------------------------------
    # 预期成果
    # --------------------------------------------------------

    def _add_expected_results(self):
        """九、预期成果（必须可量化）"""
        self.add_h1("九、预期成果")
        outcomes = self._get("expected_outcomes", default=[])
        if isinstance(outcomes, str):
            outcomes = [outcomes]
        if outcomes:
            for o in outcomes:
                self.add_para(f"• {o}", indent=False)
        else:
            self.add_para("（请填写预期成果，每项含数量+级别+平台。"
                          "如：SCI 二区论文 1 篇（拟投 IEEE TGRS）、"
                          "发明专利 1 项、原型系统 1 套。）", indent=False)

    # --------------------------------------------------------
    # 经费预算
    # --------------------------------------------------------

    def _add_budget(self):
        """十、经费预算（3 列表格：科目/金额/计算依据）"""
        self.add_h1("十、经费预算")
        items = self._get("budget_items", default=[])
        if items:
            rows = []
            total = 0
            for b in items:
                if not isinstance(b, dict):
                    continue
                amount_str = str(b.get("amount", "0"))
                try:
                    amount_num = int(amount_str)
                except ValueError:
                    amount_num = 0
                total += amount_num
                rows.append([
                    b.get("item", ""),
                    f"{amount_num} 元",
                    b.get("basis", ""),
                ])
            rows.append(["合计", f"{total} 元", ""])
            self.add_table(
                ["预算科目", "金额", "计算依据"],
                rows,
                col_widths=[3.5, 3.0, 9.5],
            )
        else:
            self.add_para("（请填写经费预算，5 类标准科目：资料费/调研差旅费/"
                          "实验材料费/会议费/印刷复印。每项金额非整数，附计算依据。"
                          "国家级 1~2 万 / 省级 5 千~1 万 / 校级 2~3 千。）")

    # --------------------------------------------------------
    # 前期工作基础
    # --------------------------------------------------------

    def _add_preliminary_work(self):
        """十一、前期工作基础（500~700 字，3 子节）"""
        self.add_h1("十一、前期工作基础")

        self.add_h2("（一）团队基础")
        team = self._get("team_foundation", default="")
        self.add_para(team if team else
                      "（请填写团队基础：成员相关课程、项目经验、技能匹配度。）")

        self.add_h2("（二）指导教师基础")
        advisor = self._get("advisor_foundation", default="")
        self.add_para(advisor if advisor else
                      "（请填写指导教师基础：主持项目、发表论文、研究方向匹配度。）")

        self.add_h2("（三）实验条件")
        lab = self._get("lab_condition", default="")
        self.add_para(lab if lab else
                      "（请填写实验条件：实验室设备、软件平台、合作单位支持。）")

    # --------------------------------------------------------
    # 签字栏
    # --------------------------------------------------------

    def _add_signature_section(self):
        """十二/十三、指导教师意见、学院评审意见"""
        self.add_h1("十二、指导教师意见")
        for _ in range(6):
            self.doc.add_paragraph()
        self.add_para(
            "指导教师签字：____________________    "
            "日期：______年____月____日",
            indent=False,
        )

        self.add_h1("十三、学院评审意见")
        for _ in range(6):
            self.doc.add_paragraph()
        self.add_para(
            "学院盖章：____________________    "
            "日期：______年____月____日",
            indent=False,
        )

        if self._get("include_school_approval", default=False):
            self.add_h1("十四、学校审批意见")
            for _ in range(6):
                self.doc.add_paragraph()
            self.add_para(
                "学校盖章：____________________    "
                "日期：______年____月____日",
                indent=False,
            )

    # --------------------------------------------------------
    # 主构建方法
    # --------------------------------------------------------

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 12 栏目 + 8 案例字段，生成 docx

        Args:
            data: 申报书字段字典
            output_path: 输出 docx 路径

        Returns:
            实际保存路径
        """
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()

            # 封面 + 基本信息 + 简介
            self._add_cover()
            self._add_basic_info_table()
            self._add_abstract()

            # 立项依据（含 v3.0 新增 4 字段：政策引用/科学挑战/文献综述/算法对比）
            self._add_background()
            self._add_policy_citations()
            self._add_scientific_challenges()
            self._add_literature_review()
            self._add_algorithm_comparison()

            # 研究内容 + 创新点
            self._add_research_content()
            self._add_innovation()

            # 研究方案（含 v3.0 新增 2 字段：技术路线图/数学公式）
            self._add_research_scheme()

            # 社会经济效益（v3.0 新增 economic_benefits 字段）
            self._add_economic_benefits()

            # 进度安排（v3.0 升级 project_schedule 4 阶段）
            self._add_implementation_plan()

            # 预期成果 + 经费 + 前期基础 + 签字栏
            self._add_expected_results()
            self._add_budget()
            self._add_preliminary_work()
            self._add_signature_section()

            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申报书已生成：{output_path}")
        return str(output_path)

    # --------------------------------------------------------
    # 数据校验
    # --------------------------------------------------------

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []
        p0_fields = [
            ("project_name", "项目名称"),
            ("leader_name", "负责人姓名"),
            ("advisor_name", "指导教师姓名"),
            ("college", "所在学院"),
        ]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        basic = self._get("basic_info", default={})
        if not isinstance(basic, dict):
            basic = {}
        if not basic.get("discipline") and not self._get("discipline"):
            warnings.append("缺少 所属学科（discipline）")
        if not basic.get("duration") and not self._get("duration"):
            warnings.append("缺少 起止时间（duration）")

        if not self._get("abstract"):
            warnings.append("缺少 项目简介（abstract），将使用占位文本")
        if not self._get("background"):
            warnings.append("缺少 立项依据（background），将使用占位文本")
        if not self._get("innovations"):
            warnings.append("缺少 创新点（innovations），将使用占位文本")

        # v3.0 案例规范字段校验
        level = str(self._get("project_level", default="国家级"))
        policy_min = 8 if level == "国家级" else (5 if level == "省级" else 3)
        policies = self._get("policy_citations", default=[])
        if isinstance(policies, list) and len(policies) < policy_min:
            warnings.append(f"政策引用仅 {len(policies)} 条，{level}需 ≥{policy_min} 条")

        challenges = self._get("scientific_challenges", default=[])
        if isinstance(challenges, list) and len(challenges) < 3:
            warnings.append(f"科学挑战仅 {len(challenges)} 段，需 ≥3 段")

        lit_min = 30 if level == "国家级" else (20 if level == "省级" else 10)
        refs = self._get("literature_review", default=[])
        if isinstance(refs, list) and len(refs) < lit_min:
            warnings.append(f"文献仅 {len(refs)} 篇，{level}需 ≥{lit_min} 篇")

        comp = self._get("algorithm_comparison", default={})
        if isinstance(comp, dict) and comp:
            rows = comp.get("rows", [])
            if isinstance(rows, list) and len(rows) < 3:
                warnings.append(f"算法对比表仅 {len(rows)} 算法，需 ≥3 算法")

        roadmaps = self._get("tech_roadmap", default=[])
        if isinstance(roadmaps, list) and len(roadmaps) < 3:
            warnings.append(f"技术路线图仅 {len(roadmaps)} 张，需 ≥3 张")

        formulas = self._get("formulas", default=[])
        if isinstance(formulas, list) and len(formulas) < 1:
            warnings.append("数学公式缺失，每个核心算法需 ≥1 个公式")

        benefits = self._get("economic_benefits", default=[])
        if isinstance(benefits, list) and len(benefits) < 10:
            warnings.append(f"社会经济效益仅 {len(benefits)} 项，需 ≥10 项")

        schedule = self._get("project_schedule", default=[])
        if isinstance(schedule, list) and len(schedule) != 4:
            warnings.append(f"进度安排 {len(schedule)} 阶段，需 4 阶段")

        items = self._get("budget_items", default=[])
        if items:
            total = 0
            for b in items:
                if isinstance(b, dict):
                    try:
                        total += int(b.get("amount", 0))
                    except (ValueError, TypeError):
                        pass
            budget_total_str = str(self._get("budget_total", default="")).strip()
            try:
                budget_total_num = int(budget_total_str)
            except ValueError:
                budget_total_num = -1
            if budget_total_num >= 0 and total != budget_total_num:
                warnings.append(
                    f"预算合计 {total} 元 与申请经费 {budget_total_num} 元不一致"
                )

        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据（v3.0 消防无人机多模态融合主题，对齐案例 1）
# ============================================================

DEFAULT_DATA = {
    "project_name": "多模态融合的无人机消防目标检测系统",
    "project_level": "国家级",
    "project_type": "创新训练项目",
    "leader_name": "张三",
    "leader_id": "202212345",
    "leader_major": "自动化",
    "leader_grade": "2022 级 大三",
    "leader_phone": "138XXXXXXXX",
    "advisor_name": "李教授",
    "advisor_title": "教授",
    "advisor_research": "多模态融合与无人机智能感知",
    "college": "自动化工程学院",
    "apply_date": "2025 年 3 月 15 日",
    "basic_info": {
        "project_name": "多模态融合的无人机消防目标检测系统",
        "project_type": "创新训练项目",
        "project_source": "A 学生自主选题",
        "discipline": "0808 自动化",
        "duration": "2025.04-2027.03（24 个月）",
        "budget": "18000",
        "leader_info": "张三 / 202212345 / 自动化 / 大三 / 138XXXXXXXX",
        "team_members": "李四（数据采集与标注）、王五（硬件搭建与部署）、"
                       "赵六（算法设计与训练）、钱七（系统测试与文档）",
        "advisor_info": "李教授 / 教授 / 多模态融合与无人机智能感知 / 139XXXXXXXX",
    },
    "abstract": "本项目针对当前消防无人机在火灾场景下目标检测响应慢、误报率高、"
                "复杂环境鲁棒性不足的痛点，开发基于多模态融合技术"
                "（视觉+红外+激光雷达+音频）的无人机异构并行移动目标定位与"
                "跟踪系统。通过 Big-Little Model 大模型串联小模型架构、改进 "
                "YOLO-LRP 轻量化卷积神经网络、CPU/GPU 异构并行运算三大核心"
                "技术，实现火场目标检测准确率 ≥92%、3D 重构关键尺寸误差 "
                "≤5cm、单次响应时间 <1 秒。预期产出：1 套原型系统（部署在"
                "合作消防中队试运行）、2 项发明专利（与合作企业共同申请）、"
                "3 篇论文（含 SCI 二区 1 篇、中文核心 2 篇）、1 套故障样本"
                "数据集（5000+ 样本）。项目已完成前期调研与 500 样本预训练，"
                "预实验准确率已达 88%，技术路线成熟可行。",
    "background": [
        "时代背景：随着我国智慧消防战略推进，消防无人机产业 2024 年市场规模"
        "突破 120 亿元（应急管理部数据），年均增速 35% 以上。但消防无人机"
        "在火场复杂环境下的目标检测仍面临响应慢、误报率高、鲁棒性不足等"
        "技术瓶颈，传统单一视觉传感器方案难以满足实战需求。",
        "现实痛点：调研 XX 市消防救援支队发现，2022 年 1-9 月我国共接报"
        "火灾 63.68 万起，直接财产损失 55 亿元，火灾死亡 1441 人、受伤 "
        "1640 人。传统人工巡检响应时间超过 48 小时，火灾期间因信息延迟"
        "导致的额外损失达 5%~8%。某大型商业综合体 2024 年因火场目标识别"
        "延迟，单次损失超 800 万元。",
        "国内外研究现状：早期方法（Smith 2020, Wang 2021）主要基于 SVM、"
        "决策树等传统机器学习，依赖人工特征提取，准确率约 75%~80%。"
        "近年来 Zhang (2022)、Li (2023) 尝试引入深度学习，但所用数据集"
        "多为实验室仿真。本项目关键差异：（1）使用真实火场 5000+ 样本；"
        "（2）覆盖 8 类典型火场目标；（3）多模态融合（视觉+红外+音频+点云）。",
        "项目意义：理论上探索多模态融合与大模型轻量化在火场极端环境下的"
        "适用边界；实践上与 XX 消防支队合作开发可落地系统，预期将目标检测"
        "响应时间从 48 小时缩短至 1 小时，单次火灾损失降低 20%~40%；"
        "社会上助力国家智慧消防与应急救援体系现代化建设。",
    ],
    "policy_citations": [
        {"issuer": "国务院办公厅", "doc_no": "国办发〔2026〕8 号",
         "title": "关于加强基层消防工作的意见", "date": "2026.03",
         "key_quote": "强调要充分发挥无人机、机器人等新型装备在基层消防救援"
                     "中的作用，推动基层消防装备智能化升级，构建『早发现、"
                     "早预警、早处置』的基层消防新格局。"},
        {"issuer": "工信部等三部门", "doc_no": "工信部联装〔2025〕198 号",
         "title": "应急装备产业重点产品发展指导目录（2025 版）", "date": "2025.12",
         "key_quote": "将多模态融合无人机、3D 火场重构系统、智能目标识别"
                     "装备列入重点发展方向，鼓励产学研协同攻关。"},
        {"issuer": "国务院新闻办公室", "doc_no": "",
         "title": "新时代的中国国家安全》白皮书", "date": "2025.05",
         "key_quote": "明确指出要构建『大安全、大应急』框架，加强应急救援"
                     "装备智能化、无人化升级，提升重特大灾害事故应对能力。"},
        {"issuer": "中共中央办公厅、国务院办公厅", "doc_no": "中办发〔2024〕23 号",
         "title": "关于推进新型城市基础设施建设打造韧性城市的意见", "date": "2024.12",
         "key_quote": "强调发展智慧消防，利用物联网、人工智能、无人机等"
                     "技术提升城市消防安全水平，推动消防装备与城市基础"
                     "设施深度融合。"},
        {"issuer": "工信部等四部门", "doc_no": "工信部联重装〔2024〕12 号",
         "title": "通用航空装备创新应用实施方案（2024-2030 年）", "date": "2024.03",
         "key_quote": "提出到 2030 年通用航空装备及相关产业形成万亿级"
                     "市场规模，重点推动无人机在应急救援、消防灭火等"
                     "公共服务领域的应用。"},
        {"issuer": "应急管理部、工业和信息化部", "doc_no": "应急〔2023〕116 号",
         "title": "关于加快应急机器人发展的指导意见", "date": "2023.12",
         "key_quote": "提到突出建设测试基地及公共服务平台，研发全国应急"
                     "管理系统无人机综合信息平台，服务应急管理需求。"},
        {"issuer": "应急管理部", "doc_no": "应急〔2022〕15 号",
         "title": "『十四五』应急救援力量建设规划", "date": "2022.06",
         "key_quote": "提出加快构建大型固定翼灭火飞机、灭火直升机与无人机"
                     "高低搭配、布局合理、功能互补的应急救援航空器体系。"},
        {"issuer": "国务院", "doc_no": "国发〔2021〕36 号",
         "title": "『十四五』国家应急体系规划", "date": "2021.12",
         "key_quote": "表明推广运用智能机器人、无人机等高技术配送装备，"
                     "推动应急物资储运设备集装单元化发展，提升应急运输"
                     "调度效率。"},
        {"issuer": "习近平", "doc_no": "",
         "title": "向国家综合性消防救援队伍授旗致辞", "date": "2018.11",
         "key_quote": "指出组建国家综合性消防救援队伍，是构建新时代国家"
                     "应急救援体系的重要举措，对提高防灾减灾救灾能力、"
                     "维护社会公共安全、保护人民生命财产安全具有重大意义。"},
        {"issuer": "习近平", "doc_no": "",
         "title": "在中共十九届中央政治局第十九次集体学习时的讲话", "date": "2019.11",
         "key_quote": "要加强风险评估和检测预警，提升多灾种和灾害链综合"
                     "监测、风险早期识别和预报预警能力。要加强应急预案"
                     "管理，健全应急预案体系，落实各环节责任和措施。"},
    ],
    "scientific_challenges": [
        {"title": "大语言模型背景下多种模态信息融合技术的研究",
         "description": "在大语言模型背景下，模态（modality）是指一种表达或"
                       "感知信息的方式，例如文本、图像、视频、音频等[21]。"
                       "单一模态模型处理技术就是用大数据训练一个小模型，"
                       "解决单一的任务。多模态模型利用不同模态之间的互补"
                       "和协同，来提高模型的性能和泛化能力[22]。然而大模型"
                       "处理技术也存在计算资源需求高、能源消耗大的弊端，"
                       "因此，大模型轻量化成为一个新的研究挑战[33-35]。",
         "sub_challenges": [
             {"name": "异构数据时空对齐",
              "detail": "多模态传感器（视觉/红外/音频/激光雷达）数据采样"
                       "频率、坐标系、时间戳均不一致，需研究高精度时空"
                       "对齐算法。火场环境强干扰进一步加大对齐难度，传统"
                       "基于卡尔曼滤波的方法在烟雾环境下误差超过 15%。"},
             {"name": "大模型轻量化",
              "detail": "多模态大模型参数量通常超过 100M，难以在无人机"
                       "嵌入式平台部署。需研究模型蒸馏、重参数化、量化"
                       "等技术，在保持精度前提下将参数量降至 25M 以下。"},
             {"name": "跨模态语义适配",
              "detail": "搜索系统协同信号与多模态大模型语义信息存在不一致"
                       "性，需研究跨模态对齐损失函数与适配网络，构建通用"
                       "的搜索多模态大模型。"},
         ]},
        {"title": "特殊场景复杂环境下 3D 移动目标检测技术的研究",
         "description": "无人机跟踪的目标是通过搭载的相机捕获运动目标在"
                       "航空视角下的信息。该跟踪性能受到目标运动、无人机"
                       "飞行状态以及运行环境的共同影响[7-9]。在基于深度"
                       "学习的无人机目标跟踪领域，存在多个待探索的方向。",
         "sub_challenges": [
             {"name": "目标跟踪技术控制",
              "detail": "无人机不断飞行导致目标处于不断变化的复杂背景中，"
                       "如何有效处理目标与背景的划分成为研究的关键。移动"
                       "目标的姿态变化和位置变化引起的遮挡现象都需要得到"
                       "合理处理。火场烟雾、强光、热辐射等极端环境进一步"
                       "加剧了目标跟踪的难度，传统基于可见光的跟踪方法在"
                       "烟雾环境下几乎失效[13-16]。"},
             {"name": "模型优化",
              "detail": "模型优化面临挑战，需要繁琐的调参和预训练过程以"
                       "发挥最佳性能。许多终端用户缺乏相关专业知识，难以"
                       "从复杂的模型和数据中选择最优资源。对个别用户而言，"
                       "由于计算资源不足，训练时间过长，使得算法调优变得"
                       "更为困难[17-20]。"},
             {"name": "运行速度及实时处理",
              "detail": "深度学习下实时跟踪的计算量巨大，对处理器性能有"
                       "较高要求。减少跟踪算法的计算量是一个重要的挑战"
                       "因素。火场救援场景要求毫秒级响应，传统算法难以"
                       "满足[15]。"},
         ]},
        {"title": "深度学习下高性能计算的能耗与实时性关键技术的研究",
         "description": "在当前高性能计算领域，能耗问题成为备受关注的焦点。"
                       "传统单核处理器的性能受到能耗因素的限制，而 CPU/GPU "
                       "异构系统则成为一种兼顾通用性和能耗的有前景高性能"
                       "计算系统[6]。新型高性能计算系统为大规模科学计算"
                       "和工程模拟类应用提供了发展机遇，但也使许多高性能"
                       "计算面临巨大挑战。",
         "sub_challenges": [
             {"name": "缺少易用、高效的编程环境",
              "detail": "目前广泛使用的 GPU 编程模型包括 AMD Stream SDK、"
                       "OpenCL 和 NVIDIA CUDA。对于大型 CPU/GPU 异构高性能"
                       "计算系统，通常采用在计算节点间使用 MPI，而在计算"
                       "节点内使用 GPU 厂商提供的编程模型的方式。亟需建立"
                       "能够有效开发整个 CPU/GPU 异构系统计算资源的并行"
                       "编程模型。"},
             {"name": "缺少准确的可扩展性模型",
              "detail": "可扩展性在并行计算中是一个关键研究领域。高性能"
                       "计算系统的可扩展性受多种因素影响，包括算法内在"
                       "并行性、各种并行开销、系统规模和问题规模等。应"
                       "发展一种可扩展性度量指标，以全面评估并行系统/问题"
                       "规模变化对高性能计算机和并行算法性能的影响。"},
         ]},
    ],
    "research_significance": [
        "（1）无人机目标跟踪与智慧消防紧密结合的重要意义：灭火救援对维护"
        "社会稳定、保障人民安全有十分重要的作用，无人机在消防救援领域的"
        "应用具有深远的意义[4]。无人机能够快速、灵活地进入事故现场，"
        "提供实时的空中视角，为消防人员提供全面的信息，有助于快速而"
        "精准地定位事故点，降低救援响应时间。",
        "（2）多模态融合技术为 3D 目标跟踪与检测开辟了新途径：多模态融合"
        "在无人机目标跟踪与检测领域的研究具有重要的理论与实际意义[5]。"
        "将异构数据进行融合，不仅能够弥补单一传感器的不足，更能够实现"
        "信息的互补与协同，为无人机目标跟踪与检测任务提供了新的解决途径。",
        "（3）构建并行运算构架对多模态深度学习的重要意义：在进行多目标"
        "任务跟踪的时候，需要实时处理视频信息，对图像处理的速率有一定"
        "要求，采用 GPU 并行运算技术[6]能够促进深度学习的运算效率，"
        "实现神经网络的优化。",
    ],
    "literature_review_text": [
        "（1）利用卷积神经网络进行目标跟踪的研究现状：随着卷积神经网络"
        "技术的发展，基于卷积神经网络的图像识别技术在各行业得到广泛应用。"
        "在深度学习领域，R-CNN、SSD 以及 YOLO 系列是目标检测领域的代表"
        "性算法[13-20]。Girshick 等人提出的 R-CNN[13] 是基于候选区域"
        "提取的目标检测算法系列奠基之作；Liu 等人提出的 SSD[15] 通过"
        "均匀采样实现单阶段检测；Redmon 创造了 YOLO[17]，将目标检测"
        "简化为单一回归问题，提升了效率和速度。",
        "（2）多模态融合技术在机器视觉领域的研究现状：面向深度学习的"
        "多模态融合技术是指机器从文本、图像、语音等领域获取信息，实现"
        "转换与融合，提升模型性能[21]。多模态融合方法分为模型无关方法"
        "[22-24]和基于模型方法[25-32]，其中基于模型方法包括多核学习"
        "（MKL）、图像模型（GM）、神经网络（NN）等。近期神经网络方法"
        "通过使用循环神经网络（RNN）和长短期记忆网络（LSTM）来融合"
        "时间多模态信息[30]，相对于 MKL 和 GM 方法表现出更优的性能。",
        "（3）近期大模型多模态融合技术发展：近一年来，基于大模型的多"
        "模态融合技术得到飞速发展。DynStatF 特征融合策略[33]通过当前"
        "单帧的精确位置信息增强多帧提供的丰富语义信息；CDDFuse 多模态"
        "特征融合法[34]利用 Restormer 模块提取跨模态的浅层特征，通过"
        "双分支 Transformer-CNN 结构处理全局和局部特征；MFDF 网络融合"
        "法[35]用于实时 RGB-D-T 显著目标检测。",
    ],
    "literature_summary": "多模态融合使得机器能够从多个感知通道获取信息，"
                         "从而提高了对环境的综合感知能力。尽管已经取得了"
                         "显著进展，但多模态融合技术仍然面临一些挑战，"
                         "如异构数据融合和模态不平衡。在火灾搜索监测系统"
                         "领域研究中，受限于搜索系统协同信号与多模态大模型"
                         "语义信息的不一致性，如何将大模型与小模型结合，"
                         "如何使协同信号与多模态语义信息相适配，构建出通用"
                         "的搜索多模态大模型是本申请的重难点问题。本项目"
                         "拟采用集成学习的方式，将大型模型与小型模型结合，"
                         "研究轻量化多模态融合技术下的无人机目标检测。",
    "literature_review": [
        {"author": "应急部", "title": "关于印发《『十四五』应急救援力量建设规划》的通知",
         "journal": "中华人民共和国国务院公报", "year": "2022", "volume": "(25)",
         "pages": "42-49", "language": "zh", "level": "核心"},
        {"author": "中国应急管理编辑部", "title": "突出重点领域聚焦实战应用——《关于加快应急机器人发展的指导意见》解读",
         "journal": "中国应急管理", "year": "2024", "volume": "(01)",
         "pages": "76-79", "language": "zh", "level": "核心"},
        {"author": "Maksymiuk R, et al.", "title": "5G Network-Based Passive Radar for Drone Detection",
         "journal": "2023 24th International Radar Symposium (IRS), IEEE",
         "year": "2023", "volume": "", "pages": "1-10", "language": "en", "level": "EI"},
        {"author": "Gupta A, Bhatnagar A, Mehta A", "title": "Application of drones in maritime industry (firefighting)",
         "journal": "Bulletin of Marine Science and Technology",
         "year": "2021", "volume": "14", "pages": "59-69", "language": "en", "level": "SCI"},
        {"author": "de Rochechouart M, et al.", "title": "Drone Tracking Based on the Fusion of Staring Radar and Camera Data: An Experimental Study",
         "journal": "2023 IEEE Radar Conference (RadarConf23)",
         "year": "2023", "volume": "", "pages": "01-06", "language": "en", "level": "EI"},
        {"author": "Baji T", "title": "GPU: the biggest key processor for AI and parallel processing",
         "journal": "Photomask Japan 2017, SPIE",
         "year": "2017", "volume": "10454", "pages": "24-29", "language": "en", "level": "EI"},
        {"author": "Wei Z, et al.", "title": "UAV-Assisted Data Collection for Internet of Things: A Survey",
         "journal": "IEEE Internet of Things Journal",
         "year": "2022", "volume": "9(17)", "pages": "15460-15483",
         "language": "en", "level": "SCI 一区"},
        {"author": "Farmani N, et al.", "title": "Tracking multiple mobile targets using cooperative Unmanned Aerial Vehicles",
         "journal": "2015 International Conference on Unmanned Aircraft Systems (ICUAS)",
         "year": "2015", "volume": "", "pages": "395-400", "language": "en", "level": "EI"},
        {"author": "Su J, et al.", "title": "AI meets UAVs: A survey on AI empowered UAV perception systems for precision agriculture",
         "journal": "Neurocomputing", "year": "2023", "volume": "518",
         "pages": "242-270", "language": "en", "level": "SCI 二区"},
        {"author": "Sarkar N I, Gul S", "title": "Artificial Intelligence-Based Autonomous UAV Networks: A Survey",
         "journal": "Drones", "year": "2023", "volume": "7(5)",
         "pages": "322", "language": "en", "level": "SCI 三区"},
        {"author": "芦艳春, 周开园, 张建杰", "title": "无人机的发展现状及其在航空应急救援领域的应用综述",
         "journal": "医疗卫生装备", "year": "2023", "volume": "44(10)",
         "pages": "108-113", "language": "zh", "level": "核心"},
        {"author": "钟映霞", "title": "无人机在森林火灾调查中的应用研究",
         "journal": "华南农业大学硕士学位论文", "year": "2023", "volume": "",
         "pages": "1-78", "language": "zh", "level": "硕论"},
        {"author": "Girshick R, et al.", "title": "Rich feature hierarchies for accurate object detection and semantic segmentation",
         "journal": "Proceedings of the IEEE conference on computer vision and pattern recognition",
         "year": "2014", "volume": "", "pages": "580-587", "language": "en", "level": "CVPR"},
        {"author": "张倩, 周平平, 王公堂, 等", "title": "基于合成图像的 Faster R-CNN 森林火灾烟雾检测",
         "journal": "山东师范大学学报: 自然科学版", "year": "2019", "volume": "34(2)",
         "pages": "180-185", "language": "zh", "level": "核心"},
        {"author": "Liu W, et al.", "title": "SSD: Single Shot MultiBox Detector",
         "journal": "European conference on computer vision, Springer",
         "year": "2016", "volume": "", "pages": "21-37", "language": "en", "level": "ECCV"},
        {"author": "刘丽娟, 陈松楠", "title": "一种基于改进 SSD 的烟雾实时检测模型",
         "journal": "信阳师范学院学报 (自然科学版)", "year": "2020", "volume": "33(2)",
         "pages": "305-311", "language": "zh", "level": "核心"},
        {"author": "Redmon J, et al.", "title": "You only look once: Unified, real-time object detection",
         "journal": "Proceedings of the IEEE conference on computer vision and pattern recognition",
         "year": "2016", "volume": "", "pages": "779-788", "language": "en", "level": "CVPR"},
        {"author": "He K, et al.", "title": "Deep residual learning for image recognition",
         "journal": "Proceedings of the IEEE conference on computer vision and pattern recognition",
         "year": "2016", "volume": "", "pages": "770-778", "language": "en", "level": "CVPR"},
        {"author": "Shen D, et al.", "title": "Flame detection using deep learning",
         "journal": "2018 4th International conference on control, automation and robotics (ICCAR)",
         "year": "2018", "volume": "", "pages": "416-420", "language": "en", "level": "EI"},
        {"author": "祖鑫萍, 李丹", "title": "基于无人机图像和改进 YOLOv3-SPP 算法的森林火灾烟雾识别方法",
         "journal": "林业工程学报", "year": "2022", "volume": "7(05)",
         "pages": "142-149", "language": "zh", "level": "核心"},
        {"author": "Jokinen K, Raike A", "title": "Multimodality-technology, visions and demands for the future",
         "journal": "Proceedings of the 1st Nordic Symposium on Multimodal Interfaces",
         "year": "2003", "volume": "", "pages": "239-251", "language": "en", "level": "其他"},
        {"author": "Castellano G, et al.", "title": "Emotion Recognition through Multiple Modalities: Face, Body Gesture, Speech",
         "journal": "Affect and Emotion in Human-Computer Interaction",
         "year": "2008", "volume": "", "pages": "92-103", "language": "en", "level": "其他"},
        {"author": "Ramírez G A, et al.", "title": "Modeling Latent Discriminative Dynamic of Multi-dimensional Affective Signals",
         "journal": "Affective Computing and Intelligent Interaction",
         "year": "2011", "volume": "", "pages": "396-406", "language": "en", "level": "其他"},
        {"author": "Lan Z, et al.", "title": "Multimedia classification and event detection using double fusion",
         "journal": "Multimedia Tools and Applications",
         "year": "2014", "volume": "71", "pages": "333-347", "language": "en", "level": "SCI"},
        {"author": "Bucak S S, et al.", "title": "Multiple Kernel Learning for Visual Object Recognition: A Review",
         "journal": "IEEE Transactions on Pattern Analysis and Machine Intelligence",
         "year": "2014", "volume": "36", "pages": "1354-1369",
         "language": "en", "level": "SCI 一区"},
        {"author": "Jaques N, et al.", "title": "Multi-task, multi-kernel learning for estimating individual wellbeing",
         "journal": "Proc. NIPS Workshop on Multimodal Machine Learning",
         "year": "2015", "volume": "898", "pages": "3", "language": "en", "level": "其他"},
        {"author": "Gurban M, et al.", "title": "Dynamic modality weighting for multi-stream in audio-visual speech recognition",
         "journal": "International Conference on Multimodal Interaction",
         "year": "2008", "volume": "", "pages": "237-240", "language": "en", "level": "其他"},
        {"author": "Baltrušaitis T, et al.", "title": "Dimensional affect recognition using Continuous Conditional Random Fields",
         "journal": "2013 10th IEEE International Conference and Workshops on Automatic Face and Gesture Recognition (FG)",
         "year": "2013", "volume": "", "pages": "1-8", "language": "en", "level": "EI"},
        {"author": "Jiang X, et al.", "title": "The classification of multi-modal data with hidden conditional random field",
         "journal": "Pattern Recognit. Lett", "year": "2015", "volume": "51",
         "pages": "63-69", "language": "en", "level": "SCI"},
        {"author": "Kahou S E, et al.", "title": "EmoNets: Multimodal deep learning approaches for emotion recognition in video",
         "journal": "Journal on Multimodal User Interfaces",
         "year": "2015", "volume": "10", "pages": "99-111", "language": "en", "level": "SCI"},
        {"author": "Ngiam J, et al.", "title": "Multimodal Deep Learning",
         "journal": "International Conference on Machine Learning",
         "year": "2011", "volume": "", "pages": "689-696", "language": "en", "level": "ICML"},
        {"author": "Chen S, Jin Q", "title": "Multi-modal Dimensional Emotion Recognition using Recurrent Neural Networks",
         "journal": "Proceedings of the 5th International Workshop on Audio/Visual Emotion Challenge",
         "year": "2015", "volume": "", "pages": "49-56", "language": "en", "level": "其他"},
        {"author": "Rong Y, et al.", "title": "DynStatF: An Efficient Feature Fusion Strategy for LiDAR 3D Object Detection",
         "journal": "2023 IEEE/CVF Conference on Computer Vision and Pattern Recognition Workshops (CVPRW)",
         "year": "2023", "volume": "", "pages": "3238-3247", "language": "en", "level": "CVPRW"},
        {"author": "Zhao Z, et al.", "title": "CDDFuse: Correlation-Driven Dual-Branch Feature Decomposition for Multi-Modality Image Fusion",
         "journal": "2023 IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR)",
         "year": "2023", "volume": "", "pages": "5906-5916", "language": "en", "level": "CVPR"},
        {"author": "Song K, et al.", "title": "Lightweight multi-level feature difference fusion network for RGB-D-T salient object detection",
         "journal": "J. King Saud Univ. Comput. Inf. Sci",
         "year": "2023", "volume": "35", "pages": "101702", "language": "en", "level": "SCI"},
    ],
    "algorithm_comparison": {
        "title": "表 1.2 基于深度学习的目标跟踪检测算法优缺点对比分析",
        "headers": ["模型", "文献", "技术特点", "优势", "劣势", "适用场景"],
        "rows": [
            ["R-CNN 系列", "[13,14]", "两阶段检测，先区域提取再分类",
             "识别精度高，对小目标敏感",
             "速度慢，每张图需上千次前向计算",
             "静态场景高精度识别"],
            ["SSD 系列", "[15,16]", "单阶段检测，多尺度特征图",
             "速度较快，精度适中",
             "小目标信息丢失，密集场景表现差",
             "实时性要求中等场景"],
            ["YOLO 系列", "[17-20]", "单阶段回归，端到端网络",
             "检测速度最快，实时性高，参数量小",
             "位置精度较差，召回率低",
             "实时嵌入式部署"],
            ["YOLO-LRP（本项目）", "[本]", "重参数化 MobileOne + GhostConv + SE 注意力",
             "参数减少 69%，速度提升 3.75×",
             "训练复杂度高",
             "火场实时检测"],
        ],
        "conclusion": "基于上表，相比 R-CNN 和 SSD 系列深度学习模型，"
                      "YOLO 系列模型检测精度最大，模型内存小，而且实时性"
                      "最好，这样可以在嵌入式设备中很好的应用，同时达到"
                      "对火灾检测实时性的要求。本项目在 YOLOv10 基础上"
                      "进一步优化，引入 MobileOne 重参数化骨干网络、OREPA "
                      "在线卷积重参数化、GhostConv 模块、SE 注意力机制，"
                      "在保持检测精度的同时大幅降低参数量与推理时延，是"
                      "火场实时检测场景的最佳选择。",
    },
    "research_content": [
        "多模态融合技术下的大模型串联小模型结构研究（针对科学问题一）："
        "开发大模型提取和融合来自不同模态的特征，利用小模型针对火灾场景"
        "进行优化，实现快速准确的目标检测。多模态模型融合技术整合声音、"
        "图像、点云等异构数据，通过双向长短期记忆网络（BiLSTM）融合语音"
        "特征序列，结合视觉、激光雷达等传感器数据，利用卷积神经网络和"
        "点云处理网络的融合实现对目标在三维空间中的检测和定位。预期产出："
        "多模态融合框架 1 套、跨模态对齐算法 1 套。",
        "轻量化自适应卷积神经网络识别与检测研究（针对科学问题二）：引入"
        "高性能低耗资的小模型，实现『降耗增效』。在主干网络中引入无参"
        "注意力机制（ECA），在特征融合网络中加入 SimAM 注意力机制，"
        "将 LAAN 与 SimAM 结合提高运算精度。对于损失函数，改进 MPDIoU "
        "算法，提出基于质心距离的快速收敛函数 FCDIoU 来增强目标回归"
        "能力。预期产出：YOLO-LRP 改进算法 1 套、模型权重 1 套。",
        "CPU/GPU 异构体系并行计算技术研究（针对科学问题三）：在异构计算"
        "方面，利用 CPU 和 GPU 异构计算资源，通过构建 GPU CUDA 并行运算"
        "平台，给不同的移动目标分配到不同的处理器进行并行运算处理，"
        "加快整个跟踪过程的运行效率。研究协同感知并行可扩展性模型，"
        "构建 MPI+OpenMP/CUDA 混合编程模型。预期产出：并行加速比 ≥8×、"
        "能效比 ≥3×。",
        "系统集成与火场场景验证：搭建多模态融合异构并行无人机信息交互"
        "火灾检测系统原型，部署在合作消防中队开展 3 个月火场场景测试，"
        "收集真实场景数据评估泛化能力。预期产出：原型系统 1 套、测试"
        "报告 1 份。",
    ],
    "research_goal": "总目标：开发准确率 ≥92%、3D 重构误差 ≤5cm、响应时间 "
                     "<1 秒的消防无人机多模态融合目标检测系统，覆盖火场"
                     "目标识别、定位、跟踪全流程。阶段目标 1：完成 5000+ "
                     "火场样本数据集构建（2025.09 前）；阶段目标 2：完成 "
                     "YOLO-LRP 改进算法训练，准确率 ≥88%（2026.03 前）；"
                     "阶段目标 3：完成多模态融合，准确率 ≥92%、3D 重构"
                     "误差 ≤5cm（2026.09 前）；阶段目标 4：原型系统部署"
                     "试运行，火场场景验证通过（2027.02 前）。",
    "key_problems": [
        "如何通过综合不同模态信息（视觉/红外/音频/点云）提高系统对火灾"
        "复杂环境的适应性和鲁棒性？涉及多模态时空对齐、跨模态语义适配、"
        "大模型与小模型协同优化等子问题。",
        "如何在多模态信息融合框架下提高无人机对 3D 目标检测的高效实时"
        "感知力？涉及轻量化网络设计、卷积重参数化、注意力机制优化、并行"
        "可扩展性度量等子问题。",
        "如何在 CPU/GPU 异构计算体系下实现深度学习模型的高效能耗比？涉及"
        "异构编程模型、任务调度策略、显存管理优化等子问题。",
    ],
    "innovations": [
        "方法创新。传统消防无人机目标检测依赖单一视觉传感器（准确率约 "
        "75%~80%），本项目采用 Big-Little Model 多模态融合架构，大模型"
        "处理视觉+红外+音频+点云异构数据，小模型 YOLO-LRP 进行轻量化识别。"
        "预实验准确率已达 88%，目标达 92%，相比传统单模态提升 12~17 个"
        "百分点。",
        "方法创新。传统 YOLOv10 主干网络参数量大、推理慢（约 80M 参数、"
        "45ms/帧），本项目引入重参数化 MobileOne 骨干网络 + OREPA 在线"
        "卷积重参数化 + GhostConv 模块，参数量降至 25M、推理速度提升至 "
        "12ms/帧，参数减少 69%、速度提升 3.75×。",
        "方法创新。传统单核 CPU 处理多目标跟踪无法满足实时性（处理 1080p "
        "视频约 200ms/帧），本项目构建 CPU/GPU 异构并行运算平台，采用 "
        "MPI+OpenMP/CUDA 混合编程模型，并行加速比达 8×，能效比提升 3×，"
        "支持 8 路 4K 视频并行处理。",
        "数据创新。公开数据集多为实验室仿真（如 MIT Fire Dataset 仅含 "
        "70 段视频），本项目与 XX 消防支队合作采集真实火场场景样本 "
        "5000+，覆盖 8 类典型火场目标（被困人员/危险品/烟雾/火源/消防员/"
        "装备/车辆/结构），是已有公开数据集规模的 70 倍。",
    ],
    "tech_route": "本申请旨在研究和开发一种新型的大模型结合小模型"
                  "（Big-Little Model）结构，通过大模型处理复杂的多模态"
                  "数据融合，小模型负责针对具体任务的高效识别与检测。这种"
                  "多模态融合方法，通过有效整合声音信息、图片信息以及点云"
                  "信息，实现对 3D 环境中目标的跟踪与检测。本课题将综合"
                  "模拟计算、实验和测试分析等多种研究方法结合开展研究工作，"
                  "对项目中拟解决的关键问题采取应对性的研究。",
    "tech_roadmap": [
        {"figure_no": "图 1", "title": "项目研究内容及其相互关系图",
         "description": "顶部：系统框架设计→数据编码译码→大模型轻量化；"
                       "中部：科学问题一（多模态融合技术下的大模型串联"
                       "小模型）→科学问题二（特殊复杂场景目标检测跟踪及"
                       "小模型轻量化优化）→科学问题三（CPU/GPU 异构体系的"
                       "并行计算理论分析）；底部：开展试验验证和应用——基于"
                       "多模态融合技术的异构并行运算信息交互搜索系统验证"
                       "与应用。节点对应章节：研究内容一/二/三 ↔ §四 研究内容子任务一/二/三。",
         "image_path": ""},
        {"figure_no": "图 2", "title": "项目研究方法和思路图",
         "description": "顶部：科学挑战一/二/三 ↔ 多模态融合技术 / 微小型"
                       "飞行器目标检测识别跟踪 / CUDA 并行计算编程模型；"
                       "中部：研究内容一/二/三。研究内容一：设计有效的融合"
                       "策略以最大化不同模态数据的互补优势；研究内容二：对 "
                       "YOLO 深度学习算法进行改进；研究内容三：构建 CPU/GPU "
                       "异构系统的并行编程 MPI+OpenMP/CUDA 模型；底部：前期"
                       "研究基础→理论工具（最小生成树/光流法/SVM/强化学习/"
                       "深度学习/迁移学习）→仿真/试验/验证工具；三类成果："
                       "理论成果（火灾环境下多模态 3D 目标跟踪体系框架）/ "
                       "实物成果（消防无人机模型原理样机 2 套 + 软硬件仿真"
                       "实验平台）/ 其他成果（发明专利 3 项以上 + SCI 论文 "
                       "6 篇以上 + 培养学生 8 名）。",
         "image_path": ""},
        {"figure_no": "图 3", "title": "项目拟采用的实施技术路线图",
         "description": "顶部：国家战略（加快推进智慧消防建设 / 建立应急"
                       "救援航空器体系 / 创新驱动发展战略）→技术需求（模型"
                       "优化 / 平台支撑 / 实时高效目标定位跟踪）→关键技术"
                       "（智能稳定性需求 / 高效实时性需求 / 大模型大数据大算力）；"
                       "中部：维度（降耗增效 / 自适应轻量化 / 准确性及鲁棒性）"
                       "→实施路线（多模态融合 MFT / YOLO-AFP / FCDIoU / "
                       "SimAM / SPConv / 多模态嵌入对齐 / 4G/5G / 深度学习"
                       "循环&迭代训练）→维度（高并发计算能力 / 能效比优势 / "
                       "可扩展模型 / 异构并行运算）；底部：应用层面（高效稳定"
                       "智能体系平台建设 / 灾害管理与救援 / 现场侦查与检测 / "
                       "灾后评估）。",
         "image_path": ""},
    ],
    "formulas": [
        {"no": "(1)", "algorithm": "算法 1：MFCC 音频特征提取",
         "expression": "MFCC(t, i) = sqrt(2/M) * Σ_j log(E_mel(t,j)) * cos[iπ(j-0.5)/M], i=1,2,...,P",
         "variables": "M 为梅尔滤波器个数，E_mel(t, j) 为 t 时刻第 j 个"
                     "滤波器输出的能量，P 为 MFCC 系数维数。",
         "flowchart": "音频数据提取 → 预处理（预加重、分帧、加窗）→ FFT 快速"
                     "傅里叶变换 → MFCC 特征提取 → 归一化处理 → BiLSTM 网络"
                     "计算 → CTC 连接时序分类 → 输出识别结果"},
        {"no": "(2)", "algorithm": "算法 2：BiLSTM 遗忘门",
         "expression": "f_t = sigmoid(W_f · [h_{t-1}, x_t] + b_f)",
         "variables": "W_f 为遗忘门权重矩阵，h_{t-1} 为上一时刻隐藏状态，"
                     "x_t 为当前输入，b_f 为偏置，sigmoid 为激活函数。",
         "flowchart": "输入 x_t 与 h_{t-1} 拼接 → 与 W_f 矩阵乘 → 加偏置 "
                     "b_f → sigmoid 激活 → 输出遗忘门 f_t"},
        {"no": "(3)", "algorithm": "算法 2：BiLSTM 输入门",
         "expression": "i_t = sigmoid(W_i · [h_{t-1}, x_t] + b_i)",
         "variables": "W_i 为输入门权重矩阵，b_i 为偏置，其他同式(2)。",
         "flowchart": "同遗忘门，仅权重矩阵换为 W_i"},
        {"no": "(4)", "algorithm": "算法 2：BiLSTM 候选状态",
         "expression": "c̃_t = tanh(W_c · [h_{t-1}, x_t] + b_c)",
         "variables": "W_c 为候选状态权重矩阵，b_c 为偏置，tanh 为双曲正切"
                     "激活函数。",
         "flowchart": "输入拼接 → 矩阵乘 → 加偏置 → tanh 激活 → 输出候选状态 c̃_t"},
        {"no": "(5)", "algorithm": "算法 2：BiLSTM 细胞状态更新",
         "expression": "c_t = f_t ⊙ c_{t-1} + i_t ⊙ c̃_t",
         "variables": "⊙ 为哈达玛积（按元素相乘），f_t 为遗忘门输出，"
                     "i_t 为输入门输出，c̃_t 为候选状态。",
         "flowchart": "遗忘门 × 上一细胞状态 + 输入门 × 候选状态 → 当前细胞状态 c_t"},
        {"no": "(6)", "algorithm": "算法 2：BiLSTM 输出门与隐藏状态",
         "expression": "o_t = sigmoid(W_o · [h_{t-1}, x_t] + b_o);  h_t = o_t ⊙ tanh(c_t)",
         "variables": "o_t 为输出门，h_t 为当前隐藏状态。BiLSTM 由前向 LSTM "
                     "与后向 LSTM 拼接而成，输出 h_t = [→h_t; ←h_t]。",
         "flowchart": "输出门 × tanh(细胞状态) → 当前隐藏状态 h_t"},
        {"no": "(7)", "algorithm": "算法 3：激光雷达点云球坐标转换",
         "expression": "c = arcsin(z / sqrt(x²+y²+z²));  r = arctan2(y, x)",
         "variables": "(x, y, z) 为激光雷达坐标系下的 3D 点坐标，(c, r) "
                     "为激光雷达图像中的位置坐标（垂直角、水平角）。",
         "flowchart": "3D 点云输入 → 球坐标转换 → 归一化 → 畸变校正 → "
                     "相机内参投影 → 形成稀疏深度图"},
        {"no": "(8)", "algorithm": "算法 3：相机畸变校正",
         "expression": "f(r) = 1 + k₁r² + k₂r⁴ + k₃r⁶",
         "variables": "r 为归一化半径，k₁, k₂, k₃ 为径向畸变参数。",
         "flowchart": "归一化坐标 → 计算半径 r → 畸变函数 f(r) → 输出校正坐标"},
        {"no": "(9)", "algorithm": "算法 3：相机内参投影",
         "expression": "P_i = K · m_d,  K = [[f_x,0,o_x],[0,f_y,o_y],[0,0,1]]",
         "variables": "K 为相机内参矩阵，f_x, f_y 为相机沿 x、y 轴的焦距，"
                     "(o_x, o_y) 为相机光学中心，m_d 为畸变校正后的归一化坐标。",
         "flowchart": "畸变校正坐标 → 内参矩阵 K 乘 → 像素坐标 P_i"},
        {"no": "(10)", "algorithm": "算法 3：邻域均值插值（稀疏深度图稠密化）",
         "expression": "d* = (1/N) · Σ_{p∈N} d_p",
         "variables": "N 为邻域点集，d_p 为邻域点的深度值，N 为邻域点数。",
         "flowchart": "稀疏深度图 → 邻域搜索 → 均值插值 → 稠密深度图 → "
                     "与 RGB 图像融合"},
        {"no": "(11)", "algorithm": "算法 4：FCDIoU 改进损失函数",
         "expression": "L_FCDIoU = 1 - IoU + ρ²(b, b^gt)/c² + d_cen²(b, b^gt)/c²",
         "variables": "b, b^gt 分别为预测框和真实框，ρ(b, b^gt) 为两框中心"
                     "点欧氏距离，d_cen 为两框质心距离，c 为两框最小外接矩形"
                     "的对角线长度。FCDIoU 相比 MPDIoU 在小目标场景下收敛"
                     "速度提升 1.8×。",
         "flowchart": "预测框 + 真实框 → 计算 IoU → 计算中心距离 ρ → 计算"
                     "质心距离 d_cen → 计算 c → 组合得 L_FCDIoU"},
    ],
    "methods": [
        "多模态融合技术下的无人机信息检测：无人机搭载摄像头（捕获高分辨率"
        "图像和视频）、麦克风（捕获声音信号）、激光雷达 LiDAR（生成高精度"
        "3D 点云数据）。多模态数据处理的第一步是同步和校准这些传感器数据，"
        "以确保数据在时间和空间上的一致性。多模态 3D 目标检测架构由两个"
        "主要分支组成：Vision-Language 分支和 Audio-Language 分支。",
        "轻量级重参数化卷积神经网络小模型设计：本项目提出 YOLO-LRP 算法，"
        "基于 YOLOv10 改进，引入 MobileOne 骨干网络、OREPA 在线卷积重"
        "参数化、GhostConv 模块、SE 注意力机制，实现全局轻量化处理。",
        "基于 CUDA GPU 处理器的异构并行运算：本研究通过可扩展的并行平台，"
        "在 GPU CUDA 中实现大规模并行训练，实现目标跟踪的加速计算。初期"
        "测试平台：CPU Intel Core i9-14900K（四核 2.60GHz，64G 内存），"
        "GPU Nvidia GeForce RTX 4090（17408 着色单元，2055MHz）。",
        "实验法：设计 CNN 模型架构，在数据集上进行训练/验证/测试，与 "
        "SVM、决策树等基线模型对比。",
        "实地测试法：将训练好的模型部署在合作消防中队，收集 3 个月实际"
        "运行数据评估泛化能力。",
    ],
    "feasibility_analysis": [
        "研究方案可行性：本项目的技术路线和研究方案是在总结项目组以及"
        "项目申请人过去的研究工作和仔细分析国内外研究现状和发展趋势的"
        "基础上提出的。在解决移动目标实时跟踪的问题上采用当前比较热门"
        "的多模态融合技术展开，并对其进行改进和优化。通过以国内外同行"
        "提出的理论和技术贡献为基础，基于大语言模型的基本框架的总体"
        "研究思路，进行展开和改进，因此本项目的研究方案具有科学可行性。",
        "工作基础可行性：项目组以及项目申请人近年来对目标跟踪、深度学习"
        "以及 GPU 并行运算等内容进行了研究，取得了较好的成果，为本项目的"
        "开展奠定了坚实的基础。在目标跟踪领域，项目申请人前期通过对多目标"
        "跟踪与匹配算法进行深入研究，相关成果已发表在 Computer Communications "
        "等学术期刊。在 GPU 并行运算的研究方面，项目组从 2020 年起就将 "
        "GPU CUDA 并行运算方法应用于各种机器学习和模式识别应用中。",
        "研究条件可行性：申请人及其所在的智能感知实验室，依托自动化工程"
        "技术研究中心、先进制造研究所、计算智能与信息处理研究所等。"
        "实验室配备 GPU 服务器（RTX 4090 ×4）、红外热成像仪 FLIR T630、"
        "激光雷达 Velodyne VLP-16、大疆经纬 M300 RTK 无人机等设备，"
        "可满足本项目需求。",
    ],
    "data_source": "数据来源：与 XX 消防支队签署数据合作协议，获取 2023-2024 "
                   "年 8 起真实火灾场景数据，含火场目标样本 5000+，覆盖 8 类"
                   "典型火场目标。实验设备：红外热成像仪 FLIR T630、激光雷达 "
                   "Velodyne VLP-16、大疆经纬 M300 RTK 无人机、GPU 服务器"
                   "（RTX 4090 ×4）。软件工具：Python 3.10、PyTorch 2.1、"
                   "OpenCV 4.8、CUDA 12.1、ROS 2 Humble。",
    "economic_benefits": [
        {"indicator": "单次城市火灾平均经济损失", "baseline": "100%（基准）",
         "expected": "60%~80%", "improvement": "降低 20%~40%（即 1.25~1.67 倍减少）"},
        {"indicator": "年均可减少直接经济损失", "baseline": "约 50 亿元/年",
         "expected": "约 30~40 亿元/年", "improvement": "减少 10~20 亿元/年（即 1.25~1.67 倍）"},
        {"indicator": "被困人员识别召回率", "baseline": "≤70%",
         "expected": "≥92%", "improvement": "提升 ≥22 个百分点（即 1.31 倍）"},
        {"indicator": "三维火场重构关键尺寸误差", "baseline": ">15cm",
         "expected": "≤5cm", "improvement": "精度提升 ≥3 倍"},
        {"indicator": "火势蔓延预测关键节点误差", "baseline": ">10m",
         "expected": "≤2m", "improvement": "精度提升 ≥5 倍"},
        {"indicator": "救援响应时间", "baseline": "100%（基准）",
         "expected": "30%~50%", "improvement": "缩短 50%~70%（即 2~3.33 倍）"},
        {"indicator": "消防员高危环境暴露频次", "baseline": "100%（基准）",
         "expected": "≤40%", "improvement": "降低 ≥60%（即 2.5 倍）"},
        {"indicator": "救援力量调度准确率", "baseline": "≤60%",
         "expected": "≥85%", "improvement": "提升 ≥25 个百分点（即 1.42 倍）"},
        {"indicator": "装备物资调配效率", "baseline": "100%（基准）",
         "expected": "130%~150%", "improvement": "提升 30%~50%（即 1.3~1.5 倍）"},
        {"indicator": "信息不对称导致的资源浪费率", "baseline": "约 25%~35%",
         "expected": "≤10%", "improvement": "降低 ≥15 个百分点（即 2.5~3.5 倍）"},
        {"indicator": "火势扩散范围控制", "baseline": "100%（基准）",
         "expected": "60%~70%", "improvement": "控制范围缩小 30%~40%（即 1.43~1.67 倍）"},
        {"indicator": "危险品识别与预警响应时间", "baseline": ">5 分钟",
         "expected": "≤1 分钟", "improvement": "缩短 ≥80%（即 5 倍）"},
    ],
    "project_schedule": [
        {"phase": "第一阶段：调研与方案设计", "time": "2025.04-2025.09（6 个月）",
         "tasks": ["针对关键科学问题和技术需求，完成项目总体研究方案设计",
                   "国内外研究现状进一步调研，搜集整理相关技术资料",
                   "研究大语言模型和多模态融合技术",
                   "调研 YOLO 系列神经网络算法",
                   "对部分研究成果进行论文发表"],
         "output": "调研报告 1 份、研究方案 1 份、文献综述 1 篇"},
        {"phase": "第二阶段：模型与算法开发", "time": "2025.10-2026.06（9 个月）",
         "tasks": ["搭建多模态融合网络结构模型",
                   "对卷积神经网络的主干网络、特征融合网络进行研究与改进",
                   "完成 YOLO-LRP 改进算法实现与训练",
                   "研究协同感知并行可扩展性模型",
                   "完成多模态融合算法验证"],
         "output": "多模态融合模型 1 套、YOLO-LRP 算法 1 套、模型权重 1 套、SCI 论文 1 篇"},
        {"phase": "第三阶段：系统集成与火场验证", "time": "2026.07-2027.01（7 个月）",
         "tasks": ["搭建无人机目标跟踪的地面控制平台",
                   "完成整个跟踪系统的代码编写及调试",
                   "在 GPU 并行系统进行完善训练",
                   "与合作消防中队开展火场场景测试",
                   "收集真实场景数据评估泛化能力"],
         "output": "原型系统 1 套、测试报告 1 份、发明专利申请 1 项、中文核心论文 1 篇"},
        {"phase": "第四阶段：总结与结题", "time": "2027.02-2027.03（2 个月）",
         "tasks": ["整理实验数据，撰写结题报告",
                   "论文返修与发表",
                   "软件著作权申请",
                   "项目总结与成果展示",
                   "培养学生毕业答辩"],
         "output": "结题报告 1 份、软件著作权 1 项、毕业论文 4 篇"},
    ],
    "expected_outcomes": [
        "SCI 二区期刊论文 1 篇（拟投 IEEE Transactions on Geoscience and Remote Sensing）",
        "中文核心期刊论文 2 篇（拟投《计算机辅助设计与图形学学报》《计算机应用》）",
        "发明专利申请 2 项（与合作企业共同申请）",
        "原型系统 1 套（部署在 XX 消防中队试运行 3 个月）",
        "火场样本数据集 1 套（5000+ 样本，8 类目标）",
        "软件著作权 1 项（《多模态融合无人机消防目标检测系统 V1.0》）",
    ],
    "budget_items": [
        {"item": "资料费", "amount": "1280",
         "basis": "图书 20 本 × 50 元 + 数据库订阅 280 元"},
        {"item": "调研差旅费", "amount": "4500",
         "basis": "3 次实地调研 × 1500 元（含交通、住宿）"},
        {"item": "实验材料费", "amount": "9320",
         "basis": "GPU 服务器租赁 1200 元/月 × 7 月 + 红外相机耗材 320 元 + 3D 打印耗材 600 元"},
        {"item": "会议费", "amount": "2400",
         "basis": "参加 ICASSP 国际会议 1 次 + 全国无人机大会 1 次"},
        {"item": "印刷复印", "amount": "500",
         "basis": "论文版面费 300 元 + 报告印刷 200 元"},
    ],
    "team_foundation": "团队 4 名成员已修读《机器学习》《数字图像处理》"
                       "《自动控制原理》《计算机视觉》等核心课程，3 人有"
                       "校级大创参与经验。负责人张三掌握 Python/PyTorch/CUDA，"
                       "曾参与校级『无人机目标检测』项目并获优秀。",
    "advisor_foundation": "指导教师李教授主持国家自然科学基金 1 项（62373011），"
                          "近 3 年发表 SCI 论文 8 篇（含 IEEE TGRS 2 篇、"
                          "Neurocomputing 3 篇），研究方向为多模态融合与"
                          "无人机智能感知，与本项目高度契合。",
    "lab_condition": "自动化工程学院智能感知实验室配备 GPU 服务器"
                     "（RTX 4090 ×4）、红外热成像仪 FLIR T630、激光雷达 "
                     "Velodyne VLP-16、大疆经纬 M300 RTK 无人机、3D 打印机，"
                     "可满足本项目研发与测试需求。",
    "budget_total": "18000",
    "include_school_approval": False,
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="大创-创新训练项目申报书 docx 生成器 v3.0",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n"
            "JSON 字段定义详见 SKILL.md 第十一章。"
            "v3.0 新增 8 字段：policy_citations / scientific_challenges / "
            "literature_review / algorithm_comparison / tech_roadmap / "
            "formulas / economic_benefits / project_schedule。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置 v3.0 示例数据（消防无人机多模态融合主题）生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
