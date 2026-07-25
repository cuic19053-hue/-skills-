#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
国家公派留学申请书（CSC Scholarship Application）docx 生成器 v1.0.0

A4 / 页边距 2.54cm/2.5cm / 标题黑体二号居中 / 正文宋体小四 1.5 倍行距首行缩进 2 字符 /
表格宋体五号居中 / 此致敬礼 / 落款右对齐。

CSC 4 类项目：phd（攻读博士）/ joint_phd（联合培养博士）/ master（攻读硕士）/ visiting_scholar（访问学者）
3 档字数版本：2500 / 3000 / 4000
6 段结构：个人基本情况 + 留学动机（4 维度）+ 留学单位选择 + 学习计划（5 要素）+ 归国计划 + 经济保障
依据：留金发〔2007〕3010 号 / 留金发〔2020〕112 号

使用：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx
    python build.py --demo --csc-type phd --out demo_phd.docx

JSON 字段详见 SKILL.md §9 JSON Schema。
"""

import argparse
import json
import os
import sys
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# ============================================================
# CSC 项目类别与字数版本配置
# ============================================================

CSC_TYPE_PHD = "phd"
CSC_TYPE_JOINT_PHD = "joint_phd"
CSC_TYPE_MASTER = "master"
CSC_TYPE_VISITING_SCHOLAR = "visiting_scholar"

CSC_TYPE_LABEL = {
    CSC_TYPE_PHD: "攻读博士学位",
    CSC_TYPE_JOINT_PHD: "联合培养博士研究生",
    CSC_TYPE_MASTER: "攻读硕士学位",
    CSC_TYPE_VISITING_SCHOLAR: "访问学者",
}

CSC_TYPE_PERIOD_LIMIT = {
    CSC_TYPE_PHD: 48,
    CSC_TYPE_JOINT_PHD: 24,
    CSC_TYPE_MASTER: 24,
    CSC_TYPE_VISITING_SCHOLAR: 12,
}

CSC_TYPE_DEFAULT_WORD_COUNT = {
    CSC_TYPE_PHD: "4000",
    CSC_TYPE_JOINT_PHD: "3000",
    CSC_TYPE_MASTER: "3000",
    CSC_TYPE_VISITING_SCHOLAR: "2500",
}

# 3 档字数版本配置（ranges 为 6 段字数分配）
WORD_COUNT_CONFIG: Dict[str, Dict[str, Any]] = {
    "2500": {"total": 2500, "ranges": [200, 600, 500, 700, 350, 150],
             "tolerance": 150, "target": "访问学者"},
    "3000": {"total": 3000, "ranges": [250, 750, 600, 900, 350, 150],
             "tolerance": 200, "target": "联合培养博士/攻读硕士"},
    "4000": {"total": 4000, "ranges": [300, 1000, 800, 1200, 500, 200],
             "tolerance": 300, "target": "攻读博士"},
}

# CSC 外语门槛
LANGUAGE_THRESHOLD = {
    "IELTS": 6.5, "TOEFL": 95, "PETS5": 55,
    "WSK": "合格", "培训部合格证": "合格",
}

# 国家留学基金委按国家/城市分档生活费（美元/月，参考 2024 年度）
LIVING_ALLOWANCE_BY_COUNTRY = {
    "美国一类": 1900, "美国二类": 1700, "英国": 1200,
    "德国": 1200, "法国": 1300, "日本": 145000,
    "澳大利亚": 1900, "加拿大": 1700,
}

# 禁用句（详见 SKILL.md §15.1）
FORBIDDEN_PHRASES = [
    "师资雄厚", "历史悠久", "从小就对",
    "填补国内空白", "国际首创", "世界领先", "独家首创", "颠覆性突破",
    "回国发展", "报效祖国", "回国效力", "为祖国做贡献", "回国创业",
    "发表论文若干", "力争取得突破", "争取取得成果",
    "永久居留", "绿卡", "移民", "定居", "入籍",
]

# 研究方向 → 国家战略政策映射
POLICY_MAP = {
    "新能源/储能": "《“十四五”新型储能发展实施方案》（发改能源〔2022〕209 号）",
    "新能源汽车": "《新能源汽车产业发展规划（2021—2035 年）》（国办发〔2020〕39 号）",
    "半导体/集成电路": "《新时期促进集成电路产业和软件产业高质量发展的若干政策》"
                       "（国发〔2020〕8 号）",
    "人工智能": "《新一代人工智能发展规划》（国发〔2017〕35 号）",
    "生物医药": "《“十四五”医药工业发展规划》（工信部联规〔2021〕208 号）",
    "双碳": "《2030 年前碳达峰行动方案》（国发〔2021〕23 号）",
}


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size: Pt = SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size: Pt = SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(doc, text, font_name=FONT_SONG,
                              font_size=SIZE_XIAO_SI, bold=False,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT,
                              first_line_indent=True, line_spacing=1.5,
                              space_before=0, space_after=0):
    """添加带格式段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号居中"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_section_heading(doc, text: str):
    """小节标题（一、二、三…）：黑体小四加粗"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
        line_spacing=1.5, space_before=6, space_after=3)


def add_subsection_heading(doc, text: str):
    """子节标题（（一）（二）…）：宋体小四加粗"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
        line_spacing=1.5, space_before=3, space_after=3)


def add_cizhi_paragraph(doc, text: str = "此致"):
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=False, line_spacing=1.5)


def add_table_from_data(doc, headers, rows, col_widths=None, caption=""):
    """从数据创建表格，自动应用规范格式"""
    if caption:
        add_paragraph_with_format(
            doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, font_size=SIZE_WU, bold=True)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def setup_page(doc):
    """A4 页面与页边距"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码"""
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


# ============================================================
# CSC 辅助函数
# ============================================================

def compute_study_months(study_start: str, study_end: str) -> int:
    """根据留学起止时间（YYYY-MM-DD 或 YYYY-MM）计算月数"""
    if not study_start or not study_end:
        return 0
    try:
        s = study_start.split("-")
        e = study_end.split("-")
        months = (int(e[0]) - int(s[0])) * 12 + (int(e[1]) - int(s[1])) + 1
        return max(months, 0)
    except (ValueError, IndexError):
        return 0


def check_language_threshold(test_type: str, score: str) -> Tuple[bool, str]:
    """检查外语成绩是否达 CSC 门槛"""
    if not test_type or not score:
        return False, "缺少外语成绩"
    threshold = LANGUAGE_THRESHOLD.get(test_type)
    if threshold is None:
        return False, f"未知外语考试类型 '{test_type}'"
    if threshold == "合格":
        if score in ("合格", "通过", "合格证"):
            return True, f"{test_type} {score}，达 CSC 门槛"
        return False, f"{test_type} 成绩 '{score}' 不达 CSC 门槛（须合格）"
    try:
        if float(score) >= threshold:
            return True, f"{test_type} {score}，达 CSC 门槛（{threshold}）"
        return False, f"{test_type} {score} 不达 CSC 门槛（{threshold}）"
    except (ValueError, TypeError):
        return False, f"{test_type} 成绩 '{score}' 格式异常"


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """国家公派留学申请书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []
        self.word_count_version: str = "3000"
        self.csc_type: str = CSC_TYPE_JOINT_PHD

    def _get(self, *keys, default=""):
        """安全取嵌套字段"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, *keys):
        """安全取列表字段"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return []
            cur = cur.get(k, [])
        if isinstance(cur, list):
            return cur
        return [cur] if isinstance(cur, str) else []

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_heading(self, text):
        return add_section_heading(self.doc, text)

    def add_subheading(self, text):
        return add_subsection_heading(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, caption=""):
        return add_table_from_data(self.doc, headers, rows, col_widths, caption)

    # --- 标题与称呼 ---

    def _add_title(self):
        title = self._get("title", default="国家公派留学申请书")
        add_title(self.doc, title)

    def _add_salutation(self):
        salutation = self._get(
            "salutation",
            default="尊敬的国家留学基金管理委员会评审专家：")
        add_salutation_paragraph(self.doc, salutation)

    # --- 段一 个人基本情况 ---

    def _add_opening(self):
        """段一 个人基本情况"""
        opening = self._get("opening", default="")
        if opening:
            self.add_para(opening)
            return

        name = self._get("name")
        gender = self._get("gender")
        birth_date = self._get("birth_date")
        ethnicity = self._get("ethnicity", default="汉族")
        political_status = self._get("political_status", default="共青团员")
        marital_status = self._get("marital_status", default="未婚")
        current_unit = self._get("current_unit")
        research_direction = self._get("research_direction")
        csc_type_label = self._get(
            "csc_type_label",
            default=CSC_TYPE_LABEL.get(self.csc_type, "联合培养博士研究生"))
        target_country = self._get("target_country")
        target_university_cn = self._get("target_university_cn")
        study_start = self._get("study_start")
        study_end = self._get("study_end")
        study_period_months = self._get("study_period_months", default=0)
        if not study_period_months and study_start and study_end:
            study_period_months = compute_study_months(study_start, study_end)
        language_test_type = self._get("language_test_type")
        language_score = self._get("language_score")

        parts = []
        personal = []
        if name:
            personal.append(f"本人{name}")
        if gender:
            personal.append(gender)
        if birth_date:
            personal.append(f"{birth_date} 出生")
        if ethnicity:
            personal.append(ethnicity)
        if political_status:
            personal.append(political_status)
        if marital_status:
            personal.append(marital_status)
        if personal:
            parts.append("，".join(personal) + "。")
        if current_unit:
            parts.append(f"现为{current_unit}。")
        if research_direction:
            parts.append(f"研究方向为{research_direction}。")

        # 学历背景
        bachelor_school = self._get("education_background", "bachelor", "school")
        master_school = self._get("education_background", "master", "school")
        phd_school = self._get("education_background", "phd", "school")
        edu = []
        if bachelor_school:
            edu.append(f"本科毕业于{bachelor_school}")
        if master_school:
            edu.append(f"硕士毕业于{master_school}")
        if phd_school:
            edu.append(f"现为{phd_school}博士研究生")
        if edu:
            parts.append("，".join(edu) + "。")

        # 学术成果简述
        pubs = self._get_list("academic_achievements", "publications")
        first_author = sum(
            1 for p in pubs
            if isinstance(p, dict) and "第一作者" in str(p.get("role", "")))
        if first_author:
            parts.append(f"博士在读期间以第一作者发表 SCI 论文 {first_author} 篇。")

        if language_test_type and language_score:
            parts.append(f"外语水平：{language_test_type} {language_score}。")

        apply = []
        if csc_type_label:
            apply.append(f"现申请国家公派留学'{csc_type_label}'项目")
        if target_country and target_university_cn:
            apply.append(f"赴{target_country}{target_university_cn}")
        if study_period_months:
            apply.append(f"联合培养/留学 {study_period_months} 个月")
        if apply:
            parts.append("，".join(apply) + "。")

        parts.append("现将本人情况汇报如下：")
        self.add_para("".join(parts))

    # --- 段二 留学动机（4 维度）---

    def _add_motivation(self):
        """段二 留学动机：学术追求 + 国家战略 + 学科差距 + 个人发展"""
        self.add_heading("二、留学动机")
        text = self._get("motivation_text", default="")
        if text:
            self.add_para(text)
            return

        motivation = self._get("motivation", default={})
        if not isinstance(motivation, dict):
            motivation = {}

        dimensions = [
            ("（一）学术追求", "academic_pursuit"),
            ("（二）国家战略", "national_strategy"),
            ("（三）学科差距", "discipline_gap"),
            ("（四）个人发展", "personal_development"),
        ]
        for heading, key in dimensions:
            content = motivation.get(key, "")
            if content:
                self.add_subheading(heading)
                self.add_para(content)

    # --- 段三 留学单位选择 ---

    def _add_university_choice(self):
        """段三 留学单位选择"""
        self.add_heading("三、留学单位选择")
        text = self._get("university_choice_text", default="")
        if text:
            self.add_para(text)
            return

        target_country = self._get("target_country")
        target_university_en = self._get("target_university_en")
        target_university_cn = self._get("target_university_cn")
        target_university_qs_rank = self._get("target_university_qs_rank")
        target_department_cn = self._get("target_department_cn")
        target_department_en = self._get("target_department_en")
        target_advisor_name = self._get("target_advisor_name")
        target_advisor_title = self._get("target_advisor_title")
        target_advisor_honor = self._get("target_advisor_honor")
        target_advisor_email = self._get("target_advisor_email")
        target_advisor_h_index = self._get("target_advisor_h_index")
        target_advisor_research = self._get("target_advisor_research")
        target_advisor_recent_papers = self._get(
            "target_advisor_recent_papers", default=0)
        target_advisor_nature_science_papers = self._get(
            "target_advisor_nature_science_papers", default=0)
        research_direction = self._get("research_direction")
        domestic_advisor_name = self._get("domestic_advisor_name")

        # 院校基本信息
        univ_parts = []
        if target_country and target_university_cn:
            univ_parts.append(
                f"本人拟赴{target_country}{target_university_cn}联合培养/留学")
        if target_university_en:
            univ_parts.append(f"（英文：{target_university_en}）")
        if target_university_qs_rank:
            univ_parts.append(f"，QS 世界大学排名第 {target_university_qs_rank} 位")
        if target_department_cn:
            univ_parts.append(f"，目标院系为{target_department_cn}")
        if target_department_en:
            univ_parts.append(f"（{target_department_en}）")
        if univ_parts:
            self.add_para("".join(univ_parts) + "。")

        # 外方导师信息
        if target_advisor_name:
            advisor_parts = [f"外方导师为{target_advisor_name}"]
            if target_advisor_title:
                advisor_parts.append(target_advisor_title)
            if target_advisor_honor:
                advisor_parts.append(target_advisor_honor)
            self.add_para("（".join(advisor_parts[:1])
                          + "（" + "，".join(advisor_parts[1:]) + "）"
                          if len(advisor_parts) > 1
                          else advisor_parts[0] + "")
            if target_advisor_research:
                self.add_para(f"研究方向为{target_advisor_research}。")
            if target_advisor_h_index:
                self.add_para(f"H 指数 {target_advisor_h_index}。")
            if target_advisor_email:
                self.add_para(f"邮箱：{target_advisor_email}。")

        # 选择理由 + 匹配度量化
        match_parts = []
        if target_advisor_name and research_direction:
            match_parts.append(
                f"{target_advisor_name} 教授研究方向与本人博士课题"
                f"'{research_direction}'高度匹配")
        if target_advisor_recent_papers:
            match_parts.append(
                f"其课题组近 3 年发表论文 {target_advisor_recent_papers} 篇，"
                f"其中相当部分与本人课题直接相关")
        if match_parts:
            self.add_para("选择理由：" + "，".join(match_parts) + "。")

        # 国内外导师合作基础
        collab = self._get("collaboration_basis", default={})
        if isinstance(collab, dict):
            joint_papers = collab.get("joint_papers", [])
            mutual_visits = collab.get("mutual_visits", 0)
            email_exchanges = collab.get("email_exchanges", 0)
            if domestic_advisor_name and target_advisor_name:
                cb = [f"本人国内导师{domestic_advisor_name}教授与外方导师"
                      f"{target_advisor_name}教授已有合作基础"]
                sub = []
                if joint_papers:
                    sub.append(f"共同发表论文 {len(joint_papers)} 篇")
                if mutual_visits:
                    sub.append(f"互访 {mutual_visits} 次")
                if email_exchanges:
                    sub.append(f"邮件往来 {email_exchanges} 封")
                if sub:
                    cb.append("，".join(sub))
                self.add_para("".join(cb) + "。")

    # --- 段四 学习计划（5 要素）---

    def _add_study_plan(self):
        """段四 学习计划：研究方向/课程选修/导师对接/合作研究/预期成果"""
        self.add_heading("四、学习计划")
        text = self._get("study_plan_text", default="")
        if text:
            self.add_para(text)
            self._add_timeline_table()
            self._add_expected_outcomes_table()
            return

        plan = self._get("research_plan", default={})
        if not isinstance(plan, dict):
            plan = {}

        # 要素一：研究方向
        self.add_subheading("（一）研究方向")
        rd = plan.get("research_direction") or self._get("research_direction")
        if rd:
            self.add_para(f"本人留学期间研究方向为'{rd}'。")
        bg = plan.get("research_background", "")
        if bg:
            self.add_para(f"研究背景与意义：{bg}")
        objs = plan.get("research_objectives", [])
        if objs:
            self.add_para("研究目标：" + "；".join(
                [f"（{i+1}）{o}" for i, o in enumerate(objs)]) + "。")
        contents = plan.get("research_contents", [])
        if contents:
            self.add_para("研究内容：" + "；".join(
                [f"（{i+1}）{c}" for i, c in enumerate(contents)]) + "。")
        methods = plan.get("research_methods", "")
        if methods:
            self.add_para(f"研究方法：{methods}。")

        # 要素二：课程选修（访问学者可省略）
        if self.csc_type != CSC_TYPE_VISITING_SCHOLAR:
            courses = plan.get("courses", [])
            if courses:
                self.add_subheading("（二）课程选修")
                course_strs = []
                for c in courses:
                    if not isinstance(c, dict):
                        continue
                    cs = ""
                    if c.get("code"):
                        cs += f"{c['code']} "
                    if c.get("name"):
                        cs += f"《{c['name']}》"
                    if c.get("credits"):
                        cs += f"（{c['credits']} 学分）"
                    if c.get("reason"):
                        cs += f"，{c['reason']}"
                    if cs:
                        course_strs.append(cs)
                if course_strs:
                    self.add_para("本人拟选修以下课程："
                                  + "；".join(course_strs) + "。")
                    self.add_para("上述课程均以旁听形式参加，剩余时间用于科研工作。")

        # 要素三：导师对接
        self.add_subheading("（三）导师对接")
        advisor_collab = plan.get("advisor_collab", "")
        if advisor_collab:
            self.add_para(advisor_collab)
        else:
            target_advisor_name = self._get("target_advisor_name")
            target_advisor_research = self._get("target_advisor_research")
            target_advisor_h_index = self._get("target_advisor_h_index")
            research_direction = self._get("research_direction")
            domestic_advisor_name = self._get("domestic_advisor_name")
            if target_advisor_name:
                parts = [f"外方导师为{target_advisor_name}"]
                if target_advisor_research:
                    parts.append(f"研究方向为{target_advisor_research}")
                if target_advisor_h_index:
                    parts.append(f"H 指数 {target_advisor_h_index}")
                self.add_para("，".join(parts) + "。")
            if research_direction and target_advisor_name:
                self.add_para(
                    f"{target_advisor_name} 教授研究方向与本人博士课题"
                    f"'{research_direction}'高度匹配。")
            if domestic_advisor_name and target_advisor_name:
                self.add_para(
                    f"本人国内导师{domestic_advisor_name}教授与外方导师"
                    f"{target_advisor_name}教授已有合作基础"
                    "（详见三、留学单位选择部分）。")
            self.add_para("本人拟每周参加课题组组会 1 次，"
                          "每两周与外方导师一对一讨论 1 次。")

        # 要素四：合作研究
        self.add_subheading("（四）合作研究")
        collab_research = plan.get("collab_research", {})
        if not isinstance(collab_research, dict):
            collab_research = {}
        my_role = collab_research.get("my_role", "")
        advisor_role = collab_research.get("advisor_role", "")
        sustainability = collab_research.get("sustainability", "")
        if my_role:
            self.add_para(f"本人负责：{my_role}。")
        if advisor_role:
            self.add_para(f"外方导师课题组负责：{advisor_role}。")
        if sustainability:
            self.add_para(f"合作研究的可持续性：{sustainability}")
        if not (my_role or advisor_role):
            rd = self._get("research_direction")
            if rd:
                self.add_para(
                    f"本人拟与外方课题组开展'{rd}'合作研究，"
                    "本人负责样品制备与数据分析，"
                    "外方课题组负责设备使用与方法指导。")

        # 要素五：预期成果
        self.add_subheading("（五）预期成果")
        outcomes = plan.get("expected_outcomes", {})
        if not isinstance(outcomes, dict):
            outcomes = {}
        papers = outcomes.get("papers", 0)
        first_author = outcomes.get("first_author_papers", 0)
        target_journals = outcomes.get("target_journals", [])
        patents_pct = outcomes.get("patents_pct", 0)
        patents_domestic = outcomes.get("patents_domestic", 0)
        conferences = outcomes.get("conferences", 0)
        oral = outcomes.get("oral_presentations", 0)
        if papers or first_author:
            ps = (f"以第一作者发表 SCI 论文 {first_author} 篇"
                  if first_author else f"发表论文 {papers} 篇")
            if target_journals:
                ps += "，目标期刊为 " + "、".join(target_journals)
            self.add_para(f"学术成果：{ps}。")
        if patents_pct or patents_domestic:
            ps = []
            if patents_pct:
                ps.append(f"国际 PCT 专利 {patents_pct} 项")
            if patents_domestic:
                ps.append(f"国内发明专利 {patents_domestic} 项")
            self.add_para("专利成果：" + "、".join(ps) + "。")
        if conferences:
            cs = f"参加国际会议 {conferences} 次"
            if oral:
                cs += f"（其中口头报告 {oral} 次）"
            self.add_para(f"国际会议：{cs}。")
        self.add_para("技术报告：向 CSC 提交季度进展报告与年度总结报告，"
                      "向国内导师提交月度报告。")
        if self.csc_type == CSC_TYPE_PHD:
            univ = self._get("target_university_cn")
            if univ:
                self.add_para(f"学位成果：取得{univ}博士学位。")
        elif self.csc_type == CSC_TYPE_MASTER:
            univ = self._get("target_university_cn")
            if univ:
                self.add_para(f"学位成果：取得{univ}硕士学位。")

        # 学习计划时间表 + 预期成果表
        self._add_timeline_table()
        self._add_expected_outcomes_table()

    def _add_timeline_table(self):
        """学习计划时间表"""
        timeline = self._get_list("research_plan", "timeline")
        if not timeline:
            timeline = self._get_list("timeline")
        if not timeline:
            return
        headers = ["时间段", "任务安排"]
        rows = []
        for item in timeline:
            if isinstance(item, dict):
                period = str(item.get("period", item.get("time", "")))
                task = str(item.get("task", item.get("content", "")))
                if period or task:
                    rows.append([period, task])
        if rows:
            self.add_table(headers, rows, col_widths=[5.0, 9.5],
                           caption="学习计划时间表：")

    def _add_expected_outcomes_table(self):
        """预期成果量化表"""
        outcomes = self._get("research_plan", "expected_outcomes", default={})
        if not isinstance(outcomes, dict):
            return
        rows = []
        first_author = outcomes.get("first_author_papers", 0)
        papers = outcomes.get("papers", 0)
        if papers or first_author:
            rows.append(["SCI 论文",
                         f"{first_author} 篇（一作）",
                         "目标期刊 IF ≥ 15"])
        patents_pct = outcomes.get("patents_pct", 0)
        patents_domestic = outcomes.get("patents_domestic", 0)
        if patents_pct or patents_domestic:
            rows.append(["专利",
                         f"PCT {patents_pct} + 国内 {patents_domestic}", ""])
        conferences = outcomes.get("conferences", 0)
        oral = outcomes.get("oral_presentations", 0)
        if conferences:
            oral_str = f"（口头 {oral}）" if oral else ""
            rows.append(["国际会议", f"{conferences} 次{oral_str}", ""])
        if not rows:
            return
        self.add_table(["成果类别", "数量", "备注"], rows,
                       col_widths=[4.0, 5.0, 5.5],
                       caption="预期成果量化表：")

    # --- 段五 归国计划 ---

    def _add_return_plan(self):
        """段五 归国计划"""
        self.add_heading("五、归国计划")
        text = self._get("return_plan_text", default="")
        if text:
            self.add_para(text)
            return

        rp = self._get("return_plan", default={})
        if not isinstance(rp, dict):
            rp = {}
        return_date = rp.get("return_date", "")
        target_unit = rp.get("target_unit") or self._get("return_target_unit")
        position = rp.get("position") or self._get("return_position")
        research_plan = rp.get("research_plan", "")
        industry_partner = rp.get("industry_partner", "")
        service_years = rp.get("service_years", 2)
        long_term_plan = rp.get("long_term_plan", "")
        return_letter_signed = rp.get("return_letter_signed", False)

        parts = []
        if return_date:
            parts.append(f"本人承诺于{return_date}学成按期回国")
        if target_unit:
            signed_str = "（已签订回国工作意向书，见附件）" \
                if return_letter_signed else ""
            parts.append(f"已与{target_unit}达成回国意向{signed_str}")
        if position:
            parts.append(f"拟任{position}岗位")
        if parts:
            self.add_para("，".join(parts) + "。")
        if research_plan:
            self.add_para(f"回国后研究设想：{research_plan}。")
        if industry_partner:
            self.add_para(f"联合产业伙伴：{industry_partner}。")
        if service_years:
            self.add_para(
                "依据《关于明确国家公派留学人员回国服务期和违约追偿相关事宜的通知》"
                f"（留金发〔2020〕112 号），"
                f"本人承诺学成回国后服务 {service_years} 年。")
        if long_term_plan:
            self.add_para(f"长期发展规划：{long_term_plan}")
        self._add_return_table(rp)

    def _add_return_table(self, rp: Dict[str, Any]):
        """归国去向表"""
        rows = []
        if rp.get("return_date"):
            rows.append(["回国时间", str(rp.get("return_date"))])
        target_unit = rp.get("target_unit") or self._get("return_target_unit")
        if target_unit:
            rows.append(["回国单位", str(target_unit)])
        position = rp.get("position") or self._get("return_position")
        if position:
            rows.append(["回国岗位", str(position)])
        if rp.get("research_plan"):
            rows.append(["研究方向", str(rp.get("research_plan"))])
        rows.append(["服务期", f"{rp.get('service_years', 2)} 年"])
        if not rows:
            return
        self.add_table(["项目", "内容"], rows,
                       col_widths=[3.5, 11.0],
                       caption="归国去向表：")

    # --- 段六 经济保障 ---

    def _add_financial(self):
        """段六 经济保障"""
        self.add_heading("六、经济保障")
        text = self._get("financial_text", default="")
        if text:
            self.add_para(text)
            return

        fa = self._get("financial_arrangement", default={})
        if not isinstance(fa, dict):
            fa = {}
        csc_allowance = fa.get("csc_allowance_monthly_usd", 0)
        tuition_waived = fa.get("tuition_waived", False)
        tuition_source = fa.get("tuition_waiver_source", "")
        advisor_stipend = fa.get("advisor_research_stipend_monthly_usd", 0)
        spouse_income = fa.get("spouse_income_monthly_cny", 0)
        emergency_fund = fa.get("emergency_fund_cny", 0)
        insurance = fa.get("insurance_arrangement", "")

        parts = []
        csc_str = "本人留学期间主要经济来源为国家留学基金委提供的奖学金生活费"
        if csc_allowance:
            csc_str += f"（按 {csc_allowance} 美元/月发放）"
        csc_str += "及一次往返国际旅费。"
        parts.append(csc_str)
        if tuition_waived:
            ts = "外方院校已承诺免除本人留学期间学费"
            if tuition_source:
                ts += f"（{tuition_source}）"
            parts.append(ts + "。")
        if advisor_stipend:
            parts.append(f"外方导师课题组提供 {advisor_stipend} 美元/月"
                         "助研津贴用于实验耗材补充。")
        if spouse_income:
            parts.append(f"本人家庭配偶月收入约 {spouse_income} 元人民币。")
        if emergency_fund:
            parts.append(f"已在银行储备 {emergency_fund} 元人民币应急金。")
        if insurance:
            parts.append(insurance + "。")
        else:
            parts.append("CSC 统一购买的国际医疗保险已覆盖留学期间医疗费用。")
        self.add_para("".join(parts))

    # --- 此致敬礼与落款 ---

    def _add_closing(self):
        add_cizhi_paragraph(self.doc, "此致")
        add_jingli_paragraph(self.doc, "敬礼！")
        name = self._get("name", default="申请人")
        apply_date = self._get("apply_date")
        if not apply_date:
            apply_date = datetime.now().strftime("%Y 年 %m 月 %d 日")
        add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        add_right_aligned_paragraph(self.doc, str(apply_date))

    # --- 基本信息表 ---

    def _add_basic_info_table(self):
        """基本信息表（4 列：项目/内容/项目/内容）"""
        rows_data = []
        fields = [
            (("姓名", "name"), ("性别", "gender")),
            (("出生日期", "birth_date"), ("民族", "ethnicity")),
            (("政治面貌", "political_status"), ("婚姻状况", "marital_status")),
            (("现单位", "current_unit"), ("联系电话", "phone")),
            (("电子邮箱", "email"),
             ("申请项目", "csc_type_label")),
            (("留学国家", "target_country"),
             ("留学院校", "target_university_cn")),
            (("外方导师", "target_advisor_name"),
             ("国内导师", "domestic_advisor_name")),
            (("推荐单位", "recommend_unit"),
             ("研究方向", "research_direction")),
        ]
        for (l_label, l_key), (r_label, r_key) in fields:
            l_val = self._get(l_key)
            r_val = self._get(r_key)
            if l_key == "csc_type_label" and not r_val:
                r_val = CSC_TYPE_LABEL.get(self.csc_type, "")
            if l_val or r_val:
                rows_data.append([l_label, str(l_val), r_label, str(r_val)])

        # 留学起止行（特殊处理）
        study_start = self._get("study_start")
        study_end = self._get("study_end")
        if study_start:
            months = self._get("study_period_months", default=0)
            if not months:
                months = compute_study_months(study_start, study_end)
            rows_data.append(["留学起止", f"{study_start} ~ {study_end}",
                              "留学月数", f"{months} 个月"])
        # 外语成绩行
        ltt = self._get("language_test_type")
        ls = self._get("language_score")
        if ltt:
            rows_data.append(["外语考试", str(ltt), "外语成绩", str(ls)])

        if not rows_data:
            return
        self.add_table(["项目", "内容", "项目", "内容"], rows_data,
                       col_widths=[2.5, 5.0, 2.5, 4.5],
                       caption="申请人基本信息表：")

    # --- 数据校验 ---

    def validate(self) -> List[str]:
        """校验数据完整性"""
        warnings: List[str] = []
        required = [
            ("name", "姓名"), ("gender", "性别"), ("birth_date", "出生日期"),
            ("current_unit", "现工作学习单位"), ("csc_type", "项目类别"),
            ("target_country", "目标留学国家"),
            ("target_university_cn", "目标院校（中文）"),
            ("target_advisor_name", "外方导师姓名"),
            ("study_start", "留学开始时间"), ("study_end", "留学结束时间"),
            ("research_direction", "研究方向"),
            ("recommend_unit", "国内推荐单位"),
            ("domestic_advisor_name", "国内导师"),
        ]
        for field, label in required:
            if not self._get(field):
                warnings.append(f"缺少必填字段 {field}（{label}）")

        if self.csc_type not in CSC_TYPE_LABEL:
            warnings.append(f"项目类别 '{self.csc_type}' 不在 4 类内"
                            f"（phd/joint_phd/master/visiting_scholar）")

        study_start = self._get("study_start")
        study_end = self._get("study_end")
        months = self._get("study_period_months", default=0)
        if not months and study_start and study_end:
            months = compute_study_months(study_start, study_end)
        if months and self.csc_type:
            limit = CSC_TYPE_PERIOD_LIMIT.get(self.csc_type, 0)
            if limit and months > limit:
                warnings.append(
                    f"留学月数 {months} 超过 {self.csc_type} 项目上限 "
                    f"{limit} 个月（详见 SKILL.md §6.9）")

        ltt = self._get("language_test_type")
        ls = self._get("language_score")
        if ltt and ls:
            passed, msg = check_language_threshold(ltt, ls)
            if not passed:
                warnings.append(f"外语成绩不达 CSC 门槛：{msg}")
        else:
            warnings.append("缺少外语水平证明（language_test_type/language_score）")

        if self.word_count_version not in WORD_COUNT_CONFIG:
            warnings.append(
                f"字数版本 '{self.word_count_version}' 不在 2500/3000/4000 内，"
                f"使用默认 3000")
            self.word_count_version = "3000"

        return_target = self._get("return_target_unit")
        rp = self._get("return_plan", default={})
        if isinstance(rp, dict):
            return_target = return_target or rp.get("target_unit")
        if not return_target:
            warnings.append("缺少归国去向单位（return_target_unit）"
                            "——归国计划须具体，详见 SKILL.md §6.2")

        # 禁用句检查
        full_text = json.dumps(self.data, ensure_ascii=False)
        for phrase in FORBIDDEN_PHRASES:
            if phrase in full_text:
                warnings.append(
                    f"数据含禁用句 '{phrase}'"
                    f"（详见 SKILL.md §6.1/§6.5/§15.1）")

        self.warnings = warnings
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings

    # --- 主入口 ---

    def build(self, data: Dict[str, Any], output_path: str) -> None:
        """构建申请书 docx"""
        self.data = data
        self.csc_type = data.get("csc_type", CSC_TYPE_JOINT_PHD)
        self.word_count_version = data.get(
            "word_count_version",
            CSC_TYPE_DEFAULT_WORD_COUNT.get(self.csc_type, "3000"))
        self.validate()

        self._add_title()
        self._add_salutation()
        self._add_basic_info_table()

        self.add_heading("一、个人基本情况")
        self._add_opening()
        self._add_motivation()
        self._add_university_choice()
        self._add_study_plan()
        self._add_return_plan()
        self._add_financial()
        self._add_closing()

        self.doc.save(output_path)
        print(f"✅ 国家公派留学申请书已生成：{output_path}")
        self._post_check_word_count()

    def _post_check_word_count(self):
        """构建后字数统计"""
        total_text = ""
        for para in self.doc.paragraphs:
            total_text += para.text
        chinese_chars = sum(1 for c in total_text if '\u4e00' <= c <= '\u9fff')
        total_chars = len(total_text)
        print(f"📊 全文字数（中文字符）：{chinese_chars}，总字符数：{total_chars}")

        config = WORD_COUNT_CONFIG.get(self.word_count_version)
        if not config:
            return
        target = config["total"]
        tolerance = config["tolerance"]
        lower, upper = target - tolerance, target + tolerance
        if chinese_chars < lower:
            print(f"⚠️ [字数] 全文仅 {chinese_chars} 字，"
                  f"建议 {self.word_count_version} 档"
                  f"（{target} 字，区间 {lower}~{upper}）", file=sys.stderr)
        elif chinese_chars > upper:
            print(f"⚠️ [字数] 全文 {chinese_chars} 字偏多，"
                  f"建议 {self.word_count_version} 档"
                  f"（{target} 字，区间 {lower}~{upper}）", file=sys.stderr)
        else:
            print(f"✅ [字数] 符合 {self.word_count_version} 档"
                  f"（{target} 字 ±{tolerance}）")


# ============================================================
# 默认示例数据
# ============================================================

DEFAULT_DATA = {
    # 基础信息
    "name": "王某某", "gender": "男", "birth_date": "1995-06-15",
    "nationality": "中国", "ethnicity": "汉族", "political_status": "中共党员",
    "marital_status": "已婚",
    "current_unit": "XX 大学材料科学与工程学院 2022 级博士研究生",
    "phone": "138XXXXXXXX", "email": "wang@example.edu.cn", "student_id": "2022B11050",
    # 申请项目
    "csc_type": "joint_phd", "csc_type_label": "联合培养博士研究生",
    "title": "国家公派留学申请书",
    # 目标留学单位
    "target_country": "美国",
    "target_university_en": "Massachusetts Institute of Technology",
    "target_university_cn": "麻省理工学院", "target_university_qs_rank": 1,
    "target_department_en": "Department of Materials Science and Engineering",
    "target_department_cn": "材料科学与工程系",
    # 外方导师
    "target_advisor_name": "Yet-Ming Chiang",
    "target_advisor_title": "Kyocera 讲席教授",
    "target_advisor_honor": "美国国家工程院院士",
    "target_advisor_email": "chiang@mit.edu", "target_advisor_h_index": 95,
    "target_advisor_recent_papers": 28, "target_advisor_nature_science_papers": 4,
    "target_advisor_research": "电化学储能与材料科学",
    # 留学时间
    "study_start": "2024-09-01", "study_end": "2026-08-31",
    "study_period_months": 24,
    # 研究方向与外语
    "research_direction": "硫碳复合正极的原位表征与界面调控",
    "research_field_category": "新能源/储能",
    "language_test_type": "IELTS", "language_score": "7.0",
    "language_test_date": "2023-03-15", "language_min_requirement": "6.5",
    # 国内推荐单位与导师
    "recommend_unit": "XX 大学", "domestic_advisor_name": "李某某",
    "domestic_advisor_title": "教授", "domestic_advisor_research": "锂硫电池正极材料",
    # 归国计划
    "return_target_unit": "XX 大学材料学院", "return_position": "副教授",
    "return_service_years": 2,

    # 教育背景
    "education_background": {
        "bachelor": {"school": "XX 大学", "major": "材料科学与工程",
                     "start": "2013-09", "end": "2017-06", "degree": "工学学士",
                     "gpa": "3.7/4.0", "rank": "5/120"},
        "master": {"school": "XX 大学", "major": "材料学",
                   "start": "2017-09", "end": "2020-06", "degree": "工学硕士",
                   "advisor": "李某某 教授", "gpa": "3.85/4.0", "rank": "2/45",
                   "thesis_title": "硫碳复合正极的制备与电化学性能研究"},
        "phd": {"school": "XX 大学", "major": "材料科学与工程",
                "start": "2020-09", "expected_end": "2024-06",
                "advisor": "李某某 教授", "research_direction": "锂硫电池正极材料",
                "gpa": "3.9/4.0", "rank": "1/30"}
    },
    # 学术成果
    "academic_achievements": {
        "publications": [
            {"authors": "Wang X, et al.", "title": "Sulfur-carbon composite cathode",
             "journal": "Advanced Energy Materials", "year": 2023, "if": 24.4, "role": "第一作者"},
            {"authors": "Wang X, et al.", "title": "In-situ XRD study of sulfur cathode",
             "journal": "Joule", "year": 2023, "if": 41.2, "role": "第一作者"},
            {"authors": "Wang X, et al.", "title": "Pore structure control of carbon host",
             "journal": "ACS Energy Letters", "year": 2024, "if": 22.0, "role": "第一作者"}
        ],
        "patents": [{"title": "一种硫碳复合正极制备方法", "patent_no": "ZL202210012345.6",
                     "year": 2023, "role": "第一发明人"}],
        "awards": [{"name": "国家奖学金", "year": 2022, "level": "国家级"},
                   {"name": "校长奖学金", "year": 2023, "level": "校级"}],
        "projects": [{"name": "国家自然科学基金面上项目", "role": "主要参与人",
                      "pi": "李某某 教授", "year": "2021-2024"}]
    },

    # 留学动机 4 维度
    "motivation": {
        "academic_pursuit": (
            "在学术追求层面，本人在国内已系统开展锂硫电池正极材料研究，"
            "博士在读期间以第一作者发表 SCI 一区论文 3 篇（影响因子累计 87.6），"
            "但仍面临硫正极容量衰减机理不清晰等技术瓶颈。"
            "MIT Chiang 课题组在锂硫电池原位表征技术上处于国际领先地位，"
            "2020~2024 年发表 Nature 主刊论文 2 篇、Joule 论文 4 篇，"
            "其首创的原位 X 射线衍射技术可对硫正极充放电过程进行原子尺度观测。"
            "本人拟在 Chiang 课题组学习该原位表征技术，"
            "预期将其应用至本人博士课题，提升对硫正极衰减机理的认识深度。"
        ),
        "national_strategy": (
            "在国家战略层面，本人研究方向与《“十四五”新型储能发展实施方案》"
            "（发改能源〔2022〕209 号）“突破高比能锂硫电池关键技术”的部署高度契合。"
            "国家发改委、能源局明确将锂硫电池列为“十四五”重点攻关方向，"
            "2023 年相关财政支持达 12 亿元。"
            "然而我国锂硫电池产业化与国际先进水平仍有差距：美国 MIT/OXIS、"
            "韩国三星 SDI 已实现 500 Wh/kg 中试，我国最高水平仅为 450 Wh/kg。"
            "本人学成回国后将联合宁德时代等国内龙头企业推动 500 Wh/kg 锂硫电池中试。"
        ),
        "discipline_gap": (
            "在学科差距层面，本人研究方向“硫碳复合正极”的国内外差距体现在三方面。"
            "设备方面：国内仅有 3 套同步辐射原位 XRD 装置，MIT 课题组自有 2 套并"
            "享有美国布鲁克海文国家实验室 NSLS-II 同步辐射光源优先使用权。"
            "方法方面：国内硫正极表征仍以离线 SEM/TEM 为主，"
            "MIT 课题组已建立原位同步辐射 + 原位电化学质谱联合表征平台。"
            "产出方面：2020~2024 年锂硫电池领域 Nature/Science 主刊论文 14 篇中，"
            "美国 8 篇、韩国 3 篇、德国 2 篇、中国 1 篇，差距显著。"
            "本人拟通过联合培养引入原位表征方法，缩小国内外方法差距。"
        ),
        "personal_development": (
            "在个人发展层面，本人在国内博士第 3 年已完成硫碳复合正极基础研究，"
            "正处于由基础研究向中试放大的关键转折期，此时赴 MIT 联合培养 24 个月恰逢其时。"
            "本人归国后拟任 XX 大学材料学院副教授，组建“先进电池材料”课题组，"
            "留学期间建立的国际合作网络（MIT Chiang 课题组、布鲁克海文国家实验室）"
            "将成为课题组国际合作基础。本人承诺留学期间每年回国 1 次向国内导师汇报"
            "并参加国内学术会议，保持与国内研究团队衔接，确保归国后研究工作顺利推进。"
        )
    },

    # 学习计划 5 要素
    "research_plan": {
        "research_direction": "硫碳复合正极的原位表征与界面调控",
        "research_background": (
            "该研究是本人国内博士课题“硫碳复合正极制备与电化学性能”的延续与深化，"
            "与外方 Chiang 课题组“锂硫电池机理研究”方向一致。"
        ),
        "research_objectives": [
            "建立硫碳复合正极充放电过程原位同步辐射表征方法",
            "揭示硫还原中间产物多硫化锂的迁移机理",
            "明确碳基底孔结构与多硫化锂吸附/转化的构效关系",
            "指导设计 500 Wh/kg 锂硫电池原型器件"
        ],
        "research_contents": [
            "硫碳复合正极原位 XRD 表征",
            "多硫化锂迁移的原位 XANES 追踪",
            "碳基底孔结构调控与电化学性能关联",
            "原型器件组装与测试"
        ],
        "research_methods": "以原位同步辐射实验为主，辅以 DFT 计算",
        "courses": [
            {"code": "3.091", "name": "Introduction to Solid State Chemistry",
             "credits": 5.0, "reason": "基础巩固材料化学基础"},
            {"code": "3.43", "name": "Electrochemical Energy Storage",
             "credits": 4.5, "reason": "与本人研究方向直接相关"},
            {"code": "3.320", "name": "Battery Materials and Devices",
             "credits": 3.0, "reason": "前沿课程"},
            {"code": "3.21", "name": "Kinetics of Materials",
             "credits": 3.0, "reason": "理论方法补充"}
        ],
        "advisor_collab": (
            "外方导师为 MIT 材料系 Yet-Ming Chiang 教授（Kyocera 讲席教授，"
            "美国国家工程院院士），研究方向为电化学储能与材料科学，H 指数 95，"
            "2020~2024 年发表论文 28 篇，其中 Nature/Science 主刊 4 篇，"
            "获美国能源部 ARPA-E 项目资助 1200 万美元。Chiang 教授研究方向与本人"
            "博士课题高度匹配，其课题组硫正极研究是国际风向标。Chiang 教授与本人国内"
            "导师李某某教授已有 6 年合作基础，共同发表 Nature Communications 论文 1 篇"
            "（2022），互访 3 次，邮件往来 200+ 封（部分见附件 7）。本人拟每周参加"
            "Chiang 课题组组会 1 次，每两周与 Chiang 教授一对一讨论 1 次。"
        ),
        "collab_research": {
            "my_role": "硫碳复合正极样品制备、原位 XRD 实验设计、数据分析与机理建模",
            "advisor_role": "同步辐射光源申请与机时安排、原位电化学池设计、DFT 计算支持",
            "sustainability": "留学结束后本课题将继续以“中-美合作”形式延续，国内导师李某某"
                              "教授与 Chiang 教授已签订 5 年合作框架协议（见附件 9）"
        },
        "expected_outcomes": {
            "papers": 3, "first_author_papers": 2,
            "target_journals": ["Advanced Energy Materials (IF 24.4)",
                                "Joule (IF 41.2)",
                                "ACS Energy Letters (IF 22.0)"],
            "patents_pct": 1, "patents_domestic": 2,
            "conferences": 2, "oral_presentations": 1
        },
        "timeline": [
            {"period": "2024.09~2024.12", "task": "硫碳复合正极制备工艺优化"},
            {"period": "2025.01~2025.06", "task": "原位 XRD 表征与数据分析"},
            {"period": "2025.07~2025.12", "task": "原型器件组装与测试"},
            {"period": "2026.01~2026.06", "task": "数据分析与论文撰写"},
            {"period": "2026.07~2026.08", "task": "归国总结与成果汇报"}
        ]
    },

    # 国内外导师合作基础
    "collaboration_basis": {
        "joint_papers": [{"title": "Mechanistic study of sulfur cathode",
                          "journal": "Nature Communications", "year": 2022}],
        "mutual_visits": 3, "email_exchanges": 200,
        "joint_projects": [{"name": "中美锂硫电池合作项目",
                             "period": "2020-2025", "funder": "NSFC-NSF"}]
    },

    # 归国计划
    "return_plan": {
        "return_date": "2026-09",
        "target_unit": "XX 大学材料学院",
        "position": "副教授",
        "research_plan": "组建“先进电池材料”课题组，继续推进锂硫电池产业化研究",
        "industry_partner": "宁德时代新能源科技股份有限公司",
        "service_years": 2,
        "long_term_plan": "3 年内申报国家自然科学基金青年项目，"
                          "5 年内组建 5 人课题组，推动 500 Wh/kg 锂硫电池中试",
        "return_letter_signed": True
    },

    # 经济保障
    "financial_arrangement": {
        "csc_allowance_monthly_usd": 1900,
        "csc_allowance_total_usd": 45600,
        "tuition_waived": True,
        "tuition_waiver_source": "MIT 材料系免学费（见邀请函附件 5）",
        "advisor_research_stipend_monthly_usd": 500,
        "spouse_income_monthly_cny": 12000,
        "emergency_fund_cny": 60000,
        "insurance_arrangement": "CSC 统一购买国际医疗保险 + 个人购买商业补充保险"
    },

    # 输出控制
    "word_count_version": "3000",
    "material_type": "B",
    "salutation": "尊敬的国家留学基金管理委员会评审专家：",
    "opening": "", "motivation_text": "", "university_choice_text": "",
    "study_plan_text": "", "return_plan_text": "", "financial_text": "",
    "apply_date": "2024 年 3 月 10 日"
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="国家公派留学申请书（CSC Scholarship Application）docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "  python build.py --demo --csc-type phd --out demo_phd.docx\n"
            "  python build.py --demo --word-count 3000 --out demo_3000.docx\n"
            "\n"
            "CSC 4 类项目：phd / joint_phd / master / visiting_scholar\n"
            "3 档字数版本：2500 / 3000 / 4000\n"
            "依据：留金发〔2007〕3010 号 + 留金发〔2020〕112 号\n"
            "\nJSON 字段详见 SKILL.md §9 JSON Schema。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    parser.add_argument("--csc-type", type=str, default=None,
                        choices=["phd", "joint_phd", "master", "visiting_scholar"],
                        help="项目类别（覆盖 data 中 csc_type）")
    parser.add_argument("--word-count", type=str, default=None,
                        choices=["2500", "3000", "4000"],
                        help="字数版本（覆盖 data 中 word_count_version）")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA.copy()
        if args.csc_type:
            data["csc_type"] = args.csc_type
            data["csc_type_label"] = CSC_TYPE_LABEL.get(
                args.csc_type, data.get("csc_type_label", ""))
        if args.word_count:
            data["word_count_version"] = args.word_count
        print(f"ℹ️ 使用内置示例数据生成演示文档"
              f"（{data.get('word_count_version', '3000')} 档，"
              f"{data.get('csc_type_label', '联合培养博士研究生')}项目）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
        if args.csc_type:
            data["csc_type"] = args.csc_type
            data["csc_type_label"] = CSC_TYPE_LABEL.get(
                args.csc_type, data.get("csc_type_label", ""))
        if args.word_count:
            data["word_count_version"] = args.word_count
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
