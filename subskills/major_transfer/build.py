#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
转专业申请书 docx 生成器

支持 3 类转专业原因（通过 transfer_reason_type 字段切换）：
- interest: 兴趣驱动型（兴趣萌芽 + 持续探索 + 已有成果）
- academic: 学业优势型（交叉课程高分 + 竞赛/科研 + 可迁移论证）
- career: 职业规划型（明确职业方向 + 行业调研 + 目标专业支撑性）

申请窗口：大一第二学期末（5~6 月）或大二第一学期初（9~10 月）。
正文 1000~1500 字，书信体格式，5 段结构（个人基本情况 + 现专业学习 +
转专业原因 + 目标专业认知 + 转后规划）。

按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 标题：黑体二号，居中（固定"转专业申请书"）
- 称呼：顶格，宋体小四，全角冒号
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 现专业主干课程表、目标专业课程表、转后学习计划表：宋体五号，居中
- "此致"另起一行空两格，"敬礼！"另起一行顶格
- 落款：右对齐

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第五章信息采集清单。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号（标题）
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四（正文）
SIZE_WU = Pt(10.5)          # 五号（表格）
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

# 转专业原因类型
REASON_INTEREST = "interest"
REASON_ACADEMIC = "academic"
REASON_CAREER = "career"
VALID_REASONS = (REASON_INTEREST, REASON_ACADEMIC, REASON_CAREER)

# 标题
TITLE_TEXT = "转专业申请书"


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(doc, text: str, font_name: str = FONT_SONG,
        font_size=SIZE_XIAO_SI, bold: bool = False,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent: bool = True,
        line_spacing: float = 1.5, space_before: float = 0, space_after: float = 0):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_ER,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进），宋体小四"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_section_heading(doc, text: str):
    """正文小节标题（一、二、三…）：黑体小四加粗，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SI,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
        line_spacing=1.5, space_before=6, space_after=3)


def add_cizhi_paragraph(doc, text: str = "此致"):
    """'此致'另起一行，空两格（首行缩进 2 字符）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格（不缩进）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=False, line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None, caption: str = ""):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    if caption:
        add_paragraph_with_format(doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def format_rank_percent(rank: str, rank_total: str) -> str:
    """根据排名与基数计算百分比，返回 '前 X.X%' 字符串，失败返回空串"""
    try:
        r_num = int(str(rank).split("/")[0])
        total_num = int(rank_total) if rank_total else 0
        if total_num > 0:
            return f"前 {round(r_num / total_num * 100, 1)}%"
    except (ValueError, IndexError):
        pass
    return ""


# ============================================================
# MajorTransferDocBuilder 主类
# ============================================================

class MajorTransferDocBuilder:
    """转专业申请书 docx 构建器（按 transfer_reason_type 切换原因段撰写）"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []
        self.reason_type = REASON_INTEREST

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段，非列表返回空列表"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_heading(self, text):
        return add_section_heading(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, caption=""):
        return add_table_from_data(self.doc, headers, rows, col_widths, caption)

    # --------------------------------------------------------
    # 标题（固定"转专业申请书"）
    # --------------------------------------------------------

    def _add_title(self):
        """标题：黑体二号居中，固定'转专业申请书'"""
        add_title(self.doc, TITLE_TEXT)

    # --------------------------------------------------------
    # 称呼
    # --------------------------------------------------------

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号"""
        salutation = self._get("salutation",
                               default="尊敬的学院领导、教务处老师：")
        add_salutation_paragraph(self.doc, salutation)

    # --------------------------------------------------------
    # 开头段落
    # --------------------------------------------------------

    def _add_opening(self):
        """开头段落（约 50 字）：身份 + 申请事项 + 汇报句"""
        opening = self._get("opening", default="")
        if opening:
            self.add_para(opening)
            return
        name = self._get("name")
        current_college = self._get("current_college")
        current_major = self._get("current_major")
        grade = self._get("grade")
        target_college = self._get("target_college")
        target_major = self._get("target_major")
        parts = []
        if name and current_college and current_major and grade:
            parts.append(f"我是{current_college}{current_major}{grade}学生{name}，现申请转入{target_college}{target_major}专业。")
        else:
            parts.append(f"现申请转入{target_college}{target_major}专业。")
        parts.append("现就转专业事宜汇报如下：")
        self.add_para("".join(parts))

    # --------------------------------------------------------
    # 一、个人基本情况
    # --------------------------------------------------------

    def _add_basic_info(self):
        """个人基本情况（约 150~225 字）：身份 + 入学背景 + 高考志愿说明"""
        self.add_heading("一、个人基本情况")
        basic_info = self._get("basic_info", default="")
        if basic_info:
            if isinstance(basic_info, list):
                for p in basic_info:
                    self.add_para(p)
            else:
                self.add_para(basic_info)
            return
        gaokao_year = self._get("gaokao_year", default="")
        gaokao_score = self._get("gaokao_score", default="")
        admission_mode = self._get("admission_mode", default="")
        original_choice = self._get("original_choice", default="")
        current_reason = self._get("current_reason", default="")
        current_college = self._get("current_college", default="")
        current_major = self._get("current_major", default="")
        target_college = self._get("target_college", default="")
        target_major = self._get("target_major", default="")
        parts = []
        if gaokao_year and gaokao_score:
            parts.append(f"{gaokao_year}高考{gaokao_score}，")
        if admission_mode:
            mode_str = f"因{admission_mode}" if "调剂" in admission_mode else admission_mode
            parts.append(f"{mode_str}进入{current_college}{current_major}专业。")
        if current_reason:
            parts.append(current_reason + "。")
        if original_choice:
            parts.append(f"高考第一志愿为{original_choice}，")
        if target_college and target_major:
            parts.append(f"经过大一一年的学习与思考，我希望转入{target_college}{target_major}专业。")
        if parts:
            self.add_para("".join(parts))

    # --------------------------------------------------------
    # 二、现专业学习情况（含课程表）
    # --------------------------------------------------------

    def _add_current_study(self):
        """现专业学习情况（约 200~300 字）：GPA + 排名 + 主干课程表 + 交叉课程"""
        self.add_heading("二、现专业学习情况")
        current_study = self._get("current_study", default="")
        if current_study and isinstance(current_study, str):
            self.add_para(current_study)
            self._add_current_courses_table()
            return
        gpa = self._get("current_gpa", default="")
        weighted = self._get("current_weighted_avg", default="")
        rank = self._get("current_rank", default="")
        rank_total = self._get("current_rank_total", default="")
        course_count = self._get("current_course_count", default="")
        high_score_count = self._get("current_high_score_count", default="")
        parts = []
        if gpa:
            gpa_str = f"大一学年 GPA {gpa}"
            if weighted:
                gpa_str += f"，加权平均分 {weighted}"
            if rank:
                gpa_str += f"，专业排名第 {rank}"
                if rank_total:
                    pct = format_rank_percent(rank, rank_total)
                    if pct:
                        gpa_str += f"（{pct}）"
            parts.append(gpa_str + "。")
        if course_count and high_score_count:
            parts.append(f"本学年修读 {course_count} 门课程，{high_score_count} 门 85 分以上。")
        if parts:
            self.add_para("".join(parts))
        self._add_current_courses_table()
        if self.reason_type == REASON_ACADEMIC:
            self._add_cross_courses_note()
        self._add_academics_summary()

    def _add_current_courses_table(self):
        """现专业主干课程表：4~6 门，含课程名/学分/成绩"""
        courses = self._get_list("current_core_courses")
        if not courses:
            return
        rows = []
        for c in courses:
            if not isinstance(c, dict):
                continue
            rows.append([str(c.get("name", "")), str(c.get("credit", "")), str(c.get("score", ""))])
        if rows:
            self.add_table(["课程名称", "学分", "成绩"], rows,
                           col_widths=[8.0, 3.0, 3.0], caption="现专业主干课程成绩：")

    def _add_cross_courses_note(self):
        """学业优势型：高亮与目标专业交叉课程"""
        cross_courses = self._get_list("cross_courses")
        if not cross_courses:
            return
        names = []
        for c in cross_courses:
            if isinstance(c, dict):
                name = c.get("name", "")
                score = c.get("score", "")
                if name and score:
                    names.append(f"{name} {score} 分")
        if names:
            self.add_para(
                "其中与目标专业交叉的高分课程包括：" + "、".join(names) +
                "，为转入目标专业后的学习奠定扎实基础。"
            )

    def _add_academics_summary(self):
        """现专业学习末尾：整体学业能力评价一句话"""
        summary = self._get("academics_summary", default="")
        if summary:
            self.add_para(summary)
        else:
            self.add_para("总体而言，我在现专业学习中保持了较好的学业表现，具备扎实的数理基础与学习能力。")

    # --------------------------------------------------------
    # 三、转专业原因（按 3 类切换）
    # --------------------------------------------------------

    def _add_transfer_reason(self):
        """转专业原因（约 250~375 字）：按 transfer_reason_type 切换主线"""
        self.add_heading("三、转专业原因")
        reason_text = self._get("transfer_reason", default="")
        if reason_text:
            if isinstance(reason_text, list):
                for p in reason_text:
                    self.add_para(p)
            else:
                self.add_para(reason_text)
            return
        # 按类型分支
        rtype = self.reason_type
        if rtype == REASON_INTEREST:
            self._build_interest_reason()
        elif rtype == REASON_ACADEMIC:
            self._build_academic_reason()
        elif rtype == REASON_CAREER:
            self._build_career_reason()
        else:
            self._build_interest_reason()  # 默认兴趣驱动

    def _build_interest_reason(self):
        """兴趣驱动型：兴趣萌芽 + 持续探索 + 已有成果 + 契合点 + 中性收束"""
        origin = self._get("interest_origin", default="")
        exploration = self._get("interest_exploration", default="")
        books = self._get("interest_books", default="")
        club = self._get("interest_club", default="")
        target_major = self._get("target_major", default="")
        current_major = self._get("current_major", default="")
        parts = []
        if origin:
            parts.append(f"自{origin}中接触到相关内容后，我萌生对{target_major}的浓厚兴趣。")
        else:
            parts.append(f"我对{target_major}产生了浓厚兴趣。")
        if exploration:
            parts.append(exploration + "。")
        if books:
            parts.append(books + "。")
        if club:
            parts.append(club + "。")
        parts.append(
            f"这些探索让我确认{target_major}是我真正热爱的方向。"
            f"经过大一一年的学习与思考，我发现自己的兴趣与"
            f"{current_major}专业的课程体系侧重存在差异，故申请转入{target_major}专业。"
        )
        self.add_para("".join(parts))

    def _build_academic_reason(self):
        """学业优势型：交叉课程高分 + 竞赛/科研 + 可迁移论证"""
        cross = self._get("academic_cross", default="")
        competitions = self._get_list("academic_competitions")
        research = self._get("academic_research", default="")
        gpa = self._get("current_gpa", default="")
        rank = self._get("current_rank", default="")
        target_major = self._get("target_major", default="")
        current_major = self._get("current_major", default="")
        parts = []
        if gpa and rank:
            parts.append(f"大一学年 GPA {gpa}，专业排名第 {rank}。")
        if cross:
            parts.append(cross + "。")
        if competitions:
            comp_parts = []
            for c in competitions:
                if not isinstance(c, dict):
                    continue
                name, award, time, role = c.get("name", ""), c.get("award", ""), c.get("time", ""), c.get("role", "")
                seg = (f"{time} " if time else "") + name
                if award:
                    seg += f" {award}"
                if role:
                    seg += f"（{role}）"
                if seg:
                    comp_parts.append(seg + "；")
            if comp_parts:
                comp_parts[-1] = comp_parts[-1].rstrip("；") + "。"
                parts.append("同时参加" + "".join(comp_parts))
        if research:
            parts.append(research + "。")
        parts.append(
            f"这些学业成果证明我具备转入{target_major}后的学业竞争力。"
            f"经过充分考虑，我希望在{target_major}专业深入发展，现就读的{current_major}专业虽提供良好基础，"
            f"但与我的学业优势方向存在差异，故申请转入{target_major}专业。"
        )
        self.add_para("".join(parts))

    def _build_career_reason(self):
        """职业规划型：明确职业方向 + 行业调研 + 目标专业支撑性 + 现专业限制"""
        target = self._get("career_target", default="")
        internship = self._get("career_internship", default="")
        interview = self._get("career_interview", default="")
        research = self._get("career_research", default="")
        target_major = self._get("target_major", default="")
        current_major = self._get("current_major", default="")
        parts = []
        if target:
            parts.append(f"我的长期职业目标是成为{target}。")
        if internship:
            parts.append(internship + "。")
        if interview:
            parts.append(interview + "后，")
        parts.append(f"我确认{target_major}是通往该职业的必经路径。")
        if research:
            parts.append(research + "。")
        parts.append(f"现就读的{current_major}专业缺乏相关课程与实践平台，无法支撑我的职业目标，故申请转入{target_major}专业。")
        self.add_para("".join(parts))

    # --------------------------------------------------------
    # 四、目标专业认知（4 维度）
    # --------------------------------------------------------

    def _add_target_understanding(self):
        """目标专业认知（约 250~375 字）：4 维度（专业课程/师资力量/就业前景/学科竞赛）"""
        self.add_heading("四、目标专业认知")
        understanding = self._get("target_understanding", default="")
        if understanding and isinstance(understanding, str):
            self.add_para(understanding)
            self._add_target_courses_table()
            return
        target_major = self._get("target_major", default="")
        self.add_para(f"通过查阅目标学院{target_major}专业培养方案与官网信息，我已系统了解其核心课程体系、师资力量、就业前景与学科竞赛。")
        self._add_target_courses_dim()
        self._add_target_teachers_dim()
        self._add_target_employment_dim()
        self._add_target_competitions_dim()

    def _add_target_courses_dim(self, table_only: bool = False):
        """维度 1：专业课程（至少 5 门，含表格 + 文字描述）；table_only=True 时仅输出表格"""
        courses = self._get_list("target_courses")
        if not courses:
            return
        rows, names = [], []
        for c in courses:
            if not isinstance(c, dict):
                continue
            name, credit, goal = str(c.get("name", "")), str(c.get("credit", "")), str(c.get("goal", ""))
            rows.append([name, credit, goal])
            if name and not table_only:
                seg = f"《{name}》"
                if credit:
                    seg += f"（{credit} 学分）"
                if goal:
                    seg += f"培养{goal}能力"
                names.append(seg)
        if rows:
            self.add_table(["课程名称", "学分", "培养目标"], rows,
                           col_widths=[6.5, 2.5, 5.0], caption="目标专业核心课程：")
        if names:
            self.add_para("**专业课程方面**，目标专业核心课程包括" + "、".join(names) + "。")

    def _add_target_courses_table(self):
        """整段覆盖模式下也输出课程表（调用 _add_target_courses_dim(table_only=True)）"""
        self._add_target_courses_dim(table_only=True)

    def _add_target_teachers_dim(self):
        """维度 2：师资力量（至少 2 位）"""
        teachers = self._get_list("target_teachers")
        if not teachers:
            return
        parts = ["**师资力量方面**，"]
        for t in teachers:
            if not isinstance(t, dict):
                continue
            name, title, research = t.get("name", ""), t.get("title", ""), t.get("research", "")
            seg = name
            if title:
                seg += f"（{title}）"
            if research:
                seg += f"研究方向为{research}"
            if seg:
                parts.append(seg + "；")
        if len(parts) > 1:
            parts[-1] = parts[-1].rstrip("；") + "。"
            self.add_para("".join(parts))

    def _add_target_employment_dim(self):
        """维度 3：就业前景（含方向/就业率/薪资/来源）"""
        emp = self._get("target_employment", default="")
        if not emp:
            return
        if isinstance(emp, str):
            self.add_para(f"**就业前景方面**，{emp}")
            return
        if not isinstance(emp, dict):
            return
        directions = emp.get("directions", [])
        rate, salary, source = emp.get("rate", ""), emp.get("salary", ""), emp.get("source", "")
        parts = ["**就业前景方面**，"]
        if source:
            parts.append(f"据{source}数据，")
        if rate:
            parts.append(f"应届生就业率{rate}")
        if salary:
            parts.append(f"，{salary}")
        parts.append("。")
        if directions and isinstance(directions, list):
            parts.append("主要就业方向包括" + "、".join(str(d) for d in directions) + "。")
        self.add_para("".join(parts))

    def _add_target_competitions_dim(self):
        """维度 4：学科竞赛（至少 2 个）"""
        competitions = self._get_list("target_competitions")
        if not competitions:
            return
        parts = ["**学科竞赛方面**，我已规划转入后参与以下竞赛："]
        for c in competitions:
            if not isinstance(c, dict):
                continue
            time, name, level, role = c.get("time", ""), c.get("name", ""), c.get("level", ""), c.get("role", "")
            seg = (f"{time} " if time else "") + name
            if level:
                seg += f"（{level}）"
            if role:
                seg += f"，{role}"
            if seg:
                parts.append(seg + "；")
        if len(parts) > 1:
            parts[-1] = parts[-1].rstrip("；") + "。"
            self.add_para("".join(parts))

    # --------------------------------------------------------
    # 五、转后规划（含学习计划表）
    # --------------------------------------------------------

    def _add_future_plan(self):
        """转后规划（约 150~225 字）：短/中/长期 + 学习计划表"""
        self.add_heading("五、转后规划")
        future_plan = self._get("future_plan", default="")
        if future_plan:
            if isinstance(future_plan, list):
                for p in future_plan:
                    self.add_para(p)
            else:
                self.add_para(future_plan)
            self._add_study_plan_table()
            return
        short_term = self._get("short_term_plan", default="")
        mid_term = self._get("mid_term_plan", default="")
        long_term = self._get("long_term_plan", default="")
        target_major = self._get("target_major", default="")
        parts = [f"转入{target_major}专业后，我的学习规划如下："]
        if short_term:
            parts.append(f"**短期**（转入后 1 学期）：{short_term}。")
        if mid_term:
            parts.append(f"**中期**（转入后 1~2 学年）：{mid_term}。")
        if long_term:
            parts.append(f"**长期**（毕业去向）：{long_term}。")
        parts.append(f"我相信通过上述规划，能够在转入{target_major}专业后迅速适应并取得优异成绩。")
        self.add_para("".join(parts))
        self._add_study_plan_table()

    def _add_study_plan_table(self):
        """转后学习计划表：短/中/长期 3 行（4 列：阶段/时间/主要内容/目标）"""
        plans = self._get_list("study_plan_table")
        if not plans:
            return
        rows = []
        for p in plans:
            if not isinstance(p, dict):
                continue
            rows.append([str(p.get("stage", "")), str(p.get("time", "")),
                         str(p.get("content", "")), str(p.get("target", ""))])
        if rows:
            self.add_table(["阶段", "时间", "主要内容", "目标"], rows,
                           col_widths=[2.0, 3.5, 5.5, 3.0], caption="转后学习计划表：")

    # --------------------------------------------------------
    # 结尾"此致 敬礼！"
    # --------------------------------------------------------

    def _add_ending(self):
        """结尾（约 50~75 字）：朴素表态 + 此致/敬礼！"""
        ending = self._get("ending", default="")
        if ending:
            self.add_para(ending)
        else:
            self.add_para(
                "以上是我的转专业申请。无论审批结果如何，我都将继续努力学习，以更高标准要求自己。恳请评审委员会予以考虑。"
            )
        add_cizhi_paragraph(self.doc, "此致")
        add_jingli_paragraph(self.doc, "敬礼！")

    # --------------------------------------------------------
    # 落款
    # --------------------------------------------------------

    def _add_signature(self):
        """落款：右对齐，含申请人 + 日期"""
        self.doc.add_paragraph()  # 空一行
        name = self._get("name", default="")
        apply_date = self._get("apply_date", default="")
        if name:
            add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if apply_date:
            add_right_aligned_paragraph(self.doc, apply_date)

    # --------------------------------------------------------
    # 主构建方法
    # --------------------------------------------------------

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/开头/个人基本情况/现专业学习/
        转专业原因/目标专业认知/转后规划/结尾/落款"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_opening()
            self._add_basic_info()
            self._add_current_study()
            self._add_transfer_reason()
            self._add_target_understanding()
            self._add_future_plan()
            self._add_ending()
            self._add_signature()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 转专业申请书已生成：{output_path}")
        return str(output_path)

    # --------------------------------------------------------
    # 数据校验
    # --------------------------------------------------------

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []
        # 转专业原因类型校验
        rtype = str(self._get("transfer_reason_type", default=REASON_INTEREST)).lower()
        if rtype in VALID_REASONS:
            self.reason_type = rtype
        else:
            warnings.append(f"转专业原因类型 {rtype} 非标准值，默认按兴趣驱动型处理")
            self.reason_type = REASON_INTEREST
        # P0 必采字段
        for key, label in [("name", "申请人姓名"), ("current_college", "现学院"),
                           ("current_major", "现专业"), ("grade", "年级"),
                           ("target_college", "目标学院"), ("target_major", "目标专业"),
                           ("current_gpa", "现专业 GPA"), ("current_rank", "现专业排名")]:
            if not self._get(key):
                warnings.append(f"缺少 {label}（{key}）")
        # 排名校验
        rank_str = str(self._get("current_rank", default=""))
        if rank_str and "/" not in rank_str and not self._get("current_rank_total"):
            warnings.append(f"排名 '{rank_str}' 缺基数，应为 'X/N' 格式或补填 current_rank_total")
        # 课程/师资/就业/竞赛/计划表数量校验
        checks = [
            ("current_core_courses", 4, "现专业主干课程", "4~6 门"),
            ("target_courses", 5, "目标专业核心课程", "至少 5 门"),
            ("target_teachers", 2, "目标专业师资", "至少 2 位"),
            ("target_competitions", 2, "学科竞赛", "至少 2 个"),
            ("study_plan_table", 3, "转后学习计划表", "短/中/长期 3 行"),
        ]
        for key, min_n, label, hint in checks:
            items = self._get_list(key)
            if not items:
                warnings.append(f"缺少 {label}（{key}），将省略对应内容")
            elif len(items) < min_n:
                warnings.append(f"{label}仅 {len(items)} 项，建议{hint}")
        # 就业前景
        if not self._get("target_employment"):
            warnings.append("缺少 就业前景（target_employment），将省略就业段落")
        # 3 类原因对应字段校验
        reason_field_checks = {
            REASON_INTEREST: [("interest_origin", "兴趣萌芽时间点"), ("interest_exploration", "探索成果")],
            REASON_ACADEMIC: [("academic_cross", "学业优势核心论据"), ("academic_competitions", "相关竞赛")],
            REASON_CAREER: [("career_target", "明确职业方向"), ("career_internship", "行业实习经历")],
        }
        reason_label = {REASON_INTEREST: "兴趣驱动型", REASON_ACADEMIC: "学业优势型", REASON_CAREER: "职业规划型"}
        for field, label in reason_field_checks.get(self.reason_type, []):
            val = self._get_list(field) if field == "academic_competitions" else self._get(field)
            if not val:
                warnings.append(f"{reason_label[self.reason_type]}建议提供 {field}（{label}）")
        self.warnings = warnings
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据（兴趣驱动型）
# ============================================================

DEFAULT_DATA = {
    "name": "张明", "student_id": "2024123456", "gender": "男",
    "current_college": "化学工程学院", "current_major": "化学工程与工艺",
    "grade": "2024 级 大一", "class_name": "化工 2401 班",
    "target_college": "计算机科学与技术学院", "target_major": "计算机科学与技术",
    "phone": "138XXXXXXXX", "apply_date": "2025 年 5 月 20 日",
    "apply_window": "大一第二学期", "transfer_reason_type": "interest",
    "salutation": "尊敬的学院领导、教务处老师：",
    "current_gpa": "3.92/4.0", "current_weighted_avg": "89.5",
    "current_rank": "3/95", "current_rank_total": "95",
    "current_course_count": "12", "current_high_score_count": "9",
    "current_core_courses": [
        {"name": "高等数学（上）", "credit": "5", "score": "95"},
        {"name": "线性代数", "credit": "3", "score": "94"},
        {"name": "无机化学", "credit": "4", "score": "88"},
        {"name": "大学物理", "credit": "4", "score": "90"},
        {"name": "C 程序设计", "credit": "3", "score": "96"},
    ],
    "cross_courses": [
        {"name": "高等数学（上）", "credit": "5", "score": "95"},
        {"name": "C 程序设计", "credit": "3", "score": "96"},
    ],
    "interest_origin": "2024.09 通识课《人工智能导论》",
    "interest_exploration": (
        "8 个月来，利用课余时间自学吴恩达《Machine Learning》课程"
        "（累计学习 120 小时），完成 3 个 Kaggle 入门项目"
        "（Titanic 分类、House Prices 回归、Digit Recognizer，最好成绩 top 30%）"
    ),
    "interest_books": "阅读《深度学习》（花书）前 6 章",
    "interest_club": (
        "加入校 ACM 协会，每周参与 2 次算法训练，累计解决 LeetCode 题 150 道"
    ),
    "gaokao_year": "2024 年", "gaokao_score": "632 分（山东）",
    "admission_mode": "普通批次",
    "original_choice": "计算机科学与技术",
    "current_reason": (
        "因分数未达计算机科学与技术专业录取线，调剂至化学工程与工艺专业"
    ),
    "target_courses": [
        {"name": "数据结构", "credit": "4", "goal": "算法设计"},
        {"name": "操作系统", "credit": "4", "goal": "系统级编程"},
        {"name": "计算机网络", "credit": "3", "goal": "网络协议分析"},
        {"name": "数据库系统原理", "credit": "3", "goal": "数据管理"},
        {"name": "编译原理", "credit": "4", "goal": "语言处理"},
    ],
    "target_teachers": [
        {"name": "李XX", "title": "教授", "research": "计算机视觉与医学影像分析"},
        {"name": "王XX", "title": "副教授", "research": "机器学习与推荐系统"},
    ],
    "target_employment": {
        "directions": ["互联网算法工程师", "金融机构 IT 岗位", "政府信息化岗位"],
        "rate": "96.8%", "salary": "8500 元/月起薪",
        "source": "麦可思《2024 中国大学生就业报告》",
    },
    "target_competitions": [
        {"name": "全国大学生数学建模竞赛", "level": "国家级", "time": "2025.09", "role": "算法负责人，目标省级一等奖"},
        {"name": "ACM-ICPC 区域赛", "level": "国家级", "time": "2026.03", "role": "队长，目标铜牌"},
        {"name": "中国大学生计算机设计大赛", "level": "省级", "time": "2026.05", "role": "后端开发，目标省级二等奖"},
    ],
    "short_term_plan": "补修《C++ 程序设计》（4 学分）、《数据结构》（4 学分）共 8 学分；现专业已修通识课 18 学分可认定",
    "mid_term_plan": "加入李XX 教授课题组参与医学影像分析项目，目标以第二作者发表 1 篇 SCI 论文；2026.03 参加 ACM-ICPC 区域赛",
    "long_term_plan": "保研至清华大学计算机系医疗 AI 方向",
    "study_plan_table": [
        {"stage": "短期", "time": "2025.09~2026.01", "content": "补修《C++》《数据结构》共 8 学分", "target": "GPA≥3.8，学分补修完毕"},
        {"stage": "中期", "time": "2026.02~2027.06", "content": "加入李XX 课题组 + ACM-ICPC 区域赛", "target": "第二作者 SCI 论文 1 篇 + 铜牌"},
        {"stage": "长期", "time": "2027.07~2028.06", "content": "保研清华计算机系医疗 AI 方向", "target": "获保研资格 + 入读清华"},
    ],
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="转专业申请书 docx 生成器（按 transfer_reason_type 切换 3 类原因）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=("示例：\n  python build.py --data data.json --out output.docx\n"
                "  python build.py --demo --out demo.docx\n\n"
                "JSON 字段定义详见 SKILL.md 第五章信息采集清单。\n"
                "transfer_reason_type: interest（兴趣驱动）/ academic（学业优势）/ career（职业规划）"),
    )
    parser.add_argument("--data", type=str, default=None, help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True, help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true", help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档（默认兴趣驱动型，化学工程与工艺 → 计算机科学与技术）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = MajorTransferDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
