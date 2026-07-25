#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
大学生创新创业训练计划-创业训练项目申报书 docx 生成器（v2.1）

按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 一级标题：黑体三号，居中
- 二级标题：黑体小三，左对齐
- 三级标题：宋体四号加粗
- 表格：宋体五号，居中

v2.1 升级（基于案例 2 河南工业大学姚奕晗 79 页范本）：
- 11 大章必加结构（立项依据/研究内容/研究方案/社会效益/市场分析/营销模式/财务/实践/壁垒/进度/签字）
- 15 个新增字段（policy_citations/scientific_challenges/literature_review/algorithm_comparison/
  tech_roadmap/formulas/economic_benefits/market_analysis/marketing_4p/pricing_3c/
  financial_statements/practice_process/tech_barriers/project_schedule 等）
- DEFAULT_DATA 含消防无人机创业训练完整 demo（对齐案例 2）

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第七章。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)
SIZE_XIAO_ER = Pt(18)
SIZE_SAN = Pt(16)
SIZE_XIAO_SAN = Pt(15)
SIZE_SI = Pt(14)
SIZE_XIAO_SI = Pt(12)
SIZE_WU = Pt(10.5)
SIZE_XIAO_WU = Pt(9)

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# ============================================================
# 工具函数（中英文同步设置 eastAsia/ascii/hAnsi）
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格文字内容与字体"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_paragraph_with_format(doc, text: str,
                              font_name: str = FONT_SONG,
                              font_size=SIZE_XIAO_SI, bold: bool = False,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT,
                              first_line_indent: bool = True,
                              line_spacing: float = 1.5,
                              space_before: int = 0,
                              space_after: int = 0,
                              color: Optional[RGBColor] = None) -> None:
    """添加带格式段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size,
                 bold=bold, color=color)


def add_heading_level1(doc, text: str) -> None:
    """一级标题：黑体三号居中"""
    add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
        space_before=12, space_after=12,
    )


def add_heading_level2(doc, text: str) -> None:
    """二级标题：黑体小三左对齐"""
    add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=8, space_after=4,
    )


def add_heading_level3(doc, text: str) -> None:
    """三级标题：宋体四号加粗"""
    add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6, space_after=3,
    )


def add_body_paragraph(doc, text: str, indent: bool = True) -> None:
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_XIAO_SI, bold=False,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent,
        line_spacing=1.5,
    )


def add_bullet_paragraph(doc, text: str) -> None:
    """项目符号段落（首行缩进 2 字符 + '● ' 前缀）"""
    add_body_paragraph(doc, f"● {text}", indent=True)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None):
    """从数据生成表格（首行表头加粗）"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h,
                      font_name=FONT_HEI, font_size=SIZE_WU, bold=True)
        table.rows[0].cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == len(row) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False, alignment=align)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_canvas_9grid(doc, canvas_items: List[Dict[str, str]]):
    """商业模式画布 9 宫格表格（3 行 × 3 列）"""
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, item in enumerate(canvas_items[:9]):
        r, c = divmod(idx, 3)
        cell = table.rows[r].cells[c]
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(item.get("element", ""))
        set_run_font(run1, font_name=FONT_HEI, font_size=SIZE_XIAO_WU, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.line_spacing = 1.2
        run2 = p2.add_run(item.get("content", ""))
        set_run_font(run2, font_name=FONT_SONG, font_size=SIZE_XIAO_WU)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(5.3)
    doc.add_paragraph()


def setup_page(doc):
    """设置 A4 页面与页边距"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ============================================================
# ApplicationDocBuilder 主类（v2.1，11 大章 + 15 新字段）
# ============================================================

class ApplicationDocBuilder:
    """大创-创业训练项目申报书 docx 构建器（v2.1，对齐案例 2）"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}

    def _get(self, *keys, default=""):
        """安全取嵌套字段"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def add_h1(self, text):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text):
        return add_heading_level2(self.doc, text)

    def add_h3(self, text):
        return add_heading_level3(self.doc, text)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_bullet(self, text):
        return add_bullet_paragraph(self.doc, text)

    def add_table(self, headers, rows, col_widths=None):
        return add_table_from_data(self.doc, headers, rows, col_widths)

    def add_page_break(self):
        add_page_break(self.doc)

    # --------------------------------------------------------
    # 封面
    # --------------------------------------------------------

    def _add_cover(self):
        """封面：黑体二号标题 + 副标题 + 5 行下划线信息"""
        for _ in range(2):
            self.doc.add_paragraph()

        level = self._get("project_level", default="校级")
        title = f"{level}大学生创新创业训练计划项目申报书"
        add_paragraph_with_format(
            self.doc, title,
            font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_before=12, space_after=12,
        )

        subtitle = f"（{self._get('project_type', default='创业训练项目')}）"
        add_paragraph_with_format(
            self.doc, subtitle,
            font_name=FONT_HEI, font_size=SIZE_SAN, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_after=24,
        )

        for _ in range(2):
            self.doc.add_paragraph()

        info_items = [
            ("项目名称", self._get("project_name")),
            ("项目负责人", self._get("leader_name")),
            ("指导教师", self._get("advisor_name")),
            ("所在学院", self._get("college")),
            ("申报日期", self._get("apply_date")),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI,
                         font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True

        self.add_page_break()

    # --------------------------------------------------------
    # 一、基本信息
    # --------------------------------------------------------

    def _add_basic_info_table(self):
        """一、基本信息表"""
        self.add_h1("一、基本信息")
        basic = self._get("basic_info", default={})
        if not isinstance(basic, dict):
            basic = {}
        budget_str = str(basic.get("budget", self._get("budget_total", "")))
        if budget_str and not budget_str.endswith("元"):
            budget_str = f"{budget_str} 元"

        rows = [
            ["项目名称", basic.get("project_name", self._get("project_name"))],
            ["项目类型", basic.get("project_type",
                                    self._get("project_type", "创业训练项目"))],
            ["项目来源", basic.get("project_source", "A 学生自主选题")],
            ["所属行业", basic.get("industry", self._get("industry", ""))],
            ["起止时间", basic.get("duration", "")],
            ["申请经费", budget_str],
            ["负责人", basic.get("leader_info",
                                  f"{self._get('leader_name')} / "
                                  f"{self._get('leader_id')} / "
                                  f"{self._get('leader_major')} / "
                                  f"{self._get('leader_grade')} / "
                                  f"{self._get('leader_phone')}")],
            ["团队成员", basic.get("team_members", "")],
            ["指导教师", basic.get("advisor_info",
                                    f"{self._get('advisor_name')} / "
                                    f"{self._get('advisor_title')} / "
                                    f"{self._get('advisor_research')}")],
        ]
        self.add_table(["项目", "内容"], rows, col_widths=[4.5, 11.5])

    # --------------------------------------------------------
    # 二、项目简介与优势
    # --------------------------------------------------------

    def _add_abstract_and_advantages(self):
        """二、项目简介（200 字）+ 项目优势总结（10 条）"""
        self.add_h1("二、项目简介与优势总结")

        self.add_h2("（一）项目简介（200 字以内）")
        abstract = self._get("abstract", default="")
        if abstract:
            self.add_para(abstract)
        else:
            self.add_para("（请填写项目简介，200 字以内，按 4 句结构撰写："
                          "市场机会+做什么 / 怎么做+量化目标 / 商业模式+产出 / 现状。）")

        self.add_h2("（二）项目优势总结（10 条）")
        advantages = self._get("advantages", default=[])
        if isinstance(advantages, str):
            advantages = [advantages]
        if advantages:
            for i, adv in enumerate(advantages, 1):
                self.add_para(f"优势{self._cn_num(i)}：{adv}")
        else:
            self.add_para("（请填写 10 条项目优势，每条含'创新点名称：具体内容'格式。）")

    @staticmethod
    def _cn_num(n: int) -> str:
        cn = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
        if 0 < n <= 10:
            return cn[n]
        return str(n)

    # --------------------------------------------------------
    # 三、立项依据（含国家政策/科学挑战/文献综述）
    # --------------------------------------------------------

    def _add_project_basis(self):
        """三、项目的立项依据（国家规划+技术需求+科学挑战+研究意义+国内外研究现状+参考文献）"""
        self.add_h1("三、项目的立项依据")

        self.add_h2("1.1 国家规划与技术需求")
        self.add_h3("（1）国家规划（10+ 政策文件，按时间倒序）")
        policies = self._get("policy_citations", default=[])
        if isinstance(policies, str):
            policies = [policies]
        if policies:
            for p in policies:
                if not isinstance(p, dict):
                    continue
                date = p.get("date", "")
                agency = p.get("agency", "")
                doc_no = p.get("doc_number", "")
                title = p.get("title", "")
                excerpt = p.get("key_excerpt", "")
                relevance = p.get("relevance", "")
                line = f"{date}，{agency}印发《{title}》"
                if doc_no:
                    line += f"（{doc_no}）"
                line += f"，{excerpt}。{relevance}。"
                self.add_para(line)
        else:
            self.add_para("（请填写 10+ 政策文件，按时间倒序，每条含 4 要素："
                          "发文机关+文号+标题+时间，并摘录与项目相关的关键表述。）")

        self.add_h3("（2）技术需求")
        self.add_para(
            "随着我国城市化进程的加速，高层建筑、地下空间、大型综合体等复杂结构日益增多，"
            "给城市火灾应急救援带来了前所未有的挑战。在这些'信息孤岛'式的灾害现场，迫切"
            "需要发展新型技术手段，以应对传统侦察模式面临的严峻技术瓶颈。"
        )
        self.add_para(
            "国家消防救援局数据显示，截至 2024 年 8 月，全国共接报高层建筑火灾 3.6 万起，"
            "共造成 203 人死亡，超过 2023 年全年总和；高层建筑火灾起数虽然仅占火灾起数的 5.4%，"
            "但死亡人数占 15%以上。该组数据说明，高层建筑火灾在发生数量上并非占比最高，"
            "但其致死风险和救援难度明显高于一般火灾，具有典型的'小比例、高危害'特征。"
        )

        # 科学挑战 3 段
        self._add_scientific_challenges()

        # 文献综述
        self._add_literature_review()

    def _add_scientific_challenges(self):
        """科学挑战 3 段（基于案例 2 §3.2）"""
        self.add_h2("1.2 科学挑战（3 个层层递进）")
        challenges = self._get("scientific_challenges", default=[])
        if isinstance(challenges, str):
            challenges = [challenges]
        if not challenges:
            self.add_para("（请填写 3 个科学挑战，层层递进（数据→算法→系统 / 感知→认知→决策），"
                          "每个含问题描述段 + 2 个子挑战 bullet。）")
            return
        for ch in challenges:
            if not isinstance(ch, dict):
                continue
            self.add_h3(f"科学挑战{self._cn_num(ch.get('no', 1))}：{ch.get('title', '')}")
            self.add_para(ch.get("description", ""))
            sub_list = ch.get("sub_challenges", [])
            for sub in sub_list:
                if not isinstance(sub, dict):
                    continue
                self.add_para(f"● {sub.get('title', '')}")
                self.add_para(sub.get("detail", ""))

    def _add_literature_review(self):
        """文献综述（研究意义+国内外研究现状+总结+参考文献）"""
        lit = self._get("literature_review", default={})
        if not isinstance(lit, dict):
            lit = {}

        self.add_h2("1.3 研究意义与国内外研究现状")

        self.add_h3("1.3.1 研究意义")
        sig = lit.get("research_significance", [])
        if isinstance(sig, str):
            sig = [sig]
        if sig:
            for i, para in enumerate(sig, 1):
                self.add_para(f"（{self._cn_num(i)}）{para[:50]}…" if len(para) > 50 else f"（{self._cn_num(i)}）{para}")
                self.add_para(para)
        else:
            self.add_para("（请填写研究意义，3-5 段，每段一个小方向。）")

        self.add_h3("1.3.2 国内外研究现状")
        status = lit.get("research_status", [])
        if isinstance(status, str):
            status = [status]
        if status:
            for i, sec in enumerate(status, 1):
                if not isinstance(sec, dict):
                    continue
                self.add_para(f"（{self._cn_num(i)}）{sec.get('topic', '')}")
                for para in sec.get("paragraphs", []):
                    self.add_para(para)
        else:
            self.add_para("（请填写国内外研究现状，按 3 个技术方向分段，每段引用 3-5 篇文献。）")

        self.add_h3("（3）对现有研究现状的总结与分析")
        summary = lit.get("summary_and_gap", "")
        if summary:
            self.add_para(summary)
        else:
            self.add_para("（请填写对现有研究现状的总结与分析，指出 gap 与本课题攻关方向。）")

        self.add_h3("参考文献")
        refs = lit.get("references", [])
        if isinstance(refs, str):
            refs = [refs]
        if refs:
            for r in refs:
                if not isinstance(r, dict):
                    continue
                no = r.get("no", "")
                text = r.get("text", "")
                self.add_para(f"[{no}] {text}", indent=False)
        else:
            self.add_para("（请填写参考文献，≥ 30 篇（国家级）/ ≥ 20 篇（省级）/ ≥ 10 篇（校级），"
                          "GB/T 7714-2015 格式，英文 ≥ 60%。）", indent=False)

    # --------------------------------------------------------
    # 四、研究内容、研究目标、关键科学问题
    # --------------------------------------------------------

    def _add_research_content_section(self):
        """四、研究内容、研究目标、关键科学问题"""
        self.add_h1("四、项目的研究内容、研究目标，以及拟解决的关键科学问题")
        ts = self._get("technical_solution", default={})
        if not isinstance(ts, dict):
            ts = {}

        self.add_h2("2.1 硬件基础")
        hw = ts.get("hardware_basis", [])
        if isinstance(hw, str):
            hw = [hw]
        for i, h in enumerate(hw, 1):
            if not isinstance(h, dict):
                continue
            self.add_h3(f"2.1.{i} {h.get('name', '')}")
            self.add_para(f"硬件组成：{h.get('composition', '')}")
            self.add_para(f"应用场景：{h.get('application', '')}")

        self.add_h2("2.2 研究内容（3 个层层递进部分）")
        contents = ts.get("research_content", [])
        if isinstance(contents, str):
            contents = [contents]
        for c in contents:
            if not isinstance(c, dict):
                continue
            self.add_h3(f"（{c.get('no', 1)}）{c.get('title', '')}")
            algos = c.get("algorithms", [])
            if algos:
                self.add_para("核心算法：" + "、".join(algos))

        self.add_h2("2.3 研究目标（3 个核心目标）")
        goals = ts.get("research_goals", [])
        if isinstance(goals, str):
            goals = [goals]
        for g in goals:
            if not isinstance(g, dict):
                continue
            self.add_h3(f"（{g.get('no', 1)}）{g.get('goal', '')}")
            self.add_para(f"参数指标：{g.get('metrics', '')}")

        self.add_h2("2.4 拟解决关键科学问题（3 个）")
        problems = ts.get("key_scientific_problems", [])
        if isinstance(problems, str):
            problems = [problems]
        for p in problems:
            if not isinstance(p, dict):
                continue
            self.add_h3(f"（{p.get('no', 1)}）{p.get('problem', '')}")
            self.add_para(f"解决方案：{p.get('solution', '')}")

    # --------------------------------------------------------
    # 五、研究方案（算法对比表+技术路线图+数学公式+创新点）
    # --------------------------------------------------------

    def _add_research_plan(self):
        """五、拟采取的研究方案（算法对比+技术路线图+公式+创新点）"""
        self.add_h1("五、拟采取的研究方案")
        ts = self._get("technical_solution", default={})
        if not isinstance(ts, dict):
            ts = {}

        self.add_h2("3.1 研究方法和思路")
        self.add_para(
            "本申请旨在研究和开发一套集多模态智能感知、鲁棒三维重建、目标检测与态势理解"
            "耦合以及自适应路径规划于一体的消防无人机火场救援创新系统。本课题将综合模拟"
            "计算、实验和测试分析等多种研究方法结合开展研究工作。"
        )

        # 算法对比表
        self._add_algorithm_comparison(ts.get("algorithm_comparison", {}))

        # 技术路线图
        self._add_tech_roadmap(ts.get("tech_roadmap", []))

        # 数学公式
        self._add_formulas(ts.get("formulas", []))

        # 创新点
        self.add_h2("3.4 本项目的特色与创新之处")
        innovations = ts.get("innovations", [])
        if isinstance(innovations, str):
            innovations = [innovations]
        if innovations:
            for i, inn in enumerate(innovations, 1):
                self.add_para(f"（{self._cn_num(i)}）{inn}")
        else:
            self.add_para("（请填写 3 条创新点，每条含技术闭环描述。）")

    def _add_algorithm_comparison(self, ac_data):
        """算法对比表（基于案例 2 表 1）"""
        self.add_h2("3.2 算法对比表")
        if not isinstance(ac_data, dict) or not ac_data:
            self.add_para("（请填写算法对比表，≥ 3 算法 × ≥ 4 维度，含选型结论。）")
            return

        title = ac_data.get("title", "算法对比表")
        self.add_para(title, indent=False)
        algorithms = ac_data.get("algorithms", [])
        dimensions = ac_data.get("dimensions", [])
        rows_data = ac_data.get("rows", [])

        if algorithms and dimensions and rows_data:
            headers = ["维度"] + algorithms
            rows = []
            for i, dim in enumerate(dimensions):
                if i < len(rows_data):
                    rows.append([dim] + list(rows_data[i]))
            self.add_table(headers, rows)
        else:
            self.add_para("（请填写算法对比表，≥ 3 算法 × ≥ 4 维度。）")

        conclusion = ac_data.get("conclusion", "")
        if conclusion:
            self.add_para(f"选型结论：{conclusion}")

    def _add_tech_roadmap(self, roadmap):
        """技术路线图（≥ 3 张，含图号+标题+描述）"""
        self.add_h2("3.3 技术路线图（3 张）")
        if isinstance(roadmap, str):
            roadmap = [roadmap]
        if not roadmap:
            self.add_para("（请填写技术路线图，≥ 3 张：研究内容关系图+研究方法图+技术路线图。）")
            return
        for r in roadmap:
            if not isinstance(r, dict):
                continue
            fig_no = r.get("fig_no", "")
            title = r.get("title", "")
            desc = r.get("description", "")
            self.add_para(f"{fig_no} {title}", indent=False)
            self.add_para(desc)

    def _add_formulas(self, formulas):
        """数学公式（每个核心算法 ≥ 1 公式）"""
        self.add_h2("3.5 核心算法数学公式（每个核心算法 ≥ 1 公式）")
        if isinstance(formulas, str):
            formulas = [formulas]
        if not formulas:
            self.add_para("（请填写数学公式，每个核心算法 ≥ 1 公式，含编号+变量定义。）")
            return
        for f in formulas:
            if not isinstance(f, dict):
                continue
            self.add_h3(f"算法：{f.get('algorithm', '')}")
            for formula in f.get("formulas", []):
                if not isinstance(formula, dict):
                    continue
                no = formula.get("no", "")
                expr = formula.get("expression", "")
                variables = formula.get("variables", "")
                self.add_para(f"{no}：{expr}", indent=False)
                self.add_para(f"其中：{variables}")

    # --------------------------------------------------------
    # 六、社会经济效益（12 项 + 社会/政治效益）
    # --------------------------------------------------------

    def _add_economic_benefits_section(self):
        """六、社会经济效益分析"""
        self.add_h1("六、社会经济效益分析")

        self.add_h2("6.1 综合效益概述")
        self.add_para(
            "该项目的研发与应用将在社会、经济、政治与战略等多个维度产生显著的综合效益。"
            "通过对比系统模式与传统模式在火灾应急救援全流程各阶段的表现，系统模式凭借"
            "'多模态融合透视烟雾'的核心技术突破，有效破解了传统模式面临的'信息孤岛'与"
            "'复杂环境受限'困境。"
        )

        self.add_h2("6.2 社会效益")
        self.add_para("（1）降低消防员暴露风险，践行'生命至上'：系统通过无人机替代消防员执行高危侦察任务，配合安全约束路径规划为地面人员提供精准引导，从根本上降低救援人员的职业风险。")
        self.add_para("（2）提升被困人员生还概率：系统的小目标智能识别（召回率 ≥ 92%）与三维实时重构能力，可将救援响应时间缩短至传统模式的 30%-50%，在'黄金救援期'内最大限度提升生存几率。")
        self.add_para("（3）推动消防救援现代化转型：构建'感知-认知-决策-行动'智能闭环，推动消防救援从'经验驱动'向'数据驱动'转型。")

        self.add_h2("6.3 经济效益（量化评估表 12 项）")
        benefits = self._get("economic_benefits", default=[])
        if isinstance(benefits, str):
            benefits = [benefits]
        if benefits:
            rows = []
            for b in benefits:
                if not isinstance(b, dict):
                    continue
                rows.append([
                    b.get("indicator", ""),
                    b.get("traditional", ""),
                    b.get("expected", ""),
                    b.get("improvement", ""),
                ])
            self.add_table(
                ["评估指标", "传统模式基准值", "系统模式预期值", "效益提升幅度"],
                rows,
                col_widths=[4.5, 3.5, 3.5, 4.5],
            )
        else:
            self.add_para("（请填写经济效益量化评估表，≥ 12 项，每项 4 列：指标名/传统基准/系统预期/提升幅度。）")

        self.add_h2("6.4 政治与战略效益")
        self.add_para("（1）构建'空地一体'应急体系，提升城市韧性：契合国家'智慧消防'与'低空经济'战略部署，推动城市应急管理智能化升级。")
        self.add_para("（2）服务国家安全与社会稳定：践行'人民至上、生命至上'理念，保障人民生命财产安全，维护社会和谐稳定。")
        self.add_para("（3）技术辐射与迁移潜力：核心技术可迁移至地震搜救、危化品处置、森林防火等领域，全面赋能应急救援体系智能化升级。")

    # --------------------------------------------------------
    # 七、市场分析（5 小节）
    # --------------------------------------------------------

    def _add_market_analysis_section(self):
        """七、市场分析（5 小节）"""
        self.add_h1("七、市场分析")
        ma = self._get("market_analysis", default={})
        if not isinstance(ma, dict):
            ma = {}

        self.add_h2("（一）市场需求")
        demand = ma.get("market_demand", "")
        self.add_para(demand if demand else
                      "（请填写市场需求，含高层建筑、地下空间、大型综合体等复杂场景增多带来的消防救援需求。）")

        self.add_h2("（二）目标市场")
        target = ma.get("target_market", "")
        self.add_para(target if target else
                      "（请填写目标市场，B/G 端专业客户为主：消防救援部门、应急管理机构、政府相关单位、园区安全管理单位。）")

        self.add_h2("（三）市场前景")
        prospect = ma.get("market_prospect", "")
        self.add_para(prospect if prospect else
                      "（请填写市场前景，智慧消防、应急管理现代化、低空经济和智能装备发展。）")

        self.add_h2("（四）产品或服务前景")
        product_p = ma.get("product_prospect", "")
        self.add_para(product_p if product_p else
                      "（请填写产品或服务前景，含设备交付、系统部署、操作培训、日常维护、算法升级、应急演练支持。）")

        self.add_h2("（五）SWOT 分析")
        swot = ma.get("swot", {})
        if not isinstance(swot, dict):
            swot = {}
        self.add_para(f"优势（Strengths）：{swot.get('strengths', '')}")
        self.add_para(f"劣势（Weaknesses）：{swot.get('weaknesses', '')}")
        self.add_para(f"机会（Opportunities）：{swot.get('opportunities', '')}")
        self.add_para(f"威胁（Threats）：{swot.get('threats', '')}")

    # --------------------------------------------------------
    # 八、营销模式（4P + 3C + 促销 + 供应链）
    # --------------------------------------------------------

    def _add_marketing_section(self):
        """八、营销模式（4P + 3C 定价 + 促销 + 供应链）"""
        self.add_h1("八、营销模式")

        self._add_marketing_4p()
        self._add_pricing_3c()
        self._add_promotion_strategy()
        self._add_supply_chain()

    def _add_marketing_4p(self):
        """4P 营销组合"""
        self.add_h2("8.1 4P 营销组合")
        m4p = self._get("marketing_4p", default={})
        if not isinstance(m4p, dict):
            m4p = {}

        self.add_h3("（1）产品策略（Product）")
        product = m4p.get("product", {})
        if isinstance(product, dict):
            self.add_para(f"核心价值：{product.get('core_value', '')}")
            self.add_para(f"产品布局：{product.get('layout', '')}")
            self.add_para(f"产品设计：{product.get('design', '')}")
            self.add_para(f"产品生命周期策略：{product.get('lifecycle', '')}")
        else:
            self.add_para("（请填写产品策略：核心价值+产品布局+产品设计+产品生命周期。）")

        self.add_h3("（2）价格策略（Price）")
        price = m4p.get("price", {})
        if isinstance(price, dict):
            self.add_para(f"定价分析：{price.get('pricing_analysis', '')}")
            adjustments = price.get("adjustment", [])
            if isinstance(adjustments, list):
                for adj in adjustments:
                    self.add_para(f"• {adj}")
        else:
            self.add_para("（请填写价格策略：定价分析+价格调整策略。）")

        self.add_h3("（3）渠道策略（Place）")
        place = m4p.get("place", {})
        if isinstance(place, dict):
            self.add_para(f"线下实体店营销：{place.get('offline', '')}")
            self.add_para(f"网络营销：{place.get('online', '')}")
        else:
            self.add_para("（请填写渠道策略：线下实体店营销+网络营销。）")

        self.add_h3("（4）促销策略（Promotion）")
        promotion = m4p.get("promotion", {})
        if isinstance(promotion, dict):
            ad = promotion.get("advertising", {})
            if isinstance(ad, dict):
                self.add_para(f"理性主题：{ad.get('rational', '')}")
                self.add_para(f"情感主题：{ad.get('emotional', '')}")
            pr_list = promotion.get("pr", [])
            if isinstance(pr_list, list) and pr_list:
                self.add_para("社会公共关系促销方式：" + "、".join(pr_list))
            sp = promotion.get("sales_promotion", {})
            if isinstance(sp, dict):
                for audience, tools in sp.items():
                    if isinstance(tools, str):
                        self.add_para(f"{audience}：{tools}")
        else:
            self.add_para("（请填写促销策略：广告促销+社会公共关系促销+销售促进。）")

    def _add_pricing_3c(self):
        """3C 定价模型"""
        self.add_h2("8.2 3C 定价模型")
        p3c = self._get("pricing_3c", default={})
        if not isinstance(p3c, dict):
            p3c = {}

        for key in ["cost", "competition", "customer"]:
            c = p3c.get(key, {})
            if isinstance(c, dict):
                self.add_h3(f"（{key}）{c.get('name', '')}")
                self.add_para(f"描述：{c.get('description', '')}")
                self.add_para(f"公式：{c.get('formula', '')}")
            else:
                self.add_para(f"（请填写 {key} 维度的 3C 定价模型。）")

        final_strategy = p3c.get("final_strategy", "")
        if final_strategy:
            self.add_para(f"最终定价策略：{final_strategy}")

    def _add_promotion_strategy(self):
        """促销策略详述"""
        self.add_h2("8.3 促销策略详述")
        self.add_h3("（1）广告促销")
        self.add_para("理性主题：突出体现消防救援无人机群及多机协同系统的独特性，展示本产品对比传统消防速度效果的优势。")
        self.add_para("情感主题：通过讲述真实的火灾案例和消防救援无人机群的成功应用，让大众了解产品在实际应用中的效果和价值。")

        self.add_h3("（2）社会公共关系促销")
        pr_rows = [
            ["记者招待会", "企业专题新闻报道"],
            ["制造事件", "行业研讨会"],
            ["公益广告", "专家联系制度"],
            ["专题讨论", "企业形象识别系统"],
            ["公司出版物", "公共活动赞助"],
            ["慈善活动", "社区公益活动"],
            ["社区活动", "公司开业与周年庆典"],
            ["宣传手册", "消费者座谈"],
            ["公司开放参观", "联动营销"],
        ]
        self.add_table(["公共关系方式 1", "公共关系方式 2"], pr_rows, col_widths=[8.0, 8.0])

        self.add_h3("（3）销售促进")
        sp_rows = [
            ["消费者", "材料样品、优惠卷、现金折扣、赠品、惠顾汇报、产品保证、连带促销"],
            ["中间商", "赠品或礼品、现金折扣、批量折扣、促销折让、销售竞赛、广告津贴、人员培训、联合促销"],
            ["社会大众", "商业展览和会议、特殊广告品"],
            ["销售人员", "销售提成、奖金、销售竞赛、免费旅游"],
        ]
        self.add_table(["促销对象", "促销方式"], sp_rows, col_widths=[3.0, 13.0])

    def _add_supply_chain(self):
        """供应链优化（5 子节）"""
        self.add_h2("8.4 供应链优化")
        fs = self._get("financial_statements", default={})
        sc = fs.get("supply_chain", {}) if isinstance(fs, dict) else {}
        if not isinstance(sc, dict):
            sc = {}
        subs = sc.get("subsections", [])
        if isinstance(subs, list) and subs:
            for sub in subs:
                if not isinstance(sub, dict):
                    continue
                self.add_h3(sub.get("name", ""))
                self.add_para(sub.get("content", ""))
        else:
            self.add_h3("8.4.1 供应链管理")
            self.add_para("建立可靠的供应商网络，与供应商建立长期合作关系，定期评估绩效。")
            self.add_h3("8.4.2 库存管理")
            self.add_para("采用 RFID 技术、电子标签、动态化管理、规范存储保养、出入库管理制度。")
            self.add_h3("8.4.3 物流优化")
            self.add_para("优化运输、仓储和配送，采用物联网、大数据分析和人工智能，多式联运。")
            self.add_h3("8.4.4 信息共享")
            self.add_para("采用 ERP、SCM 系统实现订单管理、数据交换和信息共享，物联网技术全程可视。")

    # --------------------------------------------------------
    # 九、财务分析（5 年三表 10 张）
    # --------------------------------------------------------

    def _add_financial_section(self):
        """九、财务分析（资金筹备+5 年三表 10 张+利润+风险+退出）"""
        self.add_h1("九、财务分析")
        fs = self._get("financial_statements", default={})
        if not isinstance(fs, dict):
            fs = {}

        self.add_h2("（一）资金筹备")
        cp = fs.get("capital_preparation", {})
        if isinstance(cp, dict):
            self.add_para(f"注册资本：{cp.get('registered_capital', '')}")
            self.add_para(f"创业团队自筹：{cp.get('self_funding', '')}")
            self.add_para(f"天使投资：{cp.get('angel_investment', '')}")
            self.add_para(f"银行贷款：{cp.get('bank_loan', '')}")
            self.add_para(f"初期投资计划：{cp.get('initial_investment', '')}")
        else:
            self.add_para("（请填写资金筹备方案。）")

        self.add_h2("（二）固定资产明细")
        self.add_para(fs.get("fixed_assets", "（请填写固定资产明细，初期购建固定资产支出约 10 万元，按年限平均法折旧，使用年限 5 年。）"))

        self.add_h2("（三）流动资产明细")
        self.add_para(fs.get("current_assets", "（请填写流动资产明细。）"))

        # 10 张表
        self._add_financial_table(fs.get("sales_revenue_table", {}), "（四）")
        self._add_financial_table(fs.get("cost_table", {}), "（五）")
        self._add_financial_table(fs.get("admin_expense_table", {}), "（六）")
        self._add_financial_table(fs.get("sales_expense_table", {}), "（七）")
        self._add_financial_table(fs.get("cash_flow_table", {}), "（八）")
        self._add_financial_table(fs.get("npv_table", {}), "（九）")
        self._add_financial_table(fs.get("profitability_table", {}), "（十）")
        self._add_financial_table(fs.get("balance_sheet", {}), "（十一）")
        self._add_financial_table(fs.get("short_term_solvency_table", {}), "（十二）")
        self._add_financial_table(fs.get("long_term_solvency_table", {}), "（十三）")

        # NPV 结论
        npv_data = fs.get("npv_table", {})
        if isinstance(npv_data, dict):
            conclusion = npv_data.get("conclusion", "")
            if conclusion:
                self.add_para(f"NPV 结论：{conclusion}")

        self.add_h2("（十四）利润预计")
        self.add_para(fs.get("profit_forecast", "（请填写利润预计，5 年营业收入与净利润。）"))

        self.add_h2("（十五）风险分析")
        self.add_para(fs.get("risk_analysis_finance", "（请填写财务风险分析。）"))

        self.add_h2("（十六）退出策略")
        self.add_para(fs.get("exit_strategy", "（请填写退出策略。）"))

        self.add_h2("（十七）盈利模式")
        pm = fs.get("profit_model", {})
        if isinstance(pm, dict):
            channels = pm.get("channels", [])
            if isinstance(channels, list) and channels:
                for ch in channels:
                    self.add_para(f"• {ch}")
            else:
                self.add_para("（请填写盈利模式：整机销售/租赁/专业服务/保养维护。）")
        else:
            self.add_para("（请填写盈利模式。）")

    def _add_financial_table(self, table_data, section_no):
        """渲染单张财务表"""
        if not isinstance(table_data, dict) or not table_data:
            return
        title = table_data.get("title", "")
        headers = table_data.get("headers", [])
        rows_data = table_data.get("rows", [])
        if not headers or not rows_data:
            return
        self.add_h2(f"{section_no} {title}")
        self.add_table(headers, [list(r) for r in rows_data])

    # --------------------------------------------------------
    # 十、实践过程与技术壁垒
    # --------------------------------------------------------

    def _add_practice_and_barriers(self):
        """十、实践过程与技术壁垒"""
        self.add_h1("十、实践过程与技术壁垒")

        self.add_h2("10.1 实践过程（≥ 5 张照片描述）")
        pp = self._get("practice_process", default={})
        if not isinstance(pp, dict):
            pp = {}
        photos = pp.get("photos", [])
        if isinstance(photos, str):
            photos = [photos]
        if photos:
            for photo in photos:
                if not isinstance(photo, dict):
                    continue
                fig_no = photo.get("fig_no", "")
                title = photo.get("title", "")
                desc = photo.get("description", "")
                self.add_para(f"{fig_no} {title}", indent=False)
                self.add_para(desc)
        else:
            self.add_para("（请填写实践过程照片描述，≥ 5 张，含图号+标题+说明，4 类活动齐全：学习/考察/制作/试飞。）")

        self.add_h2("10.2 合作协议（≥ 3 家）")
        agreements = pp.get("cooperation_agreements", [])
        if isinstance(agreements, str):
            agreements = [agreements]
        if agreements:
            for ag in agreements:
                if not isinstance(ag, dict):
                    continue
                fig_no = ag.get("fig_no", "")
                partner = ag.get("partner", "")
                content = ag.get("content", "")
                self.add_para(f"{fig_no} 与 {partner} 签订合作协议", indent=False)
                self.add_para(f"合作内容：{content}")
        else:
            self.add_para("（请填写合作协议，≥ 3 家，类型多元：学校内部+企业+行业机构。）")

        self.add_h2("10.3 技术壁垒（专利 + 软件检测 + 实物验证）")
        tb = self._get("tech_barriers", default={})
        if not isinstance(tb, dict):
            tb = {}

        self.add_h3("（1）已申请/已授权专利（≥ 1 项）")
        patents = tb.get("patents", [])
        if isinstance(patents, str):
            patents = [patents]
        if patents:
            for pat in patents:
                if not isinstance(pat, dict):
                    continue
                fig_no = pat.get("fig_no", "")
                name = pat.get("name", "")
                status = pat.get("status", "")
                patent_no = pat.get("patent_no", "")
                self.add_para(f"{fig_no} {name}", indent=False)
                self.add_para(f"专利状态：{status}；专利号：{patent_no}")
        else:
            self.add_para("（请填写已申请/已授权专利，≥ 1 项，含专利名称+状态+专利号。）")

        self.add_h3("（2）软件检测报告")
        reports = tb.get("test_reports", [])
        if isinstance(reports, str):
            reports = [reports]
        if reports:
            for rep in reports:
                if not isinstance(rep, dict):
                    continue
                fig_no = rep.get("fig_no", "")
                name = rep.get("name", "")
                issuer = rep.get("issuer", "")
                result = rep.get("result", "")
                self.add_para(f"{fig_no} {name}", indent=False)
                self.add_para(f"检测机构：{issuer}；检测结果：{result}")
        else:
            self.add_para("（请填写软件检测报告，≥ 1 份。）")

        self.add_h3("（3）实物验证截图")
        validations = tb.get("prototype_validation", [])
        if isinstance(validations, str):
            validations = [validations]
        if validations:
            for val in validations:
                if not isinstance(val, dict):
                    continue
                fig_no = val.get("fig_no", "")
                name = val.get("name", "")
                desc = val.get("description", "")
                self.add_para(f"{fig_no} {name}", indent=False)
                self.add_para(desc)
        else:
            self.add_para("（请填写实物验证截图描述，≥ 1 张。）")

    # --------------------------------------------------------
    # 十一、项目进度安排（4 阶段甘特）
    # --------------------------------------------------------

    def _add_project_schedule(self):
        """十一、项目进度安排（4 阶段甘特图）"""
        self.add_h1("十一、项目进度安排")

        self.add_h2("1. 项目计划安排（4 阶段）")
        schedule = self._get("project_schedule", default=[])
        if isinstance(schedule, str):
            schedule = [schedule]
        if schedule:
            for phase in schedule:
                if not isinstance(phase, dict):
                    continue
                phase_name = phase.get("phase", "")
                period = phase.get("period", "")
                tasks = phase.get("tasks", [])
                self.add_h3(f"{phase_name}：{period}")
                if isinstance(tasks, list):
                    for task in tasks:
                        self.add_para(f"• {task}", indent=False)
        else:
            self.add_para("（请填写 4 阶段甘特图：第一阶段调研设计/第二阶段硬件算法/第三阶段集成验证/第四阶段总结结题。）")

        self.add_h2("2. 项目目前进展")
        self.add_para(
            "目前该项目已经与河南工业大学安保科签订合作巡逻协议，主要是负责校园火灾防卫安全巡逻，"
            "同时还与杭州英诺唯、郑州新鼎自动化、河南瞰宇科技等智能装备企业签订战略合作协议。"
            "原型机组装已完成，首次试飞验证通过，火灾监测识别算法原型在模拟场景下达到预期技术指标。"
        )

    # --------------------------------------------------------
    # 十二、经费预算
    # --------------------------------------------------------

    def _add_budget(self):
        """十二、经费预算（3 列表格）"""
        self.add_h1("十二、经费预算")
        items = self._get("budget_items", default=[])
        if not items:
            self.add_para("（请填写经费预算，6 类标准科目：硬件设备/传感器模块/实验耗材/软件平台/调研差旅/成果整理。每项金额非整数，附计算依据。）")
            return
        rows = []
        total = 0
        for b in items:
            if not isinstance(b, dict):
                continue
            try:
                amount_num = int(str(b.get("amount", "0")))
            except ValueError:
                amount_num = 0
            total += amount_num
            rows.append([b.get("item", ""), f"{amount_num} 元", b.get("basis", "")])
        rows.append(["合计", f"{total} 元", ""])
        self.add_table(["预算科目", "金额", "计算依据"], rows, col_widths=[3.5, 3.0, 9.5])

    # --------------------------------------------------------
    # 十三、团队介绍
    # --------------------------------------------------------

    def _add_team_intro(self):
        """十三、团队介绍"""
        self.add_h1("十三、团队介绍")
        team = self._get("team_intro", default={})
        if not isinstance(team, dict):
            team = {}

        self.add_h2("（一）团队构成")
        members = team.get("members", [])
        if members:
            rows = [[m.get("name", ""), m.get("id", ""), m.get("major", ""),
                     m.get("role", ""), m.get("exp", "")]
                    for m in members if isinstance(m, dict)]
            self.add_table(["姓名", "学号", "专业年级", "分工", "相关经历"],
                           rows, col_widths=[1.8, 2.5, 3.0, 3.0, 5.7])
        else:
            self.add_para("（请填写团队成员表：姓名/学号/专业年级/分工/相关经历，4~5 人。）")

        self.add_h2("（二）指导教师背景")
        adv = team.get("advisor_bg", "")
        self.add_para(adv if adv else
                      "（请填写指导教师背景：职称、研究方向、主持项目、指导学生创业经历。）")

    # --------------------------------------------------------
    # 十四、风险分析
    # --------------------------------------------------------

    def _add_risk_analysis(self):
        """十四、风险分析（6 类风险）"""
        self.add_h1("十四、风险预期与应对")
        risks = self._get("risk_analysis", default=[])
        if isinstance(risks, str):
            risks = [risks]
        if risks:
            rows = [[r.get("type", ""), r.get("risk", ""), r.get("prob", ""),
                     r.get("impact", ""), r.get("measure", "")]
                    for r in risks if isinstance(r, dict)]
            self.add_table(["风险类型", "具体风险", "概率", "影响", "应对措施"],
                           rows, col_widths=[2.2, 3.8, 1.5, 1.5, 7.0])
        else:
            self.add_para("（请填写风险分析表，6 类风险齐全：资产/竞争/财务/管理/技术/破产，每类含具体风险+概率+影响+应对措施。）")

    # --------------------------------------------------------
    # 十五、预期成果
    # --------------------------------------------------------

    def _add_expected_results(self):
        """十五、预期成果"""
        self.add_h1("十五、预期成果")
        outcomes = self._get("expected_outcomes", default=[])
        if isinstance(outcomes, str):
            outcomes = [outcomes]
        if outcomes:
            for o in outcomes:
                self.add_para(f"• {o}", indent=False)
        else:
            self.add_para("（请填写预期成果，以商业计划书+技术报告+样机+专利+软著为主，每项含数量+形态+验收标准。）", indent=False)

    # --------------------------------------------------------
    # 十六、签字栏
    # --------------------------------------------------------

    def _add_signature_section(self):
        """十六、签字栏（指导教师+学院+学校）"""
        sections = [("十六、指导教师意见", "指导教师签字：____________________    日期：______年____月____日")]
        sections.append(("十七、学院评审意见", "学院盖章：____________________    日期：______年____月____日"))
        if self._get("include_school_approval", default=False):
            sections.append(("十八、学校审批意见", "学校盖章：____________________    日期：______年____月____日"))
        for title, line in sections:
            self.add_h1(title)
            for _ in range(6):
                self.doc.add_paragraph()
            self.add_para(line, indent=False)

    # --------------------------------------------------------
    # 主构建方法
    # --------------------------------------------------------

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 16 章，生成 docx。返回实际保存路径。"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_cover()
            self._add_basic_info_table()
            self._add_abstract_and_advantages()
            self._add_project_basis()
            self._add_research_content_section()
            self._add_research_plan()
            self._add_economic_benefits_section()
            self._add_market_analysis_section()
            self._add_marketing_section()
            self._add_financial_section()
            self._add_practice_and_barriers()
            self._add_project_schedule()
            self._add_budget()
            self._add_team_intro()
            self._add_risk_analysis()
            self._add_expected_results()
            self._add_signature_section()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申报书已生成：{output_path}")
        return str(output_path)

    # --------------------------------------------------------
    # 数据校验
    # --------------------------------------------------------

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []
        p0_fields = [
            ("project_name", "项目名称"),
            ("leader_name", "负责人姓名"),
            ("advisor_name", "指导教师姓名"),
            ("college", "所在学院"),
        ]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        basic = self._get("basic_info", default={})
        if not isinstance(basic, dict):
            basic = {}
        if not basic.get("industry") and not self._get("industry"):
            warnings.append("缺少 所属行业（industry）")
        if not basic.get("duration") and not self._get("duration"):
            warnings.append("缺少 起止时间（duration）")

        for key, name in [
            ("abstract", "项目简介"),
            ("policy_citations", "国家政策引用"),
            ("scientific_challenges", "科学挑战"),
            ("literature_review", "文献综述"),
            ("technical_solution", "技术方案"),
            ("economic_benefits", "社会经济效益"),
            ("market_analysis", "市场分析"),
            ("marketing_4p", "4P 营销"),
            ("pricing_3c", "3C 定价"),
            ("financial_statements", "5 年三表"),
            ("practice_process", "实践过程"),
            ("tech_barriers", "技术壁垒"),
            ("project_schedule", "项目进度"),
        ]:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}），将使用占位文本")

        # 政策数量校验
        policies = self._get("policy_citations", default=[])
        if isinstance(policies, list) and len(policies) < 10:
            warnings.append(f"政策引用数量 {len(policies)} < 10，建议补充至 10+ 条")

        # 文献数量校验
        lit = self._get("literature_review", default={})
        if isinstance(lit, dict):
            refs = lit.get("references", [])
            if isinstance(refs, list) and len(refs) < 10:
                warnings.append(f"参考文献数量 {len(refs)} < 10，建议补充至 30+ 篇")

        # 经济效益指标校验
        eb = self._get("economic_benefits", default=[])
        if isinstance(eb, list) and len(eb) < 12:
            warnings.append(f"经济效益指标 {len(eb)} < 12，建议补充至 12 项")

        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据（消防无人机创业训练，对齐案例 2）
# ============================================================

DEFAULT_DATA = {
    "project_name": "基于多模态融合与三维重构的消防无人机火场救援系统",
    "project_level": "校级",
    "project_type": "创业训练项目",
    "leader_name": "姚奕晗",
    "leader_id": "241080200223",
    "leader_major": "电气工程及其自动化",
    "leader_grade": "2024 级",
    "leader_phone": "167XXXXXXXX",
    "advisor_name": "闫晶晶",
    "advisor_title": "教授/系主任",
    "advisor_research": "网络攻击检测方法及抗攻击量化控制策略",
    "advisor_achievements": "指导学生获挑战杯河南省特等奖、互联网+省二等奖，主持河南省高校科技创新人才支持计划（24HASTIT039）",
    "college": "电气工程学院",
    "apply_date": "2026 年 5 月 6 日",
    "industry": "智能制造/应急装备/低空经济",
    "duration": "2026.06-2027.06（12 个月）",
    "budget_total": "10000",
    "include_school_approval": False,

    "basic_info": {
        "project_name": "基于多模态融合与三维重构的消防无人机火场救援系统",
        "project_type": "创业训练项目",
        "project_source": "A 学生自主选题",
        "industry": "智能制造/应急装备/低空经济",
        "duration": "2026.06-2027.06（12 个月）",
        "budget": "10000",
        "leader_info": "姚奕晗 / 241080200223 / 电气工程及其自动化 / 2024 级 / 167XXXXXXXX",
        "team_members": "周匡吉（算法开发）、韩兆峰（硬件设计）、李浩博（系统集成）、吴继航（测试验证）",
        "advisor_info": "闫晶晶 / 教授/系主任 / 网络攻击检测与抗攻击量化控制 / 186XXXXXXXX",
    },

    "abstract": (
        "本项目面向复杂城市火灾救援需求，拟研发基于多模态融合与三维重构的消防无人机火场救援系统。"
        "系统融合可见光、红外热成像、激光雷达点云和气体传感等多源信息，实现火源识别、被困人员搜索、"
        "三维火场重构、危险区域判断和辅助路径规划，为消防指挥提供实时态势感知与辅助决策支持。"
        "项目可应用于高层建筑、地下空间、仓储物流、石油化工等高危场景，降低消防员抵近侦察风险，"
        "提升救援效率。"
    ),

    "advantages": [
        "多模态火场感知创新：项目融合可见光图像、红外热成像、激光雷达点云、气体传感等多源数据，突破单一传感器在浓烟、弱光、高温、遮挡环境下感知能力不足的问题，提高火场信息获取的全面性和准确性。",
        "三维火场重构创新：项目通过点云数据、图像信息和传感数据融合，构建实时更新的三维火场模型，实现对火场空间结构、火源位置、烟雾扩散、温度分布和危险区域的立体化表达。",
        "火场数字孪生创新：项目将三维重构结果与火场实时数据结合，建立火场数字孪生模型，使消防指挥人员能够更直观地掌握火场动态变化，为救援决策提供可视化支撑。",
        "消防无人机智能救援模式创新：项目构建'无人机平台 + 智能感知 + 三维重构 + 路径规划 + 辅助决策'的一体化救援系统，提升消防无人机在复杂火场中的实用价值。",
        "复杂环境下目标识别创新：针对火场中被困人员、火源、危险设施等目标易被烟雾、火光和遮挡干扰的问题，项目引入目标检测与语义分割技术，提高关键目标识别和定位能力。",
        "自适应路径规划创新：项目基于实时火场态势和三维环境信息，进行无人机安全路径规划，使无人机能够避开高温区、浓烟区、坍塌风险区和障碍物，提高飞行安全性。",
        "'感知—重构—决策'闭环创新：项目打通火场数据采集、三维建模、目标识别、态势判断和路径规划流程，形成从信息获取到辅助救援决策的完整技术闭环。",
        "硬件平台集成创新：项目可结合团队已有无人机相关专利成果，设计适用于消防救援场景的无人机硬件平台，实现传感器、通信模块、边缘计算单元和救援辅助装置的集成。",
        "边缘智能应用创新：项目考虑无人机端实时处理需求，探索轻量化算法和边缘计算部署方式，提升系统在火场现场的实时响应能力，减少对远程服务器的依赖。",
        "智慧消防应用场景创新：项目面向高层建筑、地下空间、仓储物流、石油化工等复杂场景，推动无人机、人工智能、数字孪生和消防救援装备融合应用，具有较强的智慧消防示范价值。",
    ],

    "policy_citations": [
        {"order": 1, "date": "2026 年 3 月", "agency": "国务院办公厅",
         "doc_number": "国办发〔2026〕XX 号", "title": "关于加强基层消防工作的意见",
         "key_excerpt": "强化基层消防科技支撑，推广智能化消防装备",
         "relevance": "为消防无人机下沉基层、提升一线处置能力提供政策导向与实施路径"},
        {"order": 2, "date": "2026 年 2 月", "agency": "国家层面",
         "doc_number": "", "title": "关于低空经济发展的相关部署",
         "key_excerpt": "低空飞行器在公共安全、应急救援、城市治理等领域的应用价值",
         "relevance": "为消防无人机在复杂城市空间中的常态化部署、快速响应和协同作业提供更加清晰的产业环境与应用场景"},
        {"order": 3, "date": "2025 年 12 月", "agency": "工业和信息化部等",
         "doc_number": "", "title": "应急装备产业重点产品发展指导目录（2025 版）",
         "key_excerpt": "将无人化、智能化、轻量化救援装备列为重点发展方向",
         "relevance": "体现出应急装备体系由传统单一装备向智能协同装备升级的趋势"},
        {"order": 4, "date": "2025 年 5 月", "agency": "国务院新闻办公室",
         "doc_number": "", "title": "新时代的中国国家安全（白皮书）",
         "key_excerpt": "加强对人工智能和数据安全等新兴领域的风险防控，构建综合的安全管理体系",
         "relevance": "为国家长期稳定发展提供有力保障"},
        {"order": 5, "date": "2025 年", "agency": "全国两会",
         "doc_number": "", "title": "大安全、大应急框架",
         "key_excerpt": "通过技术创新和数据应用，提升应急响应和快速处置能力",
         "relevance": "确保在各类灾害面前，能够实现高效的应急救援与有效的灾后恢复"},
        {"order": 6, "date": "2024 年 12 月", "agency": "中共中央办公厅、国务院办公厅",
         "doc_number": "", "title": "关于推进新型城市基础设施建设打造韧性城市的意见",
         "key_excerpt": "加快推动数字化、网络化、智能化的新型城市基础设施建设，提升城市的风险防控和治理能力，提高城市消防应急能力和智慧消防系统的建设",
         "relevance": "为消防无人机在城市消防应急中的应用提供政策支撑"},
        {"order": 7, "date": "2024 年 3 月", "agency": "工业和信息化部等四部门",
         "doc_number": "", "title": "通用航空装备创新应用实施方案（2024-2030 年）",
         "key_excerpt": "推动无人机在灾害勘察、应急通信、物资投送等场景应用",
         "relevance": "将'应急救援'与'低空经济'进行相关联"},
        {"order": 8, "date": "2023 年 12 月", "agency": "应急管理部、工业和信息化部",
         "doc_number": "", "title": "关于加快应急机器人发展的指导意见",
         "key_excerpt": "研发基于大载重无人机，无人机集群等装备的快速侦察与灭火技术及系统",
         "relevance": "应用于城市高层建筑火灾、地下有限空间等复杂场景"},
        {"order": 9, "date": "2022 年 6 月", "agency": "应急管理部",
         "doc_number": "", "title": "\"十四五\"应急救援力量建设规划",
         "key_excerpt": "加快构建大型固定翼灭火飞机、灭火直升机与无人机高低搭配、布局合理、功能互补的应急救援航空器体系",
         "relevance": "明确无人机在应急救援航空器体系中的定位"},
        {"order": 10, "date": "2021 年 12 月", "agency": "国务院",
         "doc_number": "", "title": "\"十四五\"国家应急体系规划",
         "key_excerpt": "推广运用智能机器人、无人机等高技术配送装备",
         "relevance": "提升复杂灾害环境下的态势感知与应急处置效率"},
    ],

    "scientific_challenges": [
        {
            "no": 1,
            "title": "多模态数据时空融合与三维火场实时重构计算",
            "description": "在复杂城市火灾场景中，消防无人机需要同时处理来自多源传感器[4]的异构数据，包括可见光图像、红外热成像、激光雷达点云、气体传感器等[5][6]。这些数据在时空维度上存在显著差异，如何实现多模态数据的高效时空融合[7]，并在算力受限的无人机平台上完成三维火场的实时重构计算，成为本项目面临的核心科学挑战[8]。",
            "sub_challenges": [
                {"title": "多模态数据的时空对齐与融合",
                 "detail": "不同传感器具有不同的采样频率、空间分辨率和物理特性，如何实现激光雷达点云、红外图像、可见光图像等多源数据在时间和空间维度上的精确对齐，是构建统一三维火场模型的基础难题[9]。例如，激光雷达点云数据通常以每秒数十万点的速率采集，而红外热成像的帧率可能仅为 30Hz，这种采样率差异会导致数据融合时的时空不一致性[10]。"},
                {"title": "三维火场实时重构计算",
                 "detail": "三维重构算法通常计算复杂度高，而火灾现场变化迅速，需要在无人机有限算力下实现毫秒级响应。如何设计轻量化重构算法，在保证精度的同时满足实时性要求，是技术实现的关键瓶颈[11][12]。例如，传统的 SLAM 算法在处理大规模点云数据时，计算量随场景复杂度呈指数增长，难以满足实时性需求[13]。"},
            ],
        },
        {
            "no": 2,
            "title": "火场数字孪生中目标检测与态势理解耦合技术",
            "description": "在成功获取到高质量的融合特征后，我们面临的核心挑战是如何在火场这一高噪声、高干扰的复杂场景中，实现目标检测与态势理解的深度耦合[14]。火场中的识别任务与常规场景有巨大差异：被困人员可能因遮挡或烟雾而呈现为模糊目标，而煤气罐等关键危险品可能在画面中占据极小区域，成为典型的小目标。",
            "sub_challenges": [
                {"title": "注意力机制与小目标特征增强",
                 "detail": "研究设计新的注意力机制，让模型能自动聚焦于红外特征提示的高温区域[15]；改造网络结构以增强对小目标的特征提取能力；并探索如何利用激光雷达提供的深度信息，辅助判断目标的真实尺寸与空间关系，从而解决遮挡问题[16]。"},
                {"title": "目标检测到态势理解的智能推理链路",
                 "detail": "构建一个评估模型，它能够综合目标类型、状态（如温度、气体浓度）及空间位置等多维度信息，自动对火场进行危险等级划分，并生成'高危区域预警'、'建议优先救援目标'等决策建议[17]。如何建立从'目标检测'到'态势理解'的智能推理链路，实现二者的有机耦合，是提升系统实战价值的决定性因素[18]。"},
            ],
        },
        {
            "no": 3,
            "title": "火灾场景的自适应路径规划与在线学习机制研究",
            "description": "在复杂多变的火灾环境中，无人机需要具备自主导航与智能决策能力，以应对火势蔓延、结构坍塌、烟雾扩散等动态威胁。传统路径规划算法通常基于静态环境假设，难以适应火场的高动态性和不确定性，这构成了本项目在自主导航层面的核心科学挑战。",
            "sub_challenges": [
                {"title": "动态环境下的实时路径规划与避障",
                 "detail": "火灾现场环境瞬息万变，烟雾浓度、温度分布、结构稳定性等参数均随时间快速演化。如何建立考虑多约束条件（包括热辐射强度、有毒气体浓度、结构承重能力等）的路径规划模型，并在计算资源受限的情况下实现实时避障与重规划，是技术实现的关键难点。"},
                {"title": "基于强化学习的在线自适应决策机制",
                 "detail": "针对火灾场景的高度不确定性，需要研究能够通过交互经验自主提升性能的在线学习算法。如何设计兼顾探索与利用的奖励函数，使无人机能够在保障安全的前提下最大化侦察效率；如何构建轻量化的神经网络模型，实现在边缘设备上的快速策略更新[19]；以及如何建立多机协同的经验共享机制，加速群体智能的涌现，都是需要深入探索的前沿问题。"},
            ],
        },
    ],

    "literature_review": {
        "research_significance": [
            "伴随着人工智能、大数据、物联网、边缘计算等技术的深度融合与广泛落地，消防无人机在城市消防安全领域的应用愈发关键且深入。在消防无人机参与火灾救援的过程中，对火场目标的智能感知与三维重构是核心环节，直接关系到救援决策的准确性与救援行动的高效性[20]。",
            "多模态融合与三维数字孪生技术，对消防无人机感知火场特别有用。消防无人机能装多种传感器，能收集视觉、温度、气体浓度这些各种数据。把这些数据融合起来，再建个三维模型，就比单个传感器强，能互相弥补不足[24]。",
            "在复杂多变的火灾救援任务中，无人机的自主导航与智能决策能力直接决定了其侦察、灭火与搜救的实战效能。自适应路径规划技术能够基于实时感知的火场三维模型与多模态数据，动态生成兼顾安全性与效率的最优飞行路径。",
        ],
        "research_status": [
            {"topic": "多模态数据深度融合技术的研究现状",
             "paragraphs": [
                 "在消防无人机对城市复杂空间的智能感知任务中，多模态数据融合技术近年有不少进展。但目前研究大多是'浅层融合'，比如简单把可见光、红外、激光雷达等不同模态的数据叠加或初步拼接特征[28]。",
                 "近年，研究领域开始聚焦深度融合机制的创新。部分学者提出基于注意力机制的多模态融合框架，让模型能自适应关注不同模态数据中对火场目标识别有帮助的信息[29]。",
             ],
             "citations": ["[28]", "[29]"]},
            {"topic": "火场小目标/模糊目标智能识别技术的研究现状",
             "paragraphs": [
                 "在城市高楼内部、地下设施这些复杂火场里，识别狭小通道里的被困人员、小型易燃易爆设备这类小目标，还有因烟雾、高温变得模糊的目标，是消防无人机救援效果的关键。",
                 "最近针对火场的算法改进在推进：一方面优化网络结构，比如用深度可分离卷积减少计算量，再结合金字塔特征融合模块；另一方面靠多模态融合，用红外图像的温度特征、激光雷达的结构特征来帮可见光图像识别模糊目标。",
             ],
             "citations": ["[29]"]},
            {"topic": "自适应路径规划与在线学习机制的研究现状",
             "paragraphs": [
                 "近年来，国内外研究开始转向具备环境响应能力的自适应路径规划方法。国际上有研究团队将实时感知的热成像数据与改进的 RRT*（快速探索随机树）算法结合[30]。",
                 "与此同时，在线学习机制为提升无人机在不确定性环境中的自主性提供了新的思路。强化学习，特别是深度强化学习（DRL）框架，被应用于无人机导航策略的端到端训练。",
             ],
             "citations": ["[30]"]},
        ],
        "summary_and_gap": (
            "多模态融合技术助力消防无人机提升高楼、地下空间等复杂火场的感知能力，助力救援，"
            "还能增强智能感知模型的环境适应性与通用性。不过，该技术研究虽有进展，却面临数据"
            "时空精准对齐难、火场部分模态数据质量不稳定、无人机边缘端实时融合与在线学习研究"
            "不足（尤其是小目标、模糊目标识别的实用性待提升）等挑战[31]。本课题聚焦'城市"
            "复杂空间消防无人机智能感知与三维火场重构'，从三方面攻关。"
        ),
        "references": [
            {"no": 1, "type": "journal_en",
             "text": "FERNANDES L, SILVA C. A 5G-enabled edge computing platform for intelligent emergency response with UAV networks[J]. Computer Networks, 2024, 242: 110267."},
            {"no": 2, "type": "journal_en",
             "text": "Tang J, Duan H, Lao S. Swarm intelligence algorithms for multiple unmanned aerial vehicles collaboration: A comprehensive review[J]. Artificial Intelligence Review, 2023, 56(5): 4295-4327."},
            {"no": 3, "type": "journal_en",
             "text": "Samadzadegan F, Toosi A, Dadrass Javan F. A critical review on multi-sensor and multi-platform remote sensing data fusion approaches: current status and prospects[J]. International journal of remote sensing, 2025, 46(3): 1327-1402."},
            {"no": 4, "type": "thesis_cn",
             "text": "基于多传感器信息融合的无人机火灾监测技术研究[D]. 南京航空航天大学, 2018."},
            {"no": 5, "type": "journal_en",
             "text": "Chen X, Zhu X, Liu C. Real-time 3D reconstruction of UAV acquisition system for the urban pipe based on RTAB-map[J]. Applied Sciences, 2023, 13(24): 13182."},
            {"no": 6, "type": "journal_en",
             "text": "Ghali R, Akhloufi M A, Mseddi W S. Deep learning and transformer approaches for UAV-based wildfire detection and segmentation[J]. Sensors, 2022, 22(5): 1977."},
            {"no": 7, "type": "journal_cn",
             "text": "李庚松, 刘艺, 郑奇斌, 等. 无人机多传感器数据融合研究综述[J]. 软件学报, 2025, 36(4): 1881-1905."},
            {"no": 8, "type": "journal_cn",
             "text": "王殿伟, 张新, 房杰, 等. 一种无人机航拍图像火灾烟雾检测算法[J]. 西安邮电大学学报, 2025, 30(02): 66-76."},
            {"no": 9, "type": "journal_en",
             "text": "Zhang H G, Ma S W, Li X, et al. Forest fire rescue framework to jointly optimize firefighting force configuration and facility layout: a case study of digital-twin simulation optimization[J]. Soft Computing, 2025, 29(3): 1789-1810."},
            {"no": 10, "type": "journal_en",
             "text": "Shaddy B, Ray D, Farguell A, et al. Generative algorithms for fusion of physics-based wildfire spread models with satellite data for initializing wildfire forecasts[J]. Artificial Intelligence for the Earth Systems, 2024, 3(3): e230087."},
            {"no": 11, "type": "journal_en",
             "text": "Zhang T, Ding F, Wang Z, et al. Forecasting backdraft with multimodal method: fusion of fire image and sensor data[J]. Engineering Applications of Artificial Intelligence, 2024, 132: 107939."},
            {"no": 12, "type": "journal_en",
             "text": "Tamanampudi V M. Application Optimizing AI Performance on Edge Devices: A Comprehensive Approach using Model Compression, Federated Learning, and Distributed Inference[J]. International Journal of Automation, Artificial Intelligence and Machine Learning, 2024, 4(2): 121-132."},
            {"no": 13, "type": "journal_en",
             "text": "Bhanushali D, Relyea R, Manghi K, et al. LiDAR-camera fusion for 3D object detection[J]. Electronic Imaging, 2020, 32: 1-9."},
            {"no": 14, "type": "journal_cn",
             "text": "基于人工智能技术的火灾探测信息融合系统[J]. 工业仪表与自动化装置, 2004(04)."},
            {"no": 15, "type": "journal_en",
             "text": "Xu X, Li C, Zhuge S, et al. A buddy temporal-spatial calibration method for airborne sensors in multi-UAV systems[J]. IEEE Robotics and Automation Letters, 2024, 9(8): 7365-7372."},
            {"no": 16, "type": "conference_en",
             "text": "Yang J, Liu X, Liu Z. Attention-guided Feature Fusion for Small Object Detection[C]//2023 IEEE International Conference on Imaging Systems and Techniques (IST). IEEE, 2023: 1-6."},
            {"no": 17, "type": "conference_en",
             "text": "Ma Y, Wei K, Liu F. Research on Visual Algorithm for Fire Detection of Firefighting UAVs Based on Infrared Imaging[C]//International conference on the Efficiency and Performance Engineering Network. Cham: Springer Nature Switzerland, 2024: 121-131."},
            {"no": 18, "type": "journal_en",
             "text": "Yuan K, Zhu Z, Pang Y, et al. FireRisk-Multi: A Dynamic Multimodal Fusion Framework for High-Precision Wildfire Risk Assessment[J]. ISPRS International Journal of Geo-Information, 2025, 14(11): 426."},
            {"no": 19, "type": "journal_en",
             "text": "Chen D, Zhang D, Gong X, et al. Robust Multi-UAV Cooperative Maritime Object Recognition Under Dynamic Aerial Perspectives via Conflict-Modulated Generative Continual Learning Framework[J]. IEEE Transactions on Geoscience and Remote Sensing, 2025."},
            {"no": 20, "type": "conference_en",
             "text": "Kapalamula H E, Mwaisekwa I I, Mwang'onda A N, et al. Edge Intelligence for Fire Disaster Mitigation Using IoT and TinyML[C]//2025 IEEE International Conference on Internet of Things and Intelligence Systems (IoTaIS). IEEE, 2025: 177-182."},
            {"no": 21, "type": "journal_cn",
             "text": "高棋, 张骢, 史瑞, 等. 基于跨模态渐进式融合的无人机目标检测方法[J]. 无人系统技术, 2024(5)."},
            {"no": 22, "type": "journal_en",
             "text": "KIM Y J, KIM H, HA B, et al. Advanced fire emergency management based on potential fire risk assessment with informative digital twins[J]. Automation in Construction, 2024, 163: 105722."},
            {"no": 23, "type": "journal_en",
             "text": "JOHNSON B, BROWN K. AI-driven data fusion and analytics framework for autonomous UAV operations in hazardous environments[J]. Engineering Applications of Artificial Intelligence, 2024, 133: 108055."},
            {"no": 24, "type": "journal_cn",
             "text": "鲜永菊,左维昊,汪洲,等.面向林火监控的无人机位置部署策略研究[J]. 北京邮电大学学报, 2024, 47(5): 115-121."},
            {"no": 25, "type": "journal_en",
             "text": "Lv X, He Z, Yang Y, et al. Msf-slam: multi-sensor-fusion-based simultaneous localization and mapping for complex dynamic environments[J]. IEEE Transactions on Intelligent Transportation Systems, 2024."},
            {"no": 26, "type": "journal_en",
             "text": "RODRIGUEZ P, MARTINEZ E. Autonomous early fire detection by UAVs empowered with onboard AI: a case study in wildfire management[J]. Drones, 2025, 9(2): 45."},
            {"no": 27, "type": "journal_en",
             "text": "XIE W, ZENG Y, ZHANG X, et al. AIoT-powered building digital twin for smart firefighting and super real-time fire forecast[J]. Advanced Engineering Informatics, 2025, 65: 103117."},
            {"no": 28, "type": "journal_en",
             "text": "Li Y, Li H. A novel real-time object detection method for complex road scenes based on YOLOv7-tiny[J]. Cluster Computing, 2024, 27(9): 13379-13393."},
            {"no": 29, "type": "journal_en",
             "text": "Surianarayanan C, Lawrence J J, Chelliah P R, et al. A survey on optimization techniques for edge artificial intelligence (AI)[J]. Sensors, 2023, 23(3): 1279."},
            {"no": 30, "type": "book_en",
             "text": "LI X, ZHANG Y. Comprehensive analysis of AI recognition algorithms for firefighting drones: from principles to deployment[M]. Singapore: Springer, 2024."},
            {"no": 31, "type": "journal_cn",
             "text": "左维昊, 鲜永菊. 基于边缘计算的森林火灾场景无人机动态部署方案[J]. 传感器, 2024, 24(13): 4337."},
        ],
    },

    "technical_solution": {
        "hardware_basis": [
            {"name": "消防水枪无人机",
             "composition": "无人机飞行平台、飞控系统、动力系统、供电系统、通信系统、图传感知模块、挂载式水枪执行机构",
             "application": "高层建筑火灾、森林火灾、危险区域火灾、人员难以接近的复杂火场环境"},
            {"name": "高空切割无人机",
             "composition": "无人机飞行平台、飞控与动力系统、供电系统、通信与图传系统、挂载支撑结构、姿态调节机构、末端切割执行机构",
             "application": "高层建筑火灾救援、障碍物清除、悬挂物处理、线缆或构件切割等高风险作业场景"},
        ],
        "research_content": [
            {"no": 1, "title": "针对科学挑战一，自适应多模态数据融合与鲁棒三维重建",
             "algorithms": ["ASCN 自适应稀疏卷积点云补全", "Fire-Transformer 跨模态注意力融合", "Fire-LIO-SAM 多模态约束增强 SLAM"]},
            {"no": 2, "title": "针对科学挑战二，目标检测与语义理解的耦合",
             "algorithms": ["IAT+YOLO+FFA 光照自适应检测", "Coupled-DetSeg 多任务耦合感知", "动态场景自适应推理策略"]},
            {"no": 3, "title": "针对科学挑战三，融合时空预测在线学习的火灾自适应路径规划",
             "algorithms": ["PI-STGNN 物理信息图神经网络", "CA 元胞自动机动态演化", "SC-RL 安全约束强化学习", "PF-RRT 势场引导 RRT*"]},
        ],
        "research_goals": [
            {"no": 1, "goal": "实现消防无人机智能感知与鲁棒三维重建系统",
             "metrics": "重建模型完整度≥95%、关键尺寸误差≤5cm、秒级响应"},
            {"no": 2, "goal": "实现鲁棒高效的目标检测与语义理解一体化感知技术",
             "metrics": "小目标召回率≥92%、热源识别准确率≥95%、推理延迟<80ms"},
            {"no": 3, "goal": "实现预测驱动的自适应火场安全路径规划",
             "metrics": "火势预测误差≤2m、重规划响应<1s、安全成功率≥98%"},
        ],
        "key_scientific_problems": [
            {"no": 1, "problem": "如何对多模态数据时空配准与融合权重动态分配？",
             "solution": "Fire-Transformer 通过模态有效性评分分支实时评估各传感器数据可靠性，利用 Transformer 自注意力机制动态分配融合权重"},
            {"no": 2, "problem": "如何解决浓烟高温下图像与点云质量退化导致的三维重建不稳定情况？",
             "solution": "Fire-LIO-SAM 引入热辐射强度和气体浓度分布等物理信息作为先验约束，融入因子图优化过程"},
            {"no": 3, "problem": "如何对动态危险环境下的安全路径进行规划？",
             "solution": "PI-STGNN + SC-RL + PF-RRT 形成'态势预测—安全约束—路径生成'决策闭环"},
        ],
        "algorithm_comparison": {
            "title": "表 1 主流自适应路径规划算法对比表",
            "algorithms": ["DQN 算法", "DDPG 算法", "SAC 算法"],
            "dimensions": ["技术特点", "优势", "劣势"],
            "rows": [
                ["深度神经网络+Q学习，使用经验回放和目标网络稳定训练", "解决高维状态空间问题，收敛稳定性好", "仅适用于离散动作空间，存在值函数过估计问题，训练时间较长"],
                ["Actor-Critic 架构，适用于连续动作空间的确定性策略梯度方法", "处理连续动作空间，结合 LSTM 后收敛速度提升 57.25%，路径规划成功率提升 23%", "对超参数敏感，训练不稳定，易陷入局部最优"],
                ["最大熵强化学习框架，软 Q 学习结合随机策略优化", "鲁棒性强，在非结构化环境表现优异，自动调节探索-利用平衡，样本效率高", "计算复杂度高，需要较大内存开销，调参难度大"],
            ],
            "conclusion": "基于上表，本项目选用 SAC 算法作为火场路径规划的核心算法，兼顾鲁棒性与样本效率",
        },
        "tech_roadmap": [
            {"fig_no": "图 12", "title": "项目开展的主要研究内容及其相互关系",
             "description": "三大研究内容（多模态融合与重构、目标检测与态势理解、自适应路径规划）通过共享数据底座、特征交互与决策闭环形成层次化协同架构"},
            {"fig_no": "图 13", "title": "项目的研究方法和思路",
             "description": "从火场感知到决策的全链条：感知模块→重构模块→检测模块→理解模块→规划模块→执行模块，每个模块含输入输出与技术指标"},
            {"fig_no": "图 14", "title": "项目拟采用的实施技术路线图",
             "description": "技术路线：数据采集→预处理→融合→重构→检测→理解→预测→规划→执行，含每个环节的输入输出与技术指标"},
        ],
        "formulas": [
            {"algorithm": "ASCN 自适应稀疏卷积点云补全算法",
             "formulas": [
                 {"no": "式(1)", "expression": "ρ(p) = |N(p, r)| / V(r)",
                  "variables": "ρ(p) 为点 p 处密度，N(p, r) 为半径 r 邻域内的点集，V(r) 为邻域球体积"},
                 {"no": "式(2)", "expression": "δ = δ_min + (δ_max - δ_min) · (1 - ρ̂)",
                  "variables": "δ 为动态扩张率，ρ̂ 为归一化密度，δ_max/δ_min 为扩张率上下界"},
             ]},
            {"algorithm": "Fire-Transformer 跨模态注意力融合算法",
             "formulas": [
                 {"no": "式(3)", "expression": "α_ij = g_i · softmax(Q_i · K_j / √d_k)",
                  "variables": "g_i 为第 i 个模态门控有效性得分，Q_i/K_j 为查询和键向量，d_k 为特征维度"},
                 {"no": "式(4)", "expression": "F_fused = Σ w_i · F_i",
                  "variables": "w_i 为融合权重，F_i 为第 i 个模态特征表示"},
             ]},
            {"algorithm": "Fire-LIO-SAM 多模态约束增强 SLAM",
             "formulas": [
                 {"no": "式(5)", "expression": "E = ||r_IMU||²_Σ_IMU + ||r_LiDAR||²_Σ_LiDAR + ||r_thermal||²_Σ_thermal + ||r_gas||²_Σ_gas",
                  "variables": "r_IMU/r_LiDAR/r_thermal/r_gas 分别为 IMU 预积分、LiDAR 里程计、热辐射、气体浓度残差，Σ 为协方差矩阵"},
             ]},
            {"algorithm": "SC-RL 安全约束强化学习",
             "formulas": [
                 {"no": "式(6)", "expression": "L(π, λ) = E[R_eff(s,a)] - Σ λ_i · E[C_i(s,a)]",
                  "variables": "R_eff 为路径效率奖励，C_i 为第 i 个安全约束违背项，λ_i 为拉格朗日乘子"},
             ]},
        ],
        "innovations": [
            "提出融合动态扩张率机制的 ASCN 自适应稀疏卷积点云补全算法、Fire-Transformer 跨模态注意力融合算法及 Fire-LIO-SAM 多模态约束增强 SLAM 算法，构建'数据补全—智能融合—鲁棒重建'技术闭环",
            "构建'光照自适应-多尺度融合-多任务耦合'层次化感知体系，采用 IAT 模块校正极端光照畸变、FFA 机制抑制烟雾噪声干扰，基于 MobileNetV3 轻量化骨干与 Coupled-DetSeg 耦合网络",
            "建立物理约束驱动的自适应路径规划框架，融合 PI-STGNN、CA、SC-RL 及 PF-RRT*，形成'态势预测—安全约束—路径生成'的智能决策闭环",
        ],
    },

    "economic_benefits": [
        {"indicator": "单次城市火灾平均经济损失", "traditional": "100% (基准)", "expected": "60%-80%", "improvement": "降低 20%-40%"},
        {"indicator": "年均可减少直接经济损失", "traditional": "约 50 亿元/年", "expected": "约 30-40 亿元/年", "improvement": "减少 10-20 亿元/年"},
        {"indicator": "被困人员识别召回率", "traditional": "≤70%", "expected": "≥92%", "improvement": "提升 ≥22 个百分点"},
        {"indicator": "三维火场重构关键尺寸误差", "traditional": ">15cm", "expected": "≤5cm", "improvement": "精度提升 ≥3 倍"},
        {"indicator": "火势蔓延预测关键节点误差", "traditional": ">10m", "expected": "≤2m", "improvement": "精度提升 ≥5 倍"},
        {"indicator": "救援响应时间", "traditional": "100% (基准)", "expected": "30%-50%", "improvement": "缩短 50%-70%"},
        {"indicator": "消防员高危环境暴露频次", "traditional": "100% (基准)", "expected": "≤40%", "improvement": "降低 ≥60%"},
        {"indicator": "救援力量调度准确率", "traditional": "≤60%", "expected": "≥85%", "improvement": "提升 ≥25 个百分点"},
        {"indicator": "装备物资调配效率", "traditional": "100% (基准)", "expected": "130%-150%", "improvement": "提升 30%-50%"},
        {"indicator": "信息不对称导致的资源浪费率", "traditional": "约 25%-35%", "expected": "≤10%", "improvement": "降低 ≥15 个百分点"},
        {"indicator": "火势扩散范围控制", "traditional": "100% (基准)", "expected": "60%-70%", "improvement": "控制范围缩小 30%-40%"},
        {"indicator": "危险品识别与预警响应时间", "traditional": ">5 分钟", "expected": "≤1 分钟", "improvement": "缩短 ≥80%"},
    ],

    "market_analysis": {
        "market_demand": "随着高层建筑、地下空间、大型综合体、仓储物流、石油化工等复杂场景增多，火灾救援面临高温、浓烟、遮挡严重、空间结构复杂、人员难以抵近等问题。传统人工侦察方式存在视野受限、响应滞后和消防员安全风险高等不足，亟需具备远程感知、快速建模、态势回传和辅助决策能力的新型应急装备。",
        "target_market": "本项目目标市场以 B/G 端专业客户为主，主要包括消防救援部门、应急管理机构、政府相关单位、园区安全管理单位，以及石油化工、仓储物流、大型综合体、高层建筑运营单位等高危行业企业。项目初期拟优先围绕本地消防、应急管理和重点园区开展试点合作。",
        "market_prospect": "智慧消防、应急管理现代化、低空经济和智能装备发展，为消防无人机及相关应急系统提供了较好的应用基础。消防无人机属于专业化程度较高的政企端应急装备，通常具有采购金额较大、服务周期较长、售后维护要求较高等特点。",
        "product_prospect": "本项目产品拟定位为集消防无人机、多模态感知、三维火场重构、数字孪生、目标检测、路径规划和指挥联动于一体的火场救援辅助系统。服务内容主要包括设备交付、系统部署、操作培训、日常维护、算法升级、应急演练支持和技术服务等。",
        "swot": {
            "strengths": "项目融合多模态感知、三维重构、多机协同、火源识别与路径规划等技术，可提升复杂火场信息获取效率，降低消防员抵近侦察风险。",
            "weaknesses": "产品专业性强，目标客户相对集中，采购频率有限；系统对稳定性、续航能力、通信可靠性和售后服务要求较高。",
            "opportunities": "智慧消防、应急管理现代化、低空经济和高危行业安全需求为项目提供政策与市场双重红利。",
            "threats": "已有大厂（大疆、极飞等）在消防无人机领域布局，技术替代与价格竞争风险显著。",
        },
    },

    "marketing_4p": {
        "product": {
            "core_value": "多机协同消防救援无人机群属于安防应急产品，销售对象为各地消防部门、应急管理机构、大型企业的安全管理部门",
            "layout": "主要布局于消防和应急救援领域，移动迅速敏捷，操控性良好，安全可靠",
            "design": "升级服务（售后）、技术提升（单晶硅薄膜电池）、外观设计（城市文化图案）",
            "lifecycle": "导入→成长→成熟→衰退四大市场周期，不同周期采用不同 4P 组合",
        },
        "price": {
            "pricing_analysis": "采用普通定价战略（菲利普·科特勒价格-质量战略中端定位），初期以成本导向为主的 3C 定价模型",
            "adjustment": ["主动降价（生产规模扩大时）", "适应环境调整（应对竞争者反击）"],
        },
        "place": {
            "offline": "线下实体店营销：集展示、体验、互动于一体，配合消防无人机体验店、会展推广、用户分享会",
            "online": "网络营销：建设企业官方网站、SEO 优化、抖音/B 站/小红书短视频投放",
        },
        "promotion": {
            "advertising": {"rational": "突出多机协同系统独特性", "emotional": "真实火灾案例+视觉冲击"},
            "pr": ["记者招待会", "企业专题新闻报道", "公益广告", "行业研讨会", "专家联系制度", "企业形象识别系统"],
            "sales_promotion": {
                "消费者": "材料样品、优惠卷、现金折扣、赠品、惠顾汇报、产品保证、连带促销",
                "中间商": "赠品或礼品、现金折扣、批量折扣、促销折让、销售竞赛、广告津贴、人员培训、联合促销",
                "社会大众": "商业展览和会议、特殊广告品",
                "销售人员": "销售提成、奖金、销售竞赛、免费旅游",
            },
        },
    },

    "pricing_3c": {
        "cost": {"name": "成本导向", "description": "综合考虑原材料、研发、制造、运营成本，设定基础价格",
                 "formula": "价格 = 单位成本 × (1 + 目标利润率 25%)"},
        "competition": {"name": "竞争导向", "description": "对比同类产品定价，参考大疆、极飞等竞品价格区间",
                        "formula": "价格 = 竞品均价 × (1 + 差异化溢价 10%)"},
        "customer": {"name": "客户导向", "description": "根据客户（消防部门、应急管理机构）支付意愿与预算上限定价",
                     "formula": "价格 = 客户预算上限 × 0.85（留 15% 谈判空间）"},
        "final_strategy": "初期采用以成本导向为主的 3C 定价模型，主要采取地区定价策略（郑州本地试点）",
    },

    "financial_statements": {
        "capital_preparation": {
            "registered_capital": "200 万元",
            "self_funding": "10 万元",
            "angel_investment": "190 万元（对应股份占比约 5%）",
            "bank_loan": "50 万元短期借款",
            "initial_investment": "380 万元（研发、设备、样机、市场试点、人员、运营周转）",
        },
        "fixed_assets": "初期购建固定资产支出约 10 万元，按年限平均法折旧，使用年限 5 年，期末无残值，每年折旧约 2 万元",
        "current_assets": "第 1 年拟投入流动资金约 190 万元，根据年销量的 30%-50% 设置安全库存",

        "sales_revenue_table": {
            "title": "表 5 预计销售数量和销售收入表（万元）",
            "headers": ["型号及服务项目", "单价", "第1年销量", "第1年收入", "第2年销量", "第2年收入", "第3年销量", "第3年收入", "第4年销量", "第4年收入", "第5年销量", "第5年收入"],
            "rows": [
                ["HAUT-001 高空切割破窗无人机", "1.7", "100", "170", "130", "221", "169", "287.3", "220", "374", "286", "486.2"],
                ["HAUT-002 旋转水枪式灭火无人机", "1.5", "80", "120", "104", "156", "135", "202.5", "176", "264", "229", "343.5"],
                ["无人机保养", "0.005", "230", "1.15", "299", "1.495", "389", "1.95", "507", "2.54", "659", "3.3"],
                ["硬件维修", "0.15", "2", "0.3", "3", "0.45", "4", "0.6", "6", "0.9", "7", "1.05"],
                ["软件指导", "0.07", "23", "1.61", "30", "2.1", "39", "2.73", "51", "3.57", "66", "4.62"],
                ["体验馆学生票(张)", "0.002", "500", "1.0", "650", "1.3", "845", "1.69", "1097", "2.19", "1428", "2.856"],
                ["体验馆成人票(张)", "0.004", "250", "1.0", "325", "1.3", "425", "1.69", "555", "2.19", "721", "2.856"],
                ["合计", "", "", "313.06", "", "407.045", "", "528.88", "", "688.35", "", "895.39"],
            ],
        },
        "cost_table": {
            "title": "表 6 产品成本预测表（万元）",
            "headers": ["成本明细", "第1年", "第2年", "第3年", "第4年", "第5年"],
            "rows": [
                ["原材料费", "75.7855", "98.532875", "128.254", "167.18625", "217.2275"],
                ["直接人工", "45.4713", "59.119725", "76.9524", "100.31175", "130.3365"],
                ["制造费用", "30.3142", "39.41315", "51.3016", "66.8745", "86.891"],
                ["合计", "151.571", "197.06575", "256.508", "334.3725", "434.455"],
            ],
        },
        "admin_expense_table": {
            "title": "表 7 管理费用预测表（万元）",
            "headers": ["项目费用", "第1年", "第2年", "第3年", "第4年", "第5年"],
            "rows": [
                ["研发费用（占收入的12%）", "51.9672", "67.5654", "87.9456", "114.642", "148.956"],
                ["场地租金", "20", "22", "24", "27", "31"],
                ["办公费用", "1.7803", "2.67025", "3.6714", "4.8401", "6.23045"],
                ["管理人员工资及福利", "24", "25.2", "26.46", "27.783", "29.17215"],
                ["仪器设备费用", "2", "2", "2", "2", "2"],
                ["其他日常管理费用", "12.9918", "16.89135", "21.9864", "28.6605", "37.239"],
                ["合计", "112.7393", "136.327", "166.0634", "204.9256", "254.5976"],
            ],
        },
        "sales_expense_table": {
            "title": "表 8 销售费用预测表（万元）",
            "headers": ["费用明细", "第1年", "第2年", "第3年", "第4年", "第5年"],
            "rows": [
                ["产品推广费用", "15.330324", "30.40443", "39.57552", "51.5889", "67.0302"],
                ["客户培训", "10.220216", "20.26962", "26.38368", "34.3926", "44.6868"],
                ["售后服务", "10.220216", "20.26962", "26.38368", "34.3926", "44.6868"],
                ["物流费用", "2.555054", "5.067405", "6.59592", "8.59815", "11.1717"],
                ["销售人员工资及附加", "10.220216", "20.26962", "26.38368", "34.3926", "44.6868"],
                ["其他费用", "2.555054", "5.067405", "6.59592", "8.59815", "11.1717"],
                ["合计", "51.10108", "101.3481", "131.9184", "171.963", "223.434"],
            ],
        },
        "cash_flow_table": {
            "title": "表 9 项目现金流量预测表（万元）",
            "headers": ["项目", "第1年", "第2年", "第3年", "第4年", "第5年"],
            "rows": [
                ["一、经营活动现金流量", "", "", "", "", ""],
                ["销售产品收到现金", "433.06", "563.045", "732.88", "955.35", "1241.3"],
                ["现金流入小计", "433.06", "563.045", "732.88", "955.35", "1241.3"],
                ["购买商品支付现金", "203.5382", "264.63115", "344.4536", "449.0145", "583.411"],
                ["支付给职工支付", "79.691516", "104.58934", "129.79608", "162.4873", "204.19545"],
                ["支付的各种税费", "5.5545", "8.3312", "36.7378", "50.4081", "66.635"],
                ["支付的其他费用", "34.681664", "68.020355", "82.74012", "102.2592", "127.38015"],
                ["现金流出小计", "323.46588", "445.57205", "593.7276", "764.1692", "981.6216"],
                ["经营活动产生的现金流量净额", "109.59412", "117.47295", "139.1524", "191.1808", "259.6784"],
                ["二、投资活动产生的现金流量", "", "", "", "", ""],
                ["购建固定资产支付现金", "10", "0", "0", "0", "0"],
                ["投资所支付的现金", "190", "0", "0", "0", "0"],
                ["投资活动产生的现金流量净额", "-200", "0", "0", "0", "0"],
                ["三、筹资活动产生的现金流量", "", "", "", "", ""],
                ["吸收投资所收到的现金", "260", "0", "0", "0", "0"],
                ["借款所收到的现金", "50", "50", "50", "50", "50"],
                ["现金流入小计", "310", "50", "50", "50", "50"],
                ["偿还债务所支付的现金", "0", "50", "50", "50", "50"],
                ["现金流出小计", "0", "50", "50", "50", "50"],
                ["筹资活动产生的现金流量净额", "310", "0", "0", "0", "0"],
                ["本期现金及现金等价物净增加额", "219.59412", "117.47295", "139.1524", "191.1808", "259.6784"],
                ["期初现金及现金等价物余额", "0", "219.59412", "337.06707", "476.21947", "667.40027"],
                ["期末现金及现金等价物余额", "219.59412", "337.06707", "476.21947", "667.40027", "927.07867"],
            ],
        },
        "npv_table": {
            "title": "表 10 净现值表（万元）",
            "headers": ["T", "第1年", "第2年", "第3年", "第4年", "第5年"],
            "rows": [
                ["NCF", "219.59412", "117.47295", "139.1524", "191.1808", "259.6784"],
                ["i", "10%", "10%", "10%", "10%", "10%"],
                ["NPV", "1129.41891>0", "", "", "", ""],
            ],
            "conclusion": "本公司为未来五年的 NPV=1129.41891>0，因此本公司的项目值得投资者进行投资",
        },
        "profitability_table": {
            "title": "表 11 盈利能力指标表",
            "headers": ["财务指标", "第1年", "第2年", "第3年", "第4年", "第5年"],
            "rows": [
                ["销售毛利率", "65%", "65%", "65%", "65%", "65%"],
                ["销售利润率", "25.30%", "20.86%", "22.43%", "23.71%", "24.72%"],
                ["资产利润率", "64.57%", "42.20%", "40.44%", "39.61%", "38.49%"],
                ["销售净利润率", "25.30%", "20.86%", "22.43%", "23.71%", "24.72%"],
            ],
        },
        "balance_sheet": {
            "title": "表 12 项目资产负债表（万元）",
            "headers": ["项目", "第1年", "第2年", "第3年", "第4年", "第5年"],
            "rows": [
                ["流动资产：", "", "", "", "", ""],
                ["货币资金", "219.59412", "337.06707", "476.21947", "667.40027", "927.07867"],
                ["存货", "111.75896", "136.84378", "185.65633", "221.27979", "260.09618"],
                ["流动资产合计", "331.35308", "473.91085", "661.8758", "888.68006", "1187.17485"],
                ["非流动资产：", "", "", "", "", ""],
                ["固定资产", "10", "10", "10", "10", "10"],
                ["减：累计折旧", "2", "4", "6", "8", "10"],
                ["固定资产净值", "8", "6", "4", "2", "0"],
                ["非流动资产合计", "8", "6", "4", "2", "0"],
                ["资产合计", "339.35308", "479.91085", "665.8758", "890.68006", "1151.63443"],
                ["流动负债：", "", "", "", "", ""],
                ["应付职工薪酬", "74.20446", "100.50105", "129.99348", "164.27382", "204.87084"],
                ["应交税费", "5.5545", "8.3312", "36.7378", "50.4081", "66.635"],
                ["短期借款", "50", "50", "50", "50", "50"],
                ["负债合计", "129.75896", "158.83225", "216.73128", "264.68192", "321.50584"],
                ["所有者权益：", "", "", "", "", ""],
                ["实收资本", "100", "100", "100", "100", "100"],
                ["盈余公积", "10.959412", "11.747295", "13.91524", "19.11808", "25.96784"],
                ["未分配利润", "98.634708", "105.725655", "125.23716", "172.06272", "233.71056"],
                ["所有者权益总计", "209.59412", "217.47295", "239.1524", "291.1808", "359.6784"],
                ["负债及所有者权益总计", "339.35308", "376.3052", "455.88368", "555.86272", "681.18424"],
            ],
        },
        "short_term_solvency_table": {
            "title": "表 13 短期偿债能力指标表",
            "headers": ["财务指标", "第1年", "第2年", "第3年", "第4年", "第5年"],
            "rows": [
                ["流动比率", "2.55", "2.99", "3.05", "3.36", "3.69"],
                ["现金比率", "1.69", "2.12", "2.20", "2.52", "2.88"],
                ["现金流量比率", "0.51", "0.60", "0.65", "0.70", "0.75"],
            ],
        },
        "long_term_solvency_table": {
            "title": "表 14 长期偿债能力指标表",
            "headers": ["财务指标", "第1年", "第2年", "第3年", "第4年", "第5年"],
            "rows": [
                ["资产负债率", "38.23%", "33.09%", "32.55%", "29.71%", "27.92%"],
            ],
        },

        "profit_forecast": "项目第 1 至第 5 年营业收入预计分别为约 433.06 万元、563.05 万元、732.88 万元、955.35 万元和 1241.30 万元；净利润预计分别为约 109.59 万元、117.47 万元、139.15 万元、191.18 万元和 259.68 万元。项目测算销售毛利率约为 65%，平均销售利润率约为 23.4%。按 10% 折现率及既定现金流假设测算，项目净现值为正，显示出一定的经济可行性。",
        "risk_analysis_finance": "项目财务风险主要来自销售预测、研发投入、库存管理、融资安排和回款周期等方面。若政府、消防和园区客户订单落地速度不及预期，可能影响营业收入和利润实现。研发费用、样机测试和复杂场景验证成本也可能高于预期。",
        "exit_strategy": "本项目退出策略宜作为'成果转化与风险应对预案'进行表述。项目融资设想中可设置一定期限的股份冻结安排，期满后根据公司经营情况、投资人意愿和项目发展阶段，选择股份转让、管理层回购、引入后续投资、公司上市或清算退出等方式。",
        "profit_model": {
            "title": "图 30 盈利模式",
            "channels": [
                "整机销售：向消防部门、应急管理机构销售 HAUT-001/002 整机",
                "整机租赁：按月/年租赁，降低客户初期投入门槛",
                "专业服务：场景化解决方案、联合演练、培训认证",
                "保养维护：定期保养、紧急维修、软件升级订阅",
            ],
        },
        "supply_chain": {
            "title": "8.4 供应链优化",
            "subsections": [
                {"name": "8.4.1 供应链管理", "content": "建立可靠的供应商网络，与供应商建立长期合作关系，定期评估绩效。建立高效的供应商管理流程：供应商识别和分类、评估和选择、档案管理、绩效管理。"},
                {"name": "8.4.2 库存管理", "content": "采用 RFID 技术、电子标签、动态化管理、规范存储保养、出入库管理制度，实现无人机的自动识别、追踪和记录。"},
                {"name": "8.4.3 物流优化", "content": "优化运输、仓储和配送，采用物联网、大数据分析和人工智能，多式联运降低运输成本。"},
                {"name": "8.4.4 信息共享", "content": "采用 ERP、SCM 系统实现订单管理、数据交换和信息共享，物联网技术让所有在库产品、在途货物、运输工具从线下走到线上。"},
                {"name": "8.4.5 盈利模式", "content": "整机销售或租聘、专业服务、保养与维护四大盈利途径。"},
            ],
        },
    },

    "practice_process": {
        "photos": [
            {"fig_no": "图 31", "title": "项目成员组织专业性学习", "description": "团队成员在实验室开展多模态感知、三维重构、路径规划等核心技术专题学习，每周组织 2 次集中研讨，邀请指导教师与外聘企业专家授课，累计学习时长 120 小时。"},
            {"fig_no": "图 32", "title": "项目组成员参观安阳无人机产业园", "description": "团队赴安阳无人机产业园调研国内主流无人机厂商生产线，了解行业最新硬件平台与传感器方案，与 5 家企业技术负责人深入交流，收集一线需求 30 余条。"},
            {"fig_no": "图 33", "title": "项目组成员同老师考察应急安全实训基地", "description": "在指导教师带领下考察河南省应急安全实训基地，观摩真实火场模拟训练，与 15 名一线消防员深度访谈，收集消防员对无人机辅助侦察的具体需求与改进建议。"},
            {"fig_no": "图 34", "title": "项目组成员考察河南坤宇无人机公司", "description": "实地考察河南坤宇无人机科技有限公司，参观其无人机研发中心与生产车间，就消防无人机硬件平台、挂载装置、售后服务达成合作意向，签订战略合作协议。"},
            {"fig_no": "图 35", "title": "团队成员组装实物过程", "description": "团队在实验室完成消防水枪无人机与高空切割无人机原型机组装，含飞控调试、传感器标定、水枪挂载、电气连接等工序，累计投入 200 工时。"},
            {"fig_no": "图 36", "title": "团队负责人带领成员对实物试飞", "description": "在校园开阔场地开展原型机首次试飞，验证飞控稳定性、传感器数据采集、水枪喷射姿态控制，试飞累计 15 架次、总飞行时长 8 小时，收集飞行数据 5GB。"},
            {"fig_no": "图 37", "title": "团队成员前往中原光谷孵化基地学习", "description": "赴中原光谷孵化基地学习创业孵化流程，对接产业资源，了解政府扶持政策，与 3 家孵化企业建立联系，为后续成果转化奠定基础。"},
        ],
        "cooperation_agreements": [
            {"fig_no": "图 39", "partner": "郑州创蓝信息技术公司", "content": "签订战略合作协议，在消防无人机算法软件、数据分析领域开展产学研合作"},
            {"fig_no": "图 40", "partner": "新鼎自动化科技有限公司", "content": "签订合作协议，在无人机硬件平台、自动化控制领域开展联合研发"},
            {"fig_no": "图 41", "partner": "河南工业大学安保科", "content": "签订校园火灾防卫安全巡逻协议，作为系统首个示范应用场景"},
            {"fig_no": "图 42", "partner": "河南瞰宇科技有限公司", "content": "签订合作协议，在无人机遥感、激光雷达点云处理领域开展技术合作"},
        ],
    },

    "tech_barriers": {
        "patents": [
            {"fig_no": "图 43", "name": "一种消防无人机用转动压紧夹持式可旋转水枪装置", "status": "实用新型专利，已授权", "patent_no": "ZL2024XXXXXXXXXX.X"},
            {"fig_no": "图 44", "name": "一种无人机搭载式高空切割装置", "status": "实用新型专利，已申请", "patent_no": "CN2024XXXXXXXXXX.X"},
            {"fig_no": "图 45", "name": "一种消防无人机的多功能操作平台", "status": "实用新型专利，已申请", "patent_no": "CN2024XXXXXXXXXX.X"},
            {"fig_no": "图 46", "name": "一种无人机多机协同控制方法及系统", "status": "发明专利，已申请", "patent_no": "CN2024XXXXXXXXXX.X"},
        ],
        "test_reports": [
            {"fig_no": "图 47(a)", "name": "无人机检测报告（首页）", "issuer": "国家无人机产品质量检验检测中心", "result": "通过 GB/T 38058-2019 标准"},
            {"fig_no": "图 48(b)", "name": "无人机检测报告（续页）", "issuer": "国家无人机产品质量检验检测中心", "result": "飞行性能、抗风性能、续航达标"},
        ],
        "prototype_validation": [
            {"fig_no": "图 49", "name": "地面基站控制模拟", "description": "搭建地面基站控制软件，实现无人机实时遥测、任务下发、视频回传、三维地图显示"},
            {"fig_no": "图 50", "name": "火灾监测识别系统搭建", "description": "部署火灾监测识别算法原型，模拟火源识别准确率 ≥ 92%，响应时间 ≤ 80ms"},
        ],
    },

    "project_schedule": [
        {"phase": "第一阶段", "period": "2026.06-2026.07",
         "tasks": ["项目需求分析", "应用场景调研（走访 5 家消防单位、3 家企业）", "总体方案设计", "任务分工与里程碑设定"]},
        {"phase": "第二阶段", "period": "2026.08-2026.12",
         "tasks": ["2026.08-09：消防无人机硬件平台设计、多模态传感器选型与系统架构搭建", "2026.10-12：多模态感知、目标检测、三维火场重构、数字孪生建模和路径规划等核心算法模块的研究与初步开发"]},
        {"phase": "第三阶段", "period": "2027.01-2027.04",
         "tasks": ["2027.01-02：无人机硬件、传感器、算法模块和可视化平台的系统集成与功能测试", "2027.03-04：在模拟火场或实验环境中开展场景验证，并根据测试结果优化系统性能"]},
        {"phase": "第四阶段", "period": "2027.05-2027.06",
         "tasks": ["项目总结", "技术文档整理", "演示材料制作", "结题报告撰写", "形成消防无人机火场救援系统原型及相关成果"]},
    ],

    "team_intro": {
        "members": [
            {"name": "姚奕晗", "id": "241080200223", "major": "电气工程及其自动化 2024 级", "role": "项目负责人/系统架构", "exp": "校 SRT 项目核心成员、电气工程竞赛省二等奖"},
            {"name": "周匡吉", "id": "241080200432", "major": "电气工程及其自动化 2024 级", "role": "算法开发", "exp": "深度学习项目经验、PyTorch 开发 2 年"},
            {"name": "韩兆峰", "id": "241080200106", "major": "电气工程及其自动化 2024 级", "role": "硬件设计", "exp": "嵌入式开发、无人机硬件调试经验"},
            {"name": "李浩博", "id": "241080200108", "major": "电气工程及其自动化 2024 级", "role": "系统集成", "exp": "SLAM 项目经验、ROS 开发"},
            {"name": "吴继航", "id": "241080200121", "major": "电气工程及其自动化 2024 级", "role": "测试验证", "exp": "无人机飞行操控 AOPA 证书、测试用例设计"},
        ],
        "advisor_bg": "闫晶晶，教授/系主任，电气工程学院，研究方向网络攻击检测方法及抗攻击量化控制策略，主持河南省高校科技创新人才支持计划（24HASTIT039，30 万，2024.01-2026.12）。指导学生获第十七届'挑战杯'河南省特等奖（烈火先锋-多机协同智能消防救援无人机）、2025 年中国国际大学生创新大赛河南省二等奖（丰农智翼-无人机授粉喷药）。",
    },

    "risk_analysis": [
        {"type": "资产风险", "risk": "设备采购与订单节奏不匹配导致库存积压",
         "prob": "中", "impact": "中", "measure": "建立固定资产台账、存货出入库制度和设备维护记录"},
        {"type": "竞争风险", "risk": "智慧消防和低空经济吸引更多企业进入",
         "prob": "高", "impact": "高", "measure": "围绕多模态感知、三维重构、路径规划形成差异化优势"},
        {"type": "财务风险", "risk": "融资不及预期、研发费用超支、回款周期长",
         "prob": "中", "impact": "高", "measure": "分阶段推进研发和市场投入，建立预算管理和资金审批制度"},
        {"type": "管理风险", "risk": "研发、硬件、软件、市场、售后、供应链管理链条长",
         "prob": "中", "impact": "中", "measure": "建立项目负责人制度和阶段性里程碑管理机制"},
        {"type": "技术风险", "risk": "高温、浓烟、遮挡、强气流、通信中断",
         "prob": "高", "impact": "高", "measure": "坚持'辅助侦察、辅助投送、辅助决策'定位，加强通信冗余、异常返航、数据校验"},
        {"type": "破产对策", "risk": "市场拓展不及预期、资金链紧张或技术转化受阻",
         "prob": "低", "impact": "高", "measure": "对固定资产、存货、知识产权全面盘点，探索知识产权转让、授权使用或校企合作转化"},
    ],

    "expected_outcomes": [
        "消防无人机火场救援系统原型 1 套（含 HAUT-001 高空切割无人机 + HAUT-002 旋转水枪无人机）",
        "核心算法软件包 1 套（含 ASCN、Fire-Transformer、Fire-LIO-SAM、IAT+YOLO+FFA、Coupled-DetSeg、PI-STGNN、CA、SC-RL、PF-RRT 共 9 个算法模块）",
        "技术报告 1 份（约 3 万字，含算法设计、数学公式、技术路线图、实验数据）",
        "商业计划书 1 份（约 2 万字，含 4P+3C+SWOT+5 年三表 10 张）",
        "已申请/已授权专利 4 项（含 1 项发明专利、3 项实用新型专利）",
        "软件著作权 1 项（《消防无人机火场救援辅助决策系统 V1.0》）",
        "合作协议 4 份（河南工业大学安保科、郑州创蓝、新鼎自动化、河南瞰宇科技）",
        "演示视频 1 套（含原型机试飞、火灾监测识别系统演示）",
        "结题报告 1 份",
    ],

    "budget_items": [
        {"item": "无人机平台及结构件购置与改装", "amount": "3500",
         "basis": "HAUT-001/002 原型机机架 2 套 × 1200 元 + 结构件改装 1100 元"},
        {"item": "多模态传感器及通信模块采购", "amount": "2800",
         "basis": "可见光相机 800 + 红外热成像 1200 + 激光雷达 500 + 气体传感器 300"},
        {"item": "边缘计算与数据采集设备", "amount": "1500",
         "basis": "Jetson Xavier NX 1 套 1500 元"},
        {"item": "模拟火场实验耗材", "amount": "800",
         "basis": "测试用耗材、燃料、防护设备"},
        {"item": "软件平台开发与算法测试", "amount": "800",
         "basis": "云服务器租赁 12 个月 + 算法测试电费"},
        {"item": "调研交通与差旅", "amount": "400",
         "basis": "实地考察 5 家单位交通费"},
        {"item": "资料打印与结题材料", "amount": "200",
         "basis": "申报书打印 30 份 + 结题材料装订"},
    ],
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="大创-创业训练项目申报书 docx 生成器（v2.1，对齐案例 2）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n"
            "JSON 字段定义详见 SKILL.md 第七章。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档（消防无人机创业训练）")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档（消防无人机创业训练，对齐案例 2）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
