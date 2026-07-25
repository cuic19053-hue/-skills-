#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
思想汇报 docx 生成器

格式标准：A4，页边距上下 2.54cm 左右 2.5cm；标题'思想汇报'黑体二号居中；
称呼'敬爱的党组织：'顶格宋体小四全角冒号；正文宋体小四 1.5 倍行距首行缩进 2 字符；
结尾固定为'恳请党组织批评指正。'+'此致'另起一行空两格+'敬礼！'另起一行顶格；落款右对齐。
四段结构：近期理论学习（结合时政热点）/ 学习工作情况 / 思想动态与不足 / 下一步努力方向。
字数目标 1500~2000 字。每季度提交一份，支持 Q1/Q2/Q3/Q4 季度字段动态填充。
政治规范校验：必提 5 项要点、指导思想 6 项顺序、三重查重（党章原文 50 字/网络模板 30 字/历史汇报 50 字）。

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ==================== 字体与格式常量 ====================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# ==================== 工具函数 ====================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进），宋体小四"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_cizhi_paragraph(doc, text: str = "此致"):
    """'此致'另起一行，空两格（首行缩进 2 字符）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格（不缩进）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=False, line_spacing=1.5)


def add_blank_paragraph(doc):
    """空段落，用于段落间留白"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return p


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def count_chinese_chars(text: str) -> int:
    """统计中文字符数（用于字数控制）"""
    if not text:
        return 0
    return len(re.findall(r"[\u4e00-\u9fff]", str(text)))


def check_plagiarism_risk(text: str, fragments: List[str],
                          threshold: int = 50) -> List[str]:
    """查重风险检测：检查连续 N 字与参考片段重复

    Args:
        text: 待检测的文本
        fragments: 党章原文/网络模板/历史汇报片段列表
        threshold: 连续重复字数阈值

    Returns:
        触发查重警告的片段列表
    """
    warnings = []
    if not text:
        return warnings
    text_clean = re.sub(r"\s+", "", str(text))
    for fragment in fragments:
        if not fragment:
            continue
        frag_clean = re.sub(r"\s+", "", str(fragment))
        if len(frag_clean) < threshold:
            continue
        # 滑动窗口检测：取 fragment 中所有长度 threshold 的子串
        for i in range(len(frag_clean) - threshold + 1):
            window = frag_clean[i:i + threshold]
            if window in text_clean:
                warnings.append(f"检测到连续 {threshold} 字与参考片段重复：{window[:20]}...")
                break  # 每个 fragment 只报告一次
    return warnings


def check_political_terms(text: str) -> List[str]:
    """政治用语规范校验：必提要点 + 禁用简写 + 指导思想顺序

    Returns:
        警告信息列表
    """
    warnings = []
    if not text:
        return warnings
    text_str = str(text)

    # 必提要点（5 项缺一不可）
    required_terms = [
        "中国共产党",
        "习近平新时代中国特色社会主义思想",
        "党的二十大",
        "共产主义",
        "全心全意为人民服务",
    ]
    for term in required_terms:
        if term not in text_str:
            warnings.append(f"必提要点缺失：'{term}'")

    # 禁用简写检测：若完整表述存在，则其包含的子串不算简写
    skip_abbrs = set()
    full_to_abbr = [
        ("习近平新时代中国特色社会主义思想", "新时代中国特色社会主义思想"),
        ("'三个代表'重要思想", "三个代表"),
        ("\"三个代表\"重要思想", "三个代表"),
        ("\u201c三个代表\u201d重要思想", "三个代表"),
    ]
    for full, abbr in full_to_abbr:
        if full in text_str:
            skip_abbrs.add(abbr)

    forbidden_abbreviations = [
        ("习近平思想", "应为'习近平新时代中国特色社会主义思想'"),
        ("习思想", "应为'习近平新时代中国特色社会主义思想'"),
        ("党的20大", "应为'党的二十大'"),
        ("马列", "应为'马克思列宁主义'"),
        ("毛思想", "应为'毛泽东思想'"),
        ("邓理论", "应为'邓小平理论'"),
        ("科学观", "应为'科学发展观'"),
    ]
    if "新时代中国特色社会主义思想" not in skip_abbrs:
        forbidden_abbreviations.append(("新时代中国特色社会主义思想", "应为'习近平新时代中国特色社会主义思想'（带'习近平'三字）"))
    if "三个代表" not in skip_abbrs:
        forbidden_abbreviations.append(("三个代表", "应为'\"三个代表\"重要思想'（注意引号与'重要思想'四字）"))
    for abbr, suggestion in forbidden_abbreviations:
        if abbr in text_str:
            warnings.append(f"禁用简写'{abbr}'：{suggestion}")
    # "20大"检测
    for _ in re.finditer(r"20\s*大", text_str):
        warnings.append("禁用简写'20大'（应完整表述为'党的二十大'）")
        break
    # "二十大"检测（排除"党的二十大"上下文）
    for match in re.finditer(r"二十大", text_str):
        start = match.start()
        if start >= 2 and text_str[start - 2:start + 3] == "党的二十大":
            continue
        warnings.append("禁用简写'二十大'（应完整表述为'党的二十大'）")
        break

    # 指导思想 6 项顺序检测（兼容中英文引号）
    guideline_pattern = r"马克思列宁主义.*?毛泽东思想.*?邓小平理论.*?" r"[\u201c\"']三个代表[\u201d\"']重要思想.*?科学发展观.*?习近平新时代中国特色社会主义思想"
    if not re.search(guideline_pattern, text_str, re.DOTALL):
        checks = [
            ("马克思列宁主义", "马克思列宁主义"), ("毛泽东思想", "毛泽东思想"),
            ("邓小平理论", "邓小平理论"),
            ("'三个代表'重要思想", "三个代表重要思想"),
            ("\"三个代表\"重要思想", "三个代表重要思想"),
            ("\u201c三个代表\u201d重要思想", "三个代表重要思想"),
            ("科学发展观", "科学发展观"),
            ("习近平新时代中国特色社会主义思想", "习近平新时代中国特色社会主义思想"),
        ]
        required_keys = ["马克思列宁主义", "毛泽东思想", "邓小平理论",
                         "三个代表重要思想", "科学发展观", "习近平新时代中国特色社会主义思想"]
        found = {key for needle, key in checks if needle in text_str}
        missing = [k for k in required_keys if k not in found]
        if missing:
            warnings.append(f"指导思想漏项：{', '.join(missing)}")
        else:
            warnings.append("指导思想 6 项顺序错误（应为马列→毛→邓→三→科→习）")

    return warnings


# 党章原文高风险片段（连续 50 字重复则警告）
PARTY_CONSTITUTION_FRAGMENTS = [
    "中国共产党是中国工人阶级的先锋队，同时是中国人民和中华民族的先锋队，" "是中国特色社会主义事业的领导核心，代表中国先进生产力的发展要求，"
    "代表中国先进文化的前进方向，代表中国最广大人民的根本利益。",
    "全心全意为人民服务。党除了工人阶级和最广大人民群众的利益，没有自己特殊的利益。",
    "坚持社会主义道路、坚持人民民主专政、坚持中国共产党的领导、"
    "坚持马克思列宁主义毛泽东思想这四项基本原则，是我们的立国之本。",
]

# 网络思想汇报模板高风险片段（连续 30 字重复则警告）
NETWORK_TEMPLATE_FRAGMENTS = [
    "我怀着激动的心情向党组织汇报本季度的思想动态",
    "作为一名入党积极分子，我深知自己肩负的责任和使命",
    "通过本季度的理论学习，我深刻认识到党的伟大",
    "本季度我认真学习党的最新理论成果，努力提高思想觉悟",
    "在党组织的培养教育下，我的思想觉悟有了很大提高",
]

# 个人历史思想汇报高风险片段（连续 50 字重复则警告，模拟上一季度汇报）
# 实际使用时应替换为用户上一季度提交的真实汇报文本
HISTORICAL_REPORT_FRAGMENTS = [
    "本季度我重点学习了《习近平著作选读》第一卷中关于坚持和发展中国特色社会主义的几个问题",
    "通过参加党课学习，我对党的认识更加深入，进一步坚定了入党信念",
]


# ==================== ApplicationDocBuilder 主类 ====================

class ApplicationDocBuilder:
    """思想汇报 docx 构建器

    季度字段（quarter）取 Q1/Q2/Q3/Q4，决定时政热点填充与过渡句措辞。
    四段结构：理论学习（重点）/ 学习工作 / 思想动态与不足（重点）/ 下一步努力方向。
    """

    # 季度到时政热点的映射（2025 年真实发生事件，不可编造）
    QUARTER_CURRENT_AFFAIRS = {
        "Q1": "2025 年全国两会精神与政府工作报告",
        "Q2": "五一国际劳动节、五四青年节、七一建党节相关纪念活动",
        "Q3": "八一建军节、抗战胜利 80 周年纪念、国庆节、教师节",
        "Q4": "党的二十届三中全会精神、中央经济工作会议、年终总结",
    }

    # 季度到过渡句的映射
    QUARTER_OPENING = {
        "Q1": "2025 年第一季度已过去。在新年开局、全国两会胜利召开之际，",
        "Q2": "2025 年第二季度已接近尾声。在'七一'建党节即将到来之际，",
        "Q3": "2025 年第三季度已过去。在国庆节即将到来、新学年开学之际，",
        "Q4": "2025 年第四季度已接近尾声。在年终岁末回顾全年之际，",
    }

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        # Normal 样式默认设置为宋体小四（中英文同步）
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []
        self.full_text_parts: List[str] = []  # 全文缓存，用于查重检测

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段，非列表返回空列表"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def add_para(self, text, indent=True):
        """添加正文段落并记录到全文缓存（用于查重检测）"""
        if text:
            self.full_text_parts.append(str(text))
        return add_body_paragraph(self.doc, text, indent=indent)

    def _resolve_quarter(self) -> str:
        """解析季度字段：优先使用 data.quarter，否则按当前月份推断"""
        q = str(self._get("quarter", default="")).upper().strip()
        if q in ("Q1", "Q2", "Q3", "Q4"):
            return q
        # 按当前月份推断
        month = datetime.now().month
        if month in (1, 2, 3):
            return "Q1"
        if month in (4, 5, 6):
            return "Q2"
        if month in (7, 8, 9):
            return "Q3"
        return "Q4"

    # --- 标题 ---

    def _add_title(self):
        """标题：黑体二号居中，固定为'思想汇报'4 字

        思想汇报标题不附季度号，季度信息体现在正文开头。
        """
        custom_title = self._get("title", default="")
        title_text = custom_title if custom_title else "思想汇报"
        add_title(self.doc, title_text)
        add_blank_paragraph(self.doc)

    # --- 称呼 ---

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号，固定为'敬爱的党组织：'"""
        salutation = self._get("salutation", default="敬爱的党组织：")
        add_salutation_paragraph(self.doc, salutation)

    # --- 正文开头（过渡段，100~150 字） ---

    def _add_opening(self):
        """开头过渡段：[季度] + [身份] + [汇报内容概述] + [恳请批评指正]

        季度字段动态填充 Q1/Q2/Q3/Q4 对应的过渡句。
        身份字段取 identity_status，默认'入党积极分子'。
        """
        opening_text = self._get("opening_text", default="")
        if opening_text:
            self.add_para(opening_text)
            return

        quarter = self._resolve_quarter()
        identity = self._get("identity_status", default="入党积极分子")
        # 季度过渡句
        quarter_intro = self.QUARTER_OPENING.get(quarter, "本季度已接近尾声。")
        # 拼装完整开头
        opening = f"{quarter_intro}我作为一名{identity}，" "谨向党组织汇报本季度在理论学习、学习工作、思想动态、下一步努力方向等方面的基本情况，" "恳请党组织批评指正。"
        self.add_para(opening)

    # --- 第一段：近期理论学习情况（600~700 字，重点） ---

    def _add_theory_study(self):
        """近期理论学习段：原著学习 + 时政热点 + 心得体会

        若用户提供 theory_study_text 字段则直接使用，
        否则按"学习概况+原著心得+时政体会+理论收获"4 子段拼装。
        季度字段决定时政热点的默认填充。
        """
        theory_text = self._get("theory_study_text", default="")
        if theory_text:
            self.add_para(theory_text)
            return

        # 段首引导句
        lead = self._get("theory_lead", default="")
        if not lead:
            lead = "一、近期理论学习情况"
        self.add_para(lead)

        # 子段 1：本季度学习概况（80~100 字）
        overview = self._get("theory_overview", default="")
        if not overview:
            overview = self._build_theory_overview()
        self.add_para(overview)

        # 子段 2：原著学习心得（200~250 字）
        book_insight = self._get("theory_book_insight", default="")
        if not book_insight:
            book_insight = self._build_book_insight()
        self.add_para(book_insight)

        # 子段 3：时政热点体会（200~250 字）
        current_affairs_insight = self._get("theory_current_affairs_insight", default="")
        if not current_affairs_insight:
            current_affairs_insight = self._build_current_affairs_insight()
        self.add_para(current_affairs_insight)

        # 子段 4：理论收获总结（80~100 字）
        summary = self._get("theory_summary", default="")
        if not summary:
            summary = self._build_theory_summary()
        self.add_para(summary)

    def _build_theory_overview(self) -> str:
        """构建本季度理论学习概况段"""
        quarter = self._resolve_quarter()
        current_affairs = self.QUARTER_CURRENT_AFFAIRS.get(quarter, "党的最新理论成果")
        books = self._get_list("theory_books")
        if books:
            book_names = []
            for b in books:
                if isinstance(b, dict):
                    name, chapter = b.get("book", ""), b.get("chapter", "")
                    book_names.append(f"《{name}》{chapter}" if chapter else f"《{name}》")
                else:
                    book_names.append(f"《{str(b)}》")
            books_str = "、".join(book_names)
        else:
            books_str = "《习近平著作选读》第一卷'关于坚持和发展中国特色社会主义的几个问题'一篇"
        return f"本季度，我重点学习了{books_str}，并系统学习了{current_affairs}。" "通过原著精读与时政学习相结合，我努力用党的最新理论成果武装头脑，" "进一步坚定了对中国特色社会主义道路的信念。"

    def _build_book_insight(self) -> str:
        """构建原著学习心得段（200~250 字），优先用用户提供的 reflection"""
        for b in self._get_list("theory_books"):
            if isinstance(b, dict) and b.get("reflection"):
                return str(b["reflection"])
        return "原著学习方面，我反复研读了《习近平著作选读》第一卷中关于'道路问题是关系党的事业兴衰成败" "第一位的问题'的论述。习近平总书记指出，中国特色社会主义是社会主义，不是别的什么主义——" "这句话让我深刻认识到，我们走的道路不是简单套用马克思主义经典作家设想的模板，" "而是中国共产党人把马克思主义基本原理同中国具体实际相结合、同中华优秀传统文化相结合的伟大创造。" "在阅读中，我对照苏联解体、东欧剧变的历史教训，更深刻理解了'道路决定命运'的内涵——" "道路选择不是抽象的政治问题，而是关系国家命运的根本问题。"

    def _build_current_affairs_insight(self) -> str:
        """构建时政热点体会段（200~250 字），按季度填充，时政事件均为 2025 年真实发生"""
        quarter = self._resolve_quarter()
        current_affairs = self._get_list("current_affairs")
        if current_affairs:
            parts = [str(ca["insight"]) if isinstance(ca, dict) and ca.get("insight")
                     else (ca if isinstance(ca, str) else "") for ca in current_affairs]
            parts = [p for p in parts if p]
            if parts:
                return " ".join(parts)
        insights = {
            "Q1": "时政热点方面，我认真学习了 2025 年全国两会精神，特别是政府工作报告中" "关于'因地制宜发展新质生产力'的部署。通过学习我认识到，新质生产力不是简单的技术升级，" "而是以科技创新为核心、以战略性新兴产业和未来产业为载体的生产力跃迁。" "作为一名计算机专业学生，我深感自己所学的人工智能、大数据技术正是新质生产力的重要组成部分，" "这让我对专业学习有了更强的使命感，也更加坚定了'科技报国'的志向。",
            "Q2": "时政热点方面，我认真学习了'七一'建党节相关纪念活动，特别是重温了习近平总书记在" "庆祝中国共产党成立 100 周年大会上的重要讲话。总书记指出'江山就是人民、人民就是江山'，" "让我深刻理解了党与人民群众的血肉联系。结合五四青年节总书记对青年的寄语，" "我深感作为新时代青年入党积极分子，要把'小我'融入'大我'，" "把个人理想追求融入党和国家事业之中。",
            "Q3": "时政热点方面，我认真学习了抗战胜利 80 周年纪念相关内容，并参加了学院组织的" "'重温抗战精神'主题党日活动。通过学习我认识到，伟大抗战精神是中国人民弥足珍贵的精神财富，" "是激励我们克服一切艰难险阻、为实现中华民族伟大复兴而奋斗的强大精神力量。" "结合国庆节，我系统回顾了新中国成立 76 年来取得的伟大成就，" "更加坚定了对中国特色社会主义的道路自信、理论自信、制度自信、文化自信。",
            "Q4": "时政热点方面，我持续学习了党的二十届三中全会精神，特别是关于'进一步全面深化改革、" "推进中国式现代化'的部署。通过学习我认识到，改革是推动中国式现代化的根本动力，" "作为新时代大学生，要把对改革的理解转化为专业学习的动力。同时学习了中央经济工作会议精神，" "了解了党中央对当前经济形势的科学判断和明年经济工作的总体要求，" "增强了我对国家发展前景的信心。",
        }
        return insights.get(quarter, insights["Q1"])

    def _build_theory_summary(self) -> str:
        """构建理论收获总结段（80~100 字）"""
        summary = self._get("theory_insight", default="")
        if summary:
            return summary
        return "本季度理论学习最大的收获，是进一步坚定了'四个自信'——" "中国特色社会主义道路是党和人民在长期实践中开辟出来的正确道路，" "我将以此为指引，继续深入学习党的创新理论，自觉用党的最新理论成果武装头脑。"

    # --- 第二段：学习/工作情况（300~400 字） ---

    def _add_study_work(self):
        """学习/工作情况段：学业+学生工作+社会实践+志愿服务+科研竞赛

        若用户提供 study_work_text 字段则直接使用，
        否则按子字段拼装，数字优先使用用户数据。
        """
        text = self._get("study_work_text", default="")
        if text:
            self.add_para(text)
            return

        lead = self._get("study_work_lead", default="二、学习工作情况")
        self.add_para(lead)

        # 学业部分
        academic = self._build_academic_text()
        self.add_para(academic)

        # 学生工作+社会实践+志愿服务+科研竞赛
        practice = self._build_practice_text()
        self.add_para(practice)

    def _build_academic_text(self) -> str:
        """构建学业表现段"""
        academic = self._get("academic_progress", default="")
        if academic:
            return f"学业方面，{academic}。"
        return "学业方面，本学期共修 6 门课程，加权平均分 87.3，专业排名 12/120，" "较上学期提升 3 个名次；《机器学习》《计算机网络》两门核心课程均取得 90 分以上。" "我深感作为入党积极分子，专业学习是服务国家建设的基础，必须把学业放在首位。"

    def _build_practice_text(self) -> str:
        """构建学生工作+社会实践+志愿服务+科研竞赛段"""
        student_work = self._get("student_work", default="")
        social_practice = self._get("social_practice", default="")
        volunteer = self._get("volunteer_service", default="")
        research = self._get("research_competition", default="")

        parts = []
        if student_work:
            parts.append(f"学生工作上，{student_work}")
        else:
            parts.append("学生工作上，担任班长期间组织班级'一对一'帮扶活动 2 期，累计服务同学 18 人次；组织班级赴西柏坡红色教育基地参观 1 次，撰写班级调研报告 1 份")
        if social_practice:
            parts.append(f"社会实践方面，{social_practice}")
        else:
            parts.append("社会实践方面，参加学校'返家乡'社会实践 1 次，累计服务时长 24 小时")
        if volunteer:
            parts.append(f"志愿服务方面，{volunteer}")
        else:
            parts.append("志愿服务方面，参加校园文明引导志愿 3 次，累计 12 小时")
        if research:
            parts.append(f"科研竞赛方面，{research}")
        else:
            parts.append("科研竞赛方面，作为负责人主持校级大创项目《基于对比学习的法律问答系统》，本季度完成中期答辩，准确率较立项时提升 7 个百分点")
        return "；".join(parts) + "。"

    # --- 第三段：思想动态与不足（350~450 字，重点） ---

    def _add_thought_dynamics(self):
        """思想动态与不足段：触动事件+思想变化+个人不足剖析

        若用户提供 thought_dynamics_text 字段则直接使用，
        否则按"触动事件+思想变化+不足剖析"3 子段拼装。
        """
        text = self._get("thought_dynamics_text", default="")
        if text:
            self.add_para(text)
            return

        lead = self._get("thought_dynamics_lead", default="三、思想动态与不足")
        self.add_para(lead)

        # 子段 1：触动事件（120~150 字）
        trigger = self._get("thought_trigger", default="")
        if not trigger:
            trigger = self._build_default_trigger()
        self.add_para(trigger)

        # 子段 2：思想变化（80~100 字）
        change = self._get("thought_change", default="")
        if not change:
            change = self._build_default_change()
        self.add_para(change)

        # 子段 3：个人不足剖析（150~200 字）
        shortcomings_text = self._build_shortcomings_text()
        self.add_para(shortcomings_text)

    def _build_default_trigger(self) -> str:
        """构建本季度思想触动事件段（默认示例）"""
        return "本季度我的思想有两次明显触动。第一次是参加学院组织的'看望抗战老兵'志愿活动，" "听 96 岁的张爷爷讲述 1944 年豫湘桂战役中战友为掩护他撤离而牺牲的经历。" "张爷爷说：'我们那一代人入党，不是为了个人前途，是为了让后人不再过苦日子。'" "这句话让我震动——我开始反思自己申请入党是否也抱有'个人前途'的杂念。" "第二次是聆听学院党委书记讲授的专题党课，书记讲到焦裕禄、谷文昌、廖俊波等优秀党员" "'吃苦在前、享受在后'的具体事迹，让我对'党员身份'有了更立体的认识。"

    def _build_default_change(self) -> str:
        """构建思想变化段"""
        return "两次触动让我从'个人优秀'的视角转向'集体奉献'的视角——" "党员不只是个人表现优秀，更是要带动集体、服务群众。" "这让我对'全心全意为人民服务'的宗旨有了更具体的理解：" "服务不是抽象的口号，而是具体的行动，是关心身边每一个需要帮助的同学。"

    def _build_shortcomings_text(self) -> str:
        """构建个人不足剖析段（2~3 个真实不足，含表现+根源）"""
        shortcomings = self._get_list("self_shortcomings")
        if not shortcomings:
            # 默认 3 个真实不足
            shortcomings = [
                {"desc": "理论学习系统性不够",
                 "manifestation": "按时学习规定篇目，但对马克思主义经典著作（如《共产党宣言》《资本论》第一卷）只读过简写本，未读全本；",
                 "root_cause": "学习时间碎片化，缺乏系统规划。"},
                {"desc": "联系群众不够紧密",
                 "manifestation": "担任班长期间虽履职，但对班级经济困难同学、心理压力较大同学关注不够；",
                 "root_cause": "把'履职'等同于'完成任务'，缺乏主动关心。"},
                {"desc": "志愿服务持续性不强",
                 "manifestation": "参加活动多为被动响应，缺乏主动策划；",
                 "root_cause": "对志愿服务的认识停留在'完成任务'层面，未上升到'党员责任'高度。"},
            ]
        parts = []
        for idx, item in enumerate(shortcomings, start=1):
            if isinstance(item, dict):
                desc = item.get("desc", "")
                manifest = item.get("manifestation", "")
                root = item.get("root_cause", "")
                num_word = ["一", "二", "三", "四", "五"][idx - 1] if idx <= 5 else str(idx)
                parts.append(f"{num_word}是{desc}。{manifest}{root}")
            else:
                parts.append(str(item))
        prefix = "在反思中，我也认识到自己存在以下不足："
        return prefix + "".join(parts)

    # --- 第四段：下一步努力方向（200~250 字） ---

    def _add_next_steps(self):
        """下一步努力方向段：理论学习计划+实践计划+不足改进计划

        计划必须可执行，与上段不足一一对应。
        """
        text = self._get("next_steps_text", default="")
        if text:
            self.add_para(text)
            return

        lead = self._get("next_steps_lead", default="四、下一步努力方向")
        self.add_para(lead)

        # 理论学习计划
        theory_plan = self._get("next_theory_plan", default="")
        if not theory_plan:
            theory_plan = "理论学习方面，计划通读《习近平著作选读》第二卷，每月撰写 1 篇读书笔记；" "下季度末前通读《共产党宣言》全本，并向培养联系人汇报。"
        self.add_para(theory_plan)

        # 实践计划
        practice_plan = self._get("next_practice_plan", default="")
        if not practice_plan:
            practice_plan = "实践方面，主动报名学校'党员先锋岗'，每月至少参加 1 次志愿服务；" "暑期参加三下乡支教 1 次，重点服务乡村青少年。"
        self.add_para(practice_plan)

        # 不足改进计划
        improvement_plan = self._get("next_improvement_plan", default="")
        if not improvement_plan:
            improvement_plan = "联系群众方面，每月走访班级宿舍 2 次，重点关注经济困难和心理压力较大同学，" "及时向辅导员反馈。我将以党员标准严格要求自己，自觉接受党组织考验。"
        self.add_para(improvement_plan)

    # --- 结尾"恳请党组织批评指正。此致 敬礼！" ---

    def _add_ending(self):
        """结尾：恳请党组织批评指正 + 此致 + 敬礼！

        与入党申请书结尾不同：思想汇报结尾句为'恳请党组织批评指正。'，
        非'请党组织在实践中考验我！'。
        """
        ending_line = self._get("ending_line", default="恳请党组织批评指正。")
        self.add_para(ending_line)
        add_cizhi_paragraph(self.doc, "此致")
        add_jingli_paragraph(self.doc, "敬礼！")

    # --- 落款（右对齐） ---

    def _add_signature(self):
        """落款：汇报人 + 日期 + 培养联系人，右对齐

        与入党申请书落款不同：思想汇报落款为'汇报人'，非'申请人'。
        培养联系人一行可选，部分学校要求签字栏。
        """
        add_blank_paragraph(self.doc)
        name = self._get("name", default="汇报人")
        submit_date = self._get("submit_date", default="")
        cultivation_contact = self._get("cultivation_contact", default="")
        add_right_aligned_paragraph(self.doc, f"汇报人：{name}")
        if submit_date:
            add_right_aligned_paragraph(self.doc, submit_date)
        if cultivation_contact:
            add_right_aligned_paragraph(self.doc, f"培养联系人：{cultivation_contact}")

    # --- 主构建方法 ---

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/开头/4 段正文/结尾/落款，生成 docx

        四段正文为思想汇报核心：理论学习（重点）/ 学习工作 / 思想动态与不足（重点）/ 下一步努力方向。
        """
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_opening()
            self._add_theory_study()
            self._add_study_work()
            self._add_thought_dynamics()
            self._add_next_steps()
            self._add_ending()
            self._add_signature()
            self._post_build_checks()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 思想汇报已生成：{output_path}")
        return str(output_path)

    def _post_build_checks(self):
        """构建后检查：政治用语规范 + 查重风险检测 + 字数检查"""
        full_text = "".join(self.full_text_parts)

        # 政治用语规范检查
        for w in check_political_terms(full_text):
            self.warnings.append(f"[政治规范] {w}")

        # 三重查重检测：党章原文 50 字 / 网络模板 30 字 / 历史汇报 50 字
        plagiarism_checks = [
            ("查重风险-党章原文", PARTY_CONSTITUTION_FRAGMENTS, 50),
            ("查重风险-网络模板", NETWORK_TEMPLATE_FRAGMENTS, 30),
            ("查重风险-历史汇报", HISTORICAL_REPORT_FRAGMENTS, 50),
        ]
        for label, frags, thr in plagiarism_checks:
            for w in check_plagiarism_risk(full_text, frags, threshold=thr):
                self.warnings.append(f"[{label}] {w}")

        # 字数检查（思想汇报目标 1500~2000 字）
        char_count = count_chinese_chars(full_text)
        if char_count < 1500:
            self.warnings.append(f"[字数] 全文仅 {char_count} 字，建议 1500~2000 字")
        elif char_count > 2000:
            self.warnings.append(f"[字数] 全文 {char_count} 字偏多，建议压缩至 2000 字以内")

        if self.warnings:
            print("⚠️ 构建后检查警告：", file=sys.stderr)
            for w in self.warnings:
                print(f"  - {w}", file=sys.stderr)

    # --- 数据校验 ---

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）

        包含政治规范校验：入党动机红线、假缺点检测、个人不足数量校验。
        """
        warnings = []

        # P0 必采字段
        p0_fields = [
            ("name", "汇报人姓名"), ("college", "学院"), ("major", "专业"),
            ("grade", "年级"), ("identity_status", "政治身份（入党积极分子/预备党员）"),
            ("submit_date", "汇报日期"), ("party_branch", "党支部全称"),
            ("cultivation_contact", "培养联系人姓名"),
        ]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        # 季度字段校验
        quarter = str(self._get("quarter", default="")).upper().strip()
        if quarter and quarter not in ("Q1", "Q2", "Q3", "Q4"):
            warnings.append(f"季度字段'{quarter}'不规范，应为 Q1/Q2/Q3/Q4 之一")

        # 身份字段校验
        identity = str(self._get("identity_status", default=""))
        if identity and identity not in ("入党积极分子", "预备党员", "发展对象"):
            warnings.append(f"政治身份'{identity}'非标准表述，" "应为'入党积极分子'/'预备党员'/'发展对象'之一")

        # 个人不足数量校验
        shortcomings = self._get_list("self_shortcomings")
        if shortcomings and len(shortcomings) < 2:
            warnings.append(f"个人不足仅 {len(shortcomings)} 个，建议 2~3 个真实不足")

        # 假缺点检测（党支部一眼识破，会要求重写）
        fake_shortcomings_keywords = [
            "工作太投入", "追求完美", "学习太刻苦", "责任心太强",
            "为人太直率", "事必躬亲", "太较真", "工作太认真"]
        shortcomings_text = self._get("thought_dynamics_text", default="")
        for item in self._get_list("self_shortcomings"):
            if isinstance(item, dict):
                shortcomings_text += str(item.get("desc", "")) + str(item.get("manifestation", ""))
        for kw in fake_shortcomings_keywords:
            if kw in shortcomings_text:
                warnings.append(f"[政治红线] 个人不足出现假缺点'{kw}'，" "党支部会要求重写，请用户重新表述")
                break

        # 时政热点真实性提醒（仅警告，不阻断）
        current_affairs = self._get_list("current_affairs")
        if not current_affairs and not self._get("theory_current_affairs_insight"):
            quarter = self._resolve_quarter()
            warnings.append(f"[时政提醒] 未提供 current_affairs 字段，" f"将使用季度 {quarter} 默认时政，请用户核实替换为本人真实关注的事件")

        self.warnings.extend(warnings)
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ==================== 默认示例数据（Q2 思想汇报，2025 年 6 月提交） ====================

DEFAULT_DATA = {
    "name": "张明", "student_id": "2022123456", "gender": "男",
    "college": "计算机科学与技术学院", "major": "计算机科学与技术",
    "grade": "2022 级 大三", "class_name": "计科 2201 班",
    "identity_status": "入党积极分子", "political_status": "共青团员",
    "salutation": "敬爱的党组织：", "title": "思想汇报",
    "quarter": "Q2", "submit_date": "2025 年 6 月 25 日",
    "party_branch": "计算机科学与技术学院本科生第一党支部",
    "cultivation_contact": "李志强",
    "theory_books": [
        {"book": "习近平著作选读",
         "chapter": "第一卷'关于坚持和发展中国特色社会主义的几个问题'",
         "reflection": "原著学习方面，我反复研读了《习近平著作选读》第一卷中关于" "'道路问题是关系党的事业兴衰成败第一位的问题'的论述。" "习近平总书记指出，中国特色社会主义是社会主义，不是别的什么主义——" "这句话让我深刻认识到，我们走的道路不是简单套用马克思主义经典作家设想的模板，" "而是中国共产党人把马克思主义基本原理同中国具体实际相结合、" "同中华优秀传统文化相结合的伟大创造。" "在阅读中，我对照苏联解体、东欧剧变的历史教训，更深刻理解了'道路决定命运'的内涵。"},
    ],
    "current_affairs": [
        {"event": "2025 年'七一'建党节相关纪念活动", "approach": "重温总书记在庆祝中国共产党成立 100 周年大会上的重要讲话",
         "insight": "时政热点方面，我认真学习了'七一'建党节相关纪念活动，" "特别是重温了习近平总书记在庆祝中国共产党成立 100 周年大会上的重要讲话。" "总书记指出'江山就是人民、人民就是江山'，" "让我深刻理解了党与人民群众的血肉联系。结合五四青年节总书记对青年的寄语，" "我深感作为新时代青年入党积极分子，要把'小我'融入'大我'，" "把个人理想追求融入党和国家事业之中。"},
    ],
    "academic_progress": "本学期共修 6 门课程，加权平均分 87.3，专业排名 12/120",
    "student_work": "担任班长期间组织班级'一对一'帮扶活动 2 期，累计服务同学 18 人次",
    "social_practice": "参加学校'返家乡'社会实践 1 次，累计服务时长 24 小时",
    "volunteer_service": "参加校园文明引导志愿 3 次，累计 12 小时",
    "research_competition": "作为负责人主持校级大创项目《基于对比学习的法律问答系统》，" "本季度完成中期答辩，准确率较立项时提升 7 个百分点",
    "thought_trigger": "本季度我的思想有两次明显触动。第一次是参加学院组织的'看望抗战老兵'志愿活动，" "听 96 岁的张爷爷讲述 1944 年豫湘桂战役中战友为掩护他撤离而牺牲的经历。" "张爷爷说：'我们那一代人入党，不是为了个人前途，是为了让后人不再过苦日子。'" "这句话让我震动——我开始反思自己申请入党是否也抱有'个人前途'的杂念。",
    "thought_change": "这次触动让我从'个人优秀'的视角转向'集体奉献'的视角——" "党员不只是个人表现优秀，更是要带动集体、服务群众。" "这让我对'全心全意为人民服务'的宗旨有了更具体的理解。",
    "self_shortcomings": [
        {"desc": "理论学习系统性不够",
         "manifestation": "对马克思主义经典著作只读过简写本，未读全本；",
         "root_cause": "学习时间碎片化，缺乏系统规划。"},
        {"desc": "联系群众不够紧密",
         "manifestation": "对班级经济困难同学、心理压力较大同学关注不够；",
         "root_cause": "把'履职'等同于'完成任务'，缺乏主动关心。"},
        {"desc": "志愿服务持续性不强",
         "manifestation": "参加活动多为被动响应，缺乏主动策划。",
         "root_cause": "对志愿服务的认识停留在'完成任务'层面，未上升到'党员责任'高度。"},
    ],
    "next_theory_plan": "理论学习方面，计划通读《习近平著作选读》第二卷，每月撰写 1 篇读书笔记；" "9 月底前通读《共产党宣言》全本，并向培养联系人汇报。",
    "next_practice_plan": "实践方面，主动报名学校'党员先锋岗'，每月至少参加 1 次志愿服务；" "暑期参加三下乡支教 1 次。",
    "next_improvement_plan": "联系群众方面，每月走访班级宿舍 2 次，重点关注经济困难和心理压力较大同学，" "及时向辅导员反馈。我将以党员标准严格要求自己，自觉接受党组织考验。",
}


# ==================== CLI 入口 ====================

def main():
    parser = argparse.ArgumentParser(
        description="思想汇报 docx 生成器（每季度一份，结合时政热点）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="示例：\n" "  python build.py --data data.json --out output.docx\n" "  python build.py --demo --out demo.docx\n" "\nJSON 字段定义详见 SKILL.md 第三章信息采集清单。\n" "政治规范校验：必提 5 项要点 + 指导思想 6 项顺序 + 三重查重检测。\n" "季度字段 quarter 取 Q1/Q2/Q3/Q4，决定时政热点填充与过渡句措辞。",
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档（Q2 思想汇报，2025 年 6 月）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
