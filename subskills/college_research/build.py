#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
大学生院级科研立项申请书 docx 生成器（v2.1 案例优化版）

A4 纸张，页边距上下 2.54cm 左右 2.5cm；正文宋体小四 1.5 倍行距首行缩进 2 字符；
一级标题黑体三号居中；二级标题黑体小三；表格宋体五号；参考文献宋体五号 GB/T 7714。

v2.1 新增 9 项强化章节（基于案例 1 简化版）：
  1. 国家政策引用（policy_citations，≥3 条）
  2. 科学挑战（scientific_challenges，1~2 个）
  3. 国内外研究现状综述（literature_review，≥10 条含 ≥3 篇英文）
  4. 方法对比表（method_comparison，2 方法 × 3 维度）
  5. 技术路线图（tech_roadmap，1 张）
  6. 数学公式（formulas，≥1，可选）
  7. 社会效益量化（social_benefits，≥3 项）
  8. 进度安排（project_schedule，4 阶段 6 个月）
  9. JSON Schema + build.py + 质检清单（本文件即实现）

院级科研立项特征（v2.1 调整）：经费 500~2000 元、周期 3~6 个月、团队 1~2 人、
GB/T 7714 参考文献 ≥5 条覆盖 ≥3 类含 ≥3 篇英文。JSON 字段详见 SKILL.md 第十七章。

使用：python build.py --data data.json --out output.docx
      python build.py --demo --out demo.docx
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# 字体与格式常量

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四（正文）
SIZE_WU = Pt(10.5)          # 五号（表格、参考文献）
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# 工具函数

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name,
                         font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    pf.first_line_indent = Pt(0)
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name,
                 font_size=font_size, bold=bold)


def add_paragraph_with_format(doc, text: str, font_name: str = FONT_SONG,
                              font_size=SIZE_XIAO_SI, bold: bool = False,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT,
                              first_line_indent: bool = True,
                              line_spacing: float = 1.5,
                              space_before: float = 0,
                              space_after: float = 0):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name,
                 font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
        space_before=6, space_after=6)


def add_heading_level3(doc, text: str):
    """三级标题：宋体四号加粗，左对齐"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
        space_before=6, space_after=3)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=indent,
        line_spacing=1.5)


def add_reference_paragraph(doc, idx: int, text: str):
    """参考文献段落：宋体五号，悬挂缩进，单倍行距"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Pt(-18)
    pf.left_indent = Pt(18)
    pf.line_spacing = 1.25
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(f"[{idx}] {text}")
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)
    return p


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def merge_vertical_cells(table, col_idx: int, start_row: int, end_row: int):
    """纵向合并单元格（用于签字栏预留空白）"""
    cells = [table.rows[r].cells[col_idx] for r in range(start_row, end_row + 1)]
    merged = cells[0]
    for c in cells[1:]:
        merged = merged.merge(c)


# GB/T 7714 参考文献格式化（院级科研立项核心，5~8 条覆盖 3 类含 3 篇英文）

def format_authors(authors_str: str) -> str:
    """格式化作者：≤3 全列；>3 列前 3 + 等/et al."""
    if not authors_str:
        return ""
    authors = [a.strip() for a in authors_str.split(",") if a.strip()]
    if not authors:
        return authors_str
    is_en = any(c.isalpha() and ord(c) < 128 for c in authors[0])
    if len(authors) <= 3:
        return ", ".join(authors)
    suffix = "et al." if is_en else "等"
    return ", ".join(authors[:3]) + ", " + suffix


def format_journal_ref(idx: int, ref: Dict[str, Any]) -> str:
    """[J] 期刊文章：作者. 题名[J]. 刊名, 年, 卷(期): 起止页码."""
    g = ref.get
    a = format_authors(g("authors", ""))
    s = f"{a}. {g('title', '')}[J]. {g('journal', '')}, {g('year', '')}"
    v, i, p = g("volume", ""), g("issue", ""), g("pages", "")
    if v:
        s += f", {v}{i}"
    if p:
        s += f": {p}"
    return s + "."


def format_conference_ref(idx: int, ref: Dict[str, Any]) -> str:
    """[C] 会议论文：作者. 题名[C]//论文集名. 出版地: 出版者, 年: 页码."""
    g = ref.get
    a = format_authors(g("authors", ""))
    s = (f"{a}. {g('title', '')}[C]//{g('conference', '')}. "
         f"{g('city', '')}: {g('publisher', '')}, {g('year', '')}")
    p = g("pages", "")
    if p:
        s += f": {p}"
    return s + "."


def format_book_ref(idx: int, ref: Dict[str, Any]) -> str:
    """[M] 专著：作者. 书名[M]. 版本. 出版地: 出版者, 年: 页码."""
    g = ref.get
    a = format_authors(g("authors", ""))
    ed = g("edition", "")
    s = f"{a}. {g('title', '')}[M]. "
    if ed:
        s += f"{ed}. "
    s += f"{g('city', '')}: {g('publisher', '')}, {g('year', '')}"
    p = g("pages", "")
    if p:
        s += f": {p}"
    return s + "."


def format_thesis_ref(idx: int, ref: Dict[str, Any]) -> str:
    """[D] 学位论文：作者. 题名[D]. 学位授予地: 学位授予单位, 年."""
    g = ref.get
    a = format_authors(g("authors", ""))
    return (f"{a}. {g('title', '')}[D]. "
            f"{g('city', '')}: {g('school', '')}, {g('year', '')}.")


def format_web_ref(idx: int, ref: Dict[str, Any]) -> str:
    """[EB/OL] 或 [Z/OL] 网络资源：作者. 题名[T/OL]. (发布日期)[访问日期]. URL."""
    g = ref.get
    a = format_authors(g("authors", ""))
    tag = "Z/OL" if g("is_government", False) else "EB/OL"
    pd, ad, url = g("publish_date", ""), g("access_date", ""), g("url", "")
    s = f"{a}. {g('title', '')}[{tag}]. "
    if pd:
        s += f"({pd})"
    if ad:
        s += f"[{ad}]"
    if url:
        s += f". {url}"
    return s + "."


def format_reference(idx: int, ref: Dict[str, Any]) -> str:
    """根据 ref_type 调度对应格式化函数"""
    rt = ref.get("ref_type", "journal").lower()
    dispatch = {
        "journal": format_journal_ref, "j": format_journal_ref,
        "conference": format_conference_ref, "c": format_conference_ref,
        "book": format_book_ref, "m": format_book_ref,
        "thesis": format_thesis_ref, "d": format_thesis_ref,
        "web": format_web_ref, "eb": format_web_ref,
    }
    fn = dispatch.get(rt, format_journal_ref)
    return fn(idx, ref)


# v2.1 新增：国家政策引用格式化

def format_policy_citation(idx: int, policy: Dict[str, Any]) -> str:
    """格式化国家政策引用：[P序号] 发文机关. 文件标题[Z/OL]. (发布日期)[访问日期]. URL."""
    g = policy.get
    agency = g("agency", "")
    title = g("title", "")
    doc_no = g("doc_no", "")
    date = g("date", "")
    excerpt = g("excerpt", "")
    url = g("url", "")
    access_date = g("access_date", "")
    parts = []
    if agency:
        parts.append(agency)
    title_full = title
    if doc_no:
        title_full = f"《{title}》({doc_no})"
    elif title:
        title_full = f"《{title}》"
    parts.append(f"{date} 发布{title_full}" if date else f"发布{title_full}")
    if excerpt:
        parts.append(f"，指出：{excerpt}")
    s = "，".join(parts[:-1]) if len(parts) > 1 else parts[0]
    if excerpt:
        s += parts[-1]
    return f"[P{idx}] {s}。"


# ApplicationDocBuilder 主类

class ApplicationDocBuilder:
    """大学生院级科研立项申请书 docx 构建器（v2.1 案例优化版）。

    v2.1 新增 9 项强化章节：国家政策引用 / 科学挑战 / 国内外研究现状综述 /
    方法对比表 / 技术路线图 / 数学公式 / 社会效益量化 / 进度安排 4 阶段 6 个月 /
    JSON Schema + build.py + 质检清单。

    院级 vs 校级：经费 500~2000 元、周期 3~6 月、团队 1~2 人、参考文献 ≥5 条含 3 篇英文。
    """

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def add_h1(self, text):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text):
        return add_heading_level2(self.doc, text)

    def add_h3(self, text):
        return add_heading_level3(self.doc, text)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_table(self, headers, rows, col_widths=None):
        return add_table_from_data(self.doc, headers, rows, col_widths)

    def add_page_break(self):
        add_page_break(self.doc)

    # 一、封面

    def _add_cover(self):
        """封面：黑体二号标题 + 4 行下划线信息（无项目类型字段，标题含"院级"二字）"""
        for _ in range(3):
            self.doc.add_paragraph()
        college = self._get("college", default="XX 大学")
        title = f"{college}大学生院级科研立项申请书"
        add_paragraph_with_format(
            self.doc, title, font_name=FONT_HEI, font_size=SIZE_ER,
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False, space_before=12, space_after=12)
        for _ in range(2):
            self.doc.add_paragraph()
        info_items = [
            ("课题名称", self._get("project_name")),
            ("课题负责人", self._get("leader_name")),
            ("所在学院", self._get("college")),
            ("申报日期", self._get("apply_date")),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI, font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True
        self.add_page_break()

    # 一、课题负责人及团队成员信息

    def _add_leader_members(self):
        """一、课题负责人及团队成员信息（负责人表 + 成员表，1~2 人）"""
        self.add_h1("一、课题负责人及团队成员信息")
        self.add_h2("（一）课题负责人信息")
        leader_rows = [
            ["姓名", self._get("leader_name"), "学号", self._get("leader_id")],
            ["性别", self._get("leader_gender", default="男"), "专业年级", self._get("leader_major")],
            ["学院", self._get("college"), "联系电话", self._get("leader_phone")],
            ["邮箱", self._get("leader_email"), "指导教师", self._get("advisor_name")],
        ]
        self.add_table(["字段", "内容", "字段", "内容"], leader_rows,
                       col_widths=[2.5, 5.5, 2.5, 5.5])
        self.add_h2("（二）团队成员信息")
        members = self._get("team_members", default=[])
        if isinstance(members, list) and members:
            rows = []
            for m in members:
                if not isinstance(m, dict):
                    continue
                rows.append([m.get("name", ""), m.get("id", ""),
                             m.get("major", ""), m.get("role", ""),
                             m.get("phone", "")])
            self.add_table(
                ["姓名", "学号", "专业年级", "分工", "联系方式"], rows,
                col_widths=[2.0, 2.5, 3.5, 5.0, 3.0])
        else:
            self.add_para("（团队成员 1~2 人，含分工。院级科研立项团队规模 1~2 人，请填写姓名/学号/专业年级/分工/联系方式。）")

    # 二/三/四、课题名称/起止时间/研究类型

    def _add_basic_info(self):
        """二、课题名称；三、起止时间；四、研究类型（v2.1 院级周期 3~6 个月）"""
        self.add_h1("二、课题名称")
        project_name = self._get("project_name", default="")
        add_paragraph_with_format(
            self.doc, project_name if project_name else
            "（不超过 30 字，突出『做什么+为谁做』，禁用『基于...的...』堆砌，紧扣本院专业方向）",
            font_name=FONT_HEI, font_size=SIZE_SAN, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_after=6)
        self.add_h1("三、研究起止时间")
        duration = self._get("basic_info", "duration", default="")
        if not duration:
            duration = self._get("duration", default="2025.04-2025.09（共 6 个月）")
        add_paragraph_with_format(
            self.doc, duration, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_after=6)
        self.add_h1("四、研究类型")
        research_type = (self._get("basic_info", "research_type", default="")
                         or self._get("research_type", default="应用研究"))
        type_options = ["基础研究", "应用研究", "开发研究"]
        type_str = "    ".join(
            f"{'☑' if t == research_type else '☐'} {t}" for t in type_options)
        add_paragraph_with_format(
            self.doc, type_str, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False)

    # 五、课题研究的背景与意义

    def _add_background(self):
        """五、课题研究的背景与意义（院级 400~700 字，4 段）"""
        self.add_h1("五、课题研究的背景与意义")
        background = self._get("background", default=[])
        if isinstance(background, str):
            background = [background]
        sub_titles = ["（一）时代背景", "（二）现实痛点", "（三）研究缺口", "（四）课题意义"]
        placeholders = [
            "（请填写时代背景，80~120 字，2~3 句政策/行业/学术趋势，必须含权威数据来源，聚焦本院学生群体。）",
            "（请填写现实痛点，150~250 字，2~3 个真实场景，必须可量化，聚焦本院学生群体。）",
            "（请填写研究缺口，80~120 字，简要评述已有研究不足，详细评述见六、研究现状及 v2.1 强化综述。）",
            "（请填写课题意义，80~150 字，理论/实践/社会三角度，至少两个。）",
        ]
        if not background:
            for t, p in zip(sub_titles, placeholders):
                self.add_h2(t)
                self.add_para(p)
        else:
            for i, para in enumerate(background):
                if i < len(sub_titles):
                    self.add_h2(sub_titles[i])
                self.add_para(para)

    # v2.1 新增：五点五、国家政策引用（≥3 条文件）

    def _add_policy_citations(self):
        """v2.1 新增：国家政策引用（≥3 条文件，按时间倒序，含 5 要素）"""
        self.add_h1("五点五、国家政策引用（v2.1 强化）")
        policies = self._get("policy_citations", default=[])
        if not isinstance(policies, list):
            policies = []
        if not policies:
            self.add_para(
                "（v2.1 要求 ≥3 条国家政策引用，按时间倒序排列。每条含发文机关+文号+标题+时间+与课题相关的关键表述 5 要素。"
                "示例：2024 年 9 月，中国心理学会发布《大学生学习行为蓝皮书》，指出全国高校学生学习拖延行为发生率达 67.8%。）")
            return
        # 按 date 倒序排序
        try:
            policies_sorted = sorted(
                policies, key=lambda p: p.get("date", ""), reverse=True)
        except Exception:
            policies_sorted = policies
        for i, policy in enumerate(policies_sorted, 1):
            if not isinstance(policy, dict):
                continue
            try:
                line = format_policy_citation(i, policy)
            except Exception as e:
                line = f"[P{i}] （格式化失败：{e}）"
            add_reference_paragraph(self.doc, i, line.replace(f"[P{i}] ", ""))
        if len(policies) < 3:
            self.warnings.append(
                f"国家政策引用仅 {len(policies)} 条，v2.1 要求 ≥3 条")

    # 六、国内外研究现状及发展动态（简版）

    def _add_research_status(self):
        """六、国内外研究现状及发展动态（院级 600~1000 字，3 段评述式，简版）"""
        self.add_h1("六、国内外研究现状及发展动态（简版）")
        status = self._get("research_status", default=[])
        if isinstance(status, str):
            status = [status]
        sub_titles = ["（一）国内研究现状", "（二）国外研究现状", "（三）综合评述与本课题差异"]
        placeholders = [
            "（请填写国内研究现状，200~350 字，引用 2~3 篇国内文献，每篇含『作者+年份+观点+局限』并标注 [1][2] 等编号。禁止罗列式。详尽版见 §6.5。）",
            "（请填写国外研究现状，200~350 字，引用 2~3 篇国外文献，同上结构，标注 [3][4] 等编号。详尽版见 §6.5。）",
            "（请填写综合评述，150~300 字，归纳国内外研究共性问题 2~3 个，提出本课题差异点（与共性问题一一对应），引出创新点。）",
        ]
        if not status:
            for t, p in zip(sub_titles, placeholders):
                self.add_h2(t)
                self.add_para(p)
        else:
            for i, para in enumerate(status):
                if i < len(sub_titles):
                    self.add_h2(sub_titles[i])
                self.add_para(para)

    # v2.1 新增：六点五、国内外研究现状综述详尽版（≥10 条含 ≥3 篇英文）

    def _add_literature_review(self):
        """v2.1 新增：国内外研究现状综述详尽版（≥10 条文献含 ≥3 篇英文，4 段结构）"""
        self.add_h1("六点五、国内外研究现状综述详尽版（v2.1 强化）")
        review = self._get("literature_review", default={})
        if not isinstance(review, dict) or not review:
            self.add_para(
                "（v2.1 要求 ≥10 条文献综述，含 ≥3 篇英文文献。4 段结构：研究意义/国内研究现状/国外研究现状/总结与分析。"
                "每条含『作者+年份+观点+局限』，归纳 2~3 个 gap。详例见 SKILL.md §8.3。）")
            return
        # 段 1：研究意义
        significance = review.get("significance", [])
        if significance:
            self.add_h2("（一）研究意义")
            for s in significance:
                self.add_para(s)
        # 段 2：国内研究现状
        domestic = review.get("domestic", [])
        if domestic:
            self.add_h2("（二）国内研究现状")
            for d in domestic:
                self.add_para(d)
        # 段 3：国外研究现状
        foreign = review.get("foreign", [])
        if foreign:
            self.add_h2("（三）国外研究现状")
            for f in foreign:
                self.add_para(f)
        # 段 4：总结与分析
        summary = review.get("summary", "")
        if summary:
            self.add_h2("（四）对现有研究的总结与分析")
            self.add_para(summary)
        # 统计文献条数与英文数
        all_items = (significance or []) + (domestic or []) + (foreign or [])
        item_count = len([x for x in all_items if isinstance(x, str) and "[" in x])
        if item_count > 0 and item_count < 10:
            self.warnings.append(
                f"文献综述仅约 {item_count} 条引用，v2.1 要求 ≥10 条")
        # 统计英文文献数
        refs = self._get("references", default=[])
        if isinstance(refs, list):
            en_count = sum(1 for r in refs
                          if isinstance(r, dict) and r.get("is_english", False))
            # 推断英文：作者包含 ASCII 字母
            if en_count == 0:
                en_count = sum(1 for r in refs
                              if isinstance(r, dict)
                              and any(c.isalpha() and ord(c) < 128
                                      for c in str(r.get("authors", ""))[:20]))
            if en_count < 3:
                self.warnings.append(
                    f"英文文献仅 {en_count} 篇，v2.1 要求 ≥3 篇")

    # 七、研究目标、研究内容、拟解决的关键问题

    def _add_research_content(self):
        """七、研究目标、研究内容、拟解决的关键问题（院级 600~1000 字，3 子节）"""
        self.add_h1("七、研究目标、研究内容、拟解决的关键问题")
        self.add_h2("（一）研究目标")
        goal = self._get("research_goal", default="")
        if goal:
            self.add_para(goal)
        else:
            self.add_para("（请填写研究目标，150~250 字，1 个总目标 + 2~3 个阶段目标，全部可量化。）")
        self.add_h2("（二）研究内容")
        contents = self._get("research_content", default=[])
        if isinstance(contents, str):
            contents = [contents]
        if contents:
            for i, c in enumerate(contents, 1):
                self.add_para(f"{i}. {c}")
        else:
            self.add_para("（请填写研究内容，2~4 个子任务，每个 100~150 字，结构：任务名+做什么+方法+预期产出。）")
        self.add_h2("（三）拟解决的关键问题")
        problems = self._get("key_problems", default=[])
        if isinstance(problems, str):
            problems = [problems]
        if problems:
            for i, q in enumerate(problems, 1):
                self.add_para(f"{i}. {q}")
        else:
            self.add_para("（请填写关键问题，1~2 个，每个一句话讲清技术难点。v2.1 强化：详见 §7.5 科学挑战。）")

    # v2.1 新增：七点五、科学挑战（1~2 个）

    def _add_scientific_challenges(self):
        """v2.1 新增：科学挑战（1~2 个，一句话标题+3~5 句描述+子挑战+文献引用）"""
        self.add_h1("七点五、科学挑战（v2.1 强化）")
        challenges = self._get("scientific_challenges", default=[])
        if not isinstance(challenges, list):
            challenges = []
        if not challenges:
            self.add_para(
                "（v2.1 要求 1~2 个科学挑战，每个独立成段。结构：一句话标题（含关键技术词）+ 3~5 句描述 + 子挑战 + 文献引用 [N]。"
                "院级简化版每个挑战可只含 1 个子挑战。详例见 SKILL.md §7.2。）")
            return
        for i, ch in enumerate(challenges, 1):
            if not isinstance(ch, dict):
                continue
            title = ch.get("title", "")
            description = ch.get("description", "")
            sub_challenges = ch.get("sub_challenges", [])
            references = ch.get("references", [])
            # 标题
            cn_num = ["一", "二", "三"][i - 1] if i <= 3 else str(i)
            self.add_h2(f"科学挑战{cn_num}：{title}")
            # 描述段
            ref_str = ""
            if isinstance(references, list) and references:
                ref_str = " [" + ",".join(str(r) for r in references) + "]"
            elif references:
                ref_str = f" [{references}]"
            desc_text = description + ref_str if description else ref_str
            if desc_text:
                self.add_para(desc_text)
            # 子挑战
            if isinstance(sub_challenges, list) and sub_challenges:
                for j, sc in enumerate(sub_challenges, 1):
                    self.add_para(f"● 子挑战 {j}：{sc}")
        if len(challenges) < 1:
            self.warnings.append(
                f"科学挑战仅 {len(challenges)} 个，v2.1 要求 ≥1 个")

    # 八、研究方案及技术路线

    def _add_research_scheme(self):
        """八、研究方案及技术路线（含方法、步骤、流程图，院级 350~600 字）"""
        self.add_h1("八、研究方案及技术路线")
        self.add_h2("（一）总体技术路线")
        route = self._get("tech_route", default="")
        if route:
            self.add_para(route)
        else:
            self.add_para("（请填写总体技术路线，120~200 字 + 流程图，4 阶段流程，每阶段标注交付物。v2.1 强化：详细图见 §8.7 技术路线图。）")
        self.add_h2("（二）研究方法")
        methods = self._get("methods", default=[])
        if isinstance(methods, str):
            methods = [methods]
        if methods:
            for i, m in enumerate(methods, 1):
                self.add_para(f"{i}. {m}")
        else:
            self.add_para("（请填写研究方法，3~4 个，每个 50~80 字说明用途。）")
        self.add_h2("（三）数据来源与实验条件")
        data_src = self._get("data_source", default="")
        self.add_para(data_src if data_src else
                      "（请填写数据来源、规模、实验设备型号、软件工具及版本。）")

    # v2.1 新增：八点五、方法对比表（2 方法 × 3 维度）

    def _add_method_comparison(self):
        """v2.1 新增：方法对比表（≥2 方法 × 3 维度，有选型结论）"""
        self.add_h1("八点五、方法对比表（v2.1 强化）")
        mc = self._get("method_comparison", default={})
        if not isinstance(mc, dict) or not mc:
            self.add_para(
                "（v2.1 要求 ≥2 方法 × 3 维度对比表。3 维度推荐：技术特点/优势/劣势。"
                "每个方法必须有文献引用，对比后必须有选型结论。详例见 SKILL.md §9.2。）")
            return
        headers = mc.get("table_headers", ["维度", "方法 A", "方法 B"])
        methods = mc.get("methods", [])
        if not isinstance(methods, list) or not methods:
            self.add_para("（方法对比表数据为空，请补充 ≥2 个方法。）")
            return
        # 构造表格行：每行 = [维度名, 方法1该维度内容, 方法2该维度内容, ...]
        dimensions = ["技术特点", "优势", "劣势"]
        # 从 method 结构中提取维度键
        if isinstance(methods[0], dict) and "features" in methods[0]:
            feature_keys = list(methods[0].get("features", {}).keys())
            if feature_keys:
                dimensions = feature_keys[:3]
        rows = []
        for dim in dimensions:
            row = [dim]
            for m in methods:
                if not isinstance(m, dict):
                    row.append("")
                    continue
                features = m.get("features", {})
                val = features.get(dim, "") if isinstance(features, dict) else ""
                row.append(str(val) if val else "")
            rows.append(row)
        # 添加表号标题
        add_paragraph_with_format(
            self.doc, "表 1 候选方法对比表",
            font_name=FONT_HEI, font_size=SIZE_WU, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_before=6, space_after=3)
        self.add_table(headers, rows)
        # 选型结论
        conclusion = mc.get("conclusion", "")
        if conclusion:
            self.add_h2("（选型结论）")
            self.add_para(conclusion)
        # 校验
        if len(methods) < 2:
            self.warnings.append(
                f"方法对比仅 {len(methods)} 方法，v2.1 要求 ≥2 方法")
        if len(dimensions) < 3:
            self.warnings.append(
                f"方法对比仅 {len(dimensions)} 维度，v2.1 要求 ≥3 维度")

    # v2.1 新增：八点六、数学公式（≥1，可选）

    def _add_formulas(self):
        """v2.1 新增：数学公式（≥1 个，含编号+变量解释，可选）"""
        formulas = self._get("formulas", default=[])
        if not isinstance(formulas, list) or not formulas:
            # 公式为可选字段，缺省不警告
            return
        self.add_h1("八点六、核心方法数学公式（v2.1 强化）")
        for f in formulas:
            if not isinstance(f, dict):
                continue
            formula = f.get("formula", "")
            number = f.get("number", 0)
            variables = f.get("variables", [])
            if not formula:
                continue
            # 公式段（居中，编号右对齐）
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            pf.first_line_indent = Pt(0)
            run = p.add_run(f"{formula}        ({number})")
            set_run_font(run, font_name=FONT_TIMES, font_size=SIZE_XIAO_SI)
            # 变量解释段
            if isinstance(variables, list) and variables:
                var_parts = []
                for v in variables:
                    if not isinstance(v, dict):
                        continue
                    name = v.get("name", "")
                    meaning = v.get("meaning", "")
                    if name and meaning:
                        var_parts.append(f"{name} 为{meaning}")
                if var_parts:
                    var_text = "其中 " + "，".join(var_parts) + "。"
                    self.add_para(var_text)

    # v2.1 新增：八点七、技术路线图（1 张）

    def _add_tech_roadmap(self):
        """v2.1 新增：技术路线图（1 张图，含图号+标题+图后说明）"""
        roadmap = self._get("tech_roadmap", default={})
        if not isinstance(roadmap, dict) or not roadmap:
            return
        self.add_h1("八点七、技术路线图（v2.1 强化）")
        image_path = roadmap.get("image_path", "")
        caption = roadmap.get("caption", "图 1 课题技术路线图")
        description = roadmap.get("description", "")
        # 插入图片
        if image_path and os.path.exists(image_path):
            try:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(image_path, width=Cm(15))
            except Exception as e:
                self.add_para(f"（技术路线图插入失败：{e}）")
        else:
            # 无图时渲染文字版流程图（ASCII）
            self.add_para(
                "（如无图片，请提供 image_path 或附文字版流程图如下：）", indent=False)
            ascii_flow = (
                "┌─ 阶段 1：准备 ─┐    ┌─ 阶段 2：试点 ─┐    "
                "┌─ 阶段 3：分析 ─┐    ┌─ 阶段 4：跟踪 ─┐"
            )
            add_paragraph_with_format(
                self.doc, ascii_flow, font_name=FONT_TIMES, font_size=SIZE_WU,
                alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False)
        # 图注
        add_paragraph_with_format(
            self.doc, caption,
            font_name=FONT_HEI, font_size=SIZE_WU, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_before=3, space_after=6)
        # 图后说明
        if description:
            self.add_para(description)

    # 九、创新之处

    def _add_innovation(self):
        """九、创新之处（院级 200~400 字，至少 1 个，对比式写法）"""
        self.add_h1("九、创新之处")
        innovations = self._get("innovations", default=[])
        if isinstance(innovations, str):
            innovations = [innovations]
        if innovations:
            for i, inv in enumerate(innovations, 1):
                self.add_para(f"创新点 {i}：{inv}")
        else:
            self.add_para("（请填写创新点，至少 1 个，建议 2 个，每个 100~200 字。结构：[类型]。传统方法 [描述]，本课题 [方法]，[量化优势]。禁用『首次』『先进』『智能』等无量化支撑词。）")

    # 十、研究基础（前期成果、设备条件）

    def _add_research_basis(self):
        """十、研究基础（3 子节：团队/指导教师/实验条件）"""
        self.add_h1("十、研究基础")
        self.add_h2("（一）团队基础")
        team = self._get("team_foundation", default="")
        self.add_para(team if team else
                      "（请填写团队基础：成员相关课程、已有项目经验、技能匹配度。）")
        self.add_h2("（二）指导教师基础")
        advisor = self._get("advisor_foundation", default="")
        self.add_para(advisor if advisor else
                      "（请填写指导教师基础：主持项目、发表论文、研究方向匹配度。）")
        self.add_h2("（三）实验条件")
        lab = self._get("lab_condition", default="")
        self.add_para(lab if lab else
                      "（请填写实验条件：实验室设备、软件平台、合作单位支持。）")

    # 十一、预期研究成果

    def _add_expected_results(self):
        """十一、预期研究成果（院级：1 篇论文或 1 份调研报告，可量化）"""
        self.add_h1("十一、预期研究成果")
        outcomes = self._get("expected_outcomes", default=[])
        if isinstance(outcomes, str):
            outcomes = [outcomes]
        if outcomes:
            for o in outcomes:
                self.add_para(f"• {o}", indent=False)
        else:
            self.add_para("（请填写预期成果，每项含数量+级别+平台。院级至少 1 篇论文或 1 份调研报告，如：省级期刊论文 1 篇（拟投《校园心理》）、调研报告 1 份（约 1 万字）。）", indent=False)

    # v2.1 新增：十一点五、社会效益量化（≥3 项指标）

    def _add_social_benefits(self):
        """v2.1 新增：社会效益量化（≥3 项指标，每项 4 列：指标/基准/预期/提升幅度）"""
        self.add_h1("十一点五、社会效益量化（v2.1 强化）")
        benefits = self._get("social_benefits", default=[])
        if not isinstance(benefits, list) or not benefits:
            self.add_para(
                "（v2.1 要求 ≥3 项社会效益指标，每项含 4 列：指标名/传统基准/系统预期/提升幅度。"
                "提升幅度必须双重表达（百分比+倍数）。基准值必须有来源。详例见 SKILL.md §12.2。）")
            return
        # 表号标题
        add_paragraph_with_format(
            self.doc, "表 2 本课题社会效益量化评估表",
            font_name=FONT_HEI, font_size=SIZE_WU, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_before=6, space_after=3)
        rows = []
        for b in benefits:
            if not isinstance(b, dict):
                continue
            rows.append([
                b.get("indicator", ""),
                b.get("baseline", ""),
                b.get("expected", ""),
                b.get("improvement", ""),
            ])
        self.add_table(
            ["评估指标", "传统模式基准值", "系统模式预期值", "效益提升幅度"],
            rows, col_widths=[4.0, 4.0, 4.0, 4.0])
        if len(benefits) < 3:
            self.warnings.append(
                f"社会效益仅 {len(benefits)} 项，v2.1 要求 ≥3 项")

    # 十二、研究进度安排（v2.1 强化：4 阶段 6 个月）

    def _add_schedule(self):
        """十二、研究进度安排（v2.1 强化：4 列表格 + 4 阶段 6 个月，按月划分）"""
        self.add_h1("十二、研究进度安排")
        # v2.1 优先使用 project_schedule 字段（含 tasks 列表），否则回退到 schedule
        project_schedule = self._get("project_schedule", default=[])
        if isinstance(project_schedule, list) and project_schedule:
            add_paragraph_with_format(
                self.doc, "表 3 院级科研立项进度安排表（v2.1：4 阶段 6 个月）",
                font_name=FONT_HEI, font_size=SIZE_WU, bold=False,
                alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
                space_before=6, space_after=3)
            rows = []
            for s in project_schedule:
                if not isinstance(s, dict):
                    continue
                tasks = s.get("tasks", [])
                if isinstance(tasks, list):
                    work_str = " ".join(f"{chr(0x2460 + i) if i < 20 else i + 1}. {t}"
                                       for i, t in enumerate(tasks))
                else:
                    work_str = str(tasks)
                rows.append([
                    s.get("phase", ""),
                    s.get("time", ""),
                    work_str,
                    s.get("output", ""),
                ])
            self.add_table(["阶段", "时间", "主要工作", "阶段成果"], rows,
                           col_widths=[2.0, 3.0, 7.0, 4.0])
            # 校验阶段数
            if len(project_schedule) < 4:
                self.warnings.append(
                    f"进度安排仅 {len(project_schedule)} 阶段，v2.1 要求 ≥4 阶段")
            return
        # 回退到旧版 schedule 字段
        schedule = self._get("schedule", default=[])
        if schedule:
            add_paragraph_with_format(
                self.doc, "表 3 院级科研立项进度安排表（4 阶段 6 个月）",
                font_name=FONT_HEI, font_size=SIZE_WU, bold=False,
                alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
                space_before=6, space_after=3)
            rows = []
            for s in schedule:
                if not isinstance(s, dict):
                    continue
                rows.append([s.get("phase", ""), s.get("time", ""),
                             s.get("work", ""), s.get("output", "")])
            self.add_table(["阶段", "时间", "主要工作", "阶段成果"], rows,
                           col_widths=[2.5, 3.0, 6.5, 4.0])
            if len(schedule) < 4:
                self.warnings.append(
                    f"进度安排仅 {len(schedule)} 阶段，v2.1 要求 ≥4 阶段")
        else:
            self.add_para("（请填写进度安排，按月划分，每阶段标注交付物。v2.1 要求 ≥4 阶段，院级周期 3~6 个月，建议 4 行，留 0.5~1 月弹性时间。）")

    # 十三、经费预算

    def _add_budget(self):
        """十三、经费预算（6 类科目：资料/调研/材料/会议/印刷/其他，院级 500~2000 元）"""
        self.add_h1("十三、经费预算")
        items = self._get("budget_items", default=[])
        if items:
            rows = []
            total = 0
            for b in items:
                if not isinstance(b, dict):
                    continue
                try:
                    amount_num = int(b.get("amount", 0))
                except ValueError:
                    amount_num = 0
                total += amount_num
                rows.append([b.get("item", ""), f"{amount_num} 元",
                             b.get("basis", "")])
            rows.append(["合计", f"{total} 元", ""])
            self.add_table(["预算科目", "金额", "计算依据"], rows,
                           col_widths=[3.5, 3.0, 9.5])
        else:
            self.add_para("（请填写经费预算，6 类标准科目：资料费/调研费/材料费/会议费/印刷费/其他。金额非整数，附计算依据。院级经费 500~2000 元。示例：资料费 300/调研费 500/印刷费 200/其他 300，合计 1300 元。）")

    # 十四、参考文献（GB/T 7714）

    def _add_references(self):
        """十四、参考文献（GB/T 7714 五类格式：J/C/M/D/EB，院级 ≥5 条覆盖 ≥3 类含 ≥3 篇英文）"""
        self.add_h1("十四、参考文献")
        refs = self._get("references", default=[])
        if not isinstance(refs, list):
            refs = []
        if not refs:
            self.add_para("（请填写参考文献，v2.1 要求 ≥10 条覆盖 ≥3 类含 ≥3 篇英文。格式按 GB/T 7714-2015，详见 SKILL.md 第十四章。）", indent=False)
            return
        type_counter: Dict[str, int] = {}
        en_count = 0
        for i, ref in enumerate(refs, 1):
            if not isinstance(ref, dict):
                continue
            try:
                line = format_reference(i, ref)
            except Exception as e:
                line = f"（格式化失败：{e}）"
            add_reference_paragraph(self.doc, i, line)
            rt = ref.get("ref_type", "journal").lower()
            type_counter[rt] = type_counter.get(rt, 0) + 1
            # 统计英文文献
            if ref.get("is_english", False):
                en_count += 1
            elif any(c.isalpha() and ord(c) < 128
                     for c in str(ref.get("authors", ""))[:20]):
                en_count += 1
        # v2.1 校验
        if len(refs) < 5:
            self.warnings.append(f"参考文献仅 {len(refs)} 条，院级要求 ≥5 条")
        if len(type_counter) < 3:
            self.warnings.append(
                f"参考文献仅覆盖 {len(type_counter)} 类，院级要求至少 3 类")
        if en_count < 3:
            self.warnings.append(
                f"英文文献仅 {en_count} 篇，v2.1 要求 ≥3 篇")

    # 十五、指导教师推荐意见 / 十六、学院评审意见

    def _add_review_section(self):
        """十五/十六、指导教师推荐意见 + 学院评审意见（双栏签字）"""
        self.add_h1("十五、指导教师推荐意见")
        for _ in range(6):
            self.doc.add_paragraph()
        self.add_para("指导教师签字：____________________    日期：______年____月____日", indent=False)
        self.add_h1("十六、学院评审意见")
        for _ in range(6):
            self.doc.add_paragraph()
        self.add_para("学院盖章：____________________    日期：______年____月____日", indent=False)

    # 主构建方法

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 16+9 章节，生成 docx（v2.1 含 9 项强化章节）"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_cover()
            self._add_leader_members()
            self._add_basic_info()
            self._add_background()
            self._add_policy_citations()           # v2.1 新增
            self._add_research_status()
            self._add_literature_review()          # v2.1 新增
            self._add_research_content()
            self._add_scientific_challenges()      # v2.1 新增
            self._add_research_scheme()
            self._add_method_comparison()          # v2.1 新增
            self._add_formulas()                   # v2.1 新增
            self._add_tech_roadmap()               # v2.1 新增
            self._add_innovation()
            self._add_research_basis()
            self._add_expected_results()
            self._add_social_benefits()            # v2.1 新增
            self._add_schedule()
            self._add_budget()
            self._add_references()
            self._add_review_section()
            if self.warnings:
                print("⚠️ 数据校验警告：", file=sys.stderr)
                for w in self.warnings:
                    print(f"  - {w}", file=sys.stderr)
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 院级科研立项申请书（v2.1）已生成：{output_path}")
        return str(output_path)

    # 数据校验

    def _validate_data(self) -> List[str]:
        """校验数据完整性，记录警告但不阻断生成（含院级 + v2.1 专属校验）"""
        p0_fields = [("project_name", "课题名称"), ("leader_name", "负责人姓名"),
                     ("college", "所在学院"), ("advisor_name", "指导教师姓名")]
        for key, name in p0_fields:
            if not self._get(key):
                self.warnings.append(f"缺少 {name}（{key}）")
        basic = self._get("basic_info", default={})
        if not isinstance(basic, dict):
            basic = {}
        if not basic.get("duration") and not self._get("duration"):
            self.warnings.append("缺少 起止时间（duration）")
        if not basic.get("research_type") and not self._get("research_type"):
            self.warnings.append("缺少 研究类型（research_type），默认『应用研究』")
        if not self._get("background"):
            self.warnings.append("缺少 背景与意义（background），将使用占位文本")
        if not self._get("research_status"):
            self.warnings.append("缺少 国内外研究现状（research_status），将使用占位文本")
        if not self._get("innovations"):
            self.warnings.append("缺少 创新之处（innovations），将使用占位文本")

        # 经费校验
        items = self._get("budget_items", default=[])
        if items:
            total = 0
            for b in items:
                if isinstance(b, dict):
                    try:
                        total += int(b.get("amount", 0))
                    except (ValueError, TypeError):
                        pass
            budget_total_str = str(self._get("budget_total", default="")).strip()
            try:
                budget_total_num = int(budget_total_str)
            except ValueError:
                budget_total_num = -1
            if budget_total_num >= 0 and total != budget_total_num:
                self.warnings.append(f"预算合计 {total} 元 与申请经费 {budget_total_num} 元不一致")
            if total > 2000:
                self.warnings.append(f"预算合计 {total} 元 超院级上限 2000 元，需调整或说明必要性")
            if total < 500:
                self.warnings.append(f"预算合计 {total} 元 低于院级下限 500 元，请补充")

        # 团队人数校验
        members = self._get("team_members", default=[])
        if isinstance(members, list) and len(members) > 2:
            self.warnings.append(f"团队成员 {len(members)} 人 超院级上限 2 人，需精简")

        # 周期校验（v2.1 调整为 3~6 个月）
        duration_str = str(basic.get("duration") or self._get("duration") or "")
        if "个月" in duration_str:
            try:
                match = re.search(r'(\d+)\s*个月', duration_str)
                months = int(match.group(1)) if match else 0
                if months > 6:
                    self.warnings.append(f"周期 {months} 个月 超院级 v2.1 上限 6 个月，需调整")
                if 0 < months < 3:
                    self.warnings.append(f"周期 {months} 个月 低于院级下限 3 个月，需调整")
            except (ValueError, IndexError):
                pass

        # v2.1 强化字段校验
        policies = self._get("policy_citations", default=[])
        if isinstance(policies, list) and len(policies) < 3:
            self.warnings.append(
                f"国家政策引用仅 {len(policies)} 条，v2.1 要求 ≥3 条")

        challenges = self._get("scientific_challenges", default=[])
        if isinstance(challenges, list) and len(challenges) < 1:
            self.warnings.append(
                f"科学挑战仅 {len(challenges)} 个，v2.1 要求 ≥1 个")

        review = self._get("literature_review", default={})
        if isinstance(review, dict) and review:
            total_items = (len(review.get("significance", [])) +
                          len(review.get("domestic", [])) +
                          len(review.get("foreign", [])))
            if total_items < 10:
                self.warnings.append(
                    f"文献综述仅约 {total_items} 段，v2.1 要求 ≥10 条引用")

        mc = self._get("method_comparison", default={})
        if isinstance(mc, dict) and mc:
            methods = mc.get("methods", [])
            if isinstance(methods, list) and len(methods) < 2:
                self.warnings.append(
                    f"方法对比仅 {len(methods)} 方法，v2.1 要求 ≥2 方法")

        benefits = self._get("social_benefits", default=[])
        if isinstance(benefits, list) and len(benefits) < 3:
            self.warnings.append(
                f"社会效益仅 {len(benefits)} 项，v2.1 要求 ≥3 项")

        # project_schedule 与 schedule 至少有一个 ≥4 阶段
        ps = self._get("project_schedule", default=[])
        sc = self._get("schedule", default=[])
        ps_count = len(ps) if isinstance(ps, list) else 0
        sc_count = len(sc) if isinstance(sc, list) else 0
        if max(ps_count, sc_count) < 4:
            self.warnings.append(
                f"进度安排仅 {max(ps_count, sc_count)} 阶段，v2.1 要求 ≥4 阶段")

        # 参考文献 + 英文文献校验
        refs = self._get("references", default=[])
        if isinstance(refs, list):
            if len(refs) < 5:
                self.warnings.append(
                    f"参考文献仅 {len(refs)} 条，院级要求 ≥5 条")
            en_count = 0
            for r in refs:
                if not isinstance(r, dict):
                    continue
                if r.get("is_english", False):
                    en_count += 1
                elif any(c.isalpha() and ord(c) < 128
                         for c in str(r.get("authors", ""))[:20]):
                    en_count += 1
            if en_count < 3:
                self.warnings.append(
                    f"英文文献仅 {en_count} 篇，v2.1 要求 ≥3 篇")

        return self.warnings


# 默认示例数据（本院学生学习拖延行为干预研究，v2.1 含全部强化字段）

DEFAULT_DATA = {
    "project_name": "本院学生学习拖延行为干预研究",
    "leader_name": "王晓",
    "leader_id": "2023123456",
    "leader_gender": "女",
    "leader_major": "应用心理学 2023 级",
    "leader_phone": "138XXXXXXXX",
    "leader_email": "wangxiao@xxx.edu.cn",
    "advisor_name": "刘敏副教授",
    "college": "教育科学学院",
    "apply_date": "2025 年 3 月 10 日",
    "basic_info": {
        "research_type": "应用研究",
        "discipline": "0402 心理学",
        "duration": "2025 年 4 月 — 2025 年 9 月（共 6 个月）",
        "budget": "1300",
    },
    "team_members": [
        {"name": "陈静", "id": "2023123457", "major": "应用心理学 2023 级",
         "role": "问卷设计+数据采集", "phone": "139XXXXXXXX"},
    ],
    "background": [
        "时代背景：2024 年中国心理学会《大学生学习行为蓝皮书》指出，全国高校学生学习拖延行为发生率达 67.8%，其中\"严重拖延\"占 23%。学习拖延已成为影响大学生学业表现的重要心理因素。",
        "现实痛点：调研本院应用心理学专业 3 个年级 120 名学生发现，学习拖延发生率为 71.3%，高于全国均值。其中\"考前突击型\"占 45%，\"日常回避型\"占 26%。期末挂科学生中 82% 存在严重拖延。学生方面，68% 表示\"想改变但找不到方法\"，仅 12% 接受过系统干预。",
        "研究缺口：已有研究多停留在\"现象描述\"层面，缺乏针对本学院学生特点的本土化干预方案。少数干预研究样本量小（<50）、周期短（<2 周），缺乏效果评估。",
        "课题意义：理论上探索拖延行为干预理论在本院学生中的适用性；实践上形成可复用的干预方案，预期能将本院学生学习拖延发生率降低 20 个百分点以上；社会上助力学风建设。",
    ],
    # v2.1 新增：国家政策引用（3 条）
    "policy_citations": [
        {
            "agency": "中国心理学会",
            "title": "大学生学习行为蓝皮书",
            "doc_no": "",
            "date": "2024-09-20",
            "excerpt": "全国高校学生学习拖延行为发生率达 67.8%，建议高校建立学生心理行为早期识别与干预机制",
            "url": "http://www.cps.org.cn/xxx",
            "access_date": "2025-02-15"
        },
        {
            "agency": "教育部思想政治工作司",
            "title": "关于全面加强和改进新时代学生心理健康工作专项行动计划（2023-2025 年）的通知",
            "doc_no": "教体艺〔2023〕1 号",
            "date": "2023-04-20",
            "excerpt": "普通高等学校要常态化开展学生心理健康监测，建立心理健康状况档案",
            "url": "",
            "access_date": ""
        },
        {
            "agency": "教育部、中央宣传部等八部门",
            "title": "关于加快构建高校思想政治工作体系的意见",
            "doc_no": "教思政〔2022〕1 号",
            "date": "2022-04-15",
            "excerpt": "加强学风建设，做好学业辅导与心理疏导衔接",
            "url": "",
            "access_date": ""
        },
    ],
    "research_status": [
        "国内研究现状：张丽等（2022）对 3 所高校 600 名大学生调查发现学习拖延发生率 65.4% [1]，但仅描述现象未提出干预方案。王强（2023）设计了\"番茄工作法+同伴监督\"干预机制，样本 40 人、周期 2 周，短期效果显著（拖延率降 18%）但未做长期跟踪 [2]。",
        "国外研究现状：Steel (2020) 综述了近 20 年拖延研究，提出\"动机-时间感知-冲动性\"三因素模型 [3]。Pychyl & Sirois (2021) 验证了基于正念的干预对拖延行为的改善效果，6 周后拖延率降 24%，6 个月后保持 65% [4]。",
        "综合评述：已有研究存在两个共性问题：① 样本量普遍偏小（多数 <100 人），外推性不足；② 干预周期短（<2 周），缺乏长期效果评估。本课题的关键差异：（1）样本量 200+，覆盖本院 3 个年级；（2）周期 6 个月含 2 个月跟踪；（3）融合\"正念+同伴监督+学业规划\"三维干预，本土化适配。",
    ],
    # v2.1 新增：国内外研究现状综述详尽版（10 篇含 3 篇英文）
    "literature_review": {
        "significance": [
            "学习拖延（academic procrastination）是大学生最常见的自我调节失效行为之一，影响学业表现、心理健康与职业发展 [1]。本课题聚焦本院学生学习拖延行为，从理论意义与实践意义两方面展开综述。",
            "理论意义：本课题验证拖延行为三因素模型（动机-时间感知-冲动性）[3] 在中国大学生群体中的适用性，丰富拖延行为的本土化理论。实践意义：本课题开发的三维干预方案（正念+同伴监督+学业规划）为高校学风建设提供可复用工具。",
        ],
        "domestic": [
            "张丽等（2022）对 3 所高校 600 名大学生调查发现学习拖延发生率 65.4% [1]，但仅描述现象未提出干预方案。",
            "王强（2023）设计了\"番茄工作法+同伴监督\"干预机制，样本 40 人、周期 2 周，短期效果显著（拖延率降 18%）但未做长期跟踪 [2]。",
            "陈伟（2023）综述了近 5 年国内干预研究，指出 75% 的研究周期短于 2 周 [5]。",
            "赵敏（2023）在学位论文中提出了\"动机-时间-情绪\"三维拖延模型，但缺乏实证验证 [7]。",
            "刘洋等（2024）采用正念干预大学生拖延行为，4 周后拖延率降 21%，但未设对照组 [6]。",
        ],
        "foreign": [
            "Steel (2020) 综述了近 20 年拖延研究，提出\"动机-时间感知-冲动性\"三因素模型 [3]。",
            "Pychyl & Sirois (2021) 验证了基于正念的干预对拖延行为的改善效果，6 周后拖延率降 24%，6 个月后保持 65% [4]。",
            "Geng et al. (2022) 在 Computers & Education 发表的研究表明，基于移动应用的同伴监督干预对拖延行为有显著效果（效应量 d=0.42）[8]。",
        ],
        "summary": "已有研究存在三个共性问题：① 样本量普遍偏小（国内 75% <100 人），外推性不足；② 干预周期短（<2 周），缺乏长期效果评估；③ 缺乏针对本院学生特点（如年级差异、专业压力）的本土化设计。本课题的关键差异：（1）样本量 200+，覆盖本院 3 个年级；（2）周期 6 个月含 2 个月跟踪；（3）融合\"正念+同伴监督+学业规划\"三维干预，本土化适配。",
    },
    "research_goal": "总目标：开发能将本院学生学习拖延发生率降低 ≥20 个百分点、2 个月后效果保持率 ≥60% 的本土化干预方案。阶段目标 1：完成方案设计（2025.05 前）；阶段目标 2：完成 4 周试点实施（2025.07 前）；阶段目标 3：完成 2 个月跟踪评估（2025.09 前）。",
    "research_content": [
        "干预方案设计：基于拖延行为三因素模型设计\"正念+同伴监督+学业规划\"三维方案，含 8 项具体干预动作，产出方案文档 1 份。",
        "试点实施：在本院 3 个年级 6 个班共 200 名学生中开展 4 周试点，含干预组与对照组，产出实施记录与原始数据。",
        "效果评估：采用前后测+跟踪测设计，分析干预对学习拖延发生率、学业成绩、自我效能感的影响，产出评估报告 1 份。",
    ],
    "key_problems": [
        "三维干预中各维度的协同机制设计（避免维度间相互抵消）",
        "准实验设计中混杂变量（学生自控力、年级差异）的控制",
    ],
    # v2.1 新增：科学挑战（1 个）
    "scientific_challenges": [
        {
            "title": "三维干预中各维度的协同机制设计",
            "description": "本课题融合\"正念+同伴监督+学业规划\"三维干预，三维之间存在潜在的协同或抵消效应。如何设计各维度的干预顺序、剂量与衔接，避免维度间相互抵消，是本课题的关键科学挑战。",
            "sub_challenges": [
                "子挑战 1：维度间剂量配比的优化。已有研究表明正念干预需 ≥6 周方显效，而同伴监督短期即有效（<2 周）。如何根据本院学生学期节奏（16 周教学周期）合理配比三维干预的剂量与频次，避免学生在前期因多维度同时介入产生疲劳，是本子挑战的核心问题。"
            ],
            "references": [3, 4]
        }
    ],
    "tech_route": "总体技术路线分 4 阶段（见图 1）：① 准备阶段（2025.04），完成文献综述与方案设计，产出综述与方案；② 试点实施（2025.05-07），在本院 3 年级 6 班 200 人中开展 4 周干预，产出实施记录与原始数据；③ 数据分析（2025.08-09），SPSS 27 进行统计分析，产出分析报告；④ 跟踪研究（2025.09），2 个月后跟踪测，产出跟踪报告。",
    "methods": [
        "文献研究法：系统梳理国内外 30+ 篇相关文献，建立理论框架。",
        "准实验法：采用\"干预组-对照组\"前后测设计，控制混杂变量。",
        "问卷调查法：采用 Tuckman 拖延量表（含 16 题项），Cronbach α ≥ 0.80。",
        "深度访谈法：干预后对 15 名学生进行半结构化访谈。",
    ],
    "data_source": "数据来源：本院应用心理学专业 3 个年级 6 个班共 200 名学生（干预组 100 + 对照组 100）。实验设备：教育科学学院心理实验室生理多导仪 BioTrace Pro、行为观察记录系统 Noldus Observer XT 14。软件工具：SPSS 27、Amos 24、NVivo 14。",
    # v2.1 新增：方法对比表（2 方法 × 3 维度）
    "method_comparison": {
        "table_headers": ["维度", "单一正念干预", "三维融合干预（本课题）"],
        "methods": [
            {
                "name": "单一正念干预",
                "features": {
                    "技术特点": "仅采用正念冥想训练（每周 2 次 × 30 分钟）",
                    "优势": "实施简单、成本低、文献基础扎实 [4]",
                    "劣势": "见效慢（≥6 周）、单一维度效果有限（降 12~15%）"
                },
                "reference": 4
            },
            {
                "name": "三维融合干预（本课题）",
                "features": {
                    "技术特点": "正念+同伴监督+学业规划三维并行（每周 2+1+1 次）",
                    "优势": "多维协同、覆盖拖延行为多因素 [3]、效果更全面",
                    "劣势": "实施复杂、需多维度协调、对组织者要求高"
                },
                "reference": 3
            }
        ],
        "conclusion": "基于上表，单一正念干预实施简单但效果有限，三维融合干预虽实施复杂但覆盖拖延行为的多因素模型。结合本课题\"6 个月周期+200 人样本+指导教师行为干预方向\"的实际条件，本课题选用三维融合干预，并通过\"分阶段剂量递增\"策略降低实施复杂度。"
    },
    # v2.1 新增：数学公式（1 个，Cohen's d 效应量）
    "formulas": [
        {
            "formula": "d = (μ₁ - μ₂) / s_p",
            "number": 1,
            "variables": [
                {"name": "d", "meaning": "Cohen's d 效应量"},
                {"name": "μ₁", "meaning": "干预组后测均值"},
                {"name": "μ₂", "meaning": "对照组后测均值"},
                {"name": "s_p", "meaning": "合并标准差，s_p = √((s₁² + s₂²) / 2)，s₁ 和 s₂ 分别为干预组与对照组的标准差"}
            ]
        }
    ],
    # v2.1 新增：技术路线图（1 张，文字描述版）
    "tech_roadmap": {
        "image_path": "",
        "caption": "图 1 课题技术路线图",
        "description": "本课题技术路线分 4 阶段（详见图 1）。准备阶段（2025.04）完成文献综述与方案设计；试点阶段（2025.05-07）在本院 3 年级 6 班 200 人中开展 4 周干预；分析阶段（2025.08-09）进行 SPSS 统计分析与 Amos 中介效应分析；跟踪阶段（2025.09）开展 2 个月跟踪测与深度访谈。各阶段交付物明确，前后衔接紧密。"
    },
    "innovations": [
        "方法创新。已有干预研究多采用单一手段（仅正念/仅同伴监督），效果有限（短期降 12~18%）。本课题融合\"正念+同伴监督+学业规划\"三维干预，预实验显示三维干预短期降 25%，优于单维方案的 1.4 倍。",
        "数据创新。已有研究样本量普遍 <100、周期 <2 周。本课题样本量 200+、周期 6 个月含 2 个月跟踪，是已有研究样本量的 2 倍以上，外推性显著提升。",
    ],
    "team_foundation": "团队 2 名成员已修读《学习心理学》《心理测量学》《教育统计学》等核心课程，1 人有院级科研立项参与经验。负责人王晓掌握 SPSS、Amos，曾参与院级学生心理调查项目。",
    "advisor_foundation": "指导教师刘敏副教授主持省教育厅课题 1 项，近 3 年发表 CSSCI 论文 3 篇，研究方向为学习心理学与行为干预，与本项目高度契合。",
    "lab_condition": "教育科学学院心理实验室配备生理多导仪 BioTrace Pro、行为观察记录系统 Noldus Observer XT 14，可满足本项目行为观察与数据分析需求。已与本院 3 个年级辅导员签署调研合作协议。",
    "expected_outcomes": [
        "省级期刊论文 1 篇（拟投《校园心理》）",
        "调研报告 1 份（约 1 万字）",
        "干预方案手册 1 套（含 8 项干预动作操作指南）",
    ],
    # v2.1 新增：社会效益量化（3 项指标）
    "social_benefits": [
        {
            "indicator": "本院学生学习拖延发生率",
            "baseline": "71.3%（基线调研，本院 3 年级 120 人）",
            "expected": "≤51.3%",
            "improvement": "降低 ≥20 个百分点"
        },
        {
            "indicator": "期末挂科率（严重拖延学生）",
            "baseline": "32%（基线）",
            "expected": "≤20%",
            "improvement": "降低 ≥12 个百分点"
        },
        {
            "indicator": "学生自我效能感得分（SES 量表）",
            "baseline": "25.4 分（基线均值，满分 40）",
            "expected": "≥29.0 分",
            "improvement": "提升 ≥3.6 分（+14.2%）"
        }
    ],
    # v2.1 新增：进度安排（4 阶段 6 个月，使用 project_schedule 字段）
    "project_schedule": [
        {
            "phase": "准备",
            "time": "2025.04",
            "tasks": ["文献综述 30 篇", "三维干预方案设计", "量表选型与预测试"],
            "output": "综述 1 份 + 方案 1 份"
        },
        {
            "phase": "试点",
            "time": "2025.05-07",
            "tasks": ["招募 200 人（干预 100+对照 100）", "4 周干预实施", "实施过程记录与质控"],
            "output": "实施记录 + 原始数据"
        },
        {
            "phase": "分析",
            "time": "2025.08-09",
            "tasks": ["SPSS 27 统计分析", "Amos 24 中介效应分析", "NVivo 14 访谈编码"],
            "output": "分析报告 1 份"
        },
        {
            "phase": "跟踪",
            "time": "2025.09",
            "tasks": ["2 个月跟踪测", "深度访谈 15 人", "撰写调研报告"],
            "output": "跟踪报告 + 调研报告 1 份"
        }
    ],
    # 保留旧版 schedule 字段（向后兼容）
    "schedule": [
        {"phase": "准备", "time": "2025.04", "work": "文献综述 30 篇、方案设计", "output": "综述 1 份 + 方案 1 份"},
        {"phase": "试点", "time": "2025.05-07", "work": "3 年级 6 班 200 人 4 周干预", "output": "实施记录 + 原始数据"},
        {"phase": "分析", "time": "2025.08-09", "work": "SPSS 统计分析", "output": "分析报告 1 份"},
        {"phase": "跟踪", "time": "2025.09", "work": "2 个月后跟踪测", "output": "跟踪报告 1 份"},
    ],
    "budget_items": [
        {"item": "资料费", "amount": "300", "basis": "图书 4 本 × 50 元 + 数据库订阅 100 元"},
        {"item": "调研费", "amount": "500", "basis": "1 次本院学生集中调研 × 500 元（含问卷、场地）"},
        {"item": "材料费", "amount": "0", "basis": "问卷电子化，无印刷材料"},
        {"item": "会议费", "amount": "0", "basis": "院级一般不资助会议费"},
        {"item": "印刷费", "amount": "200", "basis": "论文版面费 200 元"},
        {"item": "其他", "amount": "300", "basis": "礼品感谢 30 份 × 10 元"},
    ],
    "budget_total": "1300",
    # 参考文献（10 条覆盖 4 类，含 3 篇英文 [3][4][8]）
    "references": [
        {"ref_type": "journal", "authors": "张丽, 王伟", "title": "大学生学习拖延行为调查研究", "journal": "心理科学", "year": "2022", "volume": "45", "issue": "(3)", "pages": "612-619", "is_english": False},
        {"ref_type": "journal", "authors": "王强", "title": "同伴监督对大学生拖延行为的影响", "journal": "校园心理", "year": "2023", "volume": "21", "issue": "(2)", "pages": "145-150", "is_english": False},
        {"ref_type": "journal", "authors": "Steel P", "title": "The nature of procrastination: A meta-analytic review", "journal": "Psychological Bulletin", "year": "2020", "volume": "146", "issue": "(4)", "pages": "318-348", "is_english": True},
        {"ref_type": "journal", "authors": "Pychyl T, Sirois F", "title": "Procrastination and mindfulness", "journal": "Journal of Behavioral Medicine", "year": "2021", "volume": "44", "issue": "(2)", "pages": "234-245", "is_english": True},
        {"ref_type": "journal", "authors": "陈伟", "title": "近5年国内大学生拖延干预研究综述", "journal": "心理发展与教育", "year": "2023", "volume": "39", "issue": "(4)", "pages": "567-578", "is_english": False},
        {"ref_type": "journal", "authors": "刘洋, 李娜, 周明", "title": "正念训练对大学生拖延行为的干预效果", "journal": "中国临床心理学杂志", "year": "2024", "volume": "32", "issue": "(1)", "pages": "156-161", "is_english": False},
        {"ref_type": "thesis", "authors": "赵敏", "title": "大学生拖延行为的干预研究", "city": "北京", "school": "北京师范大学", "year": "2023", "is_english": False},
        {"ref_type": "journal", "authors": "Geng J, Jou M, Xu Y, et al.", "title": "Mobile app-based peer monitoring intervention for academic procrastination", "journal": "Computers & Education", "year": "2022", "volume": "178", "issue": "", "pages": "104401", "is_english": True},
        {"ref_type": "book", "authors": "Bandura A", "title": "Social learning theory", "city": "Englewood Cliffs", "publisher": "Prentice-Hall", "year": "1977", "is_english": True},
        {"ref_type": "web", "authors": "中国心理学会", "title": "大学生学习行为蓝皮书", "publish_date": "2024-09-20", "access_date": "2025-02-15", "url": "http://www.cps.org.cn/xxx", "is_government": True, "is_english": False},
    ],
}


# CLI 入口

def main():
    parser = argparse.ArgumentParser(
        description="大学生院级科研立项申请书 docx 生成器（v2.1 案例优化版）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：python build.py --data data.json --out output.docx\n"
            "      python build.py --demo --out demo.docx\n"
            "院级 v2.1 特征：经费 500~2000 元、周期 3~6 个月、团队 1~2 人、"
            "GB/T 7714 参考文献 ≥5 条覆盖 ≥3 类含 ≥3 篇英文。\n"
            "v2.1 新增 9 项强化章节：国家政策引用 / 科学挑战 / 文献综述 / "
            "方法对比表 / 技术路线图 / 数学公式 / 社会效益量化 / 进度安排 4 阶段 6 个月 / "
            "JSON Schema + build.py + 质检清单。\n"
            "JSON 字段详见 SKILL.md 第十七章。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档（v2.1 含 9 项强化章节）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
