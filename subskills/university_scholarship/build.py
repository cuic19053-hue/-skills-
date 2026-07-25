#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
校级奖学金申请书 docx 生成器

支持一等/二等/三等三个等级，按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 标题：黑体二号，居中
- 称呼：顶格，宋体小四，全角冒号
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 主干课程表：宋体五号，居中
- 学期 GPA 表（4 学期）：宋体五号，居中
- "此致"另起一行空两格，"敬礼！"另起一行顶格
- 落款：右对齐

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第三章信息采集清单。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进），宋体小四"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_section_heading(doc, text: str):
    """正文中的小节标题（一、二、三…）：黑体小四加粗，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
        line_spacing=1.5, space_before=6, space_after=3)


def add_cizhi_paragraph(doc, text: str = "此致"):
    """'此致'另起一行，空两格（首行缩进 2 字符）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格（不缩进）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=False, line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None,
                        caption: str = ""):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）

    Args:
        caption: 表格上方说明文字（如"主干课程成绩："），不写则不添加
    """
    if caption:
        add_paragraph_with_format(
            doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)

    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """校级奖学金申请书 docx 构建器（支持一等/二等/三等）"""

    # 等级→加权平均分门槛
    LEVEL_THRESHOLDS = {
        "一等": (85, 5),    # (加权, 排名前 X%)
        "二等": (80, 15),
        "三等": (75, 30),
    }

    # 等级→标题后缀
    LEVEL_TITLE_MAP = {
        "一等": "校级一等奖学金申请书",
        "二等": "校级二等奖学金申请书",
        "三等": "校级三等奖学金申请书",
    }

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []
        self.level: str = ""

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段，非列表返回空列表"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def _detect_level(self) -> str:
        """根据学业数据自动匹配等级（用户未指定 level 时使用）"""
        if self._get("level"):
            return str(self._get("level"))
        weighted = self._get("weighted_avg", default="0")
        rank_str = str(self._get("rank", default=""))
        rank_total = self._get("rank_total", default="0")
        try:
            w = float(weighted)
        except (ValueError, TypeError):
            w = 0
        try:
            r_num = int(rank_str.split("/")[0]) if "/" in rank_str else 0
            r_total = int(rank_total) if rank_total else (
                int(rank_str.split("/")[1]) if "/" in rank_str else 0)
        except (ValueError, IndexError):
            r_num, r_total = 0, 0
        if not r_total:
            return "二等"  # 默认二等
        pct = r_num / r_total * 100
        if w >= 85 and pct <= 5:
            return "一等"
        if w >= 80 and pct <= 15:
            return "二等"
        if w >= 75 and pct <= 30:
            return "三等"
        return "二等"

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_heading(self, text):
        return add_section_heading(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, caption=""):
        return add_table_from_data(self.doc, headers, rows, col_widths, caption)

    # --- 标题（含奖学金等级动态填充）---

    def _add_title(self):
        """标题：黑体二号居中，动态填充等级——"校级X等奖学金申请书"；
        等级取 data.level，若未指定则按学业数据自动匹配。"""
        self.level = self._detect_level()
        title = self.LEVEL_TITLE_MAP.get(self.level, "校级奖学金申请书")
        add_title(self.doc, title)

    # --- 称呼 ---

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号"""
        salutation = self._get("salutation",
                               default="尊敬的学院领导、评审委员会：")
        add_salutation_paragraph(self.doc, salutation)

    # --- 开头段落 ---

    def _add_opening(self):
        """开头段落（80~120 字）：身份 + 申报等级 + 核心数据 2 项 + 进入正文
        若用户提供 opening 字段则直接使用，否则按模板拼装。"""
        opening = self._get("opening", default="")
        if opening:
            self.add_para(opening)
            return
        name = self._get("name")
        college = self._get("college")
        major = self._get("major")
        grade = self._get("grade")
        apply_year = self._get("apply_year", default="2024-2025 学年")
        gpa = self._get("gpa")
        rank = self._get("rank")
        level_cn = self.level or "二等"
        parts = []
        if name and college and major and grade:
            parts.append(f"我是{college}{major}{grade}学生{name}，"
                         f"特申请{apply_year}校级{level_cn}奖学金。")
        else:
            parts.append(f"特申请{apply_year}校级{level_cn}奖学金。")
        data_parts = []
        if gpa:
            data_parts.append(f"GPA {gpa}")
        if rank:
            data_parts.append(f"专业排名第 {rank}")
        if data_parts:
            parts.append("本学年" + "，".join(data_parts) + "。")
        parts.append("现将本人情况汇报如下：")
        self.add_para("".join(parts))

    # --- 思想方面 ---

    def _add_ideology(self):
        """思想方面（120~180 字）：政治立场 + 入党/团情况 + 思想觉悟 + 1 件具体事"""
        self.add_heading("一、思想方面")
        ideology = self._get("ideology", default="")
        if ideology:
            if isinstance(ideology, list):
                for p in ideology:
                    self.add_para(p)
            else:
                self.add_para(ideology)
            return
        political = self._get("political_status", default="共青团员")
        party_history = self._get("party_history", default="")
        party_activities = self._get_list("party_activities")
        parts = []
        if political:
            if "党员" in political:
                parts.append(f"作为一名{political}，我始终坚持学习党的理论，"
                             "认真学习习近平新时代中国特色社会主义思想与"
                             "党的二十大和二十届三中全会精神。")
            elif "积极分子" in political:
                parts.append("作为一名入党积极分子，我认真学习习近平新时代中国特色社会主义思想，"
                             "时刻以党员标准要求自己。")
            else:
                parts.append(f"作为一名{political}，我拥护中国共产党的领导，"
                             "认真学习党的创新理论。")
        if party_history:
            parts.append(party_history)
        if party_activities:
            parts.append("；".join(party_activities))
        parts.append("在日常生活中，我注重理论学习与实践结合，关注时政热点，提升思想觉悟。")
        self.add_para("".join(parts))

    # --- 学习方面（含主干课程表 + 学期 GPA 表，重点）---

    def _add_academics(self):
        """学习方面（350~500 字，重点）：
        GPA + 排名 + 主干课程表 + 学期 GPA 表 + 英语计算机等级 + 学习方法
        """
        self.add_heading("二、学习方面")
        academics = self._get("academics", default="")
        if academics and isinstance(academics, str):
            self.add_para(academics)
            self._add_core_courses_table()
            self._add_semester_gpa_table()
            self._add_academics_summary()
            return
        gpa = self._get("gpa")
        weighted = self._get("weighted_avg")
        rank = self._get("rank")
        rank_total = self._get("rank_total")
        course_count = self._get("course_count")
        high_score_count = self._get("high_score_count")
        parts = []
        if gpa:
            gpa_str = f"本学年 GPA {gpa}"
            if weighted:
                gpa_str += f"，加权平均分 {weighted}"
            if rank:
                gpa_str += f"，专业排名第 {rank}"
                if rank_total:
                    try:
                        r_num = int(str(rank).split("/")[0])
                        total_num = int(rank_total)
                        if total_num > 0:
                            pct = round(r_num / total_num * 100, 1)
                            gpa_str += f"（前 {pct}%）"
                    except (ValueError, IndexError):
                        pass
            parts.append(gpa_str + "。")
        if course_count and high_score_count:
            parts.append(f"修读 {course_count} 门课程，{high_score_count} 门 80 分以上。")
        if parts:
            self.add_para("".join(parts))
        self._add_core_courses_table()
        self._add_semester_gpa_table()
        self._add_academics_summary()

    def _add_core_courses_table(self):
        """主干课程表：4~6 门，含课程名/学分/成绩"""
        courses = self._get_list("core_courses")
        if not courses:
            return
        rows = []
        for c in courses:
            if not isinstance(c, dict):
                continue
            rows.append([
                str(c.get("name", "")),
                str(c.get("credit", "")),
                str(c.get("score", "")),
            ])
        if rows:
            self.add_table(
                ["课程名称", "学分", "成绩"],
                rows,
                col_widths=[8.0, 3.0, 3.0],
                caption="主干课程成绩：",
            )

    def _add_semester_gpa_table(self):
        """学期 GPA 表（4 学期）：含学期/GPA/加权平均分，展示稳定或进步趋势"""
        semesters = self._get_list("semester_gpa")
        if not semesters:
            return
        rows = []
        gpa_values = []
        weighted_values = []
        for s in semesters:
            if not isinstance(s, dict):
                continue
            sem_name = str(s.get("semester", ""))
            sem_gpa = str(s.get("gpa", ""))
            sem_weighted = str(s.get("weighted", ""))
            rows.append([sem_name, sem_gpa, sem_weighted])
            try:
                gpa_values.append(float(sem_gpa.split("/")[0]))
                weighted_values.append(float(sem_weighted))
            except (ValueError, IndexError):
                pass
        if not rows:
            return
        self.add_table(
            ["学期", "GPA", "加权平均分"],
            rows,
            col_widths=[4.0, 5.0, 5.0],
            caption="大学前 4 学期 GPA 与加权平均分：",
        )
        # 趋势描述
        trend = self._describe_gpa_trend(gpa_values, weighted_values)
        if trend:
            self.add_para(trend)

    def _describe_gpa_trend(self, gpa_values: List[float],
                            weighted_values: List[float]) -> str:
        """根据 4 学期 GPA 数据生成趋势描述（一等拔尖/二等稳定/三等进步）"""
        if len(gpa_values) < 2:
            return ""
        level = self.level or "二等"
        gpa_min, gpa_max = min(gpa_values), max(gpa_values)
        gpa_range = gpa_max - gpa_min
        is_increasing = all(gpa_values[i] <= gpa_values[i + 1]
                            for i in range(len(gpa_values) - 1))
        if level == "一等":
            avg = sum(gpa_values) / len(gpa_values)
            return (f"4 学期 GPA 均稳定在 {gpa_min:.2f} 以上，"
                    f"平均 GPA {avg:.2f}，学业表现持续拔尖。")
        if level == "三等" and is_increasing:
            w_start, w_end = weighted_values[0], weighted_values[-1]
            return (f"4 学期 GPA 由 {gpa_values[0]:.2f} 提升至 {gpa_values[-1]:.2f}，"
                    f"加权平均分由 {w_start:.1f} 提升至 {w_end:.1f}，"
                    f"呈现稳步上升趋势。")
        if level == "二等" or not is_increasing:
            return (f"4 学期 GPA 在 {gpa_min:.2f}~{gpa_max:.2f} 区间稳定波动"
                    f"（波动幅度 {gpa_range:.2f}），学业表现持续稳定。")
        return ""

    def _add_academics_summary(self):
        """学习方面末尾：主干课程亮点 + 英语计算机等级 + 学习方法"""
        summary = self._get("academics_summary", default="")
        if summary:
            self.add_para(summary)
            return
        cet4 = self._get("cet4")
        cet6 = self._get("cet6")
        computer_level = self._get("computer_level")
        course_highlight = self._get("course_highlight", default="")
        study_method = self._get("study_method",
                                  default="学习上注重课前预习与课后总结，"
                                          "建立知识体系；遇到问题主动与老师、同学讨论。")
        parts = []
        if course_highlight:
            parts.append(course_highlight + "。")
        lang_parts = []
        if cet4:
            lang_parts.append(f"CET-4 {cet4} 分")
        if cet6:
            lang_parts.append(f"CET-6 {cet6} 分")
        if lang_parts:
            parts.append("、".join(lang_parts) + "；")
        if computer_level:
            parts.append(f"计算机{computer_level}。")
        if study_method:
            parts.append(study_method)
        if parts:
            self.add_para("".join(parts))

    # --- 科研与实践方面（可选）---

    def _add_research_practice(self):
        """科研与实践方面（150~250 字，可选）：竞赛/大创 + 学生工作 + 志愿服务"""
        rp_text = self._get("research_practice", default="")
        # 校级奖学金不强制要求本段；若用户无任何科研实践数据则跳过本段标题与内容
        has_competitions = bool(self._get_list("competitions"))
        has_projects = bool(self._get_list("research_projects"))
        has_position = bool(self._get("position"))
        has_volunteer = bool(self._get("volunteer_hours"))
        if not (rp_text or has_competitions or has_projects
                or has_position or has_volunteer):
            return
        self.add_heading("三、科研与实践方面")
        if rp_text and isinstance(rp_text, str):
            self.add_para(rp_text)
            return
        self._add_competition_items()
        self._add_research_items()
        self._add_social_practice()

    def _add_competition_items(self):
        """学科竞赛列表（校级即可）"""
        competitions = self._get_list("competitions")
        if not competitions:
            return
        comp_parts = []
        for c in competitions:
            if not isinstance(c, dict):
                continue
            time = c.get("time", "")
            name = c.get("name", "")
            award = c.get("award", "")
            role = c.get("role", "")
            seg = (f"{time} " if time else "") + name
            if award:
                seg += f" {award}"
            if role:
                seg += f"（{role}）"
            if seg:
                comp_parts.append(seg + "；")
        if comp_parts:
            comp_parts[-1] = comp_parts[-1].rstrip("；") + "。"
            self.add_para("学科竞赛方面，" + "".join(comp_parts))

    def _add_research_items(self):
        """大创/科研立项（校级即可）"""
        projects = self._get_list("research_projects")
        if not projects:
            return
        proj_parts = []
        for p in projects:
            if not isinstance(p, dict):
                continue
            name = p.get("name", "")
            level = p.get("level", "")
            role = p.get("role", "")
            duration = p.get("duration", "")
            output = p.get("output", "")
            seg = f"{role}《{name}》" if role and name else f"《{name}》" if name else ""
            if level:
                seg += f"（{level}）"
            if duration:
                seg += f"，{duration}"
            if output:
                seg += f"，{output}"
            if seg:
                proj_parts.append(seg + "。")
        if proj_parts:
            self.add_para("".join(proj_parts))

    def _add_social_practice(self):
        """学生工作 + 志愿服务"""
        position = self._get("position", default="")
        position_work = self._get("position_work", default="")
        volunteer_hours = self._get("volunteer_hours", default="")
        volunteer_detail = self._get("volunteer_detail", default="")
        parts = []
        if position:
            seg = f"担任{position}"
            if position_work:
                seg += f"，主要工作：{position_work}"
            parts.append(seg + "。")
        if volunteer_hours or volunteer_detail:
            seg = f"累计志愿服务时长 {volunteer_hours} 小时" if volunteer_hours else ""
            if volunteer_detail:
                seg = (seg + "：" if seg else "") + volunteer_detail
            parts.append(seg + "。")
        if parts:
            self.add_para("".join(parts))

    # --- 生活方面 ---

    def _add_life(self):
        """生活方面（100~150 字）：生活作风 + 人际关系（校级不写家庭困难）"""
        self.add_heading("四、生活方面")
        life = self._get("life", default="")
        if life:
            if isinstance(life, list):
                for p in life:
                    self.add_para(p)
            else:
                self.add_para(life)
            return
        dorm_role = self._get("dorm_role", default="")
        dorm_activity = self._get("dorm_activity", default="")
        dorm_honor = self._get("dorm_honor", default="")
        interpersonal = self._get("interpersonal", default="")
        lifestyle = self._get("lifestyle",
                               default="生活中我注重勤俭节约，作息规律。")
        parts = [lifestyle]
        if dorm_role and dorm_activity:
            seg = f"担任{dorm_role}期间，{dorm_activity}"
            if dorm_honor:
                seg += f"，{dorm_honor}"
            parts.append(seg + "。")
        if interpersonal:
            parts.append(interpersonal + "。")
        self.add_para("".join(parts))

    # --- 结尾"此致 敬礼！"---

    def _add_ending(self):
        """结尾（60~100 字）：事实总结 + 朴素表态 + 此致/敬礼！"""
        ending = self._get("ending", default="")
        if ending:
            self.add_para(ending)
        else:
            self.add_para(
                "以上是我本学年的基本情况。无论结果如何，"
                "我都将以此为新的起点，继续努力学习。"
                "恳请评审委员会予以考虑。"
            )
        add_cizhi_paragraph(self.doc, "此致")   # "此致"另起一行，空两格
        add_jingli_paragraph(self.doc, "敬礼！")  # "敬礼！"另起一行，顶格

    # --- 落款 ---

    def _add_signature(self):
        """落款：右对齐，含申请人 + 日期"""
        self.doc.add_paragraph()  # 空一行
        name = self._get("name", default="")
        apply_date = self._get("apply_date", default="")
        if name:
            add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if apply_date:
            add_right_aligned_paragraph(self.doc, apply_date)

    # --- 主构建方法 ---

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/开头/思想/学习/科研实践/生活/结尾/落款，生成 docx"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_opening()
            self._add_ideology()
            self._add_academics()
            self._add_research_practice()
            self._add_life()
            self._add_ending()
            self._add_signature()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申请书已生成：{output_path}")
        return str(output_path)

    # --- 数据校验（含等级门槛校验）---

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []
        # P0 必采字段
        p0_fields = [("name", "申请人姓名"), ("college", "学院"),
                     ("major", "专业"), ("grade", "年级"),
                     ("gpa", "GPA"), ("rank", "专业排名")]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")
        # 主干课程
        courses = self._get_list("core_courses")
        if not courses:
            warnings.append("缺少 主干课程（core_courses），将省略主干课程表")
        elif len(courses) < 4:
            warnings.append(f"主干课程仅 {len(courses)} 门，建议 4~6 门")
        # 学期 GPA 表
        semesters = self._get_list("semester_gpa")
        if not semesters:
            warnings.append("缺少 学期 GPA 表（semester_gpa），将省略学期 GPA 表")
        elif len(semesters) < 4:
            warnings.append(f"学期 GPA 表仅 {len(semesters)} 学期，建议 4 学期")
        # 排名校验
        rank_str = str(self._get("rank", default=""))
        if rank_str and "/" not in rank_str and not self._get("rank_total"):
            warnings.append(f"排名 '{rank_str}' 缺基数，应为 'X/N' 格式或补填 rank_total")
        # 等级与数据匹配校验
        level = str(self._get("level", default=""))
        weighted = self._get("weighted_avg", default="0")
        rank_total = self._get("rank_total", default="0")
        try:
            w = float(weighted)
        except (ValueError, TypeError):
            w = 0
        try:
            r_num = int(rank_str.split("/")[0]) if "/" in rank_str else 0
            r_total = int(rank_total) if rank_total else (
                int(rank_str.split("/")[1]) if "/" in rank_str else 0)
        except (ValueError, IndexError):
            r_num, r_total = 0, 0
        if level and r_total and w:
            pct = r_num / r_total * 100
            threshold_w, threshold_pct = self.LEVEL_THRESHOLDS.get(level, (0, 100))
            if w < threshold_w:
                warnings.append(f"加权平均分 {w} 不满足{level}门槛（≥{threshold_w}），建议改申低一等级")
            if pct > threshold_pct:
                warnings.append(f"排名前 {pct:.1f}% 不满足{level}门槛（前 {threshold_pct}%），建议改申低一等级")
        # 校级奖学金最低门槛校验
        if r_total and w:
            pct = r_num / r_total * 100
            if w < 75 or pct > 30:
                warnings.append(f"加权 {w} / 排名前 {pct:.1f}% 不满足校级奖学金最低门槛"
                                f"（加权 ≥75 + 前 30%），建议改申其他奖学金")
        self.warnings = warnings
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据（校级二等奖学金）
# ============================================================

DEFAULT_DATA = {
    "name": "王晨", "student_id": "2022123456", "gender": "男",
    "college": "计算机科学与技术学院", "major": "计算机科学与技术",
    "grade": "2022 级 大三", "class_name": "计科 2201 班",
    "political_status": "共青团员", "phone": "138XXXXXXXX",
    "apply_year": "2024-2025 学年", "apply_date": "2025 年 6 月 15 日",
    "salutation": "尊敬的学院领导、评审委员会：", "level": "二等",
    "gpa": "3.65/4.0", "weighted_avg": "85.2", "rank": "10/87",
    "rank_total": "87", "course_count": "12", "high_score_count": "9",
    "core_courses": [
        {"name": "高等数学（上）", "credit": "5", "score": "95"},
        {"name": "数据结构", "credit": "4", "score": "92"},
        {"name": "操作系统", "credit": "4", "score": "89"},
        {"name": "计算机网络", "credit": "3", "score": "88"},
        {"name": "数据库原理", "credit": "3", "score": "90"},
    ],
    "course_highlight": "高等数学 95、数据结构 92，专业核心课平均 90 分，全部 85+",
    "semester_gpa": [
        {"semester": "第 1 学期", "gpa": "3.58/4.0", "weighted": "84.5"},
        {"semester": "第 2 学期", "gpa": "3.62/4.0", "weighted": "84.8"},
        {"semester": "第 3 学期", "gpa": "3.65/4.0", "weighted": "85.5"},
        {"semester": "第 4 学期", "gpa": "3.65/4.0", "weighted": "85.2"},
    ],
    "cet4": "510", "cet6": "480", "computer_level": "二级 C 语言",
    "study_method": "学习上注重课前预习与课后总结，建立知识体系；遇到问题主动与老师、同学讨论。",
    "academics_summary": "",
    "party_history": "2023.09 提交入党申请书，2024.03 列为入党积极分子。",
    "party_activities": [
        "参加学院分党校第 8 期培训班（2024.09-2024.12）结业",
        "提交思想汇报 2 篇",
    ],
    "ideology": "", "academics": "", "research_practice": "",
    "competitions": [
        {"name": "校级程序设计竞赛", "award": "二等奖",
         "time": "2025.05", "role": "队长，负责算法实现"},
    ],
    "research_projects": [
        {"name": "校园学习助手小程序", "level": "校级大创项目",
         "role": "主持", "duration": "2024.03-2025.03",
         "output": "结题评估良好"},
    ],
    "position": "班级学习委员（2024.09-2025.06）",
    "position_work": "组织学习经验交流会 5 次，服务同学 30 余人次；"
                      "建立班级'一对一'帮扶对子 6 对；推动班级平均分较上学年提升 3.2 分",
    "volunteer_hours": "60",
    "volunteer_detail": "担任图书馆管理员（2024.09-2025.06，每月 5 小时）",
    "dorm_role": "宿舍长", "dorm_activity": "组织宿舍 6 次集体活动",
    "dorm_honor": "宿舍连续两学期获评'文明宿舍'",
    "interpersonal": "与同学相处融洽，曾帮助室友完成 1 次实验调试",
    "lifestyle": "生活中我注重勤俭节约，作息规律。",
    "life": "", "ending": "",
    "honors": [
        {"time": "2025.05", "name": "校级程序设计竞赛",
         "level": "校级二等奖", "issuer": "XX 大学"},
        {"time": "2024.11", "name": "校级三好学生",
         "level": "校级", "issuer": "XX 大学"},
        {"time": "2024.10", "name": "校级优秀共青团员",
         "level": "校级", "issuer": "XX 大学"},
    ],
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="校级奖学金申请书 docx 生成器（支持一等/二等/三等）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n"
            "JSON 字段定义详见 SKILL.md 第三章信息采集清单。\n"
            "必填字段：level(一等/二等/三等)、name、major、grade、gpa、rank、\n"
            "          core_courses、semester_gpa。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档（校级二等奖学金）")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档（校级二等奖学金）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)
    # 输出校验警告数量
    if builder.warnings:
        print(f"⚠️ 共 {len(builder.warnings)} 项校验警告，详见上方输出", file=sys.stderr)


if __name__ == "__main__":
    main()
