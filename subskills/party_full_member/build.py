#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
党员转正申请书 docx 生成器（v1.0）

依据 SKILL.md v1.0 规范实现。核心要点：
1. 5 段结构：开篇说明 + 预备期表现（8 项必汇报）+ 缺点不足 + 改进措施 + 转正态度
2. 8 项表现必汇报（理论学习/政治立场/工作学习/党员义务/党费缴纳/支部活动/联系群众/廉洁自律）
3. 3 档字数版本：brief=3000 / standard=3500 / enhanced=4000
4. 预备期日期校验：probation_end_date = probation_start_date +1 年（±7 天）；
   handover_date 在 probation_end_date 前 30~60 天（±7 天）
5. 假缺点检测 + 转正态度关键表述（"按期予以转正"）+ 改进措施与缺点一一对应 + 党章义务对照（obligation_ref 1~8）
6. 兼容字段：submit_date / introducer_1 / introducer_2 / theory_study_text / study_work_text

格式标准：A4 页边距上下 2.54cm 左右 2.5cm；标题黑体二号居中；
正文宋体小四 1.5 倍行距首行缩进 2 字符；"此致"另起一行空两格，"敬礼！"另起一行顶格；落款右对齐。

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

# ============================================================
# 字数版本与 5 段配比常量（SKILL.md §3 + §5.1）
# ============================================================

# word_count_version 字段：3 档字数对应不同 5 段配比
WORD_COUNT_VERSIONS: Dict[str, Dict[str, Any]] = {
    "brief":    {"target": 3000, "label": "精简版", "min": 2850, "max": 3150,
                 "ratios": {"opening": 0.08, "performance": 0.50,
                            "shortcomings": 0.14, "measures": 0.18,
                            "attitude": 0.10}},
    "standard": {"target": 3500, "label": "标准版", "min": 3325, "max": 3675,
                 "ratios": {"opening": 0.08, "performance": 0.52,
                            "shortcomings": 0.14, "measures": 0.18,
                            "attitude": 0.08}},
    "enhanced": {"target": 4000, "label": "加强版", "min": 3800, "max": 4200,
                 "ratios": {"opening": 0.08, "performance": 0.54,
                            "shortcomings": 0.13, "measures": 0.17,
                            "attitude": 0.08}},
}

# 8 项表现各项字数（standard 档，brief/enhanced 按比例缩放）
PERFORMANCE_ITEM_BASE_WORDS = {
    "theory_study": 250,
    "political_stand": 230,
    "study_work": 280,
    "party_obligations": 240,
    "dues_payment": 130,
    "branch_activities": 230,
    "mass_contact": 230,
    "integrity_self_discipline": 180,
}

# 8 项表现的中文段首引导词（"一、"、"二、"...）
PERFORMANCE_LEADS = [
    "一、理论学习方面。",
    "二、政治立场方面。",
    "三、工作学习方面。",
    "四、党员义务履行方面。",
    "五、党费缴纳方面。",
    "六、支部活动参与方面。",
    "七、联系群众方面。",
    "八、廉洁自律方面。",
]

# 8 项表现字段名（与 PERFORMANCE_ITEM_BASE_WORDS 对应顺序）
PERFORMANCE_FIELD_NAMES = list(PERFORMANCE_ITEM_BASE_WORDS.keys())

# 党章 8 项义务简表（用于缺点不足 obligation_ref 对照）
PARTY_OBLIGATIONS_8 = [
    "认真学习马克思列宁主义、毛泽东思想、邓小平理论、'三个代表'重要思想、"
    "科学发展观、习近平新时代中国特色社会主义思想",
    "贯彻执行党的基本路线和方针政策",
    "坚持党和人民利益高于一切",
    "自觉遵守党的纪律",
    "维护党的团结和统一",
    "开展批评和自我批评",
    "密切联系群众",
    "发扬社会主义新风尚",
]

# 假缺点禁用关键词清单（SKILL.md §0.5 + §10.4）
FAKE_SHORTCOMING_KEYWORDS = [
    "工作太投入", "追求完美", "学习太刻苦", "责任心太强",
    "为人太直率", "事必躬亲", "太较真", "工作太认真",
    "对自己要求太高", "太执着",
]

# 党章原文高风险片段（连续 50 字重复则警告）
PARTY_CONSTITUTION_FRAGMENTS = [
    "中国共产党是中国工人阶级的先锋队，同时是中国人民和中华民族的先锋队，"
    "是中国特色社会主义事业的领导核心，代表中国先进生产力的发展要求，"
    "代表中国先进文化的前进方向，代表中国最广大人民的根本利益。",
    "全心全意为人民服务。党除了工人阶级和最广大人民群众的利益，"
    "没有自己特殊的利益。",
]

# 网络模板高风险片段（连续 30 字重复则警告）
NETWORK_TEMPLATE_FRAGMENTS = [
    "在党组织的培养教育下，我深刻认识到自身不足",
    "通过一年的预备期，我在各方面都有了很大提高",
    "我将继续以党员标准严格要求自己，争取早日成为一名合格的共产党员",
    "现将预备期一年来的思想、学习、工作、生活情况向党组织汇报如下",
]


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进），宋体小四"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_cizhi_paragraph(doc, text: str = "此致"):
    """'此致'另起一行，空两格（首行缩进 2 字符）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格（不缩进）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=False, line_spacing=1.5)


def add_blank_paragraph(doc):
    """空段落，用于段落间留白"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return p


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def count_chinese_chars(text: str) -> int:
    """统计中文字符数（用于字数控制）"""
    if not text:
        return 0
    return len(re.findall(r"[\u4e00-\u9fff]", str(text)))


def parse_date_cn(date_str: str) -> Optional[datetime]:
    """解析'YYYY 年 M 月 D 日'格式的中文日期"""
    if not date_str:
        return None
    m = re.match(r"\s*(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日\s*", str(date_str))
    if not m:
        return None
    try:
        return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    except ValueError:
        return None


def check_plagiarism_risk(text: str, fragments: List[str],
                          threshold: int = 50) -> List[str]:
    """查重风险检测：检查连续 N 字与参考片段重复

    Args:
        text: 待检测的文本
        fragments: 党章原文/网络模板/历史汇报片段列表
        threshold: 连续重复字数阈值

    Returns:
        触发查重警告的片段列表
    """
    warnings = []
    if not text:
        return warnings
    text_clean = re.sub(r"\s+", "", str(text))
    for fragment in fragments:
        if not fragment:
            continue
        frag_clean = re.sub(r"\s+", "", str(fragment))
        if len(frag_clean) < threshold:
            continue
        for i in range(len(frag_clean) - threshold + 1):
            window = frag_clean[i:i + threshold]
            if window in text_clean:
                warnings.append(f"检测到连续 {threshold} 字与参考片段重复：{window[:20]}...")
                break  # 每个 fragment 只报告一次
    return warnings


def check_political_terms(text: str) -> List[str]:
    """政治用语规范校验：必提要点 + 禁用简写 + 指导思想顺序

    Returns:
        警告信息列表
    """
    warnings = []
    if not text:
        return warnings
    text_str = str(text)

    # 必提要点（6 项缺一不可）
    required_terms = [
        "中国共产党", "习近平新时代中国特色社会主义思想", "党的二十大",
        "两个确立", "两个维护", "全心全意为人民服务",
    ]
    for term in required_terms:
        if term not in text_str:
            warnings.append(f"必提要点缺失：'{term}'")

    # 禁用简写检测
    forbidden_abbreviations = [
        ("习近平思想", "应为'习近平新时代中国特色社会主义思想'"),
        ("习思想", "应为'习近平新时代中国特色社会主义思想'"),
        ("党的20大", "应为'党的二十大'"),
        ("马列", "应为'马克思列宁主义'"),
        ("毛思想", "应为'毛泽东思想'"),
        ("邓理论", "应为'邓小平理论'"),
        ("科学观", "应为'科学发展观'"),
    ]
    for abbr, suggestion in forbidden_abbreviations:
        if abbr in text_str:
            warnings.append(f"禁用简写'{abbr}'：{suggestion}")
    # "20大"检测
    for _ in re.finditer(r"20\s*大", text_str):
        warnings.append("禁用简写'20大'（应完整表述为'党的二十大'）")
        break
    # "二十大"检测（排除"党的二十大"上下文）
    for match in re.finditer(r"二十大", text_str):
        start = match.start()
        if start >= 2 and text_str[start - 2:start + 3] == "党的二十大":
            continue
        warnings.append("禁用简写'二十大'（应完整表述为'党的二十大'）")
        break
    # "三中全会"检测（须带"党的二十届"前缀）
    for match in re.finditer(r"三中全会", text_str):
        start = match.start()
        if start >= 5 and text_str[start - 5:start + 4] == "党的二十届三中全会":
            continue
        warnings.append("禁用简写'三中全会'（应完整表述为'党的二十届三中全会'）")
        break

    # 指导思想 6 项顺序检测
    guideline_pattern = (
        r"马克思列宁主义.*?毛泽东思想.*?邓小平理论.*?"
        r"[\u201c\"']三个代表[\u201d\"']重要思想.*?科学发展观.*?"
        r"习近平新时代中国特色社会主义思想"
    )
    if not re.search(guideline_pattern, text_str, re.DOTALL):
        checks = [
            ("马克思列宁主义", "马克思列宁主义"),
            ("毛泽东思想", "毛泽东思想"),
            ("邓小平理论", "邓小平理论"),
            ("'三个代表'重要思想", "三个代表重要思想"),
            ("\"三个代表\"重要思想", "三个代表重要思想"),
            ("\u201c三个代表\u201d重要思想", "三个代表重要思想"),
            ("科学发展观", "科学发展观"),
            ("习近平新时代中国特色社会主义思想", "习近平新时代中国特色社会主义思想"),
        ]
        required_keys = ["马克思列宁主义", "毛泽东思想", "邓小平理论",
                         "三个代表重要思想", "科学发展观", "习近平新时代中国特色社会主义思想"]
        found = {key for needle, key in checks if needle in text_str}
        missing = [k for k in required_keys if k not in found]
        if missing:
            warnings.append(f"指导思想漏项：{', '.join(missing)}")
        else:
            warnings.append("指导思想 6 项顺序错误（应为马列→毛→邓→三→科→习）")

    return warnings


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """党员转正申请书 docx 构建器

    5 段结构：开篇说明 / 预备期表现（8 项必汇报）/ 缺点不足 /
    改进措施 / 转正态度。8 项表现缺一警告；缺点不足与改进措施须一一对应。
    """

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        # Normal 样式默认设置为宋体小四（中英文同步）
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []
        self.full_text_parts: List[str] = []  # 全文缓存，用于查重检测

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段，非列表返回空列表"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def _resolve_field(self, *keys, default=""):
        """多备选字段取值：按顺序取第一个非空字段"""
        for key_path in keys:
            if isinstance(key_path, str):
                val = self._get(key_path, default="")
            else:
                val = self._get(*key_path, default="")
            if val:
                return val
        return default

    def add_para(self, text, indent=True):
        """添加正文段落并记录到全文缓存（用于查重检测）"""
        if text:
            self.full_text_parts.append(str(text))
        return add_body_paragraph(self.doc, text, indent=indent)

    def _resolve_word_count_version(self) -> str:
        """解析字数版本字段：brief/standard/enhanced，默认 standard"""
        v = str(self._get("word_count_version", default="standard")).lower().strip()
        if v in ("brief", "standard", "enhanced"):
            return v
        return "standard"

    # --- 标题 ---

    def _add_title(self):
        """标题：黑体二号居中，固定为'转正申请书'5 字"""
        custom_title = self._get("title", default="")
        title_text = custom_title if custom_title else "转正申请书"
        add_title(self.doc, title_text)
        add_blank_paragraph(self.doc)

    # --- 称呼 ---

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号，固定为'敬爱的党组织：'"""
        salutation = self._get("salutation", default="敬爱的党组织：")
        add_salutation_paragraph(self.doc, salutation)

    # --- 开篇说明段（200~320 字） ---

    def _add_opening(self):
        """开篇说明段：姓名 + 党支部 + 预备期起止 + 申请转正意愿

        若用户提供 opening_text 字段则直接使用，否则按模板拼装。
        """
        opening_text = self._get("opening_text", default="")
        if opening_text:
            self.add_para(opening_text)
            return

        name = self._get("name", default="XXX")
        party_branch = self._get("party_branch", default="XX 党支部")
        start_date = self._resolve_field("probation_start_date", default="XXXX 年 X 月 X 日")
        end_date = self._resolve_field("probation_end_date", default="XXXX 年 X 月 X 日")

        opening = (
            f"我叫{name}，是{party_branch}于{start_date}党支部大会通过接收的预备党员，"
            f"预备期自{start_date}至{end_date}，为期 1 年。现预备期即将届满，"
            "特向党组织郑重提出转正申请，并就预备期 1 年内的思想、学习、工作、"
            "生活等情况向党组织作全面汇报，恳请党组织审阅并按期予以转正。"
        )
        self.add_para(opening)
        self.add_para("现将预备期 1 年内的表现汇报如下：")

    # --- 第一段：预备期表现（8 项必汇报，1500~2200 字） ---

    def _add_performance_overview(self):
        """8 项表现段引导句（已在 _add_opening 末尾添加，此处空实现）"""
        pass

    def _add_theory_study(self):
        """① 理论学习情况（250 字，含原著+时政+党课）"""
        text = self._resolve_field("theory_study", "theory_study_text", default="")
        if not text:
            text = (
                "一、理论学习方面。预备期 1 年内，我系统学习了《习近平著作选读》"
                "第一卷、第二卷，重点研读了'关于坚持和发展中国特色社会主义的几个"
                "问题''在庆祝中国共产党成立 100 周年大会上的讲话'等篇目；"
                "认真学习了党的二十大精神、党的二十届三中全会精神、2024 年与 2025 年"
                "全国两会精神等时政热点；按时参加党支部组织的全部党课学习 6 次，含支部书记"
                "讲授的'两个确立'专题党课、学院党委书记讲授的'七一'专题党课等。"
                "通过学习，我进一步深化了对习近平新时代中国特色社会主义思想的"
                "理解，进一步坚定了对中国特色社会主义道路的信念。"
            )
        else:
            text = self._ensure_lead(text, "一、理论学习方面。")
        self.add_para(text)

    def _add_political_stand(self):
        """② 政治立场坚定情况（230 字，含两个确立+两个维护+四个意识+四个自信）"""
        text = self._get("political_stand", default="")
        if not text:
            text = (
                "二、政治立场方面。预备期 1 年内，我始终坚定政治立场，"
                "深刻领悟'两个确立'的决定性意义——党确立习近平同志党中央的核心、"
                "全党的核心地位，确立习近平新时代中国特色社会主义思想的指导地位，"
                "是新时代党和国家事业取得历史性成就、发生历史性变革的根本保证。"
                "我自觉做到'两个维护'，坚决维护习近平总书记党中央的核心、全党的"
                "核心地位，坚决维护党中央权威和集中统一领导；不断增强'四个意识'"
                "——政治意识、大局意识、核心意识、看齐意识；坚定'四个自信'"
                "——道路自信、理论自信、制度自信、文化自信。在大是大非面前，"
                "我始终与党中央保持高度一致。"
            )
        else:
            text = self._ensure_lead(text, "二、政治立场方面。")
        self.add_para(text)

    def _add_study_work(self):
        """③ 工作学习情况（280 字，含学业+获奖+学生工作+科研竞赛）"""
        text = self._resolve_field("study_work", "study_work_text", default="")
        if not text:
            text = (
                "三、工作学习方面。学业上，预备期内两学期加权平均分分别为 87.3、"
                "89.1，专业排名从 12/120 提升至 8/120；《机器学习》《计算机网络》"
                "两门核心课程均取得 90 分以上。获奖方面，获得国家励志奖学金 1 次、"
                "校一等奖学金 1 次、全国大学生数学建模竞赛省级二等奖 1 次、"
                "'校优秀学生干部'荣誉称号 1 次。学生工作上，担任班长期间组织班级"
                "'一对一'帮扶活动 4 期、班级红色教育基地参观 2 次、班级调研报告 "
                "2 份。科研竞赛方面，作为负责人主持校级大创项目《基于对比学习的"
                "法律问答系统》，预备期内完成中期答辩与结题答辩，论文 1 篇已被 EI 会议录用。"
            )
        else:
            text = self._ensure_lead(text, "三、工作学习方面。")
        self.add_para(text)

    def _add_party_obligations(self):
        """④ 党员义务履行情况（240 字，党章 8 项义务逐项对照）"""
        text = self._get("party_obligations", default="")
        if not text:
            text = (
                "四、党员义务履行方面。对照党章第三条规定的 8 项党员义务，我逐项"
                "检视：第 1 项，认真学习党的理论，按时参加党课学习；第 2 项，"
                "贯彻执行党的基本路线和方针政策，关注国家大事；第 3 项，坚持党和"
                "人民利益高于一切，志愿服务时长累计 36 小时；第 4 项，自觉遵守党的"
                "纪律，按时参加组织生活；第 5 项，维护党的团结和统一，不传播不实"
                "信息；第 6 项，开展批评和自我批评，参加组织生活会 4 次并主动发言；"
                "第 7 项，密切联系群众，联系班级同学 6 人；第 8 项，发扬社会主义"
                "新风尚，参与校园文明引导 3 次。"
            )
        else:
            text = self._ensure_lead(text, "四、党员义务履行方面。")
        self.add_para(text)

    def _add_dues_payment(self):
        """⑤ 党费缴纳情况（130 字，含标准+方式+情况）"""
        text = self._get("dues_payment", default="")
        if not text:
            text = (
                "五、党费缴纳方面。预备期 1 年内，我按学生党员标准每月按时缴纳"
                "党费 0.2 元，共缴纳 12 次，无漏缴、无补缴。每月党费均于当月 10 日"
                "前主动交给党支部组织委员，并核对收据。我深知按时缴纳党费是党员的"
                "基本义务，也是党员组织观念的具体体现，从未因事因病延误缴纳。"
            )
        else:
            text = self._ensure_lead(text, "五、党费缴纳方面。")
        self.add_para(text)

    def _add_branch_activities(self):
        """⑥ 支部活动参与情况（230 字，含三会一课+主题党日+组织生活会+民主评议党员）"""
        text = self._get("branch_activities", default="")
        if not text:
            text = (
                "六、支部活动参与方面。预备期 1 年内，我按时参加党支部组织的全部"
                "活动，无缺席、无迟到、无早退。其中，参加支部党员大会 6 次、支委会"
                "（扩大）会议 2 次、党小组会 12 次、党课 6 次；参加主题党日活动 4 次，"
                "含赴西柏坡红色教育基地参观、走访抗战老兵、'七一'重温入党誓词等；"
                "参加组织生活会 4 次，每次均主动发言开展批评与自我批评；参加民主"
                "评议党员 1 次，自评为'合格'等次。通过参加支部活动，我进一步增强了"
                "党员意识、组织观念。"
            )
        else:
            text = self._ensure_lead(text, "六、支部活动参与方面。")
        self.add_para(text)

    def _add_mass_contact(self):
        """⑦ 联系群众情况（230 字，含班级+宿舍+团支部+群众反馈）"""
        text = self._get("mass_contact", default="")
        if not text:
            text = (
                "七、联系群众方面。预备期 1 年内，我主动联系班级同学 6 人、宿舍"
                "同学 4 人、团支部成员 8 人，累计开展谈心谈话 24 次。重点联系班级 2 名"
                "经济困难同学，协助申请助学金、提供勤工助学信息；联系 1 名学习困难"
                "同学，每周开展'一对一'学业辅导 2 次，期末其专业排名提升 5 位；联系"
                "宿舍同学组织'宿舍读书会'12 期。通过走访班级宿舍 8 次、参加团支部"
                "活动 4 次，我广泛征求同学意见，同学普遍反馈我'平易近人、乐于助人'，"
                "但也指出'有时工作方法不够灵活'。"
            )
        else:
            text = self._ensure_lead(text, "七、联系群众方面。")
        self.add_para(text)

    def _add_integrity(self):
        """⑧ 廉洁自律情况（180 字，含校规+生活作风+网络言行+党员形象）"""
        text = self._get("integrity_self_discipline", default="")
        if not text:
            text = (
                "八、廉洁自律方面。预备期 1 年内，我严格遵守校规校纪，无任何违纪"
                "记录；生活作风勤俭节约，未参与任何铺张浪费、攀比消费；网络言行规范，"
                "未在任何社交平台发表不当言论、未传播不实信息、未参与网络骂战；"
                "自觉维护党员形象，在校内外均以党员标准要求自己，未发生有损党员形象"
                "的行为。我深知党员的一言一行都代表着党的形象，必须时刻自重、自省、"
                "自警、自励。"
            )
        else:
            text = self._ensure_lead(text, "八、廉洁自律方面。")
        self.add_para(text)

    def _ensure_lead(self, text: str, lead: str) -> str:
        """若 text 不以 lead 开头，则前置 lead（保证 8 项表现段首编号统一）"""
        if not text:
            return lead
        text_str = str(text).lstrip()
        if text_str.startswith(lead):
            return text_str
        # 兼容用户提供的"一、..."与"理论学习方面。..."两种写法
        if text_str.startswith(lead[0]) and lead[0] in "一二三四五六七八":
            return text_str
        return lead + text_str

    # --- 第二段：缺点不足（420~520 字，3~4 个真实不足） ---

    def _add_shortcomings(self):
        """缺点不足段：3~4 个真实不足，含表现+根源+党章义务对照

        若用户提供 self_shortcomings 列表则按列表拼装，否则使用 3 个默认真实不足。
        """
        shortcomings = self._get_list("self_shortcomings")
        if not shortcomings:
            shortcomings = [
                {"desc": "理论学习系统性不够",
                 "manifestation": "对马克思主义经典著作（如《共产党宣言》《资本论》第一卷）"
                 "只读过简写本，未读全本，且学习时间碎片化，缺乏系统规划",
                 "root_cause": "对理论学习的认识仍停留在'完成任务'层面，未上升到'武装头脑'高度",
                 "obligation_ref": 1},
                {"desc": "联系群众深度不够",
                 "manifestation": "联系班级同学 6 人，但多为学习帮扶，对同学的思想动态、心理压力关注不够；"
                 "3 名心理压力较大同学未及时识别并提供帮助",
                 "root_cause": "把'联系群众'等同于'完成帮扶任务'，未上升到'全心全意为人民服务'高度",
                 "obligation_ref": 7},
                {"desc": "批评与自我批评勇气不够",
                 "manifestation": "参加组织生活会时，对他人批评较多、对自己批评较浅，对支部同志的不足"
                 "有时碍于情面不敢直言",
                 "root_cause": "把'和气'等同于'团结'，未理解'团结—批评—团结'的辩证关系",
                 "obligation_ref": 6},
            ]

        prefix = "在汇报成绩的同时，我也清醒地认识到自身存在以下不足："
        parts = []
        num_words = ["一", "二", "三", "四", "五"]
        for idx, item in enumerate(shortcomings):
            if isinstance(item, dict):
                desc = item.get("desc", "")
                manifest = item.get("manifestation", "")
                root = item.get("root_cause", "")
                obl_ref = item.get("obligation_ref", 0)
                num_word = num_words[idx] if idx < len(num_words) else str(idx + 1)
                obl_text = self._format_obligation_ref(obl_ref)
                parts.append(
                    f"{num_word}是{desc}。具体表现为{manifest}；"
                    f"根源在于{root}，对照党章第{obl_ref}项党员义务"
                    f"（{obl_text}），履行不够到位。"
                )
            else:
                parts.append(str(item))
        self.add_para(prefix + "".join(parts))

    def _format_obligation_ref(self, obl_ref: int) -> str:
        """根据 obligation_ref（1~8）返回党章义务简述"""
        try:
            idx = int(obl_ref) - 1
            if 0 <= idx < len(PARTY_OBLIGATIONS_8):
                return PARTY_OBLIGATIONS_8[idx]
        except (TypeError, ValueError):
            pass
        return "党员义务"

    # --- 第三段：改进措施（540~680 字，3~4 条措施，与缺点一一对应） ---

    def _add_improvement_measures(self):
        """改进措施段：3~4 条措施，与缺点一一对应

        每条措施含具体行动 + 时间节点 + 可检验标志。
        若用户提供 improvement_measures 列表则按列表拼装，否则使用 3 条默认措施。
        """
        measures = self._get_list("improvement_measures")
        if not measures:
            measures = [
                {"action": "通读《共产党宣言》《资本论》第一卷全本，每月撰写 1 篇读书笔记交培养"
                 "联系人审阅；每月参加 1 次'理论读书会'",
                 "timeline": "转正后 1 年内（2025 年 6 月 ~ 2026 年 6 月）",
                 "measurable": "2026 年 6 月前完成 2 本原著通读 + 12 篇读书笔记"},
                {"action": "每月走访班级宿舍 2 次，重点识别心理压力较大同学；与 3 名心理压力较大"
                 "同学建立'一对一'联系，每月谈心 1 次，必要时协助联系学校心理咨询中心",
                 "timeline": "2025 年 6 月 ~ 2025 年 12 月",
                 "measurable": "2025 年 12 月前完成 3 名同学的深度联系 + 谈心记录 18 次"},
                {"action": "在每次组织生活会前准备 1 份'批评与自我批评提纲'，对自己批评不少于 "
                 "3 条、对支部同志批评不少于 2 条",
                 "timeline": "2025 年 6 月 ~ 2025 年 12 月",
                 "measurable": "2025 年 12 月组织生活会上完成 1 次完整批评与自我批评，获得支部书记"
                 "与培养联系人书面认可"},
            ]

        prefix = "针对上述不足，我制定了以下改进措施："
        parts = []
        num_words = ["一", "二", "三", "四", "五"]
        for idx, item in enumerate(measures):
            if isinstance(item, dict):
                action = item.get("action", "")
                timeline = item.get("timeline", "")
                measurable = item.get("measurable", "")
                num_word = num_words[idx] if idx < len(num_words) else str(idx + 1)
                parts.append(
                    f"{num_word}、{action}；时间节点：{timeline}；"
                    f"可检验标志：{measurable}。"
                )
            else:
                parts.append(str(item))
        self.add_para(prefix + "".join(parts))

    # --- 第四段：转正态度（200~320 字，必含"按期予以转正"） ---

    def _add_conversion_attitude(self):
        """转正态度段：必含"恳请党组织按期予以转正"。用户提供 conversion_attitude / conversion_attitude_text 则直接使用。"""
        text = self._resolve_field(
            "conversion_attitude", "conversion_attitude_text", default="")
        if not text:
            text = (
                "综上所述，预备期 1 年来，我在党组织的培养教育下，在思想政治、学习工作、"
                "联系群众、廉洁自律等方面均取得了一定进步，但与正式党员标准相比仍有差距。"
                "在此，我郑重向党组织提出转正申请，恳请党组织按期予以转正。"
                "无论党组织作出何种决定，我都将坚决服从，继续以党员标准严格要求自己，"
                "自觉接受党组织考验，以更饱满的热情投入到学习、工作、生活中去，"
                "努力做一名让党组织放心、让人民群众满意的合格共产党员。"
                "恳请党组织审阅我的申请，并按期予以转正。"
            )
        self.add_para(text)

    # --- 结尾"恳请党组织按期予以转正。此致 敬礼！" ---

    def _add_ending(self):
        """结尾：此致 + 敬礼！。转正申请书结尾句为'恳请党组织按期予以转正。'。"""
        # 转正态度段已包含结尾意愿，此处仅添加此致+敬礼
        add_cizhi_paragraph(self.doc, "此致")
        add_jingli_paragraph(self.doc, "敬礼！")

    # --- 落款（右对齐） ---

    def _add_signature(self):
        """落款：申请人 + 递交日期，右对齐

        与入党申请书落款一致：'申请人：XXX'；日期为 handover_date。
        """
        add_blank_paragraph(self.doc)
        name = self._get("name", default="申请人")
        handover_date = self._resolve_field(
            "handover_date", "submit_date", default="")
        add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if handover_date:
            add_right_aligned_paragraph(self.doc, handover_date)

    # --- 主构建方法 ---

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/开篇/8 项表现/缺点/措施/转正态度/结尾/落款

        5 段正文为转正申请书核心：开篇说明 / 预备期表现（8 项必汇报）/ 缺点不足 / 改进措施 / 转正态度。
        """
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_opening()
            self._add_performance_overview()
            # 8 项表现，按 PERFORMANCE_FIELD_NAMES 顺序调用
            self._add_theory_study()
            self._add_political_stand()
            self._add_study_work()
            self._add_party_obligations()
            self._add_dues_payment()
            self._add_branch_activities()
            self._add_mass_contact()
            self._add_integrity()
            self._add_shortcomings()
            self._add_improvement_measures()
            self._add_conversion_attitude()
            self._add_ending()
            self._add_signature()
            self._post_build_checks()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 转正申请书已生成：{output_path}")
        return str(output_path)

    def _post_build_checks(self):
        """构建后检查：政治用语规范 + 查重风险检测 + 字数检查"""
        full_text = "".join(self.full_text_parts)

        # 政治用语规范检查
        for w in check_political_terms(full_text):
            self.warnings.append(f"[政治规范] {w}")

        # 三重查重检测：党章原文 50 字 / 网络模板 30 字
        plagiarism_checks = [
            ("查重风险-党章原文", PARTY_CONSTITUTION_FRAGMENTS, 50),
            ("查重风险-网络模板", NETWORK_TEMPLATE_FRAGMENTS, 30),
        ]
        for label, frags, thr in plagiarism_checks:
            for w in check_plagiarism_risk(full_text, frags, threshold=thr):
                self.warnings.append(f"[{label}] {w}")

        # 字数检查（按 word_count_version 目标字数）
        version = self._resolve_word_count_version()
        vc = WORD_COUNT_VERSIONS[version]
        char_count = count_chinese_chars(full_text)
        if char_count < vc["min"]:
            self.warnings.append(
                f"[字数] 全文仅 {char_count} 字，建议 {version} 档"
                f"（{vc['target']} 字，区间 {vc['min']}~{vc['max']}）")
        elif char_count > vc["max"]:
            self.warnings.append(
                f"[字数] 全文 {char_count} 字偏多，建议压缩至 {vc['max']} 字以内（{version} 档）")
        else:
            print(f"ℹ️ 字数：{char_count}（{version} 档目标 {vc['target']}，"
                  f"区间 {vc['min']}~{vc['max']}）")

        if self.warnings:
            print("⚠️ 构建后检查警告：", file=sys.stderr)
            for w in self.warnings:
                print(f"  - {w}", file=sys.stderr)

    # --- 数据校验 ---

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）

        包含政治规范校验：预备期日期匹配、递交窗口匹配、
        8 项表现字段缺失、缺点不足数量、改进措施与缺点对应、
        假缺点检测、转正态度关键表述。
        """
        warnings = []

        # P0 必采字段
        p0_fields = [
            ("name", "申请人姓名", None), ("college", "学院", None),
            ("major", "专业", None), ("grade", "年级", None),
            ("party_branch", "党支部全称", None),
            ("probation_start_date", "预备期起算日", None),
            ("probation_end_date", "预备期满日", None),
            ("handover_date", "转正申请书递交日", "submit_date"),
            ("introducer_primary", "第一入党介绍人", "introducer_1"),
            ("introducer_secondary", "第二入党介绍人", "introducer_2"),
        ]
        for key, name, alt_key in p0_fields:
            val = self._get(key, default="")
            if not val and alt_key:
                val = self._get(alt_key, default="")
            if not val:
                warnings.append(f"缺少 {name}（{key}）")

        # 预备期日期校验：end_date 应为 start_date +1 年（±7 天容忍）
        start_dt = parse_date_cn(self._get("probation_start_date", default=""))
        end_dt = parse_date_cn(self._get("probation_end_date", default=""))
        if start_dt and end_dt:
            expected_end = start_dt + timedelta(days=365)  # +1 年近似为 +365 天
            diff_days = abs((end_dt - expected_end).days)
            if diff_days > 7:
                warnings.append(
                    f"预备期日期不匹配：probation_end_date 应为 "
                    f"probation_start_date +1 年（±7 天），实际偏差 {diff_days} 天")

        # 递交窗口校验：handover_date 应在 probation_end_date 前 30~60 天（±7 天容忍）
        handover_dt = parse_date_cn(
            self._resolve_field("handover_date", "submit_date", default=""))
        if end_dt and handover_dt:
            days_before_end = (end_dt - handover_dt).days
            if days_before_end < 23 or days_before_end > 67:
                warnings.append(
                    f"递交窗口不匹配：handover_date 应在 probation_end_date 前 30~60 天"
                    f"（±7 天），实际为前 {days_before_end} 天")

        # 字数版本校验
        version = str(self._get("word_count_version", default="standard")).lower().strip()
        if version and version not in ("brief", "standard", "enhanced"):
            warnings.append(f"字数版本'{version}'不规范，应为 brief/standard/enhanced 之一")

        # 8 项表现字段缺失校验（缺一警告，使用默认值）
        missing_performance = [f for f in PERFORMANCE_FIELD_NAMES if not self._get(f, default="")]
        if missing_performance:
            warnings.append(
                f"8 项表现字段缺失 {len(missing_performance)} 项："
                f"{', '.join(missing_performance)}（将使用默认值，请用户核实替换）")

        # 缺点不足数量校验
        shortcomings = self._get_list("self_shortcomings")
        if shortcomings and len(shortcomings) < 3:
            warnings.append(f"缺点不足仅 {len(shortcomings)} 个，建议 3~4 个真实不足")
        if shortcomings and len(shortcomings) > 4:
            warnings.append(f"缺点不足 {len(shortcomings)} 个偏多，建议压缩至 3~4 个")

        # 改进措施与缺点数量一致性校验
        measures = self._get_list("improvement_measures")
        if shortcomings and measures and len(measures) != len(shortcomings):
            warnings.append(
                f"改进措施（{len(measures)} 条）与缺点不足（{len(shortcomings)} 个）"
                f"数量不匹配，须一一对应")

        # 假缺点检测（党支部一眼识破，会要求重写）
        shortcomings_text = "".join(
            str(item.get("desc", "")) + str(item.get("manifestation", "")) +
            str(item.get("root_cause", ""))
            for item in shortcomings if isinstance(item, dict))
        for kw in FAKE_SHORTCOMING_KEYWORDS:
            if kw in shortcomings_text:
                warnings.append(
                    f"[政治红线] 个人不足出现假缺点'{kw}'，党支部会要求重写，请用户重新表述")
                break

        # obligation_ref 取值校验（1~8）
        for item in shortcomings:
            if isinstance(item, dict):
                obl_ref = item.get("obligation_ref", 0)
                try:
                    obl_int = int(obl_ref)
                    if obl_int < 1 or obl_int > 8:
                        warnings.append(f"obligation_ref={obl_int} 不在 1~8 范围内，应为党章 8 项义务之一")
                except (TypeError, ValueError):
                    if obl_ref:
                        warnings.append(f"obligation_ref='{obl_ref}' 不是数字，应为 1~8")

        # 转正态度关键表述校验
        conversion_attitude = self._resolve_field(
            "conversion_attitude", "conversion_attitude_text", default="")
        if conversion_attitude and "按期予以转正" not in conversion_attitude:
            warnings.append("[政治红线] 转正态度段未含'按期予以转正'关键表述，党支部会要求重写")

        self.warnings.extend(warnings)
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ==================== 默认示例数据（standard 档 3500 字，2025 年 5 月递交） ====================
# 8 项表现 / 缺点不足 / 改进措施 / 转正态度字段省略，由各 _add_* 方法使用内置默认值。

DEFAULT_DATA = {
    "name": "张明",
    "student_id": "2022123456",
    "gender": "男",
    "college": "计算机科学与技术学院",
    "major": "计算机科学与技术",
    "grade": "2022 级 大三",
    "class_name": "计科 2201 班",
    "party_branch": "计算机科学与技术学院本科生第一党支部",
    "probation_start_date": "2024 年 6 月 15 日",
    "probation_end_date": "2025 年 6 月 15 日",
    "handover_date": "2025 年 5 月 8 日",
    "submit_date": "2025 年 5 月 8 日",
    "introducer_primary": "李志强",
    "introducer_secondary": "王建华",
    "introducer_1": "李志强",
    "introducer_2": "王建华",
    "cultivation_contact": "陈卫东",
    "word_count_version": "standard",
    "title": "转正申请书",
    "salutation": "敬爱的党组织：",
    "ending_line": "恳请党组织按期予以转正。",
    # 8 项表现字段（theory_study / political_stand / study_work / party_obligations /
    # dues_payment / branch_activities / mass_contact / integrity_self_discipline）
    # 缺省时由各 _add_* 方法使用内置默认值（详 §5.5.2~§5.5.9 模板示例）。
    # self_shortcomings / improvement_measures / conversion_attitude 同理。
}


# ==================== CLI 入口 ====================

def main():
    parser = argparse.ArgumentParser(
        description="党员转正申请书 docx 生成器（预备期满前 1~2 个月递交）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\nJSON 字段定义详见 SKILL.md §12 信息采集清单。\n"
            "政治规范校验：必提 6 项要点 + 指导思想 6 项顺序 + 三重查重检测。\n"
            "字数版本 word_count_version 取 brief(3000) / standard(3500) / enhanced(4000)，默认 standard。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档（standard 档 3500 字，2025 年 5 月递交）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
