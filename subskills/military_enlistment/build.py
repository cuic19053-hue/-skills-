#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
大学生应征入伍申请书 docx 生成器

按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 一级标题：黑体三号，居中
- 二级标题：黑体小三，左对齐
- 三级标题：宋体四号加粗
- 表格：宋体五号，居中

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第十一章。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)
SIZE_XIAO_ER = Pt(18)
SIZE_SAN = Pt(16)
SIZE_XIAO_SAN = Pt(15)
SIZE_SI = Pt(14)
SIZE_XIAO_SI = Pt(12)
SIZE_WU = Pt(10.5)
SIZE_XIAO_WU = Pt(9)

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# 入伍类型代码 → 中文全称映射
ENLISTMENT_TYPE_MAP = {
    "在校生入伍": "在校生入伍",
    "毕业生入伍": "毕业生入伍",
    "in_school": "在校生入伍",
    "graduate": "毕业生入伍",
}

# 入伍类型 → 适用人群描述
ENLISTMENT_TYPE_DESC = {
    "在校生入伍": "普通高校在校本专科生、研究生（保留学籍至退役后 2 年内复学）",
    "毕业生入伍": "普通高校应届本专科毕业生、研究生（已毕业无学籍问题）",
}


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
        space_before=12, space_after=12,
    )


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6, space_after=6,
    )


def add_heading_level3(doc, text: str):
    """三级标题：宋体四号加粗，左对齐"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6, space_after=3,
    )


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent,
        line_spacing=1.5,
    )


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j >= 1 and len(headers) >= 3 \
                else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False, alignment=align)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def safe_int(value, default: int = 0) -> int:
    """安全转整数"""
    try:
        return int(str(value).strip())
    except (ValueError, TypeError):
        return default


def safe_float(value, default: float = 0.0) -> float:
    """安全转浮点"""
    try:
        return float(str(value).strip())
    except (ValueError, TypeError):
        return default


def vision_to_score(vision: str) -> float:
    """视力值转数值（用于体检硬条件校验）

    支持 5 分制（如 4.8）与小数制（如 0.6）。
    5 分制范围 4.0-5.3，小数制范围 0.1-2.0。
    """
    if not vision:
        return 0.0
    try:
        v = float(str(vision).strip())
    except (ValueError, TypeError):
        return 0.0
    return v


def count_chinese_chars(text: str) -> int:
    """统计中文字符数（粗略版，用于字数校验）"""
    if not text:
        return 0
    count = 0
    for ch in str(text):
        if '\u4e00' <= ch <= '\u9fff':
            count += 1
    return count


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """大学生应征入伍申请书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def add_h1(self, text):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text):
        return add_heading_level2(self.doc, text)

    def add_h3(self, text):
        return add_heading_level3(self.doc, text)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_table(self, headers, rows, col_widths=None):
        return add_table_from_data(self.doc, headers, rows, col_widths)

    def add_page_break(self):
        add_page_break(self.doc)

    def _normalize_enlistment_type(self, et: str) -> str:
        """入伍类型规范化为中文全称"""
        if not et:
            return "在校生入伍"
        if et in ENLISTMENT_TYPE_MAP:
            return ENLISTMENT_TYPE_MAP[et]
        return et

    # 封面

    def _add_cover(self):
        """封面：黑体二号标题 + 6 行下划线信息"""
        for _ in range(3):
            self.doc.add_paragraph()

        title = "大学生应征入伍申请书"
        add_paragraph_with_format(
            self.doc, title,
            font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_before=12, space_after=12,
        )

        subtitle = f"（{self._get('apply_year', default='2025')} 年度）"
        add_paragraph_with_format(
            self.doc, subtitle,
            font_name=FONT_HEI, font_size=SIZE_SAN, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_after=24,
        )

        for _ in range(3):
            self.doc.add_paragraph()

        et_code = self._get("enlistment_type", default="在校生入伍")
        et_name = self._normalize_enlistment_type(et_code)
        info_items = [
            ("申请人姓名", self._get("applicant_name", default="")),
            ("所在学校", self._get("applicant_school", default="")),
            ("专业年级", f"{self._get('applicant_major', default='')} "
                       f"{self._get('applicant_grade', default='')}"),
            ("入伍类型", et_name),
            ("申请日期", f"{self._get('apply_year', default='2025')} 年      月      日"),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI,
                         font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True

        self.add_page_break()

    # 一、个人基本情况

    def _add_basic_info(self):
        """一、个人基本情况（200-300 字段落 + 基本信息表 + 体检政审表）"""
        self.add_h1("一、个人基本情况")

        # 段落（200-300 字）
        name = self._get("applicant_name", default="XXX")
        gender = self._get("applicant_gender", default="X")
        ethnicity = self._get("applicant_ethnicity", default="X族")
        birth = self._get("applicant_birth", default="XXXX 年 X 月")
        political = self._get("applicant_political", default="共青团员")
        native = self._get("applicant_native_place", default="XX 省 XX 市")
        school = self._get("applicant_school", default="XX 大学")
        major = self._get("applicant_major", default="XX 专业")
        grade = self._get("applicant_grade", default="XXXX 级")
        degree = self._get("applicant_degree", default="本科在读")
        graduation = self._get("applicant_graduation", default="XXXX 年 X 月")
        rank = self._get("academic_rank", default="专业前 XX%")
        honors = self._get("honors", default=[])
        honors_str = "、".join(honors) if isinstance(honors, list) and honors \
            else "校级优秀学生"
        et_code = self._get("enlistment_type", default="在校生入伍")
        et_name = self._normalize_enlistment_type(et_code)
        apply_year = self._get("apply_year", default="2025")
        physical_exam = self._get("physical_exam", default="（请填写体检结论）")
        political_review = self._get("political_review",
                                     default="（请填写政审结论）")

        intro = (
            f"本人 {name}，性别 {gender}，民族 {ethnicity}，{birth} 出生于 "
            f"{native}，政治面貌 {political}。现为 {school} {major} {grade} "
            f"{degree}，（预计）{graduation} 毕业。在校期间学业排名 {rank}，"
            f"曾获 {honors_str} 等荣誉。本人志愿报名参加 {apply_year} 年大学生"
            f"应征入伍，申请 {et_name}。本人已参加全国征兵体检，结论为"
            f"{physical_exam}；参加政治考核，结论为{political_review}。"
        )
        self.add_para(intro)

        # 基本信息表（2 列展示）
        certs = self._get("certificates", default=[])
        certs_str = "、".join(certs) if isinstance(certs, list) and certs else "—"
        rows_simple = [
            ["姓名", name],
            ["性别", gender],
            ["民族", ethnicity],
            ["政治面貌", political],
            ["出生年月", birth],
            ["籍贯", native],
            ["学校", school],
            ["学历层次", degree],
            ["专业年级", f"{major} {grade}"],
            ["学号", self._get("applicant_id", default="")],
            ["毕业时间", graduation],
            ["联系电话", self._get("applicant_phone", default="")],
            ["电子邮箱", self._get("applicant_email", default="—")],
            ["学业排名", rank],
            ["GPA", self._get("gpa", default="—")],
            ["资格证书", certs_str],
            ["入伍类型", et_name],
            ["申请年份", f"{apply_year} 年"],
        ]
        self.add_table(["项目", "内容"], rows_simple, col_widths=[4.5, 11.5])

        # 体检政审表
        self.add_h2("（附）体检与政审结论")
        age = self._get("age", default="—")
        height = self._get("height_cm", default="—")
        weight = self._get("weight_kg", default="—")
        left_v = self._get("left_eye_vision", default="—")
        right_v = self._get("right_eye_vision", default="—")
        physical_test = self._get("physical_test", default="—")
        rows_medical = [
            ["年龄", f"{age} 周岁"],
            ["身高", f"{height} cm"],
            ["体重", f"{weight} kg"],
            ["左眼裸眼视力", left_v],
            ["右眼裸眼视力", right_v],
            ["体测达标", physical_test],
            ["体检结论", physical_exam],
            ["政审结论", political_review],
        ]
        self.add_table(["项目", "内容"], rows_medical, col_widths=[4.5, 11.5])

    # 二、入伍动机

    def _add_motivation(self):
        """二、入伍动机（500-700 字，4 段）

        motivation 列表顺序约定：
        [0] 爱国情怀段（含征兵文件引用）
        [1] 责任担当段（含具体触发事件）
        [2] 个人成长段（含部队对个人的塑造）
        [3] 职业规划段（含服役期满规划 + 入伍优待政策）
        """
        self.add_h1("二、入伍动机")

        # 引用征兵政策文件
        docs = self._get("military_doc_cited", default=[])
        if docs and isinstance(docs, list) and not isinstance(docs, str):
            docs_list = list(docs)
        else:
            docs_list = []

        # 4 段入伍动机
        motivation = self._get("motivation", default=[])
        if isinstance(motivation, str):
            motivation = [motivation]

        sub_titles = ["（一）爱国情怀", "（二）责任担当",
                      "（三）个人成长", "（四）职业规划"]

        if motivation and isinstance(motivation, list) and len(motivation) >= 4:
            for i, para in enumerate(motivation[:4]):
                self.add_h2(sub_titles[i])
                self.add_para(para)
            if docs_list:
                self.add_h2("（附）引用的征兵政策文件")
                self.add_para(
                    "本人认真学习了以下征兵政策文件，深入了解应征入伍的"
                    "兵役义务、征集程序、体检标准、政治考核与优待政策："
                )
                for d in docs_list:
                    self.add_para(f"• {d}", indent=False)
        elif motivation and isinstance(motivation, list) and len(motivation) >= 1:
            for i, para in enumerate(motivation):
                idx = i if i < 4 else 4
                title = sub_titles[idx] if idx < 4 else f"（{i+1}）"
                self.add_h2(title)
                self.add_para(para)
            if docs_list:
                self.add_h2("（附）引用的征兵政策文件")
                self.add_para("本人认真学习了以下征兵政策文件：", indent=False)
                for d in docs_list:
                    self.add_para(f"• {d}", indent=False)
        else:
            self.add_h2("（一）爱国情怀")
            self.add_para(
                "（请填写爱国情怀段，150 字左右。结构：必引 1-2 项征兵政策"
                "文件 + 国防意义 + 对军队的认识。示例：本人认真学习了"
                "《中华人民共和国兵役法》《征兵工作条例》等法律法规与"
                "《2025 年大学生应征入伍优惠政策》。本人深刻认识到，依照"
                "法律服兵役和参加民兵组织是中华人民共和国公民的光荣义务，"
                "是保家卫国的神圣职责。本人愿以青春之我，护卫盛世之中华。）"
            )
            self.add_h2("（二）责任担当")
            self.add_para(
                "（请填写责任担当段，150 字左右。结构：必含 1 个具体触发"
                "事件 + 价值观形成 + 与参军选择的连接。示例：触发本人应征"
                "入伍决心的，是 2020 年 6 月中印加勒万河谷冲突中陈红军、"
                "陈祥榕、肖思远、王焯冉等烈士的英勇事迹。陈祥榕烈士牺牲"
                "时年仅 19 岁，他在战斗前写下的'清澈的爱，只为中国'八字"
                "战斗宣言深深震撼了本人。'国家兴亡，匹夫有责'，本人作为"
                "新时代大学生，理应接过烈士的钢枪，到祖国最需要的地方去。）"
            )
            self.add_h2("（三）个人成长")
            self.add_para(
                "（请填写个人成长段，150 字左右。结构：必含部队对个人品格"
                "的塑造 + 个人能力的提升 + 人生经历的丰富。示例：本人理解，"
                "部队是一座大熔炉、一所大学校。两年义务兵役将塑造本人严明"
                "的纪律性、坚韧的意志力、强烈的责任感和过硬的抗压能力。"
                "本人深知'当兵后悔两年，不当兵后悔一辈子'这句话的分量"
                "——两年军旅生涯将成为本人一生最宝贵的财富。）"
            )
            self.add_h2("（四）职业规划")
            self.add_para(
                "（请填写职业规划段，150 字左右。结构：必含服役期满规划 + "
                "入伍优待政策利用 + 长期职业愿景。示例：本人服役期满后的"
                "初步规划是退役复学完成本科学业，随后报考硕士研究生（享受"
                "退役大学生士兵专项研究生招生计划与考研加 10 分政策）。"
                "本人在校期间学费将获得每年最高 16000 元的补偿代偿；退役"
                "复学后可申请转专业。本人将以军旅生涯为起点，将军人精神"
                "带入未来职业。）"
            )

    # 三、个人素质

    def _add_abilities(self):
        """三、个人素质（400-500 字，3 子段）"""
        self.add_h1("三、个人素质")

        abilities = self._get("abilities", default=[])
        if isinstance(abilities, str):
            abilities = [abilities]
        sub_titles = ["（一）政治素质", "（二）专业能力", "（三）身体素质"]
        if abilities and isinstance(abilities, list):
            for i, para in enumerate(abilities):
                idx = i if i < 3 else 3
                title = sub_titles[idx] if idx < 3 else f"（{i+1}）"
                self.add_h2(title)
                self.add_para(para)
        else:
            self.add_h2("（一）政治素质")
            self.add_para(
                "（请填写政治素质子段，130-180 字。党员必写：入党时间 + "
                "入党介绍人 + 党内职务 + 理论学习情况；团员必写：入团时间 + "
                "团内职务 + 青年大学习完成期数 + 推优入党情况；群众必写："
                "政治学习情况 + 加入党组织意愿。）"
            )
            self.add_h2("（二）专业能力")
            self.add_para(
                "（请填写专业能力子段，130-180 字。必含：GPA + 排名 + 核心"
                "课程 + 资格证书 + 与兵种匹配度。示例：本人在校期间学业成绩"
                "GPA 3.6/4.0，排名专业前 15%。已通过 CET-6、计算机二级。"
                "所学专业计算机科学与技术与通信兵、电子对抗兵、网络空间部队"
                "等兵种高度匹配，可发挥专业特长服务部队信息化建设。）"
            )
            self.add_h2("（三）身体素质")
            self.add_para(
                "（请填写身体素质子段，130-180 字。必含：体检结论 + 体测"
                "达标 + 体能成绩 + 心理测评。示例：本人 2025 年 2 月参加"
                "成都市武侯区人民医院征兵体检，结论为合格（身高 175cm，"
                "体重 68kg，左眼裸眼视力 4.8，右眼裸眼视力 4.9）。国家学生"
                "体质健康标准良好。1000 米跑 3 分 45 秒，引体向上 12 个，"
                "立定跳远 2.45 米，体能基础良好。）"
            )

    # 四、部队适应能力

    def _add_military_adaptability(self):
        """四、部队适应能力（300-400 字，3 子段）

        这是入伍申请特有段，重点考察部队适应性。
        """
        self.add_h1("四、部队适应能力")

        adapt = self._get("military_adaptability", default=[])
        if isinstance(adapt, str):
            adapt = [adapt]
        sub_titles = ["（一）纪律性",
                      "（二）集体生活适应能力",
                      "（三）抗压能力"]
        if adapt and isinstance(adapt, list):
            for i, para in enumerate(adapt):
                idx = i if i < 3 else 3
                title = sub_titles[idx] if idx < 3 else f"（{i+1}）"
                self.add_h2(title)
                self.add_para(para)
        else:
            self.add_h2("（一）纪律性")
            self.add_para(
                "（请填写纪律性子段，100-150 字。必含：在校纪律表现 + 对"
                "部队纪律的认识 + 适应部队纪律的准备。示例：本人在校期间"
                "严格遵守校规校纪，无任何违纪记录，出勤率 100%。本人理解"
                "部队纪律是战斗力的根本保证，已通过早起作息训练（连续 100"
                "天 6 点起床）为适应部队纪律做好充分准备。）"
            )
            self.add_h2("（二）集体生活适应能力")
            self.add_para(
                "（请填写集体生活适应能力子段，100-150 字。必含：集体生活"
                "经历 + 协作能力 + 适应部队集体生活的准备。示例：本人大学"
                "四年住宿生活和谐，曾担任宿舍长 2 年。在校 ACM 算法社团"
                "担任技术副部长 1 年，与不同性格同学协作完成 3 个开源项目。"
                "本人将与战友同吃同住同训练，做到团结友爱、互帮互助。）"
            )
            self.add_h2("（三）抗压能力")
            self.add_para(
                "（请填写抗压能力子段，100-150 字。必含：抗压经历 + 心理"
                "测评 + 适应部队高强度训练的准备。示例：本人在校期间曾同时"
                "承担学业、学生工作、志愿服务三线任务，均圆满完成。2024 年"
                "9 月校级心理测评结论为良好。本人已通过长跑训练（累计 500 "
                "公里）增强体能储备，能够适应部队高强度训练。）"
            )

    # 五、家人态度与决心

    def _add_family_attitude(self):
        """五、家人态度与决心（200-300 字，2 子段）"""
        self.add_h1("五、家人态度与决心")

        family = self._get("family_attitude", default=[])
        if isinstance(family, str):
            family = [family]
        sub_titles = ["（一）家人态度与沟通过程", "（二）决心表态"]
        if family and isinstance(family, list):
            for i, para in enumerate(family):
                idx = i if i < 2 else 2
                title = sub_titles[idx] if idx < 2 else f"（{i+1}）"
                self.add_h2(title)
                self.add_para(para)
        else:
            self.add_h2("（一）家人态度与沟通过程")
            self.add_para(
                "（请填写家人态度子段，100-180 字。必含：家人总体态度 + "
                "沟通过程（何时告知、如何告知）+ 家人顾虑及化解 + 最终态度"
                "+ 是否签署家长知情同意书。）"
            )
            self.add_h2("（二）决心表态")
            self.add_para(
                "（请填写决心表态子段，100-180 字。必含：再次表态决心 + "
                "服役期承诺（服从命令、严守纪律、刻苦训练）+ 长期愿景"
                "（退伍不褪色）。示例：本人郑重表态：志愿应征入伍服 2 年"
                "义务兵役，服役期间服从命令、听从指挥、严守纪律、刻苦训练。"
                "本人将以'清澈的爱，只为中国'为座右铭，到祖国最需要的地方"
                "去，把青春奉献给国防事业。）"
            )

    # 签字栏

    def _add_signature_section(self):
        """签字栏：申请人 + 学校武装部 + 县级征兵办"""
        self.add_h1("六、申请人签字与审核意见")

        # 申请人签字
        self.add_h2("（一）申请人签字")
        for _ in range(3):
            self.doc.add_paragraph()
        self.add_para(
            "申请人签字：____________________    "
            "日期：______年____月____日",
            indent=False,
        )

        # 学校武装部审核意见
        self.add_h2("（二）学校武装部审核意见")
        for _ in range(5):
            self.doc.add_paragraph()
        self.add_para(
            "学校武装部盖章：____________________    "
            "负责人签字：____________    "
            "日期：______年____月____日",
            indent=False,
        )

        # 县级征兵办审核意见
        self.add_h2("（三）县级征兵办审核意见")
        for _ in range(5):
            self.doc.add_paragraph()
        self.add_para(
            "县级征兵办盖章：____________________    "
            "负责人签字：____________    "
            "日期：______年____月____日",
            indent=False,
        )

    # 主构建方法

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 5 段结构 + 签字栏，生成 docx

        Args:
            data: 申请书字段字典
            output_path: 输出 docx 路径

        Returns:
            实际保存路径
        """
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()

            self._add_cover()
            self._add_basic_info()
            self._add_motivation()
            self._add_abilities()
            self._add_military_adaptability()
            self._add_family_attitude()
            if self._get("include_signature", default=True):
                self._add_signature_section()

            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 应征入伍申请书已生成：{output_path}")
        return str(output_path)

    # 数据校验

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）

        校验分五类：P0 必填、入伍类型匹配、体检硬条件、家人态度、字数。
        """
        warnings = []
        p0_fields = [
            ("applicant_name", "申请人姓名"),
            ("applicant_political", "政治面貌"),
            ("applicant_degree", "学历层次"),
            ("applicant_school", "学校"),
            ("applicant_major", "专业"),
            ("applicant_grade", "年级"),
            ("apply_year", "申请年份"),
            ("age", "年龄"),
            ("height_cm", "身高"),
            ("weight_kg", "体重"),
            ("left_eye_vision", "左眼视力"),
            ("right_eye_vision", "右眼视力"),
            ("physical_exam", "体检结论"),
            ("political_review", "政审结论"),
        ]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        # 入伍类型校验
        et = self._get("enlistment_type", default="")
        if not et:
            warnings.append("缺少 入伍类型（enlistment_type），必填 在校生入伍/毕业生入伍 之一")
        elif et not in ENLISTMENT_TYPE_MAP:
            warnings.append(f"入伍类型取值 '{et}' 不在合法集合，应为 {list(ENLISTMENT_TYPE_MAP.keys())}")

        # 年龄与入伍类型匹配校验
        age = safe_int(self._get("age", default=0), default=0)
        et_norm = self._normalize_enlistment_type(et)
        gender = self._get("applicant_gender", default="")
        if age > 0:
            if et_norm == "在校生入伍" and not (18 <= age <= 22):
                warnings.append(f"在校生入伍年龄 {age} 不在 18-22 区间，不符合在校生入伍年龄硬条件")
            if et_norm == "毕业生入伍":
                if gender == "男" and not (18 <= age <= 24):
                    warnings.append(f"毕业生入伍男兵年龄 {age} 不在 18-24 区间，不符合毕业生入伍年龄硬条件")
                if gender == "女" and not (18 <= age <= 22):
                    warnings.append(f"毕业生入伍女兵年龄 {age} 不在 18-22 区间，不符合毕业生入伍年龄硬条件")

        # 体检硬条件校验
        height = safe_float(self._get("height_cm", default=0), default=0)
        if height > 0:
            if gender == "男" and height < 160:
                warnings.append(f"男兵身高 {height} cm 低于 160cm，不符合体检硬条件")
            if gender == "女" and height < 158:
                warnings.append(f"女兵身高 {height} cm 低于 158cm，不符合体检硬条件")

        left_v = vision_to_score(self._get("left_eye_vision", default=""))
        right_v = vision_to_score(self._get("right_eye_vision", default=""))
        if 0 < left_v < 4.5:
            warnings.append(f"左眼裸眼视力 {left_v} 低于陆勤标准 4.5，不符合体检硬条件（特种兵要求更高）")
        if 0 < right_v < 4.6:
            warnings.append(f"右眼裸眼视力 {right_v} 低于陆勤标准 4.6，不符合体检硬条件（特种兵要求更高）")

        # 体检政审结论关键词校验
        pe_text = str(self._get("physical_exam", default=""))
        pr_text = str(self._get("political_review", default=""))
        if pe_text and "合格" not in pe_text and "不合格" not in pe_text:
            warnings.append("体检结论未含'合格'/'不合格'关键词，建议明确为征兵办指定医院结论")
        if pr_text and "合格" not in pr_text and "不合格" not in pr_text:
            warnings.append("政审结论未含'合格'/'不合格'关键词，建议明确为县级征兵办政治考核组结论")

        # 段落数量校验
        motivation = self._get("motivation", default=[])
        if isinstance(motivation, list) and len(motivation) < 4:
            warnings.append(f"入伍动机段 {len(motivation)} 段，应至少 4 段（爱国情怀+责任担当+个人成长+职业规划）")

        abilities = self._get("abilities", default=[])
        if isinstance(abilities, list) and len(abilities) < 3:
            warnings.append(f"个人素质段 {len(abilities)} 子段，应至少 3 子段（政治+专业+身体）")

        adapt = self._get("military_adaptability", default=[])
        if isinstance(adapt, list) and len(adapt) < 3:
            warnings.append(f"部队适应能力段 {len(adapt)} 子段，应至少 3 子段（纪律+集体+抗压）")

        family = self._get("family_attitude", default=[])
        if isinstance(family, list) and len(family) < 2:
            warnings.append(f"家人态度段 {len(family)} 子段，应至少 2 子段（态度+决心）")

        # 引用文件校验
        docs = self._get("military_doc_cited", default=[])
        if not docs or (isinstance(docs, list) and len(docs) == 0):
            warnings.append("缺少 引用征兵政策文件（military_doc_cited），入伍动机段必引至少 1 项")

        # 家人态度校验
        family_text = "\n".join(family) if isinstance(family, list) else str(family)
        if "沟通" not in family_text and "告知" not in family_text and "商量" not in family_text:
            warnings.append("家人态度段未提及沟通过程（沟通/告知/商量），评审会扣大分")

        # 字数校验
        total_chars = 0
        for k in ("motivation", "abilities", "military_adaptability",
                  "family_attitude"):
            v = self._get(k, default=[])
            if isinstance(v, list):
                for s in v:
                    total_chars += count_chinese_chars(s)
        if total_chars < 1200:
            warnings.append(f"5 段合计中文字符约 {total_chars}，偏少（建议 1500-2500 字）")
        elif total_chars > 2800:
            warnings.append(f"5 段合计中文字符约 {total_chars}，偏多（建议 1500-2500 字）")

        # 政治面貌校验
        political = self._get("applicant_political", default="")
        if political and political not in ("中共党员", "共青团员", "群众", "中共预备党员"):
            warnings.append(f"政治面貌取值 '{political}' 不规范，应为 中共党员/共青团员/群众 之一")

        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        self.warnings = warnings
        return warnings


# ============================================================
# 默认示例数据
# ============================================================

DEFAULT_DATA = {
    "applicant_name": "张三",
    "applicant_gender": "男",
    "applicant_ethnicity": "汉族",
    "applicant_birth": "2003 年 5 月",
    "applicant_political": "中共党员",
    "applicant_native_place": "四川省成都市武侯区",
    "applicant_id": "2022012345",
    "applicant_major": "计算机科学与技术",
    "applicant_grade": "2022 级",
    "applicant_school": "XX 大学计算机学院",
    "applicant_degree": "本科在读",
    "applicant_graduation": "2026 年 6 月",
    "applicant_phone": "138XXXXXXXX",
    "applicant_email": "zhangsan@example.com",

    "apply_year": "2025",
    "enlistment_type": "在校生入伍",
    "age": 22,
    "height_cm": 175,
    "weight_kg": 68,
    "left_eye_vision": "4.8",
    "right_eye_vision": "4.9",
    "physical_exam": "2025 年 2 月成都市武侯区人民医院征兵体检合格（身高 175cm，体重 68kg，左眼裸眼视力 4.8，右眼裸眼视力 4.9，无重大疾病史）",
    "political_review": "2025 年 2 月成都市武侯区公安分局政治考核合格（本人及家庭主要成员无违法犯罪记录）",

    "academic_rank": "专业前 15%（8/52）",
    "gpa": "3.6/4.0",
    "core_courses": "数据结构 92 / 操作系统 89 / 计算机网络 95 / 数据库 88",
    "certificates": [
        "CET-6",
        "计算机二级（C 语言）",
        "全国计算机等级三级（网络技术）",
    ],
    "honors": [
        "2024 年国家奖学金",
        "2023 年校级一等奖学金",
        "校级优秀学生干部",
    ],
    "publications": [],
    "competitions": [
        "2024 年挑战杯省级一等奖",
        "2023 年蓝桥杯省级二等奖",
    ],

    "political_detail": "2023 年 6 月加入中国共产党，入党介绍人为辅导员李老师与班主任王老师。在校期间认真学习党的理论，研读《习近平与大学生朋友们》《习近平的七年知青岁月》等著作，完成'青年大学习'52 期，撰写思想汇报 6 篇。曾任院学生党支部宣传委员，组织主题党日活动 8 次。",

    "physical_test": "国家学生体质健康标准良好",
    "physical_fitness": {
        "run_1000m": "3 分 45 秒",
        "pull_up": "12 个",
        "standing_long_jump": "2.45 米",
    },

    "motivation": [
        "（爱国情怀）本人认真学习了《中华人民共和国兵役法》《征兵工作条例》《应征公民体格检查标准》等法律法规与《2025 年大学生应征入伍优惠政策》。本人深刻认识到，依照法律服兵役和参加民兵组织是中华人民共和国公民的光荣义务，是保家卫国的神圣职责。本人理解应征入伍是党中央、国务院、中央军委推进国防和军队现代化建设的重要举措，是优化兵员结构、提升部队战斗力的重要途径。本人愿以青春之我，护卫盛世之中华。",
        "（责任担当）触发本人应征入伍决心的，是 2020 年 6 月中印加勒万河谷冲突中陈红军、陈祥榕、肖思远、王焯冉等烈士的英勇事迹。陈祥榕烈士牺牲时年仅 19 岁，与本人年纪相仿，他在战斗前写下的'清澈的爱，只为中国'八字战斗宣言深深震撼了本人。'国家兴亡，匹夫有责'，本人作为新时代大学生，理应接过烈士的钢枪，到祖国最需要的地方去，到党和人民最需要的地方去。",
        "（个人成长）本人理解，部队是一座大熔炉、一所大学校。两年义务兵役将塑造本人严明的纪律性、坚韧的意志力、强烈的责任感和过硬的抗压能力。本人将通过新兵训练、专业训练、战术训练掌握军事技能，培养团队协作与领导力。本人深知'当兵后悔两年，不当兵后悔一辈子'这句话的分量——两年军旅生涯将成为本人一生最宝贵的财富，远胜任何奖学金与实习经历。",
        "（职业规划）本人服役期满后的初步规划是退役复学完成本科学业，随后报考硕士研究生（享受退役大学生士兵专项研究生招生计划与考研加 10 分政策）。本人已了解大学生入伍优待政策：本人在校期间学费将获得每年最高 16000 元的补偿代偿；退役复学后可申请转专业；享受专升本免试、考研加分、专项研究生计划、考公定向等政策。本人将以军旅生涯为起点，将军人精神带入未来职业，无论身在何处，都以'退伍不褪色'的标准要求自己。",
    ],

    "abilities": [
        "（政治素质）本人政治面貌为中共党员，2023 年 6 月入党，入党介绍人为辅导员李老师与班主任王老师。在校期间认真学习党的理论，研读《习近平与大学生朋友们》《习近平的七年知青岁月》等著作，完成'青年大学习'52 期，撰写思想汇报 6 篇。曾任院学生党支部宣传委员，组织主题党日活动 8 次，培养了对党忠诚、服务人民的信念。本人政治立场坚定，拥护党的领导，自觉抵制各种错误思潮。",
        "（专业能力）本人在校期间学业成绩 GPA 3.6/4.0，排名专业前 15%（8/52）。已通过 CET-6、计算机二级（C 语言）、全国计算机等级三级（网络技术），具备扎实的计算机理论基础与编程能力。所学专业计算机科学与技术与通信兵、电子对抗兵、网络空间部队等兵种高度匹配，可发挥专业特长服务部队信息化建设。曾获 2024 年国家奖学金、2024 年挑战杯省级一等奖，专业能力与所报技术兵种匹配度高。",
        "（身体素质）本人 2025 年 2 月参加成都市武侯区人民医院征兵体检，结论为合格（身高 175cm，体重 68kg，左眼裸眼视力 4.8，右眼裸眼视力 4.9，无重大疾病史）。国家学生体质健康标准良好。1000 米跑 3 分 45 秒，引体向上 12 个，立定跳远 2.45 米，体能基础良好，能适应部队高强度训练。校级心理测评结论良好。",
    ],

    "military_adaptability": [
        "（纪律性）本人在校期间严格遵守校规校纪，无任何违纪记录，出勤率 100%，诚信考试无作弊。本人理解部队纪律是战斗力的根本保证，包括条令条例、内务管理、请示报告等。本人已通过早起作息训练（连续 100 天 6 点起床）、自律打卡（连续 200 天运动）为适应部队纪律做好充分准备，能够在服从命令、听从指挥方面做到不打折扣、不搞变通。",
        "（集体生活适应能力）本人大学三年住宿生活和谐，曾担任宿舍长 2 年，组织宿舍卫生评比、节庆活动等。在校 ACM 算法社团担任技术副部长 1 年，与不同性格同学协作完成 3 个开源项目。本人理解部队是高度集体化的组织，本人将与战友同吃同住同训练，做到团结友爱、互帮互助、共同进步。",
        "（抗压能力）本人在校期间曾同时承担学业、学生工作、志愿服务三线任务（期末备考 + 学生会换届 + 凉山支教），均圆满完成，体现较强抗压能力。本人 2024 年 9 月参加校级心理测评，结论为良好。本人已通过长跑训练（累计 500 公里）、户外拉练（318 国道骑行 21 天）增强体能储备，通过心理建设提升抗压能力，能够适应部队高强度训练与突发任务。",
    ],

    "family_attitude": [
        "（家人态度与沟通过程）本人于 2025 年 1 月向家人正式告知应征入伍的决定。父母初始反应是担忧本人安全与未来出路。经过 3 次深入沟通，家人从担忧转为支持。父亲表示：'保家卫国是男子汉的本分，我们支持你的选择，但要照顾好自己。'母亲表示：'家里有我们，你安心服役。'家人最终态度为支持，已签署家长知情同意书。具体顾虑化解：（1）父亲担忧本人安全，本人向父亲展示了部队保险条款与军人优待政策后，父亲态度转为支持；（2）母亲担忧本人未来出路，本人向母亲说明入伍保留学籍、学费补偿、考研加分等政策，母亲表示理解。",
        "（决心表态）本人郑重表态：志愿应征入伍服 2 年义务兵役，服役期间服从命令、听从指挥、严守纪律、刻苦训练，以优异成绩完成各项任务。本人将以'清澈的爱，只为中国'为座右铭，到祖国最需要的地方去，把青春奉献给国防事业。退伍后无论从事何种职业，都将以'退伍不褪色'的标准要求自己，将军人精神融入血脉。",
    ],

    "version": "standard",
    "include_signature": True,
    "military_doc_cited": [
        "《中华人民共和国兵役法》（2021 年修订）",
        "《征兵工作条例》（2023 年修订）",
        "《2025 年大学生应征入伍优惠政策》",
    ],
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="大学生应征入伍申请书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n"
            "JSON 字段定义详见 SKILL.md 第十一章。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档"
              "（standard 档 2000 字，2025 年上半年应征）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
