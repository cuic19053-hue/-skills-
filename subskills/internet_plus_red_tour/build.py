#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""互联网+大学生创新创业大赛 红色之旅赛道（青年红色筑梦之旅）商业计划书 docx 生成器

格式标准：A4，页边距上下 2.54cm 左右 2.5cm；正文宋体小四 1.5 倍行距首行缩进 2 字符；
一级标题黑体三号居中；二级黑体小三左对齐；三级宋体四号加粗；表格宋体五号居中。

11 栏目：封面 / 执行摘要 / 红色基因传承与项目背景 / 市场分析与需求验证 /
产品服务介绍 / 商业模式与可持续造血 / 运营现状与红色帮扶成效 / 团队介绍 /
财务预测 / 风险与对策 / 个人成长与团队协作。

红旅赛道特色：红色基因 4 维度（党史学习/革命精神/红色文化/典型案例）+ 乡村振兴 5 维度
（产业/人才/文化/生态/组织振兴）+ 帮扶成效 6 项量化数据 + 可持续造血模式 + 6 类风险。

使用：python build.py --data data.json --out output.docx ； python build.py --demo --out demo.docx
JSON 字段详见 SKILL.md 第九章。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22); SIZE_XIAO_ER = Pt(18); SIZE_SAN = Pt(16); SIZE_XIAO_SAN = Pt(15)
SIZE_SI = Pt(14); SIZE_XIAO_SI = Pt(12); SIZE_WU = Pt(10.5); SIZE_XIAO_WU = Pt(9)

PAGE_WIDTH_CM = 21.0; PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54; MARGIN_LEFT_RIGHT_CM = 2.5

# 红旅赛道 4 大主题
RED_THEMES = ["乡村振兴", "革命老区", "民族团结", "西部开发"]

# 字数版本
WORD_COUNT_VERSIONS = {
    "brief": {"name": "8000 字版本", "min": 7000, "max": 9000,
              "desc": "校赛 / 路演版 / 初筛版"},
    "standard": {"name": "12000 字版本", "min": 11000, "max": 13000,
                 "desc": "省赛 / 国赛网评 / 标准提交版"},
    "enhanced": {"name": "15000 字版本", "min": 14000, "max": 16000,
                 "desc": "国赛现场赛 / 总决赛 / 完整版"},
}


# ============================================================
# 工具函数（中英文同步设置 eastAsia/ascii/hAnsi）
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_SAN,
                                    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    first_line_indent=False, space_before=12, space_after=12)


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SAN,
                                    bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                    first_line_indent=False, space_before=6, space_after=6)


def add_heading_level3(doc, text: str):
    """三级标题：宋体四号加粗，左对齐"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_SI,
                                    bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                    first_line_indent=False, space_before=6, space_after=3)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
                                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                    first_line_indent=indent, line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None,
                        last_col_left: bool = True):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG, font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if (j == len(row) - 1 and last_col_left) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], val, font_name=FONT_SONG, font_size=SIZE_WU,
                          bold=False, alignment=align)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_canvas_9grid(doc, canvas_items: List[Dict[str, str]]):
    """商业模式画布 9 宫格表格（3 行 × 3 列，每格含要素名+内容）"""
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, item in enumerate(canvas_items[:9]):
        r, c = divmod(idx, 3)
        cell = table.rows[r].cells[c]
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(item.get("element", ""))
        set_run_font(run1, font_name=FONT_HEI, font_size=SIZE_XIAO_WU, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.line_spacing = 1.2
        run2 = p2.add_run(item.get("content", ""))
        set_run_font(run2, font_name=FONT_SONG, font_size=SIZE_XIAO_WU)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(5.3)
    doc.add_paragraph()


def add_red_gene_table(doc, gene_items: List[Dict[str, str]]):
    """红色基因 4 维度表（4 行：维度名 + 内容摘要）"""
    if not gene_items:
        return
    headers = ["维度", "内容摘要"]
    rows = [[g.get("dimension", ""), g.get("summary", "")]
            for g in gene_items if isinstance(g, dict)]
    add_table_from_data(doc, headers, rows, col_widths=[3.0, 13.0], last_col_left=True)


def add_rural_revival_table(doc, revival_items: List[Dict[str, str]]):
    """乡村振兴 5 维度表（5 行：维度名 + 内容摘要）"""
    if not revival_items:
        return
    headers = ["振兴维度", "对应内容"]
    rows = [[r.get("dimension", ""), r.get("content", "")]
            for r in revival_items if isinstance(r, dict)]
    add_table_from_data(doc, headers, rows, col_widths=[3.0, 13.0], last_col_left=True)


def add_help_effect_table(doc, service: Dict[str, Any]):
    """红色帮扶成效 6 项量化数据表"""
    if not service:
        return
    headers = ["指标", "数值"]
    rows = [
        ["服务覆盖县数", str(service.get("service_county", ""))],
        ["服务覆盖乡镇数", str(service.get("service_township", ""))],
        ["服务覆盖行政村数", str(service.get("service_village", ""))],
        ["带动农户数", str(service.get("farmer_count", "")) + " 户"],
        ["户均增收金额/年", str(service.get("income_increase", ""))],
        ["持续服务月数", str(service.get("coverage_month", "")) + " 个月"],
        ["合作社数量", str(service.get("coop_count", "")) + " 个"],
        ["党支部结对数", str(service.get("party_branch_pair", "")) + " 个"],
    ]
    add_table_from_data(doc, headers, rows, col_widths=[5.0, 11.0], last_col_left=False)


def add_field_visit_table(doc, visits: List[Dict[str, str]]):
    """老区行/乡村行调研记录表"""
    if not visits:
        return
    headers = ["时间", "地点", "天数", "目的"]
    rows = [[v.get("time", ""), v.get("place", ""),
             str(v.get("days", "")) + " 天", v.get("purpose", "")]
            for v in visits if isinstance(v, dict)]
    add_table_from_data(doc, headers, rows, col_widths=[2.5, 4.0, 2.0, 7.5], last_col_left=True)


def add_typical_case_table(doc, cases: List[Dict[str, str]]):
    """典型农户帮扶案例表"""
    if not cases:
        return
    headers = ["农户姓名", "所在村组", "帮扶前收入", "帮扶后收入", "持续时长"]
    rows = [[c.get("name", ""), c.get("village", ""), c.get("before", ""),
             c.get("after", ""), c.get("duration", "")]
            for c in cases if isinstance(c, dict)]
    add_table_from_data(doc, headers, rows, col_widths=[2.2, 3.0, 3.0, 3.0, 3.8], last_col_left=False)


def add_milestone_table(doc, milestones: List[Dict[str, str]]):
    """关键里程碑表（3 列：时间/事件/数据）"""
    if not milestones:
        return
    headers = ["时间", "里程碑事件", "关键数据"]
    rows = [[m.get("time", ""), m.get("event", ""), m.get("data", "")]
            for m in milestones if isinstance(m, dict)]
    add_table_from_data(doc, headers, rows, col_widths=[3.0, 6.5, 6.0], last_col_left=False)


def add_kpi_table(doc, kpi_rows: List[Dict[str, str]]):
    """运营现状 KPI 表（5 列：指标/当前值/Year1/Year2/Year3）"""
    if not kpi_rows:
        return
    headers = ["指标", "当前值", "Year1 目标", "Year2 目标", "Year3 目标"]
    rows = [[r.get("metric", ""), str(r.get("current", "")), str(r.get("y1", "")),
             str(r.get("y2", "")), str(r.get("y3", ""))] for r in kpi_rows if isinstance(r, dict)]
    add_table_from_data(doc, headers, rows, col_widths=[4.0, 2.5, 2.8, 2.8, 3.4], last_col_left=False)


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM); section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM); section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM); section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ============================================================
# RedTourDocBuilder 主类
# ============================================================

class RedTourDocBuilder:
    """互联网+大学生创新创业大赛 红色之旅赛道 商业计划书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def add_h1(self, text): return add_heading_level1(self.doc, text)
    def add_h2(self, text): return add_heading_level2(self.doc, text)
    def add_h3(self, text): return add_heading_level3(self.doc, text)
    def add_para(self, text, indent=True): return add_body_paragraph(self.doc, text, indent=indent)
    def add_table(self, headers, rows, col_widths=None, last_col_left=True):
        return add_table_from_data(self.doc, headers, rows, col_widths, last_col_left)
    def add_page_break(self): add_page_break(self.doc)

    # 封面

    def _add_cover(self):
        """封面：大赛名称 + 红旅赛道副标题 + 7 行下划线信息"""
        for _ in range(2):
            self.doc.add_paragraph()
        add_paragraph_with_format(self.doc, '中国国际大学生创新大赛',
                                  font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  first_line_indent=False, space_before=12, space_after=12)
        add_paragraph_with_format(self.doc, "红色之旅赛道商业计划书",
                                  font_name=FONT_HEI, font_size=SIZE_SAN, bold=False,
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  first_line_indent=False, space_after=24)
        add_paragraph_with_format(self.doc, "——青年红色筑梦之旅——",
                                  font_name=FONT_KAI, font_size=SIZE_SI, bold=False,
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  first_line_indent=False, space_after=24)
        for _ in range(2):
            self.doc.add_paragraph()
        red_theme = self._get("red_theme", default="革命老区")
        if red_theme not in RED_THEMES:
            red_theme = "革命老区"
        info_items = [
            ("项目名称", self._get("project_name")),
            ("赛    道", "红色之旅赛道"),
            ("红旅主题", red_theme),
            ("帮扶地区", self._get("rural_county")),
            ("团队名称", self._get("team_name", default=self._get("project_name"))),
            ("负责人", self._get("leader_name")),
            ("学    校", self._get("school")),
            ("日    期", self._get("apply_date")),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0; pf.space_before = Pt(6); pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI, font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True
        self.add_page_break()

    # 一、执行摘要（800~1000 字，8 要素）

    def _add_executive_summary(self):
        """执行摘要：800~1000 字，8 要素齐全（含红色基因 + 帮扶成效）"""
        self.add_h1("一、执行摘要")
        summary = self._get("executive_summary", default="")
        if summary:
            self.add_para(summary)
        else:
            self.add_para(
                "（请填写执行摘要 800~1000 字，必含 8 要素：① 一句话项目（为 [老区/乡村谁] 解决 [什么问题] 通过 [什么方式]，≤30 字）；"
                "② 红色基因传承（项目对接什么红色资源 + 传承什么革命精神）；"
                "③ 帮扶主题与地区（4 大主题之一 + 帮扶县乡村全称）；"
                "④ 市场机会（老区/乡村产业规模 + 数据来源）；"
                "⑤ 核心产品（MVP 主要功能 + 关键指标如用户数/月活/月营收/带动农户数）；"
                "⑥ 商业模式（盈利模式 + 定价 + 农户分红比例 + 财务自给率）；"
                "⑦ 团队亮点（复合背景 + 党员/团员骨干 + 指导教师 + 党支部结对）；"
                "⑧ 帮扶成效（带动农户数 + 户均增收 + 持续月数 + 可持续造血模式）。"
                "评审 30 秒决定深读还是淘汰，是 BP 最重要栏目。）"
            )

    # 二、红色基因传承与项目背景（1200~1500 字，红旅特色章）

    def _add_red_gene(self):
        """红色基因传承与项目背景：4 维度 + 政策导向 + 团队思政"""
        self.add_h1("二、红色基因传承与项目背景")
        gene = self._get("red_gene", default={})
        if not isinstance(gene, dict):
            gene = {}

        self.add_h2("（一）红色基因 4 维度")
        gene_items = [
            ("党史学习", gene.get("red_study", ""),
             "学习内容（党史/革命史）+ 学习方式（理论+实践）+ 学习成效（时长/笔记/政治面貌发展）"),
            ("革命精神", gene.get("revolution_spirit", ""),
             "精神选型（井冈山/延安/西柏坡等）+ 精神内涵理解 + 项目中具体体现"),
            ("红色文化", gene.get("red_culture", ""),
             "红色资源对接 + 红色文化产品开发 + 红色文化传播"),
            ("典型案例", gene.get("typical_case", ""),
             "学习的老区典型（神山村/十八洞村/梁家河等）+ 典型经验提炼 + 项目对标"),
        ]
        for dim, content, hint in gene_items:
            self.add_h3(f"1. {dim}")
            self.add_para(content if content else
                          f"（请填写{dim}维度 300~400 字：{hint}）")
        # 4 维度摘要表
        gene_summary = [
            {"dimension": "党史学习",
             "summary": (gene.get("red_study", "")[:80] + "...") if gene.get("red_study") else "（待填写）"},
            {"dimension": "革命精神",
             "summary": (gene.get("revolution_spirit", "")[:80] + "...") if gene.get("revolution_spirit") else "（待填写）"},
            {"dimension": "红色文化",
             "summary": (gene.get("red_culture", "")[:80] + "...") if gene.get("red_culture") else "（待填写）"},
            {"dimension": "典型案例",
             "summary": (gene.get("typical_case", "")[:80] + "...") if gene.get("typical_case") else "（待填写）"},
        ]
        add_red_gene_table(self.doc, gene_summary)

        self.add_h2("（二）项目立项背景与红色资源对接")
        rural_county = self._get("rural_county")
        self.add_para(
            f"项目立项背景：项目对接 {rural_county if rural_county else '（请填写帮扶县乡村全称）'}，"
            f"红色资源包括（请填写：革命博物馆/红色遗址/老区党支部/红色文创 IP 等）。"
            f"项目将红色资源转化为（请填写：红色品牌包装/红色研学路线/红色文创衍生品等），"
            f"实现红色资源产业化、活态化。"
        )

        self.add_h2("（三）项目立意的政策导向")
        self.add_para(
            "项目严格对齐国家'乡村振兴'战略与革命老区振兴规划，引用《国务院关于新时代支持革命老区振兴发展的意见》"
            "（国发〔2021〕3 号）、《乡村振兴促进法》（2021 年 6 月 1 日施行）、"
            "《中共中央 国务院关于进一步深化农村改革 扎实推进乡村全面振兴的意见》（2025 年中央一号文件）。"
            "项目依托学校'红旅'活动支持，与大学生'三下乡'社会实践衔接，"
            "形成'立德树人 + 乡村振兴 + 红色基因'三位一体的助农创业实践。"
        )

        self.add_h2("（四）团队思政教育与立德树人")
        team = self._get("team_intro", default={})
        if not isinstance(team, dict):
            team = {}
        party_pair = team.get("party_branch_pair", "")
        volunteer_hours = team.get("volunteer_hours", "")
        self.add_para(
            f"团队成员参与红旅活动与党史学习教育情况：{party_pair if party_pair else '（请填写与老区/乡村党支部结对情况）'}。"
            f"团队成员累计志愿服务时长 {volunteer_hours if volunteer_hours else '（请填写）'}，"
            f"含老区支教、农户培训、红色宣讲等。项目过程中团队成员思政提升显著，"
            f"（请填写：新增入党积极分子 X 人 / 发展为预备党员 X 人 / 提交入党申请书 X 人）。"
        )

    # 三、市场分析与需求验证（1500~2000 字）

    def _add_market_analysis(self):
        """市场分析：4 子节 + TAM/SAM/SOM 表 + 竞品对比表 + 双用户画像"""
        self.add_h1("三、市场分析与需求验证")
        market = self._get("market_analysis", default={})
        if not isinstance(market, dict):
            market = {}

        self.add_h2("（一）行业背景与政策导向")
        bg = market.get("industry_background", "")
        self.add_para(bg if bg else
                      "（请填写行业背景 500~600 字：3 句话讲老区/乡村特色产业规模/增速/政策，"
                      "2 句话讲细分场景痛点。必须有权威数据来源：国家统计局/农业农村部/地方统计公报 + "
                      "政策文件名：《国务院关于新时代支持革命老区振兴发展的意见》《乡村振兴促进法》等。）")

        self.add_h2("（二）目标市场（TAM/SAM/SOM）")
        tss = market.get("tam_sam_som", [])
        if tss:
            rows = [[item.get("level", ""), item.get("definition", ""), item.get("scale", "")]
                    for item in tss if isinstance(item, dict)]
            self.add_table(["层级", "定义", "规模"], rows, col_widths=[2.5, 9.0, 4.5], last_col_left=False)
        else:
            self.add_para("（请填写 TAM/SAM/SOM 三级市场测算表，TAM 全国老区/乡村产业规模 / "
                          "SAM 本省/本市同类地区 / SOM 本项目可获取市场。每行含数据来源。）")

        self.add_h2("（三）竞品分析")
        comps = market.get("competitors", [])
        if comps:
            rows = [[c.get("name", ""), c.get("positioning", ""), c.get("users", ""),
                     c.get("advantage", ""), c.get("disadvantage", "")] for c in comps if isinstance(c, dict)]
            self.add_table(["竞品", "定位", "用户规模", "优势", "劣势"], rows,
                           col_widths=[2.2, 2.8, 2.5, 4.0, 4.5])
        else:
            self.add_para("（请填写竞品分析表，3~5 个直接竞品（同类老区助农项目 / 同类农产品电商）+ "
                          "2~3 个间接竞品（传统收购商 / 通用电商平台），5 列对比：名称/定位/用户规模/优势/劣势。）")
        differentiation = market.get("differentiation", "")
        if differentiation:
            self.add_para(f"本项目差异化定位：{differentiation}")

        self.add_h2("（四）用户画像与需求验证（双画像）")
        persona_farmer = market.get("user_persona_farmer", "")
        self.add_h3("1. 老区/乡村农户画像（供给侧）")
        self.add_para(persona_farmer if persona_farmer else
                      "（请填写农户画像 300~400 字，6 要素：基本信息+收入来源+痛点（销路/价格/品质）"
                      "+需求+付费意愿+决策因素。附调研样本量 N≥30 农户访谈。）")
        persona_consumer = market.get("user_persona_consumer", "")
        self.add_h3("2. 城市消费者画像（需求侧）")
        self.add_para(persona_consumer if persona_consumer else
                      "（请填写消费者画像 300~400 字，6 要素：基本信息+消费偏好+红色情怀+"
                      "付费意愿+决策因素+品牌忠诚度。附调研样本量 N≥50 消费者问卷。）")

    # 四、产品/服务介绍（1000~1500 字）

    def _add_product_service(self):
        """产品/服务介绍：3 子节，含红色元素与适老化"""
        self.add_h1("四、产品/服务介绍")
        ps = self._get("product_service", default={})
        if not isinstance(ps, dict):
            ps = {}

        self.add_h2("（一）产品形态与核心功能")
        features = ps.get("features", [])
        if isinstance(features, str):
            features = [features]
        if features:
            for i, f in enumerate(features, 1):
                self.add_para(f"{i}. {f}")
        else:
            self.add_para("（请填写 MVP 功能清单 5~8 个核心功能，每个含名称 + 解决什么问题 + "
                          "老区/乡村适配性 + 红色元素融入。）")

        self.add_h2("（二）核心功能演示")
        demo = ps.get("demo", "")
        self.add_para(demo if demo else
                      "（请填写 1~2 个核心功能详细描述，含农户使用流程截图说明 + "
                      "老区/乡村用户的实际使用场景。）")

        self.add_h2("（三）技术实现与知识产权")
        tech = ps.get("tech_impl", "")
        self.add_para(tech if tech else
                      "（请填写技术架构（前端/后端/数据库）+ 关键第三方服务 + 研发投入估算 + "
                      "知识产权情况（软著/专利/在申）+ 红色品牌注册情况 + 适老化/低带宽/离线场景适配。）")
        ip = ps.get("ip_status", "")
        if ip:
            self.add_para(f"知识产权情况：{ip}")

    # 五、商业模式与可持续造血（1000~1300 字 + 9 宫格画布）

    def _add_business_model(self):
        """商业模式：画布 + 盈利 + 定价 + 可持续造血模式 + 财务自给率"""
        self.add_h1("五、商业模式与可持续造血")
        bm = self._get("business_model", default={})
        if not isinstance(bm, dict):
            bm = {}

        self.add_h2("（一）商业模式画布（9 要素）")
        canvas = bm.get("canvas", [])
        if canvas and isinstance(canvas, list):
            add_canvas_9grid(self.doc, canvas)
        else:
            self.add_para("（请填写商业模式画布 9 要素：客户细分/价值主张/渠道通路/客户关系/"
                          "收入来源/核心资源/关键业务/重要伙伴/成本结构。每格 30~50 字，缺一不可。"
                          "重要伙伴必含老区党支部 + 合作社 + 农户代表；客户细分含城市消费者 + 老区农户。）")

        self.add_h2("（二）盈利模式（造血式）")
        streams = bm.get("revenue_streams", [])
        if isinstance(streams, str):
            streams = [streams]
        if streams:
            for i, s in enumerate(streams, 1):
                self.add_para(f"{i}. {s}")
        else:
            self.add_para("（请填写造血式收入来源 1~3 种，如交易抽佣/品牌溢价/研学收费/服务费，"
                          "非捐赠输血。每种含定价 + 计算依据 + Year1 占营收比 + 农户分红比例（建议≥50%）。）")

        self.add_h2("（三）定价策略")
        pricing = bm.get("pricing", "")
        self.add_para(pricing if pricing else
                      "（请填写定价策略：成本加成/竞品对标/用户价值三方法，"
                      "含农户收购价 + 城市零售价合理梯度 + 红色品牌溢价测算 + 消费者付费意愿调研数据。）")

        self.add_h2("（四）可持续造血模式")
        blood = bm.get("blood_making", "")
        self.add_para(blood if blood else
                      "（请填写可持续造血模式：①造血收入占比（不依赖外部捐赠）；"
                      "②农户分红机制（利润反哺农户比例）；③退出机制（合作社独立运营能力培养）。）")

        self.add_h2("（五）财务自给率")
        self_suf = bm.get("self_sufficiency_rate", "")
        self.add_para(self_suf if self_suf else
                      "（请填写财务自给率：年度造血收入 / 年度总支出，≥70% 为健康。"
                      "Year1 可低（依赖启动资金），Year3 ≥70% 达标。如：Year1 65% / Year2 78% / Year3 92%。）")

    # 六、运营现状与红色帮扶成效（1200~1500 字，红旅核心章）

    def _add_operation_status(self):
        """运营现状：成绩 + 红色帮扶成效 + 5 维度振兴 + 里程碑 + KPI"""
        self.add_h1("六、运营现状与红色帮扶成效")
        op = self._get("operations", default={})
        if not isinstance(op, dict):
            op = {}

        self.add_h2("（一）已取得成绩")
        achievements = op.get("achievements", "")
        self.add_para(achievements if achievements else
                      "（请填写已取得成绩 400~500 字：用户数、营收、合作客户、签约订单、"
                      "媒体报道、获奖。无数据则写 MVP 上线时间 + 内测用户数 + 首批合作意向。）")

        self.add_h2("（二）红色帮扶成效（红旅赛道必填核心子节）")
        service = self._get("service_content", default={})
        if isinstance(service, dict) and service:
            add_help_effect_table(self.doc, service)
            cases = service.get("typical_cases", [])
            if cases:
                self.add_h3("典型农户帮扶案例")
                add_typical_case_table(self.doc, cases)
        else:
            self.add_para("（请填写红色帮扶成效 6 项量化数据：①服务覆盖县乡村数；"
                          "②带动农户数；③户均增收金额/年；④持续服务月数；"
                          "⑤合作社数量；⑥党支部结对数。附 1~2 个典型农户帮扶案例，"
                          "含姓名/村组/帮扶前收入/帮扶后收入/持续时长。）")

        self.add_h2("（三）乡村振兴 5 维度对应")
        revival = self._get("rural_revival", default={})
        if not isinstance(revival, dict):
            revival = {}
        revival_dims = [
            ("产业振兴", "industry", "带动产业 + 规模 + 升级"),
            ("人才振兴", "talent", "培训农户 + 培养带头人 + 大学生返乡"),
            ("文化振兴", "culture", "红色文化 + 非遗 + 传播"),
            ("生态振兴", "ecology", "绿色生产 + 生态保护 + 可持续"),
            ("组织振兴", "organization", "党支部结对 + 合作社 + 集体经济"),
        ]
        revival_summary = []
        for dim, key, hint in revival_dims:
            content = revival.get(key, "")
            self.add_h3(f"{dim}")
            self.add_para(content if content else
                          f"（请填写{dim} 200~300 字：{hint}）")
            revival_summary.append({
                "dimension": dim,
                "content": (content[:80] + "...") if content else "（待填写）",
            })
        add_rural_revival_table(self.doc, revival_summary)

        self.add_h2("（四）关键里程碑")
        milestones = op.get("milestones", [])
        if milestones:
            add_milestone_table(self.doc, milestones)
        else:
            self.add_para("（请填写关键里程碑表：时间/事件/数据，含 MVP 上线、首批农户入驻、"
                          "首笔营收、首个党支部结对、首个合作社成立、获奖等 6~10 个节点。）")

        self.add_h2("（五）关键运营指标 KPI")
        kpi_rows = op.get("kpi", [])
        if kpi_rows:
            add_kpi_table(self.doc, kpi_rows)
        else:
            self.add_para("（请填写 KPI 表 5~8 项：月活跃交易用户、农户入驻数、城市消费者数、"
                          "月营收、农户户均增收、财务自给率等。含当前值 + Year1/2/3 目标。）")

    # 七、团队介绍（每人 150~200 字 + 表格 + 政治面貌 + 党支部结对）

    def _add_team_intro(self):
        """团队介绍：核心成员表（含政治面貌）+ 简介 + 指导教师 + 党支部结对 + 志愿服务"""
        self.add_h1("七、团队介绍")
        team = self._get("team_intro", default={})
        if not isinstance(team, dict):
            team = {}

        self.add_h2("（一）核心成员表（含政治面貌）")
        members = team.get("members", [])
        if members:
            rows = [[m.get("name", ""), m.get("role", ""),
                     m.get("background", ""), m.get("duty", ""),
                     m.get("political_status", "")]
                    for m in members if isinstance(m, dict)]
            self.add_table(["姓名", "职务", "背景", "分工", "政治面貌"], rows,
                           col_widths=[1.8, 2.8, 4.5, 4.0, 2.2], last_col_left=False)
        else:
            self.add_para("（请填写核心成员表：姓名/职务/背景/分工/政治面貌，"
                          "5~8 人，含 1 名党员/预备党员/入党积极分子骨干，商科+涉农+技术+思政复合。）")

        self.add_h2("（二）核心成员简介")
        details = team.get("member_details", [])
        if details:
            for d in details:
                if isinstance(d, dict):
                    name = d.get("name", "")
                    detail = d.get("detail", "")
                    self.add_para(f"{name}：{detail}")
        else:
            self.add_para("（请填写核心成员简介，每人 150~200 字：专业背景 + 实习/项目经历 + "
                          "在本项目中具体职责 + 已完成的关键动作 + 思政提升（如发展为入党积极分子）。）")

        self.add_h2("（三）指导教师与顾问")
        adv = team.get("advisor_bg", "")
        self.add_para(adv if adv else
                      "（请填写指导教师背景：职称、研究方向（建议含涉农/经济/思政）、主持项目、"
                      "指导学生创业经历。建议含 1 名涉农/经济专业教师 + 1 名思政教师。"
                      "外部顾问可含老区党支部书记/合作社理事长。）")

        self.add_h2("（四）党支部结对与志愿服务")
        pair = team.get("party_branch_pair", "")
        vol = team.get("volunteer_hours", "")
        if pair:
            self.add_para(f"党支部结对情况：{pair}")
        else:
            self.add_para("（请填写与老区/乡村党支部结对情况：结对支部名称 + 结对方式 + 共建内容 + 协议。）")
        if vol:
            self.add_para(f"团队成员累计志愿服务时长：{vol}，含老区支教、农户培训、红色宣讲等。")
        else:
            self.add_para("（请填写团队成员累计志愿服务时长，含老区支教、农户培训、红色宣讲等。）")

    # 八、财务预测（800~1200 字 + 3 年表格）

    def _add_financial_forecast(self):
        """财务预测：3 年表格 + 盈亏平衡 + 财务合理性 + 农户分红"""
        self.add_h1("八、财务预测")
        fin = self._get("financial_forecast", default={})
        if not isinstance(fin, dict):
            fin = {}

        self.add_h2("（一）3 年财务预测表")
        rows_data = fin.get("rows", [])
        if rows_data:
            rows = [[r.get("item", ""), str(r.get("y1", "")),
                     str(r.get("y2", "")), str(r.get("y3", ""))]
                    for r in rows_data if isinstance(r, dict)]
            self.add_table(["财务项目", "Year1", "Year2", "Year3"], rows,
                           col_widths=[5.0, 3.5, 3.5, 3.5], last_col_left=False)
        else:
            self.add_para("（请填写 3 年财务预测表，至少含：入驻农户/营业收入/营业成本/毛利率/"
                          "期间费用/净利润/净利率/农户分红/财务自给率 9 行。）")

        self.add_h2("（二）盈亏平衡点说明")
        be = fin.get("breakeven", "")
        self.add_para(be if be else
                      "（请填写盈亏平衡点：YearX 第 Y 个月月营收达 Z 万元时毛利率覆盖固定成本。"
                      "BEP = 固定成本 / (1 - 变动成本率)。红旅赛道盈亏平衡点要求比主赛道宽松，"
                      "可允许 Year3 转正。）")

        self.add_h2("（三）财务合理性说明")
        rationale = fin.get("rationale", "")
        self.add_para(rationale if rationale else
                      "（请填写财务合理性：Year1 农户数 ≤ SOM × 25%；Year3 ≤ SOM × 100%；"
                      "毛利率 30%~60%（农产品电商偏低可放宽到 25%）；净利率 Year1 可负或微利（红旅允许微利起步）、"
                      "Year3 转正或微利；财务自给率 Year1 ≥65%、Year3 ≥70% 达标。）")

    # 九、风险与对策（500~700 字 + 6 类风险表）

    def _add_risk_analysis(self):
        """风险分析：6 类风险表（含帮扶成效可持续 + 红色基因淡化）"""
        self.add_h1("九、风险与对策")
        risks = self._get("risk_analysis", default=[])
        if isinstance(risks, str):
            risks = [risks]
        if risks:
            rows = [[r.get("type", ""), r.get("risk", ""),
                     r.get("prob", ""), r.get("impact", ""),
                     r.get("measure", "")]
                    for r in risks if isinstance(r, dict)]
            self.add_table(["风险类型", "具体风险", "概率", "影响", "应对措施"], rows,
                           col_widths=[2.5, 3.5, 1.5, 1.5, 6.5])
        else:
            self.add_para("（请填写风险分析表，6 类风险齐全：①市场风险；②技术风险；③运营风险；"
                          "④财务风险；⑤帮扶成效可持续风险（红旅特色）；⑥红色基因淡化风险（红旅特色）。"
                          "每类含具体风险 + 概率 + 影响 + 应对措施。）")

        self.add_para(
            "红旅赛道额外关注两类特色风险：①帮扶成效可持续风险——团队毕业后帮扶中断、"
            "政策变动、乡村基层组织配合度变化等，应对措施包括培养合作社独立运营能力、"
            "接班人机制、学校'红旅'传承；②红色基因淡化风险——商业扩张淡化红色基因、"
            "红色资源对接中断、团队思政教育松懈等，应对措施包括保留思政顾问席位、"
            "党支部结对长期化、红色品牌内核不漂移。",
            indent=False
        )

    # 十、个人成长与团队协作（800~1200 字）

    def _add_personal_growth(self):
        """个人成长与团队协作：5 要素 + 4 维度"""
        self.add_h1("十、个人成长与团队协作")
        self.add_h2("（一）个人成长（第一人称）")
        growth = self._get("personal_growth", default={})
        if not isinstance(growth, dict):
            growth = {}
        items = [
            ("role", "角色与动作", "你在项目中承担的具体角色 + 完成的具体动作。"),
            ("transformation", "从 X 变为 Y",
             "从 X 变为 Y 的具体变化，如'从只在象牙塔里学农经的学生变为能在井冈山老区跑通蜂蜜产业链的助农创业者'。"),
            ("hard_skills", "学到的硬技能",
             "硬技能：助农政策理解/农产品供应链/电商运营/农户沟通技巧等。"),
            ("soft_skills", "学到的软技能",
             "软技能：跨文化沟通（与老区农户）/抗压能力/团队协作/决策能力。"),
            ("failure", "失败复盘", "至少 1 个具体失败案例 + 复盘收获。建议与老区/乡村场景相关。"),
        ]
        for key, label, hint in items:
            val = growth.get(key, "")
            if val:
                self.add_para(f"{label}：{val}")
            else:
                self.add_para(f"（请填写{label}：{hint}）")

        self.add_h2("（二）团队协作")
        collab = self._get("team_collaboration", default={})
        if not isinstance(collab, dict):
            collab = {}
        division = collab.get("division", "")
        conflict = collab.get("conflict", "")
        tools = collab.get("tools", "")
        if division:
            self.add_para(f"分工机制：{division}")
        else:
            self.add_para("（请填写分工机制：每人专业背景与项目分工一一对应，"
                          "红色基因相关事项建议由思政顾问一票否决。）")
        if conflict:
            self.add_para(f"冲突案例：{conflict}")
        else:
            self.add_para("（请填写冲突案例 1 个，建议与老区/乡村场景相关，如农户临时毁约/"
                          "党支部换届对接中断/3 县扩张 vs 1 县深耕分歧。）")
        if tools:
            self.add_para(f"协作工具：{tools}")
        else:
            self.add_para("（请填写协作工具：飞书/企业微信/钉钉 + 周一例会 + 老区行日站会频次。）")

    # 主构建方法

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 11 栏目，生成 docx。返回实际保存路径。"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_cover()
            self._add_executive_summary()
            self._add_red_gene()
            self._add_market_analysis()
            self._add_product_service()
            self._add_business_model()
            self._add_operation_status()
            self._add_team_intro()
            self._add_financial_forecast()
            self._add_risk_analysis()
            self._add_personal_growth()
            self._post_check()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 红旅赛道商业计划书已生成：{output_path}")
        return str(output_path)

    # 数据校验

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []
        # P0 字段
        p0_fields = [
            ("project_name", "项目名称"),
            ("red_theme", "红旅主题（乡村振兴/革命老区/民族团结/西部开发）"),
            ("rural_county", "帮扶县乡村全称"),
            ("leader_name", "负责人姓名"),
            ("school", "学校"),
            ("advisor_name", "指导教师姓名"),
        ]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")
        # 主题值校验
        theme = self._get("red_theme", default="")
        if theme and theme not in RED_THEMES:
            warnings.append(f"红旅主题值 '{theme}' 不在 4 选 1 之内（乡村振兴/革命老区/民族团结/西部开发）")
        # 关键内容字段
        for key, name in [
            ("executive_summary", "执行摘要"),
            ("red_gene", "红色基因 4 维度"),
            ("service_content", "帮扶成效 6 项数据"),
            ("market_analysis", "市场分析"),
            ("business_model", "商业模式与可持续造血"),
            ("financial_forecast", "财务预测"),
            ("personal_growth", "个人成长"),
            ("team_intro", "团队介绍"),
        ]:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}），将使用占位文本")
        # 红色基因 4 维度校验
        gene = self._get("red_gene", default={})
        if isinstance(gene, dict) and gene:
            missing = [k for k in ["red_study", "revolution_spirit", "red_culture", "typical_case"]
                       if not gene.get(k)]
            if missing:
                warnings.append(f"红色基因 4 维度缺：{', '.join(missing)}（红旅赛道核心评审要素）")
        # 帮扶成效 6 项数据校验
        service = self._get("service_content", default={})
        if isinstance(service, dict) and service:
            missing_data = []
            for k in ["service_county", "farmer_count", "income_increase",
                      "coverage_month", "coop_count", "party_branch_pair"]:
                if service.get(k) in (None, "", 0):
                    missing_data.append(k)
            if missing_data:
                warnings.append(f"帮扶成效 6 项数据缺：{', '.join(missing_data)}")
        # 商业模式画布 9 要素
        bm = self._get("business_model", default={})
        if isinstance(bm, dict):
            canvas = bm.get("canvas", [])
            if canvas and len(canvas) < 9:
                warnings.append(f"商业模式画布仅 {len(canvas)} 要素，建议补齐 9 要素")
            if not bm.get("blood_making"):
                warnings.append("缺少可持续造血模式说明（blood_making）")
            if not bm.get("self_sufficiency_rate"):
                warnings.append("缺少财务自给率（self_sufficiency_rate，应 ≥70%）")
        # 团队政治面貌校验
        team = self._get("team_intro", default={})
        if isinstance(team, dict):
            members = team.get("members", [])
            if members:
                if len(members) > 15:
                    warnings.append(f"团队成员 {len(members)} 人，超过红旅上限 15 人")
                political_ok = any(
                    isinstance(m, dict) and m.get("political_status", "") in
                    ["党员", "预备党员", "入党积极分子"] for m in members
                )
                if not political_ok:
                    warnings.append("团队中无党员/预备党员/入党积极分子骨干（红旅赛道建议含 1 名）")
                if not team.get("party_branch_pair"):
                    warnings.append("缺少党支部结对情况（party_branch_pair）")
        # 风险 6 类校验
        risks = self._get("risk_analysis", default=[])
        if isinstance(risks, list) and risks:
            risk_types = {r.get("type", "") for r in risks if isinstance(r, dict)}
            required_types = {"帮扶成效可持续风险", "红色基因淡化风险"}
            missing_types = required_types - risk_types
            if missing_types:
                warnings.append(f"风险分析缺红旅特色类别：{', '.join(missing_types)}")
        # 字数版本校验
        wc = self._get("word_count", default="standard")
        if wc not in WORD_COUNT_VERSIONS:
            warnings.append(f"字数版本 '{wc}' 不在 brief/standard/enhanced 之内，默认 standard")

        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings

    def _post_check(self) -> List[str]:
        """构建后检查，返回警告列表"""
        warnings = []
        # 统计正文字数（粗略估算：所有 paragraph run text）
        total_chars = 0
        for p in self.doc.paragraphs:
            for r in p.runs:
                total_chars += len(r.text or "")
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            total_chars += len(r.text or "")
        wc_version = self._get("word_count", default="standard")
        wc_info = WORD_COUNT_VERSIONS.get(wc_version, WORD_COUNT_VERSIONS["standard"])
        if total_chars < wc_info["min"] or total_chars > wc_info["max"]:
            warnings.append(
                f"[字数] 全文 {total_chars} 字，建议 {wc_version} 档"
                f"（{wc_info['name']}，区间 {wc_info['min']}~{wc_info['max']}）"
            )
        if warnings:
            print("⚠️ 构建后检查警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据：红土蜜语——井冈山老区蜂蜜电商助农平台
# ============================================================

DEFAULT_DATA = {
    "project_name": "红土蜜语——井冈山老区蜂蜜电商助农平台",
    "track": "red_journey",
    "red_theme": "革命老区",
    "rural_county": "江西省吉安市井冈山市茅坪镇神山村",
    "word_count": "standard",
    "team_name": "红土蜜语团队",
    "leader_name": "张三",
    "leader_id": "202212345",
    "leader_major": "农林经济管理",
    "leader_grade": "2022 级 大三",
    "leader_phone": "138XXXXXXXX",
    "advisor_name": "李教授",
    "advisor_title": "副教授",
    "advisor_research": "农村电商与乡村振兴",
    "college": "经济管理学院",
    "school": "XX 大学",
    "apply_date": "2025 年 5 月 20 日",

    "executive_summary": (
        "本项目'红土蜜语'是为江西井冈山革命老区蜂蜜养殖户解决销路窄、价格低、品牌弱问题"
        "的电商助农平台，对接井冈山茅坪镇神山村党支部，传承'坚定信念、艰苦奋斗、实事求是、"
        "敢闯新路、依靠群众、勇于胜利'的井冈山精神。帮扶主题为革命老区，落地江西省吉安市"
        "井冈山市茅坪镇神山村、坝上村、马源村 3 个行政村。中国革命老区蜂蜜产业 2024 年规模"
        "120 亿元（农业农村部），年均增速 12%，但老区蜂农户均年收入仅 1.2 万元，低于全国"
        "蜂农平均 1.8 万元。本项目 SAM 32 亿元/年，3 年内 SOM 4800 万元/年。"
        "核心产品'红土蜜语'小程序已上线，含农户入驻（党支部背书）、红色品牌包装、LBS 同村溯源、"
        "担保交易、红色研学路线 5 大功能，注册 320 用户，入驻 156 户蜂农，月营收 8 万元。"
        "采用'蜂蜜销售 80% + 红色研学 15% + 品牌授权 5%'造血式盈利模式，"
        "农户分红比例 60%（利润反哺农户），财务自给率 Year1 65% / Year3 92%。"
        "团队 5 人跨学科复合，含 1 名党员、1 名入党积极分子，指导教师为农村电商副教授 + 思政讲师，"
        "与茅坪镇 3 个党支部结对共建。帮扶成效：覆盖 3 县 8 乡 25 村，带动 156 户蜂农，"
        "户均增收 4800 元/年，持续 18 个月，合作社 4 个，党支部结对 3 个。"
        "可持续造血模式：3 年内培养合作社独立运营能力，平台退出后合作社可持续自营。"
    ),

    "red_gene": {
        "red_study": (
            "团队组建以来，累计开展党史学习教育 12 次、学习时长 36 学时，团队成员中 2 人由入党"
            "积极分子发展为预备党员，3 人新提交入党申请书。学习内容包括《中国共产党农村工作条例》"
            "《乡村振兴促进法》《井冈山斗争史》《中国共产党简史》等。学习方式采用'理论+实践'双轨制："
            "理论方面，团队联合学校马克思主义学院开展'红蜜党课'6 期，邀请井冈山老党员线上讲党史 4 次；"
            "实践方面，团队 6 次赴井冈山茅坪镇开展'行走的党史课'，在八角楼、茨坪革命旧址等红色地标"
            "开展现场教学。通过学习，团队深入理解党在农村的百年奋斗历程，将'为人民服务'的宗旨"
            "内化为助农创业的行动指南，将'实事求是'的方法论内化为蜂蜜产业链调研的工作方法。"
        ),
        "revolution_spirit": (
            "项目对接井冈山革命老区，团队深入学习'坚定信念、艰苦奋斗、实事求是、敢闯新路、"
            "依靠群众、勇于胜利'的井冈山精神，将 24 字精神内涵转化为项目运营的具体行动："
            "将'坚定信念'内化为团队在 2024 年 9 月 MVP 上线初期 8 户蜂农无法使用时的坚持——"
            "18 天紧急开发离线版本，不放弃任何一户；将'艰苦奋斗'内化为团队 6 次老区行均住农家、"
            "吃红米饭、自背设备的作风；将'实事求是'内化为基于 156 户蜂农真实调研数据设计产品功能，"
            "不闭门造车；将'敢闯新路'内化为蜂蜜电商模式创新，开创'党支部+合作社+平台+农户'"
            "四位一体新模式；将'依靠群众'内化为党支部结对+合作社运营的群众路线；"
            "将'勇于胜利'内化为校赛金奖后冲击省赛国赛的拼搏精神。"
        ),
        "red_culture": (
            "项目对接井冈山革命博物馆与茅坪镇神山村党支部，联合开发'神山红蜜'品牌，"
            "将井冈山斗争时期的'红米饭、南瓜汤'精神融入品牌叙事，品牌包装融入八角楼、"
            "黄洋界等红色地标元素，开发'神山红蜜'系列红色文创衍生品 5 款（蜂蜜+红米+南瓜籽礼盒）。"
            "联合茅坪镇文化站开发'红蜜研学'路线 1 条，含八角楼现场教学 + 蜂场体验 + 红蜜制作，"
            "通过抖音'红土蜜语'账号发布老区故事短视频 36 条，累计播放 120 万次，"
            "红色文化传播覆盖全国 30 万抖音粉丝。同时联合学校开展'红蜜宣讲'校园活动 4 场，"
            "覆盖 1200 名大学生，传播井冈山精神与乡村振兴理念。"
        ),
        "typical_case": (
            "团队赴湖南十八洞村、江西神山村、陕西梁家河三地调研学习，提炼'党建引领+产业扶贫+"
            "群众主体'三大经验。项目对标十八洞村'精准扶贫'理念，实施'一户一策'精准帮扶 156 户"
            "蜂蜜养殖户，每户建立档案（含蜂群数/年产蜜量/收入/培训记录）；对标神山村'党建+产业'"
            "模式，与茅坪镇神山村党支部结对共建，签订《党建+产业合作协议》，联合成立'红土蜜语'"
            "蜂蜜专业合作社 4 个；对标梁家河'群众路线'，团队成员 6 次老区行均住在农户家中，"
            "与蜂农同吃同住同劳动，建立深厚群众基础。三地调研共形成调研报告 2 份、访谈记录 156 份。"
        ),
    },

    "rural_revival": {
        "industry": (
            "项目带动井冈山蜂蜜产业从'散户散养、自产自销'升级为'党支部+合作社+平台+农户'"
            "四位一体标准化产业链，覆盖茅坪镇 8 个行政村 156 户蜂农，年产蜂蜜 12 吨，"
            "户均养蜂收入从 1.2 万元/年增至 1.8 万元/年。项目推动产业标准化：制定《神山红蜜"
            "生产标准》1 套（含蜂群密度/采蜜周期/品质分级/包装规范），通过国家有机产品认证。"
            "项目延伸产业链：从单一蜂蜜销售扩展到蜂蜜+红蜜研学+品牌授权三产融合。"
        ),
        "talent": (
            "项目累计开展蜂农技术培训 6 期、培训农户 156 人次，培训内容包括蜂群管理/病害防治/"
            "品质分级/电商运营。培养本土带头人 8 名，含合作社理事长 4 名（神山村李大爷、坝上村"
            "王大叔、马源村张大哥、茅坪村陈师傅）、电商运营骨干 4 名。吸引 3 名本村大学生"
            "返乡加入合作社运营（其中 1 人为本项目团队成员张三），形成'外部输血+内部造血'"
            "的人才振兴机制。3 年内计划培养合作社独立运营能力，实现平台退出后合作社可持续自营。"
        ),
        "culture": (
            "项目将井冈山红色文化与客家蜂俗文化融合，开发'神山红蜜'红色品牌 + '客家蜂俗'"
            "非遗文创衍生品 5 款，含蜂蜜+红米+南瓜籽礼盒 + 客家蜂俗竹编包装。联合茅坪镇文化站"
            "开展'红蜜节'民俗活动 2 届，含蜂俗表演+红蜜品鉴+蜂场体验，吸引游客 800 人次。"
            "红色文化传播覆盖全国 30 万抖音粉丝，短视频 36 条累计播放 120 万次。同时联合学校"
            "开展'红蜜宣讲'校园活动 4 场，覆盖 1200 名大学生，传播井冈山精神与乡村振兴理念。"
        ),
        "ecology": (
            "项目推广'林下中蜂养殖'生态模式，蜂群在井冈山原始森林中自然采蜜，不投喂糖水"
            "不使用抗生素，蜂蜜通过国家有机产品认证。同时严格控制蜂群密度（每平方公里≤30 群），"
            "避免过度养殖破坏森林生态，实现'养蜂+护林'双赢。项目联合井冈山自然保护区开展"
            "'护林蜂农'认证，蜂农承担森林巡护义务，年巡护里程 1200 公里。项目推广绿色包装，"
            "使用可降解竹编+再生纸，年减少塑料使用 800 公斤，实现生产+包装+物流全链路生态化。"
        ),
        "organization": (
            "项目与茅坪镇神山村、坝上村、马源村 3 个党支部结对共建，签订《党建+产业合作协议》，"
            "联合开展'红蜜党课'6 次、'红蜜研学'3 期。联合成立'红土蜜语'蜂蜜专业合作社 4 个，"
            "社员 156 户，建立'党支部把方向+合作社管运营+平台拓销路+农户得实惠'的治理结构。"
            "2024 年合作社集体经济收入 12 万元，按 6:3:1 比例分配（农户 60% / 公积金 30% / "
            "公益金 10%）。项目推动村集体经济从 0 到 12 万元/年的突破，为革命老区村集体经济"
            "发展提供可复制模式。"
        ),
    },

    "service_content": {
        "service_county": 3,
        "service_township": 8,
        "service_village": 25,
        "farmer_count": 156,
        "income_increase": "4800 元/户·年",
        "coverage_month": 18,
        "coop_count": 4,
        "party_branch_pair": 3,
        "typical_cases": [
            {"name": "李大爷", "village": "神山村 1 组", "before": "1.2 万元/年",
             "after": "1.8 万元/年", "duration": "18 个月"},
            {"name": "王大叔", "village": "坝上村 3 组", "before": "0.9 万元/年",
             "after": "1.5 万元/年", "duration": "15 个月"},
            {"name": "张大哥", "village": "马源村 2 组", "before": "1.5 万元/年",
             "after": "2.1 万元/年", "duration": "12 个月"},
        ],
    },

    "coverage_data": {
        "field_visit_count": 6,
        "field_visit_days": 42,
        "interview_count": 156,
        "survey_report": 2,
        "field_visit_records": [
            {"time": "2024.07", "place": "井冈山市茅坪镇", "days": 7, "purpose": "蜂蜜产业调研"},
            {"time": "2024.09", "place": "井冈山市茅坪镇", "days": 5, "purpose": "MVP 上线 + 农户入驻"},
            {"time": "2024.12", "place": "井冈山市茅坪镇", "days": 7, "purpose": "首个合作社成立"},
            {"time": "2025.02", "place": "井冈山市茅坪镇", "days": 8, "purpose": "党支部结对 + 红蜜党课"},
            {"time": "2025.04", "place": "井冈山市茅坪镇", "days": 7, "purpose": "省赛备赛 + 红蜜节"},
            {"time": "2025.05", "place": "井冈山市茅坪镇", "days": 8, "purpose": "国赛备赛 + 红蜜研学"},
        ],
    },

    "market_analysis": {
        "industry_background": (
            "2024 年中国革命老区蜂蜜产业规模达 120 亿元（农业农村部《2024 中国蜂业年鉴》），"
            "年均增速 12%，但老区蜂农户均年收入仅 1.2 万元，低于全国蜂农平均 1.8 万元。"
            "国家发改委 2021 年《国务院关于新时代支持革命老区振兴发展的意见》（国发〔2021〕3 号）"
            "明确支持革命老区特色产业发展；2025 年中央一号文件《中共中央 国务院关于进一步深化"
            "农村改革 扎实推进乡村全面振兴的意见》要求大力发展乡村特色产业。井冈山革命老区"
            "蜂蜜品质优良（森林蜜源+无污染），但传统销路窄（依赖收购商压价至 30 元/斤）、"
            "品牌弱（无统一品牌）、信息散（蜂农分散在 8 个行政村），亟需电商平台整合。"
        ),
        "tam_sam_som": [
            {"level": "TAM", "definition": "全国革命老区蜂蜜产业规模（413 个革命老区县）",
             "scale": "120 亿元/年"},
            {"level": "SAM", "definition": "中部革命老区蜂蜜产业（江西/湖南/湖北/安徽等 6 省）",
             "scale": "32 亿元/年"},
            {"level": "SOM", "definition": "井冈山及周边 3 县蜂蜜产业（井冈山/永新/遂川）",
             "scale": "4800 万元/年"},
        ],
        "competitors": [
            {"name": "拼多多助农", "positioning": "通用电商助农", "users": "8 亿",
             "advantage": "流量大、价格低", "disadvantage": "无红色品牌、品质参差"},
            {"name": "京东助农", "positioning": "品质电商助农", "users": "5 亿",
             "advantage": "物流快、品质保证", "disadvantage": "无红色品牌、抽佣高"},
            {"name": "本地蜂蜜商", "positioning": "线下收购", "users": "本地",
             "advantage": "渠道稳定、面对面", "disadvantage": "压价至 30 元/斤、农户收益低"},
            {"name": "蜂农自销", "positioning": "农户自产自销", "users": "本地散客",
             "advantage": "信任度高", "disadvantage": "销路窄、规模小"},
        ],
        "differentiation": "红色品牌（神山红蜜）+ 党支部背书 + 在地化运营（团队 6 次老区行）+ "
                          "农户分红机制（60% 利润反哺农户）+ 红色研学延伸（红色文化变现）。",
        "user_persona_farmer": (
            "典型农户'李大爷'，男，58 岁，井冈山茅坪镇神山村 1 组，养蜂 20 年，年产蜂蜜 80 斤。"
            "原销路：本地蜂蜜商压价收购 30 元/斤，年收入 2400 元，加上其他农业收入 1 万元/年。"
            "痛点：销路窄、价格低、品质无认证、不会用电商。需求：稳定销路、合理价格、"
            "技术培训、品牌赋能。决策因素：信任度 > 价格 > 培训支持。访谈 30 户蜂农，"
            "28 户愿意加入合作社，22 户愿意参加培训。"
        ),
        "user_persona_consumer": (
            "典型消费者'小张'，男，32 岁，上海互联网公司中层，月收入 2 万元，已婚有娃。"
            "消费偏好：偏好有机/红色情怀/产地溯源产品，年均蜂蜜消费 500 元。红色情怀："
            "父亲是退伍军人，从小听井冈山故事长大。付费意愿：愿意为红色品牌+党支部背书+"
            "有机认证支付 50 元/斤溢价。决策因素：红色品牌 > 有机认证 > 价格 > 物流速度。"
            "访谈 50 名目标消费者，42 人愿意为'神山红蜜'支付 80 元/斤，68% 接受 380 元/人"
            "红色研学路线。"
        ),
    },

    "product_service": {
        "features": [
            "农户入驻：实名认证 + 党支部背书 + 合作社担保，解决陌生人交易信任问题。",
            "红色品牌包装：'神山红蜜'品牌 + 井冈山元素（八角楼/黄洋界）+ 客家蜂俗文化。",
            "LBS 同村溯源：扫码可见蜂场位置（精确到行政村）+ 蜂农信息 + 党支部背书。",
            "担保交易：买家确认收货后打款，农户款项 7 日内到账，降低交易风险。",
            "红色研学路线：联合茅坪镇开发'红蜜研学'路线，含八角楼现场教学+蜂场体验+红蜜制作。",
            "适老化设计：大字体+语音播报+电话订单兜底，适配老区农户使用习惯。",
            "离线适配：弱网/无网环境下可离线下单、离线同步，适配老区网络环境。",
        ],
        "demo": (
            "核心功能'LBS 同村溯源'演示：消费者扫码蜂蜜包装上的二维码，可看到蜂场位置（精确到"
            "神山村 1 组）、蜂农信息（李大爷，养蜂 20 年）、党支部背书（神山村党支部）、"
            "品质认证（国家有机产品认证）、生产日期、采蜜周期。整个流程 3 步以内完成。"
            "核心功能'红色研学路线'演示：消费者在小程序下单'红蜜研学 1 日游'（380 元/人），"
            "含八角楼现场教学 2 小时 + 蜂场体验 2 小时（含采蜜/品鉴）+ 红蜜制作 1 小时 + "
            "午餐（红米饭+南瓜汤+蜂蜜茶）。路线由茅坪镇文化站与项目团队联合开发，已接待 800 人次。"
        ),
        "tech_impl": (
            "前端微信小程序（Taro 3.x）+ 适老化大字体模式 + 语音播报；后端 Node.js + Express；"
            "数据库 MySQL 8.0 + Redis 7.0；第三方腾讯云 LBS / 阿里云 OSS / 微信支付 / "
            "顺丰快递 API。研发投入 3 人 × 4 月 = 12 人月，约 6 万元（校内外包价折算）。"
            "针对老区网络环境做了离线适配（弱网/无网环境下可离线下单、离线同步）+ 适老化设计"
            "（大字体+语音播报+电话订单兜底）。"
        ),
        "ip_status": "软件著作权 1 项（《红土蜜语老区蜂蜜电商助农平台 V1.0》已登记，"
                     "登记号 2024SR1234567）+ '神山红蜜'商标注册申请中（申请号 2024-789012）。",
    },

    "business_model": {
        "canvas": [
            {"element": "客户细分", "content": "城市消费者（25~45 岁、有红色情怀、月收入 8000+）+ 老区农户"},
            {"element": "价值主张", "content": "红色品牌蜂蜜 + 党支部背书 + 农户直供 + 红色研学体验"},
            {"element": "渠道通路", "content": "微信小程序 + 抖音直播 + 校园社群 + 红色研学路线"},
            {"element": "客户关系", "content": "党支部背书 + 合作社担保 + 售后无忧 + 红色情怀连接"},
            {"element": "收入来源", "content": "蜂蜜销售 80% + 红色研学 15% + 品牌授权 5%（造血式，非捐赠）"},
            {"element": "核心资源", "content": "井冈山红色 IP + 156 户蜂农 + 4 个合作社 + 神山红蜜品牌"},
            {"element": "关键业务", "content": "蜂蜜收购 + 品牌包装 + 电商销售 + 红色研学 + 农户培训"},
            {"element": "重要伙伴", "content": "茅坪镇 3 个党支部 + 4 个合作社 + 顺丰快递 + 井冈山博物馆"},
            {"element": "成本结构", "content": "蜂蜜收购 50% + 物流 15% + 包装 10% + 平台运营 15% + 农户分红 10%"},
        ],
        "revenue_streams": [
            "蜂蜜销售：年销 12 吨 × 80 元/斤 = 192 万元，毛利率 35%。Year1 占营收 80%。",
            "红色研学：年接待 800 人 × 380 元/人 = 30 万元，毛利率 60%。Year1 占营收 15%。",
            "品牌授权：授权其他老区使用'神山红蜜'品牌，年授权费 5 万元。Year1 占营收 5%。",
        ],
        "pricing": (
            "蜂蜜定价 80 元/斤（成本 52 + 红色品牌溢价 18 + 利润 10）。与拼多多助农 30 元/斤、"
            "京东助农 60 元/斤对比，红色品牌溢价 50 元/斤。消费者调研 50 人，42 人（84%）"
            "愿意为'神山红蜜'品牌 + 党支部背书 + 有机认证支付 80 元/斤。农户收购价 52 元/斤"
            "（vs 老收购商 30 元/斤），农户增收 22 元/斤。农户分红比例：利润 60% 反哺农户、"
            "30% 公积金、10% 公益金。"
        ),
        "blood_making": (
            "可持续造血模式：①造血收入（蜂蜜销售 80% + 红色研学 15% + 品牌授权 5%）占总收入 100%，"
            "不依赖外部捐赠或学校经费；②农户分红机制：利润 60% 反哺农户、30% 公积金、10% 公益金，"
            "确保农户持续受益；③退出机制：3 年内培养合作社独立运营能力（已培养合作社理事长 4 名、"
            "电商运营骨干 4 名），平台退出后合作社可持续自营，已复制到永新县 1 个合作社。"
        ),
        "self_sufficiency_rate": "财务自给率：Year1 65% / Year2 78% / Year3 92%，Year3 ≥70% 达标。",
    },

    "operations": {
        "achievements": (
            "项目自 2024 年 7 月首次井冈山调研启动，已上线 MVP 小程序，注册 320 用户，"
            "入驻 156 户蜂农，月营收 8 万元，累计营收 56 万元。已与茅坪镇 3 个党支部结对共建，"
            "联合成立 4 个合作社。获 2025 年校'互联网+'红旅赛道金奖、江西省赛金奖。"
            "抖音'红土蜜语'账号粉丝 30 万，短视频累计播放 120 万次。获新华社江西分社、"
            "江西卫视、学校官网等媒体报道 5 次。"
        ),
        "milestones": [
            {"time": "2024.07", "event": "首次井冈山调研", "data": "7 天 6 人"},
            {"time": "2024.09", "event": "MVP 上线", "data": "首批 32 户入驻"},
            {"time": "2024.10", "event": "首笔营收", "data": "月营收 1.2 万元"},
            {"time": "2024.12", "event": "首个合作社成立", "data": "神山红蜜合作社"},
            {"time": "2025.02", "event": "党支部结对", "data": "3 个支部"},
            {"time": "2025.03", "event": "覆盖 156 户", "data": "月营收 8 万元"},
            {"time": "2025.04", "event": "省赛金奖", "data": "晋级国赛"},
        ],
        "kpi": [
            {"metric": "月活跃交易用户", "current": "320", "y1": "1500", "y2": "5000", "y3": "1.2 万"},
            {"metric": "入驻农户数", "current": "156", "y1": "300", "y2": "600", "y3": "1200"},
            {"metric": "月营收（万元）", "current": "8", "y1": "30", "y2": "80", "y3": "180"},
            {"metric": "农户户均增收（元/年）", "current": "4800", "y1": "6000", "y2": "8000", "y3": "1 万"},
            {"metric": "财务自给率", "current": "65%", "y1": "70%", "y2": "78%", "y3": "92%"},
            {"metric": "合作社数量", "current": "4", "y1": "8", "y2": "15", "y3": "25"},
        ],
    },

    "financial_forecast": {
        "rows": [
            {"item": "入驻农户（户）", "y1": "300", "y2": "600", "y3": "1200"},
            {"item": "营业收入（万元）", "y1": "60", "y2": "180", "y3": "420"},
            {"item": "营业成本（万元）", "y1": "42", "y2": "120", "y3": "270"},
            {"item": "毛利率", "y1": "30%", "y2": "33%", "y3": "36%"},
            {"item": "期间费用（万元）", "y1": "15", "y2": "32", "y3": "60"},
            {"item": "净利润（万元）", "y1": "3", "y2": "28", "y3": "90"},
            {"item": "净利率", "y1": "5%", "y2": "15%", "y3": "21%"},
            {"item": "农户分红（万元）", "y1": "1.8", "y2": "16.8", "y3": "54"},
            {"item": "财务自给率", "y1": "70%", "y2": "78%", "y3": "92%"},
        ],
        "breakeven": (
            "盈亏平衡点：Year1 第 8 个月。届时月营收达 5 万元，毛利率 30% 覆盖月固定成本 1.5 万元"
            "（团队运营 1 万 + 平台 0.3 万 + 物流仓储 0.2 万）。BEP = 固定成本 1.5 / (1 - 变动成本率 70%) "
            "= 5 万元/月。"
        ),
        "rationale": (
            "Year1 农户 300 ≤ SOM 1200 × 25%；Year3 农户 1200 ≤ SOM × 100%。毛利率 30%→36%"
            "（农产品电商偏低 25%~40% 区间，符合行业实际）。净利率 Year1 5%（红旅允许微利起步）、"
            "Year3 21% 转正（红旅允许微利，因社会效益优先）。财务自给率 Year1 70% 达标、Year3 92% "
            "高于 70% 红线。农户分红 Year1 1.8 万元（占利润 60%）、Year3 54 万元（占利润 60%），"
            "确保农户持续受益。"
        ),
    },

    "financing_plan": {
        "amount": "30 万元", "equity": "10%", "valuation": "300 万元",
        "usage": "产品研发 30%（9 万）：3 名工程师 6 个月薪资 + 服务器扩容 + 第三方 API；"
                 "农户拓展 30%（9 万）：3 县 8 乡地推 + 合作社培训 + 蜂农技术培训 6 期；"
                 "品牌建设 20%（6 万）：'神山红蜜'品牌注册 + 包装升级 + 抖音投放 + 校园宣讲；"
                 "运营储备 20%（6 万）：法务/财务/应急。"
    },

    "team_intro": {
        "members": [
            {"name": "张三", "role": "CEO/创始人", "background": "农林经济管理 大三，入党积极分子",
             "duty": "战略+融资+党支部对接", "political_status": "入党积极分子"},
            {"name": "李四", "role": "CPO/联合创始人", "background": "农村区域发展 大三，党员",
             "duty": "产品+农户", "political_status": "党员"},
            {"name": "王五", "role": "CTO/联合创始人", "background": "软件工程 大三，团员",
             "duty": "技术+架构", "political_status": "团员"},
            {"name": "赵六", "role": "CMO/联合创始人", "background": "市场营销 大二，团员",
             "duty": "市场+品牌", "political_status": "团员"},
            {"name": "孙七", "role": "思政顾问", "background": "马克思主义学院 大三，党员",
             "duty": "红色基因+思政教育", "political_status": "党员"},
        ],
        "member_details": [
            {"name": "张三", "detail": "农林经济管理大三，GPA 3.85/4.0，专业排名 3/87。"
             "2024 年 7 月首次井冈山调研后提交入党申请书，2025 年 3 月发展为入党积极分子。"
             "本项目负责战略规划、融资对接、3 县 8 乡地推协调、与茅坪镇 3 个党支部对接。"},
            {"name": "李四", "detail": "农村区域发展大三，GPA 3.78/4.0，2023 年 12 月入党。"
             "曾参与学校'三下乡'赴井冈山社会实践，熟悉老区蜂业。本项目负责产品规划与农户沟通，"
             "完成 156 户蜂农访谈、5 版 MVP 迭代。"},
            {"name": "王五", "detail": "软件工程大三，GPA 3.72/4.0，全栈开发 2 年。"
             "本项目负责技术架构，独立完成 MVP 全栈实现，针对老区网络环境做了离线适配与适老化改造。"},
        ],
        "advisor_bg": (
            "李教授，副教授，经济管理学院，研究方向农村电商与乡村振兴，主持国家社科基金项目 1 项"
            "（革命老区电商助农模式研究），指导学生团队获'互联网+'红旅赛道省赛金奖 2 项。"
            "思政指导教师：王老师，马克思主义学院讲师，负责团队党史学习教育指导与红色基因 4 维度把关。"
            "外部顾问：神山村党支部书记李书记，提供党支部结对与红色资源对接支持。"
        ),
        "party_branch_pair": (
            "与茅坪镇神山村党支部、坝上村党支部、马源村党支部 3 个党支部结对共建，"
            "签订《党建+产业合作协议》，联合开展'红蜜党课'6 次、'红蜜研学'3 期，"
            "联合成立 4 个合作社，建立'党支部把方向+合作社管运营+平台拓销路+农户得实惠'治理结构。"
        ),
        "volunteer_hours": "团队成员累计志愿服务时长 480 小时，含老区支教 120 小时、"
                          "农户培训 200 小时、红色宣讲 160 小时。",
    },

    "risk_analysis": [
        {"type": "市场风险", "risk": "拼多多助农进入蜂蜜品类", "prob": "中", "impact": "高",
         "measure": "强化红色品牌+党支部背书+在地化运营，构建拼多多无法复制的红色壁垒"},
        {"type": "技术风险", "risk": "老区网络不稳定影响使用", "prob": "中", "impact": "中",
         "measure": "离线适配+适老化设计+电话订单兜底，已部署 8 户蜂农离线版本"},
        {"type": "运营风险", "risk": "农户临时毁约/品质不稳定", "prob": "高", "impact": "中",
         "measure": "合作社担保+品质分级+农户培训+保证金机制（每户 500 元）"},
        {"type": "财务风险", "risk": "物流成本高于预期", "prob": "中", "impact": "中",
         "measure": "顺丰协议价+集中发货+本地仓储备货，已签顺丰协议价 7 折"},
        {"type": "帮扶成效可持续风险", "risk": "团队毕业后帮扶中断", "prob": "高", "impact": "高",
         "measure": "3 年内培养合作社独立运营能力+接班人机制+学校'红旅'传承+党支部长期对接"},
        {"type": "红色基因淡化风险", "risk": "商业扩张淡化红色基因", "prob": "中", "impact": "高",
         "measure": "保留思政顾问席位（一票否决权）+党支部结对长期化+红色品牌内核不漂移"}
    ],

    "personal_growth": {
        "role": (
            "作为 CEO，我负责战略规划、融资对接、3 县 8 乡地推协调、与茅坪镇 3 个党支部对接、"
            "与指导教师的周例会汇报。"
        ),
        "transformation": (
            "我从'只在象牙塔里学农经的学生'变为'能在井冈山老区跑通蜂蜜产业链的助农创业者'，"
            "2024 年 7 月首次井冈山调研后我提交了入党申请书，2025 年 3 月发展为入党积极分子。"
            "项目过程中我深刻理解了党在农村的百年奋斗历程，将'为人民服务'的宗旨内化为助农创业"
            "的行动指南。"
        ),
        "hard_skills": (
            "我学会了助农政策理解（《乡村振兴促进法》《革命老区振兴规划》）、农产品供应链管理"
            "（蜂蜜收购→品牌包装→电商销售→物流配送）、电商运营（小程序/抖音/社群）、"
            "农户沟通技巧（用方言+实地示范代替 PPT 演讲）。"
        ),
        "soft_skills": (
            "软技能方面提升了跨文化沟通（与老区蜂农李大爷同吃同住同劳动 7 天）、抗压能力"
            "（路演前 3 天产品崩掉重写）、决策能力（在 3 县扩张 vs 1 县深耕间做出取舍，"
            "基于 156 户蜂农调研数据决定深耕 1 县 18 个月后再扩张）。"
        ),
        "failure": (
            "在 2024 年 9 月 MVP 上线初期，我因未考虑老区网络环境导致 32 户蜂农中 8 户无法"
            "正常使用小程序，被迫紧急开发离线版本，延迟 18 天。这次失败让我学会'技术适配场景先行'，"
            "也让我意识到助农项目必须真正下到田间地头，不能闭门造车。后续我把'老区行调研'作为"
            "项目铁律，团队 6 次老区行累计 42 天蹲点。"
        ),
    },

    "team_collaboration": {
        "division": (
            "CEO 张三负责战略+融资+党支部对接；CTO 王五负责技术+架构；CPO 李四负责产品+农户；"
            "CMO 赵六负责市场+品牌；思政顾问孙七负责红色基因+思政教育。重大事项投票 3/5 通过，"
            "红色基因相关事项由思政顾问一票否决。"
        ),
        "conflict": (
            "曾就'先扩张 3 县 vs 深耕 1 县'产生分歧，CMO 主张快速扩张、CEO 主张深耕。"
            "通过 156 户蜂农调研数据（深耕 1 县留存率高 2.5 倍）+ 投票机制解决，决定深耕 1 县"
            "18 个月后再扩张。这次冲突让我学会'决策必须基于真实调研数据而非主观判断'。"
        ),
        "tools": (
            "协作工具：飞书（文档/日历/视频会议）+ 企业微信（与党支部/合作社沟通）+ 钉钉（农户培训）"
            "+ 周一例会 + 老区行日站会 15 分钟。"
        ),
    },
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="互联网+大学生创新创业大赛 红色之旅赛道 商业计划书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="示例：\n  python build.py --data data.json --out output.docx\n"
               "  python build.py --demo --out demo.docx\n\n"
               "JSON 字段定义详见 SKILL.md 第九章。"
               "红旅赛道特色：红色基因 4 维度 + 乡村振兴 5 维度 + 帮扶成效 6 项数据 + 可持续造血模式。",
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档（红土蜜语——井冈山老区蜂蜜电商助农平台）")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        wc = data.get("word_count", "standard")
        wc_info = WORD_COUNT_VERSIONS.get(wc, WORD_COUNT_VERSIONS["standard"])
        print(f"ℹ️ 使用内置示例数据生成演示文档（{wc_info['name']}，"
              f"{data.get('apply_date', '')} 递交，红旅主题：{data.get('red_theme', '')}）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = RedTourDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
