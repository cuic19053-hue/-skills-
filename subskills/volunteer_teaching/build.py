#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
大学生暑期"三下乡"社会实践-支教类立项申报书 docx 生成器

格式标准：A4 纸张，页边距上下 2.54cm 左右 2.5cm；正文宋体小四 1.5 倍行距首行缩进 2 字符；
一级标题黑体三号居中；二级标题黑体小三；三级标题宋体四号加粗；表格宋体五号居中。

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第十一章。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"

SIZE_ER = Pt(22)
SIZE_SAN = Pt(16)
SIZE_XIAO_SAN = Pt(15)
SIZE_SI = Pt(14)
SIZE_XIAO_SI = Pt(12)
SIZE_WU = Pt(10.5)

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc, text: str, font_name: str = FONT_SONG, font_size=SIZE_XIAO_SI,
    bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True, line_spacing: float = 1.5,
    space_before: float = 0, space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
        space_before=6, space_after=6)


def add_heading_level3(doc, text: str):
    """三级标题：宋体四号加粗，左对齐"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
        space_before=6, space_after=3)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=indent,
        line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None,
                        left_cols: Optional[set] = None):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if left_cols is None:
        left_cols = set()
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j in left_cols \
                else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False, alignment=align)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def merge_vertical_cells(table, col_idx: int, start_row: int, end_row: int):
    """纵向合并单元格（用于签字栏预留空白）"""
    cells = [table.rows[r].cells[col_idx] for r in range(start_row, end_row + 1)]
    merged = cells[0]
    for c in cells[1:]:
        merged = merged.merge(c)


def fmt_money(amount) -> str:
    """金额格式化：整数转字符串 + 元"""
    try:
        return f"{int(amount)} 元"
    except (ValueError, TypeError):
        return str(amount)


def safe_int(value, default: int = 0) -> int:
    """安全转整数"""
    try:
        return int(str(value).strip())
    except (ValueError, TypeError):
        return default


def compute_total_periods(schedule: List[Dict[str, Any]]) -> int:
    """从按天实施安排表中提取总课时数"""
    if not isinstance(schedule, list):
        return 0
    return sum(safe_int(s.get("periods", 0)) for s in schedule if isinstance(s, dict))


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """三下乡社会实践-支教类立项申报书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def add_h1(self, text):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text):
        return add_heading_level2(self.doc, text)

    def add_h3(self, text):
        return add_heading_level3(self.doc, text)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_table(self, headers, rows, col_widths=None, left_cols=None):
        return add_table_from_data(self.doc, headers, rows,
                                   col_widths=col_widths, left_cols=left_cols)

    def add_page_break(self):
        add_page_break(self.doc)

    # 封面

    def _add_cover(self):
        """封面：黑体二号标题 + 4 行下划线信息"""
        for _ in range(3):
            self.doc.add_paragraph()

        title = "大学生暑期\u201c三下乡\u201d社会实践活动立项申报书"
        add_paragraph_with_format(
            self.doc, title, font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_before=12, space_after=12)

        subtitle = f"（{self._get('theme', default='乡村振兴')}主题·支教类）"
        add_paragraph_with_format(
            self.doc, subtitle, font_name=FONT_HEI, font_size=SIZE_SAN,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_after=24)

        for _ in range(3):
            self.doc.add_paragraph()

        project_name = self._get("team_name", default="赴 XX 县乡村振兴支教团")
        info_items = [("项目名称", project_name), ("团队名称", self._get("team_name")),
                      ("申报单位", self._get("college")), ("申报日期", self._get("apply_date"))]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI, font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True

        self.add_page_break()

    # 一、团队基本信息

    def _add_team_info(self):
        """一、团队基本信息表（7 行 2 列）"""
        self.add_h1("一、团队基本信息")
        team = self._get("team_info", default={})
        if not isinstance(team, dict):
            team = {}
        g = self._get
        leader = team.get("leader", f"{g('leader_name')} / {g('leader_id')} / "
                          f"{g('leader_major')} / {g('leader_grade')} / {g('leader_phone')}")
        advisor = team.get("advisor", f"{g('advisor_name')} / {g('advisor_title')} / "
                           f"{g('advisor_phone')} / {g('advisor_with_team', '随队')}")
        rows = [
            ["团队名称", team.get("team_name", g("team_name", ""))],
            ["实践主题", team.get("theme", g("theme", ""))],
            ["实践地点", team.get("location", g("location", ""))],
            ["实践时间", team.get("practice_time", g("practice_time", ""))],
            ["团队人数", team.get("team_size", g("team_size", ""))],
            ["队长", leader], ["指导教师", advisor],
        ]
        self.add_table(["项目", "内容"], rows, col_widths=[4.5, 11.5])

    # 二、团队成员信息表

    def _add_members_table(self):
        """二、团队成员信息表（5 列）"""
        self.add_h1("二、团队成员信息表")
        members = self._get("members", default=[])
        if members and isinstance(members, list):
            rows = [[m.get("name", ""), m.get("id", ""), m.get("major", ""),
                     m.get("role", ""), m.get("phone", "")]
                    for m in members if isinstance(m, dict)]
            self.add_table(["姓名", "学号", "专业年级", "团队分工", "联系方式"],
                           rows, col_widths=[2.5, 2.8, 3.5, 3.5, 3.7])
        else:
            self.add_para("（请填写团队成员信息表，每人一行：姓名 / 学号 / 专业年级 / "
                          "团队分工（学科 + 班级）/ 联系方式。建议跨学院组队，师范 + 艺术 + "
                          "体育 + 心理专业互补，分工具体到\u201c教什么 + 教哪个班\u201d。团队人数 6~15 人为宜。）")

    # 三、实践主题与背景

    def _add_theme_background(self):
        """三、实践主题与背景（400~600 字，3 段）"""
        self.add_h1("三、实践主题与背景")
        bg = self._get("theme_background", default=[])
        if isinstance(bg, str):
            bg = [bg]
        if bg:
            for para in bg:
                self.add_para(para)
        else:
            self.add_h2("（一）团中央主题对应")
            self.add_para("（请填写当年团中央发布的实践主题方向，以及本项目对应切入"
                          "的维度，150 字左右。如：2025 年团中央\u201c乡村振兴 青春建功\u201d"
                          "主题，本项目切入基础教育支援/素质拓展/心理关怀三维度。）")
            self.add_h2("（二）选址理由")
            self.add_para("（请填写选址理由，150~200 字，3 句话讲清\u201c为什么去这所学校\u201d："
                          "与团中央主题契合度 + 学校实际情况（学生数/师资/特色）+ 团队已有联系。）")
            self.add_h2("（三）实践对象基本情况")
            self.add_para("（请填写学校基本情况，100~150 字：学生人数 + 师生比 + 师资缺口 + "
                          "留守儿童占比。必须有数据，如学生 90 人、教师 8 人、师生比 1:11.25、"
                          "留守儿童占比 60%。）")

    # 四、实践目的与意义

    def _add_purpose_significance(self):
        """四、实践目的与意义（300~500 字，2 段）"""
        self.add_h1("四、实践目的与意义")
        ps = self._get("purpose_significance", default=[])
        if isinstance(ps, str):
            ps = [ps]
        if ps:
            for para in ps:
                self.add_para(para)
        else:
            self.add_h2("（一）对当地")
            self.add_para("（请填写对当地的意义，150~250 字。结构：通过 14 天 84 课时系统"
                          "教学，弥补科学/艺术/体育/心理 4 门课程师资缺口，重点解决 X 个"
                          "问题。留下教案 84 份 + 教学视频 12 段，建立长期线上辅导机制。）")
            self.add_h2("（二）对学生")
            self.add_para("（请填写对学生的意义，150~250 字。结构：掌握课程设计、教案撰写、"
                          "课堂管理、学生心理疏导等教学技能 + 深入乡村小学与留守儿童同"
                          "学习同活动 + 厚植教育情怀与乡土感情。）")

    # 五、实践内容与实施方案【重点，含按天+按课时双重表格】

    def _add_implementation_plan(self):
        """五、实践内容与实施方案（含按天表 + 按课时表 + 典型教案示例）"""
        self.add_h1("五、实践内容与实施方案")
        self.add_h2("（一）实践形式")
        form = self._get("implementation_plan", default={})
        if not isinstance(form, dict):
            form = {}
        form_text = form.get("form", "")
        if form_text:
            self.add_para(form_text)
        else:
            self.add_para("本项目实践形式为支教类，采用\u201c基础教育支援（语数英科 4 学科）+ 素质"
                          "拓展（艺术体育 2 学科）+ 心理关怀（团体辅导 + 个案疏导）\u201d三段式实施。"
                          "教学对象分小班（1~2 年级）、中班（3~4 年级）、大班（5~6 年级）3 个班，"
                          "每班 30 人，共 90 人。每日 6 课时（上午 4 + 下午 2），14 天合计 84 课时。")

        self.add_h2("（二）课程表（按课时 + 学科 + 教师分工）")
        self._render_curriculum_table()

        self.add_h2("（三）按天实施安排")
        schedule = form.get("schedule", [])
        if schedule and isinstance(schedule, list):
            rows = []
            total_periods = 0
            for s in schedule:
                if not isinstance(s, dict):
                    continue
                rows.append([s.get("date", ""), s.get("work", ""),
                             s.get("periods", "0"), s.get("output", "")])
                total_periods += safe_int(s.get("periods", 0))
            self.add_table(["日期", "主要工作", "课时数", "预期产出"], rows,
                           col_widths=[2.2, 5.5, 1.8, 6.5], left_cols={1, 3})
            self.add_para(f"合计 {total_periods} 课时（含基础教育支援 66 + 素质拓展 12 + 心理关怀 6）。"
                          f"每日 16:30-17:00 团队复盘会，整理当日教学问题并调整次日教案。")
        else:
            self.add_para("（请填写按天实施安排表格，4~6 行覆盖整个 14 天实践期。每行 4 列："
                          "日期 / 主要工作 / 课时数 / 预期产出。示例：7.15 抵达 + 启动会 + 入学"
                          "测试 0 课时；7.16-26 系统教学 66 课时；7.27 成果展示 + 联欢会 6 "
                          "课时；7.28 结业典礼 + 返程 0 课时。）")

        self.add_h2("（四）典型教案示例")
        self._render_lesson_plan_example()

    def _render_curriculum_table(self):
        """渲染课程表 5 列表格：节次/时间/学科/教师/班级"""
        curriculum = self._get("curriculum", default=[])
        headers = ["节次", "时间", "学科", "教师", "班级"]
        col_widths = [2.2, 2.8, 2.5, 3.0, 5.5]
        if isinstance(curriculum, list) and curriculum:
            rows = [[c.get("period", ""), c.get("time", ""), c.get("subject", ""),
                     c.get("teacher", ""), c.get("class", "")]
                    for c in curriculum if isinstance(c, dict)]
        else:
            rows = [
                ["第 1 节", "8:00-8:40", "语文", "张三", "小班"],
                ["第 2 节", "8:50-9:30", "英语", "李四", "小班"],
                ["第 3 节", "9:50-10:30", "数学", "王五", "小班"],
                ["第 4 节", "10:40-11:20", "科学", "赵六", "小班"],
                ["第 5 节", "14:30-15:20", "艺术", "孙七", "全班合班"],
                ["第 6 节", "15:30-16:20", "体育", "周八", "全班合班"],
            ]
        self.add_table(headers, rows, col_widths=col_widths)
        self.add_para("中班、大班同步排课，每班 6 课时/日，3 班共 18 课时/日。每名教师每日承担 2 节课，团队 7 名教师轮班，14 天合计 84 课时。")

    def _render_lesson_plan_example(self):
        """渲染典型教案示例（语文《荷花》第 3 课，5 部分结构）"""
        lp = self._get("lesson_plans", default=[])
        plan = lp[0] if (isinstance(lp, list) and lp and isinstance(lp[0], dict)) else {}
        self.add_para(f"课题：{plan.get('title', '语文《荷花》（人教版二年级下册第 3 课）')}")
        self.add_para(f"班级：{plan.get('class', '小班（1~2 年级）')}    "
                      f"课时：第 1 节 8:00-8:40    教师：{plan.get('teacher', '张三')}")
        self.add_para("一、教学目标", indent=False)
        objectives = plan.get("objectives")
        if not (isinstance(objectives, list) and objectives):
            objectives = [
                "知识目标：学生掌握\u201c荷花\u201d一词的字形（共 10 画）与含义。",
                "能力目标：学生能用\u201c荷花\u201d造 2 句完整话，并画一幅荷花简笔画。",
                "情感目标：学生体会荷花的美，培养观察自然、热爱生活的兴趣。"]
        for i, obj in enumerate(objectives, 1):
            self.add_para(f"{i}. {obj}")
        self.add_para("二、重点与难点", indent=False)
        kp = plan.get("key_points")
        if isinstance(kp, list) and kp:
            self.add_para(f"重点：{'；'.join(kp)}")
        else:
            self.add_para("重点：① \u201c荷花\u201d字形笔画（10 画）；② \u201c荷花\u201d造句。")
        diff = plan.get("difficulty") or "\u201c荷\u201d与\u201c何\u201d字形区分（草字头 vs 单人旁）"
        self.add_para(f"难点：{diff}")
        self.add_para("三、教学过程", indent=False)
        process = plan.get("process")
        if process:
            self.add_para(process)
        else:
            self.add_para("1. 导入（5 分钟）：教师出示荷花图片，提问\u201c这是什么花？在哪里见过？\u201d"
                          "学生观察并回答。2. 新授（25 分钟）：① 教师范读\u201c荷花\u201d，学生跟读 3 遍；"
                          "② 讲解字形；③ 板书示范笔画，学生书空；④ 教师造句示范，学生仿写。"
                          "3. 练习（8 分钟）：学生独立造 2 句话，教师巡视指导，选取 3 名学生朗读。"
                          "4. 总结（2 分钟）：教师总结本课要点，学生复述。")
        self.add_para("四、作业", indent=False)
        hw = plan.get("homework") or ("① 抄写\u201c荷花\u201d5 遍；② 用\u201c荷花\u201d造 2 句话；"
                                      "③ 画一幅荷花简笔画。")
        self.add_para(hw)

    def _add_curriculum(self):
        """六、课程表（支教类专属重点：14 天课程进度速览表）"""
        self.add_h1("六、课程表（14 天总览）")
        self.add_para("典型日课程表已在第五章（二）中呈现。本节呈现 14 天课程进度速览，便于评审判断课程系统性。")
        rows = [
            ["7.15", "抵达 + 启动会 + 入学测试", "0", "0", "0", "测试卷 90 份"],
            ["7.16-19", "语文 + 英语 + 数学 + 科学（4 天）", "4", "4", "0", "教案 24 份"],
            ["7.20-23", "语文 + 英语 + 数学 + 科学（4 天）", "4", "4", "0", "教案 24 份"],
            ["7.24-26", "语文 + 英语 + 数学 + 科学（3 天）", "3", "3", "0", "教案 18 份"],
            ["7.27", "教学成果展示 + 联欢会", "1", "1", "0", "文艺节目 6 个"],
            ["7.28", "结业典礼 + 返程", "0", "0", "0", "反馈表 90 份"],
            ["合计", "14 天", "16", "16", "12", "84 课时 + 90 份反馈"],
        ]
        self.add_table(["日期", "主要教学", "上午课时", "下午课时", "心理关怀", "预期产出"],
                       rows, col_widths=[2.0, 5.0, 1.8, 1.8, 1.8, 3.6], left_cols={1, 5})
        self.add_para("上午 4 节 + 下午 2 节为常规课时；心理关怀团体辅导 2 次（每周 1 次）"
                      "+ 个案疏导 14 次（每日 1 次），共 12 课时计入总课时 84 节。")

    # 七、安全保障预案

    def _add_safety_plan(self):
        """七、安全保障预案（300~500 字，5 段）"""
        self.add_h1("七、安全保障预案")
        plan = self._get("safety_plan", default=[])
        if isinstance(plan, str):
            plan = [plan]
        if plan and isinstance(plan, list):
            for para in plan:
                self.add_para(para)
        else:
            self.add_h2("（一）出行安全")
            self.add_para("（请填写出行安全，80 字：交通方式（统一购买高铁票）、住宿选择"
                          "（学校教师宿舍 4 人一间男女分区）、每日 18:00 前向指导教师报平安。）")
            self.add_h2("（二）人身安全")
            self.add_para("（请填写人身安全，80 字：防暑（正午 11:00-14:00 不组织户外体育"
                          "活动）、防疫（每日早晚测温）、学生校园安全（课间轮值值守、放学"
                          "家长签字接送）。）")
            self.add_h2("（三）应急联系人")
            self.add_para("（请填写应急联系人，60 字：3 名——指导教师、队长、学院团委，"
                          "附电话。当地对接：XX 小学校长、XX 镇卫生院 120 急救点。）")
            self.add_h2("（四）应急流程")
            self.add_para("（请填写应急流程，120 字：分级响应。一般事件（学生擦伤/轻微"
                          "不适）：班主任 → 校医 → 通知家长。重大事件（学生骨折/高热/"
                          "走失）：队长 → 指导教师 → 学院 → 120/110，30 分钟首报，"
                          "2 小时书面报告。）")
            self.add_h2("（五）保险购买")
            self.add_para("（请填写保险购买，60 字：全员购买短期意外险，保额 30 万元，"
                          "保费 15 元/人，保单号 PA20250715001-010，覆盖 7 月 14-29 日。"
                          "学生意外险由校方学平险覆盖。）")

    # 八、预期成果

    def _add_expected_results(self):
        """八、预期成果（必须可量化，以课时数 + 学生反馈为主）"""
        self.add_h1("八、预期成果")
        outcomes = self._get("expected_results", default=[])
        if isinstance(outcomes, str):
            outcomes = [outcomes]
        if not outcomes:
            outcomes = [
                "支教课时数 84 课时（基础教育支援 66 + 素质拓展 12 + 心理关怀 6）",
                "学生反馈表 90 份（含学习兴趣、课程满意度、教师评价 3 维度）",
                "教案 84 份（含教学目标、重点难点、教学过程、作业 4 部分）",
                "教学日志 14 份（每日教学反思）",
                "学生作品 90 份（绘画 / 手工 / 作文）",
                "文艺节目 6 个（结业联欢会）",
                "短视频 3 个（每个 3-5 分钟，发布在 B 站、抖音）",
                "纪录短片 1 个（10 分钟，用于校内汇报）",
                "新闻稿 5 篇（中青网 1 篇、校团委公众号 2 篇、学院公众号 2 篇）",
                "物资捐赠：图书 200 册（已募集）+ 文具 50 套",
            ]
        for o in outcomes:
            self.add_para(f"• {o}", indent=False)

    # 九、经费预算

    def _add_budget(self):
        """九、经费预算（3 列表格：科目/金额/计算依据）"""
        self.add_h1("九、经费预算")
        items = self._get("budget_items", default=[])
        if items and isinstance(items, list):
            rows = []
            total = 0
            for b in items:
                if not isinstance(b, dict):
                    continue
                amount_num = safe_int(b.get("amount", 0))
                total += amount_num
                rows.append([b.get("item", ""), fmt_money(amount_num), b.get("basis", "")])
            rows.append(["合计", fmt_money(total), ""])
            self.add_table(["预算科目", "金额", "计算依据"], rows,
                           col_widths=[3.5, 3.0, 9.5], left_cols={2})
        else:
            self.add_para("（请填写经费预算，6 类标准科目：交通费 / 食宿费 / 教学物资费 / "
                          "学生奖品费 / 印刷费 / 其他。每项金额非整数，附计算依据。"
                          "示例：交通费 3640 元 = 高铁 130 × 10 × 2 + 包车 300 × 4。）")

    # 十、宣传计划

    def _add_publicity_plan(self):
        """十、宣传计划（3 类媒体）"""
        self.add_h1("十、宣传计划")
        plan = self._get("publicity_plan", default=[])
        if isinstance(plan, str):
            plan = [plan]
        if plan and isinstance(plan, list):
            for p in plan:
                self.add_para(p)
        else:
            self.add_h2("（一）校内媒体")
            self.add_para("校团委公众号 2 篇、校报 1 篇、学院公众号 2 篇。发布时间节点："
                          "7.16 启动报道、7.22 中期进展、7.28 结项报道。")
            self.add_h2("（二）校外媒体")
            self.add_para("中青网 1 篇、XX 县电视台 1 条、XX 日报 1 篇。中青网稿件由"
                          "队长审核后投递，附教学照片 5~8 张。")
            self.add_h2("（三）新媒体")
            self.add_para("B 站 3 个短视频、抖音 3 个短视频、微信视频号 2 个。每个视频 "
                          "3~5 分钟，内容覆盖教学过程、学生故事、团队风采三个角度。")

    # 十一、指导教师 + 十二、学院团委意见

    def _add_review_section(self):
        """十一、指导教师意见 / 十二、学院团委意见（可选十三、学校审批）"""
        self.add_h1("十一、指导教师意见")
        for _ in range(6):
            self.doc.add_paragraph()
        self.add_para("指导教师签字：____________________    日期：______年____月____日", indent=False)
        self.add_h1("十二、学院团委意见")
        for _ in range(6):
            self.doc.add_paragraph()
        self.add_para("学院盖章：____________________    日期：______年____月____日", indent=False)
        if self._get("include_school_approval", default=False):
            self.add_h1("十三、学校审批意见")
            for _ in range(6):
                self.doc.add_paragraph()
            self.add_para("学校盖章：____________________    日期：______年____月____日", indent=False)
        special = self._get("tuancentral_special", default="")
        if special:
            self.add_h1("附：团中央专项报送说明")
            self.add_para(f"对应专项：{special}。本项目与专项主题契合度高，预期形成优秀"
                          "支教案例 1 个、教学视频集 1 套，报送团中央专项工作组。")

    # 主构建方法

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 12 栏目，生成 docx"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_cover()
            self._add_team_info()
            self._add_members_table()
            self._add_theme_background()
            self._add_purpose_significance()
            self._add_implementation_plan()
            self._add_curriculum()
            self._add_safety_plan()
            self._add_expected_results()
            self._add_budget()
            self._add_publicity_plan()
            self._add_review_section()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申报书已生成：{output_path}")
        return str(output_path)

    # 数据校验（含支教类专属校验）

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）。"""
        warnings = []
        p0_fields = [("team_name", "团队名称"), ("theme", "实践主题"),
                     ("location", "实践地点"), ("practice_time", "实践时间"),
                     ("leader_name", "队长姓名"), ("college", "申报单位")]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        team = self._get("team_info", default={})
        if not isinstance(team, dict):
            team = {}
        if not team.get("team_size") and not self._get("team_size"):
            warnings.append("缺少 团队人数（team_size）")
        if not team.get("advisor") and not self._get("advisor_name"):
            warnings.append("缺少 指导教师（advisor_name）")
        if not self._get("theme_background"):
            warnings.append("缺少 实践主题与背景（theme_background），将使用占位文本")
        if not self._get("purpose_significance"):
            warnings.append("缺少 实践目的与意义（purpose_significance），将使用占位文本")

        # 支教类专属校验：实践形式 + 按天表 + 课程表 + 教案
        impl = self._get("implementation_plan", default={})
        if not isinstance(impl, dict):
            impl = {}
        if not impl.get("form"):
            warnings.append("缺少 实践形式（implementation_plan.form），建议明确为支教类")

        if not impl.get("schedule"):
            warnings.append("缺少 按天实施安排（implementation_plan.schedule），将使用占位文本——评审会扣大分")
        else:
            for i, s in enumerate(impl.get("schedule", []), 1):
                if isinstance(s, dict):
                    if not s.get("date") or not s.get("work"):
                        warnings.append(f"按天表第 {i} 行缺少日期或主要工作")
                    if not s.get("output"):
                        warnings.append(f"按天表第 {i} 行缺少预期产出")
            total_p = compute_total_periods(impl.get("schedule", []))
            if total_p < 56:
                warnings.append(f"总课时数 {total_p} 课时偏低（支教类建议 ≥ 56），影响持续性评分")
            elif total_p >= 100:
                warnings.append(f"总课时数 {total_p} 课时偏高（建议 ≤ 100），师生易疲劳")

        curr = self._get("curriculum", default=[])
        if not curr:
            warnings.append("缺少 课程表（curriculum），将使用典型日课程表占位——支教类核心栏目")
        elif isinstance(curr, list):
            for i, c in enumerate(curr, 1):
                if isinstance(c, dict):
                    miss = [k for k in ("period", "time", "subject",
                                        "teacher", "class") if not c.get(k)]
                    if miss:
                        warnings.append(f"课程表第 {i} 行缺少字段：{'、'.join(miss)}")

        lps = self._get("lesson_plans", default=[])
        if not lps:
            warnings.append("缺少 教案（lesson_plans），将使用典型教案示例占位——支教类核心栏目")
        elif isinstance(lps, list):
            for i, lp in enumerate(lps, 1):
                if isinstance(lp, dict) and (not lp.get("objectives") or not lp.get("process")):
                    warnings.append(f"教案 {i} 缺少教学目标或教学过程")

        plan = self._get("safety_plan", default=[])
        if not plan:
            warnings.append("缺少 安全保障预案（safety_plan），将使用占位文本——安全预案不达标可能一票否决")
        else:
            plan_text = "\n".join(plan) if isinstance(plan, list) else str(plan)
            if "保险" not in plan_text:
                warnings.append("安全预案未提及保险购买情况——一票否决项")
            if "应急联系" not in plan_text and "应急联系人" not in plan_text:
                warnings.append("安全预案未列明应急联系人")
            if "应急流程" not in plan_text:
                warnings.append("安全预案未列明应急流程")
            if "学生" not in plan_text and "校园" not in plan_text:
                warnings.append("安全预案未提及学生校园安全（支教类专属风险）")

        results = self._get("expected_results", default=[])
        if not results:
            warnings.append("缺少 预期成果（expected_results），将使用占位文本")
        else:
            results_text = "\n".join(results) if isinstance(results, list) else ""
            if "课时" not in results_text and "节" not in results_text:
                warnings.append("预期成果未以支教课时数为主——支教类核心产出缺失")
            if "反馈" not in results_text:
                warnings.append("预期成果未注明学生反馈表份数")

        items = self._get("budget_items", default=[])
        if items:
            total = sum(safe_int(b.get("amount", 0)) for b in items if isinstance(b, dict))
            budget_total_num = safe_int(str(self._get("budget_total", default="")).strip(), default=-1)
            if budget_total_num >= 0 and total != budget_total_num:
                warnings.append(f"预算合计 {total} 元 与申请经费 {budget_total_num} 元不一致")
        else:
            warnings.append("缺少 经费预算（budget_items），将使用占位文本")

        members = self._get("members", default=[])
        if isinstance(members, list):
            for i, m in enumerate(members, 1):
                if isinstance(m, dict) and not m.get("role"):
                    warnings.append(f"成员 {m.get('name', f'#{i}')} 缺少教学分工（学科 + 班级）")

        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        self.warnings = warnings
        return warnings


# ============================================================
# 默认示例数据
# ============================================================

DEFAULT_DATA = {
    "team_name": "赴 XX 县乡村振兴支教团",
    "theme": "乡村振兴",
    "location": "XX 省 XX 县 XX 镇 XX 小学",
    "practice_time": "2025.07.15-07.28（14 天）",
    "team_size": "10 人",
    "leader_name": "张三", "leader_id": "202212345",
    "leader_major": "汉语言文学（师范）2022 级", "leader_grade": "大三",
    "leader_phone": "138XXXXXXXX",
    "advisor_name": "李教授", "advisor_title": "副教授",
    "advisor_phone": "139XXXXXXXX", "advisor_with_team": "随队",
    "college": "师范学院", "apply_date": "2025 年 5 月 20 日",
    "members": [
        {"name": "张三", "id": "202212345", "major": "汉语言文学 2022 级", "role": "队长 / 语文（小班）", "phone": "138XXXXXXXX"},
        {"name": "李四", "id": "202212346", "major": "英语（师范）2022 级", "role": "英语（中班 / 大班）", "phone": "138XXXXXXXX"},
        {"name": "王五", "id": "202212347", "major": "数学与应用数学 2022 级", "role": "数学（中班）", "phone": "138XXXXXXXX"},
        {"name": "赵六", "id": "202212348", "major": "科学教育 2023 级", "role": "科学（大班）", "phone": "138XXXXXXXX"},
        {"name": "孙七", "id": "202212349", "major": "音乐学 2023 级", "role": "艺术（全年级轮教）", "phone": "138XXXXXXXX"},
        {"name": "周八", "id": "202212350", "major": "体育教育 2023 级", "role": "体育（全年级轮教）", "phone": "138XXXXXXXX"},
        {"name": "吴九", "id": "202212351", "major": "应用心理学 2022 级", "role": "心理关怀 / 班级管理", "phone": "138XXXXXXXX"},
    ],
    "theme_background": [
        "团中央主题对应：2025 年团中央发布\u201c乡村振兴 青春建功\u201d暑期社会实践主题，鼓励高校学生深入乡村开展支教助学活动，弥补乡村基础教育师资短板。本项目紧扣该主题，聚焦\u201c基础教育支援 + 素质拓展 + 心理关怀\u201d三大模块，面向 XX 县乡村小学开展 14 天系统教学。",
        "选址 XX 省 XX 县 XX 镇 XX 小学原因有三：一是该校 2024 年被列入县级乡村教育振兴重点校，办学条件改善但师资短缺；二是该校开设 1~6 年级共 6 个班，学生 90 人，结构完整便于分班教学；三是团队中 2 名成员为该校校友，已与学校校长建立联系并完成前期需求调研。",
        "实践对象基本情况：XX 小学现有在校学生 90 人（1 年级 15 人 + 2 年级 15 人 + 3 年级 15 人 + 4 年级 15 人 + 5 年级 15 人 + 6 年级 15 人），教师 8 人（含 3 名支教特岗），师生比 1:11.25。学校开齐语数英课程，但科学、艺术、体育、心理健康课程存在师资缺口。学生中留守儿童占比 60%，心理关怀需求突出。",
    ],
    "purpose_significance": [
        "对当地：通过 14 天 84 课时的系统教学，弥补 XX 小学科学、艺术、体育、心理健康 4 门课程师资缺口，重点解决\u201c科学课无实验\u201d\u201c艺术课无专业教师\u201d\u201c体育课无系统训练\u201d\u201c留守儿童心理关怀缺失\u201d四个问题。支教结束后留下教案 84 份、教学视频 12 段，建立长期线上辅导机制，为学校后续教学提供持续支持。",
        "对学生：团队成员在教学实践中掌握课程设计、教案撰写、课堂管理、学生心理疏导等教学技能，提升\u201c站上讲台\u201d的职业能力；通过深入乡村小学与留守儿童同学习同活动，深化对乡村教育现状的认识，厚植教育情怀与乡土感情，培养扎根乡村、教书育人的价值观。",
    ],
    "implementation_plan": {
        "form": "本项目实践形式为支教类，采用\u201c基础教育支援（语数英科 4 学科）+ 素质拓展（艺术体育 2 学科）+ 心理关怀（团体辅导 + 个案疏导）\u201d三段式实施。教学对象分小班（1~2 年级）、中班（3~4 年级）、大班（5~6 年级）3 个班，每班 30 人，共 90 人。每日 6 课时（上午 4 + 下午 2），14 天合计 84 课时。",
        "schedule": [
            {"date": "7.15", "work": "抵达 XX 县，与校方对接，召开启动会；下午开展入学测试与分班", "periods": "0", "output": "启动会 1 次；测试卷 90 份"},
            {"date": "7.16-26", "work": "系统教学 11 天，每日 6 课时（语数英科 + 艺体）+ 心理关怀团体辅导与个案疏导", "periods": "78", "output": "教案 66 份；学生作业 90 份/天"},
            {"date": "7.27", "work": "教学成果展示 + 结业联欢会（文艺节目 6 个）", "periods": "6", "output": "文艺节目 6 个；学生作品 90 份"},
            {"date": "7.28", "work": "结业典礼（颁发结业证 + 收集反馈表）+ 返程", "periods": "0", "output": "结业证 90 份；反馈表 90 份"},
        ],
    },
    "curriculum": [
        {"period": "第 1 节", "time": "8:00-8:40", "subject": "语文", "teacher": "张三", "class": "小班"},
        {"period": "第 2 节", "time": "8:50-9:30", "subject": "英语", "teacher": "李四", "class": "小班"},
        {"period": "第 3 节", "time": "9:50-10:30", "subject": "数学", "teacher": "王五", "class": "小班"},
        {"period": "第 4 节", "time": "10:40-11:20", "subject": "科学", "teacher": "赵六", "class": "小班"},
        {"period": "第 5 节", "time": "14:30-15:20", "subject": "艺术", "teacher": "孙七", "class": "全班合班"},
        {"period": "第 6 节", "time": "15:30-16:20", "subject": "体育", "teacher": "周八", "class": "全班合班"},
    ],
    "lesson_plans": [{
        "title": "语文《荷花》（人教版二年级下册第 3 课）",
        "class": "小班（1~2 年级）",
        "teacher": "张三",
        "objectives": [
            "知识目标：学生掌握\u201c荷花\u201d一词的字形（共 10 画）与含义（水生花卉）。",
            "能力目标：学生能用\u201c荷花\u201d造 2 句完整话，并画一幅荷花简笔画。",
            "情感目标：学生体会荷花的美，培养观察自然、热爱生活的兴趣。"],
        "key_points": ["\u201c荷花\u201d字形笔画（荷 10 画、花 7 画）", "\u201c荷花\u201d造句"],
        "difficulty": "\u201c荷\u201d与\u201c何\u201d字形区分（草字头 vs 单人旁）",
        "process": "1. 导入（5 分钟）：教师出示荷花图片，提问\u201c这是什么花？在哪里见过？\u201d学生观察并回答。2. 新授（25 分钟）：① 教师范读\u201c荷花\u201d，学生跟读 3 遍；② 讲解字形\u201c荷\u201d草字头 + 何，\u201c花\u201d草字头 + 化；③ 板书示范笔画，学生书空；④ 教师造句示范\u201c池塘里的荷花开了\u201d，学生仿写。3. 练习（8 分钟）：学生独立造 2 句话，教师巡视指导，选取 3 名学生朗读。4. 总结（2 分钟）：教师总结\u201c荷花\u201d的字形、字义、造句要点，学生复述。",
        "homework": "① 抄写\u201c荷花\u201d5 遍；② 用\u201c荷花\u201d造 2 句话；③ 画一幅荷花简笔画。",
    }],
    "safety_plan": [
        "一、出行安全：全员统一购买高铁票，不单独行动；7 月 15 日集体乘 GXXX 次列车赴 XX 县。住宿选择 XX 小学教师宿舍（已与学校对接安排），4 人一间，男女分区。每日 18:00 前向指导教师报平安。",
        "二、人身安全：携带常用药品（藿香正气水、创可贴、退烧药、止泻药）；正午 11:00-14:00 不组织户外体育活动，避免中暑；每日早晚测温；学生校园安全：教学期间教师全程在岗，课间轮值值守，放学由家长接送签字确认。",
        "三、应急联系人：指导教师李教授（139XXXXXXXX）、队长张三（138XXXXXXXX）、学院团委王老师（137XXXXXXXX）。当地对接：XX 小学校长刘老师（136XXXXXXXX）、XX 镇卫生院 120 急救点（120）。",
        "四、应急流程：一般事件（学生擦伤 / 轻微不适）：班主任处理 → 校医 → 通知家长。重大事件（学生骨折 / 高热 / 走失）：队长 → 指导教师 → 学院 → 120/110，同步报告学校校长。教师自身伤病：队长 → 指导教师 → 学院。所有事件 30 分钟内首报，2 小时内书面报告。",
        "五、保险购买：全员购买中国平安短期意外险（保额 30 万元，保费 15 元/人，保单号 PA20250715001-010），覆盖 7 月 14-29 日。学生意外险由 XX 小学校方购买（学平险），覆盖教学期间。",
    ],
    "expected_results": [
        "支教课时数 84 课时（基础教育支援 66 + 素质拓展 12 + 心理关怀 6）",
        "学生反馈表 90 份（含学习兴趣、课程满意度、教师评价 3 维度）",
        "教案 84 份（含教学目标、重点难点、教学过程、作业 4 部分）",
        "教学日志 14 份（每日教学反思）",
        "学生作品 90 份（绘画 / 手工 / 作文）",
        "文艺节目 6 个（结业联欢会）",
        "短视频 3 个（每个 3-5 分钟，发布在 B 站、抖音）",
        "纪录短片 1 个（10 分钟，用于校内汇报）",
        "新闻稿 5 篇（中青网 1 篇、校团委公众号 2 篇、学院公众号 2 篇）",
        "物资捐赠：图书 200 册（已募集）+ 文具 50 套",
    ],
    "budget_items": [
        {"item": "交通费", "amount": "3640", "basis": "高铁往返 130 元 × 10 人 × 2 次 + 县内包车 300 元/天 × 4 天"},
        {"item": "食宿费", "amount": "8400", "basis": "学校宿舍 60 元/人/天 × 10 人 × 14 天 + 餐费 40 元/人/天 × 10 人 × 14 天"},
        {"item": "教学物资费", "amount": "1280", "basis": "教具 300 元 + 学生奖品 500 元 + 实验耗材 280 元 + 体育用品 200 元"},
        {"item": "学生奖品费", "amount": "480", "basis": "学习用品 12 元 × 30 套 + 文具 4 元 × 30 套（按班发放）"},
        {"item": "印刷费", "amount": "580", "basis": "教案 84 份 × 3 元 + 测试卷 90 份 × 1 元 + 反馈表 90 份 × 1 元 + 海报 10 张 × 8 元"},
        {"item": "其他", "amount": "220", "basis": "保险 15 元 × 10 人 + 通讯补贴 7 元 × 10 人"},
    ],
    "budget_total": "14600",
    "publicity_plan": [
        "（一）校内媒体：校团委公众号 2 篇、校报 1 篇、学院公众号 2 篇。发布时间节点：7.16 启动报道、7.22 中期进展、7.28 结项报道。",
        "（二）校外媒体：中青网 1 篇、XX 县电视台 1 条、XX 日报 1 篇。中青网稿件由队长审核后投递，附教学照片 5~8 张。",
        "（三）新媒体：B 站 3 个短视频、抖音 3 个短视频、微信视频号 2 个。每个视频 3~5 分钟，内容覆盖教学过程、学生故事、团队风采三个角度。",
    ],
    "include_school_approval": False,
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="三下乡社会实践-支教类立项申报书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=("示例：\n"
                "  python build.py --data data.json --out output.docx\n"
                "  python build.py --demo --out demo.docx\n"
                "\nJSON 字段定义详见 SKILL.md 第十一章。"),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
