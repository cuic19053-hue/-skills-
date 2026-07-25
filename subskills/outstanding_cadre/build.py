#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
优秀班干部 / 优秀学生干部申请书 docx 生成器

面向担任班长/团支书/学委/宣委/体委/心理委员/楼长等学生干部的学生。
评审侧重"职务履职 + 服务同学 + 综合表现"，本学年跨度。
正文 1500~2000 字，4 段结构：个人基本情况 + 履职情况 + 工作成绩 + 不足与改进。

按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 标题：黑体二号，居中（按 honor_title 动态填充"优秀班干部申请书"等）
- 称呼：顶格，宋体小四，全角冒号
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 学生干部履职表、履职量化清单表：宋体五号，居中
- "此致"另起一行空两格，"敬礼！"另起一行顶格
- 落款：右对齐

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第六章信息采集清单。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号（标题）
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四（正文）
SIZE_WU = Pt(10.5)          # 五号（表格）
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(doc, text: str, font_name: str = FONT_SONG,
        font_size=SIZE_XIAO_SI, bold: bool = False,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent: bool = True,
        line_spacing: float = 1.5, space_before: float = 0, space_after: float = 0):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_ER,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进），宋体小四"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_section_heading(doc, text: str):
    """正文小节标题（一、二、三…）：黑体小四加粗，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SI,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
        line_spacing=1.5, space_before=6, space_after=3)


def add_cizhi_paragraph(doc, text: str = "此致"):
    """'此致'另起一行，空两格（首行缩进 2 字符）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格（不缩进）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=False, line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None, caption: str = ""):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    if caption:
        add_paragraph_with_format(doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def format_rank_percent(rank: str, rank_total: str) -> str:
    """根据排名与基数计算百分比，返回 '前 X.X%' 字符串，失败返回空串"""
    try:
        r_num = int(str(rank).split("/")[0])
        total_num = int(rank_total) if rank_total else 0
        if total_num > 0:
            return f"前 {round(r_num / total_num * 100, 1)}%"
    except (ValueError, IndexError):
        pass
    return ""


def parse_month(s: str) -> int:
    """解析 'YYYY.MM' 字符串为月份总数（用于任期时长计算）。失败返回 0。"""
    if not s:
        return 0
    try:
        parts = str(s).strip().split(".")
        if len(parts) == 2:
            y, m = int(parts[0]), int(parts[1])
            return y * 12 + m
    except (ValueError, IndexError):
        pass
    return 0


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """优秀班干部 / 优秀学生干部申请书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段，非列表返回空列表"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_heading(self, text):
        return add_section_heading(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, caption=""):
        return add_table_from_data(self.doc, headers, rows, col_widths, caption)

    # --------------------------------------------------------
    # 标题（按 honor_title 动态填充）
    # --------------------------------------------------------

    def _get_honor_title(self) -> str:
        """获取评优标题全称，默认"优秀班干部" """
        title = str(self._get("honor_title", default="优秀班干部"))
        # 标准化常见变体
        title_map = {
            "优秀班干部": "优秀班干部",
            "优秀学生干部": "优秀学生干部",
            "优秀班委": "优秀班委",
            "校级优秀学生干部": "校级优秀学生干部",
            "院级优秀学生干部": "院级优秀学生干部",
            "校级优秀班干部": "校级优秀班干部",
            "院级优秀班干部": "院级优秀班干部",
        }
        return title_map.get(title, title)

    def _add_title(self):
        """标题：黑体二号居中，按 honor_title 动态填充"""
        title = self._get_honor_title()
        add_title(self.doc, f"{title}申请书")

    # --------------------------------------------------------
    # 称呼
    # --------------------------------------------------------

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号"""
        salutation = self._get("salutation",
                               default="尊敬的学院领导、评审委员会：")
        add_salutation_paragraph(self.doc, salutation)

    # --------------------------------------------------------
    # 开头段落
    # --------------------------------------------------------

    def _add_opening(self):
        """开头段落（80~120 字）：身份 + 职务 + GPA + 排名 + 申报 + 进入句"""
        opening = self._get("opening", default="")
        if opening:
            self.add_para(opening)
            return
        name, college, major, grade = self._get("name"), self._get("college"), self._get("major"), self._get("grade")
        cadre_position = self._get("cadre_position")
        tenure_start, tenure_end = self._get("tenure_start"), self._get("tenure_end")
        apply_year = self._get("apply_year", default="2024-2025 学年")
        honor_title = self._get_honor_title()
        gpa, rank = self._get("gpa"), self._get("rank")
        parts = []
        if name and college and major and grade:
            parts.append(f"我是{college}{major}{grade}学生{name}，")
        else:
            parts.append("我")
        if cadre_position and tenure_start and tenure_end:
            parts.append(f"本学年担任{cadre_position}（{tenure_start}-{tenure_end}），")
        if gpa:
            parts.append(f"本学年 GPA {gpa}")
            if rank:
                parts.append(f"（专业排名第 {rank}），")
            else:
                parts.append("，")
        parts.append(f"特申请{apply_year}{honor_title}。现将本学年履职情况汇报如下：")
        self.add_para("".join(parts))

    # --------------------------------------------------------
    # 一、个人基本情况
    # --------------------------------------------------------

    def _add_basic_info(self):
        """个人基本情况（300~400 字，20%）：身份 + 职务 + 学业 + 政治 + 申报"""
        self.add_heading("一、个人基本情况")
        basic_info = self._get("basic_info", default="")
        if basic_info:
            if isinstance(basic_info, list):
                for p in basic_info:
                    self.add_para(p)
            else:
                self.add_para(basic_info)
            return
        name = self._get("name")
        college, major, grade, class_name = self._get("college"), self._get("major"), self._get("grade"), self._get("class_name")
        student_id, political_status = self._get("student_id"), self._get("political_status")
        cadre_position = self._get("cadre_position")
        cadre_level = self._get("cadre_level")
        tenure_start, tenure_end = self._get("tenure_start"), self._get("tenure_end")
        secondary_position = self._get("secondary_position")
        gpa, weighted, rank, rank_total = self._get("gpa"), self._get("weighted_avg"), self._get("rank"), self._get("rank_total")
        course_count, high_score_count = self._get("course_count"), self._get("high_score_count")
        party_history = self._get("party_history")
        parts = []
        # 身份介绍句
        if name and college and major and grade and class_name:
            parts.append(f"我是{college}{major}{grade}{class_name}学生{name}")
            if student_id:
                parts.append(f"，学号 {student_id}")
            if political_status:
                parts.append(f"，政治面貌 {political_status}")
            parts.append("。")
        # 职务与任期
        if cadre_position and tenure_start and tenure_end:
            seg = f"本学年担任{cadre_position}（{tenure_start}-{tenure_end}"
            if cadre_level:
                seg += f"，{cadre_level}"
            seg += "）"
            if secondary_position:
                seg += f"，同时兼任{secondary_position}"
            seg += "。"
            parts.append(seg)
        # 学业
        if gpa:
            seg = f"本学年 GPA {gpa}"
            if weighted:
                seg += f"（加权平均分 {weighted}）"
            if rank:
                seg += f"，专业排名第 {rank}"
                if rank_total:
                    pct = format_rank_percent(rank, rank_total)
                    if pct:
                        seg += f"（{pct}）"
            seg += "。"
            parts.append(seg)
        if course_count and high_score_count:
            parts.append(f"本学年修读 {course_count} 门课程，{high_score_count} 门 85 分以上。")
        # 思想政治
        if party_history:
            parts.append(party_history)
        # 申报
        apply_year = self._get("apply_year", default="2024-2025 学年")
        honor_title = self._get_honor_title()
        parts.append(f"现将本学年履职情况汇报如下，特申请{apply_year}{honor_title}。")
        self.add_para("".join(parts))
        # 主干课程表（如有）
        self._add_core_courses_table()

    def _add_core_courses_table(self):
        """主干课程表：4~6 门，含课程名/学分/成绩"""
        courses = self._get_list("core_courses")
        if not courses:
            return
        rows = []
        for c in courses:
            if not isinstance(c, dict):
                continue
            rows.append([
                str(c.get("name", "")),
                str(c.get("credit", "")),
                str(c.get("score", "")),
            ])
        if rows:
            self.add_table(
                ["课程名称", "学分", "成绩"],
                rows,
                col_widths=[8.0, 3.0, 3.0],
                caption="主干课程成绩：",
            )

    # --------------------------------------------------------
    # 二、履职情况（含学生干部履职表 + 履职量化清单表 + 服务同学案例）
    # --------------------------------------------------------

    def _add_duty_performance(self):
        """履职情况（600~800 字，40%，核心段）

        按 4 要素：① 职务履职概况 ② 学生干部履职表 ③ 履职量化清单表 ④ 服务同学案例
        """
        self.add_heading("二、履职情况")
        duty_performance = self._get("duty_performance", default="")
        if duty_performance and isinstance(duty_performance, str):
            self.add_para(duty_performance)
            return
        # ① 职务履职概况
        self._add_duty_overview()
        # ② 学生干部履职表
        self._add_position_table()
        # ③ 履职量化清单表
        self._add_duty_quantitative_table()
        # ④ 服务同学案例
        self._add_service_cases()

    def _add_duty_overview(self):
        """职务履职概况（约 100 字）：核心职责 + 工作量概述"""
        cadre_position = self._get("cadre_position")
        tenure_start, tenure_end = self._get("tenure_start"), self._get("tenure_end")
        duty_performances = self._get("duty_performances", default={})
        activity_count = ""
        service_hours = ""
        if isinstance(duty_performances, dict):
            activity_count = str(duty_performances.get("activity_count", ""))
            service_hours = str(duty_performances.get("service_hours", ""))
        parts = []
        if cadre_position and tenure_start and tenure_end:
            parts.append(f"本学年担任{cadre_position}（{tenure_start}-{tenure_end}）")
            parts.append("，核心职责为班级日常管理、协调班委分工、传达学院通知、组织主题班会与班级活动。")
        if activity_count:
            parts.append(f"本学年累计组织 {activity_count}；")
        if service_hours:
            parts.append(f"履职总时长 {service_hours}。")
        parts.append("学生干部履职情况如下：")
        self.add_para("".join(parts))

    def _add_position_table(self):
        """学生干部履职表（4 列：职务/任期/主要工作/成效）"""
        positions = self._get_list("position_history")
        if not positions:
            return
        rows = []
        for p in positions:
            if not isinstance(p, dict):
                continue
            rows.append([
                str(p.get("role", "")),
                str(p.get("term", "")),
                str(p.get("work", "")),
                str(p.get("effect", "")),
            ])
        if rows:
            self.add_table(
                ["职务", "任期", "主要工作", "成效"],
                rows,
                col_widths=[3.0, 2.5, 6.0, 3.5],
                caption="学生干部履职情况：",
            )

    def _add_duty_quantitative_table(self):
        """履职量化清单表（2 列：维度/量化数据，4 行）"""
        duty_performances = self._get("duty_performances", default={})
        if not isinstance(duty_performances, dict) or not duty_performances:
            return
        rows = [
            ["活动数", str(duty_performances.get("activity_count", ""))],
            ["服务时长", str(duty_performances.get("service_hours", ""))],
            ["获奖情况", str(duty_performances.get("awards", ""))],
            ["同学评议", str(duty_performances.get("peer_review", ""))],
        ]
        # 过滤掉全部为空的行
        if any(r[1] for r in rows):
            self.add_table(
                ["维度", "量化数据"],
                rows,
                col_widths=[3.0, 12.0],
                caption="履职量化清单：",
            )

    def _add_service_cases(self):
        """服务同学案例（2-3 个，约 200 字）"""
        cases = self._get_list("service_cases")
        if not cases:
            return
        parts = ["服务同学方面，本学年重点完成以下工作："]
        for i, c in enumerate(cases, start=1):
            if not isinstance(c, dict):
                continue
            background = str(c.get("background", ""))
            action = str(c.get("action", ""))
            effect = str(c.get("effect", ""))
            seg = f"{['①', '②', '③', '④', '⑤'][min(i-1, 4)]} "
            if background:
                seg += background + "，"
            if action:
                seg += action
            if effect:
                seg += "，" + effect
            seg += "；"
            parts.append(seg)
        # 末尾分号改句号
        if parts[-1].endswith("；"):
            parts[-1] = parts[-1].rstrip("；") + "。"
        self.add_para("".join(parts))

    # --------------------------------------------------------
    # 三、工作成绩
    # --------------------------------------------------------

    def _add_work_achievements(self):
        """工作成绩（375~500 字，25%）：整体成效 + 个人荣誉 + 跨职务统筹"""
        self.add_heading("三、工作成绩")
        work_achievements = self._get("work_achievements", default="")
        if work_achievements and isinstance(work_achievements, str):
            self.add_para(work_achievements)
            return
        class_achievements = self._get("class_achievements", default="")
        # 班级整体成效
        if class_achievements:
            self.add_para(f"班级整体成效方面，{class_achievements}")
        # 个人荣誉
        self._add_honors()
        # 跨职务统筹
        cross_position = self._get("cross_position", default="")
        if cross_position:
            self.add_para(f"跨职务统筹方面，{cross_position}")

    def _add_honors(self):
        """个人荣誉列表（含级别与时间）"""
        honors = self._get_list("honors")
        if not honors:
            return
        parts = []
        for h in honors:
            if not isinstance(h, dict):
                continue
            time = str(h.get("time", ""))
            name = str(h.get("name", ""))
            level = str(h.get("level", ""))
            issuer = str(h.get("issuer", ""))
            seg = ""
            if time:
                seg += f"{time} "
            if name:
                seg += name
            if level:
                seg += f"（{level}）"
            if issuer:
                seg += f"，{issuer}"
            if seg:
                parts.append(seg + "；")
        if parts:
            parts[-1] = parts[-1].rstrip("；") + "。"
            self.add_para("个人荣誉方面，本学年所获荣誉：" + "".join(parts))

    # --------------------------------------------------------
    # 四、不足与改进
    # --------------------------------------------------------

    def _add_shortcomings_section(self):
        """不足与改进（225~300 字，15%）：1-2 条不足 + 改进措施 + 表态"""
        self.add_heading("四、不足与改进")
        shortcomings_section = self._get("shortcomings_section", default="")
        if shortcomings_section and isinstance(shortcomings_section, str):
            self.add_para(shortcomings_section)
            return
        shortcomings = self._get_list("shortcomings")
        improvements = self._get_list("improvements")
        # 不足
        if shortcomings:
            parts = ["回顾本学年履职，存在以下不足："]
            for i, s in enumerate(shortcomings, start=1):
                parts.append(f"{i}. {s} ")
            self.add_para("".join(parts))
        # 改进措施
        if improvements:
            parts = ["改进措施："]
            for i, imp in enumerate(improvements, start=1):
                parts.append(f"{i}. {imp} ")
            self.add_para("".join(parts))

    # --------------------------------------------------------
    # 结尾"此致 敬礼！"
    # --------------------------------------------------------

    def _add_ending(self):
        """结尾（60~100 字）：本学年事实总结 + 朴素表态 + 此致/敬礼！"""
        ending = self._get("ending", default="")
        if ending:
            self.add_para(ending)
        else:
            honor_title = self._get_honor_title()
            self.add_para(
                f"以上是我本学年的基本情况。无论评选结果如何，我都将以此为新的起点，继续努力履职，以更高标准要求自己，争取获评{honor_title}。恳请评审委员会予以考虑。"
            )
        add_cizhi_paragraph(self.doc, "此致")
        add_jingli_paragraph(self.doc, "敬礼！")

    # --------------------------------------------------------
    # 落款
    # --------------------------------------------------------

    def _add_signature(self):
        """落款：右对齐，含申请人 + 日期"""
        self.doc.add_paragraph()  # 空一行
        name = self._get("name", default="")
        apply_date = self._get("apply_date", default="")
        if name:
            add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if apply_date:
            add_right_aligned_paragraph(self.doc, apply_date)

    # --------------------------------------------------------
    # 主构建方法
    # --------------------------------------------------------

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/开头/个人基本情况/履职情况/工作成绩/不足与改进/结尾/落款"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_opening()
            self._add_basic_info()
            self._add_duty_performance()
            self._add_work_achievements()
            self._add_shortcomings_section()
            self._add_ending()
            self._add_signature()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申请书已生成：{output_path}")
        return str(output_path)

    # --------------------------------------------------------
    # 数据校验
    # --------------------------------------------------------

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []
        # honor_title 校验
        title = str(self._get("honor_title", default="优秀班干部"))
        valid_titles = ["优秀班干部", "优秀学生干部", "优秀班委",
                        "校级优秀学生干部", "院级优秀学生干部",
                        "校级优秀班干部", "院级优秀班干部"]
        if title not in valid_titles:
            warnings.append(f"评优标题 '{title}' 非标准值，按默认 '优秀班干部' 处理")
        # P0 必采字段（基础）
        for key, label in [("name", "申请人姓名"), ("college", "学院"), ("major", "专业"),
                           ("grade", "年级"), ("cadre_position", "学生干部职务"),
                           ("tenure_start", "任期开始"), ("tenure_end", "任期结束"),
                           ("gpa", "GPA"), ("rank", "专业排名")]:
            if not self._get(key):
                warnings.append(f"缺少 {label}（{key}）")
        # 学生干部履职表
        positions = self._get_list("position_history")
        if not positions:
            warnings.append("缺少 position_history 学生干部履职表（至少 1 条）")
        else:
            # 检查每条是否含量化成效
            for i, p in enumerate(positions, start=1):
                if isinstance(p, dict):
                    effect = str(p.get("effect", ""))
                    if not effect or len(effect) < 5:
                        warnings.append(f"position_history 第 {i} 条 'effect' 量化成效不足，建议含数据/荣誉/影响")
        # 履职量化清单
        duty_performances = self._get("duty_performances", default={})
        if not isinstance(duty_performances, dict) or not duty_performances:
            warnings.append("缺少 duty_performances 履职量化清单（4 维度）")
        else:
            for k, label in [("activity_count", "活动数"), ("service_hours", "服务时长"),
                             ("awards", "获奖情况"), ("peer_review", "同学评议")]:
                if not duty_performances.get(k):
                    warnings.append(f"缺少 duty_performances.{k}（{label}）")
        # 不足与改进
        if not self._get_list("shortcomings"):
            warnings.append("缺少 shortcomings 不足列表（建议 1-2 条）")
        if not self._get_list("improvements"):
            warnings.append("缺少 improvements 改进措施列表")
        # 排名校验
        rank_str = str(self._get("rank", default=""))
        if rank_str and "/" not in rank_str:
            if not self._get("rank_total"):
                warnings.append(
                    f"排名 '{rank_str}' 缺基数，应为 'X/N' 格式或补填 rank_total"
                )
        # 学业门槛校验（优秀班干部要求前 50%）
        try:
            rank_num = int(rank_str.split("/")[0]) if "/" in rank_str else 0
            rank_total_num = 0
            if "/" in rank_str:
                rank_total_num = int(rank_str.split("/")[1])
            elif self._get("rank_total"):
                rank_total_num = int(self._get("rank_total"))
            if rank_num and rank_total_num:
                pct = rank_num / rank_total_num * 100
                threshold = 50
                if pct > threshold:
                    warnings.append(
                        f"排名前 {pct:.1f}% 不满足优秀班干部前 {threshold}% 门槛"
                    )
        except (ValueError, IndexError):
            pass
        # 主干课程
        courses = self._get_list("core_courses")
        if not courses:
            warnings.append("缺少 core_courses 主干课程表，将省略课程表")
        elif len(courses) < 4:
            warnings.append(f"主干课程仅 {len(courses)} 门，建议 4~6 门")
        # 任期时长校验（满 1 学年 ≈ 8 个月）
        tenure_start = self._get("tenure_start")
        tenure_end = self._get("tenure_end")
        start_m = parse_month(tenure_start)
        end_m = parse_month(tenure_end)
        if start_m and end_m:
            diff = end_m - start_m
            if diff < 8:
                warnings.append(
                    f"任期时长 {diff} 个月，不足 1 学年（约 8 个月），请核对"
                )
        # 服务同学案例
        if not self._get_list("service_cases"):
            warnings.append("建议提供 service_cases 服务同学案例（2-3 个）")
        elif len(self._get_list("service_cases")) < 2:
            warnings.append(f"服务同学案例仅 {len(self._get_list('service_cases'))} 个，建议 2-3 个")
        self.warnings = warnings
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据（优秀班干部）
# ============================================================

DEFAULT_DATA = {
    "honor_title": "优秀班干部",
    "name": "张明", "student_id": "2022123456", "gender": "男",
    "college": "计算机科学与技术学院", "major": "计算机科学与技术",
    "grade": "2022 级 大三", "class_name": "计科 2201 班",
    "political_status": "中共预备党员", "phone": "138XXXXXXXX",
    "apply_year": "2024-2025 学年", "apply_date": "2025 年 5 月 20 日",
    "salutation": "尊敬的学院领导、评审委员会：",
    "cadre_position": "班长", "cadre_level": "班级正职",
    "position_level": "院级", "tenure_start": "2024.09", "tenure_end": "2025.05",
    "secondary_position": "团支部副书记",
    "gpa": "3.65/4.0", "weighted_avg": "82.5", "rank": "30/87", "rank_total": "87",
    "course_count": "12", "high_score_count": "5",
    "core_courses": [
        {"name": "数据结构", "credit": "4", "score": "85"},
        {"name": "操作系统", "credit": "4", "score": "82"},
        {"name": "计算机网络", "credit": "3", "score": "80"},
        {"name": "数据库原理", "credit": "3", "score": "84"},
        {"name": "计算机组成原理", "credit": "4", "score": "79"},
    ],
    "party_history": "2023.09 提交入党申请书，2024.03 列为入党积极分子，2025.06 转为中共预备党员。",
    "party_activities": [
        "2024.10 主题党日活动'红色教育基地走访'",
        "2024.12 党史学习交流",
    ],
    "position_history": [
        {"role": "班长", "term": "2024.09-2025.05（1 学年）",
         "work": "组织班级主题班会 12 次、班委例会 15 次、班级活动 6 场，转发学院通知 50 余条",
         "effect": "班级加权平均分由 84.2 提升至 87.6，所在班级获评 2024 年度校级先进班集体"},
    ],
    "duty_performances": {
        "activity_count": "主题班会 12 次、班级活动 6 场、班委例会 15 次，共 33 场",
        "service_hours": "累计 320 小时（日常履职 240h、活动组织 80h）",
        "awards": "校级先进班集体（2025.05）、院级优秀学生干部（2024.11）",
        "peer_review": "满意度 94%（47/50 人评议）",
    },
    "service_cases": [
        {"background": "2024.11 班级 X 同学因家中变故情绪低落、期中考试 2 门不及格",
         "action": "组织班委集体帮扶 + 一对一谈心 3 次",
         "effect": "X 同学期末加权平均分提升 5 分、2 门课程补考通过"},
        {"background": "2025.03 班级 Y 同学出现轻度焦虑倾向",
         "action": "及时向辅导员与心理中心报告，组织宿舍长密切关注，陪同 Y 同学接受心理中心咨询 2 次",
         "effect": "Y 同学情绪逐步稳定"},
        {"background": "2024.12 宿舍 Z 与 W 因作息时间冲突产生矛盾",
         "action": "组织双方调解 2 次，制定宿舍作息公约",
         "effect": "矛盾化解"},
    ],
    "class_achievements": "班级加权平均分由 84.2 提升至 87.6（提升 3.4 分），班级挂科率由 8% 降至 2%（下降 6 个百分点），班级体测合格率 95%，班级获评 2024 年度校级先进班集体（2025.05）、校级文明班级（2025.04）。",
    "honors": [
        {"time": "2024.11", "name": "院级优秀学生干部", "level": "院级", "issuer": "XX 大学计算机学院"},
        {"time": "2025.05", "name": "校级优秀团员", "level": "校级", "issuer": "XX 大学团委"},
    ],
    "cross_position": "本学年同时担任班长与团支部副书记，统筹班级建设与团支部建设：在 2024.11 主题班会'民族团结一家亲'中联动团支部组织团员与非团员结对学习；在 2025.04 班级学习互助活动中联动团支部组织团员的'团员帮扶先锋'行动，所在班级与团支部均获评校级先进。",
    "shortcomings": [
        "班级少数同学（3-5 人）学业帮扶成效不显著，期末仍有 2 人挂科",
        "心理排查工作偏重于已知案例，对潜在危机的预防性排查不足",
    ],
    "improvements": [
        "下学年建立'一对一'长期帮扶机制，覆盖全部学业困难同学（加权平均分 <75），每月跟踪 1 次、每学期评估 1 次",
        "下学年开展全班心理排查 4 次（每学期 2 次），重点关注学业困难、经济困难、家庭变故 3 类同学，建立'心理重点关注名单'动态管理",
    ],
    "opening": "",
    "basic_info": "",
    "duty_performance": "",
    "work_achievements": "",
    "shortcomings_section": "",
    "ending": "",
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="优秀班干部 / 优秀学生干部申请书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=("示例：\n  python build.py --data data.json --out output.docx\n  python build.py --demo --out demo.docx\n\n"
                "JSON 字段定义详见 SKILL.md 第六章信息采集清单。\n"
                "honor_title: 优秀班干部 / 优秀学生干部 / 优秀班委 / 校级优秀学生干部 / 院级优秀学生干部"),
    )
    parser.add_argument("--data", type=str, default=None, help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True, help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true", help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档（默认优秀班干部）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
