#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
选调生申请书 docx 生成器 v1.0.0

格式标准：A4 / 页边距上下 2.54cm 左右 2.5cm / 正文宋体小四 1.5 倍行距首行缩进 2 字符 /
一级标题黑体三号居中 / 二级标题黑体小三左对齐 / 签字栏宋体小四右对齐。

五段结构（详见 SKILL.md §4）：
封面 / 一级标题 / 五段正文（基本情况 + 报考动机 + 能力素质 + 基层意愿 + 发展规划）/ 签字栏。

selection_type 三类：
  - central：中央选调（中组部牵头，面向中央部委，3000 字版）
  - directional：定向选调（省委组织部，面向双一流，2500 字版）
  - non_directional：非定向选调（市委组织部/省委组织部，面向省内高校，2000 字版）

word_version 三档：
  - 2000：精简版（非定向选调，2:3:3:1:1，对应 400:600:600:200:200）
  - 2500：标准版（定向选调，2:3:3:1:1，对应 500:750:750:250:250，默认）
  - 3000：详细版（中央选调，2:3:3:1:1，对应 600:900:900:300:300）

使用：
  python build.py --data data.json --out output.docx
  python build.py --demo --out demo.docx
  python build.py --demo --word-version 2000 --out demo_2000.docx
  python build.py --demo --selection-type central --out demo_central.docx

JSON 字段详见 SKILL.md §10（含 selection_type / target_province / political_status /
basic_info / academics / party_info / student_work / awards / motivation_4d /
grassroots_willingness / development_plan 等必填字段）。
"""

import argparse
import json
import os
import sys
from datetime import date
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ===== 字体与格式常量 =====

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)
SIZE_XIAO_ER = Pt(18)
SIZE_SAN = Pt(16)
SIZE_XIAO_SAN = Pt(15)
SIZE_SI = Pt(14)
SIZE_XIAO_SI = Pt(12)
SIZE_WU = Pt(10.5)
SIZE_XIAO_WU = Pt(9)

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

# 选调类型枚举
SELECTION_CENTRAL = "central"
SELECTION_DIRECTIONAL = "directional"
SELECTION_NON_DIRECTIONAL = "non_directional"
VALID_SELECTION_TYPES = {
    SELECTION_CENTRAL,
    SELECTION_DIRECTIONAL,
    SELECTION_NON_DIRECTIONAL,
}

SELECTION_TYPE_LABEL = {
    SELECTION_CENTRAL: "中央选调",
    SELECTION_DIRECTIONAL: "定向选调",
    SELECTION_NON_DIRECTIONAL: "非定向选调",
}

# 字数版本配置（详见 SKILL.md §2.1 三档版本总表）
# key: 版本号 / value: dict(total=总字数, ratios=五段比例, ranges=各段字数区间,
#                          target=适用招录类型)
WORD_VERSION_CONFIG: Dict[int, Dict[str, Any]] = {
    2000: {
        "total": 2000,
        "ratios": "2:3:3:1:1",
        "ranges": [400, 600, 600, 200, 200],
        "target": "非定向选调、字数限制严格的岗位",
        "tolerance": (1800, 2200),
    },
    2500: {
        "total": 2500,
        "ratios": "2:3:3:1:1",
        "ranges": [500, 750, 750, 250, 250],
        "target": "定向选调、多数省份选调公告",
        "tolerance": (2300, 2700),
    },
    3000: {
        "total": 3000,
        "ratios": "2:3:3:1:1",
        "ranges": [600, 900, 900, 300, 300],
        "target": "中央选调、要求详细陈述的省份",
        "tolerance": (2800, 3200),
    },
}

# 4 维度动机字数配置（政治立场 / 服务基层 / 事业追求 / 家乡情怀）
MOTIVATION_4D_RATIOS: Dict[int, List[int]] = {
    2000: [150, 150, 150, 100],
    2500: [200, 200, 200, 150],
    3000: [250, 250, 250, 150],
}

# 5 项能力素质字数配置（政治 / 学习 / 组织 / 群众 / 抗压）
CAPABILITY_5_RATIOS: Dict[int, List[int]] = {
    2000: [120, 120, 120, 120, 120],
    2500: [150, 150, 150, 150, 150],
    3000: [180, 180, 180, 180, 180],
}

# 禁用句检测列表（详见 SKILL.md §8 雷区 10 条）
# 注：'镀金'/'跳板'在否定句中合法（如"选调生不是'镀金'的跳板"），
#     故不加入自动检测，仅由人工按 SKILL.md §8.1 把关。
FORBIDDEN_PHRASES = [
    "备胎", "曲线救国",
    "考不上研", "考不上公务员", "国考太难", "保底",
    "3年回省直", "5年调中央", "基层锻炼后回部委",
    "倍感荣幸", "深感自豪", "梦寐以求",
]

# 政治规范表述检查
REQUIRED_POLITICAL_PHRASES_CHECK = {
    "习近平新时代中国特色社会主义思想": "禁止简写为'习近平思想'等",
}


# ===== 默认 demo 数据 =====

DEFAULT_DATA: Dict[str, Any] = {
    "version": "1.0.0",
    "selection_type": "directional",
    "target_province": "山东省",
    "word_version": 2500,
    "basic_info": {
        "name": "张三",
        "gender": "男",
        "ethnicity": "汉族",
        "birth_year": 2001,
        "birth_month": 5,
        "native_place": "山东省济南市",
        "political_status": "中共党员",
        "party_join_date": "2022-05-15",
        "university": "山东大学",
        "college": "政治学与公共管理学院",
        "major": "政治学与行政学",
        "degree": "本科",
        "enrollment_year": 2020,
        "graduation_year": 2024,
    },
    "academics": {
        "gpa": 3.85,
        "gpa_total": 4.0,
        "rank": 3,
        "rank_total": 120,
        "english_level": "CET-6",
        "english_score": 580,
        "core_courses": ["政治学原理", "公共管理学", "中国政府与政治", "行政法学"],
        "minor_degree": "法学",
    },
    "party_info": {
        "party_branch": "政治学与行政学专业本科生党支部",
        "party_position": "组织委员",
        "party_age_years": 2,
        "party_training": ["初级党校", "中级党校", "高级党校"],
        "study_notes_count": 12,
        "theory_exam_rank": "前5%",
    },
    "student_work": [
        {
            "position": "校学生会主席",
            "organization": "校学生会",
            "level": "校级",
            "start_date": "2023-09",
            "end_date": "2024-06",
            "main_achievements": [
                "牵头组织'青春心向党'大型校园活动，参与人次1500人",
                "推动学生会制度改革，建立议事公开机制",
            ],
        },
        {
            "position": "班长",
            "organization": "政治学与行政学专业2020级1班",
            "level": "班级",
            "start_date": "2020-09",
            "end_date": "2024-06",
            "main_achievements": [
                "所在班级获评校级先进班集体",
                "建立班级议事制度，组织班级集体活动32次",
            ],
        },
    ],
    "awards": [
        {"name": "国家奖学金", "level": "国家级", "year": 2023, "ranking": "专业第1"},
        {"name": "校级一等奖学金", "level": "校级", "year": 2022, "ranking": "连续3年"},
        {"name": "校级三好学生", "level": "校级", "year": 2023, "ranking": ""},
        {"name": "校级优秀学生干部", "level": "校级", "year": 2023, "ranking": ""},
    ],
    "social_practice": [
        {
            "name": "三下乡-山东省临沂市沂南县支教",
            "start_date": "2022-07",
            "end_date": "2022-08",
            "role": "队长",
            "achievements": "支教40天，服务学生120人，撰写调研报告获校级优秀成果奖",
        },
        {
            "name": "三下乡-山东省菏泽市曹县乡村振兴调研",
            "start_date": "2023-07",
            "end_date": "2023-08",
            "role": "队员",
            "achievements": "走访农户80户，撰写调研报告3万字",
        },
    ],
    "motivation_4d": {
        "political_standpoint": (
            "作为一名中共党员，我始终把政治建设摆在首位。大学四年，我系统学习了马克思列宁主义、"
            "毛泽东思想、邓小平理论、'三个代表'重要思想、科学发展观、习近平新时代中国特色社会主义思想，"
            "参加党校初级班、中级班、高级班培训，撰写学习心得12篇，理论考试成绩均位列学院前5%。"
            "我自觉增强'四个意识'、坚定'四个自信'、做到'两个维护'。"
            "我深刻认识到，选调生是党政领导干部后备人选，必须由政治立场坚定、对党绝对忠诚、"
            "理论素养扎实的青年党员担当。我希望通过选调生这一渠道，"
            "将个人政治信仰转化为服务党的事业、服务基层群众的实际行动。"
        ),
        "grassroots_service": (
            "基层是国家治理的'最后一公里'，是政策落地的关键环节，是青年成长最好的课堂。"
            "大学期间，我先后两次参加'三下乡'社会实践，深入山东省临沂市沂南县和菏泽市曹县开展支教和乡村振兴调研，"
            "亲眼见证了脱贫攻坚和乡村振兴给农村带来的巨大变化，"
            "也深刻体会到基层干部工作的艰辛与价值。我曾跟随驻村第一书记走访农户80户，"
            "撰写调研报告获校级社会实践优秀成果奖。这些经历让我真正理解了"
            "'纸上得来终觉浅，绝知此事要躬行'的深意。我愿扎根基层，从乡镇、村、街道一级做起，"
            "在解决群众急难愁盼问题中践行共产党员的初心使命。"
        ),
        "career_pursuit": (
            "我深知选调生不是'镀金'的跳板，而是党和国家培养青年干部的重要渠道。"
            "这一职业要求既要有基层干部的'泥土味'，又要有党政干部的'政治性'，"
            "既要做实干家，又要有政治家视野。这一职业定位与我对个人事业的追求高度契合："
            "我希望成为一名既有理论素养又有实践能力、既懂宏观政策又能解决具体问题、"
            "既能扎根基层又能胸怀大局的党政干部，"
            "在推动国家治理体系和治理能力现代化中实现个人价值。"
            "我将以长期主义心态对待选调生事业，不图短期得失，不求一时显达，"
            "扎根基层、久久为功，把青春和热血奉献给党和人民的事业。"
        ),
        "hometown_attachment": (
            "我是山东省济南市人，家乡地处鲁中南地区，是革命老区、农业大市、转型发展重点区域。"
            "大学四年我始终关注家乡发展，特别是沂南县乡村振兴和基层治理工作，每年寒暑假都主动返乡调研。"
            "毕业后回到家乡、扎根家乡、建设家乡，是我一直以来的心愿。"
            "我希望以定向选调生身份回到山东省，将所学所悟用于家乡建设，"
            "为家乡的现代化事业贡献青年力量。"
        ),
    },
    "grassroots_willingness": {
        "accept_rural": True,
        "accept_remote": True,
        "accept_cross_city": True,
        "service_years_commitment": 5,
        "psychological_readiness": "对艰苦环境有充分心理准备和应对预案",
    },
    "development_plan": {
        "first_two_years": (
            "前两年扎根基层，深入学习基层治理、群众工作、政策执行等基础业务，"
            "掌握第一手民情民意，做一名合格的基层干部，"
            "重点提升调查研究、群众沟通、矛盾调解三项核心能力。"
        ),
        "middle_two_years": (
            "中间两年成长提升，在熟悉业务基础上独立承担重点工作，"
            "争取在乡村振兴、基层党建、产业发展等某一领域形成专长，逐步成为单位业务骨干。"
        ),
        "fifth_year": (
            "第五年担当作为，力争成为本部门或本单位的业务骨干，"
            "能够独当一面处理复杂问题，并为后续发展奠定坚实基础。"
        ),
    },
    "policy_references": [
        {
            "name": "关于做好选调生工作的意见",
            "issuer": "中组部",
            "year": 2014,
            "quoted_in": "报考动机段",
        },
        {
            "name": "关于进一步加强和改进选调生工作的意见",
            "issuer": "中组部",
            "year": 2018,
            "quoted_in": "报考动机段",
        },
    ],
    "signature": {
        "applicant_name": "张三",
        "applicant_date": "2024-09-15",
        "college_reviewer": "李四（学院党委副书记）",
        "college_review_date": "2024-09-16",
        "college_seal": "（学院党委盖章）",
    },
}


# ===== 工具函数 =====

def set_cell_font(cell, text: str, font_name: str = FONT_SONG,
                  font_size: Any = SIZE_XIAO_WU, bold: bool = False,
                  align: int = WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格字体、字号、对齐方式。"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, font_name: str = FONT_SONG, font_size: Any = SIZE_XIAO_SI,
                 bold: bool = False) -> None:
    """设置 run 字体（中英文分别处理）。"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def set_page_format(doc: Document) -> None:
    """设置 A4 纸张 + 页边距（上下 2.54cm，左右 2.5cm）。"""
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
        section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
        section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
        section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def set_paragraph_format(p, first_line_indent: bool = True,
                         line_spacing: float = 1.5,
                         align: int = WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_before: float = 0,
                         space_after: float = 0) -> None:
    """设置段落格式（首行缩进 2 字符 / 行距 / 对齐）。"""
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(SIZE_XIAO_SI.pt * 2)
    pf.line_spacing = line_spacing
    p.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """添加标题（一级黑体三号居中，二级黑体小三左对齐）。"""
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, FONT_HEI, SIZE_SAN, bold=True)
        set_paragraph_format(p, first_line_indent=False,
                             align=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=12, space_after=6)
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, FONT_HEI, SIZE_XIAO_SAN, bold=True)
        set_paragraph_format(p, first_line_indent=False,
                             align=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=6, space_after=3)
    else:
        run = p.add_run(text)
        set_run_font(run, FONT_HEI, SIZE_SI, bold=True)
        set_paragraph_format(p, first_line_indent=False,
                             align=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=3, space_after=3)


def add_body_paragraph(doc: Document, text: str, indent: bool = True) -> None:
    """添加正文段落（宋体小四 1.5 倍行距首行缩进 2 字符）。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, FONT_SONG, SIZE_XIAO_SI, bold=False)
    set_paragraph_format(p, first_line_indent=indent,
                         line_spacing=1.5,
                         align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def add_signature_paragraph(doc: Document, text: str) -> None:
    """添加签字段落（宋体小四右对齐）。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, FONT_SONG, SIZE_XIAO_SI, bold=False)
    set_paragraph_format(p, first_line_indent=False,
                         line_spacing=1.5,
                         align=WD_ALIGN_PARAGRAPH.RIGHT,
                         space_before=6, space_after=6)


# ===== 内容生成函数 =====

def _format_basic_info_paragraph(data: Dict[str, Any]) -> str:
    """生成段一：个人基本情况正文。"""
    bi = data["basic_info"]
    ac = data.get("academics", {})
    pi = data.get("party_info", {})
    sw_list = data.get("student_work", [])
    awards_list = data.get("awards", [])

    name = bi["name"]
    gender = bi["gender"]
    ethnicity = bi.get("ethnicity", "汉族")
    birth_year = bi.get("birth_year", "")
    birth_month = bi.get("birth_month", "")
    native = bi.get("native_place", "")
    political = bi.get("political_status", "中共党员")
    party_date = bi.get("party_join_date", "")
    university = bi.get("university", "")
    college = bi.get("college", "")
    major = bi.get("major", "")
    degree = bi.get("degree", "本科")
    grad_year = bi.get("graduation_year", "")

    # 学业信息
    gpa = ac.get("gpa", "")
    gpa_total = ac.get("gpa_total", 4.0)
    rank = ac.get("rank", "")
    rank_total = ac.get("rank_total", "")
    english_level = ac.get("english_level", "")
    english_score = ac.get("english_score", "")

    # 主要奖项
    award_names = []
    for aw in awards_list[:3]:
        award_names.append(aw.get("name", ""))
    awards_str = "、".join([n for n in award_names if n])

    # 学生干部
    sw_desc = ""
    if sw_list:
        sw0 = sw_list[0]
        sw_desc = (
            f"担任{sw0.get('organization', '')}{sw0.get('position', '')}"
            f"（{sw0.get('start_date', '')}至{sw0.get('end_date', '')}）"
        )
        if sw0.get("main_achievements"):
            sw_desc += f"，{sw0['main_achievements'][0]}"

    sw_extra = ""
    if len(sw_list) > 1:
        sw1 = sw_list[1]
        sw_extra = (
            f"同时担任{sw1.get('organization', '')}{sw1.get('position', '')}"
        )

    # 拼接段落
    parts = []
    parts.append(
        f"我叫{name}，{gender}，{ethnicity}，{birth_year}年{birth_month}月生，"
        f"{native}人，{political}（{party_date}入党），"
        f"现为{university}{college}{major}专业应届{degree}毕业生。"
    )
    parts.append(
        f"在校期间学业成绩优异，前三年GPA {gpa}/{gpa_total}，"
        f"专业排名第{rank}/{rank_total}，{english_level} {english_score}分，"
        f"曾获{awards_str}等荣誉。"
    )
    if sw_desc:
        parts.append(sw_desc + "。")
    if sw_extra:
        parts.append(sw_extra + "，全面锻炼了组织协调与群众工作能力。")

    return "".join(parts)


def _format_motivation_paragraph(data: Dict[str, Any]) -> List[Tuple[str, str]]:
    """生成段二：报考动机正文，返回 [(二级标题, 正文), ...]。"""
    m4d = data["motivation_4d"]
    selection_type = data.get("selection_type", "directional")
    type_label = SELECTION_TYPE_LABEL.get(selection_type, "选调")

    intro = (
        f"我报考{type_label}生，是基于政治立场、服务基层、事业追求、家乡情怀"
        "四方面综合考量的理性选择。"
    )

    sections = [
        ("政治立场方面", m4d.get("political_standpoint", "")),
        ("服务基层方面", m4d.get("grassroots_service", "")),
        ("事业追求方面", m4d.get("career_pursuit", "")),
        ("家乡情怀方面", m4d.get("hometown_attachment", "")),
    ]
    return [("intro", intro)] + sections


def _format_capability_paragraph(data: Dict[str, Any]) -> List[Tuple[str, str]]:
    """生成段三：个人能力素质正文，返回 [(二级标题, 正文), ...]。"""
    bi = data.get("basic_info", {})
    ac = data.get("academics", {})
    pi = data.get("party_info", {})
    sw_list = data.get("student_work", [])

    # 政治素质
    political_quality = (
        f"作为中共党员，我始终把政治学习摆在首位，认真参加'三会一课'、主题党日、"
        f"组织生活会等党内政治生活，理论修养不断提高。"
        f"撰写学习心得{pi.get('study_notes_count', 12)}篇，"
        f"在学院'两学一做'知识竞赛中获一等奖。"
        f"我能够自觉用习近平新时代中国特色社会主义思想武装头脑，"
        f"在重大原则问题上立场坚定，在大是大非面前旗帜鲜明。"
    )

    # 学习能力
    gpa = ac.get("gpa", "")
    gpa_total = ac.get("gpa_total", 4.0)
    rank = ac.get("rank", "")
    rank_total = ac.get("rank_total", "")
    core_courses = ac.get("core_courses", [])
    core_courses_str = "、".join(core_courses[:3]) if core_courses else "主修核心课程"
    minor = ac.get("minor_degree", "")
    learning = (
        f"大学四年，我保持了严谨求实的学习态度，前三年GPA {gpa}/{gpa_total}，"
        f"专业排名第{rank}/{rank_total}，主修的{core_courses_str}等核心课程成绩均在90分以上。"
    )
    if minor:
        learning += f"同时辅修{minor}双学位，拓宽了知识结构。"
    learning += (
        "具备较强的文字综合能力，撰写论文获校级学术论坛一等奖，"
        "曾在校报、学院网站发表文章多篇。"
    )

    # 组织协调
    sw0 = sw_list[0] if sw_list else {}
    achievements = sw0.get("main_achievements", [])
    achievement_str = achievements[0] if achievements else "组织策划多项校园活动"
    organization = (
        f"担任{sw0.get('organization', '校学生会')}{sw0.get('position', '干部')}期间，"
        f"牵头组织大型校园活动，{achievement_str}，"
        f"组织策划能力、团队协作能力、应急处理能力得到全面锻炼，"
        f"能够胜任基层党政工作的组织协调任务。"
    )

    # 群众工作
    sw_other = sw_list[1] if len(sw_list) > 1 else {}
    sw_other_str = ""
    if sw_other:
        sw_other_str = (
            f"担任{sw_other.get('organization', '')}{sw_other.get('position', '')}"
        )
        other_achievements = sw_other.get("main_achievements", [])
        if other_achievements:
            sw_other_str += f"，{other_achievements[0]}"
    mass_work = (
        f"担任班长四年，我坚持'以同学为本'，建立班级议事制度，组织班级集体活动多次。"
        f"{sw_other_str}。" if sw_other_str else
        f"担任班长四年，我坚持'以同学为本'，建立班级议事制度，组织班级集体活动多次。"
    )
    mass_work += (
        "深入学生群体解决实际问题，积累了较为丰富的群众工作经验。"
    )

    # 抗压能力
    pressure = (
        "大学期间，我同时承担学业、学生工作、社会实践三重任务，"
        "最忙时曾连续3个月每天工作学习14小时以上。"
        "在实习期间独立完成调研报告，加班加点保质保量完成任务。"
        "我能够承受基层工作的强度与压力，把艰苦环境作为成长的磨刀石。"
    )

    return [
        ("政治素质方面", political_quality),
        ("学习能力方面", learning),
        ("组织协调方面", organization),
        ("群众工作方面", mass_work),
        ("抗压能力方面", pressure),
    ]


def _format_grassroots_paragraph(data: Dict[str, Any]) -> str:
    """生成段四：基层工作意愿正文。"""
    gw = data.get("grassroots_willingness", {})
    service_years = gw.get("service_years_commitment", 5)
    accept_rural = gw.get("accept_rural", True)
    accept_remote = gw.get("accept_remote", True)

    rural_clause = "无论是城市街道还是农村乡镇" if accept_rural else "主要接受城市街道工作"
    remote_clause = "无论是经济发达地区还是艰苦边远地区" if accept_remote else ""

    return (
        "我郑重承诺：自愿到基层工作，服从组织分配，"
        "接受到乡镇、村、街道等基层单位任职。"
        f"{remote_clause}，{rural_clause}，"
        "我都将以平常心对待，扎根基层、安心工作。"
        "我深知基层条件相对艰苦、工作强度大、生活节奏快，"
        "远离家人、远离都市，对此我有充分的心理准备和应对预案。"
        f"我承诺在基层工作满{service_years}年服务期，"
        "期间不调动、不跳槽、不向组织提非分要求，"
        "把基层作为人生事业的起点和根基，与基层群众同甘共苦、共同成长。"
    )


def _format_development_paragraph(data: Dict[str, Any]) -> str:
    """生成段五：个人发展规划正文。"""
    dp = data.get("development_plan", {})
    first_two = dp.get("first_two_years", "")
    middle_two = dp.get("middle_two_years", "")
    fifth_year = dp.get("fifth_year", "")

    return (
        f"关于未来五年发展规划，我的设想是：{first_two}"
        f"{middle_two}{fifth_year}"
        "我将始终以'功成不必在我'的境界和'功成必定有我'的担当，"
        "扎根基层、奉献青春，为党和人民的事业贡献毕生力量，"
        "不辜负党组织培养和人民期望。"
    )


# ===== 文档构建主流程 =====

def build_document(data: Dict[str, Any]) -> Document:
    """根据 data 构建选调生申请书 docx。"""
    doc = Document()
    set_page_format(doc)

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = FONT_SONG
    style.font.size = SIZE_XIAO_SI
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_SONG)

    selection_type = data.get("selection_type", "directional")
    type_label = SELECTION_TYPE_LABEL.get(selection_type, "选调")

    # ===== 一级标题 =====
    add_heading(doc, f"{type_label}生申请书", level=1)

    # ===== 段一：个人基本情况 =====
    add_heading(doc, "一、个人基本情况", level=2)
    para1 = _format_basic_info_paragraph(data)
    add_body_paragraph(doc, para1)

    # ===== 段二：报考选调生动机 =====
    add_heading(doc, "二、报考选调生动机", level=2)
    motivation_sections = _format_motivation_paragraph(data)
    for title, content in motivation_sections:
        if title == "intro":
            add_body_paragraph(doc, content)
        else:
            add_heading(doc, title, level=2)
            add_body_paragraph(doc, content)

    # ===== 段三：个人能力素质 =====
    add_heading(doc, "三、个人能力素质", level=2)
    capability_sections = _format_capability_paragraph(data)
    for title, content in capability_sections:
        add_heading(doc, title, level=2)
        add_body_paragraph(doc, content)

    # ===== 段四：基层工作意愿 =====
    add_heading(doc, "四、基层工作意愿", level=2)
    para4 = _format_grassroots_paragraph(data)
    add_body_paragraph(doc, para4)

    # ===== 段五：个人发展规划 =====
    add_heading(doc, "五、个人发展规划", level=2)
    para5 = _format_development_paragraph(data)
    add_body_paragraph(doc, para5)

    # ===== 签字栏 =====
    sig = data.get("signature", {})
    applicant_name = sig.get("applicant_name", "")
    applicant_date = sig.get("applicant_date", "")
    college_reviewer = sig.get("college_reviewer", "")
    college_review_date = sig.get("college_review_date", "")
    college_seal = sig.get("college_seal", "")

    # 空行
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=False, line_spacing=1.5)

    add_signature_paragraph(doc, f"申请人：{applicant_name}")
    add_signature_paragraph(doc, f"日期：{applicant_date}")
    if college_reviewer:
        add_signature_paragraph(doc, f"学院审核：{college_reviewer}")
        add_signature_paragraph(doc, f"审核日期：{college_review_date}")
    if college_seal:
        add_signature_paragraph(doc, college_seal)

    return doc


# ===== 数据校验 =====

def _validate_data(data: Dict[str, Any]) -> List[str]:
    """校验 data，返回错误信息列表。空列表表示通过。"""
    errors: List[str] = []

    # 顶层字段
    if "selection_type" not in data:
        errors.append("[字段缺失] selection_type 必填（central/directional/non_directional）")
    elif data["selection_type"] not in VALID_SELECTION_TYPES:
        errors.append(
            f"[字段错误] selection_type 必须为 central/directional/non_directional 之一，"
            f"当前为 {data['selection_type']}"
        )

    if "target_province" not in data or not data["target_province"]:
        errors.append("[字段缺失] target_province 必填（目标省份或部委）")

    if "word_version" not in data:
        errors.append("[字段缺失] word_version 必填（2000/2500/3000）")
    elif data["word_version"] not in WORD_VERSION_CONFIG:
        errors.append(
            f"[字段错误] word_version 必须为 2000/2500/3000 之一，"
            f"当前为 {data['word_version']}"
        )

    # basic_info
    bi = data.get("basic_info", {})
    if not bi:
        errors.append("[字段缺失] basic_info 必填")
    else:
        required_bi = ["name", "gender", "political_status", "party_join_date",
                       "university", "college", "major", "degree"]
        for field in required_bi:
            if field not in bi or not bi[field]:
                errors.append(f"[字段缺失] basic_info.{field} 必填")

        # 政治面貌必须为党员或预备党员
        political = bi.get("political_status", "")
        if political and political not in ("中共党员", "中共预备党员"):
            errors.append(
                f"[字段错误] basic_info.political_status 必须为'中共党员'或'中共预备党员'，"
                f"当前为'{political}'（非定向选调部分岗位可放宽，请确认）"
            )

    # academics
    ac = data.get("academics", {})
    if not ac:
        errors.append("[字段缺失] academics 必填")
    else:
        if "gpa" not in ac or "rank" not in ac:
            errors.append("[字段缺失] academics.gpa 和 academics.rank 必填")

    # party_info
    pi = data.get("party_info", {})
    if not pi:
        errors.append("[字段缺失] party_info 必填（党员信息）")

    # student_work
    sw_list = data.get("student_work", [])
    if not sw_list:
        errors.append("[字段缺失] student_work 必填（至少 1 项学生干部经历）")

    # awards
    awards = data.get("awards", [])
    if not awards:
        errors.append("[字段缺失] awards 必填（至少 1 项校级以上奖励）")

    # motivation_4d（4 维度缺一不可）
    m4d = data.get("motivation_4d", {})
    if not m4d:
        errors.append("[字段缺失] motivation_4d 必填")
    else:
        required_dims = ["political_standpoint", "grassroots_service",
                         "career_pursuit", "hometown_attachment"]
        for dim in required_dims:
            content = m4d.get(dim, "")
            if not content or len(content) < 50:
                errors.append(
                    f"[字段不足] motivation_4d.{dim} 内容过短（应 ≥ 100 字），"
                    f"当前 {len(content)} 字"
                )

    # grassroots_willingness
    gw = data.get("grassroots_willingness", {})
    if not gw:
        errors.append("[字段缺失] grassroots_willingness 必填")

    # development_plan
    dp = data.get("development_plan", {})
    if not dp:
        errors.append("[字段缺失] development_plan 必填")

    # signature
    sig = data.get("signature", {})
    if not sig or "applicant_name" not in sig:
        errors.append("[字段缺失] signature.applicant_name 必填")

    return errors


def _check_forbidden_phrases(data: Dict[str, Any]) -> List[str]:
    """检测禁用句，返回命中列表。"""
    hits: List[str] = []

    def _check_text(text: str, source: str) -> None:
        if not text:
            return
        for phrase in FORBIDDEN_PHRASES:
            if phrase in text:
                hits.append(f"[禁用句] {source} 命中'{phrase}'")

    m4d = data.get("motivation_4d", {})
    _check_text(m4d.get("political_standpoint", ""), "动机-政治立场")
    _check_text(m4d.get("grassroots_service", ""), "动机-服务基层")
    _check_text(m4d.get("career_pursuit", ""), "动机-事业追求")
    _check_text(m4d.get("hometown_attachment", ""), "动机-家乡情怀")

    dp = data.get("development_plan", {})
    _check_text(dp.get("first_two_years", ""), "规划-前两年")
    _check_text(dp.get("middle_two_years", ""), "规划-中间两年")
    _check_text(dp.get("fifth_year", ""), "规划-第五年")

    _check_text(_format_basic_info_paragraph(data), "段一基本情况")
    _check_text(_format_grassroots_paragraph(data), "段四基层意愿")

    return hits


def _check_grassroots_keyword(data: Dict[str, Any]) -> List[str]:
    """检查'基层'一词出现频次（应 ≥ 5 次）。"""
    warnings: List[str] = []

    all_text = ""
    all_text += _format_basic_info_paragraph(data)
    m4d = data.get("motivation_4d", {})
    all_text += m4d.get("political_standpoint", "")
    all_text += m4d.get("grassroots_service", "")
    all_text += m4d.get("career_pursuit", "")
    all_text += m4d.get("hometown_attachment", "")
    cap_sections = _format_capability_paragraph(data)
    for _, content in cap_sections:
        all_text += content
    all_text += _format_grassroots_paragraph(data)
    all_text += _format_development_paragraph(data)

    count = all_text.count("基层")
    if count < 5:
        warnings.append(
            f"[警告] '基层'一词仅出现 {count} 次（应 ≥ 5 次），"
            "请检查 §8 雷区 4"
        )
    return warnings


# ===== 主入口 =====

def main() -> int:
    """命令行入口。"""
    parser = argparse.ArgumentParser(
        description="选调生申请书 docx 生成器 v1.0.0",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --demo --out demo.docx\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --word-version 2000 --out demo_2000.docx\n"
            "  python build.py --demo --selection-type central --out demo_central.docx\n"
        ),
    )
    parser.add_argument("--demo", action="store_true",
                        help="使用内置 DEFAULT_DATA 生成 demo 文档")
    parser.add_argument("--data", type=str, default=None,
                        help="输入 JSON 文件路径")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--word-version", type=int, default=None,
                        choices=[2000, 2500, 3000],
                        help="覆盖字数版本（2000/2500/3000）")
    parser.add_argument("--selection-type", type=str, default=None,
                        choices=["central", "directional", "non_directional"],
                        help="覆盖选调类型（central/directional/non_directional）")
    parser.add_argument("--skip-validation", action="store_true",
                        help="跳过数据校验（仅用于调试）")

    args = parser.parse_args()

    # 加载数据
    if args.demo:
        data = json.loads(json.dumps(DEFAULT_DATA))  # deep copy
        print("[INFO] 使用内置 DEFAULT_DATA 生成 demo 文档")
    elif args.data:
        data_path = Path(args.data)
        if not data_path.exists():
            print(f"[ERROR] 数据文件不存在：{data_path}")
            return 1
        try:
            with open(data_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            print(f"[INFO] 已加载数据文件：{data_path}")
        except json.JSONDecodeError as e:
            print(f"[ERROR] JSON 解析失败：{e}")
            return 1
    else:
        print("[ERROR] 必须指定 --demo 或 --data 参数")
        parser.print_help()
        return 1

    # 覆盖参数
    if args.word_version is not None:
        data["word_version"] = args.word_version
        print(f"[INFO] 覆盖 word_version = {args.word_version}")
    if args.selection_type is not None:
        data["selection_type"] = args.selection_type
        print(f"[INFO] 覆盖 selection_type = {args.selection_type}")

    # 数据校验
    if not args.skip_validation:
        errors = _validate_data(data)
        if errors:
            print("[ERROR] 数据校验失败：")
            for err in errors:
                print(f"  {err}")
            return 1
        print("[INFO] 数据校验通过")

        # 禁用句检测
        forbidden_hits = _check_forbidden_phrases(data)
        if forbidden_hits:
            print("[ERROR] 禁用句检测命中（详见 SKILL.md §8 雷区 10 条）：")
            for hit in forbidden_hits:
                print(f"  {hit}")
            return 1
        print("[INFO] 禁用句检测通过")

        # 基层关键词检查
        grassroots_warnings = _check_grassroots_keyword(data)
        for warn in grassroots_warnings:
            print(f"[WARN] {warn}")

    # 生成文档
    try:
        doc = build_document(data)
    except Exception as e:
        print(f"[ERROR] 文档生成失败：{e}")
        return 1

    # 输出
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(out_path))
        print(f"[INFO] 文档已生成：{out_path}")
    except Exception as e:
        print(f"[ERROR] 文档保存失败：{e}")
        return 1

    # 字数统计
    word_version = data.get("word_version", 2500)
    config = WORD_VERSION_CONFIG.get(word_version, {})
    print(f"[INFO] 字数版本：{word_version} 字（{config.get('target', '')}）")
    print(f"[INFO] 五段比例：{config.get('ratios', '')}")
    print(f"[INFO] 各段字数区间：{config.get('ranges', '')}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
