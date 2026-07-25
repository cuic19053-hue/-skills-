#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
国家励志奖学金申请书 docx 生成器

格式标准：A4 纸张，页边距上下 2.54cm 左右 2.5cm；标题黑体二号居中；
称呼顶格宋体小四全角冒号；正文宋体小四 1.5 倍行距首行缩进 2 字符；
主干课程表/家庭成员表/勤工助学表宋体五号居中；
"此致"另起一行空两格，"敬礼！"另起一行顶格；落款右对齐。

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第三章信息采集清单。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进），宋体小四"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_section_heading(doc, text: str):
    """正文中的小节标题（一、二、三…）：黑体小四加粗，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
        line_spacing=1.5, space_before=6, space_after=3)


def add_cizhi_paragraph(doc, text: str = "此致"):
    """'此致'另起一行，空两格（首行缩进 2 字符）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格（不缩进）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=False, line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None,
                        caption: str = ""):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）

    Args:
        caption: 表格上方说明文字（如"主干课程成绩："），不写则不添加
    """
    if caption:
        add_paragraph_with_format(
            doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)

    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """国家励志奖学金申请书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段，非列表返回空列表"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_heading(self, text):
        return add_section_heading(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, caption=""):
        return add_table_from_data(self.doc, headers, rows, col_widths, caption)

    # 标题

    def _add_title(self):
        """标题：黑体二号居中，固定为"国家励志奖学金申请书"9 字"""
        add_title(self.doc, "国家励志奖学金申请书")

    # 称呼

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号"""
        salutation = self._get("salutation",
                               default="尊敬的学院领导、评审委员会：")
        add_salutation_paragraph(self.doc, salutation)

    # 开头段落

    def _add_opening(self):
        """开头段落（100~150 字）：身份 + 申报奖项 + 困难认定 + 学业核心数据 + 进入正文"""
        opening = self._get("opening", default="")
        if opening:
            self.add_para(opening)
            return
        name = self._get("name")
        college = self._get("college")
        major = self._get("major")
        grade = self._get("grade")
        apply_year = self._get("apply_year", default="2024-2025 学年")
        difficulty_level = self._get("difficulty_level", default="")
        gpa = self._get("gpa")
        rank = self._get("rank")
        parts = []
        if name and college and major and grade:
            parts.append(f"我是{college}{major}{grade}学生{name}，"
                         f"特申请{apply_year}国家励志奖学金。")
        else:
            parts.append(f"特申请{apply_year}国家励志奖学金。")
        data_parts = []
        if difficulty_level:
            data_parts.append(f"已获学校家庭经济困难认定（{difficulty_level}等级）")
        if gpa:
            data_parts.append(f"GPA {gpa}")
        if rank:
            data_parts.append(f"专业排名第 {rank}")
        if data_parts:
            parts.append("；".join(data_parts) + "。")
        parts.append("现将本人情况汇报如下：")
        self.add_para("".join(parts))

    # 思想方面

    def _add_ideology(self):
        """思想方面（150~250 字）：政治立场 + 入党/团情况 + 思想觉悟 + 具体活动"""
        self.add_heading("一、思想方面")
        ideology = self._get("ideology", default="")
        if ideology:
            if isinstance(ideology, list):
                for p in ideology:
                    self.add_para(p)
            else:
                self.add_para(ideology)
            return
        political = self._get("political_status", default="共青团员")
        party_history = self._get("party_history", default="")
        party_activities = self._get_list("party_activities")
        parts = []
        if political:
            if "党员" in political:
                parts.append(f"作为一名{political}，我始终坚持学习党的理论，"
                             "认真学习习近平新时代中国特色社会主义思想与党的二十大和二十届三中全会精神。")
            elif "积极分子" in political:
                parts.append("作为一名入党积极分子，我认真学习习近平新时代中国特色社会主义思想，"
                             "时刻以党员标准要求自己。")
            else:
                parts.append(f"作为一名{political}，我拥护中国共产党的领导，"
                             "认真学习党的创新理论。")
        if party_history:
            parts.append(party_history)
        if party_activities:
            parts.append("；".join(party_activities))
        parts.append("在日常生活中，我注重理论学习与实践结合，关注时政热点，提升思想觉悟。")
        self.add_para("".join(parts))

    # 学习方面（含主干课程表）

    def _add_academics(self):
        """学习方面（250~350 字）：GPA + 排名 + 主干课程表 + 英语计算机等级"""
        self.add_heading("二、学习方面")
        academics = self._get("academics", default="")
        if academics and isinstance(academics, str):
            self.add_para(academics)
            self._add_core_courses_table()
            self._add_academics_summary()
            return
        gpa = self._get("gpa")
        weighted = self._get("weighted_avg")
        rank = self._get("rank")
        rank_total = self._get("rank_total")
        course_count = self._get("course_count")
        high_score_count = self._get("high_score_count")
        parts = []
        if gpa:
            gpa_str = f"本学年 GPA {gpa}"
            if weighted:
                gpa_str += f"，加权平均分 {weighted}"
            if rank:
                gpa_str += f"，专业排名第 {rank}"
                if rank_total:
                    try:
                        r_num = int(str(rank).split("/")[0])
                        total_num = int(rank_total)
                        if total_num > 0:
                            pct = round(r_num / total_num * 100, 1)
                            gpa_str += f"（前 {pct}%）"
                    except (ValueError, IndexError):
                        pass
            parts.append(gpa_str + "。")
        if course_count and high_score_count:
            parts.append(f"修读 {course_count} 门课程，{high_score_count} 门 80 分以上。")
        if parts:
            self.add_para("".join(parts))
        self._add_core_courses_table()
        self._add_academics_summary()

    def _add_core_courses_table(self):
        """主干课程表：4~6 门，含课程名/学分/成绩"""
        courses = self._get_list("core_courses")
        if not courses:
            return
        rows = []
        for c in courses:
            if not isinstance(c, dict):
                continue
            rows.append([
                str(c.get("name", "")),
                str(c.get("credit", "")),
                str(c.get("score", "")),
            ])
        if rows:
            self.add_table(
                ["课程名称", "学分", "成绩"],
                rows,
                col_widths=[8.0, 3.0, 3.0],
                caption="主干课程成绩：",
            )

    def _add_academics_summary(self):
        """学习方面末尾：主干课程亮点 + 英语计算机等级 + 学习方法"""
        summary = self._get("academics_summary", default="")
        if summary:
            self.add_para(summary)
            return
        cet4 = self._get("cet4")
        cet6 = self._get("cet6")
        computer_level = self._get("computer_level")
        course_highlight = self._get("course_highlight", default="")
        study_method = self._get("study_method",
                                  default="学习上注重课前预习与课后总结，"
                                          "遇到问题主动与老师、同学讨论。")
        parts = []
        if course_highlight:
            parts.append(course_highlight + "。")
        lang_parts = []
        if cet4:
            lang_parts.append(f"CET-4 {cet4} 分")
        if cet6:
            lang_parts.append(f"CET-6 {cet6} 分")
        if lang_parts:
            parts.append("、".join(lang_parts) + "；")
        if computer_level:
            parts.append(f"计算机{computer_level}。")
        if study_method:
            parts.append(study_method)
        if parts:
            self.add_para("".join(parts))

    # 家庭经济情况【励志类核心段，含家庭成员表 + 勤工助学表】

    def _add_family_economy(self):
        """家庭经济情况（300~450 字，重点）：
        家庭基本信息 + 年收入 + 供养 + 困难认定 + 家庭成员表 + 勤工助学 + 变故"""
        self.add_heading("三、家庭经济情况")
        family_text = self._get("family_economy", default="")
        if family_text and isinstance(family_text, str):
            self.add_para(family_text)
            self._add_family_members_table()
            self._add_work_study_table()
            self._add_family_event()
            return

        hometown = self._get("hometown", default="")
        income_year = self._get("family_income_year", default="")
        income_source = self._get("family_income_source", default="")
        difficulty_level = self._get("difficulty_level", default="")
        difficulty_cert_no = self._get("difficulty_cert_no", default="")
        dependents_count = self._get("dependents_count", default="")
        dependents_detail = self._get("dependents_detail", default="")

        parts = []
        if hometown:
            seg = f"我来自{hometown}"
            if income_source:
                seg += f"，{income_source}"
            if income_year:
                seg += f"，家庭年收入约 {self._format_income(income_year)} 元"
            seg += "。"
            parts.append(seg)
        if dependents_count:
            seg = f"家庭需供养 {dependents_count} 人"
            if dependents_detail:
                seg += f"：{dependents_detail}"
            seg += "。"
            parts.append(seg)
        if difficulty_level:
            seg = f"本人已获学校家庭经济困难认定（{difficulty_level}等级）"
            if difficulty_cert_no:
                seg += f"，{difficulty_cert_no}"
            seg += "。"
            parts.append(seg)
        if parts:
            self.add_para("".join(parts))

        self._add_family_members_table()
        self._add_work_study_table()
        self._add_family_event()

    def _format_income(self, income) -> str:
        """格式化家庭年收入：整数元或万元"""
        s = str(income).strip()
        if not s:
            return ""
        # 5 位数以上且为整数，转换为万元表述
        try:
            n = float(s)
            if n >= 10000:
                return f"{n / 10000:.1f} 万"
            return f"{int(n) if n == int(n) else n}"
        except ValueError:
            return s

    def _add_family_members_table(self):
        """家庭成员表：5 列（姓名/关系/工作/年收入/政治面貌）"""
        members = self._get_list("family_members")
        if not members:
            return
        rows = []
        for m in members:
            if not isinstance(m, dict):
                continue
            income_val = m.get("income", "")
            income_str = str(income_val) if income_val else "-"
            rows.append([
                str(m.get("name", "")),
                str(m.get("relation", "")),
                str(m.get("job", "")),
                income_str,
                str(m.get("political", "")),
            ])
        if rows:
            self.add_table(
                ["姓名", "与本人关系", "工作单位/职务", "年收入（元）", "政治面貌"],
                rows,
                col_widths=[2.0, 2.5, 4.5, 2.5, 2.0],
                caption="家庭成员情况：",
            )

    def _add_work_study_table(self):
        """勤工助学经历表：4 列（岗位/地点/时间/月收入）"""
        work_study = self._get_list("work_study")
        if not work_study:
            return
        # 若用户提供的勤工助学为字符串列表，直接转表格行
        rows = []
        for w in work_study:
            if isinstance(w, dict):
                rows.append([
                    str(w.get("position", "")),
                    str(w.get("place", "")),
                    str(w.get("duration", "")),
                    str(w.get("income", "")),
                ])
            elif isinstance(w, str):
                rows.append([w, "", "", ""])
        if rows:
            self.add_table(
                ["岗位", "地点", "时间", "月收入"],
                rows,
                col_widths=[4.0, 3.5, 4.0, 2.5],
                caption="勤工助学经历：",
            )
        # 勤工助学小结
        work_summary = self._get("work_study_summary", default="")
        if work_summary:
            self.add_para(work_summary)
        elif rows:
            self.add_para("勤工助学收入用于覆盖教材费与生活日用品支出，减轻家庭负担。")

    def _add_family_event(self):
        """家庭重大变故（如有）"""
        family_event = self._get("family_event", default="")
        if family_event:
            self.add_para(family_event)

    # 科研与实践方面（弱化科研要求）

    def _add_research_practice(self):
        """科研与实践方面（200~300 字）：学生工作 + 志愿服务 + 竞赛/大创"""
        self.add_heading("四、科研与实践方面")
        rp_text = self._get("research_practice", default="")
        if rp_text and isinstance(rp_text, str):
            self.add_para(rp_text)
        self._add_competition_items()
        self._add_social_practice()

    def _add_competition_items(self):
        """学科竞赛 + 大创列表（励志类不要求国家级，校级即可）"""
        parts = []
        for c in self._get_list("competitions"):
            if not isinstance(c, dict):
                continue
            seg = (f"{c.get('time', '')} " if c.get("time") else "") + c.get("name", "")
            if c.get("award"):
                seg += f" {c['award']}"
            if c.get("role"):
                seg += f"（{c['role']}）"
            if seg:
                parts.append(seg + "；")
        for p in self._get_list("research_projects"):
            if not isinstance(p, dict):
                continue
            name = p.get("name", "")
            seg = f"{p.get('role', '')}《{name}》" if p.get("role") and name else f"《{name}》" if name else ""
            if p.get("level"):
                seg += f"（{p['level']}）"
            if p.get("duration"):
                seg += f"，{p['duration']}"
            if seg:
                parts.append(seg + "；")
        if parts:
            parts[-1] = parts[-1].rstrip("；") + "。"
            self.add_para("".join(parts))

    def _add_social_practice(self):
        """学生工作 + 志愿服务"""
        position = self._get("position", default="")
        position_work = self._get("position_work", default="")
        volunteer_hours = self._get("volunteer_hours", default="")
        volunteer_detail = self._get("volunteer_detail", default="")
        parts = []
        if position:
            seg = f"担任{position}"
            if position_work:
                seg += f"，主要工作：{position_work}"
            parts.append(seg + "。")
        if volunteer_hours or volunteer_detail:
            seg = f"累计志愿服务时长 {volunteer_hours} 小时" if volunteer_hours else ""
            if volunteer_detail:
                seg = (seg + "：" if seg else "") + volunteer_detail
            parts.append(seg + "。")
        if parts:
            self.add_para("".join(parts))

    # 生活方面

    def _add_life(self):
        """生活方面（120~200 字）：生活作风 + 勤俭具体表现 + 人际关系"""
        self.add_heading("五、生活方面")
        life = self._get("life", default="")
        if life:
            if isinstance(life, list):
                for p in life:
                    self.add_para(p)
            else:
                self.add_para(life)
            return
        dorm_role = self._get("dorm_role", default="")
        dorm_activity = self._get("dorm_activity", default="")
        dorm_honor = self._get("dorm_honor", default="")
        interpersonal = self._get("interpersonal", default="")
        lifestyle = self._get("lifestyle",
                               default="生活中我注重勤俭节约，每月生活费控制在 800 元以内，"
                                       "教材均通过二手或图书馆借阅获取。")
        parts = [lifestyle]
        if dorm_role and dorm_activity:
            seg = f"担任{dorm_role}期间，{dorm_activity}"
            if dorm_honor:
                seg += f"，{dorm_honor}"
            parts.append(seg + "。")
        if interpersonal:
            parts.append(interpersonal + "。")
        self.add_para("".join(parts))

    # 结尾"此致 敬礼！"

    def _add_ending(self):
        """结尾（80~150 字）：事实总结 + 朴素表态 + 此致/敬礼！"""
        ending = self._get("ending", default="")
        if ending:
            self.add_para(ending)
        else:
            self.add_para(
                "以上是我本学年的基本情况。无论结果如何，"
                "我都将以此为新的起点，继续努力学习、自强不息，"
                "以实际行动回报学校与社会。恳请评审委员会予以考虑。"
            )
        add_cizhi_paragraph(self.doc, "此致")   # "此致"另起一行，空两格
        add_jingli_paragraph(self.doc, "敬礼！")  # "敬礼！"另起一行，顶格

    # 落款

    def _add_signature(self):
        """落款：右对齐，含申请人 + 日期"""
        self.doc.add_paragraph()  # 空一行
        name = self._get("name", default="")
        apply_date = self._get("apply_date", default="")
        if name:
            add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if apply_date:
            add_right_aligned_paragraph(self.doc, apply_date)

    # 主构建方法

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/开头/思想/学习/家庭经济/科研实践/生活/结尾/落款"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_opening()
            self._add_ideology()
            self._add_academics()
            self._add_family_economy()
            self._add_research_practice()
            self._add_life()
            self._add_ending()
            self._add_signature()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申请书已生成：{output_path}")
        return str(output_path)

    # 数据校验

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []

        # P0 必采字段
        for key, name in [("name", "申请人姓名"), ("college", "学院"),
                          ("major", "专业"), ("grade", "年级"),
                          ("gpa", "GPA"), ("rank", "专业排名")]:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        # 主干课程
        courses = self._get_list("core_courses")
        if not courses:
            warnings.append("缺少 主干课程（core_courses），将省略主干课程表")
        elif len(courses) < 4:
            warnings.append(f"主干课程仅 {len(courses)} 门，建议 4~6 门")

        # 排名校验
        rank_str = str(self._get("rank", default=""))
        if rank_str and "/" not in rank_str and not self._get("rank_total"):
            warnings.append(
                f"排名 '{rank_str}' 缺基数，应为 'X/N' 格式或补填 rank_total")

        # 励志类硬门槛校验：家庭经济困难认定
        if not self._get("difficulty_level", default=""):
            warnings.append("缺少 家庭经济困难认定等级（difficulty_level），励志类硬门槛")

        # 家庭年收入校验
        income = self._get("family_income_year", default="")
        if not income:
            warnings.append("缺少 家庭年收入（family_income_year），励志类核心数据")
        else:
            try:
                n = float(str(income))
                if n >= 60000:
                    warnings.append(
                        f"家庭年收入 {n:.0f} 元偏高，可能不符合困难认定（建议 < 6 万）")
            except ValueError:
                pass

        # 家庭成员表校验
        members = self._get_list("family_members")
        if not members:
            warnings.append("缺少 家庭成员表（family_members），励志类必备")
        elif len(members) < 3:
            warnings.append(f"家庭成员仅 {len(members)} 人，建议至少 3 人（父母+本人）")

        # 勤工助学校验
        if not self._get_list("work_study"):
            warnings.append("缺少 勤工助学经历（work_study），励志类核心佐证")

        # 困难认定编号校验
        if not self._get("difficulty_cert_no", default=""):
            warnings.append("缺少 困难认定编号/建档立卡号/低保号（difficulty_cert_no）")

        # 学业门槛校验（励志类宽于国奖）
        try:
            rank_num = int(rank_str.split("/")[0]) if "/" in rank_str else 0
            rank_total_num = 0
            if "/" in rank_str:
                rank_total_num = int(rank_str.split("/")[1])
            elif self._get("rank_total"):
                rank_total_num = int(self._get("rank_total"))
            if rank_num and rank_total_num:
                pct = rank_num / rank_total_num * 100
                if pct > 30:
                    warnings.append(
                        f"排名前 {pct:.1f}% 不满足励志类前 30% 门槛，建议改申助学金")
        except (ValueError, IndexError):
            pass

        self.warnings = warnings
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据
# ============================================================

DEFAULT_DATA = {
    "name": "李华", "student_id": "2022123456", "gender": "男",
    "college": "计算机科学与技术学院", "major": "计算机科学与技术",
    "grade": "2022 级 大三", "class_name": "计科 2201 班",
    "political_status": "共青团员", "phone": "138XXXXXXXX",
    "apply_year": "2024-2025 学年", "apply_date": "2025 年 6 月 15 日",
    "salutation": "尊敬的学院领导、评审委员会：",
    "gpa": "3.65/4.0", "weighted_avg": "85.5", "rank": "8/87", "rank_total": "87",
    "course_count": "12", "high_score_count": "9",
    "core_courses": [
        {"name": "高等数学（上）", "credit": "5", "score": "88"},
        {"name": "数据结构", "credit": "4", "score": "90"},
        {"name": "操作系统", "credit": "4", "score": "87"},
        {"name": "计算机网络", "credit": "3", "score": "86"},
        {"name": "数据库原理", "credit": "3", "score": "89"},
    ],
    "course_highlight": "高等数学 88、数据结构 90、操作系统 87，专业核心课平均 88 分，全部 80+",
    "cet4": "510", "cet6": "", "computer_level": "二级 C 语言（合格）",
    "study_method": "学习上注重课前预习与课后总结，遇到问题主动与老师、同学讨论。",
    "academics_summary": "",
    "party_history": "2024.03 提交入党申请书，2024.09 列为入党积极分子。",
    "party_activities": [
        "参加学院分党校第 6 期培训班（2024.10-2024.12）结业",
        "2024.11 参与主题党日活动'社区志愿服务'，撰写心得 1 篇",
        "提交思想汇报 3 篇",
    ],
    "ideology": "", "academics": "", "family_economy": "",
    "hometown": "XX 省 XX 县农村家庭", "family_income_year": "25000",
    "family_income_source": "父亲在家务农，母亲身体不好长期服药，偶尔做零工",
    "difficulty_level": "特别困难",
    "difficulty_cert_no": "家庭为建档立卡贫困户（编号：XX 省 XX 县建档立卡户 XXXXXXXX）",
    "dependents_count": "4",
    "dependents_detail": "父亲李 XX、母亲王 XX、本人（XX 大学大三在读）、妹妹李 X（XX 县一中高二在读），妹妹学费与生活费年支出约 1.2 万元",
    "family_members": [
        {"name": "李 XX", "relation": "父亲", "job": "XX 县务农",
         "income": "18000", "political": "群众"},
        {"name": "王 XX", "relation": "母亲", "job": "在家（长期服药）",
         "income": "7000", "political": "群众"},
        {"name": "李 X", "relation": "妹妹", "job": "XX 县一中高二",
         "income": "", "political": "共青团员"},
        {"name": "李华", "relation": "本人", "job": "XX 大学大三",
         "income": "", "political": "共青团员"},
    ],
    "work_study": [
        {"position": "图书馆管理员", "place": "校图书馆",
         "duration": "2024.09-2025.06", "income": "每月 400 元"},
        {"position": "食堂帮厨", "place": "校第二食堂",
         "duration": "2024.03-2024.07", "income": "每月 600 元"},
        {"position": "家教", "place": "校外",
         "duration": "2024.10-2025.05", "income": "每月 600 元"},
    ],
    "work_study_summary": "勤工助学月收入约 1600 元，用于覆盖教材费与生活日用品支出，减轻家庭负担。",
    "family_event": "2023.06 父亲因突发脑溢血住院治疗 3 个月，累计医疗支出约 8 万元，出院后丧失部分劳动能力，家庭收入骤减 60%。",
    "research_practice": "",
    "research_projects": [
        {"name": "校园闲置物品流转平台", "level": "校级大创项目",
         "role": "主持", "duration": "2024.05-2025.05", "output": ""},
    ],
    "competitions": [
        {"name": "校级程序设计竞赛", "award": "二等奖",
         "time": "2024.11", "role": "个人参赛"},
    ],
    "position": "班级学习委员（2024.09-2025.06）",
    "position_work": "组织 5 次学习经验交流会，服务同学 30 余人次；建立班级学习互助小组，覆盖 6 门主干课程",
    "volunteer_hours": "80",
    "volunteer_detail": "参与三下乡 1 次（2024.07），支教 8 课时；担任校园开放日志愿者 2 次",
    "dorm_role": "宿舍长", "dorm_activity": "组织宿舍 6 次集体活动",
    "dorm_honor": "宿舍连续两学期获评'文明宿舍'",
    "interpersonal": "与同学相处融洽，曾帮助室友完成 1 次重要实验调试",
    "lifestyle": "生活中我注重勤俭节约，每月生活费控制在 800 元以内，教材均通过二手或图书馆借阅获取。",
    "life": "", "ending": "",
    "honors": [
        {"time": "2024.11", "name": "校级二等奖学金",
         "level": "校级（专业前 15%）", "issuer": "XX 大学"},
        {"time": "2024.10", "name": "校级三好学生",
         "level": "校级", "issuer": "XX 大学"},
    ],
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="国家励志奖学金申请书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n"
            "JSON 字段定义详见 SKILL.md 第三章信息采集清单。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
