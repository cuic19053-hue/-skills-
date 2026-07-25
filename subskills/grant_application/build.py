#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
国家助学金申请书 docx 生成器

格式标准：A4 纸张，页边距上下 2.54cm 左右 2.5cm；标题黑体二号居中；
称呼顶格宋体小四全角冒号；正文宋体小四 1.5 倍行距首行缩进 2 字符；
家庭成员表 5 列 / 勤工助学表 4 列 / 主干课程表 3 列 宋体五号居中；
"此致"另起一行空两格，"敬礼！"另起一行顶格；落款右对齐。

助学金档次：一等 4400 元/年（特别困难）/ 二等 3300 元/年（困难）/ 三等 2200 元/年（一般困难）
依据：财教〔2007〕92 号 + 教财函〔2019〕106 号 + 教财〔2024〕8 号（2024.7.1 起新标准）

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第九章 JSON Schema。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# ============================================================
# 助学金档次与金额对照（2024 年 7 月 1 日起执行）
# ============================================================

GRANT_LEVEL_AMOUNT = {
    "一等": 4400,
    "二等": 3300,
    "三等": 2200,
}

GRANT_LEVEL_DIFFICULTY = {
    "一等": "特别困难",
    "二等": "困难",
    "三等": "一般困难",
}

DIFFICULTY_LEVEL_LINE = {
    "特别困难": 400,
    "困难": 600,
    "一般困难": 800,
}

POVERTY_CATEGORIES = [
    "建档立卡", "低保", "重大疾病", "单亲", "多子女上大学", "失业",
]


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size: Pt = SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size: Pt = SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size: Pt = SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进），宋体小四"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_section_heading(doc, text: str):
    """正文中的小节标题（一、二、三…）：黑体小四加粗，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
        line_spacing=1.5, space_before=6, space_after=3)


def add_cizhi_paragraph(doc, text: str = "此致"):
    """'此致'另起一行，空两格（首行缩进 2 字符）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格（不缩进）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=False, line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None,
                        caption: str = ""):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    if caption:
        add_paragraph_with_format(
            doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)

    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


# ============================================================
# 助学金档次与人均收入推算
# ============================================================

def compute_grant_amount(grant_level: str) -> int:
    """根据 grant_level 推算 grant_amount（一等 4400 / 二等 3300 / 三等 2200）"""
    return GRANT_LEVEL_AMOUNT.get(grant_level, 0)


def compute_per_capita(family_income_year: float, family_size: int) -> Tuple[float, float]:
    """根据年总收入与家庭人口推算人均年收入与人均月收入"""
    if family_size <= 0:
        return 0.0, 0.0
    per_capita_year = family_income_year / family_size
    per_capita_month = per_capita_year / 12
    return per_capita_year, per_capita_month


def get_school_line(difficulty_level: str) -> int:
    """根据困难等级获取学校认定线（特别困难 400 / 困难 600 / 一般困难 800）"""
    return DIFFICULTY_LEVEL_LINE.get(difficulty_level, 0)


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """国家助学金申请书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []
        self.word_count_version: str = "1200"

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段，非列表返回空列表"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def _get_fs(self, *keys, default=""):
        """从 family_situation 子结构安全取字段"""
        return self._get("family_situation", *keys, default=default)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_heading(self, text):
        return add_section_heading(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, caption=""):
        return add_table_from_data(self.doc, headers, rows, col_widths, caption)

    # 标题

    def _add_title(self):
        """标题：黑体二号居中，固定为'国家助学金申请书'8 字"""
        add_title(self.doc, "国家助学金申请书")

    # 称呼

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号"""
        salutation = self._get("salutation",
                               default="尊敬的学院领导、评审委员会：")
        add_salutation_paragraph(self.doc, salutation)

    # 段一 个人基本情况

    def _add_opening(self):
        """段一 个人基本情况：姓名/学院/专业/年级 + 申请档次与金额 + 困难认定等级与编号"""
        opening = self._get("opening", default="")
        if opening:
            self.add_para(opening)
            return

        name = self._get("name")
        college = self._get("college")
        major = self._get("major")
        grade = self._get("grade")
        apply_year = self._get("apply_year", default="2024-2025 学年")
        grant_level = self._get("grant_level", default="")
        grant_amount = self._get("grant_amount", default=0)
        # 自动推算 grant_amount
        if not grant_amount and grant_level:
            grant_amount = compute_grant_amount(grant_level)
        difficulty_level = self._get("difficulty_level", default="")
        difficulty_cert_no = self._get("difficulty_cert_no", default="")
        poverty_category = self._get("poverty_category", default="")

        parts = []
        if name and college and major and grade:
            parts.append(f"我是{college}{major}{grade}学生{name}，"
                         f"特申请{apply_year}国家助学金{grant_level}"
                         f"（{grant_amount} 元/年）。")
        else:
            parts.append(f"特申请{apply_year}国家助学金{grant_level}。")

        # 困难认定等级与编号
        cert_parts = []
        if difficulty_level:
            cert_parts.append(f"等级为'{difficulty_level}'")
        if difficulty_cert_no:
            cert_parts.append(f"认定编号：{difficulty_cert_no}")
        if cert_parts:
            parts.append("本人已获学校家庭经济困难认定，"
                         + "，".join(cert_parts) + "。")

        # 6 类困难类别
        if poverty_category:
            parts.append(f"家庭为{poverty_category}户。")

        parts.append("现将本人情况汇报如下：")
        self.add_para("".join(parts))

    # 段二 家庭经济情况

    def _add_family_economy(self):
        """段二 家庭经济情况：6 大要素（家庭人口/收入来源/年总收入与人均/大额支出/债务/当地对照）"""
        family_economy = self._get("family_economy", default="")
        if family_economy:
            # 用户直接提供整段文本
            self.add_heading("二、家庭经济情况")
            self.add_para(family_economy)
            return

        self.add_heading("二、家庭经济情况")

        family_size = self._get_fs("family_size", default=0)
        college_students = self._get_fs("college_students", default=0)
        elderly_count = self._get_fs("elderly_count", default=0)
        father_name = self._get_fs("father_name", default="")
        father_job = self._get_fs("father_job", default="")
        father_income_detail = self._get_fs("father_income_detail", default="")
        mother_name = self._get_fs("mother_name", default="")
        mother_job = self._get_fs("mother_job", default="")
        mother_income_detail = self._get_fs("mother_income_detail", default="")
        income_source = self._get_fs("income_source", default="")
        family_income_year = self._get_fs("family_income_year", default=0)
        local_min_living = self._get_fs("local_min_living", default=0)
        school_line = self._get_fs("school_line", default=0)
        if not school_line:
            difficulty_level = self._get("difficulty_level", default="")
            school_line = get_school_line(difficulty_level)
        large_expenses = self._get_fs("large_expenses", default="")
        debt_total = self._get_fs("debt_total", default=0)
        debt_relative = self._get_fs("debt_relative", default=0)
        debt_loan = self._get_fs("debt_loan", default=0)
        debt_reason = self._get_fs("debt_reason", default="")
        poverty_category = self._get("poverty_category", default="")
        difficulty_cert_no = self._get("difficulty_cert_no", default="")
        difficulty_level = self._get("difficulty_level", default="")
        hometown = self._get_fs("hometown", default="")

        # 要素 1 家庭人口
        pop_parts = []
        if family_size:
            pop_parts.append(f"我家庭共 {family_size} 口人")
        if college_students:
            pop_parts.append(f"在校大学生 {college_students} 人")
        if elderly_count:
            pop_parts.append(f"需赡养老人 {elderly_count} 位")
        if pop_parts:
            self.add_para("，".join(pop_parts) + "。"
                          + (f"家庭住址：{hometown}。" if hometown else ""))

        # 要素 2 收入来源
        income_parts = []
        if father_name:
            father_str = f"{father_name}"
            if father_job:
                father_str += f"，{father_job}"
            if father_income_detail:
                father_str += f"，{father_income_detail}"
            income_parts.append(father_str)
        if mother_name:
            mother_str = f"{mother_name}"
            if mother_job:
                mother_str += f"，{mother_job}"
            if mother_income_detail:
                mother_str += f"，{mother_income_detail}"
            income_parts.append(mother_str)
        if income_parts:
            self.add_para("；".join(income_parts) + "。"
                          + (f"家庭主要收入来源为{income_source}。" if income_source else ""))

        # 要素 3 年度总收入与人均
        if family_income_year and family_size:
            per_capita_year, per_capita_month = compute_per_capita(
                float(family_income_year), int(family_size))
            compare_parts = [f"家庭年总收入约 {int(family_income_year)} 元",
                             f"人均年收入约 {int(per_capita_year)} 元",
                             f"人均月收入约 {int(per_capita_month)} 元"]
            if local_min_living:
                compare_parts.append(
                    f"低于当地最低生活保障标准 {int(local_min_living)} 元/月")
            if school_line:
                compare_parts.append(
                    f"低于学校家庭经济困难认定线 {int(school_line)} 元/月")
            self.add_para("，".join(compare_parts) + "。")

        # 6 类困难类别 + 编号
        if poverty_category:
            self.add_para(
                f"家庭为{poverty_category}户"
                + (f"（编号：{difficulty_cert_no}）" if difficulty_cert_no else "")
                + "。")

        # 要素 4 大额支出
        if large_expenses:
            self.add_para(f"近 1~2 年内大额支出：{large_expenses}。")

        # 要素 5 债务
        if debt_total:
            debt_parts = [f"家庭累计借款约 {int(debt_total)} 元"]
            debt_sub = []
            if debt_relative:
                debt_sub.append(f"亲戚借款 {int(debt_relative)} 元")
            if debt_loan:
                debt_sub.append(f"生源地助学贷款 {int(debt_loan)} 元")
            if debt_sub:
                debt_parts.append("其中 " + "、".join(debt_sub))
            if debt_reason:
                debt_parts.append(f"借贷原因为{debt_reason}")
            self.add_para("，".join(debt_parts) + "。")

        # 要素 6 当地生活水平对照（收尾）
        if difficulty_level and difficulty_cert_no:
            self.add_para(
                f"本人已获学校家庭经济困难认定，"
                f"等级为'{difficulty_level}'"
                f"（认定编号：{difficulty_cert_no}）。")

    # 家庭成员表

    def _add_family_table(self):
        """家庭成员表（5 列：姓名/关系/工作单位职务/年收入/政治面貌）"""
        members = self._get_list("family_members")
        if not members:
            return
        headers = ["姓名", "与本人关系", "工作单位/职务", "年收入（元）", "政治面貌"]
        rows = []
        for m in members:
            if not isinstance(m, dict):
                continue
            rows.append([
                str(m.get("name", "")),
                str(m.get("relation", "")),
                str(m.get("job", "")),
                str(m.get("income", "")),
                str(m.get("political", "")),
            ])
        if rows:
            self.add_table(headers, rows,
                           col_widths=[2.0, 2.5, 4.5, 2.5, 2.5],
                           caption="家庭成员情况：")

    # 勤工助学表

    def _add_work_study_table(self):
        """勤工助学表（4 列：岗位/地点/时长/月收入）"""
        # 800 字版本省略勤工助学表
        if self.word_count_version == "800":
            return
        items = self._get_list("work_study")
        if not items:
            return
        headers = ["岗位", "地点", "时长", "月收入（元）"]
        rows = []
        for it in items:
            if not isinstance(it, dict):
                continue
            rows.append([
                str(it.get("position", "")),
                str(it.get("place", "")),
                str(it.get("duration", "")),
                str(it.get("monthly_income", "")),
            ])
        if rows:
            self.add_table(headers, rows,
                           col_widths=[4.0, 3.5, 4.5, 2.5],
                           caption="勤工助学经历：")

    # 段三 在校表现

    def _add_campus_performance(self):
        """段三 在校表现：思想 + 学习（不强调 GPA 排名）+ 勤工助学 + 综合素质 + 生活"""
        self.add_heading("三、在校表现")

        ideology = self._get("ideology", default="")
        party_history = self._get("party_history", default="")
        study_situation = self._get("study_situation", default="")
        gpa = self._get("gpa", default="")
        rank = self._get("rank", default="")
        rank_total = self._get("rank_total", default="")
        course_failed = self._get("course_failed", default="")
        position = self._get("position", default="")
        position_work = self._get("position_work", default="")
        volunteer_hours = self._get("volunteer_hours", default="")
        volunteer_detail = self._get("volunteer_detail", default="")
        lifestyle = self._get("lifestyle", default="")
        dorm_role = self._get("dorm_role", default="")
        dorm_honor = self._get("dorm_honor", default="")

        # 思想方面
        thought_parts = []
        if ideology:
            thought_parts.append(ideology)
        if party_history:
            thought_parts.append(party_history)
        if thought_parts:
            self.add_para("思想上，本人" + "，".join(thought_parts) + "。")

        # 学习方面（不强调 GPA 排名，简述及格情况）
        study_parts = []
        if study_situation:
            study_parts.append(study_situation)
        else:
            if course_failed and course_failed != "无":
                study_parts.append(f"本学年{course_failed}")
            else:
                study_parts.append("本学年必修课全部及格")
        if study_parts:
            self.add_para("学习上，本人勤奋上进，" + "，".join(study_parts) + "。")
            # 注意：助学金申请书不写 GPA 与排名，避免让评审委员会误以为应改申励志奖学金

        # 勤工助学（在段三中以文字描述，详细列表见勤工助学表）
        work_study_items = self._get_list("work_study")
        if work_study_items:
            ws_summary_parts = []
            for it in work_study_items[:3]:
                if isinstance(it, dict):
                    pos = it.get("position", "")
                    inc = it.get("monthly_income", "")
                    if pos and inc:
                        ws_summary_parts.append(f"{pos}（每月 {inc} 元）")
            if ws_summary_parts:
                monthly_total = sum(
                    int(it.get("monthly_income", 0) or 0)
                    for it in work_study_items if isinstance(it, dict))
                self.add_para(
                    "课余参加勤工助学 " + str(len(work_study_items)) + " 项："
                    + "、".join(ws_summary_parts)
                    + f"，月收入约 {monthly_total} 元用于覆盖教材费与生活支出，"
                    + "减轻家庭负担。")

        # 学生干部与志愿服务
        overall_parts = []
        if position:
            overall_parts.append(f"担任{position}")
            if position_work:
                overall_parts[-1] += f"，{position_work}"
        if volunteer_hours:
            volunteer_str = f"累计志愿服务 {volunteer_hours} 小时"
            if volunteer_detail:
                volunteer_str += f"（{volunteer_detail}）"
            overall_parts.append(volunteer_str)
        if overall_parts:
            self.add_para("、".join(overall_parts) + "。")

        # 生活方式
        life_parts = []
        if lifestyle:
            life_parts.append(lifestyle)
        if dorm_role:
            life_parts.append(f"担任{dorm_role}")
        if dorm_honor:
            life_parts.append(f"宿舍荣誉：{dorm_honor}")
        if life_parts:
            self.add_para("生活中" + "，".join(life_parts) + "。")

    # 段四 申请理由

    def _add_application_reason(self):
        """段四 申请理由：助学金用途（学费/教材费/生活费）+ 自我解困意识"""
        application_reason = self._get("application_reason", default="")
        if application_reason:
            self.add_heading("四、申请理由")
            self.add_para(application_reason)
            return

        self.add_heading("四、申请理由")

        grant_level = self._get("grant_level", default="")
        grant_amount = self._get("grant_amount", default=0)
        if not grant_amount and grant_level:
            grant_amount = compute_grant_amount(grant_level)
        apply_year = self._get("apply_year", default="2024-2025 学年")
        tuition = self._get("tuition", default=0)
        textbook_fee = self._get("textbook_fee", default=0)
        monthly_food = self._get("monthly_food", default=0)
        monthly_income = self._get("monthly_income", default=0)
        loan_amount = self._get("loan_amount", default=0)

        parts = []
        parts.append(f"鉴于家庭经济困难，本人特申请国家助学金{grant_level}"
                     + (f" {grant_amount} 元" if grant_amount else "")
                     + "，")

        # 用途明细
        usage_parts = []
        if tuition:
            usage_parts.append(f"{apply_year}学费 {int(tuition)} 元")
        if textbook_fee:
            usage_parts.append(f"教材费 {int(textbook_fee)} 元")
        if monthly_food:
            usage_parts.append(f"月伙食费 {int(monthly_food)} 元")
        if usage_parts:
            parts.append("用于支付" + "、".join(usage_parts) + "等支出。")
        else:
            parts.append("用于支付学费、教材费与生活日用品支出。")

        parts.append("所申请助学金将严格按照学校规定用途使用，绝不挪作他用。")

        # 自我解困意识
        self_help_parts = []
        if monthly_income:
            self_help_parts.append(f"勤工助学（月收入 {int(monthly_income)} 元）")
        if loan_amount:
            self_help_parts.append(f"生源地助学贷款（{int(loan_amount)} 元/年）")
        if self_help_parts:
            parts.append("本人在领取助学金的同时，将继续通过"
                         + "与".join(self_help_parts)
                         + "自筹剩余学费与生活费，努力减轻家庭经济负担。")
        else:
            parts.append("本人将继续通过勤工助学努力减轻家庭经济负担。")

        self.add_para("".join(parts))

    # 段五 感恩表态

    def _add_gratitude(self):
        """段五 感恩表态：感谢国家资助政策（引用文号）+ 立志回报 + 此致敬礼"""
        gratitude = self._get("gratitude", default="")
        if gratitude:
            self.add_heading("五、感恩表态")
            self.add_para(gratitude)
            return

        self.add_heading("五、感恩表态")

        self.add_para(
            "国家助学金政策（依据财教〔2007〕92 号与教财〔2024〕8 号）"
            "充分体现了党和政府对家庭经济困难学生的关怀与支持，"
            "本人衷心感谢国家与学校的资助。若获评助学金，"
            "本人将以此为动力，继续勤奋学习、自强不息，"
            "立志以优异成绩回报学校与社会，将来为国家建设贡献力量。"
        )

    # 此致敬礼

    def _add_cizhi_jingli(self):
        """'此致'另起一行空两格，'敬礼！'另起一行顶格"""
        add_cizhi_paragraph(self.doc, "此致")
        add_jingli_paragraph(self.doc, "敬礼！")

    # 落款

    def _add_signature(self):
        """落款：右对齐，含申请人 + 日期"""
        self.doc.add_paragraph()  # 空一行
        name = self._get("name", default="")
        apply_date = self._get("apply_date", default="")
        if name:
            add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if apply_date:
            add_right_aligned_paragraph(self.doc, apply_date)

    # 主构建方法

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/段一~段五/家庭成员表/勤工助学表/落款"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self.word_count_version = str(self._get("word_count_version",
                                                    default="1200"))
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_opening()
            self._add_family_economy()
            self._add_family_table()
            self._add_work_study_table()
            self._add_campus_performance()
            self._add_application_reason()
            self._add_gratitude()
            self._add_cizhi_jingli()
            self._add_signature()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申请书已生成：{output_path}")
        return str(output_path)

    # 数据校验

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []

        # P0 必采字段
        for key, name in [("name", "申请人姓名"), ("college", "学院"),
                          ("major", "专业"), ("grade", "年级"),
                          ("student_id", "学号"),
                          ("apply_year", "申请学年")]:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        # grant_level 校验
        grant_level = self._get("grant_level", default="")
        if not grant_level:
            warnings.append("缺少 助学金档次（grant_level）——必填字段")
        elif grant_level not in GRANT_LEVEL_AMOUNT:
            warnings.append(
                f"助学金档次 '{grant_level}' 不在 一等/二等/三等 内")

        # grant_amount 自动推算与校验
        grant_amount = self._get("grant_amount", default=0)
        if grant_level:
            expected_amount = compute_grant_amount(grant_level)
            if grant_amount and grant_amount != expected_amount:
                warnings.append(
                    f"助学金金额 {grant_amount} 与档次 {grant_level}（应 {expected_amount} 元）不对应")

        # difficulty_level 校验
        difficulty_level = self._get("difficulty_level", default="")
        if not difficulty_level:
            warnings.append("缺少 困难认定等级（difficulty_level）——必填字段")
        elif difficulty_level not in DIFFICULTY_LEVEL_LINE:
            warnings.append(
                f"困难认定等级 '{difficulty_level}' 不在 特别困难/困难/一般困难 内")

        # grant_level 与 difficulty_level 对应校验
        if grant_level and difficulty_level:
            expected_difficulty = GRANT_LEVEL_DIFFICULTY.get(grant_level)
            if expected_difficulty and difficulty_level != expected_difficulty:
                warnings.append(
                    f"助学金档次 {grant_level}（对应 {expected_difficulty}）"
                    f"与困难等级 {difficulty_level} 不对应，跨档申请")

        # difficulty_cert_no 校验
        if not self._get("difficulty_cert_no", default=""):
            warnings.append("缺少 困难认定编号/建档立卡号/低保号（difficulty_cert_no）")

        # poverty_category 校验
        poverty_category = self._get("poverty_category", default="")
        if not poverty_category:
            warnings.append("缺少 家庭困难类别（poverty_category）——6 类之一")
        elif poverty_category not in POVERTY_CATEGORIES:
            warnings.append(
                f"家庭困难类别 '{poverty_category}' 不在 6 类内"
                f"（建档立卡/低保/重大疾病/单亲/多子女上大学/失业）")

        # 家庭经济子结构校验
        family_size = self._get_fs("family_size", default=0)
        if not family_size:
            warnings.append("缺少 family_situation.family_size（家庭人口数）")
        family_income_year = self._get_fs("family_income_year", default=0)
        if not family_income_year:
            warnings.append("缺少 family_situation.family_income_year（家庭年总收入）")
        else:
            try:
                n = float(str(family_income_year))
                if n >= 60000:
                    warnings.append(
                        f"家庭年总收入 {n:.0f} 元偏高，可能不符合困难认定（建议 < 6 万）")
            except (ValueError, TypeError):
                pass

        local_min_living = self._get_fs("local_min_living", default=0)
        if not local_min_living:
            warnings.append("缺少 family_situation.local_min_living（当地最低生活保障标准）")

        large_expenses = self._get_fs("large_expenses", default="")
        if not large_expenses:
            warnings.append("缺少 family_situation.large_expenses（大额支出）")

        if not self._get_fs("income_source", default=""):
            warnings.append("缺少 family_situation.income_source（收入来源）")

        # 家庭成员表校验
        members = self._get_list("family_members")
        if not members:
            warnings.append("缺少 家庭成员表（family_members），必备")
        elif len(members) < 3:
            warnings.append(f"家庭成员仅 {len(members)} 人，建议至少 3 人（父母+本人）")

        # 勤工助学校验
        if not self._get_list("work_study"):
            warnings.append("缺少 勤工助学经历（work_study），建议填 2~3 项")

        # 字数版本校验
        if self.word_count_version not in ("800", "1200", "1500"):
            warnings.append(
                f"字数版本 '{self.word_count_version}' 不在 800/1200/1500 内，使用默认 1200")
            self.word_count_version = "1200"

        self.warnings = warnings
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据
# ============================================================

DEFAULT_DATA = {
    # 基础信息
    "name": "张三",
    "gender": "男",
    "ethnicity": "汉族",
    "college": "计算机科学与技术学院",
    "major": "计算机科学与技术",
    "grade": "2024 级 大一",
    "class_name": "计科 2401 班",
    "student_id": "2024010101",
    "political_status": "共青团员",
    "phone": "138XXXXXXXX",

    # 申请信息
    "apply_year": "2024-2025 学年",
    "apply_date": "2024 年 9 月 20 日",
    "grant_level": "一等",
    "grant_amount": 4400,
    "difficulty_level": "特别困难",
    "difficulty_cert_no": "GZ2024-00123",
    "poverty_category": "建档立卡",

    # 家庭经济子结构
    "family_situation": {
        "family_size": 5,
        "college_students": 2,
        "elderly_count": 1,
        "hometown": "XX 省 XX 县农村家庭",
        "zip_code": "XXXXXX",
        "father_name": "张 XX",
        "father_job": "XX 县 XX 厂下岗工人",
        "father_income_detail": "下岗失业，打零工月收入约 1800 元",
        "mother_name": "李 XX",
        "mother_job": "在家务农 + 镇上家政",
        "mother_income_detail": "种植玉米 3 亩年产值 4500 元 + 家政月收入 1200 元",
        "income_source": "父亲打零工 + 母亲务农与家政",
        "family_income_year": 41000,
        "local_min_living": 750,
        "school_line": 400,
        "large_expenses": "2024.03 祖父确诊肺癌自付医疗费 4.5 万元；2024.09 妹妹考入大学学费 6000 元/年",
        "debt_total": 68000,
        "debt_relative": 48000,
        "debt_loan": 20000,
        "debt_bank": 0,
        "debt_reason": "支付医疗费与两人学费",
    },

    # 家庭成员表（5 列）
    "family_members": [
        {"name": "张 XX", "relation": "父亲", "job": "XX 县下岗工人",
         "income": 21600, "political": "群众"},
        {"name": "李 XX", "relation": "母亲", "job": "在家务农 + 家政",
         "income": 18000, "political": "群众"},
        {"name": "张 X", "relation": "妹妹", "job": "XX 大学在读",
         "income": 0, "political": "共青团员"},
        {"name": "王 XX", "relation": "祖父", "job": "在家休养（肺癌术后）",
         "income": 0, "political": "群众"},
        {"name": "张三", "relation": "本人", "job": "XX 大学大一",
         "income": 0, "political": "共青团员"},
    ],

    # 勤工助学表（4 列）
    "work_study": [
        {"position": "图书馆管理员", "place": "学校图书馆",
         "duration": "2024.09~至今 每周 8 小时", "monthly_income": 400},
        {"position": "食堂帮厨", "place": "学校一食堂",
         "duration": "2024.10~2024.12 每周 6 小时", "monthly_income": 600},
        {"position": "家教（高中数学）", "place": "学生家中",
         "duration": "2024.03~2024.07 每周 4 小时", "monthly_income": 600},
    ],

    # 在校表现
    "study_situation": "本学年必修课全部及格，主干课平均成绩 82 分，CET-4 510 分",
    "gpa": "",
    "rank": "",
    "rank_total": "",
    "course_failed": "无",
    "ideology": "拥护中国共产党的领导，2024.09 递交入党申请书，2024.10 列为入党积极分子",
    "party_history": "参加学院分党校第 38 期培训班结业，提交思想汇报 2 篇",
    "position": "班级学习委员",
    "position_work": "组织 5 次学习经验交流会",
    "volunteer_hours": 50,
    "volunteer_detail": "敬老院慰问 3 次、三下乡 1 次",
    "lifestyle": "月伙食费 800 元以内，教材通过二手或图书馆获取",
    "dorm_role": "宿舍长",
    "dorm_honor": "宿舍连续两学期获评'文明宿舍'",

    # 申请理由
    "tuition": 6000,
    "textbook_fee": 800,
    "monthly_food": 800,
    "loan_amount": 8000,
    "monthly_income": 1600,

    # 输出控制
    "material_type": "B",
    "word_count_version": "1200",
    "salutation": "尊敬的学院领导、评审委员会：",
    "opening": "",
    "family_economy": "",
    "application_reason": "",
    "gratitude": "",
    "apply_date": "2024-09-20",
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="国家助学金申请书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n"
            "助学金档次：一等 4400 元（特别困难）/ 二等 3300 元（困难）/ "
            "三等 2200 元（一般困难）\n"
            "依据：财教〔2007〕92 号 + 教财函〔2019〕106 号 + 教财〔2024〕8 号\n"
            "\n"
            "JSON 字段定义详见 SKILL.md 第九章 JSON Schema。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
