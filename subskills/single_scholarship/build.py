#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""单项奖学金申请书 docx 生成器

支持 4 个子方向：科研（research）/社工（social_work）/文体（art_sport）/学习进步（progress）
格式：A4，页边距上下 2.54cm 左右 2.5cm；标题黑体二号居中；正文宋体小四 1.5 倍行距首行缩进 2 字符；
单项表格宋体五号居中；"此致"另起一行空两格；"敬礼！"另起一行顶格；落款右对齐。

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

# 子方向代码 → (中文标题后缀, 中文短称)
CATEGORY_MAP = {
    "research": ("科研单项奖学金申请书", "科研"),
    "social_work": ("社工单项奖学金申请书", "社会工作"),
    "art_sport": ("文体单项奖学金申请书", "文体"),
    "progress": ("学习进步单项奖学金申请书", "学习进步"),
}


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc, text: str, font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI, bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent: bool = True,
    line_spacing: float = 1.5, space_before: float = 0, space_after: float = 0,
):
    """添加带格式段落"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_section_heading(doc, text: str):
    """小节标题：黑体小四加粗，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
        line_spacing=1.5, space_before=6, space_after=3)


def add_cizhi_paragraph(doc, text: str = "此致"):
    """"此致"另起一行，空两格"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """"敬礼！"另起一行，顶格"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=False, line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None,
                        caption: str = ""):
    """从数据创建表格（表头加粗居中、数据居中），caption 为表前说明"""
    if caption:
        add_paragraph_with_format(
            doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def setup_page(doc):
    """A4 页面与页边距"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码"""
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def safe_float(value, default: float = 0.0) -> float:
    """安全转浮点数"""
    try:
        return float(value)
    except (ValueError, TypeError):
        return default


def parse_rank(rank_str: str) -> Tuple[int, int]:
    """解析 'X/N' → (X, N)"""
    if not rank_str:
        return (0, 0)
    s = str(rank_str)
    if "/" in s:
        parts = s.split("/")
        try:
            return (int(parts[0]), int(parts[1]))
        except (ValueError, IndexError):
            return (0, 0)
    try:
        return (int(s), 0)
    except ValueError:
        return (0, 0)


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """单项奖学金申请书 docx 构建器（支持 4 个子方向）"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        # 设置 Normal 样式默认字体（中英文同步）
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []
        self.category: str = ""

    def _get(self, *keys, default=""):
        """安全取嵌套字段"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def _detect_category(self) -> str:
        """根据数据自动判定子方向"""
        if self._get("category"):
            cat = str(self._get("category")).strip()
            if cat in CATEGORY_MAP:
                return cat
        # 优先级：research > social_work > art_sport > progress
        if (self._get_list("papers") or self._get_list("patents")
                or self._get_list("innovations")):
            return "research"
        if (self._get("position") and
                (self._get_list("organized_activities")
                 or self._get("volunteer_hours"))):
            return "social_work"
        if self._get_list("art_sport_awards"):
            return "art_sport"
        if self._get("gpa_improvement") or self._get("rank_improvement"):
            return "progress"
        return "research"  # 兜底默认

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_heading(self, text):
        return add_section_heading(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, caption=""):
        return add_table_from_data(self.doc, headers, rows, col_widths, caption)

    # --- 标题（含单项类别动态填充）---

    def _add_title(self):
        """标题：黑体二号居中，按子方向动态填充"X 单项奖学金申请书" """
        self.category = self._detect_category()
        title = CATEGORY_MAP.get(self.category, ("单项奖学金申请书", ""))[0]
        add_title(self.doc, title)

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号"""
        salutation = self._get("salutation",
                               default="尊敬的学院领导、评审委员会：")
        add_salutation_paragraph(self.doc, salutation)

    def _add_opening(self):
        """开头段落（80~120 字）：身份+申报子方向+单项核心数据+进入正文"""
        opening = self._get("opening", default="")
        if opening:
            self.add_para(opening)
            return
        name = self._get("name")
        college = self._get("college")
        major = self._get("major")
        grade = self._get("grade")
        apply_year = self._get("apply_year", default="2024-2025 学年")
        cat_short = CATEGORY_MAP.get(self.category, ("", ""))[1]
        parts = []
        if name and college and major and grade:
            parts.append(f"我是{college}{major}{grade}学生{name}，"
                         f"特申请{apply_year}{cat_short}单项奖学金。")
        else:
            parts.append(f"特申请{apply_year}{cat_short}单项奖学金。")
        core_data = self._build_opening_core_data()
        if core_data:
            parts.append(core_data)
        parts.append("现将本人情况汇报如下：")
        self.add_para("".join(parts))

    def _build_opening_core_data(self) -> str:
        """按子方向拼装开头核心数据句"""
        if self.category == "research":
            papers = self._get_list("papers")
            patents = self._get_list("patents")
            innovations = self._get_list("innovations")
            items = []
            if papers:
                first = papers[0] if isinstance(papers[0], dict) else {}
                order = first.get("author_order", "第一作者")
                level = first.get("level", "中文核心")
                items.append(f"以{order}在《{first.get('journal', 'XX')}》"
                             f"发表{level}论文 1 篇")
            if patents:
                items.append("申请发明专利 1 项")
            if innovations:
                first = innovations[0] if isinstance(innovations[0], dict) else {}
                level = first.get("level", "校级")
                role = first.get("role", "主持")
                items.append(f"{role}{level}大创 1 项")
            return "本学年" + "，".join(items) + "。" if items else ""
        if self.category == "social_work":
            position = self._get("position", default="")
            activities = self._get_list("organized_activities")
            volunteer_hours = self._get("volunteer_hours", default="")
            beneficiaries = self._get("service_beneficiaries", default="")
            parts = []
            if position:
                pos_main = position.split("（")[0].split("(")[0]
                parts.append(f"担任{pos_main}")
            if activities:
                parts.append(f"组织大型活动 {len(activities)} 场")
            if beneficiaries:
                parts.append(f"服务 {beneficiaries} 人次")
            elif volunteer_hours:
                parts.append(f"志愿服务 {volunteer_hours} 小时")
            return "本学年" + "，".join(parts) + "。" if parts else ""
        if self.category == "art_sport":
            awards = self._get_list("art_sport_awards")
            duration = self._get("art_sport_duration", default="")
            parts = []
            if awards:
                first = awards[0] if isinstance(awards[0], dict) else {}
                level = first.get("level", "校级")
                award_str = first.get("name", "XX 比赛")
                parts.append(f"获{level}奖项（{award_str}）")
            if duration:
                parts.append(duration)
            return "本学年" + "，".join(parts) + "。" if parts else ""
        if self.category == "progress":
            gpa_before = self._get("gpa_before", default="")
            gpa_after = self._get("gpa_after", default="")
            rank_before = self._get("rank_before", default="")
            rank_after = self._get("rank_after", default="")
            parts = []
            if gpa_before and gpa_after:
                parts.append(f"GPA 由 {gpa_before} 提升至 {gpa_after}")
            if rank_before and rank_after:
                parts.append(f"专业排名由 {rank_before} 提升至 {rank_after}")
            return "本学年" + "，".join(parts) + "。" if parts else ""
        return ""

    # --- 一、单项突出表现（重点，按类别分支）---

    def _add_single_highlight(self):
        """一、单项突出表现（450~800 字，重点）
        research：论文表+专利表+大创表+科研小结；
        social_work：履职表+志愿服务表+履职亮点；
        art_sport：获奖作品表+训练与团队经历；
        progress：GPA 提升表+进步课程对比表+进步方法。"""
        self.add_heading("一、单项突出表现")
        highlight_text = self._get("single_highlight", default="")
        if highlight_text and isinstance(highlight_text, str):
            self.add_para(highlight_text)
            return
        if self.category == "research":
            self._add_research_highlight()
        elif self.category == "social_work":
            self._add_social_work_highlight()
        elif self.category == "art_sport":
            self._add_art_sport_highlight()
        elif self.category == "progress":
            self._add_progress_highlight()
        else:
            self.add_para("（请补充单项突出表现内容）")

    # ---- 科研单项分支 ----

    def _add_research_highlight(self):
        """科研单项：总述+论文表+专利表+大创表+科研小结"""
        summary = self._get("research_summary", default="")
        if summary:
            self.add_para(summary)
        else:
            papers = self._get_list("papers")
            patents = self._get_list("patents")
            innovations = self._get_list("innovations")
            field = self._get("research_field", default="XXX")
            counts = []
            if papers:
                counts.append(f"已发表 {len(papers)} 篇论文")
            if patents:
                counts.append(f"申请专利 {len(patents)} 项")
            if innovations:
                counts.append(f"主持/参与大创 {len(innovations)} 项")
            count_str = "、".join(counts) if counts else "已开展系列科研工作"
            self.add_para(f"本学年科研工作围绕{field}方向展开，{count_str}。")
        # 论文表
        papers = self._get_list("papers")
        rows = [[str(p.get("title", "")), str(p.get("journal", "")),
                 str(p.get("level", "")), str(p.get("author_order", "")),
                 str(p.get("time", ""))]
                for p in papers if isinstance(p, dict)]
        if rows:
            self.add_table(["论文题目", "期刊/会议", "级别", "作者排序", "时间"],
                           rows, col_widths=[4.5, 3.5, 2.5, 2.0, 2.5],
                           caption="论文发表情况：")
        # 专利表
        patents = self._get_list("patents")
        rows = [[str(p.get("name", "")), str(p.get("type", "")),
                 str(p.get("application_no", "")), str(p.get("status", "")),
                 str(p.get("inventor_order", ""))]
                for p in patents if isinstance(p, dict)]
        if rows:
            self.add_table(["专利名称", "类型", "申请号/授权号", "状态", "发明人排序"],
                           rows, col_widths=[4.0, 2.5, 3.5, 2.5, 2.5],
                           caption="发明专利申请情况：")
        # 大创表
        innovations = self._get_list("innovations")
        rows = [[str(p.get("name", "")), str(p.get("level", "")),
                 str(p.get("role", "")), str(p.get("duration", "")),
                 str(p.get("output", ""))]
                for p in innovations if isinstance(p, dict)]
        if rows:
            self.add_table(["项目名称", "级别", "角色", "起止时间", "产出"],
                           rows, col_widths=[4.0, 2.5, 2.0, 3.5, 3.0],
                           caption="大创/科研立项参与情况：")
        # 指导教师与科研小结
        advisor = self._get("advisor", default="")
        if isinstance(advisor, dict):
            advisor_name = advisor.get("name", "")
            advisor_title = advisor.get("title", "")
            if advisor_name:
                self.add_para(f"指导教师：{advisor_title}{advisor_name}。")
        reflection = self._get("research_reflection", default="")
        if reflection:
            self.add_para(reflection)
        else:
            self.add_para(
                "通过上述科研工作，我系统掌握了相关研究方法，"
                "理解了从问题定义到论文撰写的完整科研流程，"
                "为后续研究生阶段奠定基础。")

    # ---- 社工单项分支 ----

    def _add_social_work_highlight(self):
        """社工单项：总述+履职表+志愿服务表+履职亮点"""
        position = self._get("position", default="")
        position_work = self._get("position_work", default="")
        activities = self._get_list("organized_activities")
        volunteer_hours = self._get("volunteer_hours", default="")
        beneficiaries = self._get("service_beneficiaries", default="")
        # 总述句
        summary_parts = []
        if position:
            summary_parts.append(f"本学年担任{position}")
        if activities:
            summary_parts.append(f"组织大型活动 {len(activities)} 场")
        if beneficiaries:
            summary_parts.append(f"累计参与 {beneficiaries} 人次")
        if volunteer_hours:
            summary_parts.append(f"志愿服务 {volunteer_hours} 小时")
        if summary_parts:
            self.add_para("，".join(summary_parts) + "。")
        # 履职表
        rows = [[str(a.get("name", "")), str(a.get("time", "")),
                 str(a.get("scale", "")), str(a.get("role", ""))]
                for a in activities if isinstance(a, dict)]
        if rows:
            self.add_table(["活动名称", "时间", "规模", "本人角色"], rows,
                           col_widths=[5.0, 3.0, 3.0, 4.0],
                           caption="学生干部履职情况：")
        # 志愿服务表
        volunteer_projects = self._get_list("volunteer_projects")
        rows = [[str(v.get("name", "")), str(v.get("hours", "")),
                 str(v.get("beneficiary", ""))]
                for v in volunteer_projects if isinstance(v, dict)]
        if rows:
            self.add_table(["项目名称", "时长", "服务对象"], rows,
                           col_widths=[5.0, 4.0, 6.0],
                           caption="志愿服务情况：")
        # 履职亮点
        if position_work:
            self.add_para("履职亮点：" + position_work + "。")
        else:
            self.add_para(
                "履职亮点：组织协调活动从策划到落地全流程参与，"
                "注重活动实效与服务对象反馈。")

    # ---- 文体单项分支 ----

    def _add_art_sport_highlight(self):
        """文体单项：总述+获奖作品表+训练与团队经历"""
        art_sport_summary = self._get("art_sport_summary", default="")
        duration = self._get("art_sport_duration", default="")
        awards = self._get_list("art_sport_awards")
        # 总述句
        if art_sport_summary:
            self.add_para(art_sport_summary)
        else:
            parts = []
            if duration:
                parts.append(duration)
            if awards:
                first = awards[0] if isinstance(awards[0], dict) else {}
                parts.append(f"本学年获{first.get('level', '校级')}奖项")
            if parts:
                self.add_para("，".join(parts) + "。")
        # 获奖作品表
        rows = [[str(a.get("name", "")), str(a.get("level", "")),
                 str(a.get("time", "")), str(a.get("role", "")),
                 str(a.get("work_name", ""))]
                for a in awards if isinstance(a, dict)]
        if rows:
            self.add_table(["竞赛/展演名称", "级别", "时间", "角色", "作品名"],
                           rows, col_widths=[4.0, 2.5, 2.5, 2.5, 3.5],
                           caption="获奖/作品情况：")
        # 训练与团队经历
        team_role = self._get("team_role", default="")
        training = self._get("art_sport_training", default="")
        parts = []
        if team_role:
            parts.append(f"担任{team_role}")
        if training:
            parts.append(training)
        if parts:
            self.add_para("训练与团队经历：" + "，".join(parts) + "。")
        else:
            self.add_para(
                "训练与团队经历：日常坚持训练，注重团队协作与责任担当，"
                "在比赛与展演中不断提升专业能力。")

    # ---- 学习进步单项分支 ----

    def _add_progress_highlight(self):
        """学习进步单项：总述+GPA 提升表+进步课程对比表+进步方法"""
        gpa_before = self._get("gpa_before", default="")
        gpa_after = self._get("gpa_after", default="")
        weighted_before = self._get("weighted_before", default="")
        weighted_after = self._get("weighted_after", default="")
        rank_before = self._get("rank_before", default="")
        rank_after = self._get("rank_after", default="")
        # 总述句
        parts = []
        if gpa_before and gpa_after:
            gpa_imp = safe_float(gpa_after.split("/")[0]) - \
                safe_float(gpa_before.split("/")[0])
            parts.append(f"GPA 由 {gpa_before} 提升至 {gpa_after}"
                         f"（提升 {gpa_imp:.2f}）")
        if weighted_before and weighted_after:
            parts.append(f"加权平均分由 {weighted_before} 提升至 {weighted_after}")
        if rank_before and rank_after:
            r_b, _ = parse_rank(rank_before)
            r_a, _ = parse_rank(rank_after)
            if r_b and r_a:
                parts.append(f"专业排名由 {rank_before} 提升至 {rank_after}"
                             f"（提升 {r_b - r_a} 位）")
            else:
                parts.append(f"专业排名由 {rank_before} 提升至 {rank_after}")
        if parts:
            self.add_para("本学年" + "，".join(parts) + "。")
        # GPA 提升表
        semesters = self._get_list("gpa_progress")
        rows = []
        for s in semesters:
            if isinstance(s, dict):
                rows.append([str(s.get("semester", "")), str(s.get("gpa", "")),
                             str(s.get("weighted", "")), str(s.get("rank", ""))])
        if rows:
            self.add_table(["学期", "GPA", "加权平均分", "排名"], rows,
                           col_widths=[4.0, 4.0, 4.0, 3.0],
                           caption="GPA 提升轨迹：")
        # 进步课程对比表
        courses = self._get_list("progress_courses")
        rows = []
        for c in courses:
            if isinstance(c, dict):
                rows.append([str(c.get("name", "")), str(c.get("before", "")),
                             str(c.get("after", ""))])
        if rows:
            self.add_table(["课程名称", "上学年成绩", "本学年成绩"], rows,
                           col_widths=[6.0, 4.5, 4.5],
                           caption="进步课程对比：")
        # 进步方法
        method = self._get("improvement_method", default="")
        if method:
            self.add_para("进步方法：" + method + "。")
        else:
            self.add_para(
                "进步方法：调整学习方法，建立错题本与知识体系；"
                "合理分配主干课程学习时间；主动向老师和助教请教；"
                "与成绩优秀同学结对学习。")

    # --- 二、其他方面（简略）---

    def _add_other_aspects(self):
        """二、其他方面（150~250 字，简略）：学业/思想+生活+综合荣誉"""
        self.add_heading("二、其他方面")
        other_text = self._get("other_aspects", default="")
        if other_text:
            if isinstance(other_text, list):
                for p in other_text:
                    self.add_para(p)
            else:
                self.add_para(other_text)
            return
        parts = []
        # 学业（学习进步单项不重复 GPA 数据）
        if self.category != "progress":
            gpa = self._get("gpa", default="")
            rank = self._get("rank", default="")
            cet4 = self._get("cet4", default="")
            cet6 = self._get("cet6", default="")
            academic_parts = []
            if gpa:
                academic_parts.append(f"GPA {gpa}")
            if rank:
                academic_parts.append(f"专业排名第 {rank}")
            if academic_parts:
                parts.append("学业方面，本学年" + "，".join(academic_parts) + "；")
            lang_parts = []
            if cet4:
                lang_parts.append(f"CET-4 {cet4} 分")
            if cet6:
                lang_parts.append(f"CET-6 {cet6} 分")
            if lang_parts:
                parts.append("、".join(lang_parts) + "。")
        # 思想
        political = self._get("political_status", default="共青团员")
        party_history = self._get("party_history", default="")
        party_activities = self._get_list("party_activities")
        thought_parts = []
        if political:
            if "党员" in political:
                thought_parts.append(f"作为{political}，认真学习习近平新时代中国特色社会主义思想")
            elif "积极分子" in political:
                thought_parts.append("作为入党积极分子，认真学习党的创新理论")
            else:
                thought_parts.append(f"作为{political}，拥护中国共产党的领导")
        if party_history:
            thought_parts.append(party_history)
        if party_activities:
            thought_parts.append("；".join(party_activities))
        if thought_parts:
            parts.append("".join(thought_parts) + "。")
        # 生活
        lifestyle = self._get("lifestyle",
                              default="生活中勤俭节约，作息规律，与同学相处融洽")
        dorm_role = self._get("dorm_role", default="")
        dorm_honor = self._get("dorm_honor", default="")
        life_parts = [lifestyle]
        if dorm_role and dorm_honor:
            life_parts.append(f"担任{dorm_role}期间，{dorm_honor}")
        parts.append("".join(life_parts) + "。")
        if parts:
            self.add_para("".join(parts))

    # --- 结尾"此致 敬礼！"---

    def _add_ending(self):
        """结尾（60~100 字）：事实总结+朴素表态+此致/敬礼！"""
        ending = self._get("ending", default="")
        if ending:
            self.add_para(ending)
        else:
            cat_short = CATEGORY_MAP.get(self.category, ("", ""))[1]
            self.add_para(
                f"以上是我本学年在{cat_short}方面的基本情况。"
                "无论结果如何，我都将以此为新的起点，继续努力。"
                "恳请评审委员会予以考虑。")
        add_cizhi_paragraph(self.doc, "此致")   # "此致"另起一行，空两格
        add_jingli_paragraph(self.doc, "敬礼！")  # "敬礼！"另起一行，顶格

    # --- 落款 ---

    def _add_signature(self):
        """落款：右对齐，含申请人+日期"""
        self.doc.add_paragraph()  # 空一行
        name = self._get("name", default="")
        apply_date = self._get("apply_date", default="")
        if name:
            add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if apply_date:
            add_right_aligned_paragraph(self.doc, apply_date)

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/开头/单项突出表现/其他方面/结尾/落款"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_opening()
            self._add_single_highlight()
            self._add_other_aspects()
            self._add_ending()
            self._add_signature()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申请书已生成：{output_path}")
        return str(output_path)

    # --- 数据校验（含单项类别校验）---

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []
        # P0 必采字段
        p0_fields = [("name", "申请人姓名"), ("college", "学院"),
                     ("major", "专业"), ("grade", "年级"),
                     ("apply_year", "申报学年")]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")
        # 子方向代码校验
        category = str(self._get("category", default="")).strip()
        if category and category not in CATEGORY_MAP:
            warnings.append(f"子方向代码 '{category}' 不在 "
                            f"{list(CATEGORY_MAP)} 中，将自动判定")
        # 子方向硬门槛校验（按数据自动判定后比对）
        detected = self._detect_category()
        if category and category != detected:
            warnings.append(f"指定子方向 {category} 与自动判定 {detected} 不一致，"
                            f"按指定执行；建议核对数据")
        # 各子方向硬门槛
        if detected == "research":
            if not (self._get_list("papers") or self._get_list("patents")
                    or self._get_list("innovations")):
                warnings.append("科研单项需至少 1 项论文/专利/大创，当前数据均缺失")
        elif detected == "social_work":
            position = self._get("position", default="")
            volunteer_hours = self._get("volunteer_hours", default="")
            activities = self._get_list("organized_activities")
            if not position:
                warnings.append("社工单项建议提供学生干部职务（position）")
            vh = safe_float(volunteer_hours) if volunteer_hours else 0
            if not activities and vh < 80:
                warnings.append("社工单项需组织活动 ≥5 场或志愿服务 ≥80 小时，当前数据不足")
        elif detected == "art_sport":
            if not self._get_list("art_sport_awards"):
                warnings.append("文体单项需至少 1 项文艺/体育竞赛获奖（art_sport_awards），当前缺失")
        elif detected == "progress":
            gpa_before = self._get("gpa_before", default="")
            gpa_after = self._get("gpa_after", default="")
            rank_before = self._get("rank_before", default="")
            rank_after = self._get("rank_after", default="")
            if not (gpa_before and gpa_after):
                warnings.append("学习进步单项需 GPA before/after 数据，当前缺失")
            else:
                gpa_imp = safe_float(gpa_after.split("/")[0]) - \
                    safe_float(gpa_before.split("/")[0])
                if gpa_imp < 0.5 and not (rank_before and rank_after):
                    warnings.append(f"GPA 提升 {gpa_imp:.2f} 不足 0.5，需配合排名提升 ≥30% 数据")
                if rank_before and rank_after:
                    r_b, n_b = parse_rank(rank_before)
                    r_a, _ = parse_rank(rank_after)
                    if r_b and r_a and n_b and gpa_imp < 0.5:
                        rank_imp_pct = (r_b - r_a) / n_b * 100
                        if rank_imp_pct < 30:
                            warnings.append(f"排名提升 {rank_imp_pct:.1f}% 不足 30%")
        # 排名格式校验（如提供）
        rank_str = str(self._get("rank", default=""))
        if rank_str and "/" not in rank_str:
            warnings.append(f"排名 '{rank_str}' 应为 'X/N' 格式")
        self.warnings = warnings
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据（科研单项）
# ============================================================

DEFAULT_DATA = {
    "name": "张明", "student_id": "2022123456", "gender": "男",
    "college": "计算机科学与技术学院", "major": "计算机科学与技术",
    "grade": "2022 级 大三", "class_name": "计科 2201 班",
    "political_status": "共青团员", "phone": "138XXXXXXXX",
    "apply_year": "2024-2025 学年", "apply_date": "2025 年 6 月 15 日",
    "salutation": "尊敬的学院领导、评审委员会：",
    "category": "research",
    "gpa": "3.65/4.0", "rank": "10/87",
    "cet4": "510", "cet6": "480",
    "research_field": "自然语言处理",
    "research_summary": "本学年科研工作围绕自然语言处理方向展开，"
                       "已发表 SCI 论文 1 篇、申请发明专利 1 项、主持国家级大创 1 项。",
    "papers": [
        {"title": "基于对比学习的法律问答系统",
         "journal": "计算机研究与发展", "level": "SCI（CCF-A 类）",
         "author_order": "第一作者", "time": "2025.03"},
    ],
    "patents": [
        {"name": "一种基于深度学习的光伏故障诊断方法",
         "type": "发明专利", "application_no": "CN2024XXXXXXXX.X",
         "status": "实质审查", "inventor_order": "第一发明人"},
    ],
    "innovations": [
        {"name": "分布式光伏故障智能诊断系统", "level": "国家级大创",
         "role": "主持", "duration": "2024.03-2025.03",
         "output": "结题评估优秀"},
    ],
    "advisor": {"name": "李教授", "title": "教授",
                "research_field": "自然语言处理"},
    "research_reflection": "通过上述科研工作，我系统掌握了深度学习方法，"
                          "理解了从问题定义到论文撰写的完整科研流程，"
                          "为后续研究生阶段奠定基础。",
    "party_history": "2023.09 提交入党申请书，2024.03 列为入党积极分子。",
    "party_activities": ["参加学院分党校第 8 期培训班（2024.09-2024.12）结业"],
    "lifestyle": "生活中勤俭节约，作息规律，与同学相处融洽",
    "dorm_role": "宿舍长", "dorm_honor": "宿舍连续两学期获评'文明宿舍'",
    "single_highlight": "", "other_aspects": "", "ending": "", "opening": "",
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="单项奖学金申请书 docx 生成器"
                    "（支持科研/社工/文体/学习进步 4 个子方向）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n"
            "JSON 字段定义详见 SKILL.md 第三章信息采集清单。\n"
            "必填字段：category（research/social_work/art_sport/progress）、\n"
            "          name、major、grade、college、apply_year。\n"
            "各子方向专属字段见 SKILL.md 3.3~3.6。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档（科研单项）")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档（科研单项）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)
    # 输出校验警告数量
    if builder.warnings:
        print(f"⚠️ 共 {len(builder.warnings)} 项校验警告，详见上方输出",
              file=sys.stderr)


if __name__ == "__main__":
    main()
