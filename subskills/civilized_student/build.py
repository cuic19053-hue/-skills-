#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""文明大学生申请书 docx 生成器

评审跨度为本学年，正文 1500~2000 字，书信体格式。与优秀学生区别：
评审重点为文明素养（35%）+道德品质（30%），合计 65%；学业仅 15%。
必填：文明素养 4 维度表（课堂纪律/宿舍卫生/网络文明/公共秩序）、
道德事例表（时间/地点/事件/影响/证明人）。文明行为必须有具体时间地点人物事例。

格式：A4 页边距上下 2.54cm 左右 2.5cm；标题黑体二号居中；正文宋体小四 1.5 倍行距首行缩进 2 字符；
表格宋体五号居中；"此致"空两格+"敬礼！"顶格；落款右对齐。

使用：python build.py --data data.json --out output.docx 或 python build.py --demo --out demo.docx
JSON 字段详见 SKILL.md 第三章信息采集清单。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号（标题）
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四（正文）
SIZE_WU = Pt(10.5)          # 五号（表格）
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

# 文明素养 4 维度固定顺序
CIVILITY_DIMENSIONS = ["课堂纪律", "宿舍卫生", "网络文明", "公共秩序"]

# 文明素养表列宽（cm）
CIVILITY_COL_WIDTHS = [2.5, 6.0, 3.0, 3.5]
# 道德事例表列宽（cm）
MORALITY_COL_WIDTHS = [2.0, 2.5, 4.0, 3.5, 3.0]


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(doc, text: str, font_name: str = FONT_SONG,
        font_size=SIZE_XIAO_SI, bold: bool = False,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent: bool = True,
        line_spacing: float = 1.5, space_before: float = 0, space_after: float = 0):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_ER,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进），宋体小四"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_section_heading(doc, text: str):
    """正文小节标题（一、二、三…）：黑体小四加粗，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SI,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
        line_spacing=1.5, space_before=6, space_after=3)


def add_cizhi_paragraph(doc, text: str = "此致"):
    """'此致'另起一行，空两格（首行缩进 2 字符）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格（不缩进）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=False, line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None, caption: str = ""):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    if caption:
        add_paragraph_with_format(doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 2 and len(str(val)) > 12 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False, alignment=align)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def format_rank_percent(rank: str, rank_total: str) -> str:
    """根据排名与基数计算百分比，返回 '前 X.X%' 字符串，失败返回空串"""
    try:
        r_num = int(str(rank).split("/")[0])
        total_num = int(rank_total) if rank_total else 0
        if total_num > 0:
            return f"前 {round(r_num / total_num * 100, 1)}%"
    except (ValueError, IndexError):
        pass
    return ""


def safe_str(val: Any) -> str:
    """安全转字符串，None 返回空串"""
    return "" if val is None else str(val)


def join_non_empty(parts: List[str], sep: str = "；") -> str:
    """用分隔符连接非空字符串"""
    return sep.join(p for p in parts if p)


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """文明大学生申请书 docx 构建器。

    评审重点为文明素养（35%）与道德品质（30%），合计 65%。
    编排顺序：标题 → 称呼 → 开头 → 一、思想品德 → 二、文明素养【重点】 →
    三、道德品质【重点】 → 四、学业与生活 → 结尾 → 落款。
    文明素养段必含 4 维度表（课堂纪律/宿舍卫生/网络文明/公共秩序），
    道德品质段必含事例表（时间/地点/事件/影响/证明人）。
    """

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段，非列表返回空列表"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_heading(self, text):
        return add_section_heading(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, caption=""):
        return add_table_from_data(self.doc, headers, rows, col_widths, caption)

    # --------------------------------------------------------
    # 标题
    # --------------------------------------------------------

    def _add_title(self):
        """标题：黑体二号居中，固定为"文明大学生申请书" """
        title = "文明大学生申请书"
        custom_title = self._get("title", default="")
        if custom_title:
            title = custom_title
        add_title(self.doc, title)

    # --------------------------------------------------------
    # 称呼
    # --------------------------------------------------------

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号"""
        salutation = self._get("salutation",
                               default="尊敬的学院领导、评审委员会：")
        add_salutation_paragraph(self.doc, salutation)

    # --------------------------------------------------------
    # 开头段落
    # --------------------------------------------------------

    def _add_opening(self):
        """开头段落（80~120 字）：身份 + 申报奖项 + 文明素养亮点 + 道德事例数量 + 进入句"""
        opening = self._get("opening", default="")
        if opening:
            self.add_para(opening)
            return
        name = self._get("name")
        college = self._get("college")
        major = self._get("major")
        grade = self._get("grade")
        apply_year = self._get("apply_year", default="2024-2025 学年")
        honor_name = "文明大学生"
        gpa = self._get("gpa")
        rank = self._get("rank")
        position = self._get("position", default="")
        # 文明素养 4 维度亮点
        civility_highlights = self._build_civility_highlights()
        # 道德事例数量
        morality_count = len(self._get_list("morality_cases"))
        parts = []
        if name and college and major and grade:
            parts.append(f"我是{college}{major}{grade}学生{name}，特申请{apply_year}{honor_name}。")
        else:
            parts.append(f"特申请{apply_year}{honor_name}。")
        data_parts = []
        if gpa:
            data_parts.append(f"GPA {gpa}")
        if rank:
            data_parts.append(f"专业排名第 {rank}")
        if position:
            data_parts.append(f"担任{position}")
        if civility_highlights:
            data_parts.append(civility_highlights)
        if data_parts:
            parts.append(join_non_empty(data_parts, "；") + "。")
        if morality_count > 0:
            parts.append(f"本学年道德事例 {morality_count} 件。")
        parts.append("现将本学年情况汇报如下：")
        self.add_para("".join(parts))

    def _build_civility_highlights(self) -> str:
        """构建文明素养 4 维度亮点概述（用于开头段）"""
        highlights = []
        classroom = self._get_list("civility_classroom")
        if classroom:
            highlights.append("课堂零违规")
        dormitory = self._get_list("civility_dormitory")
        if dormitory:
            for d in dormitory:
                if isinstance(d, dict) and d.get("honor"):
                    highlights.append(d.get("honor"))
                    break
            else:
                highlights.append("宿舍文明")
        network = self._get_list("civility_network")
        if network:
            highlights.append("网络零违规")
        public = self._get_list("civility_public")
        if public:
            highlights.append("公共秩序良好")
        # 去重保留前 4 项
        seen = set()
        unique = []
        for h in highlights:
            if h and h not in seen:
                seen.add(h)
                unique.append(h)
            if len(unique) >= 4:
                break
        return "、".join(unique) if unique else ""

    # --------------------------------------------------------
    # 一、思想品德
    # --------------------------------------------------------

    def _add_ideology(self):
        """思想品德（200~280 字）：政治面貌 + 入党/团情况 + 价值观践行 + 思想觉悟

        文明大学生无强制政治面貌要求（党员/团员/群众均可），但需写明
        社会主义核心价值观践行活动（1~2 项具体活动）。
        """
        self.add_heading("一、思想品德")
        ideology = self._get("ideology", default="")
        if ideology:
            if isinstance(ideology, list):
                for p in ideology:
                    self.add_para(p)
            else:
                self.add_para(ideology)
            return
        political = self._get("political_status", default="共青团员")
        party_history = self._get("party_history", default="")
        ideology_activities = self._get_list("ideology_activities")
        core_values = self._get("core_values", default="")
        parts = []
        # 政治面貌开头
        if "党员" in political and "预备" not in political:
            parts.append(f"作为一名{political}，我认真学习习近平新时代中国特色社会主义思想，深入学习党的二十大和二十届三中全会精神，时刻以正式党员标准要求自己。")
        elif "预备党员" in political:
            parts.append("作为一名中共预备党员，我认真学习习近平新时代中国特色社会主义思想，深入学习党的二十大和二十届三中全会精神，时刻以正式党员标准要求自己。")
        elif "积极分子" in political:
            parts.append("作为一名入党积极分子，我认真学习习近平新时代中国特色社会主义思想，深入学习党的二十大和二十届三中全会精神，时刻以党员标准要求自己。")
        elif "共青团员" in political:
            parts.append(f"作为一名{political}，我拥护中国共产党的领导，认真学习习近平新时代中国特色社会主义思想，深入学习党的二十大和二十届三中全会精神。")
        else:
            parts.append(f"作为一名{political}，我拥护中国共产党的领导，认真学习党的创新理论，关注党的二十大和二十届三中全会精神。")
        if party_history:
            parts.append(party_history + "。")
        # 价值观践行
        if ideology_activities:
            parts.append("本学年参与" + "、".join(str(a) for a in ideology_activities) + "等活动。")
        elif core_values:
            parts.append(core_values + "。")
        parts.append("在日常生活中，我自觉践行社会主义核心价值观，注重理论学习与文明实践结合。")
        self.add_para("".join(parts))

    # --------------------------------------------------------
    # 二、文明素养【重点】含 4 维度表
    # --------------------------------------------------------

    def _add_civility(self):
        """文明素养（400~500 字，重点）：4 维度具体事例 + 4 维度表

        4 维度固定顺序：课堂纪律 / 宿舍卫生 / 网络文明 / 公共秩序。
        每维度至少 1 条具体事例，含时间与证明人。4 维度表 4 列：
        维度 / 具体表现 / 时间 / 证明人。
        """
        self.add_heading("二、文明素养")
        civility = self._get("civility", default="")
        if civility and isinstance(civility, str):
            self.add_para(civility)
            self._add_civility_table()
            return
        # 概述句
        self.add_para("本学年我在课堂纪律、宿舍卫生、网络文明、公共秩序 4 个维度持续践行文明标准。")
        # 4 维度详述
        self._add_civility_classroom()
        self._add_civility_dormitory()
        self._add_civility_network()
        self._add_civility_public()
        # 4 维度表
        self._add_civility_table()

    def _add_civility_classroom(self):
        """课堂纪律维度：零迟到零早退零手机违规 + 主动维护课堂秩序事例"""
        self._add_civility_dimension("civility_classroom", "课堂纪律", with_honor=False)

    def _add_civility_dormitory(self):
        """宿舍卫生维度：卫生检查得分 + 文明宿舍荣誉 + 宿舍长履职"""
        self._add_civility_dimension("civility_dormitory", "宿舍卫生", with_honor=True)

    def _add_civility_network(self):
        """网络文明维度：网络零违规 + 主动举报不实信息 + 网络文明志愿活动"""
        self._add_civility_dimension("civility_network", "网络文明", with_honor=False)

    def _add_civility_public(self):
        """公共秩序维度：食堂光盘 + 图书馆安静 + 礼让师生"""
        self._add_civility_dimension("civility_public", "公共秩序", with_honor=False)

    def _add_civility_dimension(self, key: str, dim_label: str, with_honor: bool = False):
        """通用文明维度渲染（课堂/网络/公共走同一逻辑；宿舍多 honor 字段）"""
        items = self._get_list(key)
        if not items:
            return
        parts = []
        for item in items:
            if not isinstance(item, dict):
                continue
            detail = safe_str(item.get("detail", ""))
            honor = safe_str(item.get("honor", "")) if with_honor else ""
            time = safe_str(item.get("time", ""))
            witness = safe_str(item.get("witness", ""))
            seg = detail
            if honor:
                seg += f"，{honor}"
            if time:
                seg += f"（{time}）"
            if witness:
                seg += f"，证明人：{witness}"
            if seg:
                parts.append(seg + "。")
        if parts:
            self.add_para(f"{dim_label}方面，" + "".join(parts))

    def _add_civility_table(self):
        """文明素养 4 维度表（4 列：维度/具体表现/时间/证明人，固定 4 行）

        - 维度列固定 4 行：课堂纪律 / 宿舍卫生 / 网络文明 / 公共秩序（不可省略任一维度）
        - 具体表现列含本学年量化数据 + 1~2 件具体事例（不写"等"等模糊词）
        - 时间列精确到学期或日期，不写"本学年"等模糊时间
        - 证明人列含姓名+职务，至少 1 名教师/管理人员
        - 若某维度本学年无事例，仍需列该维度行，具体表现写"本学年无具体事例"
        - 仅当至少 1 行有具体事例时才渲染表格，避免空表
        """
        rows = []
        dimension_keys = [
            ("civility_classroom", "课堂纪律"),
            ("civility_dormitory", "宿舍卫生"),
            ("civility_network", "网络文明"),
            ("civility_public", "公共秩序"),
        ]
        for key, dim_name in dimension_keys:
            items = self._get_list(key)
            if not items:
                # 无事例维度仍需列行，作为底线证明
                rows.append([dim_name, "本学年无具体事例", "", ""])
                continue
            details = []
            times = []
            witnesses = []
            for item in items:
                if not isinstance(item, dict):
                    continue
                d = safe_str(item.get("detail", ""))
                h = safe_str(item.get("honor", ""))
                seg = d
                if h:
                    # 宿舍维度含 honor 字段，拼接到 detail 后
                    seg = (seg + "；" if seg else "") + h
                if seg:
                    details.append(seg)
                t = safe_str(item.get("time", ""))
                if t:
                    times.append(t)
                w = safe_str(item.get("witness", ""))
                if w:
                    witnesses.append(w)
            # 多件事例用"；"分隔合并为 1 行
            detail_str = "；".join(details) if details else ""
            time_str = "+".join(times) if times else ""
            witness_str = "、".join(witnesses) if witnesses else ""
            rows.append([dim_name, detail_str, time_str, witness_str])
        # 仅当至少 1 行有具体事例时才渲染表格，避免空表
        has_content = any(row[1] for row in rows)
        if has_content:
            self.add_table(
                ["维度", "具体表现", "时间", "证明人"],
                rows,
                col_widths=CIVILITY_COL_WIDTHS,
                caption="文明素养 4 维度情况：",
            )

    # --------------------------------------------------------
    # 三、道德品质【重点】含事例表
    # --------------------------------------------------------

    def _add_morality(self):
        """道德品质（250~350 字，重点）：具体事例 + 事例表

        事例类型：拾金不昧 / 见义勇为 / 助人为乐等。每件事例含五要素：
        时间 / 地点 / 事件 / 影响 / 证明人。事例表 5 列固定。
        """
        self.add_heading("三、道德品质")
        morality = self._get("morality", default="")
        if morality and isinstance(morality, str):
            self.add_para(morality)
            self._add_morality_table()
            return
        cases = self._get_list("morality_cases")
        morality_summary = self._get("morality_summary", default="")
        if not cases:
            if morality_summary:
                self.add_para(morality_summary)
            else:
                self.add_para("本学年本人注重日常道德修养，无突出道德事件，但坚持在小事中践行传统美德。")
            return
        # 概述句
        if morality_summary:
            self.add_para(morality_summary)
        else:
            self.add_para(f"本学年我在日常生活中践行拾金不昧、助人为乐等传统美德，具体事例有：")
        # 详述每件事例
        for case in cases:
            if not isinstance(case, dict):
                continue
            seg = self._format_morality_case_text(case)
            if seg:
                self.add_para(seg)
        # 事例表
        self._add_morality_table()

    def _format_morality_case_text(self, case: Dict[str, Any]) -> str:
        """将单件道德事例格式化为正文段落（含时间地点事件影响证明人）"""
        time = safe_str(case.get("time", ""))
        place = safe_str(case.get("place", ""))
        event = safe_str(case.get("event", ""))
        impact = safe_str(case.get("impact", ""))
        witness = safe_str(case.get("witness", ""))
        parts = []
        if time:
            parts.append(time)
        if place:
            parts.append(f"在{place}")
        if event:
            parts.append(event + "，")
        else:
            parts.append("，")
        if impact:
            parts.append(impact + "。")
        else:
            parts.append("。")
        if witness:
            parts.append(f"证明人：{witness}。")
        return "".join(parts)

    def _add_morality_table(self):
        """道德事例表（5 列：时间/地点/事件/影响/证明人，至少 1 行）

        - 表头固定 5 列，不可增减
        - 至少 1 行，建议 2~3 行（事例数量影响评分）
        - 时间列精确到日（YYYY.MM.DD），不写"某天""上周"
        - 地点列具体到楼层/房间/站台，不写"校园内""某处"
        - 事件列含金额/人数/伤情等可量化要素
        - 影响列含失主反馈/医院诊断/同学恢复情况
        - 证明人列含姓名+职务，至少 1 名教师/管理人员
        """
        cases = self._get_list("morality_cases")
        if not cases:
            return
        rows = []
        for case in cases:
            if not isinstance(case, dict):
                continue
            # 五要素严格按顺序：时间/地点/事件/影响/证明人
            rows.append([
                safe_str(case.get("time", "")),
                safe_str(case.get("place", "")),
                safe_str(case.get("event", "")),
                safe_str(case.get("impact", "")),
                safe_str(case.get("witness", "")),
            ])
        if rows:
            self.add_table(
                ["时间", "地点", "事件", "影响", "证明人"],
                rows,
                col_widths=MORALITY_COL_WIDTHS,
                caption="道德事例情况：",
            )

    # --------------------------------------------------------
    # 四、学业与生活
    # --------------------------------------------------------

    def _add_academics_life(self):
        """学业与生活（200~280 字）：GPA + 加权 + 排名 + 宿舍长履职 + 人际关系

        文明大学生学业门槛低（无 GPA 前 30% 强制要求，无挂科即可），
        本段为辅助段，重点在宿舍长履职与人际关系具体事例。
        """
        self.add_heading("四、学业与生活")
        academics_life = self._get("academics_life", default="")
        if academics_life:
            if isinstance(academics_life, list):
                for p in academics_life:
                    self.add_para(p)
            else:
                self.add_para(academics_life)
            return
        gpa = self._get("gpa")
        weighted = self._get("weighted_avg")
        rank = self._get("rank")
        course_count = self._get("course_count")
        high_score_count = self._get("high_score_count")
        cet4 = self._get("cet4")
        cet6 = self._get("cet6")
        dorm_role = self._get("dorm_role", default="")
        dorm_activity = self._get("dorm_activity", default="")
        dorm_honor = self._get("dorm_honor", default="")
        interpersonal = self._get("interpersonal", default="")
        lifestyle = self._get("lifestyle", default="生活中我注重勤俭节约，作息规律。")
        parts = []
        # 学业句
        acad_parts = []
        if gpa:
            seg = f"本学年 GPA {gpa}"
            if weighted:
                seg += f"，加权平均分 {weighted}"
            if rank:
                seg += f"，专业排名第 {rank}"
            acad_parts.append(seg)
        if course_count and high_score_count:
            acad_parts.append(f"修读 {course_count} 门课程，{high_score_count} 门 80 分以上")
        lang_parts = []
        if cet4:
            lang_parts.append(f"CET-4 {cet4} 分")
        if cet6:
            lang_parts.append(f"CET-6 {cet6} 分")
        if lang_parts:
            acad_parts.append("、".join(lang_parts))
        if acad_parts:
            parts.append("。".join(acad_parts) + "。")
        if lifestyle:
            parts.append(lifestyle)
        # 宿舍长履职
        if dorm_role and dorm_activity:
            seg = f"担任{dorm_role}期间，{dorm_activity}"
            if dorm_honor:
                seg += f"，{dorm_honor}"
            parts.append(seg + "。")
        elif dorm_honor:
            parts.append(dorm_honor + "。")
        # 人际关系
        if interpersonal:
            parts.append(interpersonal + "。")
        if parts:
            self.add_para("".join(parts))

    # --------------------------------------------------------
    # 结尾"此致 敬礼！"
    # --------------------------------------------------------

    def _add_ending(self):
        """结尾（60~100 字）：本学年事实总结 + 朴素表态 + 此致/敬礼！"""
        ending = self._get("ending", default="")
        if ending:
            self.add_para(ending)
        else:
            self.add_para(
                "以上是我本学年的基本情况。无论评选结果如何，我都将继续以文明标准要求自己，"
                "在课堂、宿舍、网络、公共秩序 4 个维度持续践行，争取获评文明大学生。"
                "恳请评审委员会予以考虑。"
            )
        add_cizhi_paragraph(self.doc, "此致")   # "此致"另起一行，空两格
        add_jingli_paragraph(self.doc, "敬礼！")  # "敬礼！"另起一行，顶格

    # --------------------------------------------------------
    # 落款
    # --------------------------------------------------------

    def _add_signature(self):
        """落款：右对齐，含申请人 + 日期"""
        self.doc.add_paragraph()  # 空一行
        name = self._get("name", default="")
        apply_date = self._get("apply_date", default="")
        if name:
            add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if apply_date:
            add_right_aligned_paragraph(self.doc, apply_date)

    # --------------------------------------------------------
    # 主构建方法
    # --------------------------------------------------------

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/开头/思想品德/文明素养/道德品质/学业与生活/结尾/落款"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_opening()
            self._add_ideology()
            self._add_civility()
            self._add_morality()
            self._add_academics_life()
            self._add_ending()
            self._add_signature()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申请书已生成：{output_path}")
        return str(output_path)

    # --------------------------------------------------------
    # 数据校验（含文明事例校验）
    # --------------------------------------------------------

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）。重点校验文明素养 4 维度与道德事例的证明人字段（评审需向证明人核实）。"""
        warnings = []
        # P0 必采字段
        for key, label in [("name", "申请人姓名"), ("college", "学院"), ("major", "专业"),
                           ("grade", "年级"), ("apply_year", "申报学年")]:
            if not self._get(key):
                warnings.append(f"缺少 {label}（{key}）")
        # 文明素养 4 维度校验
        civility_keys = [
            ("civility_classroom", "课堂纪律"),
            ("civility_dormitory", "宿舍卫生"),
            ("civility_network", "网络文明"),
            ("civility_public", "公共秩序"),
        ]
        civility_with_content = 0
        for key, label in civility_keys:
            items = self._get_list(key)
            if not items:
                warnings.append(f"缺少 文明素养{label}维度（{key}），将填充为'本学年无具体事例'")
                continue
            civility_with_content += 1
            # 校验证明人字段
            for i, item in enumerate(items):
                if not isinstance(item, dict):
                    continue
                witness = item.get("witness", "")
                if not witness:
                    warnings.append(f"{label}维度第 {i+1} 条事例缺少证明人（witness），评审无法核实")
                detail = item.get("detail", "")
                if not detail:
                    warnings.append(f"{label}维度第 {i+1} 条事例缺少具体表现（detail）")
        if civility_with_content < 2:
            warnings.append(f"文明素养 4 维度仅 {civility_with_content} 项有具体事例，建议至少 2 项（硬门槛）")
        # 道德事例校验
        cases = self._get_list("morality_cases")
        if not cases:
            warnings.append("缺少 道德事例（morality_cases），建议至少 1 件（含五要素：时间/地点/事件/影响/证明人）")
        else:
            for i, case in enumerate(cases):
                if not isinstance(case, dict):
                    continue
                for field, label in [("time", "时间"), ("place", "地点"),
                                     ("event", "事件"), ("impact", "影响"),
                                     ("witness", "证明人")]:
                    if not case.get(field):
                        warnings.append(f"道德事例第 {i+1} 件缺少 {label}（{field}）")
        # 排名校验（如有排名需带基数）
        rank_str = safe_str(self._get("rank", default=""))
        if rank_str and "/" not in rank_str:
            if not self._get("rank_total"):
                warnings.append(f"排名 '{rank_str}' 缺基数，应为 'X/N' 格式或补填 rank_total")
        # 文明大学生无 GPA 前 30% 强制要求，仅需无挂科，此处不强制校验 GPA 阈值
        self.warnings = warnings
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据（文明大学生）
# ============================================================

DEFAULT_DATA = {
    "category": "civilized_student", "name": "张明", "student_id": "2022123456",
    "gender": "男", "college": "计算机科学与技术学院", "major": "计算机科学与技术",
    "grade": "2022 级 大三", "class_name": "计科 2201 班", "political_status": "共青团员",
    "phone": "138XXXXXXXX", "apply_year": "2024-2025 学年", "apply_date": "2025 年 5 月 20 日",
    "salutation": "尊敬的学院领导、评审委员会：",
    "gpa": "3.50/4.0", "weighted_avg": "82.5", "rank": "15/87", "rank_total": "87",
    "course_count": "12", "high_score_count": "8", "cet4": "510", "cet6": "",
    "party_history": "2019.05 在高中入团，2024.09 转入大学团支部，2024.09 递交入党申请书",
    "ideology_activities": ["2024.10 主题党日活动「红色教育基地走访」", "2024.11 校园反诈宣传志愿活动", "提交思想汇报 3 篇"],
    "core_values": "", "ideology": "",
    "civility_classroom": [{"detail": "12 门课程零迟到零早退零手机违规；2024.10 在《数据结构》课堂上发现前排同学遗留手机，主动转交任课教师并联系失主", "time": "2024.09-2025.05", "witness": "任课教师李老师"}],
    "civility_dormitory": [{"detail": "担任宿舍长期间，组织宿舍 6 次集体打扫，建立宿舍值日制度，本学年 6 次卫生检查平均 95 分", "honor": "宿舍连续两学期获评「文明宿舍」（院级）", "time": "2024 秋+2025 春", "witness": "辅导员王老师、宿管阿姨张师傅"}],
    "civility_network": [{"detail": "本学年网络言行零违规；2024.11 在校园论坛发现不实兼职招聘信息 1 条，主动向学校网络中心举报，协助避免 2 名同学上当；2025.03 参与校园网络文明志愿活动 1 次，撰写文明用帖倡议 1 份", "time": "2024.11+2025.03", "witness": "辅导员王老师"}],
    "civility_public": [{"detail": "食堂就餐光盘行动累计参与 60 余次；图书馆自习保持安静，主动提醒喧哗同学 4 次；上下楼梯礼让师生，主动搀扶行动不便老教师 2 次", "time": "2024.09-2025.05", "witness": "宿管阿姨张师傅、图书馆管理员刘老师"}],
    "civility": "",
    "morality_cases": [
        {"time": "2024.05.18", "place": "校园二食堂二楼", "event": "拾到钱包 1 个，内有现金 800 元、银行卡 3 张、校园卡 1 张", "impact": "立即交至学校保卫处，配合保卫处联系失主（数学学院 2021 级赵同学），失主领回时表示感谢", "witness": "保卫处周老师"},
        {"time": "2024.11.07", "place": "校园东门外公交站台", "event": "发现一位老人摔倒在地，立即上前扶起并拨打 120，陪同送至校医院检查", "impact": "老人轻微擦伤，已联系老人家属到场", "witness": "校医院值班护士孙医生"},
        {"time": "2025.03.15", "place": "图书馆自习室", "event": "发现邻座同学低血糖晕倒，立即取出随身携带的糖果让其含服，并通知图书馆管理员拨打校医院电话，陪同送医", "impact": "同学恢复健康，无需住院", "witness": "图书馆管理员刘老师"},
    ],
    "morality_summary": "", "morality": "",
    "position": "宿舍长", "dorm_role": "宿舍长",
    "dorm_activity": "组织宿舍 6 次集体活动，建立宿舍值日制度",
    "dorm_honor": "宿舍连续两学期获评「文明宿舍」（院级）",
    "interpersonal": "与同学相处融洽，曾帮助室友完成 1 次重要实验调试，调解同学间因作息时间产生的矛盾 2 次",
    "lifestyle": "生活中我注重勤俭节约，作息规律。",
    "academics_life": "", "ending": "",
    "honors": [{"time": "2024.11", "name": "校级文明宿舍（成员）", "level": "院级", "issuer": "计算机学院"}],
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="文明大学生申请书 docx 生成器（含 4 维度文明素养表与道德事例表）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=("示例：\n  python build.py --data data.json --out output.docx\n  python build.py --demo --out demo.docx\n\n"
                "JSON 字段定义详见 SKILL.md 第三章信息采集清单。\n"
                "必填：civility_classroom / civility_dormitory / civility_network / civility_public（4 维度文明素养）\n"
                "必填：morality_cases（道德事例表，至少 1 条含五要素）"),
    )
    parser.add_argument("--data", type=str, default=None, help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True, help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true", help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档（文明大学生）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
