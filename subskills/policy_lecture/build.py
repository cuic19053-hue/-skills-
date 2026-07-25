#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
大学生暑期"三下乡"社会实践-宣讲类立项申报书 docx 生成器

按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 一级标题：黑体三号，居中
- 二级标题：黑体小三，左对齐
- 三级标题：宋体四号加粗
- 表格：宋体五号，居中
- 实施方案使用按场次表（6 列：日期/地点/对象/人数/时长/主讲人）
- 宣讲大纲使用结构化表格（5 列：环节/时长/内容要点/通俗化表达/案例素材）

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第十二章。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)
SIZE_SAN = Pt(16)
SIZE_XIAO_SAN = Pt(15)
SIZE_SI = Pt(14)
SIZE_XIAO_SI = Pt(12)
SIZE_WU = Pt(10.5)

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
        space_before=12, space_after=12,
    )


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6, space_after=6,
    )


def add_heading_level3(doc, text: str):
    """三级标题：宋体四号加粗，左对齐"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6, space_after=3,
    )


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent,
        line_spacing=1.5,
    )


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None,
                        left_align_cols: Optional[List[int]] = None):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            if left_align_cols and j in left_align_cols:
                align = WD_ALIGN_PARAGRAPH.LEFT
            else:
                align = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False, alignment=align)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def fmt_money(amount) -> str:
    """金额格式化：整数转字符串 + 元"""
    try:
        return f"{int(amount)} 元"
    except (ValueError, TypeError):
        return str(amount)


def safe_int(value, default: int = 0) -> int:
    """安全转整数"""
    try:
        return int(str(value).strip())
    except (ValueError, TypeError):
        return default


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """三下乡社会实践-宣讲类立项申报书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def add_h1(self, text):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text):
        return add_heading_level2(self.doc, text)

    def add_h3(self, text):
        return add_heading_level3(self.doc, text)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_table(self, headers, rows, col_widths=None, left_align_cols=None):
        return add_table_from_data(self.doc, headers, rows, col_widths,
                                   left_align_cols)

    def add_page_break(self):
        add_page_break(self.doc)

    # 封面

    def _add_cover(self):
        """封面：黑体二号标题 + 4 行下划线信息"""
        for _ in range(3):
            self.doc.add_paragraph()
        title = "大学生暑期\u201c三下乡\u201d社会实践活动立项申报书"
        add_paragraph_with_format(self.doc, title, font_name=FONT_HEI, font_size=SIZE_ER,
                                  bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  first_line_indent=False, space_before=12, space_after=12)
        theme = self._get("theme", default="乡村振兴")
        subtitle = f"（{theme}主题·宣讲类）"
        add_paragraph_with_format(self.doc, subtitle, font_name=FONT_HEI, font_size=SIZE_SAN,
                                  bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  first_line_indent=False, space_after=24)
        for _ in range(3):
            self.doc.add_paragraph()
        project_name = self._get("team_name", default="赴 XX 县党的二十大精神宣讲团")
        info_items = [
            ("项目名称", project_name), ("团队名称", self._get("team_name")),
            ("申报单位", self._get("college")), ("申报日期", self._get("apply_date")),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI, font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True
        self.add_page_break()

    # 一、团队基本信息

    def _add_team_info(self):
        """一、团队基本信息表（7 行 2 列）"""
        self.add_h1("一、团队基本信息")
        team = self._get("team_info", default={})
        if not isinstance(team, dict):
            team = {}
        g = self._get
        leader = team.get("leader", f"{g('leader_name')} / {g('leader_id')} / "
                          f"{g('leader_major')} / {g('leader_grade')} / {g('leader_phone')}")
        advisor = team.get("advisor", f"{g('advisor_name')} / {g('advisor_title')} / "
                           f"{g('advisor_phone')} / {g('advisor_with_team', '随队')}")
        rows = [
            ["团队名称", team.get("team_name", g("team_name", ""))],
            ["实践主题", team.get("theme", g("theme", ""))],
            ["实践地点", team.get("location", g("location", ""))],
            ["实践时间", team.get("practice_time", g("practice_time", ""))],
            ["团队人数", team.get("team_size", g("team_size", ""))],
            ["队长", leader],
            ["指导教师", advisor],
        ]
        self.add_table(["项目", "内容"], rows, col_widths=[4.5, 11.5])

    # 二、团队成员信息表

    def _add_members_table(self):
        """二、团队成员信息表（5 列）"""
        self.add_h1("二、团队成员信息表")
        members = self._get("members", default=[])
        if members and isinstance(members, list):
            rows = []
            for m in members:
                if not isinstance(m, dict):
                    continue
                rows.append([m.get("name", ""), m.get("id", ""), m.get("major", ""),
                             m.get("role", ""), m.get("phone", "")])
            self.add_table(["姓名", "学号", "专业年级", "团队分工", "联系方式"],
                           rows, col_widths=[2.5, 2.8, 3.5, 3.5, 3.7])
        else:
            self.add_para("（请填写团队成员信息表，每人一行：姓名 / 学号 / 专业年级 / 团队分工 / 联系方式。"
                          "宣讲类建议分工：主讲 / 大纲撰写 / 案例采集 / 材料设计 / 后勤保障。"
                          "分工必须具体到\u201c做什么\u201d，不能只写\u201c成员\u201d。团队人数 6~15 人为宜。）")

    # 三、实践主题与背景

    def _add_theme_background(self):
        """三、实践主题与背景（400~600 字，3 段）"""
        self.add_h1("三、实践主题与背景")
        bg = self._get("theme_background", default=[])
        if isinstance(bg, str):
            bg = [bg]
        if bg:
            for para in bg:
                self.add_para(para)
        else:
            self.add_h2("（一）团中央主题对应")
            self.add_para("（请填写当年团中央发布的实践主题方向，以及本项目对应切入的维度，"
                          "150 字左右。如：2025 年团中央\u201c乡村振兴 青春建功\u201d主题，"
                          "本项目切入\u201c乡风文明\u201d维度开展政策宣讲。）")
            self.add_h2("（二）宣讲主题选定理由")
            self.add_para("（请填写主题选定理由，200 字左右，回答 3 问：①为什么选这个主题"
                          "（政策依据）②为什么是这个切入点（与当地相关性）③为什么由本团队讲"
                          "（团队专业契合度）。如：依据《中共中央关于认真学习宣传贯彻党的"
                          "二十大精神的决定》，切入点为乡村振兴维度，团队成员 4 人为"
                          "马克思主义理论专业。）")
            self.add_h2("（三）实践对象基本情况")
            self.add_para("（请填写实践对象基本情况，150 字左右。包括常住人口、年龄结构、"
                          "文化程度、关注点等。必须有数据，如常住人口 860 人、18~65 岁 520 人、"
                          "文化程度以小学和初中为主。）")

    # 四、实践目的与意义

    def _add_purpose_significance(self):
        """四、实践目的与意义（300~500 字，2 段）"""
        self.add_h1("四、实践目的与意义")
        ps = self._get("purpose_significance", default=[])
        if isinstance(ps, str):
            ps = [ps]
        if ps:
            for para in ps:
                self.add_para(para)
        else:
            self.add_h2("（一）对当地")
            self.add_para("（请填写对当地的意义，200~250 字。结构：宣讲 N 场覆盖 M 人次，"
                          "让 N1 名村干部掌握 X 政策要点、N2 名党员深化 Y 认识、"
                          "N3 名村民了解 Z 法律/技术。重点解决 X 个问题。）")
            self.add_h2("（二）对学生")
            self.add_para("（请填写对学生的意义，150~250 字。结构：①政治素养——通过备讲深化"
                          "对党的二十大精神理解 ②表达能力——脱稿宣讲锻炼 ③基层国情——与"
                          "村民互动深化认识。避免\u201c提升综合素质\u201d等空话。）")

    # 五、实践内容与实施方案【重点】含按场次表

    def _add_implementation_plan(self):
        """五、实践内容与实施方案（800~1200 字，3 子节 + 按场次表）"""
        self.add_h1("五、实践内容与实施方案")

        self.add_h2("（一）实践形式")
        form = self._get("implementation_plan", default={})
        if not isinstance(form, dict):
            form = {}
        form_text = form.get("form", "")
        if form_text:
            self.add_para(form_text)
        else:
            self.add_para("本项目实践形式为宣讲类，采用\u201c集中宣讲 + 互动答疑 + 材料发放\u201d"
                          "三段式实施流程。集中宣讲面向党员干部、普通村民、青少年三类对象"
                          "分场开展；互动答疑每场设置 10 道有奖问答；材料发放包含宣讲手册、"
                          "宣传单页、易拉宝、互动礼品四类。")

        self.add_h2("（二）宣讲对象分析")
        audience = form.get("audience_analysis", "")
        if audience:
            if isinstance(audience, list):
                for p in audience:
                    self.add_para(p)
            else:
                self.add_para(audience)
        else:
            self.add_para("（请填写对象分析，320 字左右，4 维度：①年龄结构——按 4 档划分"
                          "（18~35/36~55/56~70/70+）②文化程度——按 4 档划分（小学及以下/"
                          "初中/高中中专/大专及以上）③关注点——按对象类型列举（党员/村民/"
                          "妇女/青少年/老年）④接受方式——听讲/提问/讨论/情景模拟的时长占比。"
                          "必须基于预调研 10 户数据，不能凭空想象。）")

        self.add_h2("（三）按场次实施安排")
        schedule = form.get("schedule", [])
        if schedule and isinstance(schedule, list):
            rows = []
            total_count = 0
            for s in schedule:
                if not isinstance(s, dict):
                    continue
                count = safe_int(s.get("count", 0))
                total_count += count
                rows.append([
                    s.get("date", ""),
                    s.get("location", ""),
                    s.get("audience", ""),
                    f"{s.get('count', '')} 人" if s.get("count") != "" else "",
                    s.get("duration", ""),
                    s.get("speaker", ""),
                ])
            rows.append([
                "合计",
                f"{len(rows)} 场",
                "三类对象",
                f"{total_count} 人次",
                "—",
                "主讲轮换",
            ])
            self.add_table(
                ["日期", "地点", "对象", "人数", "时长", "主讲人"],
                rows,
                col_widths=[1.8, 3.5, 2.8, 2.0, 2.0, 3.9],
            )
            self.add_para("备注：每场结束当晚 19:00-20:00 召开日例会，整理当日反馈，"
                          "调整次日大纲。雨天启用备用场所（村委会会议室）。若某场到场人数"
                          "不足 60%（< 30 人），启用应急宣讲（入户小范围宣讲 + 留材料 30 份）。")
        else:
            self.add_para("（请填写按场次实施安排表，3~5 场。每行 6 列：日期 / 地点（精确到"
                          "具体场所）/ 对象（分类，不能笼统\u201c村民\u201d）/ 人数（按场地容量"
                          "80% 估算）/ 时长（党员 90 分钟、村民 60 分钟）/ 主讲人（2~3 人轮换）。"
                          "底部加合计行：N 场 × M 人次。示例：7.16 A 村党员活动室 党员+村干部"
                          " 35 人 90 分钟 张三；7.17 B 村文化礼堂 普通村民 60 人 60 分钟 李四；"
                          "7.18 C 村小学操场 青少年+家长 55 人 60 分钟 王五；合计 3 场 150 人次。）")

    # 六、宣讲大纲【宣讲类专属】

    def _add_lecture_outline(self):
        """六、宣讲大纲（宣讲类专属，结构化表格 5 列 × 7 行）"""
        self.add_h1("六、宣讲大纲")
        outline = self._get("lecture_outline", default={})
        if not isinstance(outline, dict):
            outline = {}
        theme = outline.get("theme", "")
        total_min = outline.get("total_minutes", "")
        if theme:
            self.add_para(f"宣讲主题：{theme}", indent=False)
        if total_min:
            self.add_para(f"总时长：{total_min} 分钟", indent=False)
        segments = outline.get("segments", [])
        if segments and isinstance(segments, list):
            rows = []
            for seg in segments:
                if not isinstance(seg, dict):
                    continue
                rows.append([
                    seg.get("phase", ""),
                    f"{seg.get('minutes', '')} 分钟",
                    seg.get("content", ""),
                    seg.get("plain", ""),
                    seg.get("case", ""),
                ])
            self.add_table(
                ["环节", "时长", "内容要点", "通俗化表达", "案例素材"],
                rows,
                col_widths=[1.8, 1.6, 5.0, 4.0, 3.6],
                left_align_cols=[2, 3, 4],
            )
            remark = outline.get("remark", "")
            if remark:
                self.add_para(remark)
            else:
                self.add_para("备注：针对不同对象做时长调整——党员干部主体加到 20 分钟、互动减到 5 分钟；"
                              "普通村民案例加到 15 分钟、互动加到 15 分钟；青少年案例加到 20 分钟、"
                              "主体每点减到 10 分钟。")
        else:
            self.add_para("（请填写宣讲大纲，结构化表格 5 列 × 7 行：环节 / 时长 / 内容要点 / "
                          "通俗化表达 / 案例素材。7 环节：开场 5 分钟 + 主体第 1 点 15 分钟 + "
                          "主体第 2 点 15 分钟 + 主体第 3 点 15 分钟 + 案例 10 分钟 + 互动 10 分钟 "
                          "+ 结尾 5 分钟，总时长 75 分钟。每点必须配\u201c通俗化表达\u201d列内容，"
                          "否则评审扣通俗化分（权重 20%）。案例素材必须本地化。）")

    # 七、安全保障预案【重点】

    def _add_safety_plan(self):
        """七、安全保障预案（300~500 字，5 段）"""
        self.add_h1("七、安全保障预案")
        plan = self._get("safety_plan", default=[])
        if isinstance(plan, str):
            plan = [plan]
        if plan and isinstance(plan, list):
            for para in plan:
                self.add_para(para)
        else:
            self.add_h2("（一）出行安全")
            self.add_para("（请填写出行安全，80 字：交通方式（统一购买高铁票）、住宿选择（政府招待所"
                          " 2 人一间男女分区）、每日 18:00 前向指导教师报平安。）")
            self.add_h2("（二）人身安全")
            self.add_para("（请填写人身安全，80 字：防暑（正午 11:00-14:00 不外出）、防疫（每日早晚"
                          "测温）、防骗（不携带大量现金、单人不离队）。）")
            self.add_h2("（三）应急联系人")
            self.add_para("（请填写应急联系人，60 字：3 名——指导教师、队长、学院团委，附电话。"
                          "当地对接：XX 县团委刘书记 136XXXXXXXX。）")
            self.add_h2("（四）应急流程")
            self.add_para("（请填写应急流程，120 字：分级响应。一般事件（< 1000 元）：队长 → 指导教师"
                          " → 学院存档。重大事件（≥ 1000 元 / 治安 / 灾害）：队长 → 指导教师 → "
                          "学院 → 110/120，30 分钟首报，2 小时书面报告。）")
            self.add_h2("（五）保险购买")
            self.add_para("（请填写保险购买，60 字：全员购买短期意外险，保额 30 万元，保费 15 元/人，"
                          "保单号 PA20250715001-010，覆盖 7 月 14-22 日。）")

    # 八、预期成果

    def _add_expected_results(self):
        """八、预期成果（必须可量化，场次+人次+材料印制为主）"""
        self.add_h1("八、预期成果")
        outcomes = self._get("expected_results", default=[])
        if isinstance(outcomes, str):
            outcomes = [outcomes]
        if outcomes:
            for o in outcomes:
                self.add_para(f"• {o}", indent=False)
        else:
            self.add_para("• 宣讲 3 场，覆盖 150 人次（村干部 15 + 党员 35 + 普通村民 100）",
                          indent=False)
            self.add_para("• 印制宣讲手册 200 份、宣传单页 500 份、易拉宝 3 套",
                          indent=False)
            self.add_para("• 新闻稿 5 篇（中青网 1 + 校团委公众号 2 + 学院公众号 2）",
                          indent=False)
            self.add_para("• 短视频 3 个（每个 3-5 分钟，发布在抖音、B 站）",
                          indent=False)
            self.add_para("• 纪录短片 1 个（10 分钟，用于校内汇报）", indent=False)
            self.add_para("• 反馈问卷 150 份（覆盖每场到场村民）", indent=False)

    # 九、经费预算

    def _add_budget(self):
        """九、经费预算（3 列表格：科目/金额/计算依据）"""
        self.add_h1("九、经费预算")
        items = self._get("budget_items", default=[])
        if items and isinstance(items, list):
            rows = []
            total = 0
            for b in items:
                if not isinstance(b, dict):
                    continue
                amount_num = safe_int(b.get("amount", 0))
                total += amount_num
                rows.append([
                    b.get("item", ""),
                    fmt_money(amount_num),
                    b.get("basis", ""),
                ])
            rows.append(["合计", fmt_money(total), ""])
            self.add_table(
                ["预算科目", "金额", "计算依据"],
                rows,
                col_widths=[3.5, 3.0, 9.5],
                left_align_cols=[2],
            )
        else:
            self.add_para("（请填写经费预算，5 类标准科目：交通费 / 食宿费 / 物资费 / 印刷费 / 其他。"
                          "宣讲类印刷费占比建议 30~40%。每项金额非整数，附计算依据。"
                          "示例：印刷费 3680 元 = 手册 200×8 + 单页 500×1.2 + 易拉宝 3×360。）")

    # 十、宣传计划

    def _add_publicity_plan(self):
        """十、宣传计划（3 类媒体）"""
        self.add_h1("十、宣传计划")
        plan = self._get("publicity_plan", default=[])
        if isinstance(plan, str):
            plan = [plan]
        if plan and isinstance(plan, list):
            for p in plan:
                self.add_para(p)
        else:
            self.add_h2("（一）校内媒体")
            self.add_para("校团委公众号 2 篇、校报 1 篇、学院公众号 2 篇。"
                          "发布时间节点：7.16 启动报道、7.19 中期进展、7.21 结项报道。")
            self.add_h2("（二）校外媒体")
            self.add_para("中青网 1 篇（团中央指定平台，团队队长审核后投递，附宣讲"
                          "照片 5~8 张）、XX 县电视台 1 条、XX 日报 1 篇。")
            self.add_h2("（三）新媒体")
            self.add_para("B 站 3 个短视频、抖音 3 个短视频、微信视频号 2 个。"
                          "每个视频 3~5 分钟，内容覆盖宣讲过程、村民互动、"
                          "团队风采三个角度。")

    # 十一、指导教师 + 学院团委意见

    def _add_review_section(self):
        """十一/十二/十三 签字栏（双栏 + 可选学校审批）"""
        def _sign_block(title, sign_label):
            self.add_h1(title)
            for _ in range(6):
                self.doc.add_paragraph()
            self.add_para(
                f"{sign_label}：____________________    日期：______年____月____日",
                indent=False,
            )
        _sign_block("十一、指导教师意见", "指导教师签字")
        _sign_block("十二、学院团委意见", "学院盖章")
        if self._get("include_school_approval", default=False):
            _sign_block("十三、学校审批意见", "学校盖章")

        special = self._get("tuancentral_special", default="")
        if special:
            self.add_h1("附：团中央专项报送说明")
            self.add_para(f"对应专项：{special}。本项目与专项主题契合度高，预期形成"
                          "优秀宣讲案例 1 个、宣讲视频 1 个，报送团中央专项工作组。")

    # 主构建方法

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 11 栏目，生成 docx。返回实际保存路径。"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_cover()
            self._add_team_info()
            self._add_members_table()
            self._add_theme_background()
            self._add_purpose_significance()
            self._add_implementation_plan()
            self._add_lecture_outline()
            self._add_safety_plan()
            self._add_expected_results()
            self._add_budget()
            self._add_publicity_plan()
            self._add_review_section()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申报书已生成：{output_path}")
        return str(output_path)

    # 数据校验（含宣讲类专属校验）

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）

        校验分四类：P0 必填字段、宣讲专属校验、安全预案校验、经费预算校验。
        """
        warnings = []
        for key, name in [("team_name", "团队名称"), ("theme", "实践主题"),
                          ("location", "实践地点"), ("practice_time", "实践时间"),
                          ("leader_name", "队长姓名"), ("college", "申报单位")]:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        team = self._get("team_info", default={})
        if not isinstance(team, dict):
            team = {}
        if not team.get("team_size") and not self._get("team_size"):
            warnings.append("缺少 团队人数（team_size）")
        if not team.get("advisor") and not self._get("advisor_name"):
            warnings.append("缺少 指导教师（advisor_name）")
        if not self._get("theme_background"):
            warnings.append("缺少 实践主题与背景（theme_background），将使用占位文本")
        if not self._get("purpose_significance"):
            warnings.append("缺少 实践目的与意义（purpose_significance），将使用占位文本")

        impl = self._get("implementation_plan", default={})
        if not isinstance(impl, dict):
            impl = {}
        if not impl.get("schedule"):
            warnings.append("缺少 按场次实施安排（implementation_plan.schedule），将使用占位文本——评审会扣大分")
        else:
            total_count = 0
            for i, s in enumerate(impl.get("schedule", []), 1):
                if not isinstance(s, dict):
                    continue
                if not s.get("date") or not s.get("location"):
                    warnings.append(f"按场次表第 {i} 行缺少日期或地点")
                if not s.get("audience"):
                    warnings.append(f"按场次表第 {i} 行缺少对象（不能笼统\u201c村民\u201d）")
                if not s.get("speaker"):
                    warnings.append(f"按场次表第 {i} 行缺少主讲人")
                total_count += safe_int(s.get("count", 0))
            if total_count > 0:
                self._session_total = total_count
        if not impl.get("audience_analysis"):
            warnings.append("缺少 宣讲对象分析（implementation_plan.audience_analysis），评审会扣针对性分")
        if not impl.get("form"):
            warnings.append("缺少 实践形式（implementation_plan.form），建议明确为宣讲类")

        outline = self._get("lecture_outline", default={})
        if not isinstance(outline, dict) or not outline.get("segments"):
            warnings.append("缺少 宣讲大纲（lecture_outline.segments），将使用占位文本——宣讲类核心栏目")
        else:
            segs = outline.get("segments", [])
            if len(segs) != 7:
                warnings.append(f"宣讲大纲环节应为 7 个（开场/3 点/案例/互动/结尾），当前 {len(segs)} 个")
            for i, seg in enumerate(segs, 1):
                if isinstance(seg, dict) and not seg.get("plain"):
                    warnings.append(f"大纲第 {i} 环节缺少通俗化表达——评审会扣通俗化分")

        plan = self._get("safety_plan", default=[])
        if not plan:
            warnings.append("缺少 安全保障预案（safety_plan），将使用占位文本——安全预案不达标可能一票否决")
        else:
            plan_text = "\n".join(plan) if isinstance(plan, list) else str(plan)
            if "保险" not in plan_text:
                warnings.append("安全预案未提及保险购买情况——一票否决项")
            if "应急联系" not in plan_text and "应急联系人" not in plan_text:
                warnings.append("安全预案未列明应急联系人")
            if "应急流程" not in plan_text:
                warnings.append("安全预案未列明应急流程")

        results = self._get("expected_results", default=[])
        if not results:
            warnings.append("缺少 预期成果（expected_results），将使用占位文本")
        else:
            results_text = "\n".join(results) if isinstance(results, list) else ""
            if "场次" not in results_text and "场" not in results_text:
                warnings.append("预期成果未以宣讲场次为主——宣讲类核心产出缺失")
            if "人次" not in results_text:
                warnings.append("预期成果未注明覆盖人次")
            if "印制" not in results_text and "手册" not in results_text and "单页" not in results_text:
                warnings.append("预期成果未列明材料印制——宣讲类核心产出缺失")

        items = self._get("budget_items", default=[])
        if items:
            total = sum(safe_int(b.get("amount", 0)) for b in items if isinstance(b, dict))
            budget_total_num = safe_int(str(self._get("budget_total", default="")).strip(), default=-1)
            if budget_total_num >= 0 and total != budget_total_num:
                warnings.append(f"预算合计 {total} 元 与申请经费 {budget_total_num} 元不一致")
        else:
            warnings.append("缺少 经费预算（budget_items），将使用占位文本")

        members = self._get("members", default=[])
        if members and isinstance(members, list):
            for i, m in enumerate(members, 1):
                if isinstance(m, dict) and not m.get("role"):
                    warnings.append(f"成员 {m.get('name', f'#{i}')} 缺少分工")

        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        self.warnings = warnings
        return warnings


# ============================================================
# 默认示例数据
# ============================================================

DEFAULT_DATA = {
    "team_name": "赴 XX 县党的二十大精神宣讲团", "theme": "乡村振兴",
    "location": "XX 省 XX 县 XX 镇 3 个行政村（A/B/C 村）",
    "practice_time": "2025.07.15-07.21（7 天）", "team_size": "10 人",
    "leader_name": "张三", "leader_id": "202212345",
    "leader_major": "马克思主义理论 2022 级", "leader_grade": "大三",
    "leader_phone": "138XXXXXXXX", "advisor_name": "李教授",
    "advisor_title": "副教授", "advisor_phone": "139XXXXXXXX",
    "advisor_with_team": "随队", "college": "马克思主义学院",
    "apply_date": "2025 年 5 月 20 日",
    "team_info": {
        "team_name": "赴 XX 县党的二十大精神宣讲团",
        "theme": "乡村振兴（对应 2025 年团中央“乡村振兴 青春建功”主题）",
        "location": "XX 省 XX 县 XX 镇 3 个行政村（A/B/C 村）",
        "practice_time": "2025.07.15-07.21（7 天）", "team_size": "10 人",
        "leader": "张三 / 202212345 / 马克思主义理论 2022 级 / 大三 / 138XXXXXXXX",
        "advisor": "李教授 / 副教授 / 139XXXXXXXX / 随队",
    },
    "members": [
        {"name": "张三", "id": "202212345", "major": "马克思主义理论 2022 级", "role": "队长 / 主讲（第 1、3 场）+ 主体第 1 点撰写", "phone": "138XXXXXXXX"},
        {"name": "李四", "id": "202212346", "major": "法学 2022 级", "role": "主讲（第 2 场）+ 案例采集", "phone": "138XXXXXXXX"},
        {"name": "王五", "id": "202212347", "major": "新闻学 2022 级", "role": "宣传报道 + 短视频制作", "phone": "138XXXXXXXX"},
        {"name": "赵六", "id": "202212348", "major": "设计学 2023 级", "role": "材料设计（手册/单页/易拉宝）", "phone": "138XXXXXXXX"},
        {"name": "孙七", "id": "202212349", "major": "社会学 2023 级", "role": "对象分析 + 反馈问卷", "phone": "138XXXXXXXX"},
    ],
    "theme_background": [
        "团中央主题对应：2025 年团中央发布“乡村振兴 青春建功”暑期社会实践主题，鼓励高校学生深入乡村开展政策宣讲、调研服务等活动。本项目紧扣该主题，聚焦“全面推进乡村振兴”维度，开展党的二十大精神宣讲，对应“乡风文明”切入方向。",
        "宣讲主题选定理由：依据《中共中央关于认真学习宣传贯彻党的二十大精神的决定》，本项目选择宣讲党的二十大精神，切入点为“乡村振兴”维度——XX 县 2024 年被列入省级乡村振兴示范县，但 5 个调研村中 3 个村村民对“中国式现代化”概念认知模糊（预调研数据），需重点宣讲“全面推进乡村振兴”部署。团队成员 4 人为马克思主义理论专业，已修读《习近平新时代中国特色社会主义思想概论》并获优秀。",
        "实践对象基本情况：XX 镇常住人口 860 人，其中 18~65 岁 520 人。预调研 10 户显示：文化程度以小学（占 50%）和初中（占 30%）为主；关注点集中在“医保报销比例”“土地流转政策”“防诈骗”三项；接受方式以听讲为主，60% 受访者希望主讲人用方言辅助讲解。",
    ],
    "purpose_significance": [
        "对当地：3 场宣讲覆盖 150 人次，让 15 名村干部掌握“乡村振兴战略”5 大维度（产业兴旺、生态宜居、乡风文明、治理有效、生活富裕），35 名党员深化对“中国式现代化”理解，100 名村民了解“医保报销比例 + 土地流转政策 + 防诈骗”3 项核心政策。重点解决“政策落地最后一公里”问题，提升村民对乡村振兴战略的认知度。",
        "对学生：①政治素养——通过 2 周备讲深化对党的二十大精神理解，全员通读《习近平著作选读》第一卷、第二卷；②表达能力——脱稿宣讲锻炼，全员完成 3 轮试讲互评；③基层国情——与 150 户村民互动深化认识，形成 30 份村民反馈记录，厚植家国情怀与三农感情。",
    ],
    "implementation_plan": {
        "form": "本项目实践形式为宣讲类，采用“集中宣讲 + 互动答疑 + 材料发放”三段式实施流程。集中宣讲面向党员干部、普通村民、青少年三类对象分场开展；互动答疑每场设置 10 道有奖问答；材料发放包含宣讲手册、宣传单页、易拉宝、互动礼品四类。",
        "audience_analysis": "年龄结构：宣讲对象以 56~70 岁老年人为主（占 60%），36~55 岁中年人占 30%，18~35 岁青年人占 10%——以老年人为主，需放大字号、放慢语速、多用本地案例。文化程度：小学及以下占 50%、初中占 30%、高中/中专占 15%、大专及以上占 5%。文化程度以小学和初中为主，宣讲大纲通俗化表达占比 ≥ 60%，所有专业术语必须配打比方解释。关注点：预调研显示，本村村民关注点集中在“医保报销比例”（80% 受访者提及）、“土地流转政策”（60%）、“防诈骗”（50%）。宣讲大纲主体 3 点对应覆盖。接受方式：宣讲采用“70% 讲解 + 30% 互动”模式。讲解部分用方言辅助（团队成员 2 人为本地籍贯），互动部分设置 10 道有奖问答（答对发印有政策标语的环保袋）。",
        "schedule": [
            {"date": "7.16", "location": "A 村党员活动室", "audience": "党员 + 村干部", "count": 35, "duration": "90 分钟", "speaker": "张三（主讲）+ 李四（助教）"},
            {"date": "7.17", "location": "B 村文化礼堂", "audience": "普通村民", "count": 60, "duration": "60 分钟", "speaker": "李四（主讲）+ 王五（助教）"},
            {"date": "7.18", "location": "C 村小学操场", "audience": "青少年 + 家长", "count": 55, "duration": "60 分钟", "speaker": "张三（主讲）+ 王五（助教）"},
        ],
    },
    "lecture_outline": {
        "theme": "党的二十大精神——乡村振兴维度", "total_minutes": 75,
        "segments": [
            {"phase": "开场", "minutes": 5, "content": "自我介绍 + 主题引出 + 提问“什么是乡村振兴”", "plain": "—", "case": "—"},
            {"phase": "主体 1", "minutes": 15, "content": "党的二十大对乡村振兴的部署（5 大维度）", "plain": "让村里“有钱挣、住得舒、风气正、办事顺、日子美”", "case": "—"},
            {"phase": "主体 2", "minutes": 15, "content": "5 大维度详解（产业/生态/乡风/治理/生活）", "plain": "5 句话 5 个字，村民一听就懂", "case": "—"},
            {"phase": "主体 3", "minutes": 15, "content": "实施路径（党建引领 + 产业升级 + 人才回引）", "plain": "党员带头干、产业做强干、年轻人回来干", "case": "—"},
            {"phase": "案例", "minutes": 10, "content": "本村致富带头人王某案例剖析", "plain": "王某 2024 年收入 18 万元", "case": "王某土地流转 + 特色种植"},
            {"phase": "互动", "minutes": 10, "content": "有奖问答 10 题（5 题主体 + 3 题案例 + 2 题防诈）", "plain": "答对发印有政策标语的环保袋", "case": "—"},
            {"phase": "结尾", "minutes": 5, "content": "总结 + 号召“乡村振兴从我做起” + 留联系方式", "plain": "—", "case": "—"},
        ],
        "remark": "备注：针对不同对象做时长调整——党员干部主体加到 20 分钟、互动减到 5 分钟；普通村民案例加到 15 分钟、互动加到 15 分钟；青少年案例加到 20 分钟、主体每点减到 10 分钟。",
    },
    "safety_plan": [
        "一、出行安全：全员统一购买高铁票，不单独行动；7 月 15 日集体乘 GXXX 次列车赴 XX 县。住宿选择 XX 县政府招待所（已与县团委对接预订），2 人一间，男女分区。每日 18:00 前向指导教师报平安。",
        "二、人身安全：携带常用药品（藿香正气水、创可贴、退烧药、止泻药）；正午 11:00-14:00 不外出，避免中暑；每日早晚测温，体温 ≥ 37.3℃ 立即隔离观察；不携带大量现金，单人不离队。",
        "三、应急联系人：指导教师李教授（139XXXXXXXX）、队长张三（138XXXXXXXX）、学院团委王老师（137XXXXXXXX）。当地对接：XX 县团委刘书记（136XXXXXXXX）。",
        "四、应急流程：一般事件（伤病 < 1000 元）：队长处理 → 指导教师报备 → 学院团委存档。重大事件（伤病 ≥ 1000 元 / 治安事件 / 自然灾害）：队长 → 指导教师 → 学院 → 110/120，同步报告当地团委。所有事件 30 分钟内首报，2 小时内书面报告。",
        "五、保险购买：全员购买中国平安短期意外险（保额 30 万元，保费 15 元/人，保单号 PA20250715001-010），覆盖 7 月 14-22 日。",
    ],
    "expected_results": [
        "宣讲 3 场，覆盖 150 人次（村干部 15 + 党员 35 + 普通村民 100）",
        "印制宣讲手册 200 份、宣传单页 500 份、易拉宝 3 套",
        "新闻稿 5 篇（中青网 1 + 校团委公众号 2 + 学院公众号 2）",
        "短视频 3 个（每个 3-5 分钟，发布在抖音、B 站）",
        "纪录短片 1 个（10 分钟，用于校内汇报）",
        "反馈问卷 150 份（覆盖每场到场村民）",
    ],
    "budget_items": [
        {"item": "交通费", "amount": "2860", "basis": "高铁往返 130 元 × 10 人 × 2 次 + 县内包车 300 元/天 × 4 天"},
        {"item": "食宿费", "amount": "5600", "basis": "招待所 120 元/间 × 5 间 × 7 天 + 餐费 40 元/人/天 × 10 人 × 7 天"},
        {"item": "物资费", "amount": "850", "basis": "互动礼品 50 个 × 4 元 + 文具 50 套 × 5 元 + 药品 150 元"},
        {"item": "印刷费", "amount": "3680", "basis": "手册 200 × 8 元 + 单页 500 × 1.2 元 + 易拉宝 3 × 360 元"},
        {"item": "其他", "amount": "310", "basis": "保险 15 元 × 10 人 + 通讯补贴 16 元 × 10 人"},
    ],
    "budget_total": "13300",
    "publicity_plan": [
        "（一）校内媒体：校团委公众号 2 篇、校报 1 篇、学院公众号 2 篇。发布时间节点：7.16 启动报道、7.19 中期进展、7.21 结项报道。",
        "（二）校外媒体：中青网 1 篇（团中央指定平台，团队队长审核后投递，附宣讲照片 5~8 张）、XX 县电视台 1 条、XX 日报 1 篇。",
        "（三）新媒体：B 站 3 个短视频、抖音 3 个短视频、微信视频号 2 个。每个视频 3~5 分钟，内容覆盖宣讲过程、村民互动、团队风采三个角度。",
    ],
    "include_school_approval": False,
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="三下乡社会实践-宣讲类立项申报书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="示例：\n  python build.py --data data.json --out output.docx\n"
               "  python build.py --demo --out demo.docx\n\nJSON 字段定义详见 SKILL.md 第十二章。",
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True, help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true", help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
