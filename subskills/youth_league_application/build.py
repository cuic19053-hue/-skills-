#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""入团申请书 docx 生成器（v1.0）

依据 SKILL.md v1.0 规范实现。核心特性：
1. 开篇首句固定为「我申请加入中国共产主义青年团。」，全文禁出现「志愿书」三字
   （与《入团志愿书》区分；单独「志愿」二字合规）
2. 字数 3 档分级（word_count_version）：compact=1500 / standard=1800 / enhanced=2000
3. 段落结构按字数版本路由（4 段配比 8%/38%/38%/16%）
4. 4 项政治理论表述完整版必引（共青团性质/团章/团史/五四精神）
5. 申请人年龄校验（须满 14 周岁且不满 28 周岁，按递交日期计算）
6. school_tier 字段用于字数下限校验（loose/normal/strict）
7. 家庭成员附页（require_family_appendix=true 时附录生成）
8. 落款日期汉字格式自动规范化（"2025.6.15" → "2025 年 6 月 15 日"）

格式标准：A4，页边距上下 2.54cm 左右 2.5cm；标题黑体二号居中；
称呼"敬爱的团组织："顶格宋体小四全角冒号；正文宋体小四 1.5 倍行距首行缩进 2 字符；
"此致"另起一行空两格，"敬礼！"另起一行顶格；落款右对齐。

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_KAI = "楷体"
SIZE_ER = Pt(22)            # 二号
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TB_CM = 2.54
MARGIN_LR_CM = 2.5

# 3 档字数版本（compact/standard/enhanced）
WORD_COUNT_VERSIONS: Dict[str, Dict[str, Any]] = {
    "compact":   {"target": 1500, "label": "精简版", "min": 1200,
                  "story_count": 2, "shortcoming_count": 2},
    "standard":  {"target": 1800, "label": "标准版", "min": 1500,
                  "story_count": 3, "shortcoming_count": 3},
    "enhanced":  {"target": 2000, "label": "加强版", "min": 1700,
                  "story_count": 3, "shortcoming_count": 3},
}

# 3 档学校档位（loose/normal/strict）
SCHOOL_TIERS: Dict[str, Dict[str, Any]] = {
    "loose":   {"min": 1200, "label": "宽松校（中学/大学新生）",
                "default_version": "compact"},
    "normal":  {"min": 1500, "label": "普通本科（默认档位）",
                "default_version": "standard"},
    "strict":  {"min": 1700, "label": "985/211 严格校",
                "default_version": "enhanced"},
}

# 4 项必引政治理论表述 ID（SKILL.md §6）
THEORY_QUOTE_IDS = [
    "league_nature",          # 共青团性质
    "league_constitution",    # 团章
    "league_history",         # 团史
    "may_fourth_spirit",      # 五四精神
]

# 入团誓词（80 字，团章原文）——仅当 include_oath=true 时引用
LEAGUE_OATH_TEXT = (
    "我志愿加入中国共产主义青年团，坚决拥护中国共产党的领导，"
    "遵守团的章程，执行团的决议，履行团员义务，严守团的纪律，"
    "勤奋学习，积极工作，吃苦在前，享受在后，为共产主义事业而奋斗。"
)

# ============================================================
# 工具函数：段落与字体
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_para_fmt(doc, text: str, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 alignment=WD_ALIGN_PARAGRAPH.LEFT,
                 first_line_indent: bool = True,
                 line_spacing: float = 1.5,
                 space_before: float = 0, space_after: float = 0):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """黑体二号居中标题，段前段后 12pt"""
    return add_para_fmt(doc, text, font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        first_line_indent=False, space_before=12, space_after=12)


def add_subtitle(doc, text: str):
    """黑体三号居中（家庭成员附页标题）"""
    return add_para_fmt(doc, text, font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        first_line_indent=False, space_before=12, space_after=12)


def add_body(doc, text: str, indent: bool = True):
    """宋体小四正文，1.5 倍行距，首行缩进 2 字符"""
    return add_para_fmt(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        first_line_indent=indent, line_spacing=1.5)


def add_salutation(doc, text: str):
    """顶格（不缩进）称呼"""
    return add_para_fmt(doc, text, first_line_indent=False, line_spacing=1.5)


def add_cizhi(doc, text: str = "此致"):
    """'此致'另起一行，空两格"""
    return add_para_fmt(doc, text, first_line_indent=True, line_spacing=1.5)


def add_jingli(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格"""
    return add_para_fmt(doc, text, first_line_indent=False, line_spacing=1.5)


def add_right(doc, text: str):
    """右对齐段落（落款用）"""
    return add_para_fmt(doc, text, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                        first_line_indent=False, line_spacing=1.5)


def add_blank(doc):
    """空段落（段落间留白）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return p


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TB_CM)
    section.bottom_margin = Cm(MARGIN_TB_CM)
    section.left_margin = Cm(MARGIN_LR_CM)
    section.right_margin = Cm(MARGIN_LR_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for fld_type in ("begin", None, "end"):
        if fld_type:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), fld_type)
            run._r.append(fld)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = "PAGE"
            run._r.append(instr)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_table(doc, headers: List[str], rows: List[List[str]]):
    """添加宋体小四的表格（家庭成员附页用）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, font_name=FONT_HEI, font_size=SIZE_XIAO_SI, bold=True)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = table.rows[ri].cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run_font(run, font_name=FONT_SONG, font_size=SIZE_XIAO_SI)
    return table


# ============================================================
# 工具函数：日期/年龄/字数
# ============================================================

def count_chinese_chars(text: str) -> int:
    """统计中文字符数"""
    if not text:
        return 0
    return len(re.findall(r"[\u4e00-\u9fff]", str(text)))


def normalize_date(s: str) -> str:
    """规范化日期为汉字格式（SKILL.md §8.4）

    - "2025 年 6 月 15 日"（已规范）
    - "2025.6.15" / "2025-06-15" / "2025/6/15"（自动转换并警告）
    """
    if not s:
        return s
    s = str(s).strip()
    if re.match(r"^\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日$", s):
        return re.sub(r"\s+", "", s)
    if re.match(r"^\d{4}\s*年\s*\d{1,2}\s*月$", s):
        return re.sub(r"\s+", "", s)
    m = re.match(r"^(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})$", s)
    if m:
        y, mo, d = m.groups()
        normalized = f"{y} 年 {int(mo)} 月 {int(d)} 日"
        print(f"⚠️ 日期格式不规范：'{s}'，已自动转换为 '{normalized}'", file=sys.stderr)
        return normalized
    m = re.match(r"^(\d{4})[.\-/](\d{1,2})$", s)
    if m:
        y, mo = m.groups()
        normalized = f"{y} 年 {int(mo)} 月"
        print(f"⚠️ 日期格式不规范：'{s}'，已自动转换为 '{normalized}'", file=sys.stderr)
        return normalized
    return s


def _parse_date(s: str) -> Optional[Tuple[int, int, int]]:
    """解析日期字符串为 (y, m, d) 元组，无法解析返回 None"""
    if not s:
        return None
    s = str(s).strip()
    m = re.match(r"^(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日$", s)
    if m:
        return tuple(int(x) for x in m.groups())
    m = re.match(r"^(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})$", s)
    if m:
        return tuple(int(x) for x in m.groups())
    m = re.match(r"^(\d{4})\s*年\s*(\d{1,2})\s*月$", s)
    if m:
        return (int(m.group(1)), int(m.group(2)), 1)
    m = re.match(r"^(\d{4})[.\-/](\d{1,2})$", s)
    if m:
        return (int(m.group(1)), int(m.group(2)), 1)
    return None


def compute_age(birth_date: str, submit_date: str) -> Optional[int]:
    """根据出生日期与递交日期计算周岁（SKILL.md §0.2）"""
    b = _parse_date(birth_date)
    s = _parse_date(submit_date)
    if not b or not s:
        return None
    age = s[0] - b[0]
    if (s[1], s[2]) < (b[1], b[2]):
        age -= 1
    return age


# ============================================================
# 工具函数：政治规范检查
# ============================================================

def check_political_terms(text: str) -> List[str]:
    """政治用语规范校验：必提要点 + 禁用简写检测"""
    warnings = []
    if not text:
        return warnings
    # 必提要点（5 项缺一不可）
    required_terms = [
        "中国共产主义青年团",
        "习近平新时代中国特色社会主义思想",
        "党的二十大",
        "共产主义",
        "助手和后备军",
    ]
    for term in required_terms:
        if term not in text:
            warnings.append(f"必提要点缺失：'{term}'")
    # 禁用简写检测
    forbidden = [
        ("习近平思想", "应为'习近平新时代中国特色社会主义思想'"),
        ("习思想", "应为'习近平新时代中国特色社会主义思想'"),
        ("党的20大", "应为'党的二十大'"),
        ("马列", "应为'马克思列宁主义'"),
        ("毛思想", "应为'毛泽东思想'"),
        ("邓理论", "应为'邓小平理论'"),
        ("科学观", "应为'科学发展观'"),
    ]
    for abbr, suggestion in forbidden:
        if abbr in text:
            warnings.append(f"禁用简写'{abbr}'：{suggestion}")
    # 指导思想 6 项顺序检测
    guideline = (
        r"马克思列宁主义.*?毛泽东思想.*?邓小平理论.*?"
        r"[\u201c\"']三个代表[\u201d\"']重要思想.*?科学发展观.*?习近平新时代中国特色社会主义思想"
    )
    if not re.search(guideline, text, re.DOTALL):
        required_keys = ["马克思列宁主义", "毛泽东思想", "邓小平理论",
                         "三个代表重要思想", "科学发展观", "习近平新时代中国特色社会主义思想"]
        found = {k for k in required_keys if k in text}
        missing = [k for k in required_keys if k not in found]
        if missing:
            warnings.append(f"指导思想漏项：{', '.join(missing)}")
        else:
            warnings.append("指导思想 6 项顺序错误（应为马列→毛→邓→三→科→习）")
    return warnings


def check_4_theory_expressions(text: str) -> List[str]:
    """4 项必引政治理论表述校验（SKILL.md §6.5）"""
    checks = {
        "共青团性质": "中国共产主义青年团是中国共产党领导的先进青年的群团组织",
        "团章-团员义务-马列": "马克思列宁主义",
        "团章-团员义务-毛": "毛泽东思想",
        "团章-团员义务-邓": "邓小平理论",
        "团章-团员义务-三": "\u201c三个代表\u201d重要思想",
        "团章-团员义务-科": "科学发展观",
        "团章-团员义务-习": "习近平新时代中国特色社会主义思想",
        "团史-1922": "1922 年 5 月 5 日",
        "团史-1925": "1925 年 1 月",
        "团史-1957": "1957 年 5 月",
        "五四精神-爱国": "爱国",
        "五四精神-进步": "进步",
        "五四精神-民主": "民主",
        "五四精神-科学": "科学",
    }
    return [name for name, pattern in checks.items() if pattern not in text]


def check_plagiarism_risk(text: str, fragments: List[str],
                          threshold: int = 50) -> List[str]:
    """查重风险检测：检查连续 N 字与团章原文/模板片段重复"""
    warnings = []
    if not text:
        return warnings
    text_clean = re.sub(r"\s+", "", str(text))
    for fragment in fragments:
        if not fragment:
            continue
        frag_clean = re.sub(r"\s+", "", str(fragment))
        if len(frag_clean) < threshold:
            continue
        for i in range(len(frag_clean) - threshold + 1):
            window = frag_clean[i:i + threshold]
            if window in text_clean:
                warnings.append(
                    f"检测到连续 {threshold} 字与参考片段重复：{window[:20]}...")
                break
    return warnings


# 团章原文高风险片段（连续 50 字重复则警告）
# 注意：共青团性质完整表述、团员义务条款按 SKILL.md §6 必引，不入此列表
LEAGUE_CONSTITUTION_FRAGMENTS = [
    "中国共产主义青年团要加强全国各族青年之间的团结，"
    "为把我国建设成为富强民主文明和谐美丽的社会主义现代化强国而奋斗。",
    "对于不执行团的决议、违反团章的团员，团的组织应当本着惩前毖后、"
    "治病救人的精神，进行批评和帮助，情节严重的，给予纪律处分。",
]

# 网络入团申请书模板高风险片段（连续 30 字重复则警告）
NETWORK_TEMPLATE_FRAGMENTS = [
    "我怀着十分激动和诚恳的心情，郑重向团组织提出申请",
    "作为一名当代中学生，我深知自己肩负的责任和使命",
    "我决心以实际行动争取早日入团",
]


# ============================================================
# 4 项政治理论表述文本（4 项必引，单一标准版）
# ============================================================

THEORY_TEXTS: Dict[str, str] = {
    "league_nature": (
        "中国共产主义青年团是中国共产党领导的先进青年的群团组织，"
        "是广大青年在实践中学习中国特色社会主义和共产主义的学校，"
        "是中国共产党的助手和后备军。"
        "共青团自成立以来，始终在党的领导下团结带领广大青年"
        "为争取民族独立、人民解放和实现国家富强、人民幸福而不懈奋斗。"
        "这种'党的助手和后备军'性质，决定了共青团既具有鲜明的政治性，"
        "又具有广泛的群众性，是广大青年学习中国特色社会主义和共产主义的学校。"
    ),
    "league_constitution": (
        "我认识到，《中国共产主义青年团章程》是共青团的根本行为准则。"
        "团章规定了团员必须履行的义务：学习马克思列宁主义、毛泽东思想、"
        "邓小平理论、\u201c三个代表\u201d重要思想、科学发展观、习近平新时代中国特色社会主义思想；"
        "宣传、执行党的基本路线和各项方针政策；积极参加改革开放和社会主义现代化建设；"
        "自觉遵守团的纪律；虚心向群众学习；开展批评与自我批评。"
        "党的二十大以来，共青团以习近平新时代中国特色社会主义思想为指导，"
        "深入贯彻党的二十大精神，团结带领广大青年听党话、跟党走。"
        "团章同时规定团员享有参加团内活动、在团内享有选举权被选举权和表决权等权利。"
        "我愿意自觉履行团员义务、行使团员权利。"
        "团章还规定了团的民主集中制组织制度，团的全国代表大会每 5 年举行一次，"
        "团员个人服从组织，少数服从多数，下级组织服从上级组织。"
    ),
    "league_history": (
        "我了解到，中国共产主义青年团有着光荣的历史。"
        "1922 年 5 月 5 日，中国社会主义青年团第一次全国代表大会在广州召开，"
        "标志着中国社会主义青年团的正式成立。"
        "1925 年 1 月，团的第三次全国代表大会决定将名称改为中国共产主义青年团。"
        "1957 年 5 月，新民主主义青年团第三次全国代表大会决定将名称改为中国共产主义青年团，并沿用至今。"
        "百余年来，共青团始终在党的领导下团结带领广大青年在革命、建设、改革各个历史时期作出了重要贡献。"
        "1922 年 5 月 5 日是马克思诞辰纪念日，团的一大特意选在这一天召开，"
        "象征着共青团自诞生之日起就以马克思主义为指导。"
    ),
    "may_fourth_spirit": (
        "我深刻认识到五四精神的内涵。"
        "1919 年 5 月 4 日爆发的五四运动，是一场伟大的爱国革命运动，"
        "孕育了以'爱国、进步、民主、科学'为主要内容的伟大五四精神。"
        "爱国是五四精神的核心，进步是五四精神的动力，"
        "民主与科学是五四精神的旗帜。"
        "作为新时代青年，我应当继承和发扬五四精神，"
        "把个人理想融入民族复兴的伟大事业之中。"
    ),
}


def get_theory_text(theory_id: str) -> str:
    """按 theory_id 获取理论表述文本"""
    return THEORY_TEXTS.get(theory_id, "")


# ============================================================
# 文档构建器
# ============================================================

class ApplicationDocBuilder:
    """入团申请书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []
        self.full_text_parts: List[str] = []
        self.oath_paragraph_text: str = ""
        # 配置（在 _resolve_config 中赋值）
        self.word_count_version: str = "standard"
        self.word_count_target: int = 1800
        self.school_tier: str = "normal"
        self.story_count: int = 3
        self.shortcoming_count: int = 3
        self.theory_quotes_selected: List[str] = list(THEORY_QUOTE_IDS)
        self.include_oath: bool = False
        self.require_family_appendix: bool = True
        self.applicant_age: Optional[int] = None
        setup_page(self.doc)
        add_page_number(self.doc)

    # --------------------------------------------------------
    # 数据访问辅助
    # --------------------------------------------------------

    def _get(self, key: str, default: Any = None) -> Any:
        return self.data.get(key, default)

    def _get_str(self, key: str, default: str = "") -> str:
        v = self.data.get(key, default)
        return str(v) if v is not None else default

    def _get_list(self, key: str) -> List[Any]:
        v = self.data.get(key, [])
        return list(v) if isinstance(v, list) else []

    def add_para(self, text: str, indent: bool = True):
        """添加正文段落并记录到 full_text_parts"""
        add_body(self.doc, text, indent=indent)
        self.full_text_parts.append(text)

    # --------------------------------------------------------
    # 配置解析
    # --------------------------------------------------------

    def _resolve_config(self):
        """解析 word_count_version 与 school_tier 等配置"""
        wcv = str(self._get("word_count_version", default="standard")).strip().lower()
        if wcv not in WORD_COUNT_VERSIONS:
            if wcv:
                self.warnings.append(
                    f"[字段] word_count_version='{wcv}' 不合法，"
                    f"应为 {'/'.join(WORD_COUNT_VERSIONS.keys())} 之一，已回退到 'standard'")
            wcv = "standard"
        self.word_count_version = wcv
        cfg = WORD_COUNT_VERSIONS[wcv]
        self.word_count_target = cfg["target"]
        self.story_count = cfg["story_count"]
        self.shortcoming_count = cfg["shortcoming_count"]

        tier = str(self._get("school_tier", default="normal")).strip().lower()
        if tier not in SCHOOL_TIERS:
            if tier:
                self.warnings.append(
                    f"[字段] school_tier='{tier}' 不合法，"
                    f"应为 {'/'.join(SCHOOL_TIERS.keys())} 之一，已回退到 'normal'")
            tier = "normal"
        self.school_tier = tier

        quotes = self._get_list("theory_quotes")
        if quotes:
            invalid = [q for q in quotes if q not in THEORY_QUOTE_IDS]
            if invalid:
                self.warnings.append(
                    f"[字段] theory_quotes 含不合法 ID：{invalid}，"
                    f"合法 ID：{THEORY_QUOTE_IDS}")
            valid = [q for q in quotes if q in THEORY_QUOTE_IDS]
            if valid:
                self.theory_quotes_selected = valid

        self.include_oath = bool(self._get("include_oath", default=False)) or \
                            bool(self._get("require_oath", default=False))
        self.require_family_appendix = bool(
            self._get("require_family_appendix", default=True))

        explicit_age = self._get("applicant_age")
        if isinstance(explicit_age, (int, float)) and explicit_age > 0:
            self.applicant_age = int(explicit_age)
        else:
            self.applicant_age = compute_age(
                self._get_str("birth_date"),
                self._get_str("submit_date"))

    # --------------------------------------------------------
    # 标题与称呼
    # --------------------------------------------------------

    def _add_title(self):
        add_title(self.doc, "入团申请书")

    def _add_salutation(self):
        add_salutation(self.doc, self._get_str("salutation", default="敬爱的团组织："))

    # --------------------------------------------------------
    # 第一段：开篇志愿表达（占 8%）
    # --------------------------------------------------------

    def _add_will_statement(self):
        """第一段：开篇志愿表达"""
        will_text = self._get_str("will_statement", default="")
        if will_text:
            self.add_para(will_text)
            return

        self.add_para("我申请加入中国共产主义青年团。")

        if self.word_count_version == "compact":
            body = (
                "我愿意为共产主义事业奋斗，愿意接受团组织的教育与考验。"
                "在成长过程中亲眼见证了党领导全国人民取得的伟大成就，"
                "对中国共产主义青年团产生了由衷向往。"
                "我郑重向团组织递交入团申请书，"
                "希望能在团组织的教育和培养下，成为一名合格的共青团员。"
            )
        elif self.word_count_version == "standard":
            body = (
                "我愿意为共产主义事业奋斗，愿意接受团组织的教育与考验。"
                "我是一名普通的在校大学生，在成长过程中亲眼见证了党领导全国人民取得的伟大成就。"
                "从家乡脱贫攻坚的扎实进展，到抗击新冠疫情的举国同心，"
                "再到全面建设社会主义现代化国家新征程的稳健开局，"
                "每一件大事都让我深切感受到中国共产主义青年团是中国共产党的助手和后备军。"
                "基于对团组织的初步认识与由衷向往，我郑重向团组织递交入团申请书，"
                "希望能在团组织的教育和培养下，成为一名合格的共青团员。"
            )
        else:  # enhanced
            body = (
                "我愿意为共产主义事业奋斗，愿意接受团组织的教育与考验。"
                "我是一名普通的在校大学生，在成长过程中亲眼见证了党领导全国人民取得的伟大成就。"
                "从家乡脱贫攻坚的扎实进展，到抗击新冠疫情的举国同心，"
                "再到全面建设社会主义现代化国家新征程的稳健开局，"
                "每一件大事都让我深切感受到中国共产主义青年团是中国共产党的助手和后备军，"
                "是党联系青年的桥梁和纽带。"
                "基于对团组织的初步认识与由衷向往，我郑重向团组织递交入团申请书，"
                "希望能在团组织的教育和培养下，成为一名合格的共青团员。"
            )
        self.add_para(body)

    # --------------------------------------------------------
    # 第二段：对团的认识（占 38%）
    # --------------------------------------------------------

    def _add_league_understanding(self):
        """第二段：对团的认识"""
        understanding_text = self._get_str("league_understanding", default="")
        if understanding_text:
            self.add_para(understanding_text)
            return
        self.add_para("我对中国共产主义青年团有如下认识。")
        for tid in ["league_nature", "league_constitution",
                    "league_history", "may_fourth_spirit"]:
            if tid in self.theory_quotes_selected:
                self.add_para(get_theory_text(tid))

    # --------------------------------------------------------
    # 第三段：个人情况与成长经历（占 38%）
    # --------------------------------------------------------

    def _add_personal_experience(self):
        """第三段：个人情况与成长经历"""
        experience_text = self._get_str("personal_experience", default="")
        if experience_text:
            self.add_para(experience_text)
            return
        self._add_personal_basic_info()
        self._add_education_history()
        self._add_motivation_and_stories()
        self._add_current_cognition()

    def _add_personal_basic_info(self):
        """基本信息与家庭背景"""
        name = self._get_str("name", default="")
        birth_date = self._get_str("birth_date", default="")
        family_origin = self._get_str("family_origin", default="普通")
        family_members = self._get_list("family_members")

        def _member_desc(prefix: str) -> str:
            m = next((x for x in family_members
                      if isinstance(x, dict)
                      and str(x.get("relation", "")).startswith(prefix)), {})
            parts = []
            if m.get("political_status"):
                parts.append(str(m["political_status"]))
            if m.get("work_unit"):
                parts.append(str(m["work_unit"]))
                if m.get("position"):
                    parts.append(str(m["position"]))
            return "、".join(parts)

        intro_parts = []
        if name:
            intro_parts.append(f"我叫{name}")
        if birth_date:
            intro_parts.append(f"{birth_date}出生")
        if family_origin:
            intro_parts.append(f"于一个{family_origin}家庭")
        intro = "，".join(intro_parts) + "。" if intro_parts else ""

        parents = []
        for prefix in ("父亲", "母亲"):
            d = _member_desc(prefix)
            if d:
                parents.append(f"{prefix}{d}")
        parents_desc = "，".join(parents) + "。" if parents else ""

        longing = "在家庭熏陶下，我从小对团组织有着朴素的向往。"
        text = (intro + parents_desc + longing).strip()
        if text:
            self.add_para(text)

    def _add_education_history(self):
        """教育经历"""
        education_history = self._get_list("education_history")
        team_join_date = self._get_str("team_join_date", default="")
        if not education_history:
            return
        parts = ["我的教育经历如下："]
        for edu in education_history:
            if not isinstance(edu, dict):
                continue
            seg = f"{edu.get('start', '')}至{edu.get('end', '')}在{edu.get('school', '')}就读"
            pos = edu.get("position", "")
            if pos and pos != "学生":
                seg += f"，曾担任{pos}"
            parts.append(seg + "；")
        if team_join_date:
            parts.append(f"{team_join_date}加入中国少年先锋队；")
        if parts and parts[-1].endswith("；"):
            parts[-1] = parts[-1][:-1] + "。"
        text = "".join(parts)
        if text:
            self.add_para(text)

    def _add_motivation_and_stories(self):
        """入团动机与思想发展事例"""
        motivation = self._get_str("motivation", default="")
        if motivation:
            self.add_para(motivation)
            return
        stories = self._get_list("ideology_story")
        if not stories:
            stories = self._build_default_stories()
        max_count = min(self.story_count, len(stories))
        intro = "我之所以申请加入共青团，源于几件触动我心灵的事情。"
        ordinals = ["第一件事", "第二件事", "第三件事", "第四件事"]
        story_texts = []
        for idx, story in enumerate(stories[:max_count]):
            if isinstance(story, dict):
                ordinal = ordinals[idx] if idx < len(ordinals) else f"第{idx+1}件事"
                seg = f"{ordinal}是{story.get('time', '')}{story.get('event', '')}。"
                if story.get("thought"):
                    seg += story["thought"]
                if story.get("change"):
                    seg += story["change"]
                story_texts.append(seg)
            else:
                story_texts.append(str(story))
        if story_texts:
            self.add_para(intro + "".join(story_texts))

    def _build_default_stories(self) -> List[Dict[str, str]]:
        """构建默认的 3 个思想发展事例"""
        return [
            {"time": "2024 年 9 月大一军训期间，",
             "event": "教官是一名退伍军人团员，在烈日下将军训物资一件件搬到学生宿舍，全程没有一句怨言。",
             "thought": "这件事让我开始思考：是什么样的精神能让一个青年愿意为他人默默付出？",
             "change": ""},
            {"time": "2024 年 11 月，",
             "event": "我参加了学院青年志愿者协会组织的'暖冬行动'，为山区孩子捐书 200 余册。在整理图书时，我看到一本书扉页上写着'愿这本书带你看更远的世界'，",
             "thought": "那一刻我深切体会到青年志愿者的责任与担当。",
             "change": ""},
            {"time": "2025 年 3 月，",
             "event": "我旁听了学校'青年马克思主义者培养工程'的公开课，听到了 95 后驻村第一书记的分享，他放弃城市工作机会扎根山村两年，帮助村民脱贫致富。",
             "thought": "这让我深刻认识到共青团员不仅是身份，更是责任与担当。",
             "change": ""},
        ]

    def _add_current_cognition(self):
        """当前认识"""
        current_cognition = self._get_str("current_cognition", default="")
        if not current_cognition:
            current_cognition = (
                "通过这些实践，我更加深刻地认识到："
                "共青团员不仅是一个身份，更是一份责任；不是光环，而是奉献。"
                "当前，我已具备成为一名共青团员的思想基础，"
                "恳请团组织在实践中考验我。"
            )
        self.add_para(current_cognition)

    # --------------------------------------------------------
    # 第四段：入团态度与今后努力方向（占 16%）
    # --------------------------------------------------------

    def _add_shortcomings(self):
        """第四段：入团态度与今后努力方向"""
        shortcomings_text = self._get_str("shortcomings_text", default="")
        if shortcomings_text:
            self.add_para(shortcomings_text)
            return
        shortcomings = self._get_list("shortcomings")
        if not shortcomings:
            shortcomings = self._build_default_shortcomings()

        self.add_para(
            "当然，我也清醒地认识到自己还存在不少不足，主要表现在以下几个方面：")

        max_count = min(self.shortcoming_count, len(shortcomings))
        for idx, item in enumerate(shortcomings[:max_count], start=1):
            if isinstance(item, dict):
                num_word = ["一", "二", "三", "四", "五"][idx - 1] if idx <= 5 else str(idx)
                para_text = f"{num_word}是{item.get('description', '')}。{item.get('manifestation', '')}{item.get('improvement', '')}"
            else:
                para_text = str(item)
            self.add_para(para_text)

        ending_statement = self._get_str("shortcomings_ending", default="")
        if not ending_statement:
            ending_statement = self._build_shortcomings_ending()
        self.add_para(ending_statement)

    def _build_default_shortcomings(self) -> List[Dict[str, str]]:
        """构建默认的 3 项真实不足"""
        return [
            {"description": "理论学习系统性不够",
             "manifestation": "我对团史团章的学习还停留在知识点层面，没有形成体系；对团的十九大精神的学习还不够系统。",
             "improvement": "今后我准备系统研读《中国共产主义青年团章程》《新时代团员教育读本》等书籍，并坚持每周写一篇学习笔记。"},
            {"description": "实践能力仍需加强",
             "manifestation": "我参加的志愿服务和社会实践还不够多，对国情民情的了解还不够深入。",
             "improvement": "今后我准备每学期至少参加 2 次志愿服务，每次不少于 8 小时，并积极参加暑期'三下乡'社会实践。"},
            {"description": "与同学沟通方式有待改进",
             "manifestation": "在班级工作中，我有时过于直接，没有充分考虑同学的接受能力。",
             "improvement": "今后我准备多向辅导员和高年级团员请教，学习群众工作的方法，提高沟通能力。"},
        ]

    def _build_shortcomings_ending(self) -> str:
        """构建个人不足段结尾（今后努力方向）"""
        if self.word_count_version == "compact":
            return (
                "今后，我将以团员标准严格要求自己，"
                "在学习中坚定理想信念，在实践中提升能力本领，"
                "自觉接受团组织的教育和考验，"
                "争取早日成为一名合格的共青团员。"
                "若团组织暂时未批准我的申请，我也绝不气馁，"
                "将继续努力，以更高标准要求自己，接受团组织的长期考验。"
            )
        return (
            "今后，我将以团员标准严格要求自己："
            "在学习方面，坚持每学期精读 2 本政治理论书籍并撰写学习笔记；"
            "在思想方面，主动参加'青年大学习'网上主题团课，"
            "每月向团组织汇报一次思想情况；"
            "在实践方面，每学期至少参加 2 次志愿服务，每次不少于 8 小时，"
            "并积极参加暑期'三下乡'社会实践。"
            "若团组织暂时未批准我的申请，我也绝不气馁，"
            "将继续努力，以更高标准要求自己，接受团组织的长期考验。"
            "请团组织在实践中考验我！"
        )

    # --------------------------------------------------------
    # 入团誓词段（仅当 include_oath=true 时插入）
    # --------------------------------------------------------

    def _add_oath_if_required(self):
        """若学校规范要求引用入团誓词，则在第四段末尾插入"""
        if not self.include_oath:
            return
        oath_text = "入团誓词：" + LEAGUE_OATH_TEXT
        self.add_para(oath_text)
        self.oath_paragraph_text = oath_text

    # --------------------------------------------------------
    # 结尾"此致 敬礼！"与落款
    # --------------------------------------------------------

    def _add_ending(self):
        add_cizhi(self.doc, "此致")
        add_jingli(self.doc, "敬礼！")

    def _add_signature(self):
        """落款：申请人 + 日期，右对齐"""
        add_blank(self.doc)
        name = self._get_str("name", default="申请人")
        info_parts = []
        for k, fmt in [("college", None), ("major", None), ("grade", None),
                       ("class_name", None)]:
            v = self._get_str(k, default="")
            if v:
                info_parts.append(v)
        student_id = self._get_str("student_id", default="")
        if student_id:
            info_parts.append(f"学号 {student_id}")
        line1 = f"申请人：{name}"
        if info_parts:
            line1 += "（" + " ".join(info_parts) + "）"
        add_right(self.doc, line1)
        submit_date = self._get_str("submit_date", default="")
        if submit_date:
            add_right(self.doc, normalize_date(submit_date))

    # --------------------------------------------------------
    # 附页：家庭主要成员和主要社会关系
    # --------------------------------------------------------

    def _add_family_appendix(self):
        """生成家庭主要成员和主要社会关系附页（独立成页）"""
        if not self.require_family_appendix:
            return
        family_members = self._get_list("family_members")
        social_relations = self._get_list("social_relations")
        if not family_members and not social_relations:
            self.warnings.append(
                "[附页] require_family_appendix=true 但 family_members 为空，跳过家庭主要成员附页")
            return

        add_page_break(self.doc)
        add_subtitle(self.doc, "家庭主要成员和主要社会关系")
        add_blank(self.doc)
        add_subtitle(self.doc, "一、家庭主要成员")
        headers = ["称谓", "姓名", "出生年月", "政治面貌", "工作单位及职务"]

        def _row(m: Dict[str, str]) -> List[str]:
            return [str(m.get("relation", "")), str(m.get("name", "")),
                    str(m.get("birth_date", "")), str(m.get("political_status", "")),
                    f"{m.get('work_unit', '')} {m.get('position', '')}".strip()]

        rows = [_row(m) for m in family_members if isinstance(m, dict)]
        # 添加本人一行
        self_info = " ".join([s for s in [self._get_str("college"),
                                          self._get_str("major"),
                                          self._get_str("grade")] if s])
        rows.append(["本人", self._get_str("name"), self._get_str("birth_date"),
                     self._get_str("political_status"), self_info])
        add_table(self.doc, headers, rows)

        if social_relations:
            add_blank(self.doc)
            add_subtitle(self.doc, "二、主要社会关系")
            add_table(self.doc, headers,
                      [_row(m) for m in social_relations if isinstance(m, dict)])

        add_blank(self.doc)
        add_body(self.doc,
                 "注：以上家庭主要成员和主要社会关系信息均如实填写，"
                 "若有变动将及时向团组织报告。")

    # --------------------------------------------------------
    # 主构建方法
    # --------------------------------------------------------

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/4 段正文/誓词/结尾/落款/附页，生成 docx"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._resolve_config()
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_will_statement()
            self._add_league_understanding()
            self._add_personal_experience()
            self._add_shortcomings()
            self._add_oath_if_required()
            self._add_ending()
            self._add_signature()
            self._add_family_appendix()
            self._post_build_checks()
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            self.doc.save(str(output_path))
            print(f"✅ 入团申请书已生成：{output_path}")
            return str(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _post_build_checks(self):
        """构建后检查：政治规范 + 查重 + 字数 + 4 项理论 + 「志愿书」三字"""
        full_text = "".join(self.full_text_parts)

        for w in check_political_terms(full_text):
            self.warnings.append(f"[政治规范] {w}")
        for w in check_plagiarism_risk(full_text, LEAGUE_CONSTITUTION_FRAGMENTS, 50):
            self.warnings.append(f"[查重风险-团章原文] {w}")
        for w in check_plagiarism_risk(full_text, NETWORK_TEMPLATE_FRAGMENTS, 30):
            self.warnings.append(f"[查重风险-网络模板] {w}")

        char_count = count_chinese_chars(full_text)
        target = self.word_count_target
        tier_min = SCHOOL_TIERS[self.school_tier]["min"]
        version_min = WORD_COUNT_VERSIONS[self.word_count_version]["min"]
        effective_min = max(tier_min, version_min)
        if char_count < effective_min:
            self.warnings.append(
                f"[字数] 全文 {char_count} 字，低于 {self.school_tier} 档下限 "
                f"{effective_min} 字（{WORD_COUNT_VERSIONS[self.word_count_version]['label']}目标 {target} 字）")
        elif char_count < int(target * 0.70):
            self.warnings.append(
                f"[字数] 全文 {char_count} 字，低于目标 {target} 字的 70%，建议补充内容")

        for m in check_4_theory_expressions(full_text):
            self.warnings.append(f"[4 项理论] 缺失：{m}")

        # 「志愿书」三字检查（与《入团志愿书》混淆）
        if self.include_oath and self.oath_paragraph_text:
            text_to_check = full_text.replace(self.oath_paragraph_text, "")
        else:
            text_to_check = full_text
        if "志愿书" in text_to_check:
            self.warnings.append(
                "[「志愿书」三字] 出现禁用'志愿书'三字"
                "（与《入团志愿书》混淆；单独'志愿加入''志愿为'合规）")

        if self.warnings:
            print("⚠️ 构建后检查警告：", file=sys.stderr)
            for w in self.warnings:
                print(f"  - {w}", file=sys.stderr)

    # --------------------------------------------------------
    # 数据校验（SKILL.md §16.5）
    # --------------------------------------------------------

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []

        # P0 必采字段
        p0_fields = [
            ("name", "申请人姓名"), ("student_id", "学号"),
            ("gender", "性别"), ("birth_date", "出生年月"),
            ("college", "学院全称"), ("major", "专业全称"),
            ("grade", "年级"), ("league_branch", "团支部全称"),
            ("submit_date", "递交日期"), ("political_status", "政治面貌"),
        ]
        for key, label in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {label}（{key}）")

        # 1. 年龄校验（SKILL.md §0.2）
        if self.applicant_age is not None:
            if self.applicant_age < 14:
                warnings.append(
                    f"[年龄红线] 申请人 {self.applicant_age} 周岁，未满 14 周岁，"
                    "根据团章第一条规定不能申请加入中国共产主义青年团。")
            elif self.applicant_age >= 28:
                warnings.append(
                    f"[年龄红线] 申请人 {self.applicant_age} 周岁，已超过 28 周岁，"
                    "根据团章第一条规定不能再申请入团。")
        elif self._get("birth_date") and self._get("submit_date"):
            warnings.append("[年龄校验] birth_date 或 submit_date 格式无法识别")

        # 2. 政治面貌校验
        ps = self._get_str("political_status").strip()
        valid = ["群众", "少先队员", "共青团员", "中共党员"]
        if ps and ps not in valid:
            warnings.append(f"[政治面貌] political_status='{ps}'，应为 {'/'.join(valid)} 之一")
        if ps == "共青团员":
            warnings.append("[政治面貌] 申请人已是共青团员，无需再递交入团申请书")
        if ps == "中共党员":
            warnings.append("[政治面貌] 申请人已是中共党员，自动保留团籍至 28 周岁，无需再申请入团")

        # 3. 日期格式校验
        sd = self._get_str("submit_date")
        if sd and not re.match(r"^\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日$", sd.strip()):
            normalized = normalize_date(sd)
            self.data["submit_date"] = normalized
            warnings.append(f"[日期格式] submit_date='{sd}' 已自动规范化为 '{normalized}'")

        # 4. 家庭主要成员校验
        fm = self._get_list("family_members")
        if not fm:
            warnings.append("缺少 家庭主要成员（family_members），将使用默认家庭背景")
        else:
            has_father = any(isinstance(m, dict)
                             and str(m.get("relation", "")).startswith("父亲") for m in fm)
            has_mother = any(isinstance(m, dict)
                             and str(m.get("relation", "")).startswith("母亲") for m in fm)
            if not has_father:
                warnings.append("家庭主要成员缺少 父亲 信息")
            if not has_mother:
                warnings.append("家庭主要成员缺少 母亲 信息")
            for m in fm:
                if isinstance(m, dict) and not m.get("political_status"):
                    warnings.append(f"家庭主要成员 {m.get('relation', '?')} 缺少政治面貌")

        # 5. 入团动机校验
        motivation = self._get_str("motivation")
        if motivation:
            mc = count_chinese_chars(motivation)
            if mc < 200:
                warnings.append(f"[入团动机] motivation 字数 {mc} 字偏少，建议不少于 200 字")
            if any(kw in motivation for kw in
                   ["求职", "就业", "公务员", "加分", "评优", "评先"]):
                warnings.append(
                    "[政治红线] 入团动机出现'求职/就业/公务员/加分/评优/评先'字眼，"
                    "动机不端正风险，请用户重新表述")

        # 6. 个人不足校验
        shortcomings = self._get_list("shortcomings")
        if shortcomings and len(shortcomings) < 2:
            warnings.append(f"个人不足仅 {len(shortcomings)} 个，建议 2~3 个真实不足")

        # 7. 个人不足假缺点检测
        fake_kws = ["工作太投入", "追求完美", "学习太刻苦",
                    "责任心太强", "为人太直率", "事必躬亲",
                    "过于追求完美", "过于认真"]
        s_text = self._get_str("shortcomings_text")
        for item in shortcomings:
            if isinstance(item, dict):
                s_text += str(item.get("description", "")) + str(item.get("manifestation", ""))
        for kw in fake_kws:
            if kw in s_text:
                warnings.append(
                    f"[政治红线] 个人不足出现假缺点'{kw}'，"
                    "团支部会要求重写，请用户重新表述")
                break

        self.warnings.extend(warnings)
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据（standard 标准版 1800 字）
# ============================================================

DEFAULT_DATA = {
    "name": "张明",
    "student_id": "2024123456",
    "gender": "男",
    "birth_date": "2009 年 5 月",
    "college": "计算机科学与技术学院",
    "major": "计算机科学与技术",
    "grade": "2024 级 大一",
    "class_name": "计科 2401 班",
    "family_origin": "工人",
    "identity": "学生",
    "political_status": "少先队员",
    "salutation": "敬爱的团组织：",
    "submit_date": "2025 年 5 月 4 日",
    "league_branch": "计算机科学与技术学院 2024 级 1 班团支部",
    # 配置字段
    "word_count_version": "standard",
    "school_tier": "normal",
    "include_oath": False,
    "require_family_appendix": True,
    "theory_quotes": list(THEORY_QUOTE_IDS),
    "family_members": [
        {"relation": "父亲", "name": "张建国", "birth_date": "1977 年 3 月",
         "political_status": "中共党员", "work_unit": "XX 县机械厂",
         "position": "车间主任"},
        {"relation": "母亲", "name": "李秀英", "birth_date": "1979 年 8 月",
         "political_status": "共青团员", "work_unit": "XX 县人民医院",
         "position": "护士长"},
    ],
    "social_relations": [
        {"relation": "祖父", "name": "张文才", "birth_date": "1950 年 2 月",
         "political_status": "中共党员", "work_unit": "XX 县农业局（已退休）",
         "position": "原副局长"},
    ],
    "team_join_date": "2018 年 5 月",
    "education_history": [
        {"stage": "小学", "start": "2015 年 9 月", "end": "2021 年 6 月",
         "school": "XX 县第一小学", "position": "学习委员"},
        {"stage": "初中", "start": "2021 年 9 月", "end": "2024 年 6 月",
         "school": "XX 县第一初级中学", "position": "团支部书记助理"},
        {"stage": "大学", "start": "2024 年 9 月", "end": "至今",
         "school": "XX 大学", "position": "学习委员"},
    ],
    "shortcomings": [
        {"description": "理论学习系统性不够",
         "manifestation": "我对团史团章的学习还停留在知识点层面，没有形成体系；对团的十九大精神的学习还不够系统。",
         "improvement": "今后我准备系统研读《中国共产主义青年团章程》《新时代团员教育读本》等书籍，并坚持每周写一篇学习笔记。"},
        {"description": "实践能力仍需加强",
         "manifestation": "我参加的志愿服务和社会实践还不够多，对国情民情的了解还不够深入。",
         "improvement": "今后我准备每学期至少参加 2 次志愿服务，每次不少于 8 小时，并积极参加暑期'三下乡'社会实践。"},
        {"description": "与同学沟通方式有待改进",
         "manifestation": "在班级工作中，我有时过于直接，没有充分考虑同学的接受能力。",
         "improvement": "今后我准备多向辅导员和高年级团员请教，学习群众工作的方法，提高沟通能力。"},
    ],
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="入团申请书 docx 生成器（v1.0，依据 SKILL.md v1.0 规范）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n关键字段：\n"
            "  word_count_version: compact/standard/enhanced（3 档字数路由）\n"
            "  school_tier:        loose/normal/strict（字数下限校验）\n"
            "  include_oath:       true 时附录生成入团誓词（默认 false）\n"
            "  require_family_appendix: true 时附录生成家庭成员附页（默认 true）\n"
            "  applicant_age:      显式年龄（否则从 birth_date + submit_date 计算）\n"
            "  theory_quotes:      必引理论 ID 列表（默认 4 项全引）\n"
            "\n校验：必提 5 项要点 + 4 项理论表述 + 指导思想 6 项顺序 + "
            "「志愿书」三字校验 + 查重风险检测 + 14-28 周岁年龄校验。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档（standard 标准版，目标 1800 字）")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档（standard 标准版，目标 1800 字）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
