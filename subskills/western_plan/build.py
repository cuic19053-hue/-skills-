#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
大学生志愿服务西部计划申请书 docx 生成器

按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 一级标题：黑体三号，居中
- 二级标题：黑体小三，左对齐
- 三级标题：宋体四号加粗
- 表格：宋体五号，居中

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第十一章。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)
SIZE_XIAO_ER = Pt(18)
SIZE_SAN = Pt(16)
SIZE_XIAO_SAN = Pt(15)
SIZE_SI = Pt(14)
SIZE_XIAO_SI = Pt(12)
SIZE_WU = Pt(10.5)
SIZE_XIAO_WU = Pt(9)

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# 8 大服务专项代码 → 中文全称映射
SERVICE_PREFERENCE_MAP = {
    "education_service": "教育服务专项",
    "health_service": "卫生服务专项",
    "agritech_service": "农技推广专项",
    "rural_revitalization": "乡村振兴专项",
    "youth_work": "青年工作专项",
    "grassroots_governance": "基层社会治理专项",
    "rural_construction": "乡村建设专项",
    "youth_league_work": "共青团工作专项",
}

# 8 大专项简要描述（用于专项对照表）
SERVICE_PREFERENCE_DESC = {
    "education_service": "中小学教师 / 教务管理",
    "health_service": "乡镇卫生院医护 / 公共卫生",
    "agritech_service": "农技推广员 / 农业技术员",
    "rural_revitalization": "乡村振兴工作队员",
    "youth_work": "团县委干部助理 / 青年事务社工",
    "grassroots_governance": "社区工作者 / 网格员助理",
    "rural_construction": "乡村规划员 / 村镇建设助理",
    "youth_league_work": "团县委干部助理 / 少先队辅导员",
}


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
        space_before=12, space_after=12,
    )


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6, space_after=6,
    )


def add_heading_level3(doc, text: str):
    """三级标题：宋体四号加粗，左对齐"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6, space_after=3,
    )


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent,
        line_spacing=1.5,
    )


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j >= 1 and len(headers) >= 3 \
                else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False, alignment=align)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def safe_int(value, default: int = 0) -> int:
    """安全转整数"""
    try:
        return int(str(value).strip())
    except (ValueError, TypeError):
        return default


def count_chinese_chars(text: str) -> int:
    """统计中文字符数（粗略版，用于字数校验）"""
    if not text:
        return 0
    count = 0
    for ch in str(text):
        if '\u4e00' <= ch <= '\u9fff':
            count += 1
    return count


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """大学生志愿服务西部计划申请书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def add_h1(self, text):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text):
        return add_heading_level2(self.doc, text)

    def add_h3(self, text):
        return add_heading_level3(self.doc, text)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_table(self, headers, rows, col_widths=None):
        return add_table_from_data(self.doc, headers, rows, col_widths)

    def add_page_break(self):
        add_page_break(self.doc)

    # 封面

    def _add_cover(self):
        """封面：黑体二号标题 + 6 行下划线信息"""
        for _ in range(3):
            self.doc.add_paragraph()

        title = "大学生志愿服务西部计划申请书"
        add_paragraph_with_format(
            self.doc, title,
            font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_before=12, space_after=12,
        )

        subtitle = f"（{self._get('apply_year', default='2025')} 年度）"
        add_paragraph_with_format(
            self.doc, subtitle,
            font_name=FONT_HEI, font_size=SIZE_SAN, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_after=24,
        )

        for _ in range(3):
            self.doc.add_paragraph()

        sp_code = self._get("service_preference", default="education_service")
        sp_name = SERVICE_PREFERENCE_MAP.get(sp_code, sp_code)
        info_items = [
            ("申请人姓名", self._get("applicant_name", default="")),
            ("所在学校", self._get("applicant_school", default="")),
            ("专业年级", f"{self._get('applicant_major', default='')} "
                       f"{self._get('applicant_grade', default='')}"),
            ("申请服务专项", sp_name),
            ("意向服务地", self._get("service_region_preference", default="")),
            ("申请日期", f"{self._get('apply_year', default='2025')} 年      月      日"),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI,
                         font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True

        self.add_page_break()

    # 一、个人基本情况

    def _add_basic_info(self):
        """一、个人基本情况（200-300 字段落 + 11 行基本信息表）"""
        self.add_h1("一、个人基本情况")

        # 段落（200-300 字）
        name = self._get("applicant_name", default="XXX")
        gender = self._get("applicant_gender", default="X")
        ethnicity = self._get("applicant_ethnicity", default="X族")
        birth = self._get("applicant_birth", default="XXXX 年 X 月")
        political = self._get("applicant_political", default="共青团员")
        native = self._get("applicant_native_place", default="XX 省 XX 市")
        school = self._get("applicant_school", default="XX 大学")
        major = self._get("applicant_major", default="XX 专业")
        grade = self._get("applicant_grade", default="XXXX 级")
        degree = self._get("applicant_degree", default="本科应届毕业")
        graduation = self._get("applicant_graduation", default="XXXX 年 X 月")
        rank = self._get("academic_rank", default="专业前 XX%")
        honors = self._get("honors", default=[])
        honors_str = "、".join(honors) if isinstance(honors, list) and honors else "校级优秀学生"
        sp_code = self._get("service_preference", default="education_service")
        sp_name = SERVICE_PREFERENCE_MAP.get(sp_code, sp_code)
        years = self._get("service_years", default=1)

        intro = (
            f"本人 {name}，性别 {gender}，民族 {ethnicity}，{birth} 出生于 "
            f"{native}，政治面貌 {political}。现为 {school} {major} {grade} "
            f"{degree}，预计 {graduation} 毕业。在校期间学业排名 {rank}，"
            f"曾获 {honors_str} 等荣誉。本人志愿报名参加 "
            f"{self._get('apply_year', default='2025')} 年大学生志愿服务西部计划，"
            f"申请 {sp_name}，赴西部基层服务 {years} 年。"
        )
        self.add_para(intro)

        # 基本信息表（2 列展示，便于阅读）
        certs = self._get("certificates", default=[])
        certs_str = "、".join(certs) if isinstance(certs, list) and certs else "—"
        rows_simple = [
            ["姓名", name],
            ["性别", gender],
            ["民族", ethnicity],
            ["政治面貌", political],
            ["出生年月", birth],
            ["籍贯", native],
            ["学校", school],
            ["学历层次", degree],
            ["专业年级", f"{major} {grade}"],
            ["学号", self._get("applicant_id", default="")],
            ["毕业时间", graduation],
            ["联系电话", self._get("applicant_phone", default="")],
            ["电子邮箱", self._get("applicant_email", default="—")],
            ["学业排名", rank],
            ["GPA", self._get("gpa", default="—")],
            ["资格证书", certs_str],
            ["申请服务专项", sp_name],
            ["服务年限", f"{years} 年"],
            ["意向服务地", self._get("service_region_preference", default="")],
            ["是否服从调剂",
             "是" if self._get("accept_redeployment", default=True) else "否"],
        ]
        self.add_table(["项目", "内容"], rows_simple, col_widths=[4.5, 11.5])

    # 二、报考西部计划动机

    def _add_motivation(self):
        """二、报考西部计划动机（500-700 字，3 段）

        motivation 列表顺序约定：
        [0] 政策认知段（含团中央文件引用与项目区分）
        [1] 个人志向段（含具体触发事件）
        [2] 服务地情感段（含与服务地连接）
        """
        self.add_h1("二、报考西部计划动机")

        # 引用团中央文件
        docs = self._get("tuancentral_doc_cited", default=[])
        if docs and isinstance(docs, list) and not isinstance(docs, str):
            docs_list = list(docs)
        else:
            docs_list = []

        # 3 段报考动机
        motivation = self._get("motivation", default=[])
        if isinstance(motivation, str):
            motivation = [motivation]
        if motivation and isinstance(motivation, list) and len(motivation) >= 3:
            # （一）政策认知段：先输出 motivation[0] 正文，再列出引用文件
            self.add_h2("（一）政策认知")
            self.add_para(motivation[0])
            if docs_list:
                self.add_para(
                    "本人认真学习了以下政策文件，深入了解西部计划的项目性质、"
                    "服务期限、待遇保障与政策激励："
                )
                for d in docs_list:
                    self.add_para(f"• {d}", indent=False)
            # （二）个人志向段
            self.add_h2("（二）个人志向")
            self.add_para(motivation[1])
            # （三）服务地情感段
            self.add_h2("（三）服务地情感")
            self.add_para(motivation[2])
        elif motivation and isinstance(motivation, list) and len(motivation) >= 1:
            # 不足 3 段，按顺序输出
            for i, para in enumerate(motivation, 1):
                self.add_h2(
                    f"（{['一', '二', '三', '四', '五'][i-1] if i <= 5 else i}）"
                )
                self.add_para(para)
            if docs_list:
                self.add_para(
                    "本人认真学习了以下政策文件：", indent=False
                )
                for d in docs_list:
                    self.add_para(f"• {d}", indent=False)
        else:
            self.add_h2("（一）政策认知")
            self.add_para(
                "（请填写政策认知段，200 字左右。结构：必引 1-2 项团中央文件 + "
                "区分西部计划与三下乡/三支一扶/特岗教师 + 明确本人报考的是西部计划。"
                "示例：本人认真学习了《2025 年大学生志愿服务西部计划实施方案》"
                "及《关于做好 2025 年大学生志愿服务西部计划工作的通知》，深入了解"
                "西部计划自 2003 年实施以来累计选派 50 余万名志愿者的项目历程，"
                "明确西部计划是 1-3 年全职志愿服务，与三下乡（7-15 天暑期实践）、"
                "三支一扶（2-3 年人社部项目）、特岗教师（3 年教育部项目）性质不同。"
                "本人报考的是西部计划，志愿赴西部基层服务 1-3 年。）"
            )
            self.add_h2("（二）个人志向")
            self.add_para(
                "（请填写个人志向段，200 字左右。结构：必含 1 个具体触发事件 + "
                "价值观形成 + 与服务专项的连接。示例：触发本人报考西部计划的，"
                "是 2023 年暑期到 XX 省 XX 县支教时结识的一位彝族女孩阿依。阿依家在"
                "大凉山腹地，每天走 2 小时山路到镇中心校上学，她说她的梦想是'考上"
                "大学回村当老师'。这一刻让本人深刻理解了'用一年时间做一件终生难忘"
                "的事'这句话的分量，也坚定了本人报考西部计划教育服务专项的志向。）"
            )
            self.add_h2("（三）服务地情感")
            self.add_para(
                "（请填写服务地情感段，200 字左右。结构：必含与服务地的连接（籍贯/"
                "实践/家庭）+ 对服务地的认知（含数据）+ 服务期贡献设想。示例：本人"
                "生于四川、长于四川，对凉山有天然的情感认同。在校期间两次到凉山州"
                "参加社会实践，亲眼看到凉山州师生比 1:25（低于全国平均 1:16）、"
                "普通话普及率不足 60% 的现状。本人愿用 1 年时间扎根西昌市某乡镇"
                "小学，从一名普通语文教师做起，为凉山教育事业发展贡献力量。）"
            )

    # 三、个人能力素质

    def _add_abilities(self):
        """三、个人能力素质（500-700 字，4 子段）"""
        self.add_h1("三、个人能力素质")

        abilities = self._get("abilities", default=[])
        if isinstance(abilities, str):
            abilities = [abilities]
        if abilities and isinstance(abilities, list):
            for i, para in enumerate(abilities, 1):
                self.add_h2(f"（{['一', '二', '三', '四', '五'][i-1] if i <= 5 else i}）"
                            f"{['政治素质', '专业能力', '实践经历', '身心素质'][i-1] if i <= 4 else ''}")
                self.add_para(para)
        else:
            self.add_h2("（一）政治素质")
            self.add_para(
                "（请填写政治素质子段，120-180 字。党员必写：入党时间 + 入党介绍人 + "
                "党内职务 + 理论学习情况；团员必写：入团时间 + 团内职务 + 青年大学习"
                "完成期数 + 推优入党情况；群众必写：政治学习情况 + 加入党组织意愿。）"
            )
            self.add_h2("（二）专业能力")
            self.add_para(
                "（请填写专业能力子段，120-180 字。必含：GPA + 排名 + 核心课程 + "
                "资格证书（与服务专项匹配）+ 学术成果（论文/竞赛）。示例：本人在校"
                "期间学业成绩 GPA 3.6/4.0，排名专业前 15%。已取得高级中学教师资格证"
                "（语文）、普通话二级甲等证书，具备从事中小学语文教学的能力。）"
            )
            self.add_h2("（三）实践经历")
            self.add_para(
                "（请填写实践经历子段，120-180 字。必含：学生工作 + 志愿服务（含注册"
                "志愿者编号与累计时长）+ 实习经历（与服务专项匹配）。示例：本人在校"
                "期间担任院学生会主席 1 年，组织活动 12 次。注册志愿者编号 "
                "SC5101234567890，累计志愿服务时长 280 小时。曾于 2024 年 7-8 月在"
                "XX 师大附中实习 2 个月，承担初二语文教学 30 课时。）"
            )
            self.add_h2("（四）身心素质")
            self.add_para(
                "（请填写身心素质子段，120-180 字。必含：体检结论 + 体测达标 + 心理"
                "测评 + 适应能力（含跨文化适应事例）。示例：本人 2025 年 3 月校医院"
                "体检合格，国家学生体质健康标准良好，心理测评结论良好。曾独立骑行"
                "318 国道（成都—拉萨，21 天），具备较强的抗压能力与跨文化适应能力。）"
            )

    # 四、服务岗位意愿

    def _add_service_preference(self):
        """四、服务岗位意愿（400-600 字，3 子段 + 8 专项表）"""
        self.add_h1("四、服务岗位意愿")

        # 8 专项对照表（参考）
        self.add_h2("（〇）8 大服务专项对照（参考）")
        rows = []
        for code, name in SERVICE_PREFERENCE_MAP.items():
            desc = SERVICE_PREFERENCE_DESC.get(code, "")
            rows.append([name, desc])
        self.add_table(["专项名称", "服务岗位"], rows, col_widths=[5.0, 11.0])

        # 3 子段服务岗位意愿
        sp_detail = self._get("service_preference_detail", default=[])
        if isinstance(sp_detail, str):
            sp_detail = [sp_detail]
        if sp_detail and isinstance(sp_detail, list):
            for i, para in enumerate(sp_detail, 1):
                self.add_h2(f"（{['一', '二', '三'][i-1] if i <= 3 else i}）"
                            f"{['专项选择', '服务地选择', '服务年限'][i-1] if i <= 3 else ''}")
                self.add_para(para)
        else:
            self.add_h2("（一）专项选择")
            self.add_para(
                "（请填写专项选择子段，120-200 字。必含：明确所报专项全称（如'教育"
                "服务专项'）+ 选择理由（专业匹配 + 能力匹配 + 服务地需求匹配）+ "
                "第二志愿专项。）"
            )
            self.add_h2("（二）服务地选择")
            self.add_para(
                "（请填写服务地选择子段，120-200 字。必含：明确意向服务省（区、市）"
                "+ 选择理由（籍贯/情感/实践）+ 是否服从调剂。）"
            )
            self.add_h2("（三）服务年限")
            self.add_para(
                "（请填写服务年限子段，120-200 字。必含：明确服务年限（1/2/3 年）"
                "+ 选择理由（1 年→考研加分；2 年→完整项目；3 年→考公定向）+ "
                "服务期满规划。）"
            )

    # 五、家人态度与决心

    def _add_family_attitude(self):
        """五、家人态度与决心（300-500 字，2 子段）"""
        self.add_h1("五、家人态度与决心")

        family = self._get("family_attitude", default=[])
        if isinstance(family, str):
            family = [family]
        if family and isinstance(family, list):
            for i, para in enumerate(family, 1):
                self.add_h2(f"（{['一', '二'][i-1] if i <= 2 else i}）"
                            f"{['家人态度与沟通过程', '决心表态'][i-1] if i <= 2 else ''}")
                self.add_para(para)
        else:
            self.add_h2("（一）家人态度与沟通过程")
            self.add_para(
                "（请填写家人态度子段，150-250 字。必含：家人总体态度 + 沟通过程"
                "（何时告知、如何告知）+ 家人顾虑及化解（家庭责任/安全担忧/未来"
                "出路）+ 最终态度 + 是否签署家长知情同意书。）"
            )
            self.add_h2("（二）决心表态")
            self.add_para(
                "（请填写决心表态子段，150-250 字。必含：再次表态决心 + 服务期承诺"
                "（服从管理、完成本职、不擅自离岗）+ 长期愿景。示例：本人郑重表态："
                "志愿赴西部基层服务 1 年，服务期间服从组织安排，遵守志愿者管理办法，"
                "完成本职工作，不擅自离岗。本人将以'用一年时间做一件终生难忘的事'"
                "为座右铭，扎根西部、服务基层。）"
            )

    # 签字栏

    def _add_signature_section(self):
        """签字栏：申请人 + 学校项目办 + 省项目办"""
        self.add_h1("六、申请人签字与审核意见")

        # 申请人签字
        self.add_h2("（一）申请人签字")
        for _ in range(3):
            self.doc.add_paragraph()
        self.add_para(
            "申请人签字：____________________    "
            "日期：______年____月____日",
            indent=False,
        )

        # 学校项目办审核意见
        self.add_h2("（二）学校项目办审核意见")
        for _ in range(5):
            self.doc.add_paragraph()
        self.add_para(
            "学校项目办盖章：____________________    "
            "负责人签字：____________    "
            "日期：______年____月____日",
            indent=False,
        )

        # 省项目办审核意见
        self.add_h2("（三）省级项目办审核意见")
        for _ in range(5):
            self.doc.add_paragraph()
        self.add_para(
            "省级项目办盖章：____________________    "
            "负责人签字：____________    "
            "日期：______年____月____日",
            indent=False,
        )

    # 主构建方法

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 5 段结构 + 签字栏，生成 docx

        Args:
            data: 申请书字段字典
            output_path: 输出 docx 路径

        Returns:
            实际保存路径
        """
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()

            self._add_cover()
            self._add_basic_info()
            self._add_motivation()
            self._add_abilities()
            self._add_service_preference()
            self._add_family_attitude()
            if self._get("include_signature", default=True):
                self._add_signature_section()

            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申请书已生成：{output_path}")
        return str(output_path)

    # 数据校验

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）

        校验分五类：P0 必填、专业匹配、年限匹配、家人态度、字数。
        """
        warnings = []
        p0_fields = [
            ("applicant_name", "申请人姓名"),
            ("applicant_political", "政治面貌"),
            ("applicant_degree", "学历层次"),
            ("applicant_school", "学校"),
            ("applicant_major", "专业"),
            ("applicant_grade", "年级"),
            ("apply_year", "申请年份"),
            ("service_region_preference", "意向服务地"),
        ]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        # 服务专项校验
        sp = self._get("service_preference", default="")
        if not sp:
            warnings.append("缺少 服务专项（service_preference），必填 8 专项之一")
        elif sp not in SERVICE_PREFERENCE_MAP:
            warnings.append(f"服务专项取值 {sp} 不在 8 专项之列，"
                            f"应为 {list(SERVICE_PREFERENCE_MAP.keys())}")

        # 服务年限校验
        years = safe_int(self._get("service_years", default=0), default=0)
        if years not in (1, 2, 3):
            warnings.append(f"服务年限 {years} 不合法，应为 1/2/3 之一")

        # 段落数量校验
        motivation = self._get("motivation", default=[])
        if isinstance(motivation, list) and len(motivation) < 3:
            warnings.append(f"报考动机段 {len(motivation)} 段，应至少 3 段"
                            "（政策认知+个人志向+服务地情感）")

        abilities = self._get("abilities", default=[])
        if isinstance(abilities, list) and len(abilities) < 4:
            warnings.append(f"能力素质段 {len(abilities)} 子段，应至少 4 子段"
                            "（政治+专业+实践+身心）")

        sp_detail = self._get("service_preference_detail", default=[])
        if isinstance(sp_detail, list) and len(sp_detail) < 3:
            warnings.append(f"服务岗位意愿段 {len(sp_detail)} 子段，应至少 3 子段"
                            "（专项+服务地+年限）")

        family = self._get("family_attitude", default=[])
        if isinstance(family, list) and len(family) < 2:
            warnings.append(f"家人态度段 {len(family)} 子段，应至少 2 子段"
                            "（态度+决心）")

        # 引用文件校验
        docs = self._get("tuancentral_doc_cited", default=[])
        if not docs or (isinstance(docs, list) and len(docs) == 0):
            warnings.append("缺少 引用团中央文件（tuancentral_doc_cited），"
                            "报考动机段必引至少 1 项")

        # 专业匹配校验（警告级）
        sp = self._get("service_preference", default="")
        certs = self._get("certificates", default=[])
        certs_text = "、".join(certs) if isinstance(certs, list) else str(certs)
        major = str(self._get("applicant_major", default=""))
        if sp == "education_service" and "教师资格证" not in certs_text:
            warnings.append("教育服务专项但未取得教师资格证，专业匹配度低")
        if sp == "health_service" and "执业" not in certs_text:
            warnings.append("卫生服务专项但未取得执业医师/护士证，专业匹配度低")
        if sp == "agritech_service" and not any(
            k in major for k in ["农学", "园艺", "植保", "畜牧", "农业"]
        ):
            warnings.append("农技推广专项但专业不含农学/园艺/植保/畜牧，专业匹配度低")

        # 年限匹配校验（警告级）
        years = safe_int(self._get("service_years", default=0), default=0)
        all_text = json.dumps(self.data, ensure_ascii=False)
        if years == 1 and "考研" not in all_text and "深造" not in all_text:
            warnings.append("服务年限 1 年但全文未提及考研加分/深造规划，"
                            "建议在服务岗位意愿段说明")
        if years == 3 and "考公" not in all_text and "定向" not in all_text \
                and "留在西部" not in all_text:
            warnings.append("服务年限 3 年但全文未提及考公定向/留在西部规划，"
                            "建议在服务岗位意愿段说明")

        # 家人态度校验
        family_text = "\n".join(family) if isinstance(family, list) else str(family)
        if "沟通" not in family_text and "告知" not in family_text \
                and "商量" not in family_text:
            warnings.append("家人态度段未提及沟通过程（沟通/告知/商量），"
                            "评审会扣大分")

        # 字数校验
        total_chars = 0
        for k in ("motivation", "abilities", "service_preference_detail",
                  "family_attitude"):
            v = self._get(k, default=[])
            if isinstance(v, list):
                for s in v:
                    total_chars += count_chinese_chars(s)
        if total_chars < 1500:
            warnings.append(f"5 段合计中文字符约 {total_chars}，"
                            "偏少（建议 2000-3000 字）")
        elif total_chars > 3500:
            warnings.append(f"5 段合计中文字符约 {total_chars}，"
                            "偏多（建议 2000-3000 字）")

        # 政治面貌校验
        political = self._get("applicant_political", default="")
        if political and political not in ("中共党员", "共青团员", "群众",
                                            "中共预备党员"):
            warnings.append(f"政治面貌取值 '{political}' 不规范，"
                            "应为 中共党员/共青团员/群众 之一")

        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        self.warnings = warnings
        return warnings


# ============================================================
# 默认示例数据
# ============================================================

DEFAULT_DATA = {
    "applicant_name": "张三",
    "applicant_gender": "男",
    "applicant_ethnicity": "汉族",
    "applicant_birth": "2002 年 5 月",
    "applicant_political": "中共党员",
    "applicant_native_place": "四川省凉山彝族自治州西昌市",
    "applicant_id": "2022012345",
    "applicant_major": "汉语言文学（师范）",
    "applicant_grade": "2022 级",
    "applicant_school": "XX 师范大学文学院",
    "applicant_degree": "本科应届毕业",
    "applicant_graduation": "2026 年 6 月",
    "applicant_phone": "138XXXXXXXX",
    "applicant_email": "zhangsan@example.com",

    "apply_year": "2025",
    "service_preference": "education_service",
    "second_preference": "youth_work",
    "service_years": 1,
    "service_region_preference": "四川省凉山彝族自治州",
    "accept_redeployment": True,

    "academic_rank": "专业前 15%（8/52）",
    "gpa": "3.6/4.0",
    "core_courses": "现代汉语 92 / 古代文学 89 / 语文教学论 95 / 教育学 88",
    "certificates": [
        "高级中学教师资格证（语文）",
        "普通话二级甲等",
        "CET-6",
    ],
    "honors": [
        "2024 年国家奖学金",
        "2023 年校级一等奖学金",
        "校级优秀学生干部",
    ],
    "publications": [],
    "competitions": [
        "2024 年挑战杯省级一等奖",
        "2023 年师范生教学技能大赛省级二等奖",
    ],

    "political_detail": (
        "2023 年 6 月加入中国共产党，入党介绍人为辅导员李老师与班主任王老师。"
        "在校期间认真学习党的理论，研读《习近平与大学生朋友们》《习近平的七年"
        "知青岁月》等著作，完成'青年大学习'52 期，撰写思想汇报 6 篇。"
        "曾任院学生党支部宣传委员，组织主题党日活动 8 次。"
    ),

    "practice_experience": [
        {
            "role": "院学生会主席",
            "period": "2023.09-2024.06",
            "work": "组织学院迎新晚会、运动会、辩论赛等大型活动 12 次",
            "result": "覆盖同学 1500 人次，获评校级优秀学生干部",
        },
        {
            "role": "暑期支教志愿者",
            "period": "2023.07-2023.08",
            "work": "赴四川省凉山州西昌市安宁镇中心小学支教 21 天，"
                    "承担语文、音乐教学",
            "result": "教学课时 60 节，受益学生 45 人",
        },
    ],
    "volunteer_id": "SC5101234567890",
    "volunteer_hours": 280,

    "physical_check": "2025 年 3 月 XX 师范大学校医院体检合格，无重大疾病史",
    "physical_test": "国家学生体质健康标准良好",
    "psychological_test": "校级心理测评结论良好",
    "adaptability": (
        "曾独立骑行 318 国道（成都—拉萨，21 天），具备跨文化适应基础"
    ),

    "motivation": [
        "本人认真学习了团中央、教育部、财政部、人社部联合印发的"
        "《2025 年大学生志愿服务西部计划实施方案》及"
        "《关于做好 2025 年大学生志愿服务西部计划工作的通知》，深入了解西部计划"
        "自 2003 年实施以来累计选派 50 余万名志愿者的项目历程，明确西部计划与"
        "暑期'三下乡'社会实践、三支一扶、特岗教师等项目的区别——西部计划是"
        "1-3 年全职志愿服务，三下乡是 7-15 天暑期实践，三支一扶是 2-3 年人社部"
        "项目，特岗教师是 3 年教育部项目。本人报考的是西部计划，志愿赴西部基层"
        "服务 1 年。",

        "触发本人报考西部计划的，是 2023 年暑期到四川省凉山州西昌市安宁镇中心"
        "小学支教时结识的一位彝族女孩阿依。阿依家在大凉山腹地，每天走 2 小时"
        "山路到镇中心校上学，她说她的梦想是'考上大学回村当老师'。这一刻让本人"
        "深刻理解了'用一年时间做一件终生难忘的事'这句话的分量，也坚定了本人"
        "报考西部计划教育服务专项的志向。",

        "本人生于四川、长于四川，对凉山有天然的情感认同。在校期间两次到凉山州"
        "参加社会实践，亲眼看到凉山州师生比 1:25（低于全国平均 1:16）、普通话"
        "普及率不足 60% 的现状。本人愿用 1 年时间扎根西昌市某乡镇小学，从一名"
        "普通语文教师做起，为凉山教育事业发展贡献力量。",
    ],

    "abilities": [
        "（政治素质）本人政治面貌为中共党员，2023 年 6 月入党，入党介绍人为"
        "辅导员李老师与班主任王老师。在校期间认真学习党的理论，研读《习近平与"
        "大学生朋友们》《习近平的七年知青岁月》等著作，完成'青年大学习'52 期，"
        "撰写思想汇报 6 篇。曾任院学生党支部宣传委员，组织主题党日活动 8 次，"
        "培养了对党忠诚、服务人民的信念。",

        "（专业能力）本人在校期间学业成绩 GPA 3.6/4.0，排名专业前 15%（8/52）。"
        "已取得高级中学教师资格证（语文）、普通话二级甲等证书，具备从事中小学"
        "语文教学的能力。曾获 2024 年国家奖学金、2023 年校级一等奖学金、"
        "2024 年挑战杯省级一等奖，专业能力与所报教育服务专项匹配度高。",

        "（实践经历）本人在校期间担任院学生会主席 1 年，组织迎新晚会、运动会等"
        "大型活动 12 次，覆盖同学 1500 人次。注册志愿者编号 SC5101234567890，"
        "累计志愿服务时长 280 小时，参与 2023 年暑期凉山支教（21 天，教学 60 "
        "课时）、2024 年社区防疫志愿服务（48 小时）等典型项目。曾于 2024 年"
        "7-8 月在 XX 师范大学附属中学实习 2 个月，承担初二语文教学 30 课时，"
        "积累了实际教学经验。",

        "（身心素质）本人 2025 年 3 月校医院体检合格，无重大疾病史，国家学生"
        "体质健康标准良好。心理测评结论良好。曾独立骑行 318 国道（成都—拉萨，"
        "21 天），具备较强的抗压能力与跨文化适应能力。生活自理能力强，能适应"
        "西部基层生活条件。",
    ],

    "service_preference_detail": [
        "（专项选择）本人志愿报考西部计划教育服务专项。选择该专项的原因有三："
        "一是所学专业汉语言文学（师范）与教育服务专项高度匹配；二是本人已取得"
        "高级中学教师资格证（语文）与普通话二级甲等证书，具备从事中小学语文"
        "教学的能力；三是四川省 2024 年该项目办招募教育服务专项志愿者 280 名，"
        "缺口 80 名，需求量大。本人第二志愿为青年工作专项（曾任院学生党支部"
        "宣传委员），可服从调剂。",

        "（服务地选择）本人意向服务地为四川省凉山彝族自治州。选择该地的原因有二："
        "一是本人生于四川、长于四川，对凉山有天然的情感认同，且 2023 年暑期"
        "曾在凉山州西昌市安宁镇中心小学支教，建立了感情连接；二是凉山州是"
        "国家乡村振兴重点帮扶县集中区域，2024 年该州招募教育服务专项志愿者"
        "60 名，重点分布在西昌市、昭觉县、布拖县等 7 个县（市）。本人服从"
        "服务地调剂，愿意赴国家乡村振兴重点帮扶县服务。",

        "（服务年限）本人申请服务 1 年。选择 1 年期的原因是本人计划 2026 年"
        "服务期满后报考硕士研究生（享受西部计划志愿者考研加分 10 分政策），"
        "1 年期已满足政策要求。服务期满后，本人计划报考教育学硕士（学科教学·"
        "语文方向），将西部服务经历转化为长期教育事业。",
    ],

    "family_attitude": [
        "（家人态度与沟通过程）本人于 2025 年 3 月向家人正式告知报考西部计划"
        "的决定。父母初始反应是担忧本人安全与未来出路。经过 3 次深入沟通，"
        "家人从担忧转为支持。父亲表示：'西部计划是国家项目，我们支持你的"
        "选择，但要照顾好自己。'母亲表示：'家里有我们，你安心服务。'家人"
        "最终态度为支持，已签署家长知情同意书。具体顾虑化解：（1）父亲担忧"
        "本人安全，本人向父亲展示了西部计划志愿者保险条款（保额 50 万元）"
        "与服务地住宿条件后，父亲态度转为支持；（2）母亲担忧本人未来出路，"
        "本人向母亲说明西部计划服务期满后享受考研加分、考公定向、学费代偿"
        "等政策，母亲表示理解。",

        "（决心表态）本人郑重表态：志愿赴四川凉山基层服务 1 年，服务期间"
        "服从组织安排，遵守志愿者管理办法，完成本职工作，不擅自离岗，不挑拣"
        "岗位。本人将以'用一年时间做一件终生难忘的事'为座右铭，扎根凉山、"
        "服务基层，把青春奉献给祖国最需要的地方，让西部服务经历成为人生最"
        "宝贵的财富。",
    ],

    "version": "standard",
    "include_signature": True,
    "tuancentral_doc_cited": [
        "《2025 年大学生志愿服务西部计划实施方案》",
        "《关于做好 2025 年大学生志愿服务西部计划工作的通知》",
    ],
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="大学生志愿服务西部计划申请书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n"
            "JSON 字段定义详见 SKILL.md 第十一章。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
