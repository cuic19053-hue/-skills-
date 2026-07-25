#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""互联网+大学生创新创业大赛商业计划书 docx 生成器

格式标准：A4，页边距上下 2.54cm 左右 2.5cm；正文宋体小四 1.5 倍行距首行缩进 2 字符；
一级标题黑体三号居中；二级黑体小三左对齐；三级宋体四号加粗；表格宋体五号居中。

11 栏目：封面 / 执行摘要 / 项目背景与市场分析 / 产品服务介绍 / 商业模式 / 运营现状 /
财务预测 / 融资计划 / 团队介绍 / 风险分析 / 个人成长与团队协作（2025 新增重点）。
三大赛道：高教主赛道 / 红色之旅 / 职教赛道。

使用：python build.py --data data.json --out output.docx ； python build.py --demo --out demo.docx
JSON 字段详见 SKILL.md 第十一章。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22); SIZE_XIAO_ER = Pt(18); SIZE_SAN = Pt(16); SIZE_XIAO_SAN = Pt(15)
SIZE_SI = Pt(14); SIZE_XIAO_SI = Pt(12); SIZE_WU = Pt(10.5); SIZE_XIAO_WU = Pt(9)

PAGE_WIDTH_CM = 21.0; PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54; MARGIN_LEFT_RIGHT_CM = 2.5

TRACK_LABEL = {"main": "高教主赛道", "red_journey": "红色之旅赛道", "vocational": "职教赛道"}


# ============================================================
# 工具函数（中英文同步设置 eastAsia/ascii/hAnsi）
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_SAN,
                                    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    first_line_indent=False, space_before=12, space_after=12)


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SAN,
                                    bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                    first_line_indent=False, space_before=6, space_after=6)


def add_heading_level3(doc, text: str):
    """三级标题：宋体四号加粗，左对齐"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_SI,
                                    bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                    first_line_indent=False, space_before=6, space_after=3)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
                                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                    first_line_indent=indent, line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG, font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == len(row) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], val, font_name=FONT_SONG, font_size=SIZE_WU,
                          bold=False, alignment=align)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_canvas_9grid(doc, canvas_items: List[Dict[str, str]]):
    """商业模式画布 9 宫格表格（3 行 × 3 列，每格含要素名+内容）"""
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, item in enumerate(canvas_items[:9]):
        r, c = divmod(idx, 3)
        cell = table.rows[r].cells[c]
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(item.get("element", ""))
        set_run_font(run1, font_name=FONT_HEI, font_size=SIZE_XIAO_WU, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.line_spacing = 1.2
        run2 = p2.add_run(item.get("content", ""))
        set_run_font(run2, font_name=FONT_SONG, font_size=SIZE_XIAO_WU)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(5.3)
    doc.add_paragraph()


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM); section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM); section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM); section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_kpi_table(doc, kpi_rows: List[Dict[str, str]]):
    """运营现状 KPI 表（5 列：指标/当前值/Year1/Year2/Year3）"""
    if not kpi_rows:
        return
    headers = ["指标", "当前值", "Year1 目标", "Year2 目标", "Year3 目标"]
    rows = [[r.get("metric", ""), str(r.get("current", "")), str(r.get("y1", "")),
             str(r.get("y2", "")), str(r.get("y3", ""))] for r in kpi_rows if isinstance(r, dict)]
    add_table_from_data(doc, headers, rows, col_widths=[4.0, 2.5, 2.8, 2.8, 3.4])


def add_milestone_table(doc, milestones: List[Dict[str, str]]):
    """关键里程碑表（3 列：时间/事件/数据）"""
    if not milestones:
        return
    headers = ["时间", "里程碑事件", "关键数据"]
    rows = [[m.get("time", ""), m.get("event", ""), m.get("data", "")]
            for m in milestones if isinstance(m, dict)]
    add_table_from_data(doc, headers, rows, col_widths=[3.0, 6.5, 6.0])


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """互联网+大学生创新创业大赛商业计划书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def add_h1(self, text): return add_heading_level1(self.doc, text)
    def add_h2(self, text): return add_heading_level2(self.doc, text)
    def add_h3(self, text): return add_heading_level3(self.doc, text)
    def add_para(self, text, indent=True): return add_body_paragraph(self.doc, text, indent=indent)
    def add_table(self, headers, rows, col_widths=None): return add_table_from_data(self.doc, headers, rows, col_widths)
    def add_page_break(self): add_page_break(self.doc)

    # 封面

    def _add_cover(self):
        """封面：大赛名称 + 文档类型副标题 + 6 行下划线信息"""
        for _ in range(2):
            self.doc.add_paragraph()
        add_paragraph_with_format(self.doc, '中国国际"互联网+"大学生创新创业大赛',
                                  font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  first_line_indent=False, space_before=12, space_after=12)
        add_paragraph_with_format(self.doc, "商业计划书",
                                  font_name=FONT_HEI, font_size=SIZE_SAN, bold=False,
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  first_line_indent=False, space_after=24)
        for _ in range(2):
            self.doc.add_paragraph()
        track_key = self._get("track", default="main")
        track_label = TRACK_LABEL.get(track_key, "高教主赛道")
        group = self._get("group", default="本科生组")
        info_items = [
            ("项目名称", self._get("project_name")),
            ("赛    道", track_label),
            ("组    别", group),
            ("团队名称", self._get("team_name", default=self._get("project_name"))),
            ("负责人", self._get("leader_name")),
            ("学校", self._get("school")),
            ("日    期", self._get("apply_date")),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0; pf.space_before = Pt(6); pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI, font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True
        self.add_page_break()

    # 一、执行摘要（500~800 字，最重要）

    def _add_executive_summary(self):
        """执行摘要：500~800 字，六要素齐全"""
        self.add_h1("一、执行摘要")
        summary = self._get("executive_summary", default="")
        if summary:
            self.add_para(summary)
        else:
            self.add_para(
                "（请填写执行摘要 500~800 字，必含六要素：① 一句话项目（为 [谁] 解决 [什么问题] 通过 [什么方式]，≤30 字）；"
                "② 市场机会（TAM/SAM/SOM 关键数字 + 数据来源）；③ 核心产品（MVP 主要功能 + 关键指标如用户数/月活/月营收）；"
                "④ 商业模式（盈利模式 + 定价 + CAC/LTV）；⑤ 团队亮点（复合背景 + 已有成绩 + 指导教师）；"
                "⑥ 融资需求（本轮金额 + 出让股权 + 用途分配）。评审 30 秒决定深读还是淘汰，是 BP 最重要栏目。）"
            )

        # 个人成长摘要小节（2025 新增重点，30% 权重）
        growth_summary = self._get("personal_growth", default={})
        if isinstance(growth_summary, dict) and growth_summary.get("transformation"):
            self.add_h2("项目过程中我的成长（2025 评审 30% 维度）")
            transform = growth_summary.get("transformation", "")
            self.add_para(transform if transform else
                          "（请填写第一人称成长叙事：从 X 变为 Y + 学到的硬技能 + 软技能 + 失败复盘。）")

    # 二、项目背景与市场分析（1500~2000 字）

    def _add_market_analysis(self):
        """项目背景与市场分析：4 子节 + TAM/SAM/SOM 表 + 竞品对比表"""
        self.add_h1("二、项目背景与市场分析")
        market = self._get("market_analysis", default={})
        if not isinstance(market, dict):
            market = {}
        self.add_h2("（一）行业背景与政策导向")
        bg = market.get("industry_background", "")
        self.add_para(bg if bg else
                      "（请填写行业背景 300~400 字：3 句话讲行业规模/增速/政策，2 句话讲细分场景痛点。"
                      "必须有权威数据来源：艾瑞/IDC/国家统计局/行业白皮书 + 政策文件名。）")
        self.add_h2("（二）目标市场（TAM/SAM/SOM）")
        tss = market.get("tam_sam_som", [])
        if tss:
            rows = [[item.get("level", ""), item.get("definition", ""), item.get("scale", "")]
                    for item in tss if isinstance(item, dict)]
            self.add_table(["层级", "定义", "规模"], rows, col_widths=[2.5, 9.0, 4.5])
        else:
            self.add_para("（请填写 TAM/SAM/SOM 三级市场测算表，每行含数据来源。"
                          "TAM 总市场 / SAM 可服务市场 / SOM 可获取市场。）")
        self.add_h2("（三）竞品分析")
        comps = market.get("competitors", [])
        if comps:
            rows = [[c.get("name", ""), c.get("positioning", ""), c.get("users", ""),
                     c.get("advantage", ""), c.get("disadvantage", "")] for c in comps if isinstance(c, dict)]
            self.add_table(["竞品", "定位", "用户规模", "优势", "劣势"], rows,
                           col_widths=[2.2, 2.8, 2.5, 4.0, 4.5])
        else:
            self.add_para("（请填写竞品分析表，3~5 个直接竞品 + 2~3 个间接竞品，"
                          "5 列对比：名称/定位/用户规模/优势/劣势。）")
        differentiation = market.get("differentiation", "")
        if differentiation:
            self.add_para(f"本项目差异化定位：{differentiation}")
        self.add_h2("（四）用户画像与需求验证")
        persona = market.get("user_persona", "")
        self.add_para(persona if persona else
                      "（请填写用户画像 300~400 字，6 要素齐全：基本信息+行为+痛点+需求+付费意愿+决策因素。"
                      "附调研样本量 N≥50。）")

    # 三、产品/服务介绍（1000~1500 字）

    def _add_product_service(self):
        """产品/服务介绍：3 子节"""
        self.add_h1("三、产品/服务介绍")
        ps = self._get("product_service", default={})
        if not isinstance(ps, dict):
            ps = {}
        self.add_h2("（一）产品形态与核心功能")
        features = ps.get("features", [])
        if isinstance(features, str):
            features = [features]
        if features:
            for i, f in enumerate(features, 1):
                self.add_para(f"{i}. {f}")
        else:
            self.add_para("（请填写 MVP 功能清单 5~8 个核心功能，每个含名称 + 解决什么问题 + 关键指标。）")
        self.add_h2("（二）核心功能演示")
        demo = ps.get("demo", "")
        self.add_para(demo if demo else "（请填写 1~2 个核心功能详细描述，含用户操作流程截图说明。）")
        self.add_h2("（三）技术实现与知识产权")
        tech = ps.get("tech_impl", "")
        self.add_para(tech if tech else
                      "（请填写技术架构（前端/后端/数据库）+ 关键第三方服务 + 研发投入估算 + 知识产权情况（软著/专利/在申）。）")
        ip = ps.get("ip_status", "")
        if ip:
            self.add_para(f"知识产权情况：{ip}")

    # 四、商业模式（800~1200 字 + 9 宫格画布）

    def _add_business_model(self):
        """商业模式：画布 + 盈利 + 定价 + CAC/LTV

商业模式画布 9 要素：客户细分 / 价值主张 / 渠道通路 / 客户关系 /
收入来源 / 核心资源 / 关键业务 / 重要伙伴 / 成本结构。
每要素 30~50 字，缺一不可。LTV/CAC ≥ 3 为健康商业模式。
"""
        self.add_h1("四、商业模式")
        bm = self._get("business_model", default={})
        if not isinstance(bm, dict):
            bm = {}

        self.add_h2("（一）商业模式画布（9 要素）")
        canvas = bm.get("canvas", [])
        if canvas and isinstance(canvas, list):
            add_canvas_9grid(self.doc, canvas)
        else:
            self.add_para("（请填写商业模式画布 9 要素：客户细分/价值主张/渠道通路/客户关系/"
                          "收入来源/核心资源/关键业务/重要伙伴/成本结构。每格 30~50 字，缺一不可。）")

        self.add_h2("（二）盈利模式")
        streams = bm.get("revenue_streams", [])
        if isinstance(streams, str):
            streams = [streams]
        if streams:
            for i, s in enumerate(streams, 1):
                self.add_para(f"{i}. {s}")
        else:
            self.add_para("（请填写盈利模式 1~3 种收入来源，"
                          "每种含定价 + 计算依据 + Year1 占营收比。）")

        self.add_h2("（三）定价策略")
        pricing = bm.get("pricing", "")
        self.add_para(pricing if pricing else
                      "（请填写定价策略：成本加成/竞品对标/用户价值三方法，"
                      "含单品价格/月费/年费/竞品对照/用户付费意愿调研数据。）")

        self.add_h2("（四）CAC/LTV 测算")
        cac_ltv = bm.get("cac_ltv", "")
        self.add_para(cac_ltv if cac_ltv else
                      "（请填写 CAC/LTV 测算：CAC = 推广预算/新增用户数；"
                      "LTV = 客单价 × 复购次数 × 毛利率；LTV/CAC ≥ 3 为健康。）")

    # 五、运营现状（800~1200 字 + KPI 表 + 里程碑表）

    def _add_operation_status(self):
        """运营现状：成绩 + 里程碑 + KPI"""
        self.add_h1("五、运营现状")
        op = self._get("operations", default={})
        if not isinstance(op, dict):
            op = {}

        self.add_h2("（一）已取得成绩")
        achievements = op.get("achievements", "")
        self.add_para(achievements if achievements else
                      "（请填写已取得成绩 300~400 字：用户数、营收、合作客户、签约订单、"
                      "媒体报道、获奖。无数据则写 MVP 上线时间 + 内测用户数 + 首批合作意向。）")

        self.add_h2("（二）关键里程碑")
        milestones = op.get("milestones", [])
        if milestones:
            add_milestone_table(self.doc, milestones)
        else:
            self.add_para("（请填写关键里程碑表：时间/事件/数据，含 MVP 上线、首批用户、"
                          "首笔营收、首个合作、获奖等 6~10 个节点。）")

        self.add_h2("（三）关键运营指标 KPI")
        kpi_rows = op.get("kpi", [])
        if kpi_rows:
            add_kpi_table(self.doc, kpi_rows)
        else:
            self.add_para("（请填写 KPI 表：北极星指标 + 次级指标 + 当前值 + Year1/2/3 目标，"
                          "如月活跃交易用户、注册→认证转化率、月复购率、付费会员率。）")

        # 红色之旅赛道额外子节
        track = self._get("track", default="main")
        if track == "red_journey":
            self.add_h2("（四）红色帮扶成效")
            red = op.get("red_help", "")
            self.add_para(red if red else
                          "（红色之旅赛道必填：帮扶老区/乡村名称、带动农户数、"
                          "增收数据、合作社数量、社会效益。）")
        elif track == "vocational":
            self.add_h2("（四）技能岗位对接")
            voc = op.get("skill_jobs", "")
            self.add_para(voc if voc else
                          "（职教赛道必填：对接企业数、提供岗位数、学生实训人数、"
                          "技能认证通过率。）")

    # 六、财务预测（3 年表格 + 说明）

    def _add_financial_forecast(self):
        """财务预测：3 年表格 + 盈亏平衡点说明 + 财务合理性说明

3 年财务表必备项：注册用户 / 月活 / 营收 / 成本 / 毛利率 / 期间费用 /
净利润 / 净利率。每行需有算法（营收 = 用户数 × 客单价 × 复购次数）。
Year1 净利率可为负，Year3 应转正。
"""
        self.add_h1("六、财务预测")
        fin = self._get("financial_forecast", default={})
        if not isinstance(fin, dict):
            fin = {}

        rows_data = fin.get("rows", [])
        if rows_data:
            rows = [[r.get("item", ""), str(r.get("y1", "")),
                     str(r.get("y2", "")), str(r.get("y3", ""))]
                    for r in rows_data if isinstance(r, dict)]
            self.add_table(["财务项目", "Year1", "Year2", "Year3"], rows,
                           col_widths=[5.0, 3.5, 3.5, 3.5])
        else:
            self.add_para("（请填写 3 年财务预测表，至少含：注册用户/月活用户/营业收入/"
                          "营业成本/毛利率/期间费用/净利润/净利率 8 行。）")

        self.add_h2("盈亏平衡点说明")
        be = fin.get("breakeven", "")
        self.add_para(be if be else
                      "（请填写盈亏平衡点：YearX 第 Y 个月月营收达 Z 万元时毛利率覆盖固定成本。"
                      "BEP = 固定成本 / (1 - 变动成本率)。）")

        self.add_h2("财务合理性说明")
        rationale = fin.get("rationale", "")
        self.add_para(rationale if rationale else
                      "（请填写财务合理性：Year1 用户数 ≤ SOM × 5%；Year3 ≤ SOM × 30%；"
                      "毛利率 30%~60%；净利率 Year1 可负、Year3 转正；增长率 Year2/Year1 ≥ 3 倍。）")

    # 七、融资计划（500~800 字）

    def _add_financing_plan(self):
        """融资计划：金额 + 用途 + 下一轮"""
        self.add_h1("七、融资计划")
        fp = self._get("financing_plan", default={})
        if not isinstance(fp, dict):
            fp = {}

        self.add_h2("（一）本轮融资需求")
        amount = fp.get("amount", "")
        equity = fp.get("equity", "")
        valuation = fp.get("valuation", "")
        if amount or equity or valuation:
            self.add_para(f"本轮融资需求：{amount}，出让股权 {equity}，估值 {valuation}。")
        else:
            self.add_para("（请填写本轮融资需求：金额 + 出让股权 + 估值，"
                          "如'拟融资 50 万元，出让 10% 股权，估值 500 万元'。）")

        self.add_h2("（二）资金用途分配")
        usage = fp.get("usage", "")
        self.add_para(usage if usage else
                      "（请填写资金用途分配：百分比 + 金额 + 说明，"
                      "如'产品研发 40%（20 万）/ 用户获取 30%（15 万）/ "
                      "团队扩张 20%（10 万）/ 运营储备 10%（5 万）'。）")

        self.add_h2("（三）下一轮规划")
        nxt = fp.get("next_round", "")
        self.add_para(nxt if nxt else
                      "（请填写下一轮规划：18~24 个月后拟启动 A 轮，预期估值、"
                      "融资额、用途。估值方法用'成本法 + 行业对标'双轨验证。）")

    # 八、团队介绍（每人 100~200 字 + 表格）

    def _add_team_intro(self):
        """团队介绍：核心成员表 + 成员简介 + 指导教师"""
        self.add_h1("八、团队介绍")
        team = self._get("team_intro", default={})
        if not isinstance(team, dict):
            team = {}

        self.add_h2("（一）核心成员表")
        members = team.get("members", [])
        if members:
            rows = [[m.get("name", ""), m.get("role", ""),
                     m.get("background", ""), m.get("duty", "")]
                    for m in members if isinstance(m, dict)]
            self.add_table(["姓名", "职务", "背景", "分工"], rows,
                           col_widths=[2.0, 3.0, 5.5, 5.0])
        else:
            self.add_para("（请填写核心成员表：姓名/职务/背景/分工，"
                          "5~8 人，商科+技术+设计+运营复合，总团队 ≤15 人。）")

        self.add_h2("（二）核心成员简介")
        details = team.get("member_details", [])
        if details:
            for d in details:
                if isinstance(d, dict):
                    name = d.get("name", "")
                    detail = d.get("detail", "")
                    self.add_para(f"{name}：{detail}")
        else:
            self.add_para("（请填写核心成员简介，每人 100~200 字：专业背景 + 实习/项目经历 + "
                          "在本项目中具体职责 + 已完成的关键动作。突出'复合背景 + 互补能力'。）")

        self.add_h2("（三）指导教师与顾问")
        adv = team.get("advisor_bg", "")
        self.add_para(adv if adv else
                      "（请填写指导教师背景：职称、研究方向、主持项目、"
                      "指导学生创业经历；外部顾问（如有）：行业专家/投资人/校友。）")

    # 九、风险分析（500 字 + 表格）

    def _add_risk_analysis(self):
        """风险分析：4 类风险表"""
        self.add_h1("九、风险分析")
        risks = self._get("risk_analysis", default=[])
        if isinstance(risks, str):
            risks = [risks]
        if risks:
            rows = [[r.get("type", ""), r.get("risk", ""),
                     r.get("prob", ""), r.get("impact", ""),
                     r.get("measure", "")]
                    for r in risks if isinstance(r, dict)]
            self.add_table(["风险类型", "具体风险", "概率", "影响", "应对措施"], rows,
                           col_widths=[2.2, 3.8, 1.5, 1.5, 7.0])
        else:
            self.add_para("（请填写风险分析表，4 类风险齐全：市场/技术/运营/财务，"
                          "每类含具体风险 + 概率 + 影响 + 应对措施。）")

        track = self._get("track", default="main")
        if track == "red_journey":
            self.add_para("红色之旅赛道额外关注：帮扶成效不可持续风险、政策变动风险、"
                          "乡村基层组织配合度风险。", indent=False)
        elif track == "vocational":
            self.add_para("职教赛道额外关注：校企合作断裂风险、技能认证政策变动风险、"
                          "实训基地合规风险。", indent=False)

    # 十、个人成长与团队协作（2025 新增重点，500~800 字）

    def _add_personal_growth(self):
        """个人成长与团队协作：2025 新增重点栏目"""
        self.add_h1("十、个人成长与团队协作")
        self.add_h2("（一）个人成长（第一人称）")
        growth = self._get("personal_growth", default={})
        if not isinstance(growth, dict):
            growth = {}
        items = [
            ("role", "角色与动作", "你在项目中承担的具体角色 + 完成的具体动作。"),
            ("transformation", "从 X 变为 Y", "从 X 变为 Y 的具体变化，如'从只会做 PPT 的工商管理学生变为能独立完成 LTV/CAC 模型的初创团队 CEO'。"),
            ("hard_skills", "学到的硬技能", "硬技能：市场测算/财务建模/产品原型/数据分析/编程等。"),
            ("soft_skills", "学到的软技能", "软技能：沟通/协调/抗压/决策/领导力。"),
            ("failure", "失败复盘", "至少 1 个具体失败案例 + 复盘收获。"),
        ]
        for key, label, hint in items:
            val = growth.get(key, "")
            if val:
                self.add_para(f"{label}：{val}")
            else:
                self.add_para(f"（请填写{label}：{hint}）")
        self.add_h2("（二）团队协作")
        collab = self._get("team_collaboration", default={})
        if not isinstance(collab, dict):
            collab = {}
        division = collab.get("division", "")
        conflict = collab.get("conflict", "")
        tools = collab.get("tools", "")
        if division:
            self.add_para(f"分工机制：{division}")
        if conflict:
            self.add_para(f"冲突案例：{conflict}")
        if tools:
            self.add_para(f"协作工具：{tools}")
        if not (division or conflict or tools):
            self.add_para("（请填写团队协作 3 要素：分工机制 + 冲突案例（至少 1 个）+ "
                          "协作工具（飞书/钉钉/Notion + 周会/日站会频次）。）")

    # 主构建方法

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 11 栏目，生成 docx。返回实际保存路径。"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_cover()
            self._add_executive_summary()
            self._add_market_analysis()
            self._add_product_service()
            self._add_business_model()
            self._add_operation_status()
            self._add_financial_forecast()
            self._add_financing_plan()
            self._add_team_intro()
            self._add_risk_analysis()
            self._add_personal_growth()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 商业计划书已生成：{output_path}")
        return str(output_path)

    # 数据校验

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []
        p0_fields = [("project_name", "项目名称"), ("track", "赛道"),
                     ("leader_name", "负责人姓名"), ("school", "学校"),
                     ("advisor_name", "指导教师姓名"), ("college", "学院")]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")
        track = self._get("track", default="main")
        if track not in TRACK_LABEL:
            warnings.append(f"赛道值 {track} 不在 main/red_journey/vocational 之内")
        for key, name in [("executive_summary", "执行摘要"), ("market_analysis", "市场分析"),
                          ("business_model", "商业模式"), ("financial_forecast", "财务预测"),
                          ("personal_growth", "个人成长（2025 30% 维度）")]:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}），将使用占位文本")
        bm = self._get("business_model", default={})
        if isinstance(bm, dict):
            canvas = bm.get("canvas", [])
            if canvas and len(canvas) < 9:
                warnings.append(f"商业模式画布仅 {len(canvas)} 要素，建议补齐 9 要素")
        pg = self._get("personal_growth", default={})
        if isinstance(pg, dict) and pg:
            missing = [k for k in ["role", "transformation", "hard_skills",
                                   "soft_skills", "failure"] if not pg.get(k)]
            if missing:
                warnings.append(f"个人成长段落缺要素：{', '.join(missing)}（2025 30% 维度）")
        team = self._get("team_intro", default={})
        if isinstance(team, dict):
            members = team.get("members", [])
            if members and len(members) > 15:
                warnings.append(f"团队成员 {len(members)} 人，超过互联网+上限 15 人")
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据
# ============================================================

DEFAULT_DATA = {
    "project_name": "校园闲置物品流转平台'易舍'",
    "track": "main", "group": "本科生组", "team_name": "易舍团队",
    "leader_name": "张三", "leader_id": "202212345",
    "leader_major": "工商管理", "leader_grade": "2022 级 大三", "leader_phone": "138XXXXXXXX",
    "advisor_name": "李教授", "advisor_title": "副教授", "advisor_research": "创业管理",
    "college": "工商管理学院", "school": "XX 大学", "apply_date": "2025 年 4 月 20 日",
    "executive_summary": (
        "本项目针对高校校园闲置物品流转效率低、传统跳蚤市场信息不对称的痛点，"
        "开发基于微信小程序的校园闲置物品流转平台'易舍'，"
        "为高校 18~24 岁在校生提供同校面交 + 信用背书的二手交易服务。"
        "中国二手交易市场 2024 年规模达 5000 亿元（艾瑞咨询），"
        "校园场景占比约 5%（约 250 亿元），年均增速 20%+。"
        "本项目 SAM 6.4 亿元/年，3 年内 SOM 1272 万元/年。"
        "核心产品'易舍'小程序已上线，含校园认证、LBS 同校匹配、信用积分、一键发布、"
        "担保交易 5 大功能，注册 5000，月活 1500，月交易额 5 万元。"
        "采用'交易抽佣 5% + 增值会员 9.9 元/月'双轨盈利模式，"
        "CAC 5 元/LTV 48 元，LTV/CAC = 9.6。"
        "团队 5 人跨学科复合，含 2 名有创业实习经历的核心成员，指导教师为创业管理副教授。"
        "本轮拟融资 50 万元，出让 10% 股权，资金用于研发 40% / 获客 30% / 团队 20% / 储备 10%。"
        "预期 Year3 月活 6 万、年营收 200 万元、净利率 30%。"
    ),
    "market_analysis": {
        "industry_background": (
            "2024 年中国二手交易市场规模达 5000 亿元（艾瑞咨询《2024 中国二手电商报告》），"
            "年均增速 20%+，校园场景占比约 5%（约 250 亿元）。"
            "国家发改委 2024 年《促进绿色消费实施方案》鼓励校园二手物品循环利用。"
            "传统校园跳蚤市场信息散乱、跨城交易邮费高、陌生人交易信任度低，"
            "校园闲置物品流转率不足 15%。"
        ),
        "tam_sam_som": [
            {"level": "TAM", "definition": "全国高校 4700 万在校生 × 人均年闲置消费 530 元", "scale": "250 亿元/年"},
            {"level": "SAM", "definition": "本省 50 所高校 120 万在校生", "scale": "6.4 亿元/年"},
            {"level": "SOM", "definition": "本市 5 所高校 12 万在校生 × 渗透率 20%", "scale": "1272 万元/年"},
        ],
        "competitors": [
            {"name": "闲鱼", "positioning": "全品类二手", "users": "5 亿注册", "advantage": "流量大、品牌强", "disadvantage": "校园场景弱、信任度低"},
            {"name": "转转", "positioning": "数码 3C 二手", "users": "2 亿注册", "advantage": "验机服务完善", "disadvantage": "非校园场景"},
            {"name": "校园跳蚤群", "positioning": "线下/微信群", "users": "不可统计", "advantage": "信任度高", "disadvantage": "信息散乱、效率低"},
        ],
        "differentiation": "聚焦校园场景、LBS 同校面交、信用积分体系、AI 一键发布。",
        "user_persona": (
            "典型用户'小张'，男，21 岁，大三，月生活费 2000 元，年闲置物品价值约 500 元。"
            "曾用闲鱼卖旧教材，因跨城交易邮费 12 元占售价 30% 放弃。"
            "理想方案是同校面交、无需邮费、信用背书，可接受 9.9 元/月会员费。"
            "决策因素：信任度 > 价格 > 便利性。访谈 50 名目标用户，42 人愿意为'同校面交+信用背书'付费，"
            "68% 接受 9.9 元/月会员费。"
        ),
    },
    "product_service": {
        "features": [
            "校园认证：学信网/校园邮箱认证，解决陌生人交易信任问题，预期认证率 ≥80%。",
            "LBS 同校匹配：基于地理位置匹配同校买卖双方，预期面交转化率 ≥60%。",
            "信用积分：交易评价+实名认证+校园认证累计积分，积分高者优先展示。",
            "一键发布：拍照识别+AI 描述生成，发布时间从 5 分钟降至 30 秒。",
            "担保交易：买家确认收货后打款，降低交易风险。",
        ],
        "demo": "核心功能'LBS 同校匹配'演示：用户发布闲置物品后系统自动匹配同校潜在买家，买家可查看物品位置（精确到教学楼）、发起面交邀约、双方约定时间地点，完成后互相评价。整个流程 5 步以内完成，平均成交时间 2.3 天。",
        "tech_impl": "前端微信小程序（Taro 3.x）；后端 Node.js + Express；数据库 MySQL 8.0 + Redis 7.0；第三方腾讯云 LBS / 阿里云 OSS / 微信支付。研发投入 3 人 × 4 月 = 12 人月，约 6 万元（校内外包价折算）。",
        "ip_status": "软件著作权 1 项（《易舍校园闲置流转平台 V1.0》已登记，登记号 2024SR1234567）。",
    },
    "business_model": {
        "canvas": [
            {"element": "客户细分", "content": "高校 18~24 岁在校生，月生活费 1500~3000 元"},
            {"element": "价值主张", "content": "同校面交、信用背书、一键发布"},
            {"element": "渠道通路", "content": "微信小程序、校园社团合作、社群裂变"},
            {"element": "客户关系", "content": "自助服务+社群运营+客服介入"},
            {"element": "收入来源", "content": "交易抽佣 5%+增值会员 9.9 元/月+广告位"},
            {"element": "核心资源", "content": "小程序代码、用户数据、校园合作关系"},
            {"element": "关键业务", "content": "产品开发、用户运营、商家拓展"},
            {"element": "重要伙伴", "content": "高校社团、腾讯云、顺丰快递"},
            {"element": "成本结构", "content": "服务器 1500/月、获客 5 元/人、人力 0（学生）"},
        ],
        "revenue_streams": [
            "交易抽佣：每笔成交收取 5% 服务费，参考闲鱼免费、转转 1%，校园场景溢价定 5%。Year1 占营收 70%。",
            "增值会员：9.9 元/月，含优先展示、信用加权、免抽佣券。68% 用户接受。Year1 占营收 20%。",
            "校园商家广告：本地商家投放，CPM 50 元。Year1 占营收 10%。",
        ],
        "pricing": "对照竞品定价：闲鱼免费、转转 1% 抽佣。本项目定 5% 抽佣 + 9.9 元/月会员。基于 50 份用户调研，68% 用户接受 9.9 元/月，付费转化率预期 5%。",
        "cac_ltv": "CAC = 推广 5000 元 / 新增 1000 用户 = 5 元/人；LTV = 客单价 30 × 复购 4 × 毛利率 40% = 48 元；LTV/CAC = 9.6（健康，≥3）；回收期 5 个月。",
    },
    "operations": {
        "achievements": "项目自 2024 年 9 月启动，已上线 MVP 小程序，注册 5000，月活 1500，月交易额 5 万元，月营收 2500 元。已与本市 5 所高校 8 个社团建立合作，签约 3 家本地商家广告投放。获 2025 年校'互联网+'金奖。",
        "milestones": [
            {"time": "2024.09", "event": "团队组建", "data": "5 人"},
            {"time": "2024.11", "event": "MVP 上线", "data": "注册 500"},
            {"time": "2024.12", "event": "首批 3 校推广", "data": "注册 2000"},
            {"time": "2025.02", "event": "首笔营收", "data": "月营收 1 万元"},
            {"time": "2025.04", "event": "软著登记", "data": "1 项"},
            {"time": "2025.06", "event": "校赛金奖", "data": "晋级省赛"},
            {"time": "2025.08", "event": "5 校覆盖", "data": "注册 5000"},
        ],
        "kpi": [
            {"metric": "月活跃交易用户 MATU", "current": "800", "y1": "5000", "y2": "2 万", "y3": "6 万"},
            {"metric": "注册→认证转化率", "current": "65%", "y1": "80%", "y2": "85%", "y3": "90%"},
            {"metric": "发布→成交转化率", "current": "22%", "y1": "30%", "y2": "35%", "y3": "40%"},
            {"metric": "月复购率", "current": "18%", "y1": "25%", "y2": "30%", "y3": "35%"},
            {"metric": "付费会员率", "current": "2%", "y1": "5%", "y2": "8%", "y3": "12%"},
            {"metric": "月营收（万元）", "current": "0.25", "y1": "1", "y2": "5", "y3": "17"},
        ],
    },
    "financial_forecast": {
        "rows": [
            {"item": "注册用户（万）", "y1": "2", "y2": "8", "y3": "20"},
            {"item": "月活用户（万）", "y1": "0.5", "y2": "2", "y3": "6"},
            {"item": "营业收入（万元）", "y1": "12", "y2": "60", "y3": "200"},
            {"item": "营业成本（万元）", "y1": "8", "y2": "35", "y3": "100"},
            {"item": "毛利率", "y1": "33%", "y2": "42%", "y3": "50%"},
            {"item": "期间费用（万元）", "y1": "7", "y2": "17", "y3": "40"},
            {"item": "净利润（万元）", "y1": "-3", "y2": "8", "y3": "60"},
            {"item": "净利率", "y1": "-25%", "y2": "13%", "y3": "30%"},
        ],
        "breakeven": "盈亏平衡点：Year2 第 6 个月。届时月营收达 7.5 万元，毛利率 42% 覆盖月固定成本 3 万元（服务器 1500 + 客服 500 + 推广 2000 + 杂项 1000）。固定成本/毛利率 = 3/42% = 7.14 万元。",
        "rationale": "Year1 用户 2 万 ≤ SOM 12.7 万 × 5%；Year3 用户 20 万 ≤ SOM × 30%。毛利率 33%→50%（互联网平均 30%~60%）。净利率 Year1 -25% 可接受，Year3 30% 转正。增长率 Year2/Year1 = 5 倍，Year3/Year2 = 3.3 倍。",
    },
    "financing_plan": {
        "amount": "50 万元", "equity": "10%", "valuation": "500 万元",
        "usage": "产品研发 40%（20 万）：3 名工程师 6 个月薪资 + 服务器扩容 + 第三方 API；用户获取 30%（15 万）：5 校地推 + 社群裂变 + KOL 投放；团队扩张 20%（10 万）：1 名市场总监 + 1 名客户成功；运营储备 10%（5 万）：法务/财务/应急。",
        "next_round": "下一轮规划：18~24 个月后拟启动 A 轮，预期估值 2000 万元，融资 300 万元用于跨城复制（覆盖 10 城 50 所高校）。估值方法：成本法 + 行业 PS 对标（闲鱼 PS 8 倍，本项目保守 PS 4 倍 × Year2 营收 60 万 = 240 万，加用户资产溢价至 500 万）。",
    },
    "team_intro": {
        "members": [
            {"name": "张三", "role": "CEO/创始人", "background": "工商管理 大三，校 SRT 项目负责人、电商实习", "duty": "战略+融资"},
            {"name": "李四", "role": "CPO/联合创始人", "background": "计算机 大三，字节产品实习、App 上线 1 款", "duty": "产品+用户"},
            {"name": "王五", "role": "CTO/联合创始人", "background": "软件工程 大三，全栈 2 年、GitHub 1k star", "duty": "技术+架构"},
            {"name": "赵六", "role": "CMO/联合创始人", "background": "市场营销 大二，校园社群运营 5000+ 用户", "duty": "市场+获客"},
            {"name": "孙七", "role": "设计负责人", "background": "视觉传达 大三，红点设计奖入围", "duty": "UI+品牌"},
        ],
        "member_details": [
            {"name": "张三", "detail": "工商管理大三，GPA 3.85/4.0，专业排名 3/87。曾任校 SRT 项目负责人，完成 2 万元预算项目管理。京东电商实习 3 个月，熟悉 GMV 与转化率模型。本项目负责战略规划、融资对接、5 校地推协调。"},
            {"name": "李四", "detail": "计算机大三，ACM 校队成员。字节跳动产品实习 6 个月，独立负责过日活百万级 App 的功能迭代。本项目负责产品规划与用户增长，完成 50 份用户访谈、5 版 MVP 迭代。"},
            {"name": "王五", "detail": "软件工程大三，GitHub 1k star 开源项目作者。全栈开发 2 年，熟悉 Taro + Node.js + MySQL 技术栈。本项目负责技术架构与开发，独立完成 MVP 全栈实现。"},
        ],
        "advisor_bg": "李教授，副教授，工商管理学院，研究方向创业管理与商业模式创新，主持教育部人文社科项目 1 项，指导学生团队获'互联网+'省赛金奖 2 项、挑战杯国赛银奖 1 项。为本项目提供商业模式打磨与融资对接支持。",
    },
    "risk_analysis": [
        {"type": "市场风险", "risk": "闲鱼/转转进入校园", "prob": "中", "impact": "高", "measure": "强化校园社团合作、深耕单校、建立校园认证壁垒"},
        {"type": "技术风险", "risk": "高并发性能瓶颈", "prob": "低", "impact": "中", "measure": "腾讯云弹性扩容、Redis 缓存、读写分离"},
        {"type": "运营风险", "risk": "用户增长不及预期", "prob": "中", "impact": "高", "measure": "多渠道获客、降低 CAC、加强社群运营"},
        {"type": "财务风险", "risk": "现金流断裂", "prob": "低", "impact": "高", "measure": "控制 burn rate ≤ 1 万/月、申请大创延期经费、提前启动天使轮"},
    ],
    "personal_growth": {
        "role": "作为 CEO，我负责战略规划、融资对接、5 校地推协调、与指导教师的周例会汇报。",
        "transformation": "我从'只会做 PPT 的工商管理学生'变为'能独立完成 LTV/CAC 模型、跑通 5 校地推、与投资人对接的初创团队 CEO'。",
        "hard_skills": "我学会了市场测算（TAM/SAM/SOM 三级测算）、财务建模（3 年现金流表）、产品原型设计（Axure/Figma）、数据分析（SQL/Excel 数据透视）。",
        "soft_skills": "软技能方面提升了跨专业沟通（与技术同学对需求）、抗压能力（路演前 3 天产品崩掉重写）、决策能力（在 5 校 vs 10 校扩张间做出取舍）。",
        "failure": "在用户调研阶段，我曾因问卷设计偏差导致 30 份样本作废，复盘后改为深度访谈+问卷混合方法，重新拿到 50 份有效数据。这次失败让我学会'调研设计先行'，也让我意识到产品决策必须基于真实用户声音而非自己假设。",
    },
    "team_collaboration": {
        "division": "CEO 张三负责战略+融资；CTO 王五负责技术+架构；CPO 李四负责产品+用户；CMO 赵六负责市场+获客；设计孙七负责 UI+品牌。重大事项投票 3/5 通过，技术架构争议由 CTO 最终裁决。",
        "conflict": "曾就'先扩张 10 校 vs 深耕 5 校'产生分歧，CMO 主张快速扩张、CEO 主张深耕。通过用户调研数据（深耕 5 校留存率高 2 倍）+ 投票机制解决，决定深耕 5 校。",
        "tools": "协作工具：飞书（文档/日历/视频会议）+ Notion（产品需求）+ Figma（设计协作）+ 周一例会 + 日站会 15 分钟。",
    },
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="互联网+大学生创新创业大赛商业计划书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="示例：\n  python build.py --data data.json --out output.docx\n"
               "  python build.py --demo --out demo.docx\n\nJSON 字段定义详见 SKILL.md 第十一章。",
    )
    parser.add_argument("--data", type=str, default=None, help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True, help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true", help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
