#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
班集体 / 团支部集体申报书 docx 生成器

支持四类集体荣誉申报（通过 collective_type 字段切换）：
- advanced_class: 先进班集体申请书（综合 5 维度全面优秀）
- good_study_class: 优良学风班申请书（学风维度主导）
- may_fourth_flag: 五四红旗团支部申请书（团建维度主导）
- civilized_class: 文明班级申请书（文明维度主导）

面向集体而非个人，正文 3000~5000 字，6 段结构：
  班级概况 + 思想建设 + 学风建设 + 班级活动 + 班级荣誉 + 工作展望

按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 标题：黑体二号，居中
- 段落标题：黑体小三，加粗
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 表格：宋体五号，居中，表头加粗
- 落款：右对齐

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx
    python build.py --data data.json --out output.docx --word-count 5000

JSON 字段详见 SKILL.md 第六章信息采集清单与 JSON Schema。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号（主标题）
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三（段落标题）
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四（正文）
SIZE_WU = Pt(10.5)          # 五号（表格）
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

# 类别 -> 标题文本
COLLECTIVE_TYPE_TITLE_MAP = {
    "advanced_class": "先进班集体申请书",
    "good_study_class": "优良学风班申请书",
    "may_fourth_flag": "五四红旗团支部申请书",
    "civilized_class": "文明班级申请书",
}

# 5 维度权重表（按 collective_type）
DIMENSION_WEIGHT_MAP = {
    "advanced_class":    {"ideology": 20, "study": 25, "activity": 20, "discipline": 20, "hygiene": 15},
    "good_study_class":  {"ideology": 10, "study": 50, "activity": 15, "discipline": 15, "hygiene": 10},
    "may_fourth_flag":   {"ideology": 50, "study": 15, "activity": 20, "discipline": 10, "hygiene": 5},
    "civilized_class":   {"ideology": 10, "study": 10, "activity": 10, "discipline": 30, "hygiene": 40},
}

DIMENSION_LABELS = {
    "ideology": "思想建设",
    "study": "学习建设",
    "activity": "活动建设",
    "discipline": "纪律建设",
    "hygiene": "卫生建设",
}


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = bold
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), font_name)
            rFonts.set(qn("w:ascii"), font_name)
            rFonts.set(qn("w:hAnsi"), font_name)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_page_format(doc: Document) -> None:
    """设置 A4 纸张与页边距"""
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
        section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
        section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
        section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number_footer(doc: Document) -> None:
    """添加页脚页码（居中）"""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        set_run_font(run, FONT_SONG, SIZE_XIAO_WU)


def add_main_title(doc: Document, title: str) -> None:
    """添加主标题（黑体二号，居中）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(18)
    pf.line_spacing = 1.5
    run = p.add_run(title)
    set_run_font(run, FONT_HEI, SIZE_ER, bold=True)


def add_section_title(doc: Document, title: str) -> None:
    """添加段落标题（黑体小三，加粗）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    run = p.add_run(title)
    set_run_font(run, FONT_HEI, SIZE_XIAO_SAN, bold=True)


def add_body_paragraph(doc: Document, text: str, indent: bool = True) -> None:
    """添加正文段（宋体小四，1.5 倍行距，首行缩进 2 字符）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    if indent:
        # 首行缩进 2 字符（小四 12pt * 2 = 24pt）
        pf.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_run_font(run, FONT_SONG, SIZE_XIAO_SI)


def add_signature_paragraph(doc: Document, text: str) -> None:
    """添加落款段（右对齐，宋体小四）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, FONT_SONG, SIZE_XIAO_SI)


def add_blank_line(doc: Document, size: int = 1) -> None:
    """添加空行"""
    for _ in range(size):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)


def add_table_with_header(doc: Document, headers: List[str],
                          rows: List[List[str]], col_widths: Optional[List[float]] = None) -> None:
    """添加带表头的表格（宋体五号，居中，表头加粗）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_font(cell, FONT_SONG, SIZE_WU, bold=True)

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            set_cell_font(cell, FONT_SONG, SIZE_WU, bold=False)

    # 列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)


# ============================================================
# 数据校验
# ============================================================

REQUIRED_FIELDS = [
    "class_name", "college", "major", "class_size", "collective_type",
    "male_count", "female_count", "member_count", "party_member_count",
    "class_committee", "head_teacher",
    "ideology_score", "study_score", "activity_score", "discipline_score", "hygiene_score",
    "avg_gpa", "pass_rate", "competition_awards", "cet4_pass_rate", "cet6_pass_rate",
    "overview_text", "ideology_text", "study_text", "activity_text", "honor_text", "outlook_text",
]

SCORE_FIELDS = ["ideology_score", "study_score", "activity_score", "discipline_score", "hygiene_score"]

SCORE_MAX_MAP = {
    "ideology_score": 20,
    "study_score": 25,
    "activity_score": 20,
    "discipline_score": 20,
    "hygiene_score": 15,
}

VALID_COLLECTIVE_TYPES = list(COLLECTIVE_TYPE_TITLE_MAP.keys())
VALID_WORD_COUNTS = [3000, 4000, 5000]


def validate_data(data: Dict[str, Any]) -> Tuple[bool, List[str]]:
    """校验数据完整性，返回 (是否通过, 错误列表)"""
    errors: List[str] = []

    # 必填字段
    for field in REQUIRED_FIELDS:
        if field not in data:
            errors.append(f"缺失必填字段: {field}")
        elif data[field] is None or data[field] == "":
            errors.append(f"必填字段为空: {field}")

    # collective_type 校验
    ct = data.get("collective_type")
    if ct and ct not in VALID_COLLECTIVE_TYPES:
        errors.append(
            f"collective_type 非法: {ct}（必须为 "
            f"{'/'.join(VALID_COLLECTIVE_TYPES)}）"
        )

    # word_count_version 校验
    wcv = data.get("word_count_version", 4000)
    if wcv not in VALID_WORD_COUNTS:
        errors.append(
            f"word_count_version 非法: {wcv}（必须为 "
            f"{'/'.join(str(x) for x in VALID_WORD_COUNTS)}）"
        )

    # 人数一致性
    cs = data.get("class_size", 0)
    mc = data.get("male_count", 0)
    fc = data.get("female_count", 0)
    if cs and mc + fc != cs:
        errors.append(
            f"人数不一致: class_size({cs}) != male_count({mc}) + female_count({fc})"
        )

    # 团员 + 党员 ≤ 班级总人数（仅警告，因部分学校党员同时计入团员统计）
    mcount = data.get("member_count", 0)
    pcount = data.get("party_member_count", 0)
    if cs and mcount + pcount > cs:
        errors.append(
            f"⚠ 警告: 团员({mcount}) + 党员({pcount}) > 班级总人数({cs})（部分党员同时计入团员统计，请核实）"
        )

    # 5 维度分数校验
    for sf in SCORE_FIELDS:
        score_obj = data.get(sf)
        if score_obj and isinstance(score_obj, dict):
            total = score_obj.get("total", 0)
            max_val = score_obj.get("max", SCORE_MAX_MAP[sf])
            if max_val != SCORE_MAX_MAP[sf]:
                errors.append(
                    f"{sf}.max 应为 {SCORE_MAX_MAP[sf]}，实际为 {max_val}"
                )
            if total > max_val:
                errors.append(
                    f"{sf}.total({total}) 超过 max({max_val})"
                )
            # 子项校验
            sub_keys = [k for k in score_obj.keys() if k not in ("total", "max")]
            sub_sum = sum(score_obj.get(k, 0) for k in sub_keys)
            if abs(sub_sum - total) > 0.01:
                errors.append(
                    f"{sf}: 子项之和({sub_sum}) != total({total})"
                )

    # 硬门槛校验（警告，不阻塞）
    if cs and cs < 30:
        errors.append(
            f"⚠ 警告: 班级人数({cs}) < 30，未达四类集体荣誉硬门槛"
        )

    pr = data.get("pass_rate", 100)
    if pr < 95:
        errors.append(
            f"⚠ 警告: 课程及格率({pr}%) < 95%，未达先进班集体/五四红旗团支部硬门槛"
        )

    # 5 维度总分门槛
    total_score = sum(
        data.get(sf, {}).get("total", 0) for sf in SCORE_FIELDS
    )
    if total_score and total_score < 85:
        errors.append(
            f"⚠ 警告: 5 维度总分({total_score}) < 85，未达申报门槛"
        )

    return (len([e for e in errors if not e.startswith("⚠")]) == 0, errors)


# ============================================================
# 5 维度量化数据计算
# ============================================================

def compute_total_score(data: Dict[str, Any]) -> int:
    """计算 5 维度总分"""
    return sum(data.get(sf, {}).get("total", 0) for sf in SCORE_FIELDS)


def compute_max_score() -> int:
    """5 维度满分"""
    return sum(SCORE_MAX_MAP.values())


def get_dimension_score(data: Dict[str, Any], dim: str) -> Tuple[int, int]:
    """获取指定维度的 (得分, 满分)"""
    score_field = f"{dim}_score"
    if score_field in data:
        score_obj = data[score_field]
        return (score_obj.get("total", 0), score_obj.get("max", SCORE_MAX_MAP[score_field]))
    return (0, SCORE_MAX_MAP[score_field])


# ============================================================
# 内容渲染
# ============================================================

def render_cover_page(doc: Document, data: Dict[str, Any]) -> None:
    """渲染封面"""
    ct = data["collective_type"]
    title_text = COLLECTIVE_TYPE_TITLE_MAP[ct]

    add_blank_line(doc, 4)
    add_main_title(doc, title_text)
    add_blank_line(doc, 2)

    # 申报信息
    info_lines = [
        f"申报班级：{data['class_name']}",
        f"所属学院：{data['college']}",
        f"专业：{data['major']}",
        f"班级人数：{data['class_size']} 人",
        f"班主任：{data['head_teacher']['name']}（{data['head_teacher']['title']}）",
        f"申报日期：{data.get('submit_date', '2025 年 5 月')}    （可调整）",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing = 1.75
        pf.space_after = Pt(6)
        run = p.add_run(line)
        set_run_font(run, FONT_SONG, SIZE_SI)

    # 班训
    if data.get("class_slogan"):
        add_blank_line(doc, 2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing = 1.75
        run = p.add_run(f"班级口号：{data['class_slogan']}")
        set_run_font(run, FONT_HEI, SIZE_XIAO_SAN, bold=True)

    # 分页
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()
    from docx.enum.text import WD_BREAK
    p2 = doc.add_paragraph()
    run2 = p2.add_run()
    run2.add_break(WD_BREAK.PAGE)


def render_overview_section(doc: Document, data: Dict[str, Any]) -> None:
    """§1 班级概况"""
    add_section_title(doc, "一、班级概况")
    add_body_paragraph(doc, data["overview_text"])

    # 班委建制表
    add_blank_line(doc, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run("【班委建制表】")
    set_run_font(run, FONT_HEI, SIZE_SI, bold=True)

    committee = data.get("class_committee", [])
    if committee:
        headers = ["序号", "职务", "姓名", "任期", "主要工作"]
        rows = []
        for i, member in enumerate(committee, 1):
            rows.append([
                str(i),
                member.get("post", ""),
                member.get("name", ""),
                member.get("term", ""),
                member.get("duty", "—"),
            ])
        add_table_with_header(doc, headers, rows, col_widths=[1.2, 2.5, 2.0, 3.5, 5.5])


def render_ideology_section(doc: Document, data: Dict[str, Any]) -> None:
    """§2 思想建设"""
    add_section_title(doc, "二、思想建设")
    add_body_paragraph(doc, data["ideology_text"])

    # 五四红旗团支部专项：补充团建数据
    if data["collective_type"] == "may_fourth_flag":
        add_body_paragraph(
            doc,
            f"本学年团支部团费缴纳率达 100%，智慧团建系统团员信息录入率 100%，"
            f"组织关系转接完成率 100%。共开展主题团日 8 次（覆盖'党史学习''国家安全'"
            f"'五四精神''志愿服务'等主题），支部大会 4 次、支委会 12 次、团小组会 8 次、"
            f"团员教育评议 1 次、团籍注册 1 次、团课 6 次，符合三会两制一课要求。"
            f"本学年新发展团员 2 人、推优入党 4 人，入党积极分子 8 人、发展对象 3 人、"
            f"预备党员 4 人、正式党员 0 人。"
        )


def render_study_section(doc: Document, data: Dict[str, Any]) -> None:
    """§3 学风建设"""
    add_section_title(doc, "三、学风建设")
    add_body_paragraph(doc, data["study_text"])

    # 学习成绩汇总表
    add_blank_line(doc, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run("【学习成绩汇总表】")
    set_run_font(run, FONT_HEI, SIZE_SI, bold=True)

    awards = data.get("competition_awards", {})
    headers = ["指标", "数值", "指标", "数值"]
    rows = [
        ["班级平均学分绩点", f"{data['avg_gpa']:.2f}/4.0", "课程及格率", f"{data['pass_rate']:.1f}%"],
        ["英语四级通过率", f"{data['cet4_pass_rate']:.1f}%", "英语六级通过率", f"{data['cet6_pass_rate']:.1f}%"],
        ["国家级竞赛获奖", f"{awards.get('national', 0)} 人次", "省级竞赛获奖", f"{awards.get('provincial', 0)} 人次"],
        ["校级竞赛获奖", f"{awards.get('school', 0)} 人次", "团员比例", f"{data['member_count'] / data['class_size'] * 100:.1f}%"],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[4.0, 3.5, 4.0, 3.5])


def render_activity_section(doc: Document, data: Dict[str, Any]) -> None:
    """§4 班级活动"""
    add_section_title(doc, "四、班级活动")
    add_body_paragraph(doc, data["activity_text"])


def render_honor_section(doc: Document, data: Dict[str, Any]) -> None:
    """§5 班级荣誉"""
    add_section_title(doc, "五、班级荣誉")
    add_body_paragraph(doc, data["honor_text"])

    # 集体荣誉表
    add_blank_line(doc, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run("【集体荣誉表】")
    set_run_font(run, FONT_HEI, SIZE_SI, bold=True)

    collective_honors = data.get("collective_honors", [])
    if collective_honors:
        headers = ["序号", "时间", "荣誉名称", "级别"]
        rows = []
        for i, h in enumerate(collective_honors, 1):
            rows.append([
                str(i),
                h.get("year", ""),
                h.get("name", ""),
                h.get("level", ""),
            ])
        add_table_with_header(doc, headers, rows, col_widths=[1.2, 2.5, 8.5, 2.5])
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("（暂无集体荣誉记录）")
        set_run_font(run, FONT_SONG, SIZE_WU)

    # 个人荣誉汇总表
    add_blank_line(doc, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run("【个人荣誉汇总表】")
    set_run_font(run, FONT_HEI, SIZE_SI, bold=True)

    phs = data.get("personal_honors_summary", {})
    headers = ["级别", "获奖人数", "累计人次"]
    rows = [
        ["国家级", str(phs.get("national", 0)), str(max(phs.get("national", 0), 0))],
        ["省级", str(phs.get("provincial", 0)), str(max(phs.get("provincial", 0), 0))],
        ["校级", str(phs.get("school", 0)), str(max(phs.get("school", 0), 0))],
        ["合计", str(phs.get("total_count", 0)), str(phs.get("total_count", 0))],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[4.0, 5.0, 5.0])

    # 5 维度量化自评表
    add_blank_line(doc, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run("【5 维度量化自评表】")
    set_run_font(run, FONT_HEI, SIZE_SI, bold=True)

    headers = ["维度", "满分", "自评得分", "达标率"]
    rows = []
    for dim_key, dim_label in DIMENSION_LABELS.items():
        score, max_v = get_dimension_score(data, dim_key)
        rate = f"{score / max_v * 100:.1f}%" if max_v else "0%"
        rows.append([dim_label, str(max_v), str(score), rate])
    total = compute_total_score(data)
    max_total = compute_max_score()
    rows.append(["合计", str(max_total), str(total), f"{total / max_total * 100:.1f}%"])
    add_table_with_header(doc, headers, rows, col_widths=[4.0, 3.0, 3.5, 3.5])


def render_outlook_section(doc: Document, data: Dict[str, Any]) -> None:
    """§6 工作展望"""
    add_section_title(doc, "六、工作展望")
    add_body_paragraph(doc, data["outlook_text"])


def render_signature(doc: Document, data: Dict[str, Any]) -> None:
    """渲染落款（右对齐）"""
    add_blank_line(doc, 2)
    sigs = [
        "                                    全体班委",
        f"                                    班  长：________（签字）",
        f"                                    团支书：________（签字）",
        f"                                    班主任：{data['head_teacher']['name']}（签字）",
        "                                    学  院：________（盖章）",
        f"                                    {data.get('submit_date', '2025 年 5 月 20 日')}",
    ]
    for sig in sigs:
        add_signature_paragraph(doc, sig)


# ============================================================
# 文档构建
# ============================================================

def build_document(data: Dict[str, Any], output_path: str) -> Tuple[bool, List[str]]:
    """构建 docx 文档，返回 (是否成功, 错误/警告列表)"""
    ok, errors = validate_data(data)
    if not ok:
        return (False, errors)

    doc = Document()
    set_page_format(doc)
    add_page_number_footer(doc)

    # 封面
    render_cover_page(doc, data)

    # 正文 6 段
    render_overview_section(doc, data)
    render_ideology_section(doc, data)
    render_study_section(doc, data)
    render_activity_section(doc, data)
    render_honor_section(doc, data)
    render_outlook_section(doc, data)

    # 落款
    render_signature(doc, data)

    # 保存
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return (True, errors)


# ============================================================
# DEFAULT_DATA（完整 demo）
# ============================================================

DEFAULT_DATA: Dict[str, Any] = {
    "class_name": "计算机科学与技术 2023 级 1 班",
    "college": "计算机学院",
    "major": "计算机科学与技术",
    "class_size": 38,
    "collective_type": "advanced_class",
    "male_count": 26,
    "female_count": 12,
    "member_count": 36,
    "party_member_count": 4,
    "class_committee": [
        {"post": "班长", "name": "张三", "term": "2024.09-2025.06", "duty": "全面负责班级工作"},
        {"post": "副班长", "name": "王五", "term": "2024.09-2025.06", "duty": "协助班长，分管学习"},
        {"post": "团支书", "name": "李四", "term": "2024.09-2025.06", "duty": "主持团支部工作"},
        {"post": "学习委员", "name": "赵六", "term": "2024.09-2025.06", "duty": "组织学习互助"},
        {"post": "生活委员", "name": "钱七", "term": "2024.09-2025.06", "duty": "班费管理、宿舍卫生"},
        {"post": "文体委员", "name": "孙八", "term": "2024.09-2025.06", "duty": "文体活动组织"},
        {"post": "心理委员", "name": "周九", "term": "2024.09-2025.06", "duty": "心理健康教育"},
        {"post": "组织委员", "name": "吴十", "term": "2024.09-2025.06", "duty": "团组织生活"},
        {"post": "宣传委员", "name": "郑十一", "term": "2024.09-2025.06", "duty": "团支部宣传"},
    ],
    "head_teacher": {
        "name": "王老师",
        "title": "副教授",
        "unit": "计算机学院软件工程系",
    },
    "ideology_score": {
        "party_ratio": 5,
        "member_ratio": 3,
        "youth_study": 3,
        "theme_meeting": 4,
        "ideology_grade": 3,
        "total": 18,
        "max": 20,
    },
    "study_score": {
        "gpa": 6,
        "top10": 5,
        "pass_rate": 4,
        "competition": 4,
        "cet": 3,
        "total": 22,
        "max": 25,
    },
    "activity_score": {
        "collective_activity": 4,
        "school_participation": 3,
        "volunteer": 3,
        "culture_award": 4,
        "social_practice": 3,
        "total": 17,
        "max": 20,
    },
    "discipline_score": {
        "attendance": 5,
        "late": 2,
        "violation": 5,
        "cheating": 4,
        "network": 2,
        "total": 18,
        "max": 20,
    },
    "hygiene_score": {
        "excellent_rate": 5,
        "civilized_dormitory": 3,
        "pass_rate": 3,
        "safety": 2,
        "total": 13,
        "max": 15,
    },
    "avg_gpa": 3.42,
    "pass_rate": 98.5,
    "competition_awards": {"national": 3, "provincial": 5, "school": 12},
    "cet4_pass_rate": 88.5,
    "cet6_pass_rate": 65.8,
    "overview_text": (
        "计算机学院计算机科学与技术 2023 级 1 班，成立于 2023 年 9 月，"
        "现有学生 38 人，其中男生 26 人、女生 12 人，团员 36 人（占比 94.7%）、"
        "党员（含预备）4 人（占比 10.5%）、入党积极分子 8 人。班主任王老师"
        "（计算机学院软件工程系副教授），辅导员张老师。班级建制完整，设班长 1 人、"
        "副班长 1 人、团支书 1 人、学习委员 1 人、生活委员 1 人、文体委员 1 人、"
        "心理委员 1 人、组织委员 1 人、宣传委员 1 人，共 9 人组成班委团队。"
        "班级口号\"团结奋进，求实创新\"源自校训\"明德至善、博学笃行\"，"
        "体现\"团结、勤奋、求实、创新\"的班风精神。本班以\"建设 5 维度全面优秀的"
        "先进班集体\"为目标，秉持\"勤学、守纪、团结、创新\"的班风，本学年在思想、"
        "学习、活动、纪律、卫生 5 维度全面推进班级建设，5 维度量化自评 88 分，"
        "达到先进班集体申报门槛（≥85 分）。"
    ),
    "ideology_text": (
        "本班高度重视思想建设，构建\"党建带团建、团建带班级\"的思想引领体系。"
        "本学年共有团员 36 人（占比 94.7%）、党员（含预备）4 人（占比 10.5%）、"
        "入党积极分子 8 人，本学年新发展团员 2 人、推优入党 4 人。班级共开展"
        "主题班会 8 次、主题团日 8 次，覆盖\"党史学习\"\"国家安全\"\"心理健康\""
        "\"职业规划\"等主题，平均出勤率 96%。青年大学习完成率学年累计达 96.5%，"
        "连续 28 期位列学院前 3。本班思政课（马克思主义基本原理、毛泽东思想和"
        "中国特色社会主义理论体系概论、形势与政策）平均成绩 86 分，位列专业第 2。"
        "班级开展\"红色经典诵读\"分享会 3 次、参观本地革命纪念馆 2 次、观看"
        "\"长津湖\"\"我和我的祖国\"等爱国主义影片 4 部，引导学生树立正确的世界观、"
        "人生观、价值观。在班主任王老师指导下，班级建立\"党员联系宿舍\"制度，"
        "4 名党员分别联系 8 个宿舍，定期开展谈心谈话，及时了解同学思想动态，"
        "本学年共开展谈心谈话 32 人次。"
    ),
    "study_text": (
        "本班将学风建设作为班级建设的核心，构建\"互助 + 自律 + 督导\"三位一体"
        "学习体系。本学年班级平均学分绩点 3.42/4.0，位列专业第 1，较上学年提升 0.18；"
        "GPA ≥3.5 的同学 8 人（占比 21.1%），GPA ≥3.0 的同学 32 人（占比 84.2%），"
        "专业前 10% 共 5 人（专业共 76 人），位列专业前列。本学年所有课程及格率 98.5%，"
        "核心课程（高等数学、大学英语、数据结构、操作系统）及格率 100%，无重修。"
        "本学年班级同学在\"挑战杯\"\"互联网+\"\"数学建模\"\"英语竞赛\"等学科竞赛中获"
        "国家级 3 项、省级 5 项、校级 12 项，累计 20 人次获奖。英语四级通过率 88.5%，"
        "六级通过率 65.8%，计算机二级通过率 92.1%。班级建立\"学习互助小组\" 6 个，"
        "开展\"一帮一\"结对 12 对，组织学习经验分享会 5 次，编制复习资料 8 套。"
        "针对 2 名学业预警学生，建立\"班主任 + 学习委员 + 党员\"三级帮扶机制，"
        "本学年 2 名预警学生均成功脱警，GPA 提升 0.5 以上。学习委员组织\"每日学习打卡\""
        "活动，连续 180 天打卡率 92%，营造了浓厚的学习氛围。"
    ),
    "activity_text": (
        "本班将班级活动作为凝聚班级向心力、提升同学综合素质的重要载体，构建"
        "\"常规活动 + 主题特色活动 + 志愿服务 + 社会实践\"四维活动体系。本学年共组织"
        "班级集体活动 9 次，包括\"迎新晚会\"\"冬至包饺子\"\"春季踏青\"\"班级生日会\""
        "\"读书分享会\"\"心理素质拓展\"\"趣味运动会\"\"毕业季送别\"\"元旦联欢\"等，"
        "参与率平均 92%。班级组队参加校第 38 届运动会，获男子 4×100 米接力第 2 名、"
        "女子跳远第 3 名；参加\"校园十佳歌手\"比赛，李四同学获优胜奖；参加院篮球联赛，"
        "获第 2 名。全班同学注册志愿者 38 人（占比 100%），本学年累计志愿服务时长"
        "1080 小时，人均 28.4 小时，参与\"阳光助残\"\"乡村支教\"\"社区敬老\"\"无偿献血\""
        "等项目，其中\"乡村支教\"项目获校志愿服务大赛铜奖。寒暑假社会实践参与率 84%，"
        "组建实践团队 3 支，赴云南怒江、贵州黔东南、河南兰考等地开展调研，撰写调研报告"
        "5 份，获校级优秀团队 1 项。班级打造\"红色经典诵读\"特色活动品牌，已连续开展"
        "3 期，获学院官网报道 2 次、校团委公众号报道 1 次。"
    ),
    "honor_text": (
        "本学年班级获校级荣誉 3 项、院级荣誉 2 项，包括\"2024 年度校级先进班集体\""
        "\"2024 年度院级优秀团支部\"\"2025 年春季校级文明宿舍楼\"等。班级同学获国家级"
        "荣誉 3 项（含国家奖学金 1 人、\"挑战杯\"国家级二等奖 2 人次）、省级 5 项"
        "（含省级\"优秀共青团员\"1 人、\"互联网+\"省级银奖 4 人次）、校级 18 项"
        "（含校级奖学金 8 人、校级竞赛 6 人次、校级优秀团员 4 人），累计 26 人次获奖，"
        "占班级总人数的 68.4%。其中，张三同学作为学生代表在 2024 年校学生代表大会发言，"
        "李四同学获 2024 年度省级\"优秀共青团员\"称号，王五同学团队获\"挑战杯\"国家级"
        "二等奖（专业首个国家级竞赛奖项）。依据本 skill 5 维度量化体系，本班自评如下："
        "思想建设 18/20、学习建设 22/25、活动建设 17/20、纪律建设 18/20、卫生建设 13/15，"
        "合计 88/100，达到先进班集体申报门槛（≥85 分），任一维度均达 70% 满分以上，"
        "无严重短板。"
    ),
    "outlook_text": (
        "回顾本学年，班级建设仍存在以下不足：① 学习互助覆盖面不够广，仍有 2 名同学"
        "GPA 在 3.0 以下；② 班级特色活动品牌尚未成型，\"红色经典诵读\"活动影响力有限；"
        "③ 团日活动形式较为单一，缺乏沉浸式 / 实践式创新；④ 宿舍卫生优秀率有提升空间，"
        "仍有 30% 宿舍未达优秀标准；⑤ 班级与校友、企业联系不足，职业规划资源有限。"
        "下一学年，班级将围绕\"建设 5 维度全优的省级先进班集体\"建设目标，重点推进："
        "① 学习维度冲刺 GPA 3.5，2 名后进同学 GPA 提升 0.3；② 打造\"红色经典诵读\""
        "特色活动品牌，扩大至全院范围；③ 创新团日活动形式，引入沉浸式 / 实践式活动 4 次；"
        "④ 宿舍卫生优秀率提升至 80%，新增文明宿舍 2 个；⑤ 建立\"校友导师\"制度，"
        "邀请 3~5 名校友回校交流。具体措施：① 建立\"学业预警 + 一帮一\"长效机制；"
        "② 设立\"班级活动基金\"，专项支持特色活动；③ 落实\"团日活动考核\"制度，"
        "纳入团员评议；④ 推行\"宿舍长负责制 + 月评比 + 季奖励\"；⑤ 建立\"校友联络员\""
        "制度，定期开展校友活动。本班全体师生将以本次申报为新的起点，珍惜荣誉、戒骄戒躁，"
        "争取在下一学年取得更大进步，恳请评审委员会予以批准。"
    ),
    "collective_honors": [
        {"year": "2024.10", "name": "校级先进班集体", "level": "校级"},
        {"year": "2024.12", "name": "院级优秀团支部", "level": "院级"},
        {"year": "2025.04", "name": "校级文明宿舍楼", "level": "校级"},
        {"year": "2025.05", "name": "校级活力团支部", "level": "校级"},
    ],
    "personal_honors_summary": {
        "national": 3,
        "provincial": 5,
        "school": 18,
        "total_count": 26,
    },
    "word_count_version": 4000,
    "class_slogan": "团结奋进，求实创新",
    "target_honor": "2024-2025 学年校级先进班集体",
    "submit_date": "2025 年 5 月 20 日",
    "defense_needed": True,
}


# ============================================================
# 主入口
# ============================================================

def parse_args() -> argparse.Namespace:
    """解析命令行参数"""
    parser = argparse.ArgumentParser(
        description="班集体 / 团支部集体申报书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  python build.py --demo --out demo.docx
  python build.py --data data.json --out output.docx
  python build.py --data data.json --out output.docx --word-count 5000
        """,
    )
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--demo", action="store_true", help="使用内置 demo 数据生成示例文档")
    group.add_argument("--data", type=str, help="用户数据 JSON 文件路径")
    parser.add_argument("--out", type=str, required=True, help="输出 docx 文件路径")
    parser.add_argument(
        "--word-count", type=int, choices=VALID_WORD_COUNTS,
        help="字数版本（3000/4000/5000），覆盖 JSON 中的 word_count_version",
    )
    return parser.parse_args()


def load_data(data_path: str) -> Dict[str, Any]:
    """加载 JSON 数据"""
    with open(data_path, "r", encoding="utf-8") as f:
        return json.load(f)


def main() -> int:
    """主函数"""
    args = parse_args()

    # 加载数据
    if args.demo:
        data = DEFAULT_DATA.copy()
        # 深拷贝嵌套结构
        data = json.loads(json.dumps(DEFAULT_DATA))
    else:
        try:
            data = load_data(args.data)
        except FileNotFoundError:
            print(f"❌ 错误: 数据文件不存在: {args.data}", file=sys.stderr)
            return 1
        except json.JSONDecodeError as e:
            print(f"❌ 错误: JSON 解析失败: {e}", file=sys.stderr)
            return 1

    # 覆盖字数版本
    if args.word_count:
        data["word_count_version"] = args.word_count

    # 构建文档
    ok, messages = build_document(data, args.out)

    # 输出校验信息
    if messages:
        for msg in messages:
            print(msg, file=sys.stderr)

    if not ok:
        print(f"❌ 数据校验失败，未生成文档", file=sys.stderr)
        return 1

    # 验证文档可打开
    try:
        from docx import Document as Doc2
        doc = Doc2(args.out)
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        print(f"✅ 文档生成成功: {args.out}")
        print(f"   段落数: {para_count}")
        print(f"   表格数: {table_count}")
        print(f"   申报类别: {COLLECTIVE_TYPE_TITLE_MAP.get(data['collective_type'], '未知')}")
        print(f"   班级名称: {data['class_name']}")
        print(f"   字数版本: {data.get('word_count_version', 4000)} 字")
        total_score = compute_total_score(data)
        max_total = compute_max_score()
        print(f"   5 维度量化: {total_score}/{max_total} ({total_score / max_total * 100:.1f}%)")
    except Exception as e:
        print(f"⚠ 警告: 文档生成但验证失败: {e}", file=sys.stderr)
        return 2

    return 0


if __name__ == "__main__":
    sys.exit(main())
