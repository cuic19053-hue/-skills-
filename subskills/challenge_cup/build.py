#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
挑战杯大学生课外学术科技作品竞赛申报书 docx 生成器 (v2.1)

v2.1 新增字段（基于案例 1+2 提炼）：
- policy_citations: 国家政策引用（8+）
- scientific_challenges: 科学挑战 3 段
- literature_review: 国内外研究现状综述（30+ 文献）
- algorithm_comparison: 算法/方法对比表
- tech_roadmap: 技术路线图（3 张）
- formulas: 数学公式（自然/发明类必备）
- economic_benefits: 社会经济效益量化（10 项）
- defense_ppt: 答辩 PPT 框架（10 页）
- work_type: 三大类代码（与 category 等价）
- DEFAULT_DATA: 消防无人机主题 demo（对齐案例）

3 档字数版本：
- natural_science: 8 千字档
- social_science: 1.5 万字档
- tech_invention: 1 万字档

格式标准：A4，页边距上下 2.54cm 左右 2.5cm；正文宋体小四 1.5 倍行距首行缩进 2 字符；
一级标题黑体三号居中；二级标题黑体小三左对齐；三级标题宋体四号加粗；表格宋体五号居中；
参考文献宋体五号；创新点表格 4 列（序号/类型/创新点描述/对比优势）。

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段定义详见 SKILL.md 第十九章。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)
SIZE_XIAO_ER = Pt(18)
SIZE_SAN = Pt(16)
SIZE_XIAO_SAN = Pt(15)
SIZE_SI = Pt(14)
SIZE_XIAO_SI = Pt(12)
SIZE_WU = Pt(10.5)
SIZE_XIAO_WU = Pt(9)

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

# 三大类标签
CATEGORY_LABEL = {
    "natural_science": "自然科学类学术论文",
    "social_science": "哲学社会科学类社会调查报告和学术论文",
    "tech_invention": "科技发明制作",
}
APPLICANT_LABEL = {"individual": "个人项目", "collective": "集体项目"}

# 3 档字数版本（v2.1）
WORD_LIMITS = {
    "natural_science": {
        "abstract_max": 500,
        "background": (1200, 1500),
        "research_method": (1500, 2000),
        "results_discussion": (2500, 3000),
        "conclusion": (400, 500),
        "total": 8000,
        "label": "8 千字档",
    },
    "social_science": {
        "abstract_max": 500,
        "background": (2000, 2500),
        "research_method": (3500, 4500),
        "results_discussion": (3500, 4500),
        "conclusion": (500, 600),
        "total": 15000,
        "label": "1.5 万字档",
    },
    "tech_invention": {
        "abstract_max": 500,
        "background": (1500, 2000),
        "research_method": (2500, 3000),
        "results_discussion": (2000, 2500),
        "conclusion": (400, 500),
        "total": 10000,
        "label": "1 万字档",
    },
}


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
        space_before=12, space_after=12,
    )


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_HEI, font_size=SIZE_XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=6, space_after=6,
    )


def add_heading_level3(doc, text: str):
    """三级标题：宋体四号加粗，左对齐"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False,
        space_before=4, space_after=4,
    )


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        first_line_indent=indent,
        line_spacing=1.5,
    )


def add_reference_paragraph(doc, text: str):
    """参考文献条目：宋体五号，单倍行距，悬挂缩进"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    pf.left_indent = Pt(20)
    pf.first_line_indent = Pt(-20)
    run = p.add_run(text)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)
    return p


def add_formula_paragraph(doc, text: str):
    """公式段落：Times New Roman 小四，居中"""
    return add_paragraph_with_format(
        doc, text,
        font_name=FONT_TIMES, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=False,
        line_spacing=1.5,
        space_before=4, space_after=4,
    )


def add_page_number(doc):
    """添加页脚居中页码"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def setup_page(doc):
    """设置 A4 页面与页边距"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def count_chinese_chars(text: str) -> int:
    """统计中文字符数（含全角标点）"""
    return sum(1 for c in text if '\u4e00' <= c <= '\u9fff' or c in "，。、；：？！""''（）【】《》")


def join_paragraphs(paragraphs: list) -> str:
    return "".join(paragraphs) if paragraphs else ""


def add_table_from_data(doc, headers, rows, col_widths=None, first_col_bold=False):
    """添加表格，可控制列宽与首列加粗"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, font_name=FONT_HEI, font_size=SIZE_WU, bold=True)

    for r_idx, row in enumerate(rows, 1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            bold = first_col_bold and c_idx == 0
            align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == len(row) - 1 and len(str(val)) > 20 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cell, val, font_name=FONT_SONG, font_size=SIZE_WU,
                          bold=bold, alignment=align)

    return table


def format_innovation_row(idx: int, inv) -> list:
    if isinstance(inv, dict):
        return [
            str(idx),
            inv.get("type", ""),
            inv.get("description", ""),
            inv.get("advantage", ""),
        ]
    return [str(idx), "", str(inv), ""]


def format_reference(idx: int, ref: dict) -> str:
    """格式化参考文献条目（GB/T 7714）"""
    if not isinstance(ref, dict):
        return f"[{idx}] {ref}"
    authors, title = ref.get("authors", ""), ref.get("title", "")
    ref_type = ref.get("type", "J")
    source, year = ref.get("source", ""), ref.get("year", "")
    volume, issue = ref.get("volume", ""), ref.get("issue", "")
    pages, publisher = ref.get("pages", ""), ref.get("publisher", "")
    city = ref.get("city", "")

    if ref_type in ("J", "C"):
        vol_iss = f"{volume}({issue})" if volume and issue else (volume or issue)
        page_part = f": {pages}" if pages else ""
        return f"[{idx}] {authors}. {title}[{ref_type}]. {source}, {year}{', ' + vol_iss if vol_iss else ''}{page_part}."
    elif ref_type == "M":
        page_part = f": {pages}" if pages else ""
        return f"[{idx}] {authors}. {title}[M]. {city}: {publisher}, {year}{page_part}."
    elif ref_type == "D":
        return f"[{idx}] {authors}. {title}[D]. {city}: {source}, {year}."
    elif ref_type == "S":
        return f"[{idx}] {authors}. {title}: {source}[S]. {city}: {publisher}, {year}."
    elif ref_type == "EB/OL":
        url = ref.get("url", "")
        access = ref.get("access_date", "")
        return f"[{idx}] {authors}. {title}[EB/OL]. ({year})[{access}]. {url}."
    else:
        return f"[{idx}] {authors}. {title}[{ref_type}]. {source}, {year}."


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """挑战杯课外学术科技作品竞赛申报书 docx 构建器（v2.1）"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, *keys, default=None):
        """安全取嵌套 list 字段"""
        if default is None:
            default = []
        v = self._get(*keys, default=default)
        if isinstance(v, str):
            return [v]
        if isinstance(v, list):
            return v
        return default

    def _work_type(self) -> str:
        """获取 work_type（与 category 等价）"""
        wt = self._get("work_type") or self._get("category", default="tech_invention")
        return wt if wt in CATEGORY_LABEL else "tech_invention"

    def _word_limits(self) -> dict:
        return WORD_LIMITS.get(self._work_type(), WORD_LIMITS["tech_invention"])

    def add_h1(self, text):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text):
        return add_heading_level2(self.doc, text)

    def add_h3(self, text):
        return add_heading_level3(self.doc, text)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_formula(self, text):
        return add_formula_paragraph(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, first_col_bold=False):
        return add_table_from_data(self.doc, headers, rows, col_widths, first_col_bold)

    def add_page_break(self):
        add_page_break(self.doc)

    # --------------------------------------------------------
    # 封面
    # --------------------------------------------------------

    def _add_cover(self):
        """封面：黑体二号标题 + 副标题 + 5 行下划线信息"""
        for _ in range(2):
            self.doc.add_paragraph()

        title = "第十八届" + "「挑战杯」大学生课外学术科技作品竞赛申报书"
        add_paragraph_with_format(
            self.doc, title,
            font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_before=12, space_after=12,
        )

        subtitle = "（课外学术科技作品 · v2.1）"
        add_paragraph_with_format(
            self.doc, subtitle,
            font_name=FONT_HEI, font_size=SIZE_SAN, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent=False,
            space_after=24,
        )

        for _ in range(2):
            self.doc.add_paragraph()

        category_label = CATEGORY_LABEL.get(self._work_type(), "科技发明制作")
        applicant_label = APPLICANT_LABEL.get(self._get("applicant_type"), "集体项目")

        info_items = [
            ("作品全称", self._get("work_full_name")),
            ("作品类别", category_label),
            ("申报者代表", self._get("leader_name")),
            ("所在学校", self._get("school")),
            ("申报日期", self._get("apply_date")),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI,
                         font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True

        self.add_page_break()

    # --------------------------------------------------------
    # 一、作品全称 + 二、作品类别
    # --------------------------------------------------------

    def _add_work_info(self):
        self.add_h1("一、作品全称")
        work_name = self._get("work_full_name")
        if work_name:
            add_paragraph_with_format(
                self.doc, work_name,
                font_name=FONT_HEI, font_size=SIZE_XIAO_ER, bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_indent=False,
                space_before=6, space_after=6,
            )
            if len(work_name) > 30:
                self.warnings.append(f"作品全称 {len(work_name)} 字超过 30 字上限")
        else:
            self.add_para("（请填写作品全称，不超过 30 字，突出『做什么+为谁做』，禁三层『基于...的...』堆砌。）")

        self.add_h1("二、作品类别")
        category = self._work_type()
        category_options = [
            ("natural_science", "自然科学类学术论文"),
            ("social_science", "哲学社会科学类社会调查报告和学术论文"),
            ("tech_invention", "科技发明制作"),
        ]
        for key, label in category_options:
            mark = "☑" if category == key else "☐"
            add_paragraph_with_format(
                self.doc, f"{mark} {label}",
                font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                first_line_indent=False,
                line_spacing=1.5,
            )

        # v2.1 字数档说明
        limits = self._word_limits()
        self.add_h2(f"（字数档：{limits['label']}，正文约 {limits['total']} 字）")

    # --------------------------------------------------------
    # 三、申报者情况
    # --------------------------------------------------------

    def _add_applicant_info(self):
        self.add_h1("三、申报者情况")
        applicant_type = self._get("applicant_type", default="collective")
        applicant_label = APPLICANT_LABEL.get(applicant_type, "集体项目")

        self.add_h2(f"（{applicant_label}）")

        leader_rows = [
            ["姓名", self._get("leader_name")],
            ["学号", self._get("leader_id")],
            ["性别", self._get("leader_gender")],
            ["年龄", self._get("leader_age")],
            ["专业", self._get("leader_major")],
            ["年级", self._get("leader_grade")],
            ["学院", self._get("college")],
            ["学校", self._get("school")],
            ["联系电话", self._get("leader_phone")],
        ]
        self.add_table(["字段", "内容"], leader_rows, col_widths=[4.0, 12.0],
                       first_col_bold=True)

        if applicant_type == "collective":
            self.add_h2("（集体项目团队成员）")
            members = self._get_list("team_members")
            if members:
                rows = []
                for i, m in enumerate(members, 1):
                    if isinstance(m, dict):
                        rows.append([
                            str(i),
                            m.get("name", ""),
                            m.get("id", ""),
                            m.get("major", ""),
                            m.get("grade", ""),
                            m.get("role", ""),
                        ])
                    else:
                        rows.append([str(i), str(m), "", "", "", ""])
                self.add_table(
                    ["序号", "姓名", "学号", "专业", "年级", "分工"],
                    rows,
                    col_widths=[1.5, 2.0, 2.5, 3.0, 2.0, 5.0],
                )
            else:
                self.add_para("（请填写集体项目团队成员 2~7 人，每人含姓名/学号/专业/年级/分工，禁只列姓名不写分工。）")

        self.add_h2("（指导教师信息）")
        advisor_rows = [
            ["姓名", self._get("advisor_name")],
            ["职称", self._get("advisor_title")],
            ["研究方向", self._get("advisor_research")],
            ["联系电话", self._get("advisor_phone")],
        ]
        self.add_table(["字段", "内容"], advisor_rows, col_widths=[4.0, 12.0],
                       first_col_bold=True)

    # --------------------------------------------------------
    # 四、作品简介
    # --------------------------------------------------------

    def _add_abstract(self):
        self.add_h1("四、作品简介")
        abstract = self._get("abstract")
        if abstract:
            self.add_para(abstract)
            char_count = count_chinese_chars(abstract)
            if char_count > 500:
                self.warnings.append(f"作品简介 {char_count} 字超过 500 字上限")
        else:
            self.add_para("（请填写作品简介，≤500 字，按 4 句结构：痛点+做什么 / 怎么做+量化目标 / 产出什么 / 研究现状。）")

    # --------------------------------------------------------
    # 五、选题背景与意义
    # --------------------------------------------------------

    def _add_background(self):
        self.add_h1("五、选题背景与意义")
        background = self._get_list("background")
        limits = self._word_limits()
        bg_range = limits["background"]
        if background:
            for para in background:
                self.add_para(para)
            total = count_chinese_chars(join_paragraphs(background))
            if total < bg_range[0] or total > bg_range[1]:
                self.warnings.append(f"选题背景 {total} 字不在 {bg_range[0]}~{bg_range[1]} 区间")
        else:
            self.add_h2("（一）时代背景")
            self.add_para("（请填写时代背景，必须含权威数据来源。）")
            self.add_h2("（二）现实痛点")
            self.add_para("（请填写现实痛点，2~3 个真实场景，必须可量化。）")
            self.add_h2("（三）国内外研究现状")
            self.add_para("（请填写研究现状，评述已有方案不足，引出本项目差异。）")
            self.add_h2("（四）项目意义")
            self.add_para("（请填写项目意义，理论/实践/社会三角度，至少两个。）")

    # --------------------------------------------------------
    # 六、研究方法与过程【重点】
    # --------------------------------------------------------

    def _add_research_method(self):
        self.add_h1("六、研究方法与过程")
        methods = self._get_list("research_method")
        limits = self._word_limits()
        rm_range = limits["research_method"]
        if methods:
            for para in methods:
                self.add_para(para)
            total = count_chinese_chars(join_paragraphs(methods))
            if total < rm_range[0] or total > rm_range[1]:
                self.warnings.append(f"研究方法 {total} 字不在 {rm_range[0]}~{rm_range[1]} 区间")
        else:
            self.add_h2("（一）文献调研")
            self.add_para("（请填写文献调研，覆盖篇数+时间范围+核心结论。）")
            self.add_h2("（二）数据/样本采集")
            self.add_para("（请填写数据采集，来源+数量+采集方式+伦理审查。）")
            self.add_h2("（三）实验/调研/仿真设计")
            self.add_para("（请填写实验设计，变量控制+对照组+重复次数+统计方法。）")
            self.add_h2("（四）实地测试/验证")
            self.add_para("（请填写实地测试，地点+时长+评估指标+误差分析。）")

    # --------------------------------------------------------
    # 七、研究结果与讨论【重点】
    # --------------------------------------------------------

    def _add_results_discussion(self):
        self.add_h1("七、研究结果与讨论")
        results = self._get_list("results_discussion")
        limits = self._word_limits()
        rd_range = limits["results_discussion"]
        if results:
            for para in results:
                self.add_para(para)
            total = count_chinese_chars(join_paragraphs(results))
            if total < rd_range[0] or total > rd_range[1]:
                self.warnings.append(f"研究结果 {total} 字不在 {rd_range[0]}~{rd_range[1]} 区间")
        else:
            self.add_h2("（一）主要发现")
            self.add_para("（请填写主要发现，按重要性排序，每发现含数据+图表编号引用。）")
            self.add_h2("（二）与已有研究对比")
            self.add_para("（请填写对比，量化差异，如准确率从 80% 提升至 92%，p<0.01。）")
            self.add_h2("（三）结果讨论")
            self.add_para("（请填写讨论，理论/实践含义，可推广边界。）")
            self.add_h2("（四）局限性分析")
            self.add_para("（请填写局限性，2~3 条，必须真实可改进。）")

    # --------------------------------------------------------
    # 八、结论
    # --------------------------------------------------------

    def _add_conclusion(self):
        self.add_h1("八、结论")
        conclusion = self._get("conclusion")
        limits = self._word_limits()
        c_range = limits["conclusion"]
        if conclusion:
            self.add_para(conclusion)
            total = count_chinese_chars(conclusion)
            if total < c_range[0] or total > c_range[1]:
                self.warnings.append(f"结论 {total} 字不在 {c_range[0]}~{c_range[1]} 区间")
        else:
            self.add_para("（请填写结论，3 句话结构：核心发现+学术贡献+下一步计划。）")

    # --------------------------------------------------------
    # 九、创新点
    # --------------------------------------------------------

    def _add_innovations(self):
        self.add_h1("九、创新点")
        innovations = self._get_list("innovations")
        if innovations:
            rows = []
            for i, inv in enumerate(innovations, 1):
                rows.append(format_innovation_row(i, inv))
            self.add_table(
                ["序号", "类型", "创新点描述", "对比优势"],
                rows,
                col_widths=[1.5, 2.5, 6.0, 6.0],
            )
            if len(innovations) < 3 or len(innovations) > 5:
                self.warnings.append(f"创新点 {len(innovations)} 个不在 3~5 区间")
        else:
            self.add_para("（请填写创新点 3~5 个，每个 50~100 字，必须含对比+量化。）")

    # --------------------------------------------------------
    # 十、参考文献
    # --------------------------------------------------------

    def _add_references(self):
        self.add_h1("十、参考文献")
        references = self._get_list("references")
        if references:
            for i, ref in enumerate(references, 1):
                ref_text = format_reference(i, ref)
                add_reference_paragraph(self.doc, ref_text)
            if len(references) < 5:
                self.warnings.append(f"参考文献 {len(references)} 条少于 5 条下限")
            if len(references) < 30:
                self.warnings.append(f"参考文献 {len(references)} 条少于 v2.1 推荐 30 条")
        else:
            add_reference_paragraph(self.doc, "[1] 张三, 李四. 示例文献题名[J]. 期刊名, 2024, 48(3): 45-52.")
            add_reference_paragraph(self.doc, "（请按 GB/T 7714 格式补 5~30 条参考文献，v2.1 推荐 30+。）")

    # --------------------------------------------------------
    # 十一、附录
    # --------------------------------------------------------

    def _add_appendix(self):
        self.add_h1("十一、附录")
        appendix = self._get_list("appendix")
        if appendix:
            for i, item in enumerate(appendix, 1):
                if isinstance(item, dict):
                    title = item.get("title", f"附录 {i}")
                    content = item.get("content", "")
                    self.add_h3(f"附录 {i}：{title}")
                    if content:
                        self.add_para(content)
                else:
                    self.add_h3(f"附录 {i}")
                    self.add_para(str(item))
        else:
            self.add_h3("附录 1：图表清单")
            self.add_para("（请补充图表清单。）")
            self.add_h3("附录 2：问卷（社科类必备）")
            self.add_para("（请补充调研问卷。）")
            self.add_h3("附录 3：原始数据")
            self.add_para("（请补充原始数据字段说明+样本量。）")
            self.add_h3("附录 4：代码/算法（科技发明类必备）")
            self.add_para("（请补充 CNN 模型核心代码。）")
            self.add_h3("附录 5：实物照片（科技发明类必备）")
            self.add_para("（请补充实物照片。）")

    # ------------------------------------------------========
    # v2.1 新增：A. 国家政策引用
    # --------------------------------------------------------

    def _add_policy_citations(self):
        """v2.1：国家政策引用（8+ 项）"""
        self.add_h1("A. 国家政策引用（v2.1 必加）")
        policies = self._get_list("policy_citations")
        if policies:
            self.add_para(f"本项目共引用 {len(policies)} 项国家政策文件，按时间倒序排列，"
                         "每项含发文机关、文号、标题、时间 4 要素，并摘录与项目相关的关键表述。")
            rows = []
            for i, p in enumerate(policies, 1):
                if isinstance(p, dict):
                    rows.append([
                        str(i),
                        p.get("publish_date", ""),
                        p.get("issuer", ""),
                        p.get("title", ""),
                        p.get("key_excerpt", "")[:80] + ("..." if len(p.get("key_excerpt", "")) > 80 else ""),
                    ])
                else:
                    rows.append([str(i), "", "", str(p), ""])
            self.add_table(
                ["序号", "时间", "发文机关", "标题", "关键表述摘录"],
                rows,
                col_widths=[1.0, 1.8, 2.5, 4.5, 5.5],
            )
            if len(policies) < 8:
                self.warnings.append(f"v2.1：政策引用 {len(policies)} 项少于 8 项下限")

            # 详细摘录
            self.add_h2("（详细摘录）")
            for i, p in enumerate(policies, 1):
                if isinstance(p, dict):
                    self.add_h3(f"政策 {i}：{p.get('title', '')}")
                    detail = f"{p.get('publish_date', '')}，{p.get('issuer', '')}发布"
                    if p.get("doc_no"):
                        detail += f"（{p['doc_no']}）"
                    detail += f"《{p.get('title', '')}》"
                    self.add_para(detail)
                    if p.get("key_excerpt"):
                        self.add_para(f"关键表述：{p['key_excerpt']}")
                    if p.get("relevance"):
                        self.add_para(f"与项目相关性：{p['relevance']}")
        else:
            self.add_para("（v2.1：请补充国家政策引用 ≥ 8 项，按时间倒序，4 要素齐全。）")

    # --------------------------------------------------------
    # v2.1 新增：B. 科学挑战 3 段
    # --------------------------------------------------------

    def _add_scientific_challenges(self):
        """v2.1：科学挑战 3 段"""
        self.add_h1("B. 科学挑战（v2.1 必加）")
        challenges = self._get_list("scientific_challenges")
        if challenges:
            self.add_para(f"本项目凝练 {len(challenges)} 个科学挑战，按数据层 → 算法层 → 系统层层层递进，"
                         "每个挑战含文献支撑与 2~3 个子挑战。")
            for i, ch in enumerate(challenges, 1):
                if isinstance(ch, dict):
                    self.add_h2(f"科学挑战 {i}：{ch.get('title', '')}")
                    if ch.get("description"):
                        self.add_para(ch["description"])
                    subs = ch.get("sub_challenges", []) or []
                    for j, sub in enumerate(subs, 1):
                        if isinstance(sub, dict):
                            self.add_h3(f"子挑战 {i}.{j}：{sub.get('title', '')}")
                            if sub.get("detail"):
                                self.add_para(sub["detail"])
                            if sub.get("reference"):
                                self.add_para(f"文献支撑：{sub['reference']}")
            if len(challenges) != 3:
                self.warnings.append(f"v2.1：科学挑战 {len(challenges)} 段，必须 3 段")
        else:
            self.add_para("（v2.1：请补充 3 段科学挑战，层层递进，含文献支撑。）")

    # --------------------------------------------------------
    # v2.1 新增：C. 国内外研究现状综述（30+ 文献）
    # --------------------------------------------------------

    def _add_literature_review(self):
        """v2.1：国内外研究现状综述"""
        self.add_h1("C. 国内外研究现状综述（v2.1 必加）")
        review = self._get_list("literature_review")
        if review:
            self.add_para(f"本综述共引用 {len(review)} 段，覆盖 30+ 篇文献（英文 ≥ 60%，含 SCI 一区/二区期刊），"
                         "按研究意义 → 国内外现状 → 总结分析三段结构撰写。")
            current_section = ""
            for seg in review:
                if isinstance(seg, dict):
                    section = seg.get("section", "")
                    if section and section != current_section:
                        current_section = section
                        section_labels = {
                            "research_significance": "（一）研究意义",
                            "domestic_international": "（二）国内外研究现状",
                            "summary_gap": "（三）现有研究总结与分析",
                        }
                        self.add_h2(section_labels.get(section, section))
                    if seg.get("topic"):
                        self.add_h3(seg["topic"])
                    if seg.get("content"):
                        self.add_para(seg["content"])
                    cites = seg.get("citations", [])
                    if cites:
                        self.add_para(f"本段引用文献：{', '.join(cites)}")
            if len(review) < 5:
                self.warnings.append(f"v2.1：文献综述 {len(review)} 段，建议 ≥ 5 段（覆盖 30+ 文献）")
        else:
            self.add_h2("（一）研究意义")
            self.add_para("（v2.1：请补充研究意义，3~5 段，每段一个小方向。）")
            self.add_h2("（二）国内外研究现状")
            self.add_para("（v2.1：请补充国内外现状，按技术方向分段，每段引 3~5 篇文献。）")
            self.add_h2("（三）现有研究总结与分析")
            self.add_para("（v2.1：请补充总结分析，指出 gap。）")

    # --------------------------------------------------------
    # v2.1 新增：D. 算法/方法对比表
    # --------------------------------------------------------

    def _add_algorithm_comparison(self):
        """v2.1：算法/方法对比表"""
        self.add_h1("D. 算法/方法对比表（v2.1 必加）")
        ac = self._get("algorithm_comparison")
        if isinstance(ac, dict) and ac:
            self.add_h2(ac.get("title", "算法对比表"))
            dimensions = ac.get("dimensions", [])
            algorithms = ac.get("algorithms", [])
            if dimensions and algorithms:
                headers = ["维度"] + [a.get("name", "") + (f" {a.get('ref', '')}" if a.get("ref") else "")
                                      for a in algorithms]
                rows = []
                for d_idx, dim in enumerate(dimensions):
                    row = [dim]
                    for alg in algorithms:
                        values = alg.get("values", [])
                        row.append(values[d_idx] if d_idx < len(values) else "")
                    rows.append(row)
                self.add_table(headers, rows, col_widths=[3.0] + [3.0] * len(algorithms))

            if ac.get("conclusion"):
                self.add_h2("（选型结论）")
                self.add_para(ac["conclusion"])
        else:
            self.add_para("（v2.1：请补充算法对比表，至少 1 张，≥3 维度，含选型结论。）")

    # --------------------------------------------------------
    # v2.1 新增：E. 技术路线图（3 张）
    # --------------------------------------------------------

    def _add_tech_roadmap(self):
        """v2.1：技术路线图（3 张）"""
        self.add_h1("E. 技术路线图（v2.1 必加）")
        roadmaps = self._get_list("tech_roadmap")
        if roadmaps:
            self.add_para(f"本项目共绘制 {len(roadmaps)} 张技术路线图，分别为研究内容关系图、"
                         "研究方法图、实施技术路线图。")
            for rm in roadmaps:
                if isinstance(rm, dict):
                    self.add_h2(f"{rm.get('figure_no', '')} {rm.get('title', '')}")
                    if rm.get("description"):
                        self.add_para(rm["description"])
                    nodes = rm.get("nodes", []) or []
                    edges = rm.get("edges", []) or []
                    if nodes:
                        self.add_h3("（节点列表）")
                        rows = []
                        for n in nodes:
                            if isinstance(n, dict):
                                rows.append([n.get("id", ""), n.get("label", "")])
                        if rows:
                            self.add_table(["节点编号", "节点名称"], rows,
                                          col_widths=[3.0, 12.0])
                    if edges:
                        self.add_h3("（关联关系）")
                        rows = []
                        for e in edges:
                            if isinstance(e, dict):
                                rows.append([e.get("from", ""), e.get("to", "")])
                        if rows:
                            self.add_table(["起始节点", "目标节点"], rows,
                                          col_widths=[6.0, 6.0])
            if len(roadmaps) < 3:
                self.warnings.append(f"v2.1：技术路线图 {len(roadmaps)} 张，必须 ≥ 3 张")
        else:
            self.add_para("（v2.1：请补充 3 张技术路线图：研究内容关系图+研究方法图+实施技术路线图。）")

    # --------------------------------------------------------
    # v2.1 新增：F. 数学公式（自然/发明类必备）
    # --------------------------------------------------------

    def _add_formulas(self):
        """v2.1：数学公式"""
        wt = self._work_type()
        if wt == "social_science":
            self.add_h1("F. 数学公式（v2.1：哲学社会科学类可选）")
            self.add_para("（哲学社会科学类可不写数学公式，但建议含统计模型公式。）")
            return

        self.add_h1("F. 数学公式（v2.1 必加）")
        formulas = self._get_list("formulas")
        if formulas:
            min_count = 3 if wt == "natural_science" else 2
            self.add_para(f"本项目共定义 {len(formulas)} 个数学公式"
                         f"（{wt} 类要求 ≥ {min_count} 个），每个公式含编号、表达式、变量解释。")
            for f in formulas:
                if isinstance(f, dict):
                    self.add_h2(f.get("equation_no", ""))
                    if f.get("context"):
                        self.add_para(f"所属算法：{f['context']}")
                    if f.get("expression"):
                        self.add_formula(f["expression"])
                    variables = f.get("variables", []) or []
                    if variables:
                        self.add_h3("（变量解释）")
                        rows = []
                        for v in variables:
                            if isinstance(v, dict):
                                rows.append([v.get("symbol", ""), v.get("meaning", "")])
                        if rows:
                            self.add_table(["符号", "含义"], rows,
                                          col_widths=[3.0, 12.0])
            if len(formulas) < min_count:
                self.warnings.append(f"v2.1：{wt} 类公式 {len(formulas)} 个少于 {min_count} 个下限")
        else:
            self.add_para(f"（v2.1：请补充 ≥ {3 if wt == 'natural_science' else 2} 个数学公式，含编号+变量解释。）")

    # --------------------------------------------------------
    # v2.1 新增：G. 社会经济效益量化（10 项）
    # --------------------------------------------------------

    def _add_economic_benefits(self):
        """v2.1：社会经济效益量化（10 项）"""
        self.add_h1("G. 社会经济效益量化（v2.1 必加）")
        benefits = self._get_list("economic_benefits")
        if benefits:
            self.add_para(f"本项目量化评估 {len(benefits)} 项社会经济效益指标，"
                         "对比传统模式与系统模式，每项含传统基准、系统预期、提升幅度 4 列。")
            rows = []
            for i, b in enumerate(benefits, 1):
                if isinstance(b, dict):
                    rows.append([
                        str(i),
                        b.get("indicator", ""),
                        b.get("traditional_baseline", ""),
                        b.get("system_expected", ""),
                        b.get("improvement", ""),
                    ])
            if rows:
                self.add_table(
                    ["序号", "评估指标", "传统模式基准值", "系统模式预期值", "效益提升幅度"],
                    rows,
                    col_widths=[1.0, 3.5, 3.0, 3.0, 4.5],
                )
            if len(benefits) < 10:
                self.warnings.append(f"v2.1：经济效益 {len(benefits)} 项少于 10 项下限")
        else:
            self.add_para("（v2.1：请补充 ≥ 10 项社会经济效益指标，每项 4 列对比。）")

    # --------------------------------------------------------
    # v2.1 新增：H. 答辩 PPT 框架（10 页）
    # --------------------------------------------------------

    def _add_defense_ppt(self):
        """v2.1：答辩 PPT 框架（10 页）"""
        self.add_h1("H. 答辩 PPT 框架（v2.1 必加）")
        ppt = self._get_list("defense_ppt")
        if ppt:
            self.add_para(f"答辩采用 5 分钟陈述 + 5 分钟答辩形式，共 {len(ppt)} 页 PPT 大纲。")
            for slide in ppt:
                if isinstance(slide, dict):
                    self.add_h2(f"第 {slide.get('slide_no', '?')} 页：{slide.get('title', '')}")
                    if slide.get("duration"):
                        self.add_para(f"时长建议：{slide['duration']}")
                    key_points = slide.get("key_points", []) or []
                    if key_points:
                        self.add_h3("（要点）")
                        for kp in key_points:
                            self.add_para(f"● {kp}", indent=False)
                    if slide.get("visual"):
                        self.add_para(f"视觉建议：{slide['visual']}")
                    if slide.get("speaker_notes"):
                        self.add_para(f"演讲备注：{slide['speaker_notes']}")
            if len(ppt) != 10:
                self.warnings.append(f"v2.1：答辩 PPT {len(ppt)} 页，必须 10 页")
        else:
            self.add_para("（v2.1：请补充 10 页 PPT 大纲，5+5 分钟。）")

    # --------------------------------------------------------
    # 三级评审意见
    # --------------------------------------------------------

    def _add_review_sections(self):
        review_specs = [
            ("十二、指导教师推荐意见", "指导教师签字"),
            ("十三、学校评审意见", "学校盖章"),
            ("十四、省级/国家级评审意见", "评审委员会盖章"),
        ]
        for title, sign_label in review_specs:
            self.add_h1(title)
            for _ in range(6):
                self.doc.add_paragraph()
            self.add_para(
                f"{sign_label}：____________________    日期：______年____月____日",
                indent=False,
            )

    # --------------------------------------------------------
    # 主构建方法
    # --------------------------------------------------------

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 14+10 栏目，生成 docx"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()

            self._add_cover()
            self._add_work_info()
            self._add_applicant_info()
            self._add_abstract()
            self._add_background()
            self._add_research_method()
            self._add_results_discussion()
            self._add_conclusion()
            self._add_innovations()
            self._add_references()
            self._add_appendix()

            # v2.1 新增 10 个章节
            self._add_policy_citations()
            self._add_scientific_challenges()
            self._add_literature_review()
            self._add_algorithm_comparison()
            self._add_tech_roadmap()
            self._add_formulas()
            self._add_economic_benefits()
            self._add_defense_ppt()

            self._add_review_sections()

            if self.warnings:
                print("⚠️ 数据校验警告：", file=sys.stderr)
                for w in self.warnings:
                    print(f"  - {w}", file=sys.stderr)

            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申报书已生成：{output_path}")
        return str(output_path)

    # --------------------------------------------------------
    # 数据校验（v2.1 增强）
    # --------------------------------------------------------

    def _validate_data(self) -> List[str]:
        p0_fields = [
            ("work_full_name", "作品全称"),
            ("work_type", "作品类别 work_type"),
            ("applicant_type", "申报者类型"),
            ("leader_name", "申报者代表姓名"),
            ("leader_id", "学号"),
            ("college", "学院"),
            ("school", "学校"),
        ]
        for key, name in p0_fields:
            if not self._get(key):
                # work_type 与 category 等价，二选一即可
                if key == "work_type" and self._get("category"):
                    continue
                self.warnings.append(f"缺少 {name}（{key}）")

        wt = self._work_type()
        if wt not in CATEGORY_LABEL:
            self.warnings.append(f"作品类别 {wt} 不在三大类中")

        if self._get("applicant_type") and self._get("applicant_type") not in APPLICANT_LABEL:
            self.warnings.append(f"申报者类型 {self._get('applicant_type')} 不在 individual/collective 中")

        if not self._get("abstract"):
            self.warnings.append("缺少 作品简介（abstract）")
        if not self._get("background"):
            self.warnings.append("缺少 选题背景（background）")
        if not self._get("research_method"):
            self.warnings.append("缺少 研究方法（research_method）")
        if not self._get("results_discussion"):
            self.warnings.append("缺少 研究结果（results_discussion）")
        if not self._get("conclusion"):
            self.warnings.append("缺少 结论（conclusion）")
        if not self._get("innovations"):
            self.warnings.append("缺少 创新点（innovations）")
        if not self._get("references"):
            self.warnings.append("缺少 参考文献（references）")

        # v2.1 必加字段校验
        if not self._get("policy_citations"):
            self.warnings.append("v2.1：缺少 国家政策引用（policy_citations），必须 ≥ 8 项")
        else:
            n = len(self._get_list("policy_citations"))
            if n < 8:
                self.warnings.append(f"v2.1：政策引用 {n} 项少于 8 项下限")

        if not self._get("scientific_challenges"):
            self.warnings.append("v2.1：缺少 科学挑战（scientific_challenges），必须 3 段")
        else:
            n = len(self._get_list("scientific_challenges"))
            if n != 3:
                self.warnings.append(f"v2.1：科学挑战 {n} 段，必须 3 段")

        if not self._get("literature_review"):
            self.warnings.append("v2.1：缺少 文献综述（literature_review），建议 30+ 文献")

        if not self._get("tech_roadmap"):
            self.warnings.append("v2.1：缺少 技术路线图（tech_roadmap），必须 ≥ 3 张")
        else:
            n = len(self._get_list("tech_roadmap"))
            if n < 3:
                self.warnings.append(f"v2.1：技术路线图 {n} 张少于 3 张下限")

        if wt in ("natural_science", "tech_invention"):
            if not self._get("formulas"):
                min_n = 3 if wt == "natural_science" else 2
                self.warnings.append(f"v2.1：{wt} 类缺少 数学公式（formulas），必须 ≥ {min_n} 个")
            else:
                n = len(self._get_list("formulas"))
                min_n = 3 if wt == "natural_science" else 2
                if n < min_n:
                    self.warnings.append(f"v2.1：{wt} 类公式 {n} 个少于 {min_n} 个下限")

        if not self._get("economic_benefits"):
            self.warnings.append("v2.1：缺少 经济效益量化（economic_benefits），必须 ≥ 10 项")
        else:
            n = len(self._get_list("economic_benefits"))
            if n < 10:
                self.warnings.append(f"v2.1：经济效益 {n} 项少于 10 项下限")

        if not self._get("defense_ppt"):
            self.warnings.append("v2.1：缺少 答辩 PPT（defense_ppt），必须 10 页")
        else:
            n = len(self._get_list("defense_ppt"))
            if n != 10:
                self.warnings.append(f"v2.1：答辩 PPT {n} 页，必须 10 页")

        if self._get("applicant_type") == "collective":
            members = self._get_list("team_members")
            if not members:
                self.warnings.append("集体申报缺少团队成员（team_members）")
            elif len(members) > 7:
                self.warnings.append(f"集体申报团队成员 {len(members)} 人超过 7 人上限")

        work_name = self._get("work_full_name")
        if work_name and len(work_name) > 30:
            self.warnings.append(f"作品全称 {len(work_name)} 字超过 30 字上限")

        return self.warnings


# ============================================================
# 默认示例数据（v2.1：消防无人机主题，对齐案例 1+2）
# ============================================================

DEFAULT_DATA = {
    # 基础信息
    "work_full_name": "基于多模态融合与三维重构的消防无人机火场救援系统",
    "work_type": "tech_invention",
    "applicant_type": "collective",
    "leader_name": "姚奕晗", "leader_id": "241080200223", "leader_gender": "男",
    "leader_age": "21", "leader_major": "电气工程及其自动化",
    "leader_grade": "2024 级 大二", "leader_phone": "167XXXXXXXX",
    "college": "电气工程学院", "school": "XX 工业大学",
    "apply_date": "2026 年 5 月 6 日",
    "discipline": "0808 电气工程",
    "duration": "2026.06-2027.06（12 个月）",
    "advisor_name": "闫晶晶", "advisor_title": "副教授",
    "advisor_research": "多模态融合与三维重构", "advisor_phone": "186XXXXXXXX",
    "team_members": [
        {"name": "周匡吉", "id": "241080200432", "major": "电气工程及其自动化",
         "grade": "2024 级", "role": "数据采集与算法实现"},
        {"name": "韩兆峰", "id": "241080200106", "major": "电气工程及其自动化",
         "grade": "2024 级", "role": "硬件搭建与现场测试"},
    ],

    # 作品简介
    "abstract": (
        "本项目针对城市火灾救援响应慢、被困人员识别难、火场态势感知弱的痛点，"
        "开发基于多模态融合与三维重构的消防无人机火场救援系统。"
        "融合可见光、红外、激光雷达、IMU 四种模态数据，采用 Fire-Transformer "
        "多模态融合算法与 Fire-LIO-SAM 三维重构算法，构建火场数字孪生。"
        "训练真实火场 5000+ 样本，覆盖 8 类典型火场目标，目标检测准确率目标 92%+，"
        "三维重构误差 ≤5 cm，单次响应 <1 秒。"
        "预期产出 1 套原型系统、2 篇中文核心论文、3 项发明专利。"
        "已完成前期调研与 500 样本预训练，预实验准确率 89.5%。"
    ),

    # 选题背景与意义
    "background": [
        "（一）时代背景：随着我国『大安全、大应急』框架构建，消防救援装备智能化成为国家战略。"
        "2024 年 12 月中办国办《关于推进新型城市基础设施建设打造韧性城市的意见》明确提出"
        "'发展智慧消防，利用物联网、人工智能等技术提升城市消防安全水平'。"
        "据国家消防救援局统计，2024 年全国接报火灾 74.5 万起，死亡 2338 人，直接财产损失 67.6 亿元。"
        "传统消防救援依赖人工侦察，响应慢、风险高，亟需智能化装备升级。",

        "（二）现实痛点：调研 XX 市消防救援支队发现，城市火灾救援存在三大痛点："
        "① 响应慢，从接警到救援力量到达平均 8 分钟，被困人员搜救平均 25 分钟；"
        "② 识别难，浓烟环境下人工侦察被困人员召回率 ≤70%，错失黄金救援时间；"
        "③ 态势弱，指挥员对火势蔓延方向、危险品位置等关键态势感知不足，"
        "导致 2024 年全国消防员伤亡事故 23 起、伤亡 35 人。"
        "若按全国 660 个城市估算，每年因响应慢导致的人员伤亡与财产损失超 200 亿元。",

        "（三）国内外研究现状：国外方面，Smith 等[5]提出基于 Transformer 的多模态融合方法，"
        "在 KITTI 数据集上准确率 92.3%。Wang 等[6]进一步引入注意力机制，准确率 94.1%。"
        "国内方面，张三等[7]针对火场环境提出自适应权重融合，准确率 89.5%。"
        "但上述方法均未考虑浓烟环境，且三维重构误差 >15 cm，无法满足消防实战需求。"
        "本项目关键差异在于：（1）针对浓烟环境优化多模态融合策略；"
        "（2）三维重构误差降至 5 cm 以内；（3）火场目标检测覆盖 8 类典型目标。",

        "（四）项目意义：理论上探索多模态融合在极端环境（浓烟、高温）下的鲁棒性边界，"
        "验证 Fire-Transformer 在小样本火场数据上的泛化能力；"
        "实践上与 XX 市消防救援支队合作开发可落地系统，"
        "预期将救援响应时间从 8 分钟缩短至 3 分钟，被困人员识别召回率从 70% 提升至 92%+；"
        "社会上助力『大安全、大应急』框架构建，推动消防救援装备智能化升级，"
        "降低消防员高危环境暴露频次，年减少直接经济损失 10-20 亿元。"
    ],

    # 研究方法与过程
    "research_method": [
        "（一）文献调研：系统梳理 2020~2026 年国内外消防无人机、多模态融合、三维重构、"
        "目标检测文献 50 篇（中文 18 篇、英文 32 篇），覆盖 IEEE TGRS、CVPR、ICCV、"
        "Drones、Sensors 等 SCI 期刊与顶会。核心结论为：传统方法[1-3]在浓烟环境下"
        "检测准确率降至 60%-70%，多模态融合方法[5-9]可提升至 85%-90% 但三维重构误差"
        "仍 >15 cm。本项目以此为研究空白，聚焦浓烟环境下的高精度多模态融合与三维重构。",

        "（二）数据采集：与 XX 市消防救援支队签署数据合作协议，"
        "获取 2024~2025 年真实火场数据 5000+ 样本，含可见光、红外、激光雷达、IMU 四模态。"
        "采用专家交叉标注（3 人独立标注+多数表决）保证标注质量，"
        "标注一致性 Kappa 系数 0.87。涉及火场数据已签署保密协议，"
        "未涉及个人信息，经 XX 大学伦理委员会审查通过（编号 ETH-2025-018）。"
        "样本覆盖 8 类典型火场目标：被困人员、火源、烟雾、危险品、楼梯、门窗、消防通道、结构损伤。"
        "每类含 4 模态同步数据，时间戳对齐误差 <10 ms。",

        "（三）实验设计：在 XX 大学消防工程实验室搭建火场模拟平台，"
        "控制环境温度 25-200℃、能见度 0.5-5 m，对 8 类目标各采集 500+ 样本。"
        "以传统 YOLOv5 单模态方法为对照组，本组 Fire-Transformer 多模态融合方法"
        "随机初始化训练 3 次取平均，评估指标为准确率/召回率/F1。"
        "模型架构采用 Fire-Transformer（改进版 Vision Transformer + 跨模态注意力），"
        "学习率 0.0001，批量大小 16，训练 200 epoch，使用 4×NVIDIA A100 GPU。"
        "三维重构采用 Fire-LIO-SAM（改进版 LIO-SAM + 火场自适应权重），"
        "评估指标为关键尺寸误差/姿态误差/实时性。统计差异用配对 t 检验（p<0.05 视为显著）。",

        "（四）实地测试：将训练好的 Fire-Transformer + Fire-LIO-SAM 系统部署在"
        "XX 市消防救援支队 3 个消防站，自 2026 年 1 月起试运行 3 个月，"
        "累计参与真实火场救援 12 次，处理实时数据 240 万条。"
        "评估指标为准确率/召回率/响应时间/三维重构误差，"
        "对比基线为消防员人工侦察记录。误差分析显示浓烟环境（能见度 <1 m）下"
        "红外模态贡献度达 65%，可见光模态贡献度降至 15%，验证多模态融合的必要性。"
        "其余 7 类目标准确率均≥90%，仅浓烟极端环境下『被困人员』准确率 85.2%，"
        "需进一步增加浓烟样本。"
    ],

    # 研究结果与讨论
    "results_discussion": [
        "（一）主要发现：本方法在测试集上目标检测准确率 92.3%，召回率 89.5%，"
        "F1 0.908（详见图 3、表 4）。其中被困人员识别准确率 89.5%，火源识别 96.8%，"
        "危险品识别 93.2%，浓烟环境下平均准确率 87.6%。三维火场重构关键尺寸误差 ≤5 cm"
        "（详见图 5），姿态误差 ≤2°，单次重构响应时间 0.78 秒。"
        "多模态融合比单模态（仅可见光）准确率提升 22.3 个百分点。"
        "在 3 个消防站实地测试中累计识别火场目标 487 次，"
        "其中 451 次为真实目标，准确率 92.6%（详见表 6）。"
        "救援响应时间从平均 8 分钟缩短至 3.2 分钟（缩短 60%）。",

        "（二）与已有研究对比：与 Smith (2024)[5] 的 Transformer 方法准确率 92.3% 相比，"
        "本方法准确率持平（92.3% vs 92.3%），但在浓烟环境下提升 8.5 个百分点（87.6% vs 79.1%）。"
        "与 Zhang (2025)[7] 的自适应权重融合方法准确率 89.5% 相比，"
        "本方法准确率提升 2.8 个百分点（p<0.01，配对 t 检验）。"
        "三维重构方面，本方法误差 5 cm，相比 Wang (2024)[6] 的 15 cm 精度提升 3 倍。"
        "本方法覆盖 8 类火场目标，已有研究多覆盖 3~5 类，覆盖度为 1.6~2.7 倍。",

        "（三）结果讨论：本方法验证了 Fire-Transformer 多模态融合架构在极端环境"
        "（浓烟、高温）下的鲁棒性，证明跨模态注意力机制能有效补偿单一模态退化。"
        "Fire-LIO-SAM 三维重构算法在火场动态环境下精度显著优于传统 LIO-SAM，"
        "说明火场自适应权重策略对动态点云处理至关重要。"
        "多模态融合提升微小目标（如被困人员）识别精度，"
        "说明单一可见光模态不足以刻画火场全貌。"
        "本方法可推广至其他应急救援场景（地震、洪灾、危化品泄漏），"
        "推广边界为单无人机载荷 ≤5 kg、目标类别 ≤15 类、环境温度 ≤300℃。",

        "（四）局限性分析：本研究存在 3 点局限："
        "（1）样本仅来自华东地区城市火灾，未涵盖西北高粉尘环境与森林火灾场景，"
        "模型在森林火灾场景下泛化能力待验证；"
        "（2）浓烟极端环境（能见度 <0.5 m）下被困人员样本仅 200 个，"
        "小样本类别精度仍有提升空间，未来计划采集 1000+ 浓烟样本；"
        "（3）模型在边缘设备（Jetson Orin Nano）部署推理延迟待优化，"
        "当前 RTX 4090 推理 0.78 秒，移植到 Jetson Orin Nano 后延迟升至 2.1 秒，"
        "需进行 INT8 量化+TensorRT 加速，目标优化至 1 秒以内。"
    ],

    # 结论
    "conclusion": (
        "本研究开发了基于多模态融合与三维重构的消防无人机火场救援系统，"
        "在真实火场 5000+ 样本上目标检测准确率 92.3%，三维重构误差 ≤5 cm，"
        "单次响应 0.78 秒，覆盖 8 类典型火场目标。"
        "学术贡献在于验证了 Fire-Transformer 跨模态注意力机制在浓烟极端环境下的鲁棒性，"
        "并提供真实火场数据基准与 Fire-LIO-SAM 火场自适应三维重构算法。"
        "下一步计划扩展样本至 20000+，覆盖森林火灾与危化品泄漏场景，"
        "并优化边缘设备部署推理延迟至 1 秒以内，"
        "推动消防救援装备智能化升级，助力『大安全、大应急』框架构建。"
    ),

    # 创新点
    "innovations": [
        {"type": "方法创新",
         "description": "提出 Fire-Transformer 多模态融合算法，引入跨模态注意力机制补偿浓烟环境模态退化。",
         "advantage": "vs 传统 YOLOv5 单模态，准确率 70%→92.3%，浓烟环境提升 22.3 个百分点"},
        {"type": "方法创新",
         "description": "提出 Fire-LIO-SAM 火场自适应三维重构算法，动态调整点云权重。",
         "advantage": "vs 传统 LIO-SAM，三维重构误差 15cm→5cm，精度提升 3 倍"},
        {"type": "数据创新",
         "description": "与 XX 市消防救援支队合作获取真实火场 5000+ 四模态样本，覆盖 8 类目标。",
         "advantage": "vs 实验室仿真数据集 3~5 类，覆盖度提升 1.6~2.7 倍"},
        {"type": "视角创新",
         "description": "融合可见光+红外+激光雷达+IMU 四模态，特征级与决策级双融合。",
         "advantage": "vs 单模态准确率提升 22.3 个百分点，浓烟环境鲁棒性显著增强"},
        {"type": "应用创新",
         "description": "原型系统部署 3 个消防站试运行 3 个月，参与真实火场救援 12 次。",
         "advantage": "vs 实验室原型无实地验证，救援响应时间 8 分钟→3.2 分钟，缩短 60%"}
    ],

    # 参考文献（v2.1 推荐 30+，此处列 8 条核心，完整 30+ 在 literature_review 中体现）
    "references": [
        {"authors": "张三, 李四", "title": "基于多模态融合的消防无人机目标检测研究",
         "type": "J", "source": "中国应急管理", "year": "2025", "volume": "15", "issue": "3", "pages": "45-52"},
        {"authors": "Smith J, Brown K", "title": "Transformer-based multi-modal fusion for fire detection",
         "type": "J", "source": "IEEE Transactions on Geoscience and Remote Sensing",
         "year": "2024", "volume": "62", "issue": "", "pages": "1-12"},
        {"authors": "Wang Y, Liu Z", "title": "3D fire scene reconstruction with adaptive weights",
         "type": "J", "source": "Engineering Applications of Artificial Intelligence",
         "year": "2024", "volume": "127", "issue": "", "pages": "107-118"},
        {"authors": "Zhang Q, Li M", "title": "Fire-Transformer: cross-modal attention for fire detection",
         "type": "C", "source": "Proc. of CVPR", "year": "2025", "volume": "", "issue": "",
         "pages": "234-240"},
        {"authors": "Chen X, Wang H", "title": "Fire-LIO-SAM: adaptive LiDAR-Inertial Odometry for fire scenes",
         "type": "J", "source": "Drones", "year": "2025", "volume": "9", "issue": "2", "pages": "55-70"},
        {"authors": "李华, 王明", "title": "消防无人机技术与应用", "type": "M",
         "source": "", "year": "2024", "publisher": "中国消防出版社",
         "city": "北京", "pages": "88-105"},
        {"authors": "赵六", "title": "多模态融合与三维重构研究", "type": "D",
         "source": "清华大学", "year": "2024", "city": "北京", "publisher": ""},
        {"authors": "全国消防标准化技术委员会", "title": "消防无人机技术规范: GB/T 41234-2024",
         "type": "S", "source": "GB/T 41234-2024", "year": "2024",
         "city": "北京", "publisher": "中国标准出版社"}
    ],

    # 附录
    "appendix": [
        {"title": "8 类火场目标样本示例图",
         "content": "图 1~图 8 展示被困人员、火源、烟雾、危险品、楼梯、门窗、消防通道、结构损伤等 "
                   "8 类典型火场目标的可见光、红外、激光雷达、IMU 四模态同步样本。"
                   "每类含 5 个代表性样本，标注目标位置与特征。"},
        {"title": "原始数据字段说明",
         "content": "原始数据 5000+ 样本，每样本含 4 模态同步数据。"
                   "字段含 sample_id/target_type/timestamp/visible_img/ir_img/lidar_pts/imu_data/"
                   "temperature/visibility 等。时间戳对齐误差 <10 ms。"},
        {"title": "Fire-Transformer 模型核心代码",
         "content": "Python 3.10+PyTorch 2.1 实现，含数据加载（4 模态 DataLoader）+"
                   "模型定义（Fire-Transformer，含跨模态注意力）+训练循环（200 epoch，"
                   "学习率 0.0001，批量 16，4×A100）+评估脚本。"},
        {"title": "实物照片",
         "content": "原型系统实物照片含 5 视图（正/侧/背/内部/部署现场），"
                   "无人机型号 DJI Matrice 300 RTK，载荷 2.5 kg，"
                   "含可见光相机+红外热成像+激光雷达+IMU 四种传感器。"}
    ],

    # ==================== v2.1 新增字段 ====================

    # A. 国家政策引用（8 项）
    "policy_citations": [
        {"issuer": "国务院办公厅", "doc_no": "国办发〔2026〕12 号",
         "title": "关于加强基层消防工作的意见", "publish_date": "2026.03",
         "key_excerpt": "加强基层消防力量建设，推广无人机、机器人等智能装备在基层消防救援中的应用",
         "relevance": "直接对应项目无人机消防应用方向"},
        {"issuer": "国家层面", "doc_no": "",
         "title": "关于低空经济发展部署", "publish_date": "2026.02",
         "key_excerpt": "推动低空经济高质量发展，鼓励无人机在应急救援等公共服务领域创新应用",
         "relevance": "项目消防无人机属低空经济应用场景"},
        {"issuer": "工信部等", "doc_no": "",
         "title": "应急装备产业重点产品发展指导目录（2025 版）", "publish_date": "2025.12",
         "key_excerpt": "将消防无人机、多模态感知装备列入重点发展产品目录",
         "relevance": "项目核心产品列入国家重点发展目录"},
        {"issuer": "国新办", "doc_no": "",
         "title": "《新时代的中国国家安全》白皮书", "publish_date": "2025.05",
         "key_excerpt": "构建大安全大应急框架，提升应急救援智能化水平",
         "relevance": "项目响应大安全大应急框架构建"},
        {"issuer": "中办国办", "doc_no": "",
         "title": "关于推进新型城市基础设施建设打造韧性城市的意见", "publish_date": "2024.12",
         "key_excerpt": "发展智慧消防，利用物联网、人工智能等技术提升城市消防安全水平",
         "relevance": "项目智慧消防技术直接对应政策要求"},
        {"issuer": "工信部等四部门", "doc_no": "",
         "title": "通用航空装备创新应用实施方案（2024-2030 年）", "publish_date": "2024.03",
         "key_excerpt": "推动无人驾驶航空器在应急救援等场景创新应用",
         "relevance": "项目消防无人机属通用航空装备创新应用"},
        {"issuer": "应急管理部、工信部", "doc_no": "",
         "title": "关于加快应急机器人发展的指导意见", "publish_date": "2023.12",
         "key_excerpt": "突出建设测试基地及公共服务平台，研发全国应急管理系统无人机综合信息平台",
         "relevance": "项目无人机系统对接应急管理部信息化建设"},
        {"issuer": "应急管理部", "doc_no": "",
         "title": "《\"十四五\"应急救援力量建设规划》", "publish_date": "2022.06",
         "key_excerpt": "加快构建大型固定翼灭火飞机、灭火直升机与无人机高低搭配的应急救援航空器体系",
         "relevance": "项目消防无人机属应急救援航空器体系"},
        {"issuer": "国务院", "doc_no": "国发〔2021〕109 号",
         "title": "《\"十四五\"国家应急体系规划》", "publish_date": "2021.12",
         "key_excerpt": "推广运用智能机器人、无人机等高技术配送装备，提升应急运输调度效率",
         "relevance": "项目消防无人机属国家应急体系高技术装备"},
        {"issuer": "习近平", "doc_no": "",
         "title": "国家综合性消防救援队伍授旗致辞", "publish_date": "2018.11",
         "key_excerpt": "组建国家综合性消防救援队伍是构建新时代国家应急救援体系的重要举措",
         "relevance": "项目响应国家应急救援体系建设"}
    ],

    # B. 科学挑战（3 段）
    "scientific_challenges": [
        {
            "title": "多模态数据时空融合与三维火场实时重构计算",
            "description": (
                "多模态融合是消防无人机火场救援的基础。现有方法[5,6]在可见光、红外、"
                "激光雷达、IMU 等异构传感器数据时空对齐方面存在显著不足，特征表达不充分。"
                "本项目需在 0.5 m 能见度浓烟环境下实现多模态数据的高效融合与实时三维重构，"
                "面临算力约束（无人机载荷功率 ≤500 W）与实时性要求（≤1 秒）的双重挑战。"
            ),
            "sub_challenges": [
                {"title": "异构传感器数据时空配准",
                 "detail": "可见光帧率 30 fps、红外 25 fps、激光雷达 10 Hz、IMU 200 Hz，"
                          "四种传感器采样频率差异导致时间戳难以对齐。空间上不同传感器视场角与坐标系不一致，"
                          "需进行外参标定。传统方法[10]依赖固定标定板，在动态火场环境下精度下降明显。"
                          "本项目拟提出基于 ICP+IMU 预积分的在线动态标定方法，目标时间戳对齐误差 <10 ms。",
                 "reference": "[10] Wang Y, et al. Online calibration for multi-modal sensors. Sensors, 2024."},
                {"title": "跨模态特征表示学习",
                 "detail": "可见光特征为 2D 像素，红外为 2D 热分布，激光雷达为 3D 点云，"
                          "特征维度不一致导致融合困难。BEV 转换[11]虽统一表示但损失高度信息。"
                          "如何设计跨模态特征映射函数 φ(·) 使不同模态特征在同一空间可比，"
                          "是关键科学问题。本项目拟提出 Fire-Transformer 跨模态注意力机制，"
                          "在 Token 级别实现模态间信息交互。",
                 "reference": "[11] Chen X, et al. BEV-based 3D detection. CVPR, 2024."}
            ]
        },
        {
            "title": "火场数字孪生中目标检测与态势理解耦合技术",
            "description": (
                "火场环境具有浓烟、高温、低照度、动态变化等特点，传统目标检测方法在此环境下"
                "性能急剧下降。本项目需在 0.5 m 能见度下实现 90%+ 检测准确率[15]，"
                "并将检测结果与三维火场数字孪生耦合，实现态势理解（火势蔓延方向、"
                "危险品位置、被困人员位置等），为指挥决策提供支撑。"
            ),
            "sub_challenges": [
                {"title": "浓烟环境下视觉特征退化补偿",
                 "detail": "浓烟导致可见光图像对比度下降 70%+，传统 CNN 检测准确率从 95% 降至 60%[16]。"
                          "需引入红外模态补偿，但红外图像分辨率低、纹理弱，跨模态融合策略待优化。"
                          "本项目拟提出可见光-红外特征级融合 + 注意力引导的退化补偿方法。",
                 "reference": "[16] Smith J, et al. Smoke-robust detection. IEEE TGRS, 2024."},
                {"title": "目标检测结果与数字孪生耦合",
                 "detail": "传统方法将目标检测与三维重构独立处理，丢失时空关联信息。"
                          "本项目拟提出 Fire-LIO-SAM + 目标检测联合优化框架，"
                          "将检测结果作为数字孪生的语义标签，实现态势理解。",
                 "reference": "[17] Zhang Q, et al. Joint detection and reconstruction. ICCV, 2025."}
            ]
        },
        {
            "title": "火灾场景的自适应路径规划与在线学习机制研究",
            "description": (
                "火场环境动态变化（火势蔓延、烟雾扩散、结构坍塌），传统离线路径规划方法"
                "无法适应实时变化。本项目需研究基于 SAC（Soft Actor-Critic）强化学习的"
                "自适应路径规划方法，并设计在线学习机制，使无人机能在飞行中持续优化策略。"
            ),
            "sub_challenges": [
                {"title": "动态环境下的自适应路径规划",
                 "detail": "传统 DQN/DDPG 算法在动态环境下探索能力不足，易陷局部最优。"
                          "本项目拟采用 SAC 算法（最大熵强化学习），平衡探索与利用，"
                          "在火场动态环境中实现稳健路径规划。仿真实验显示 SAC 比 DPG 样本效率提升 50%。",
                 "reference": "[20] Haarnoja T, et al. Soft Actor-Critic. ICML, 2018."},
                {"title": "在线学习与模型更新",
                 "detail": "离线训练的模型在火场真实环境下泛化能力有限。"
                          "本项目拟设计基于元学习的在线适应机制，无人机在飞行中收集新样本，"
                          "每 5 分钟微调一次模型，目标是将新环境下的检测准确率从 75% 提升至 88%+。",
                 "reference": "[21] Finn C, et al. Model-agnostic meta-learning. ICML, 2017."}
            ]
        }
    ],

    # C. 国内外研究现状综述（30+ 文献）
    "literature_review": [
        {"section": "research_significance", "topic": "理论意义",
         "citations": ["[5]", "[6]", "[7]"],
         "content": "多模态融合与三维重构是机器人感知领域的核心科学问题[5-7]。"
                   "本研究在浓烟极端环境下探索多模态融合的鲁棒性边界，"
                   "对推动机器人感知理论发展具有重要意义。"},
        {"section": "research_significance", "topic": "方法意义",
         "citations": ["[8]", "[9]", "[10]"],
         "content": "现有方法[8-10]多在理想环境下验证，缺乏极端环境下的方法创新。"
                   "本研究提出的 Fire-Transformer 跨模态注意力机制为极端环境多模态融合提供新思路。"},
        {"section": "research_significance", "topic": "应用意义",
         "citations": ["[15]", "[16]", "[17]"],
         "content": "消防救援是高风险场景，智能化装备需求迫切[15-17]。"
                   "本研究开发的系统可显著降低消防员高危环境暴露，提升救援效率。"},
        {"section": "research_significance", "topic": "社会意义",
         "citations": ["[18]", "[19]"],
         "content": "据国家消防救援局统计，2024 年全国火灾直接财产损失 67.6 亿元[18]。"
                   "本研究有望将损失降低 20%-40%，产生显著社会效益[19]。"},

        {"section": "domestic_international", "topic": "（1）多模态融合研究",
         "citations": ["[5]", "[6]", "[7]", "[8]", "[9]"],
         "content": "国外方面，Smith 等[5]提出基于 Transformer 的多模态融合方法，"
                   "在 KITTI 数据集上准确率 92.3%。Wang 等[6]进一步引入注意力机制，"
                   "准确率提升至 94.1%。国内方面，张三等[7]针对火场环境提出自适应权重融合，"
                   "准确率 89.5%。Chen 等[8]提出 BEV 统一表示，但损失高度信息。"
                   "Liu 等[9]提出跨模态知识蒸馏，小样本场景下提升 5 个百分点。"
                   "但上述方法均未考虑浓烟环境。"},
        {"section": "domestic_international", "topic": "（2）三维重构研究",
         "citations": ["[10]", "[11]", "[12]", "[13]"],
         "content": "三维重构方面，传统 LIO-SAM[10]在静态环境下精度高（误差 3 cm），"
                   "但火场动态环境下精度下降至 15 cm+。Wang 等[11]提出动态点云过滤，"
                   "误差降至 10 cm。Zhang 等[12]引入语义信息，误差 8 cm。"
                   "本项目提出的 Fire-LIO-SAM 引入火场自适应权重，误差降至 5 cm 以内。"},
        {"section": "domestic_international", "topic": "（3）火场目标检测研究",
         "citations": ["[14]", "[15]", "[16]", "[17]"],
         "content": "火场目标检测方面，传统 YOLOv5[14]在浓烟环境下准确率降至 60%。"
                   "Smith 等[15]提出烟雾鲁棒检测，准确率 75%。Zhang 等[16]引入红外模态，"
                   "准确率 82%。Chen 等[17]提出多模态融合，准确率 88%。"
                   "本项目覆盖 8 类目标，准确率 92.3%，优于已有研究。"},
        {"section": "domestic_international", "topic": "（4）路径规划研究",
         "citations": ["[18]", "[19]", "[20]", "[21]"],
         "content": "路径规划方面，传统 DQN[18]适用于离散动作空间，DDPG[19]适用于连续动作。"
                   "Haarnoja 等[20]提出 SAC 算法，平衡探索与利用。Finn 等[21]提出元学习，"
                   "实现快速适应新环境。本项目综合 SAC + 元学习，实现火场动态环境自适应路径规划。"},

        {"section": "summary_gap", "topic": "现有研究总结与分析",
         "citations": ["[5]", "[6]", "[7]", "[10]", "[15]", "[20]"],
         "content": "综上所述，国内外学者在多模态融合、三维重构、火场目标检测、路径规划"
                   "取得了一系列成果，但仍有以下不足："
                   "① 现有方法[5-7]未考虑浓烟极端环境，准确率降至 60%-70%；"
                   "② 三维重构误差 >15 cm，无法满足消防实战需求[10]；"
                   "③ 火场目标检测覆盖类别少（3~5 类）[15]；"
                   "④ 路径规划缺乏在线学习能力[20]。"
                   "本项目针对上述不足，提出 Fire-Transformer + Fire-LIO-SAM + SAC + 元学习"
                   "四位一体方案，预期实现浓烟环境准确率 92%+、三维重构误差 ≤5 cm、"
                   "目标覆盖 8 类、在线学习适应能力 88%+。"}
    ],

    # D. 算法对比表
    "algorithm_comparison": {
        "title": "主流自适应路径规划算法对比表",
        "dimensions": ["技术特点", "优势", "劣势", "适用场景", "效率（样本数）"],
        "algorithms": [
            {"name": "DQN", "ref": "[18]",
             "values": ["离散动作空间，Q 学习", "收敛稳定，理论保证",
                       "不适合连续动作", "网格化路径规划", "10^6"]},
            {"name": "DDPG", "ref": "[19]",
             "values": ["连续动作空间，Actor-Critic", "适合连续控制，样本效率高",
                       "探索能力弱，易陷局部最优", "机械臂控制", "10^5"]},
            {"name": "SAC", "ref": "[20]",
             "values": ["最大熵强化学习，连续动作", "探索能力强，收敛快",
                       "超参数多，调参复杂", "复杂动态环境", "5×10^4"]}
        ],
        "conclusion": (
            "基于上表对比，火场环境复杂动态，需强探索能力，本项目选用 SAC 算法。"
            "相比 DQN 样本效率提升 20 倍，相比 DDPG 探索能力增强 40%。"
            "同时引入元学习机制，实现新环境快速适应。"
        )
    },

    # E. 技术路线图（3 张）
    "tech_roadmap": [
        {
            "figure_no": "图 1",
            "title": "项目研究内容关系图",
            "description": (
                "本图展示项目 3 大研究内容（多模态融合、三维重构、目标检测与态势理解）"
                "及其相互关系。其中内容 1（多模态融合）为内容 2（三维重构）提供输入数据，"
                "内容 2 为内容 3（目标检测与态势理解）提供场景模型，三者构成递进关系。"
                "最终汇总至预期成果（消防无人机火场救援系统原型）。"
            ),
            "nodes": [
                {"id": "1.1", "label": "国家规划与技术需求"},
                {"id": "1.2", "label": "研究意义与国内外研究现状"},
                {"id": "2.1", "label": "多模态数据时空融合算法"},
                {"id": "2.2", "label": "三维火场实时重构方法"},
                {"id": "2.3", "label": "火场目标检测与态势理解"},
                {"id": "3.1", "label": "算法实现与系统集成"},
                {"id": "3.2", "label": "实地测试与性能评估"}
            ],
            "edges": [
                {"from": "1.1", "to": "1.2"},
                {"from": "1.2", "to": "2.1"},
                {"from": "2.1", "to": "2.2"},
                {"from": "2.2", "to": "2.3"},
                {"from": "2.3", "to": "3.1"},
                {"from": "3.1", "to": "3.2"}
            ]
        },
        {
            "figure_no": "图 2",
            "title": "项目研究方法图",
            "description": (
                "本图展示项目的研究方法和思路，按问题定义 → 文献调研 → 理论建模 → "
                "算法设计 → 系统实现 → 实验仿真 → 实地验证 → 结果分析 8 步闭环。"
                "每步对应研究方法章节的具体内容，确保方法严谨性。"
            ),
            "nodes": [
                {"id": "M1", "label": "问题定义（消防无人机火场救援）"},
                {"id": "M2", "label": "文献调研（50 篇 2020~2026）"},
                {"id": "M3", "label": "理论建模（Fire-Transformer）"},
                {"id": "M4", "label": "算法设计（Fire-LIO-SAM）"},
                {"id": "M5", "label": "系统实现（原型开发）"},
                {"id": "M6", "label": "实验仿真（5000+ 样本）"},
                {"id": "M7", "label": "实地验证（3 消防站 3 个月）"},
                {"id": "M8", "label": "结果分析（准确率/误差/响应时间）"}
            ],
            "edges": [
                {"from": "M1", "to": "M2"}, {"from": "M2", "to": "M3"},
                {"from": "M3", "to": "M4"}, {"from": "M4", "to": "M5"},
                {"from": "M5", "to": "M6"}, {"from": "M6", "to": "M7"},
                {"from": "M7", "to": "M8"}
            ]
        },
        {
            "figure_no": "图 3",
            "title": "项目拟采用的实施技术路线图",
            "description": (
                "本图展示项目实施的具体技术路线，分 4 阶段："
                "阶段 1（M1-M3）：需求分析 + 数据采集；"
                "阶段 2（M4-M9）：算法开发 + 模型训练；"
                "阶段 3（M10-M11）：系统集成 + 场景验证；"
                "阶段 4（M12）：总结 + 文档 + 演示。"
                "总周期 12 个月（2026.06-2027.06）。"
            ),
            "nodes": [
                {"id": "P1", "label": "阶段 1（2026.6-7）：需求分析+数据采集"},
                {"id": "P2", "label": "阶段 2（2026.8-12）：算法开发+模型训练"},
                {"id": "P3", "label": "阶段 3（2027.1-4）：系统集成+场景验证"},
                {"id": "P4", "label": "阶段 4（2027.5-6）：总结+文档+演示"}
            ],
            "edges": [
                {"from": "P1", "to": "P2"}, {"from": "P2", "to": "P3"},
                {"from": "P3", "to": "P4"}
            ]
        }
    ],

    # F. 数学公式
    "formulas": [
        {
            "equation_no": "式(1) 密度计算公式",
            "expression": "ρ(p) = |N(p, r)| / V(r)",
            "variables": [
                {"symbol": "ρ(p)", "meaning": "点 p 处的局部密度"},
                {"symbol": "N(p, r)", "meaning": "以点 p 为中心、半径为 r 的邻域内的点集"},
                {"symbol": "V(r)", "meaning": "邻域球体积，V(r) = (4/3)πr³"}
            ],
            "context": "Fire-LIO-SAM 三维重构算法 - 点云密度估计"
        },
        {
            "equation_no": "式(2) 动态扩张率",
            "expression": "δ = δ_min + (δ_max - δ_min) · (1 - ρ̂)",
            "variables": [
                {"symbol": "δ", "meaning": "动态扩张率"},
                {"symbol": "δ_min, δ_max", "meaning": "扩张率的上下界，δ_min=0.5, δ_max=2.0"},
                {"symbol": "ρ̂", "meaning": "归一化密度值，ρ̂ = (ρ - ρ_min) / (ρ_max - ρ_min)"}
            ],
            "context": "Fire-LIO-SAM 三维重构算法 - 自适应权重计算"
        },
        {
            "equation_no": "式(3) 跨模态注意力",
            "expression": "Attention(Q, K, V) = softmax(Q·K^T / √d_k) · V",
            "variables": [
                {"symbol": "Q", "meaning": "查询矩阵（query），来自可见光特征"},
                {"symbol": "K", "meaning": "键矩阵（key），来自红外特征"},
                {"symbol": "V", "meaning": "值矩阵（value），来自红外特征"},
                {"symbol": "d_k", "meaning": "键向量维度，d_k=64"}
            ],
            "context": "Fire-Transformer 多模态融合算法 - 跨模态注意力机制"
        }
    ],

    # G. 社会经济效益量化（10 项）
    "economic_benefits": [
        {"indicator": "单次城市火灾平均经济损失",
         "traditional_baseline": "100%（基准）",
         "system_expected": "60%-80%",
         "improvement": "降低 20%-40%，约 1.5 倍"},
        {"indicator": "年均可减少直接经济损失",
         "traditional_baseline": "约 50 亿元/年",
         "system_expected": "约 30-40 亿元/年",
         "improvement": "减少 10-20 亿元/年，约 1.5 倍"},
        {"indicator": "被困人员识别召回率",
         "traditional_baseline": "≤70%",
         "system_expected": "≥92%",
         "improvement": "提升 ≥22 个百分点，约 1.31 倍"},
        {"indicator": "三维火场重构关键尺寸误差",
         "traditional_baseline": ">15 cm",
         "system_expected": "≤5 cm",
         "improvement": "精度提升 ≥3 倍"},
        {"indicator": "火势蔓延预测关键节点误差",
         "traditional_baseline": ">10 m",
         "system_expected": "≤2 m",
         "improvement": "精度提升 ≥5 倍"},
        {"indicator": "救援响应时间",
         "traditional_baseline": "100%（基准，8 分钟）",
         "system_expected": "30%-50%（3.2 分钟）",
         "improvement": "缩短 50%-70%，约 2.5 倍"},
        {"indicator": "消防员高危环境暴露频次",
         "traditional_baseline": "100%（基准）",
         "system_expected": "≤40%",
         "improvement": "降低 ≥60%，约 2.5 倍"},
        {"indicator": "救援力量调度准确率",
         "traditional_baseline": "≤60%",
         "system_expected": "≥85%",
         "improvement": "提升 ≥25 个百分点，约 1.42 倍"},
        {"indicator": "装备物资调配效率",
         "traditional_baseline": "100%（基准）",
         "system_expected": "130%-150%",
         "improvement": "提升 30%-50%，约 1.4 倍"},
        {"indicator": "危险品识别与预警响应时间",
         "traditional_baseline": ">5 分钟",
         "system_expected": "≤1 分钟",
         "improvement": "缩短 ≥80%，约 5 倍"}
    ],

    # H. 答辩 PPT（10 页）
    "defense_ppt": [
        {"slide_no": 1, "title": "封面",
         "duration": "30 秒",
         "key_points": ["作品全称：基于多模态融合与三维重构的消防无人机火场救援系统",
                       "团队：姚奕晗（负责人）+ 周匡吉 + 韩兆峰",
                       "指导教师：闫晶晶 副教授",
                       "学校：XX 工业大学"],
         "visual": "学校 LOGO + 无人机火场救援实拍图",
         "speaker_notes": "自我介绍 15 秒，作品名称 15 秒"},
        {"slide_no": 2, "title": "研究背景与意义",
         "duration": "30 秒",
         "key_points": ["国家政策：『大安全、大应急』框架，2024 年中办国办文件",
                       "现实痛点：2024 年全国火灾 74.5 万起，损失 67.6 亿元",
                       "研究意义：理论+实践+社会三角度"],
         "visual": "政策时间轴 + 痛点数据图",
         "speaker_notes": "开门见山，30 秒讲清『为什么做』"},
        {"slide_no": 3, "title": "科学挑战与研究目标",
         "duration": "30 秒",
         "key_points": ["科学挑战 1：多模态数据时空融合与三维重构",
                       "科学挑战 2：火场目标检测与态势理解耦合",
                       "科学挑战 3：自适应路径规划与在线学习",
                       "研究目标：准确率 92%+，误差 ≤5 cm，响应 <1 秒"],
         "visual": "3 段挑战递进图",
         "speaker_notes": "突出『层层递进』逻辑"},
        {"slide_no": 4, "title": "研究方法与技术路线 1",
         "duration": "30 秒",
         "key_points": ["文献调研：50 篇 2020~2026",
                       "数据采集：5000+ 真实火场样本，4 模态",
                       "技术路线图 1：研究内容关系图"],
         "visual": "技术路线图 1",
         "speaker_notes": "方法严谨性是评审重点"},
        {"slide_no": 5, "title": "研究方法与技术路线 2",
         "duration": "30 秒",
         "key_points": ["技术路线图 2：研究方法图（8 步闭环）",
                       "技术路线图 3：实施技术路线图（4 阶段 12 个月）"],
         "visual": "技术路线图 2+3 并排",
         "speaker_notes": "突出『完整闭环』"},
        {"slide_no": 6, "title": "核心算法 1：Fire-Transformer",
         "duration": "30 秒",
         "key_points": ["Fire-Transformer：跨模态注意力机制",
                       "式(3) Attention(Q, K, V) = softmax(Q·K^T/√d_k)·V",
                       "Q 来自可见光，K/V 来自红外",
                       "创新点：跨模态信息交互"],
         "visual": "公式 + 算法流程图",
         "speaker_notes": "算法创新是核心创新点"},
        {"slide_no": 7, "title": "核心算法 2：Fire-LIO-SAM + 算法对比",
         "duration": "30 秒",
         "key_points": ["Fire-LIO-SAM：火场自适应三维重构",
                       "式(1) 密度计算 + 式(2) 动态扩张率",
                       "算法对比表：DQN/DDPG/SAC，本项目选 SAC",
                       "选型结论：SAC 样本效率提升 20 倍"],
         "visual": "算法对比表",
         "speaker_notes": "突出对比优势"},
        {"slide_no": 8, "title": "实验结果与对比",
         "duration": "30 秒",
         "key_points": ["准确率 92.3%，召回率 89.5%，F1 0.908",
                       "三维重构误差 ≤5 cm（vs 传统 15 cm，提升 3 倍）",
                       "实地测试：3 消防站 3 个月，12 次真实救援",
                       "响应时间 8 分钟→3.2 分钟，缩短 60%"],
         "visual": "结果对比柱状图 + 实地照片",
         "speaker_notes": "用数据说话"},
        {"slide_no": 9, "title": "创新点与社会经济效益",
         "duration": "30 秒",
         "key_points": ["5 个创新点（方法+方法+数据+视角+应用）",
                       "经济损失：降低 20%-40%",
                       "救援响应：缩短 50%-70%",
                       "消防员暴露：降低 ≥60%"],
         "visual": "创新点表格 + 经济效益雷达图",
         "speaker_notes": "突出『对比+量化』"},
        {"slide_no": 10, "title": "结论与未来工作",
         "duration": "30 秒",
         "key_points": ["核心发现：准确率 92.3%，误差 ≤5 cm",
                       "学术贡献：Fire-Transformer + Fire-LIO-SAM",
                       "未来工作：扩展样本至 20000+，覆盖森林火灾",
                       "致谢"],
         "visual": "路线图 + 致谢",
         "speaker_notes": "留 5 秒致谢"}
    ]
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="挑战杯课外学术科技作品竞赛申报书 docx 生成器（v2.1）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n"
            "v2.1 新增字段：policy_citations / scientific_challenges / "
            "literature_review / algorithm_comparison / tech_roadmap / "
            "formulas / economic_benefits / defense_ppt / work_type\n"
            "\n"
            "JSON 字段定义详见 SKILL.md 第十九章。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档（消防无人机主题）")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档（v2.1 消防无人机主题，对齐案例 1+2）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


# 入口：python build.py --demo --out demo.docx
# 或：  python build.py --data data.json --out output.docx
if __name__ == "__main__":
    main()
