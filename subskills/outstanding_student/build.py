#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
优秀学生 / 优秀团员申请书 docx 生成器

支持两种评优类别（通过 category 字段切换）：
- outstanding_student: 优秀学生申请书（学业+综合素质主导，学生工作为重点段）
- outstanding_youth_league: 优秀团员申请书（政治表现+团组织生活+志愿活动主导）

评审跨度：本学年（不是 4 年）。正文 1500~2000 字，书信体格式。
按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 标题：黑体二号，居中（按 category 动态填充"优秀学生/优秀团员申请书"）
- 称呼：顶格，宋体小四，全角冒号
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 主干课程表、学生干部履职表：宋体五号，居中
- "此致"另起一行空两格，"敬礼！"另起一行顶格
- 落款：右对齐

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第四章信息采集清单。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号（标题）
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四（正文）
SIZE_WU = Pt(10.5)          # 五号（表格）
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

# 类别 -> 标题文本
CATEGORY_TITLE_MAP = {
    "outstanding_student": "优秀学生申请书",
    "outstanding_youth_league": "优秀团员申请书",
}


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_font(cell, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(doc, text: str, font_name: str = FONT_SONG,
        font_size=SIZE_XIAO_SI, bold: bool = False,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent: bool = True,
        line_spacing: float = 1.5, space_before: float = 0, space_after: float = 0):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_ER,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进），宋体小四"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_section_heading(doc, text: str):
    """正文小节标题（一、二、三…）：黑体小四加粗，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SI,
        bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
        line_spacing=1.5, space_before=6, space_after=3)


def add_cizhi_paragraph(doc, text: str = "此致"):
    """'此致'另起一行，空两格（首行缩进 2 字符）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格（不缩进）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=False, line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None, caption: str = ""):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    if caption:
        add_paragraph_with_format(doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def format_rank_percent(rank: str, rank_total: str) -> str:
    """根据排名与基数计算百分比，返回 '前 X.X%' 字符串，失败返回空串"""
    try:
        r_num = int(str(rank).split("/")[0])
        total_num = int(rank_total) if rank_total else 0
        if total_num > 0:
            return f"前 {round(r_num / total_num * 100, 1)}%"
    except (ValueError, IndexError):
        pass
    return ""


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """优秀学生 / 优秀团员申请书 docx 构建器（按 category 切换内容侧重）"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []
        # 缓存类别，默认优秀学生
        self.category = "outstanding_student"

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段，非列表返回空列表"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_heading(self, text):
        return add_section_heading(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, caption=""):
        return add_table_from_data(self.doc, headers, rows, col_widths, caption)

    def _is_youth_league(self) -> bool:
        """判断是否为优秀团员申请书"""
        cat = str(self._get("category", default="outstanding_student")).lower()
        return cat in ("outstanding_youth_league", "youth_league", "团员", "优秀团员")

    # --------------------------------------------------------
    # 标题（按 category 动态填充）
    # --------------------------------------------------------

    def _add_title(self):
        """标题：黑体二号居中，按 category 动态填充"优秀学生/优秀团员申请书" """
        cat = str(self._get("category", default="outstanding_student")).lower()
        if cat in ("outstanding_youth_league", "youth_league", "团员", "优秀团员"):
            self.category = "outstanding_youth_league"
            title = CATEGORY_TITLE_MAP["outstanding_youth_league"]
        else:
            self.category = "outstanding_student"
            title = CATEGORY_TITLE_MAP["outstanding_student"]
        add_title(self.doc, title)

    # --------------------------------------------------------
    # 称呼
    # --------------------------------------------------------

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号"""
        salutation = self._get("salutation",
                               default="尊敬的学院领导、评审委员会：")
        add_salutation_paragraph(self.doc, salutation)

    # --------------------------------------------------------
    # 开头段落
    # --------------------------------------------------------

    def _add_opening(self):
        """开头段落（80~120 字）：身份 + 申报奖项 + 本学年核心数据 + 进入正文"""
        opening = self._get("opening", default="")
        if opening:
            self.add_para(opening)
            return
        name, college, major, grade = self._get("name"), self._get("college"), self._get("major"), self._get("grade")
        apply_year = self._get("apply_year", default="2024-2025 学年")
        honor_name = "优秀团员" if self._is_youth_league() else "优秀学生"
        gpa, rank, position = self._get("gpa"), self._get("rank"), self._get("position", default="")
        # 取最高级别荣誉作为开头亮点
        honor_top = ""
        honors = self._get_list("honors")
        if honors and isinstance(honors[0], dict):
            h = honors[0]
            honor_top = f"获{h.get('name', '')}（{h.get('level', '')}）"
        parts = []
        if name and college and major and grade:
            parts.append(f"我是{college}{major}{grade}学生{name}，特申请{apply_year}{honor_name}。")
        else:
            parts.append(f"特申请{apply_year}{honor_name}。")
        data_parts = []
        if gpa:
            data_parts.append(f"GPA {gpa}")
        if rank:
            data_parts.append(f"专业排名第 {rank}")
        if honor_top:
            data_parts.append(honor_top)
        if position:
            data_parts.append(f"担任{position}")
        if data_parts:
            parts.append("；".join(data_parts) + "。")
        parts.append("现将本学年情况汇报如下：")
        self.add_para("".join(parts))

    # --------------------------------------------------------
    # 思想方面
    # --------------------------------------------------------

    def _add_ideology(self):
        """思想方面（200~300 字）：政治立场 + 入党/团情况 + 思想觉悟 + 具体活动

        优秀团员侧重团组织生活与政治表现；优秀学生侧重思想觉悟与党员发展。
        """
        self.add_heading("一、思想方面")
        ideology = self._get("ideology", default="")
        if ideology:
            if isinstance(ideology, list):
                for p in ideology:
                    self.add_para(p)
            else:
                self.add_para(ideology)
            return
        political = self._get("political_status", default="共青团员")
        party_history = self._get("party_history", default="")
        party_activities = self._get_list("party_activities")
        parts = []
        if political:
            if "党员" in political and "预备" not in political:
                parts.append(f"作为一名{political}，我始终坚持学习党的理论，认真学习习近平新时代中国特色社会主义思想与党的二十大和二十届三中全会精神。")
            elif "预备党员" in political:
                parts.append("作为一名中共预备党员，我认真学习习近平新时代中国特色社会主义思想，深入学习党的二十大和二十届三中全会精神，时刻以正式党员标准要求自己。")
            elif "积极分子" in political:
                parts.append("作为一名入党积极分子，我认真学习习近平新时代中国特色社会主义思想，时刻以党员标准要求自己。")
            else:
                parts.append(f"作为一名{political}，我拥护中国共产党的领导，认真学习党的创新理论。")
        if party_history:
            parts.append(party_history)
        # 优秀团员专属：团组织生活
        if self._is_youth_league():
            league_role = self._get("league_role", default="")
            league_activities = self._get_list("league_activities")
            if league_role:
                parts.append(f"在团组织中担任{league_role}，带头参加团的组织生活，按时足额缴纳团费，完成「智慧团建」系统信息维护。")
            if league_activities:
                parts.append("本学年参与" + "、".join(league_activities) + "等团日活动。")
        if party_activities:
            parts.append("；".join(party_activities))
        parts.append("在日常生活中，我注重理论学习与实践结合，关注时政热点，提升思想觉悟。")
        self.add_para("".join(parts))

    # --------------------------------------------------------
    # 学习方面（含主干课程表）
    # --------------------------------------------------------

    def _add_academics(self):
        """学习方面（300~400 字，重点）：GPA + 排名 + 主干课程表 + 英语计算机 + 学习方法"""
        self.add_heading("二、学习方面")
        academics = self._get("academics", default="")
        if academics and isinstance(academics, str):
            self.add_para(academics)
            self._add_core_courses_table()
            self._add_academics_summary()
            return
        gpa, weighted = self._get("gpa"), self._get("weighted_avg")
        rank, rank_total = self._get("rank"), self._get("rank_total")
        course_count = self._get("course_count")
        high_score_count = self._get("high_score_count")
        parts = []
        if gpa:
            gpa_str = f"本学年 GPA {gpa}"
            if weighted:
                gpa_str += f"，加权平均分 {weighted}"
            if rank:
                gpa_str += f"，专业排名第 {rank}"
                if rank_total:
                    pct = format_rank_percent(rank, rank_total)
                    if pct:
                        gpa_str += f"（{pct}）"
            parts.append(gpa_str + "。")
        if course_count and high_score_count:
            parts.append(f"本学年修读 {course_count} 门课程，{high_score_count} 门 85 分以上。")
        if parts:
            self.add_para("".join(parts))
        self._add_core_courses_table()
        self._add_academics_summary()

    def _add_core_courses_table(self):
        """主干课程表：5~8 门，含课程名/学分/成绩"""
        courses = self._get_list("core_courses")
        if not courses:
            return
        rows = []
        for c in courses:
            if not isinstance(c, dict):
                continue
            rows.append([
                str(c.get("name", "")),
                str(c.get("credit", "")),
                str(c.get("score", "")),
            ])
        if rows:
            self.add_table(
                ["课程名称", "学分", "成绩"],
                rows,
                col_widths=[8.0, 3.0, 3.0],
                caption="主干课程成绩：",
            )

    def _add_academics_summary(self):
        """学习方面末尾：课程亮点 + 英语计算机等级 + 学习方法"""
        summary = self._get("academics_summary", default="")
        if summary:
            self.add_para(summary)
            return
        cet4, cet6 = self._get("cet4"), self._get("cet6")
        computer_level = self._get("computer_level")
        course_highlight = self._get("course_highlight", default="")
        study_method = self._get("study_method",
            default="学习上注重课前预习与课后总结，建立知识体系；遇到问题主动与老师、同学讨论。")
        parts = []
        if course_highlight:
            parts.append(course_highlight + "。")
        lang_parts = []
        if cet4:
            lang_parts.append(f"CET-4 {cet4} 分")
        if cet6:
            lang_parts.append(f"CET-6 {cet6} 分")
        if lang_parts:
            parts.append("、".join(lang_parts) + "，能熟练阅读英文文献；")
        if computer_level:
            parts.append(f"计算机{computer_level}。")
        if study_method:
            parts.append(study_method)
        if parts:
            self.add_para("".join(parts))

    # --------------------------------------------------------
    # 科研与实践方面
    # --------------------------------------------------------

    def _add_research_practice(self):
        """科研与实践方面（250~400 字）：大创/论文/竞赛 + 志愿服务"""
        self.add_heading("三、科研与实践方面")
        rp_text = self._get("research_practice", default="")
        if rp_text and isinstance(rp_text, str):
            self.add_para(rp_text)
            self._add_research_items()
            self._add_competition_items()
            self._add_volunteer_practice()
            return
        self._add_research_items()
        self._add_competition_items()
        self._add_volunteer_practice()

    def _add_research_items(self):
        """大创/科研立项 + 论文"""
        projects = self._get_list("research_projects")
        papers = self._get_list("papers")
        if projects:
            proj_parts = []
            for p in projects:
                if not isinstance(p, dict):
                    continue
                name, level = p.get("name", ""), p.get("level", "")
                role, duration, output = p.get("role", ""), p.get("duration", ""), p.get("output", "")
                seg = f"{role}《{name}》" if role and name else (f"《{name}》" if name else "")
                if level:
                    seg += f"（{level}）"
                if duration:
                    seg += f"，{duration}"
                if output:
                    seg += f"，{output}"
                if seg:
                    proj_parts.append(seg + "。")
            if proj_parts:
                self.add_para("".join(proj_parts))
        if papers:
            paper_parts = []
            for p in papers:
                if not isinstance(p, dict):
                    continue
                title, journal, level = p.get("title", ""), p.get("journal", ""), p.get("level", "")
                author_order, time = p.get("author_order", ""), p.get("time", "")
                seg = (f"以{author_order}在" if author_order else "") + (f"《{journal}》" if journal else "")
                if level:
                    seg += f"（{level}）"
                seg += "发表论文"
                if title:
                    seg += f"《{title}》"
                if time:
                    seg += f"，{time} 见刊"
                if seg:
                    paper_parts.append(seg + "；")
            if paper_parts:
                paper_parts[-1] = paper_parts[-1].rstrip("；") + "。"
                self.add_para("".join(paper_parts))

    def _add_competition_items(self):
        """学科竞赛列表"""
        competitions = self._get_list("competitions")
        if not competitions:
            return
        comp_parts = []
        for c in competitions:
            if not isinstance(c, dict):
                continue
            time, name, award, role = c.get("time", ""), c.get("name", ""), c.get("award", ""), c.get("role", "")
            seg = (f"{time} " if time else "") + name
            if award:
                seg += f" {award}"
            if role:
                seg += f"（{role}）"
            if seg:
                comp_parts.append(seg + "；")
        if comp_parts:
            comp_parts[-1] = comp_parts[-1].rstrip("；") + "。"
            self.add_para("学科竞赛方面，" + "".join(comp_parts))

    def _add_volunteer_practice(self):
        """志愿服务段（优秀团员为侧重点，优秀学生为常规段）"""
        volunteer_hours = self._get("volunteer_hours", default="")
        volunteer_detail = self._get("volunteer_detail", default="")
        sanxiaxiang = self._get_list("sanxiaxiang")
        parts = []
        if volunteer_hours or volunteer_detail:
            seg = f"累计志愿服务时长 {volunteer_hours} 小时" if volunteer_hours else ""
            if volunteer_detail:
                seg = (seg + "：" if seg else "") + volunteer_detail
            parts.append(seg + "。")
        if sanxiaxiang:
            sx_parts = []
            for s in sanxiaxiang:
                if not isinstance(s, dict):
                    continue
                time, place, role, output = s.get("time", ""), s.get("place", ""), s.get("role", ""), s.get("output", "")
                seg = (f"{time} " if time else "") + (f"赴{place}" if place else "")
                if role:
                    seg += f"，{role}"
                if output:
                    seg += f"，{output}"
                if seg:
                    sx_parts.append(seg + "；")
            if sx_parts:
                sx_parts[-1] = sx_parts[-1].rstrip("；") + "。"
                parts.append("参与三下乡" + str(len(sanxiaxiang)) + " 次：" + "".join(sx_parts))
        if parts:
            self.add_para("".join(parts))

    # --------------------------------------------------------
    # 学生工作【优秀学生专属重点】
    # --------------------------------------------------------

    def _add_student_work(self):
        """学生工作段（150~250 字）：学生干部履职表（职务/任期/主要工作/成效）

        优秀学生：学生干部履职为必填重点段，含表格化呈现。
        优秀团员：转为"团务工作"段，强调团组织职务与团日活动组织。
        """
        if self._is_youth_league():
            self._add_league_work()
            return
        self.add_heading("四、学生工作")
        student_work = self._get("student_work", default="")
        if student_work:
            self.add_para(student_work)
            self._add_position_table()
            return
        position = self._get("position", default="")
        position_work = self._get("position_work", default="")
        if position:
            seg = f"担任{position}"
            if position_work:
                seg += f"，主要工作：{position_work}"
            self.add_para(seg + "。")
        self._add_position_table()

    def _add_position_table(self):
        """学生干部履职表（4 列：职务/任期/主要工作/成效）"""
        positions = self._get_list("position_history")
        if not positions:
            return
        rows = []
        for p in positions:
            if not isinstance(p, dict):
                continue
            rows.append([
                str(p.get("role", "")),
                str(p.get("term", "")),
                str(p.get("work", "")),
                str(p.get("effect", "")),
            ])
        if rows:
            self.add_table(
                ["职务", "任期", "主要工作", "成效"],
                rows,
                col_widths=[3.0, 2.5, 6.0, 3.5],
                caption="学生干部履职情况：",
            )

    def _add_league_work(self):
        """团务工作段（优秀团员专属）"""
        self.add_heading("四、团务工作")
        league_work = self._get("league_work", default="")
        if league_work:
            self.add_para(league_work)
            self._add_position_table()
            return
        league_role = self._get("league_role", default="")
        league_work_detail = self._get("league_work_detail", default="")
        if league_role:
            seg = f"担任{league_role}"
            if league_work_detail:
                seg += f"，主要工作：{league_work_detail}"
            self.add_para(seg + "。")
        self._add_position_table()

    # --------------------------------------------------------
    # 生活方面
    # --------------------------------------------------------

    def _add_life(self):
        """生活方面（120~200 字）：生活作风 + 宿舍长履职 + 人际关系"""
        self.add_heading("五、生活方面")
        life = self._get("life", default="")
        if life:
            if isinstance(life, list):
                for p in life:
                    self.add_para(p)
            else:
                self.add_para(life)
            return
        dorm_role = self._get("dorm_role", default="")
        dorm_activity = self._get("dorm_activity", default="")
        dorm_honor = self._get("dorm_honor", default="")
        interpersonal = self._get("interpersonal", default="")
        lifestyle = self._get("lifestyle", default="生活中我注重勤俭节约，作息规律。")
        parts = [lifestyle]
        if dorm_role and dorm_activity:
            seg = f"担任{dorm_role}期间，{dorm_activity}"
            if dorm_honor:
                seg += f"，{dorm_honor}"
            parts.append(seg + "。")
        if interpersonal:
            parts.append(interpersonal + "。")
        self.add_para("".join(parts))

    # --------------------------------------------------------
    # 结尾"此致 敬礼！"
    # --------------------------------------------------------

    def _add_ending(self):
        """结尾（60~100 字）：本学年事实总结 + 朴素表态 + 此致/敬礼！"""
        ending = self._get("ending", default="")
        if ending:
            self.add_para(ending)
        else:
            honor_name = "优秀团员" if self._is_youth_league() else "优秀学生"
            self.add_para(
                f"以上是我本学年的基本情况。无论评选结果如何，我都将以此为新的起点，继续努力学习、全面发展，以更高标准要求自己，争取获评{honor_name}。恳请评审委员会予以考虑。"
            )
        add_cizhi_paragraph(self.doc, "此致")   # "此致"另起一行，空两格
        add_jingli_paragraph(self.doc, "敬礼！")  # "敬礼！"另起一行，顶格

    # --------------------------------------------------------
    # 落款
    # --------------------------------------------------------

    def _add_signature(self):
        """落款：右对齐，含申请人 + 日期"""
        self.doc.add_paragraph()  # 空一行
        name = self._get("name", default="")
        apply_date = self._get("apply_date", default="")
        if name:
            add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if apply_date:
            add_right_aligned_paragraph(self.doc, apply_date)

    # --------------------------------------------------------
    # 主构建方法
    # --------------------------------------------------------

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/开头/思想/学习/科研实践/学生工作/生活/结尾/落款"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_opening()
            self._add_ideology()
            self._add_academics()
            self._add_research_practice()
            self._add_student_work()
            self._add_life()
            self._add_ending()
            self._add_signature()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申请书已生成：{output_path}")
        return str(output_path)

    # --------------------------------------------------------
    # 数据校验
    # --------------------------------------------------------

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []
        # 类别校验
        cat = str(self._get("category", default="outstanding_student")).lower()
        if cat not in ("outstanding_student", "outstanding_youth_league", "youth_league", "团员", "优秀团员"):
            warnings.append(f"类别 {cat} 非标准值，默认按优秀学生处理")
        # P0 必采字段
        for key, label in [("name", "申请人姓名"), ("college", "学院"), ("major", "专业"),
                           ("grade", "年级"), ("gpa", "GPA"), ("rank", "专业排名")]:
            if not self._get(key):
                warnings.append(f"缺少 {label}（{key}）")
        # 主干课程
        courses = self._get_list("core_courses")
        if not courses:
            warnings.append("缺少 主干课程（core_courses），将省略主干课程表")
        elif len(courses) < 5:
            warnings.append(f"主干课程仅 {len(courses)} 门，建议 5~8 门")
        # 排名校验
        rank_str = str(self._get("rank", default=""))
        if rank_str and "/" not in rank_str:
            if not self._get("rank_total"):
                warnings.append(
                    f"排名 '{rank_str}' 缺基数，应为 'X/N' 格式或补填 rank_total"
                )
        # 学业门槛校验（优秀学生要求前 30%，优秀团员放宽至前 50%）
        try:
            rank_num = int(rank_str.split("/")[0]) if "/" in rank_str else 0
            rank_total_num = 0
            if "/" in rank_str:
                rank_total_num = int(rank_str.split("/")[1])
            elif self._get("rank_total"):
                rank_total_num = int(self._get("rank_total"))
            if rank_num and rank_total_num:
                pct = rank_num / rank_total_num * 100
                threshold = 30 if not self._is_youth_league() else 50
                if pct > threshold:
                    warnings.append(
                        f"排名前 {pct:.1f}% 不满足{'优秀学生前 30%' if not self._is_youth_league() else '优秀团员前 50%'}门槛"
                    )
        except (ValueError, IndexError):
            pass
        # 学生干部履职表（优秀学生要求必填）
        if not self._is_youth_league() and not self._get_list("position_history"):
            warnings.append("优秀学生建议提供 position_history 学生干部履职表（至少 1 条）")
        # 优秀团员要求团务工作与志愿服务
        if self._is_youth_league():
            if not self._get("league_role") and not self._get("league_work"):
                warnings.append("优秀团员建议提供 league_role 团内职务或 league_work 团务工作段")
            if not self._get("volunteer_hours"):
                warnings.append("优秀团员建议提供 volunteer_hours 志愿服务时长（建议≥30 小时）")
        self.warnings = warnings
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据（优秀学生）
# ============================================================

DEFAULT_DATA = {
    "category": "outstanding_student", "name": "张明", "student_id": "2022123456",
    "gender": "男", "college": "计算机科学与技术学院", "major": "计算机科学与技术",
    "grade": "2022 级 大三", "class_name": "计科 2201 班", "political_status": "中共预备党员",
    "phone": "138XXXXXXXX", "apply_year": "2024-2025 学年", "apply_date": "2025 年 5 月 20 日",
    "salutation": "尊敬的学院领导、评审委员会：", "gpa": "3.88/4.0", "weighted_avg": "90.2",
    "rank": "2/87", "rank_total": "87", "course_count": "12", "high_score_count": "9",
    "core_courses": [
        {"name": "高等数学（下）", "credit": "5", "score": "94"},
        {"name": "数据结构", "credit": "4", "score": "93"},
        {"name": "操作系统", "credit": "4", "score": "92"},
        {"name": "计算机网络", "credit": "3", "score": "90"},
        {"name": "数据库原理", "credit": "3", "score": "91"},
        {"name": "计算机组成原理", "credit": "4", "score": "89"},
    ],
    "course_highlight": "高等数学 94、数据结构 93、操作系统 92，专业核心课平均 91.5 分，5 门 90+",
    "cet4": "568", "cet6": "532", "computer_level": "二级 C 语言（优秀）、三级数据库技术",
    "study_method": "学习上注重课前预习与课后总结，建立知识体系；遇到问题主动与老师、同学讨论。",
    "academics_summary": "",
    "party_history": "2023.09 提交入党申请书，2024.03 列为入党积极分子，2025.06 转为中共预备党员。",
    "party_activities": [
        "参加学院分党校第 8 期培训班（2024.09-2024.12）结业",
        "2024.10 参与主题党日活动'红色教育基地走访'，撰写调研报告 1 份",
        "提交思想汇报 4 篇",
    ],
    "ideology": "", "academics": "", "research_practice": "",
    "research_projects": [
        {"name": "分布式光伏故障智能诊断系统", "level": "校级大创项目", "role": "主持",
         "duration": "2024.03-2025.03", "output": "项目结题评估优秀，申请软件著作权 1 项"},
    ],
    "papers": [],
    "competitions": [
        {"name": "全国大学生数学建模竞赛", "award": "省级一等奖",
         "time": "2024.11", "role": "队长，负责整体建模与论文主笔"},
        {"name": "中国大学生计算机设计大赛", "award": "省级二等奖",
         "time": "2024.05", "role": "核心成员，负责前端开发"},
    ],
    "volunteer_hours": "80",
    "volunteer_detail": ("担任图书馆管理员（2024.09-2025.05，每月 6 小时）；"
                        "组织班级'一对一'帮扶活动，服务同学 30 余人次"),
    "sanxiaxiang": [
        {"time": "2024.07", "place": "甘肃定西", "role": "电商助农调研",
         "output": "走访 5 个村、访谈 60 户，撰写调研报告 1 份"},
    ],
    "student_work": "", "position": "学院学生会主席（2024.09-2025.05）",
    "position_work": ("组织学院迎新晚会、运动会等大型活动 8 场，累计参与 2000+ 人次；"
                     "牵头建立学院学生服务平台，覆盖学习辅导、心理咨询、就业指导 3 大模块"),
    "position_history": [
        {"role": "班长", "term": "2024.09-2025.05（1 学年）",
         "work": "组织班级主题班会 12 次、学习经验交流 5 场，建立班级学习互助群",
         "effect": "班级加权平均分由 84.2 提升至 87.6，所在班级获评 2024 年度校级先进班集体"},
        {"role": "学院学生会主席", "term": "2024.09-2025.05（1 学年）",
         "work": "组织大型活动 8 场，累计参与 2000+ 人次；建立学生服务平台覆盖 3 大模块",
         "effect": "推动学院学生会获评 2024 年度'校级优秀学生组织'"},
    ],
    "dorm_role": "宿舍长", "dorm_activity": "组织宿舍 6 次集体活动",
    "dorm_honor": "宿舍连续两学期获评'文明宿舍'",
    "interpersonal": "与同学相处融洽，曾帮助室友完成 1 次重要实验调试",
    "lifestyle": "生活中我注重勤俭节约，作息规律。", "life": "", "ending": "",
    "honors": [
        {"time": "2024.11", "name": "校级一等奖学金", "level": "校级（专业前 5%）", "issuer": "XX 大学"},
        {"time": "2024.10", "name": "校级优秀学生干部", "level": "校级", "issuer": "XX 大学"},
    ],
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="优秀学生 / 优秀团员申请书 docx 生成器（按 category 切换）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=("示例：\n  python build.py --data data.json --out output.docx\n  python build.py --demo --out demo.docx\n\n"
                "JSON 字段定义详见 SKILL.md 第四章信息采集清单。\n"
                "category: outstanding_student（优秀学生）/ outstanding_youth_league（优秀团员）"),
    )
    parser.add_argument("--data", type=str, default=None, help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True, help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true", help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档（默认优秀学生）")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
