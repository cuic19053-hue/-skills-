#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
入党申请书 docx 生成器

格式标准：A4，页边距上下 2.54cm 左右 2.5cm；标题黑体二号居中；
称呼"敬爱的党组织："顶格宋体小四全角冒号；正文宋体小四 1.5 倍行距首行缩进 2 字符；
"此致"另起一行空两格，"敬礼！"另起一行顶格；落款右对齐。

政治规范校验：必提 5 项要点（中国共产党/习近平新时代中国特色社会主义思想/党的二十大/共产主义/全心全意为人民服务），
指导思想 6 项顺序正确，查重风险检测（连续 50 字与党章原文重复则警告）。

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号
SIZE_XIAO_WU = Pt(9)        # 小五

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进），宋体小四"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_cizhi_paragraph(doc, text: str = "此致"):
    """'此致'另起一行，空两格（首行缩进 2 字符）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格（不缩进）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=False, line_spacing=1.5)


def add_blank_paragraph(doc):
    """空段落，用于段落间留白"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return p


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def count_chinese_chars(text: str) -> int:
    """统计中文字符数（用于字数控制）"""
    if not text:
        return 0
    return len(re.findall(r"[\u4e00-\u9fff]", str(text)))


def check_plagiarism_risk(text: str, party_constitution_fragments: List[str],
                          threshold: int = 50) -> List[str]:
    """查重风险检测：检查连续 N 字与党章原文/模板片段重复

    Args:
        text: 待检测的文本
        party_constitution_fragments: 党章原文/网络模板片段列表
        threshold: 连续重复字数阈值，默认 50

    Returns:
        触发查重警告的片段列表
    """
    warnings = []
    if not text:
        return warnings
    text_clean = re.sub(r"\s+", "", str(text))
    for fragment in party_constitution_fragments:
        if not fragment:
            continue
        frag_clean = re.sub(r"\s+", "", str(fragment))
        if len(frag_clean) < threshold:
            continue
        # 滑动窗口检测：取 fragment 中所有长度 threshold 的子串
        for i in range(len(frag_clean) - threshold + 1):
            window = frag_clean[i:i + threshold]
            if window in text_clean:
                warnings.append(
                    f"检测到连续 {threshold} 字与参考片段重复：{window[:20]}..."
                )
                break  # 每个 fragment 只报告一次
    return warnings


def check_political_terms(text: str) -> List[str]:
    """政治用语规范校验：必提要点 + 禁用简写检测

    Returns:
        警告信息列表
    """
    warnings = []
    if not text:
        return warnings
    text_str = str(text)

    # 必提要点（5 项缺一不可）
    required_terms = [
        "中国共产党",
        "习近平新时代中国特色社会主义思想",
        "党的二十大",
        "共产主义",
        "全心全意为人民服务",
    ]
    for term in required_terms:
        if term not in text_str:
            warnings.append(f"必提要点缺失：'{term}'")

    # 禁用简写检测：若完整表述存在，则其包含的子串不算简写
    # （如"习近平新时代中国特色社会主义思想"含"新时代中国特色社会主义思想"）
    skip_abbrs = set()
    full_to_abbr = [
        ("习近平新时代中国特色社会主义思想", "新时代中国特色社会主义思想"),
        ("'三个代表'重要思想", "三个代表"),
        ("\"三个代表\"重要思想", "三个代表"),
        ("\u201c三个代表\u201d重要思想", "三个代表"),
    ]
    for full, abbr in full_to_abbr:
        if full in text_str:
            skip_abbrs.add(abbr)

    forbidden_abbreviations = [
        ("习近平思想", "应为'习近平新时代中国特色社会主义思想'"),
        ("习思想", "应为'习近平新时代中国特色社会主义思想'"),
        ("党的20大", "应为'党的二十大'"),
        ("马列", "应为'马克思列宁主义'"),
        ("毛思想", "应为'毛泽东思想'"),
        ("邓理论", "应为'邓小平理论'"),
        ("科学观", "应为'科学发展观'"),
    ]
    if "新时代中国特色社会主义思想" not in skip_abbrs:
        forbidden_abbreviations.append(
            ("新时代中国特色社会主义思想", "应为'习近平新时代中国特色社会主义思想'（带'习近平'三字）"))
    if "三个代表" not in skip_abbrs:
        forbidden_abbreviations.append(
            ("三个代表", "应为'\"三个代表\"重要思想'（注意引号与'重要思想'四字）"))
    for abbr, suggestion in forbidden_abbreviations:
        if abbr in text_str:
            warnings.append(f"禁用简写'{abbr}'：{suggestion}")
    # "20大"检测
    for _ in re.finditer(r"20\s*大", text_str):
        warnings.append("禁用简写'20大'（应完整表述为'党的二十大'）")
        break
    # "二十大"检测（排除"党的二十大"上下文）
    for match in re.finditer(r"二十大", text_str):
        start = match.start()
        if start >= 2 and text_str[start - 2:start + 3] == "党的二十大":
            continue
        warnings.append("禁用简写'二十大'（应完整表述为'党的二十大'）")
        break

    # 指导思想 6 项顺序检测（兼容中英文引号）
    guideline_pattern = (
        r"马克思列宁主义.*?毛泽东思想.*?邓小平理论.*?"
        r"[\u201c\"']三个代表[\u201d\"']重要思想.*?科学发展观.*?习近平新时代中国特色社会主义思想"
    )
    if not re.search(guideline_pattern, text_str, re.DOTALL):
        # 检查 6 项是否齐全（兼容多种引号写法）
        checks = [
            ("马克思列宁主义", "马克思列宁主义"), ("毛泽东思想", "毛泽东思想"),
            ("邓小平理论", "邓小平理论"),
            ("'三个代表'重要思想", "三个代表重要思想"),
            ("\"三个代表\"重要思想", "三个代表重要思想"),
            ("\u201c三个代表\u201d重要思想", "三个代表重要思想"),
            ("科学发展观", "科学发展观"),
            ("习近平新时代中国特色社会主义思想", "习近平新时代中国特色社会主义思想"),
        ]
        required_keys = ["马克思列宁主义", "毛泽东思想", "邓小平理论",
                         "三个代表重要思想", "科学发展观", "习近平新时代中国特色社会主义思想"]
        found = {key for needle, key in checks if needle in text_str}
        missing = [k for k in required_keys if k not in found]
        if missing:
            warnings.append(f"指导思想漏项：{', '.join(missing)}")
        else:
            warnings.append("指导思想 6 项顺序错误（应为马列→毛→邓→三→科→习）")

    return warnings


# 党章原文高风险片段（用于查重检测，连续 50 字重复则警告）
# 注意：入党誓词必须一字不差地写入申请书，故不入此列表
PARTY_CONSTITUTION_FRAGMENTS = [
    "中国共产党是中国工人阶级的先锋队，同时是中国人民和中华民族的先锋队，"
    "是中国特色社会主义事业的领导核心，代表中国先进生产力的发展要求，"
    "代表中国先进文化的前进方向，代表中国最广大人民的根本利益。",
    "全心全意为人民服务。党除了工人阶级和最广大人民群众的利益，没有自己特殊的利益。",
    "坚持社会主义道路、坚持人民民主专政、坚持中国共产党的领导、"
    "坚持马克思列宁主义毛泽东思想这四项基本原则，是我们的立国之本。",
]

# 网络入党申请书模板高风险片段（连续 30 字重复则警告）
NETWORK_TEMPLATE_FRAGMENTS = [
    "我怀着十分激动和诚恳的心情，郑重向党组织提出申请",
    "作为一名当代大学生，我深知自己肩负的责任和使命",
    "中国共产党是中国工人阶级的先锋队",
    "我决心以实际行动争取早日入党",
]


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """入党申请书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []
        self.full_text_parts: List[str] = []  # 用于查重检测

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段，非列表返回空列表"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def add_para(self, text, indent=True):
        """添加正文段落并记录到全文缓存（用于查重检测）"""
        if text:
            self.full_text_parts.append(str(text))
        return add_body_paragraph(self.doc, text, indent=indent)

    # --------------------------------------------------------
    # 标题
    # --------------------------------------------------------

    def _add_title(self):
        """标题：黑体二号居中，固定为"入党申请书"5 字"""
        add_title(self.doc, "入党申请书")
        add_blank_paragraph(self.doc)

    # --------------------------------------------------------
    # 称呼
    # --------------------------------------------------------

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号"""
        salutation = self._get("salutation", default="敬爱的党组织：")
        add_salutation_paragraph(self.doc, salutation)

    # --------------------------------------------------------
    # 入党志愿表达段（200~300 字）
    # --------------------------------------------------------

    def _add_volunteer_statement(self):
        """入党志愿表达段：入党誓词 75 字 + 入党动机简述

        若用户提供 volunteer_statement 字段则直接使用，
        否则按入党誓词 + 自定义动机拼装。
        """
        volunteer_text = self._get("volunteer_statement", default="")
        if volunteer_text:
            self.add_para(volunteer_text)
            return

        # 入党誓词（固定 75 字，一字不差）
        oath = (
            "我志愿加入中国共产党，拥护党的纲领，遵守党的章程，履行党员义务，"
            "执行党的决定，严守党的纪律，保守党的秘密，对党忠诚，积极工作，"
            "为共产主义奋斗终身，随时准备为党和人民牺牲一切，永不叛党。"
        )

        # 入党动机（用户提供 motivation 字段或使用默认）
        motivation = self._get("motivation", default="")
        if not motivation:
            motivation = (
                "我之所以申请加入中国共产党，源于大学军训期间的一次触动。"
                "当时教官是退伍军人党员，在烈日下将军训物资一件件搬到学生宿舍，全程没有一句怨言。"
                "这件事让我开始思考：是什么样的信仰能让一个人愿意为他人默默付出？"
                "此后我开始阅读《习近平的七年知青岁月》，参加了学院党课学习，逐渐认识到党员不仅是一个身份，更是一份责任与担当。"
            )

        full_text = f"我郑重地向党组织提出申请：{oath}{motivation}"
        self.add_para(full_text)

    # --------------------------------------------------------
    # 对党的认识段（1000~1500 字，重点）
    # --------------------------------------------------------

    def _add_party_understanding(self):
        """对党的认识段：5 层结构（性质→宗旨→指导思想→历史→最终目标）

        若用户提供 party_understanding 字段则直接使用，
        否则按 5 层结构从子字段拼装。
        """
        understanding_text = self._get("party_understanding", default="")
        if understanding_text:
            self.add_para(understanding_text)
            return

        # 第 1 层：党的性质（200~300 字）
        nature = self._get("nature_text", default="")
        if not nature:
            nature = (
                "通过对党章的学习，我认识到中国共产党既是中国工人阶级的先锋队，"
                "也是中国人民和中华民族的先锋队，是中国特色社会主义事业的领导核心。"
                "这种'两个先锋队'的性质决定了党必须始终代表最广大人民的根本利益。"
                "从新民主主义革命到社会主义建设，从改革开放到新时代脱贫攻坚，"
                "党始终坚持'人民至上'——这是党能赢得人民群众衷心拥护的根本原因。"
                "我深刻认识到，党的'人民性'决定了它必须也必然为最广大人民的根本利益而奋斗，"
                "党始终走在时代前列，引领中国发展进步。"
            )
        self.add_para(nature)

        # 第 2 层：党的宗旨（200~300 字）
        purpose = self._get("purpose_text", default="")
        if not purpose:
            purpose = (
                "党的根本宗旨是全心全意为人民服务。党除了工人阶级和最广大人民群众的利益，"
                "没有自己特殊的利益。新时代以来，我感受最深的是党领导的脱贫攻坚战——"
                "8 年时间，近 1 亿农村贫困人口脱贫，832 个贫困县摘帽，这是人类减贫史上的奇迹。"
                "这场战役让我深刻理解了'全心全意'四个字的分量：党不是把'为人民服务'挂在嘴上，"
                "而是真真切切地解决人民群众的急难愁盼问题。从抗击新冠疫情到乡村振兴，"
                "党始终把人民生命安全与幸福生活放在第一位，这种'人民至上'的立场让我深受触动。"
            )
        self.add_para(purpose)

        # 第 3 层：党的指导思想（300~400 字，重点）
        guideline = self._get("guideline_text", default="")
        if not guideline:
            guideline = (
                "党的指导思想是马克思列宁主义、毛泽东思想、邓小平理论、'三个代表'重要思想、"
                "科学发展观、习近平新时代中国特色社会主义思想。这一思想体系是党在不同历史时期"
                "带领人民进行革命、建设、改革的智慧结晶。马克思列宁主义揭示了人类社会发展的普遍规律；"
                "毛泽东思想指导党取得了新民主主义革命胜利；邓小平理论开辟了中国特色社会主义道路；"
                "'三个代表'重要思想深化了对党的先进性的认识；科学发展观强调了以人为本、全面协调可持续发展。"
                "习近平新时代中国特色社会主义思想是当代中国马克思主义、二十一世纪马克思主义，"
                "是中华文化和中国精神的时代精华，实现了马克思主义中国化时代化新的飞跃。"
                "我深入学习党的二十大和二十届三中全会精神，努力用党的最新理论成果武装头脑。"
            )
        self.add_para(guideline)

        # 第 4 层：党的历史（200~300 字）
        history = self._get("party_history_text", default="")
        if not history:
            history = (
                "党的百年奋斗历程是一部不懈奋斗史、思想探索史、自身建设史。"
                "从 1921 年嘉兴南湖红船上 13 名代表宣告建党，到领导人民完成新民主主义革命、建立新中国；"
                "从社会主义改造确立基本制度，到改革开放让中国大踏步赶上时代；再到新时代取得脱贫攻坚全面胜利、"
                "全面建成小康社会——党的每一步都印证了'为中国人民谋幸福、为中华民族谋复兴'的初心使命。"
                "党能始终赢得人民拥护，根本在于党始终坚持真理、修正错误，始终保持先进性和纯洁性。"
            )
        self.add_para(history)

        # 第 5 层：党的最终目标（100~200 字）
        final_goal = self._get("final_goal_text", default="")
        if not final_goal:
            final_goal = (
                "党的最高理想和最终目标是实现共产主义。我认识到，共产主义不是空想，"
                "而是马克思、恩格斯基于对人类社会发展规律的科学分析得出的必然结论。"
                "虽然实现共产主义需要一代又一代人的长期奋斗，但作为共产党人，必须坚定这个远大理想。"
                "我会把共产主义远大理想与中国特色社会主义共同理想统一起来，立足本职学习工作，为推进中国式现代化贡献青春力量。"
            )
        self.add_para(final_goal)

    # --------------------------------------------------------
    # 个人经历与思想变化段（1000~1500 字，重点）
    # --------------------------------------------------------

    def _add_personal_experience(self):
        """个人经历与思想变化段：4 阶段时间轴

        阶段 1：家庭影响与早期启蒙（200~300 字）
        阶段 2：入队入团经历（200~300 字）
        阶段 3：大学期间思想变化（500~700 字，重点）
        阶段 4：实践行动与当前认识（200~300 字）
        """
        experience_text = self._get("personal_experience", default="")
        if experience_text:
            self.add_para(experience_text)
            return

        # 阶段 1：家庭影响与早期启蒙
        family_intro = self._get("family_intro", default="")
        if not family_intro:
            family_intro = self._build_family_intro()
        self.add_para(family_intro)

        # 阶段 2：入队入团经历
        youth_team = self._get("youth_team_text", default="")
        if not youth_team:
            youth_team = self._build_youth_team_text()
        self.add_para(youth_team)

        # 阶段 3：大学期间思想变化
        college_ideology = self._get("college_ideology_text", default="")
        if not college_ideology:
            college_ideology = self._build_college_ideology_text()
        self.add_para(college_ideology)

        # 阶段 4：实践行动与当前认识
        practice = self._get("practice_text", default="")
        if not practice:
            practice = self._build_practice_text()
        self.add_para(practice)

    def _build_family_intro(self) -> str:
        """构建家庭影响与早期启蒙段"""
        family_members = self._get_list("family_members")
        default_intro = (
            "我出生在一个普通家庭，父母都是勤恳踏实的普通人。他们虽然不是党员，"
            "但常对我说'做人要本分、做事要踏实'。这种朴素的家风让我从小形成了勤勉正直的品格，"
            "也为我后来理解党员'吃苦在前、享受在后'的精神打下了基础。"
        )
        if not family_members:
            return default_intro

        father = next((m for m in family_members if isinstance(m, dict)
                       and str(m.get("relation", "")).startswith("父亲")), None)
        if father and str(father.get("political_status", "")) == "中共党员":
            father_name = father.get("name", "父亲")
            father_unit = father.get("work_unit", "")
            return (
                f"我的{father_name}是一名中共党员，{father_unit}工作。"
                "从小父亲就给我讲焦裕禄、谷文昌、张富清等优秀党员的故事。家中书架上摆着《毛泽东选集》"
                "《邓小平文选》和一套《中国共产党简史》。这些早期的耳濡目染，让我对党产生了朴素的向往。"
                "父亲常说'党员就是要在关键时刻站得出来、危难关头豁得出来'，这句话深深印在我的心里。"
            )
        return default_intro

    def _build_youth_team_text(self) -> str:
        """构建入队入团经历段"""
        team_join_date = self._get("team_join_date", default="2015 年 5 月")
        league_join_date = self._get("league_join_date", default="2018 年 5 月")
        return (
            f"{team_join_date}，我在小学加入了中国少年先锋队。那天我戴着红领巾站在国旗下宣誓，老师告诉我们，"
            "红领巾是国旗的一角，是革命先烈的鲜血染成的。当时的我虽然不完全理解这句话的含义，但一种庄严感油然而生。"
            f"{league_join_date}，我在初中加入了中国共产主义青年团。入团仪式上，团支书带领我们学习团章，"
            "让我第一次系统了解到共青团是党的助手和后备军。从此，'听党话、跟党走'成为我的行动准则。"
        )

    def _build_college_ideology_text(self) -> str:
        """构建大学期间思想变化段（500~700 字，重点）"""
        military_touch = self._get("military_touch", default="")
        book_reading = self._get("book_reading", default="")
        party_class = self._get("party_class", default="")
        if not military_touch:
            military_touch = (
                "大一军训期间，我的教官是一名退伍军人党员。他每天最早到训练场、最晚离开，"
                "将军训物资一件件搬到学生宿舍，全程没有一句怨言。当时我不理解，后来才意识到，"
                "这正是党员'吃苦在前、享受在后'的具象化表达。"
            )
        if not book_reading:
            book_reading = (
                "大一下学期，我在思政课'中国近现代史纲要'上系统学习了党的奋斗史，认识到党是历史的选择、人民的选择。"
                "同一时期，我阅读了《习近平的七年知青岁月》，书中青年习近平在梁家河打坝淤地、建沼气池的经历让我深受震撼——"
                "一名 15 岁的北京知青能扎根陕北农村 7 年，与群众同吃同住同劳动，这种'扎根群众'的精神，正是党能始终与人民血脉相连的根源。"
            )
        if not party_class:
            party_class = (
                "大二上学期，我参加了学院分党校第 8 期入党积极分子培训班，系统学习了党章、党史、党的纪律等内容，并顺利结业。"
                "通过培训，我从理论上弄清了'为什么入党'这个根本问题：入党不是为了个人利益，而是为了更好地为人民服务、为共产主义事业奋斗。"
            )
        return military_touch + book_reading + party_class

    def _build_practice_text(self) -> str:
        """构建实践行动与当前认识段（200~300 字）"""
        practice_detail = self._get("practice_detail", default="")
        if not practice_detail:
            practice_detail = (
                "担任班长期间，我组织班级'一对一'帮扶活动，服务同学 30 余人次；"
                "参与暑期三下乡社会实践 2 次，累计支教 16 课时；志愿服务时长 120 小时。"
                "2024 年 10 月，我参加主题党日活动'红色教育基地走访'，赴西柏坡参观学习，撰写调研报告 1 份（约 3000 字）。"
            )
        current_cognition = self._get("current_cognition", default="")
        if not current_cognition:
            current_cognition = (
                "通过这些实践，我更加深刻地认识到：党员的身份不是光环，而是责任；不是特权，而是奉献。"
                "当前，我已具备成为一名入党积极分子的思想基础，恳请党组织在实践中考验我。"
            )
        return "在向党组织靠拢的过程中，我努力以实际行动践行党员标准。" + practice_detail + current_cognition

    # --------------------------------------------------------
    # 个人不足与今后努力方向段（500~800 字）
    # --------------------------------------------------------

    def _add_shortcomings(self):
        """个人不足与今后努力方向段：2~3 个真实不足 + 结尾表达"""
        shortcomings_text = self._get("shortcomings_text", default="")
        if shortcomings_text:
            self.add_para(shortcomings_text)
            return

        shortcomings = self._get_list("shortcomings")
        if not shortcomings:
            # 默认 3 个真实不足
            shortcomings = [
                {"description": "理论学习系统性不够",
                 "manifestation": (
                     "虽然通读了党章，但对党史中的某些重要节点"
                     "（如十一届三中全会前后的历史细节、延安整风运动的来龙去脉）理解不深；"
                     "对马克思主义经典著作（如《共产党宣言》《资本论》第一卷）只读过简写本，未读全本。"
                 ),
                 "improvement": (
                     "今后我将系统学习《中国共产党简史》《中国共产党的一百年》"
                     "《马克思恩格斯选集》等权威著作，每月撰写 1 篇读书笔记，并向党组织汇报。"
                 )},
                {"description": "实践能力仍需加强",
                 "manifestation": (
                     "参加党课学习较多，但参与党员志愿服务的次数偏少，本学期仅参加 2 次；"
                     "与群众联系不够紧密，对班级同学的急难愁盼问题了解不够深入。"
                 ),
                 "improvement": (
                     "今后我将主动报名学校组织的'党员先锋岗'等活动，每月至少参加 1 次志愿服务；"
                     "定期走访宿舍，了解同学学习生活情况，及时向辅导员反馈。"
                 )},
                {"description": "与同学沟通方式有待改进",
                 "manifestation": (
                     "在担任班长期间，处理宿舍矛盾时方式较为生硬，曾因直接批评导致同学情绪反弹。"
                 ),
                 "improvement": (
                     "今后我将学习更圆融的沟通技巧，多倾听同学意见；遇到矛盾先了解双方诉求，"
                     "再寻找折中方案；定期参加学校组织的学生干部沟通能力培训。"
                 )},
            ]

        for idx, item in enumerate(shortcomings, start=1):
            if isinstance(item, dict):
                desc = item.get("description", "")
                manifest = item.get("manifestation", "")
                improve = item.get("improvement", "")
                num_word = ["一", "二", "三", "四", "五"][idx - 1] if idx <= 5 else str(idx)
                para_text = f"{num_word}是{desc}。{manifest}{improve}"
            else:
                para_text = str(item)
            self.add_para(para_text)

        # 结尾表达
        ending_statement = self._get("shortcomings_ending", default="")
        if not ending_statement:
            ending_statement = (
                "以上是我对自己的不全面剖析，恳请党组织批评指正。无论党组织是否批准我的申请，"
                "我都将继续以党员标准严格要求自己，自觉接受党组织考验，争取早日成为一名合格的中国共产党党员。"
            )
        self.add_para(ending_statement)

    # --------------------------------------------------------
    # 结尾"请党组织在实践中考验我！此致 敬礼！"
    # --------------------------------------------------------

    def _add_ending(self):
        """结尾：请党组织在实践中考验我！+ 此致 + 敬礼！"""
        ending_line = self._get("ending_line", default="请党组织在实践中考验我！")
        self.add_para(ending_line)
        add_cizhi_paragraph(self.doc, "此致")
        add_jingli_paragraph(self.doc, "敬礼！")

    # --------------------------------------------------------
    # 落款（右对齐）
    # --------------------------------------------------------

    def _add_signature(self):
        """落款：申请人 + 日期 + 党支部全称，右对齐"""
        add_blank_paragraph(self.doc)
        name = self._get("name", default="申请人")
        submit_date = self._get("submit_date", default="")
        party_branch = self._get("party_branch", default="")
        add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if submit_date:
            add_right_aligned_paragraph(self.doc, submit_date)
        if party_branch:
            add_right_aligned_paragraph(self.doc, f"递交：{party_branch}")

    # --------------------------------------------------------
    # 主构建方法
    # --------------------------------------------------------

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/4 段正文/结尾/落款，生成 docx"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_volunteer_statement()
            self._add_party_understanding()
            self._add_personal_experience()
            self._add_shortcomings()
            self._add_ending()
            self._add_signature()
            self._post_build_checks()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 入党申请书已生成：{output_path}")
        return str(output_path)

    def _post_build_checks(self):
        """构建后检查：政治用语规范 + 查重风险检测 + 字数检查"""
        full_text = "".join(self.full_text_parts)

        # 政治用语规范检查
        for w in check_political_terms(full_text):
            self.warnings.append(f"[政治规范] {w}")

        # 查重风险检测：与党章原文比对（阈值 50 字）
        for w in check_plagiarism_risk(
            full_text, PARTY_CONSTITUTION_FRAGMENTS, threshold=50
        ):
            self.warnings.append(f"[查重风险-党章原文] {w}")

        # 查重风险检测：与网络模板比对（阈值 30 字）
        for w in check_plagiarism_risk(
            full_text, NETWORK_TEMPLATE_FRAGMENTS, threshold=30
        ):
            self.warnings.append(f"[查重风险-网络模板] {w}")

        # 字数检查
        char_count = count_chinese_chars(full_text)
        if char_count < 3500:
            self.warnings.append(f"[字数] 全文仅 {char_count} 字，建议 3500~4500 字")
        elif char_count > 5000:
            self.warnings.append(f"[字数] 全文 {char_count} 字偏多，建议压缩至 4500 字以内")

        if self.warnings:
            print("⚠️ 构建后检查警告：", file=sys.stderr)
            for w in self.warnings:
                print(f"  - {w}", file=sys.stderr)

    # --------------------------------------------------------
    # 数据校验
    # --------------------------------------------------------

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []

        # P0 必采字段
        p0_fields = [
            ("name", "申请人姓名"), ("college", "学院"), ("major", "专业"),
            ("grade", "年级"), ("submit_date", "递交日期"),
            ("party_branch", "递交党支部全称"),
        ]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        # 家庭主要成员校验
        family_members = self._get_list("family_members")
        if not family_members:
            warnings.append("缺少 家庭主要成员（family_members），将使用默认家庭背景")
        else:
            has_father = any(
                isinstance(m, dict) and str(m.get("relation", "")).startswith("父亲")
                for m in family_members
            )
            has_mother = any(
                isinstance(m, dict) and str(m.get("relation", "")).startswith("母亲")
                for m in family_members
            )
            if not has_father:
                warnings.append("家庭主要成员缺少 父亲 信息")
            if not has_mother:
                warnings.append("家庭主要成员缺少 母亲 信息")
            for m in family_members:
                if isinstance(m, dict) and not m.get("political_status"):
                    warnings.append(f"家庭主要成员 {m.get('relation', '?')} 缺少政治面貌")

        # 个人不足校验
        shortcomings = self._get_list("shortcomings")
        if shortcomings and len(shortcomings) < 2:
            warnings.append(f"个人不足仅 {len(shortcomings)} 个，建议 2~3 个真实不足")

        # 入党动机校验
        motivation = self._get("motivation", default="")
        if motivation and ("求职" in motivation or "就业" in motivation or "公务员" in motivation):
            warnings.append("[政治红线] 入党动机出现'求职/就业/公务员'字眼，动机不端正风险，请用户重新表述")

        # 个人不足假缺点检测（党支部一眼识破，会要求重写）
        fake_shortcomings_keywords = [
            "工作太投入", "追求完美", "学习太刻苦",
            "责任心太强", "为人太直率", "事必躬亲",
        ]
        shortcomings_text = self._get("shortcomings_text", default="")
        for item in self._get_list("shortcomings"):
            if isinstance(item, dict):
                shortcomings_text += str(item.get("description", "")) + str(item.get("manifestation", ""))
        for kw in fake_shortcomings_keywords:
            if kw in shortcomings_text:
                warnings.append(
                    f"[政治红线] 个人不足出现假缺点'{kw}'，"
                    "党支部会要求重写，请用户重新表述"
                )
                break

        self.warnings.extend(warnings)
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据
# ============================================================

DEFAULT_DATA = {
    "name": "张明", "student_id": "2022123456", "gender": "男",
    "birth_date": "2004 年 5 月", "college": "计算机科学与技术学院",
    "major": "计算机科学与技术", "grade": "2022 级 大三", "class_name": "计科 2201 班",
    "family_origin": "干部", "identity": "学生", "political_status": "共青团员",
    "salutation": "敬爱的党组织：", "submit_date": "2025 年 6 月 15 日",
    "party_branch": "计算机科学与技术学院本科生第一党支部",
    "family_members": [
        {"relation": "父亲", "name": "张建国", "political_status": "中共党员",
         "work_unit": "XX 县教育局", "position": "副局长"},
        {"relation": "母亲", "name": "李秀英", "political_status": "群众",
         "work_unit": "XX 县人民医院", "position": "护士长"},
    ],
    "team_join_date": "2015 年 5 月", "league_join_date": "2018 年 5 月",
    "motivation": (
        "我之所以申请加入中国共产党，源于大学军训期间的一次触动。"
        "当时教官是退伍军人党员，在烈日下将军训物资一件件搬到学生宿舍，全程没有一句怨言。"
        "这件事让我开始思考：是什么样的信仰能让一个人愿意为他人默默付出？"
        "此后我开始阅读《习近平的七年知青岁月》，参加了学院党课学习，"
        "逐渐认识到党员不仅是一个身份，更是一份责任与担当。"
    ),
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="入党申请书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\nJSON 字段定义详见 SKILL.md 第三章信息采集清单。\n"
            "政治规范校验：必提 5 项要点 + 指导思想 6 项顺序 + 查重风险检测。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True, help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
