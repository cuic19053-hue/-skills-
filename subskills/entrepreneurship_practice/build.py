#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
大学生创新创业训练计划-创业实践项目申报书 docx 生成器

按统一格式标准生成 Word 文档：
- A4 纸张，页边距上下 2.54cm 左右 2.5cm
- 正文：宋体小四，1.5 倍行距，首行缩进 2 字符
- 一级标题：黑体三号，居中
- 二级标题：黑体小三，左对齐
- 三级标题：宋体四号加粗
- 表格：宋体五号，居中

14 栏目结构（创业实践专属）：
封面 / 基本信息 / 项目简介 / 公司概况 / 产品服务与市场 /
商业模式与运营现状 / 财务数据与预测 / 团队组织 / 融资计划 /
风险与应对 / 社会效益 / 预期成果 / 经费预算 / 签字栏

要求项目已注册公司或运营 6 个月以上，含工商信息表、股权结构表、
3 年财务报表、运营 KPI 表、融资历史表。

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第十一章。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 字体与格式常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"

SIZE_ER = Pt(22)
SIZE_SAN = Pt(16)
SIZE_XIAO_SAN = Pt(15)
SIZE_SI = Pt(14)
SIZE_XIAO_SI = Pt(12)
SIZE_WU = Pt(10.5)

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5


# ============================================================
# 工具函数（中英文同步设置 eastAsia/ascii/hAnsi）
# ============================================================

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc, text: str, font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI, bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True, line_spacing: float = 1.5,
    space_before: float = 0, space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
        space_before=6, space_after=6)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=indent,
        line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None,
                        last_col_left: bool = True):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if (last_col_left and j == len(row) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False, alignment=align)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ============================================================
# ApplicationDocBuilder 主类
# ============================================================

class ApplicationDocBuilder:
    """大创-创业实践项目申报书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def add_h1(self, text):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text):
        return add_heading_level2(self.doc, text)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_table(self, headers, rows, col_widths=None, last_col_left=True):
        return add_table_from_data(self.doc, headers, rows, col_widths, last_col_left)

    def add_page_break(self):
        add_page_break(self.doc)

    # --------------------------------------------------------
    # 封面
    # --------------------------------------------------------

    def _add_cover(self):
        """封面：黑体二号标题 + 副标题 + 5 行下划线信息"""
        for _ in range(2):
            self.doc.add_paragraph()
        add_paragraph_with_format(
            self.doc, "国家级大学生创新创业训练计划项目申报书",
            font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_before=12, space_after=12)
        add_paragraph_with_format(
            self.doc, f"（{self._get('project_type', default='创业实践项目')}）",
            font_name=FONT_HEI, font_size=SIZE_SAN, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_after=24)
        for _ in range(2):
            self.doc.add_paragraph()
        info_items = [
            ("项目名称", self._get("project_name")),
            ("项目负责人", self._get("leader_name")),
            ("指导教师", self._get("advisor_name")),
            ("所在学院", self._get("college")),
            ("申报日期", self._get("apply_date")),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI, font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True
        self.add_page_break()

    # --------------------------------------------------------
    # 基本信息
    # --------------------------------------------------------

    def _add_basic_info_table(self):
        """一、基本信息表（11 行 2 列，含公司信息）"""
        self.add_h1("一、基本信息")
        basic = self._get("basic_info", default={})
        if not isinstance(basic, dict):
            basic = {}
        budget_str = str(basic.get("budget", self._get("budget_total", "")))
        if budget_str and not budget_str.endswith("元"):
            budget_str = f"{budget_str} 元"
        company = self._get("company_overview", default={})
        if not isinstance(company, dict):
            company = {}
        company_name = basic.get("company_name", company.get("company_name", ""))
        credit_code = basic.get("credit_code", company.get("credit_code", ""))
        rows = [
            ["项目名称", basic.get("project_name", self._get("project_name"))],
            ["项目类型", basic.get("project_type", self._get("project_type", "创业实践项目"))],
            ["项目来源", basic.get("project_source", "A 学生自主选题")],
            ["所属行业", basic.get("industry", self._get("industry", ""))],
            ["公司名称", company_name],
            ["统一社会信用代码", credit_code],
            ["起止时间", basic.get("duration", "")],
            ["申请经费", budget_str],
            ["负责人", basic.get("leader_info",
                                  f"{self._get('leader_name')} / {self._get('leader_id')} / "
                                  f"{self._get('leader_major')} / {self._get('leader_grade')} / "
                                  f"{self._get('leader_phone')}")],
            ["团队成员", basic.get("team_members", "")],
            ["指导教师", basic.get("advisor_info",
                                    f"{self._get('advisor_name')} / {self._get('advisor_title')} / "
                                    f"{self._get('advisor_research')}")],
        ]
        self.add_table(["项目", "内容"], rows, col_widths=[4.5, 11.5])

    # --------------------------------------------------------
    # 项目简介
    # --------------------------------------------------------

    def _add_abstract(self):
        """二、项目简介（300~500 字）"""
        self.add_h1("二、项目简介")
        abstract = self._get("abstract", default="")
        self.add_para(abstract if abstract else
                      "（请填写项目简介，300~500 字，按 4 句结构撰写：公司概况+做什么 / "
                      "运营数据+商业模式 / 融资现状+团队 / 本轮目标。）")

    # --------------------------------------------------------
    # 公司概况（含工商信息 + 股权结构）
    # --------------------------------------------------------

    def _add_company_overview(self):
        """三、公司概况（500~800 字，含工商信息表 + 股权结构表 + 治理）"""
        self.add_h1("三、公司概况")
        company = self._get("company_overview", default={})
        if not isinstance(company, dict):
            company = {}

        self.add_h2("（一）工商信息")
        biz = company.get("business_info", [])
        if biz:
            rows = [[item.get("field", ""), item.get("value", "")]
                    for item in biz if isinstance(item, dict)]
            self.add_table(["项目", "内容"], rows, col_widths=[4.5, 11.5])
        else:
            self.add_para("（请填写工商信息表，10 项：公司全称 / 统一社会信用代码 / "
                          "注册日期 / 注册资本 / 企业类型 / 法定代表人 / 注册地址 / "
                          "经营范围 / 经营状态 / 实际运营月数。）")

        self.add_h2("（二）股权结构")
        shareholders = company.get("shareholders", [])
        if shareholders:
            rows = [[s.get("name", ""), s.get("amount", ""), s.get("ratio", ""), s.get("form", "")]
                    for s in shareholders if isinstance(s, dict)]
            self.add_table(["股东", "出资额（万元）", "出资比例", "出资形式"],
                           rows, col_widths=[4.0, 3.5, 3.5, 5.0])
        else:
            self.add_para("（请填写股权结构表：股东姓名 / 出资额 / 出资比例 / 出资形式。"
                          "如预留期权池需单列一行。）")

        self.add_h2("（三）公司治理与合规")
        governance = company.get("governance", "")
        self.add_para(governance if governance else
                      "（请填写公司治理：股东会构成、董事会构成、监事设置、"
                      "税务合规（增值税/企业所得税/个税申报）、社保缴纳情况、"
                      "知识产权归属（软著/商标/域名归属公司）。）")

    # --------------------------------------------------------
    # 产品服务与市场
    # --------------------------------------------------------

    def _add_product_market(self):
        """四、产品服务与市场（800~1200 字，4 子节）"""
        self.add_h1("四、产品服务与市场")
        pm = self._get("product_market", default={})
        if not isinstance(pm, dict):
            pm = {}

        self.add_h2("（一）产品形态")
        features = pm.get("features", [])
        if isinstance(features, str):
            features = [features]
        if features:
            for i, f in enumerate(features, 1):
                self.add_para(f"{i}. {f}")
        else:
            self.add_para("（请填写 MVP 功能清单，5~8 个核心功能，"
                          "每个含名称+解决什么问题+关键指标。）")

        self.add_h2("（二）市场分析")
        market = pm.get("market_analysis", "")
        self.add_para(market if market else
                      "（请填写市场分析：TAM/SAM/SOM 三级测算 + 3~5 个直接竞品 5 列对比表 + "
                      "用户画像。所有数据标注来源。）")

        self.add_h2("（三）技术实现")
        tech = pm.get("tech_impl", "")
        self.add_para(tech if tech else
                      "（请填写技术架构（前端/后端/数据库）+ 关键第三方服务 + 研发投入估算。）")

        self.add_h2("（四）知识产权")
        ip = pm.get("ip_status", "")
        self.add_para(ip if ip else
                      "（请填写已申请/已授权的软件著作权、专利、商标清单，"
                      "归属公司而非个人。）")

    # --------------------------------------------------------
    # 商业模式与运营现状（含 KPI 表）
    # --------------------------------------------------------

    def _add_business_operation(self):
        """五、商业模式与运营现状（800~1200 字，3 子节 + KPI 表）"""
        self.add_h1("五、商业模式与运营现状")
        bo = self._get("business_operation", default={})
        if not isinstance(bo, dict):
            bo = {}

        self.add_h2("（一）商业模式")
        model = bo.get("model", "")
        self.add_para(model if model else
                      "（请填写盈利模式 1~3 种 + 定价策略 + 渠道通路。"
                      "每种含定价+计算依据+营收占比。）")

        self.add_h2("（二）运营 KPI")
        kpi_table = bo.get("kpi_table", [])
        if kpi_table:
            rows = [[k.get("name", ""), k.get("current", ""), k.get("yoy", "")]
                    for k in kpi_table if isinstance(k, dict)]
            self.add_table(["KPI", "当前值", "同比"], rows,
                           col_widths=[5.5, 5.5, 5.0], last_col_left=False)
        else:
            self.add_para("（请填写运营 KPI 表，8 项齐全：累计用户/月活/MRR/ARR/"
                          "毛利率/复购率/CAC/LTV，含同比变化。）")

        self.add_h2("（三）运营节奏与里程碑")
        milestones = bo.get("milestones", [])
        if milestones:
            rows = [[m.get("time", ""), m.get("event", ""), m.get("status", "")]
                    for m in milestones if isinstance(m, dict)]
            self.add_table(["时间", "里程碑事件", "完成状态"], rows,
                           col_widths=[3.5, 8.5, 4.0])
        else:
            self.add_para("（请填写按季度里程碑表：已完成里程碑 + 未来 4 季度计划，"
                          "每项含时间+事件+状态。）")

    # --------------------------------------------------------
    # 财务数据与预测（3 年利润表 + 现金流 + 资产负债）
    # --------------------------------------------------------

    def _add_financial_data(self):
        """六、财务数据与预测（3 年 3 表 + Burn/Runway）"""
        self.add_h1("六、财务数据与预测")
        fin = self._get("financial_data", default={})
        if not isinstance(fin, dict):
            fin = {}

        self.add_h2("（一）利润表（3 年）")
        inc = fin.get("income_statement", [])
        if inc:
            rows = [[r.get("item", ""), str(r.get("y1", "")), str(r.get("y2", "")), str(r.get("y3", ""))]
                    for r in inc if isinstance(r, dict)]
            self.add_table(["财务项目", "Year1", "Year2", "Year3"], rows,
                           col_widths=[5.0, 3.5, 3.5, 3.5], last_col_left=False)
        else:
            self.add_para("（请填写 3 年利润表，9 行：营收/营业成本/毛利率/"
                          "销售/管理/研发/财务费用/营业利润/净利润。）")

        self.add_h2("（二）现金流量表要点")
        cf = fin.get("cash_flow", [])
        if cf:
            rows = [[r.get("item", ""), str(r.get("y1", "")), str(r.get("y2", "")), str(r.get("y3", ""))]
                    for r in cf if isinstance(r, dict)]
            self.add_table(["现金流项目", "Year1", "Year2", "Year3"], rows,
                           col_widths=[5.0, 3.5, 3.5, 3.5], last_col_left=False)
        else:
            self.add_para("（请填写 3 年现金流量表 4 行：经营性/投资性/筹资性现金流净额 + 期末现金余额。）")

        self.add_h2("（三）资产负债表要点")
        bs = fin.get("balance_sheet", [])
        if bs:
            rows = [[r.get("item", ""), str(r.get("y1", "")), str(r.get("y2", "")), str(r.get("y3", ""))]
                    for r in bs if isinstance(r, dict)]
            self.add_table(["资产负债项目", "Year1 末", "Year2 末", "Year3 末"], rows,
                           col_widths=[5.0, 3.5, 3.5, 3.5], last_col_left=False)
        else:
            self.add_para("（请填写 3 年资产负债表 8 行：货币资金/应收/固定资产/无形资产/"
                          "资产合计/短期借款/实收资本/未分配利润。）")

        self.add_h2("（四）Burn rate / Runway / 盈亏平衡点")
        br = fin.get("burn_runway", "")
        self.add_para(br if br else
                      "（请填写：Burn rate 8 万/月、Runway 12 个月、现金余额 96 万、"
                      "盈亏平衡点 Year2 第 8 个月月营收达 15 万元时毛利率覆盖固定成本。）")

    def _add_team_organization(self):
        """七、团队组织（400~600 字 + 表格 + 组织架构）"""
        self.add_h1("七、团队组织")
        team = self._get("team_organization", default={})
        if not isinstance(team, dict):
            team = {}

        self.add_h2("（一）团队构成")
        members = team.get("members", [])
        if members:
            rows = [[m.get("name", ""), m.get("id", ""), m.get("major", ""),
                     m.get("role", ""), m.get("exp", "")]
                    for m in members if isinstance(m, dict)]
            self.add_table(["姓名", "学号", "专业年级", "公司职务", "相关经历"], rows,
                           col_widths=[1.8, 2.5, 3.0, 3.0, 5.7])
        else:
            self.add_para("（请填写团队表：姓名/学号/专业年级/公司职务/相关经历，"
                          "4~6 人，商科+技术复合。）")

        self.add_h2("（二）组织架构")
        org = team.get("org_structure", "")
        self.add_para(org if org else
                      "（请填写组织架构：CEO 下设 CTO/CMO/CFO/COO 等部门，"
                      "全职 X 人 + 兼职 X 人 + 外部顾问 X 人。）")

        self.add_h2("（三）薪资与社保")
        salary = team.get("salary", "")
        self.add_para(salary if salary else
                      "（请填写薪资：全职 X 人月薪 X 元（基本生活费）、兼职 X 人按项目计酬、"
                      "已为全职员工缴纳社保。）")

    def _add_financing_plan(self):
        """八、融资计划（500~800 字 + 融资历史表）"""
        self.add_h1("八、融资计划")
        fp = self._get("financing_plan", default={})
        if not isinstance(fp, dict):
            fp = {}

        self.add_h2("（一）融资历史")
        history = fp.get("history", [])
        if history:
            rows = [[h.get("round", ""), h.get("time", ""), h.get("amount", ""),
                     h.get("valuation", ""), h.get("equity", ""), h.get("investor", "")]
                    for h in history if isinstance(h, dict)]
            self.add_table(["轮次", "时间", "金额", "投后估值", "出让股权", "投资方"], rows,
                           col_widths=[2.0, 2.5, 2.0, 2.5, 2.0, 5.0])
        else:
            self.add_para("（请填写融资历史表：轮次/时间/金额/投后估值/出让股权/投资方。"
                          "无历史融资可写『无』。）")

        self.add_h2("（二）本轮融资计划")
        current = fp.get("current_round", "")
        self.add_para(current if current else
                      "（请填写本轮融资：金额 + 估值方法（DCF/可比公司/成本加成）+ "
                      "出让股权 + 资金用途分解（4~6 类，含百分比+金额+用途）。）")

        self.add_h2("（三）Term sheet 关键条款")
        ts = fp.get("term_sheet", "")
        self.add_para(ts if ts else
                      "（请填写 Term sheet 关键条款：投前/投后估值、清算优先权、"
                      "反稀释条款、董事会席位、回购条款、业绩对赌、一票否决权。）")

    def _add_risk_response(self):
        """九、风险与应对（400~600 字 + 5 类风险表）"""
        self.add_h1("九、风险与应对")
        risks = self._get("risk_response", default=[])
        if isinstance(risks, str):
            risks = [risks]
        if risks:
            rows = [[r.get("type", ""), r.get("risk", ""), r.get("prob", ""),
                     r.get("impact", ""), r.get("measure", "")]
                    for r in risks if isinstance(r, dict)]
            self.add_table(["风险类型", "具体风险", "概率", "影响", "应对措施"], rows,
                           col_widths=[2.2, 3.8, 1.5, 1.5, 7.0])
        else:
            self.add_para("（请填写风险分析表，5 类风险齐全：市场/技术/运营/财务/法务，"
                          "每类含具体风险+概率+影响+应对措施。法务风险为创业实践专属。）")

    def _add_social_benefit(self):
        """十、社会效益（300~500 字）"""
        self.add_h1("十、社会效益")
        sb = self._get("social_benefit", "")
        self.add_para(sb if sb else
                      "（请填写社会效益，3 个角度：就业带动（已雇佣 X 人，预期新增 Y 人）、"
                      "校园赋能（覆盖 X 所高校、服务 Y 学生）、ESG 价值（环保/节能/公益贡献）。）")

    def _add_expected_results(self):
        """十一、预期成果（创业实践型：营收/用户/融资额/就业）"""
        self.add_h1("十一、预期成果")
        outcomes = self._get("expected_outcomes", default=[])
        if isinstance(outcomes, str):
            outcomes = [outcomes]
        if outcomes:
            for o in outcomes:
                self.add_para(f"• {o}", indent=False)
        else:
            self.add_para("（请填写预期成果，以真实数据为主：营收目标（18 个月累计 X 万）、"
                          "用户目标（X 万）、融资目标（Pre-A 轮 X 万）、就业岗位（新增 X 个）、"
                          "知识产权（软著 X 项 + 商标 X 项）、合作高校（X 所）。）", indent=False)

    def _add_budget(self):
        """十二、经费预算（3 列表格）"""
        self.add_h1("十二、经费预算")
        items = self._get("budget_items", default=[])
        if not items:
            self.add_para("（请填写经费预算，6~8 类标准科目：服务器与域名/用户调研/"
                          "推广获客/物料与会议/印刷复印/知识产权申请/团队培训。"
                          "每项金额非整数，附计算依据。）")
            return
        rows = []
        total = 0
        for b in items:
            if not isinstance(b, dict):
                continue
            try:
                amount_num = int(str(b.get("amount", "0")))
            except ValueError:
                amount_num = 0
            total += amount_num
            rows.append([b.get("item", ""), f"{amount_num} 元", b.get("basis", "")])
        rows.append(["合计", f"{total} 元", ""])
        self.add_table(["预算科目", "金额", "计算依据"], rows, col_widths=[3.5, 3.0, 9.5])

    def _add_signature_section(self):
        """十三/十四、指导教师意见、学院评审意见"""
        sections = [("十三、指导教师意见",
                     "指导教师签字：____________________    日期：______年____月____日")]
        sections.append(("十四、学院评审意见",
                         "学院盖章：____________________    日期：______年____月____日"))
        if self._get("include_school_approval", default=False):
            sections.append(("十五、学校审批意见",
                             "学校盖章：____________________    日期：______年____月____日"))
        for title, line in sections:
            self.add_h1(title)
            for _ in range(6):
                self.doc.add_paragraph()
            self.add_para(line, indent=False)

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 14 栏目，生成 docx。返回实际保存路径。"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_cover()
            self._add_basic_info_table()
            self._add_abstract()
            self._add_company_overview()
            self._add_product_market()
            self._add_business_operation()
            self._add_financial_data()
            self._add_team_organization()
            self._add_financing_plan()
            self._add_risk_response()
            self._add_social_benefit()
            self._add_expected_results()
            self._add_budget()
            self._add_signature_section()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申报书已生成：{output_path}")
        return str(output_path)

    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）"""
        warnings = []
        p0_fields = [("project_name", "项目名称"), ("leader_name", "负责人姓名"),
                     ("advisor_name", "指导教师姓名"), ("college", "所在学院")]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        basic = self._get("basic_info", default={})
        if not isinstance(basic, dict):
            basic = {}
        if not basic.get("industry") and not self._get("industry"):
            warnings.append("缺少 所属行业（industry）")
        if not basic.get("duration") and not self._get("duration"):
            warnings.append("缺少 起止时间（duration）")

        company = self._get("company_overview", default={})
        if not isinstance(company, dict):
            company = {}
        if not company.get("company_name") and not basic.get("company_name"):
            warnings.append("缺少 公司名称（company_name）——创业实践硬门槛")
        if (not company.get("credit_code") and not basic.get("credit_code")
                and not company.get("business_info")):
            warnings.append("缺少 统一社会信用代码（credit_code）——创业实践硬门槛")
        if not company.get("shareholders"):
            warnings.append("缺少 股权结构（shareholders）——创业实践硬门槛")

        for key, name in [("abstract", "项目简介"),
                          ("business_operation", "商业模式与运营现状"),
                          ("financial_data", "财务数据"),
                          ("financing_plan", "融资计划")]:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}），将使用占位文本")

        items = self._get("budget_items", default=[])
        if items:
            total = 0
            for b in items:
                if isinstance(b, dict):
                    try:
                        total += int(b.get("amount", 0))
                    except (ValueError, TypeError):
                        pass
            try:
                budget_total_num = int(str(self._get("budget_total", default="")).strip())
            except ValueError:
                budget_total_num = -1
            if budget_total_num >= 0 and total != budget_total_num:
                warnings.append(f"预算合计 {total} 元 与申请经费 {budget_total_num} 元不一致")

        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ============================================================
# 默认示例数据（创业实践：杭州课座科技有限公司）
# ============================================================

DEFAULT_DATA = {
    "project_name": "校园自习室智能预约平台'课座'",
    "project_level": "国家级",
    "project_type": "创业实践项目",
    "leader_name": "张三", "leader_id": "202212345",
    "leader_major": "工商管理", "leader_grade": "2022 级 大三",
    "leader_phone": "138XXXXXXXX",
    "advisor_name": "李教授", "advisor_title": "副教授",
    "advisor_research": "创业管理",
    "college": "工商管理学院", "apply_date": "2025 年 3 月 15 日",
    "basic_info": {
        "project_name": "校园自习室智能预约平台'课座'",
        "project_type": "创业实践项目",
        "project_source": "A 学生自主选题",
        "industry": "互联网/教育科技",
        "company_name": "杭州课座科技有限公司",
        "credit_code": "91330108MA2XXXXXX5",
        "duration": "2025.04-2026.09（18 个月）",
        "budget": "20000",
        "leader_info": "张三 / 202212345 / 工商管理 / 大三 / 138XXXXXXXX",
        "team_members": "李四（CTO）、王五（CMO）、赵六（CFO）、孙七（COO）",
        "advisor_info": "李教授 / 副教授 / 创业管理 / 139XXXXXXXX",
    },
    "abstract": "杭州课座科技有限公司成立于 2024 年 6 月，针对高校自习室座位抢占、利用率低、管理粗放的痛点，开发校园自习室智能预约平台'课座'。已运营 9 个月，覆盖本市 5 所高校，累计用户 5800 人，月营收 8.2 万元，毛利率 45%，30 日复购率 32%，采用'会员订阅 19.9 元/月 + 商家广告'盈利模式。已完成种子轮 80 万元融资（杭州 XX 投资），团队 5 人含工商管理 2 人 + 计算机 2 人 + 设计 1 人复合背景，全职 3 人兼职 2 人。本轮申请大创经费 2 万元，预期 18 个月内覆盖 20 所高校、月营收 50 万元、用户 5 万、启动 Pre-A 轮 500 万元融资。",
    "company_overview": {
        "business_info": [
            {"field": "公司全称", "value": "杭州课座科技有限公司"},
            {"field": "统一社会信用代码", "value": "91330108MA2XXXXXX5"},
            {"field": "注册日期", "value": "2024 年 6 月 18 日"},
            {"field": "注册资本", "value": "100 万元（实缴 30 万元）"},
            {"field": "企业类型", "value": "有限责任公司"},
            {"field": "法定代表人", "value": "张三"},
            {"field": "注册地址", "value": "浙江省杭州市余杭区 XX 路 XX 号 X 室"},
            {"field": "经营范围", "value": "技术服务、软件开发、互联网信息服务"},
            {"field": "经营状态", "value": "在营 / 存续"},
            {"field": "实际运营月数", "value": "9 个月（截至 2025 年 3 月）"},
        ],
        "shareholders": [
            {"name": "张三", "amount": "60", "ratio": "60%", "form": "货币"},
            {"name": "李四", "amount": "25", "ratio": "25%", "form": "货币"},
            {"name": "王五", "amount": "15", "ratio": "15%", "form": "技术（软件著作权作价）"},
            {"name": "期权池（预留）", "amount": "10", "ratio": "10%", "form": "待发"},
        ],
        "governance": "公司设股东会（4 名股东）、执行董事 1 名（张三）、监事 1 名（赵六）。税务合规：增值税按月申报，2024 年 7 月至今已申报 9 期；企业所得税按季申报，已申报 3 季度；个税按月代扣代缴。社保缴纳：3 名全职员工已在杭州余杭社保局参保（养老/医疗/失业/工伤/生育五险）。知识产权归属：1 项软件著作权《课座智能预约系统 V1.0》、1 项商标'课座'（第 42 类）、域名 kezuo.com 均登记在公司名下。",
    },
    "product_market": {
        "features": [
            "座位实时预约：按时间段预约自习室座位，预约成功率 ≥85%。",
            "签到核销：扫码签到，未签到 15 分钟自动释放，签到率 ≥90%。",
            "信用积分：爽约扣分、按时签到加分，积分高者优先预约。",
            "学习数据：累计学习时长、连续打卡、学习效率分析。",
            "校园社群：自习室社群、组队学习、PK 排行榜。",
        ],
        "market_analysis": "TAM 总市场：全国 3000 所高校 4700 万在校生 × 人均年自习消费 200 元 = 94 亿元/年（数据源：教育部 2024 教育公报 + 自调研）。SAM 可服务市场：本省 50 所高校 120 万在校生 × 200 元 = 2.4 亿元/年。SOM 可获取市场：3 年内覆盖 50 所高校 60 万在校生 × 渗透率 15% × 200 元 = 1800 万元/年。竞品分析：① 超级课程表（5 千万用户，优势：用户量大，劣势：自习场景弱）；② 美团校园（5000 万用户，优势：本地服务全，劣势：非垂直）；③ 各校自有系统（不可统计，优势：免费，劣势：体验差）。本项目差异化：垂直自习场景 + 信用积分 + 跨校通用。",
        "tech_impl": "前端：微信小程序（Taro 3.x）+ React Native（校园管理端）；后端：Node.js + Express + Python（AI 推理）；数据库：MySQL 8.0 + Redis 7.0；第三方：腾讯云 LBS、阿里云 OSS、微信支付、阿里云短信。研发投入：3 人 × 9 月 = 27 人月，约 13.5 万元（按校内外包价折算）。",
        "ip_status": "软件著作权 1 项：《课座智能预约系统 V1.0》（登记号 2024SR10XXXXX，归属公司）；商标 1 项：'课座'第 42 类（注册号 78XXXXX，归属公司）；域名：kezuo.com（归属公司）。正在申请：发明专利 1 项（基于信用积分的自习室调度算法）。",
    },
    "business_operation": {
        "model": "盈利模式 3 类：① 会员订阅 19.9 元/月，含优先预约、信用加权、免爽约券 1 张，付费率 18%，月营收 4.2 万元（占 51%）；② 商家广告 CPM 50 元，校园周边商家投放，月营收 2.2 万元（占 27%）；③ 增值服务（学习数据深度分析、跨校 PK 入场券），月营收 1.8 万元（占 22%）。渠道：微信小程序 + 校园社团合作 + 社群裂变。CAC 12 元，LTV 56 元，LTV/CAC = 4.7。",
        "kpi_table": [
            {"name": "累计用户", "current": "5800 人", "yoy": "+220%"},
            {"name": "月活用户", "current": "2300 人", "yoy": "+180%"},
            {"name": "月营收（MRR）", "current": "8.2 万元", "yoy": "+150%"},
            {"name": "ARR（年化）", "current": "98 万元", "yoy": "+200%"},
            {"name": "毛利率", "current": "45%", "yoy": "+5pp"},
            {"name": "30 日复购率", "current": "32%", "yoy": "+8pp"},
            {"name": "月流失率", "current": "8%", "yoy": "-3pp"},
            {"name": "CAC", "current": "12 元", "yoy": "-25%"},
            {"name": "LTV", "current": "56 元", "yoy": "+40%"},
            {"name": "LTV/CAC", "current": "4.7", "yoy": "+85%"},
        ],
        "milestones": [
            {"time": "2024.06", "event": "公司注册成立", "status": "已完成"},
            {"time": "2024.07", "event": "MVP 上线、首批 1 校试点", "status": "已完成"},
            {"time": "2024.08", "event": "种子轮 80 万元融资到账", "status": "已完成"},
            {"time": "2024.12", "event": "覆盖 3 校、月营收破 5 万", "status": "已完成"},
            {"time": "2025.03", "event": "覆盖 5 校、月营收 8.2 万", "status": "已完成"},
            {"time": "2025.06", "event": "覆盖 10 校、启动 Pre-A 接洽", "status": "进行中"},
            {"time": "2025.12", "event": "覆盖 20 校、Pre-A 500 万到账", "status": "计划"},
            {"time": "2026.09", "event": "月营收 50 万、用户 5 万", "status": "计划"},
        ],
    },
    "financial_data": {
        "income_statement": [
            {"item": "营业收入（万元）", "y1": "78", "y2": "280", "y3": "800"},
            {"item": "营业成本（万元）", "y1": "43", "y2": "140", "y3": "360"},
            {"item": "毛利率", "y1": "45%", "y2": "50%", "y3": "55%"},
            {"item": "销售费用（万元）", "y1": "18", "y2": "50", "y3": "120"},
            {"item": "管理费用（万元）", "y1": "12", "y2": "25", "y3": "50"},
            {"item": "研发费用（万元）", "y1": "15", "y2": "40", "y3": "100"},
            {"item": "财务费用（万元）", "y1": "0.5", "y2": "1", "y3": "2"},
            {"item": "营业利润（万元）", "y1": "-10.5", "y2": "24", "y3": "168"},
            {"item": "净利润（万元）", "y1": "-10.5", "y2": "22", "y3": "150"},
        ],
        "cash_flow": [
            {"item": "经营性现金流净额", "y1": "-8", "y2": "15", "y3": "120"},
            {"item": "投资性现金流净额", "y1": "-5", "y2": "-10", "y3": "-30"},
            {"item": "筹资性现金流净额", "y1": "80", "y2": "500", "y3": "0"},
            {"item": "期末现金余额", "y1": "67", "y2": "572", "y3": "662"},
        ],
        "balance_sheet": [
            {"item": "货币资金", "y1": "67", "y2": "572", "y3": "662"},
            {"item": "应收账款", "y1": "8", "y2": "35", "y3": "100"},
            {"item": "固定资产", "y1": "3", "y2": "8", "y3": "20"},
            {"item": "无形资产", "y1": "5", "y2": "12", "y3": "30"},
            {"item": "资产合计", "y1": "83", "y2": "627", "y3": "812"},
            {"item": "短期借款", "y1": "0", "y2": "0", "y3": "0"},
            {"item": "实收资本", "y1": "100", "y2": "100", "y3": "100"},
            {"item": "未分配利润", "y1": "-10", "y2": "12", "y3": "162"},
        ],
        "burn_runway": "Burn rate（月度净消耗）8 万元/月（人力 4 万 + 服务器 0.5 万 + 推广 2 万 + 杂项 1.5 万 - 月营收 8.2 万 = 净 Burn 8 万）。现金余额 96 万元。Runway = 96/8 = 12 个月。盈亏平衡点：月营收达 15 万元时毛利率 50% 覆盖月固定成本 7.5 万元，预期 Year2 第 8 个月达到。",
    },
    "team_organization": {
        "members": [
            {"name": "张三", "id": "202212345", "major": "工商管理 大三", "role": "CEO / 项目负责人", "exp": "字节实习、校创业大赛一等奖"},
            {"name": "李四", "id": "202212346", "major": "计算机 大三", "role": "CTO", "exp": "字节技术实习、GitHub 2k star"},
            {"name": "王五", "id": "202212347", "major": "市场营销 大三", "role": "CMO", "exp": "校园社群运营 5000+ 用户"},
            {"name": "赵六", "id": "202212348", "major": "财务管理 大三", "role": "CFO", "exp": "会计事务所实习、CPA 在考"},
            {"name": "孙七", "id": "202212349", "major": "视觉传达 大三", "role": "COO / 设计", "exp": "红点设计奖入围"},
        ],
        "org_structure": "CEO 张三下设 CTO（李四，负责技术）/ CMO（王五，负责市场）/ CFO（赵六，负责财务）/ COO（孙七，负责运营与设计）四部门。全职 3 人（张三、李四、王五，已办离校创业手续），兼职 2 人（赵六、孙七，在校）。外部顾问 2 人：律师（杭州 XX 律所合伙人，负责公司法务）+ 财务顾问（前普华永道审计，负责融资财务）。",
        "salary": "全职 3 人月薪 3000 元（基本生活费，待 Pre-A 后调整至市场价 70%，即 CEO 1.2 万 / CTO 1.5 万 / CMO 1 万）。兼职 2 人按项目计酬（赵六负责月度结账 800 元/月，孙七负责设计任务 1000 元/月）。已为 3 名全职员工在杭州余杭社保局参保（五险），兼职学生保留学籍社保。",
    },
    "financing_plan": {
        "history": [
            {"round": "种子轮", "time": "2024.08", "amount": "80 万元", "valuation": "800 万投后", "equity": "10%", "investor": "杭州 XX 投资合伙企业"},
        ],
        "current_round": "本轮计划启动 Pre-A 轮融资 500 万元，投后估值 2000 万元（出让 25%）。估值采用可比公司法：参照同阶段教育 SaaS 公司 PS 倍数 8 倍 × Year2 营收预测 280 万 = 2240 万，让步至 2000 万投后。资金用途分解：① 产品研发 40%（200 万，后端架构升级、AI 算法研发）；② 团队扩张 30%（150 万，招聘工程师 5 人、市场 3 人、运营 2 人）；③ 市场推广 20%（100 万，5 城市推广、校园合作）；④ 运营储备 10%（50 万，应急资金）。",
        "term_sheet": "Term sheet 关键条款：① 投后估值 2000 万元、出让 25%；② 清算优先权 1 倍非参与型；③ 反稀释加权平均；④ 董事会 3 席（创始人 2 + 投资方 1）；⑤ 回购条款：5 年内未上市或被并购按本金 + 8% 年息回购；⑥ 业绩对赌：Year2 营收未达 200 万，创始人让渡 5%；⑦ 一票否决权：公司并购、重大资产处置、关键人事。",
    },
    "risk_response": [
        {"type": "市场风险", "risk": "美团/抖音进入校园预约", "prob": "中", "impact": "高", "measure": "深耕自习室垂直场景、强化校园合作壁垒"},
        {"type": "技术风险", "risk": "高并发签到性能瓶颈", "prob": "低", "impact": "中", "measure": "腾讯云弹性扩容、Redis 缓存"},
        {"type": "运营风险", "risk": "学生放假导致营收季节性下滑", "prob": "高", "impact": "中", "measure": "拓展成人自考/考研自习室、寒暑假保留 30% 营收"},
        {"type": "财务风险", "risk": "现金流断裂", "prob": "中", "impact": "高", "measure": "控制 burn rate 8 万/月、提前 6 个月启动 Pre-A"},
        {"type": "法务风险", "risk": "知识产权归属纠纷", "prob": "低", "impact": "高", "measure": "软件著作权已登记归属公司、签订员工 IP 转让协议"},
    ],
    "social_benefit": "① 就业带动：已雇佣全职 3 人 + 兼职 5 人（在校生），预期 Year2 新增 8 人、Year3 新增 20 人，缓解大学生就业压力。② 校园赋能：已覆盖 5 所高校、服务 5800 学生，提升自习室利用率 35%、降低占座纠纷 80%，被杭州 3 所高校列为'智慧校园'试点项目。③ ESG 价值：减少纸质预约单（年节约 12 万张纸）、降低空调能耗（智能错峰，年节电 8 万度）、助力绿色校园建设，已与杭州市余杭区团委签订绿色校园共建协议。",
    "expected_outcomes": [
        "营收：18 个月内累计营收 600 万元、月营收达 50 万元",
        "用户：累计用户 5 万人、月活 1.5 万人",
        "融资：完成 Pre-A 轮 500 万元融资",
        "就业：新增全职岗位 8 个、兼职岗位 20 个",
        "知识产权：软件著作权 2 项、商标 1 项",
        "合作高校：覆盖 20 所高校、签订正式合作协议",
        "报告：年度运营报告 1 份（含完整财务报表）",
    ],
    "budget_items": [
        {"item": "服务器与域名", "amount": "3200", "basis": "阿里云 ECS 200 元/月 × 12 月 + 域名 180 元 + SSL 620 元"},
        {"item": "用户调研", "amount": "1500", "basis": "100 份深度访谈 × 15 元 + 数据分析工具 0 元（开源）"},
        {"item": "推广获客", "amount": "5000", "basis": "5 校 × 800 元线下活动 + 1000 元社群裂变红包"},
        {"item": "物料与会议", "amount": "2300", "basis": "海报印刷 1200 + 路演物料 800 + 行业会议 300"},
        {"item": "印刷复印", "amount": "800", "basis": "商业计划书印刷 30 份 × 25 元 + 报告 50 元"},
        {"item": "知识产权申请", "amount": "4200", "basis": "软件著作权 2 项 × 1500 元 + 商标 1 项 × 1200 元"},
        {"item": "团队培训", "amount": "3000", "basis": "创业培训 2 人 × 1500 元"},
    ],
    "budget_total": "20000",
    "include_school_approval": False,
}


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="大创-创业实践项目申报书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n"
            "JSON 字段定义详见 SKILL.md 第十一章。\n"
            "创业实践硬门槛：已注册公司或运营 6 个月以上。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
