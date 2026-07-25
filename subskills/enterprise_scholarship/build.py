#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""企业专项奖学金申请书 docx 生成器

格式标准：A4 纸张，页边距上下 2.54cm 左右 2.5cm；标题黑体二号居中（动态填充"XX 奖学金申请书"）；
称呼顶格宋体小四全角冒号；正文宋体小四 1.5 倍行距首行缩进 2 字符；主干课程表与三阶段职业规划表
宋体五号居中（后者 4 列：阶段/时间/目标/路径）；"此致"另起一行空两格，"敬礼！"另起一行顶格；落款右对齐。

使用：python build.py --data data.json --out output.docx  或  python build.py --demo --out demo.docx
JSON 字段详见 SKILL.md 第三章。企业专项专属字段：enterprise_name/scholarship_name/enterprise_industry/
enterprise_position/industry_status/industry_pain/industry_trend/enterprise_value/
career_short/career_mid/career_long/enterprise_match/internship/project_cooperation。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)

def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)

def add_paragraph_with_format(
    doc, text: str, font_name: str = FONT_SONG, font_size=SIZE_XIAO_SI,
    bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True, line_spacing: float = 1.5,
    space_before: float = 0, space_after: float = 0,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p

def add_title(doc, text: str):
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)

def add_body_paragraph(doc, text: str, indent: bool = True):
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=indent, line_spacing=1.5)

def add_salutation_paragraph(doc, text: str):
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)

def add_section_heading(doc, text: str):
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
        line_spacing=1.5, space_before=6, space_after=3)

def add_cizhi_paragraph(doc, text: str = "此致"):
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=True, line_spacing=1.5)

def add_jingli_paragraph(doc, text: str = "敬礼！"):
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)

def add_right_aligned_paragraph(doc, text: str):
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=False, line_spacing=1.5)

def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None,
                        caption: str = ""):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    if caption:
        add_paragraph_with_format(
            doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)

    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table

def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)

def add_page_number(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)

class ApplicationDocBuilder:
    """企业专项奖学金申请书 docx 构建器"""

    # 企业简称对照表（全称 → 简称）
    ENTERPRISE_SHORT = {
        "华为技术有限公司": "华为",
        "腾讯科技（深圳）有限公司": "腾讯", "腾讯计算机系统有限公司": "腾讯",
        "字节跳动有限公司": "字节跳动",
        "阿里巴巴集团控股有限公司": "阿里巴巴", "阿里巴巴（中国）有限公司": "阿里巴巴",
        "百度在线网络技术（北京）有限公司": "百度", "百度股份有限公司": "百度",
        "中兴通讯股份有限公司": "中兴",
        "中国石油化工集团有限公司": "中石化", "中国石化集团有限公司": "中石化",
        "中国石油天然气集团有限公司": "中石油", "中国海洋石油集团有限公司": "中海油",
        "国家电网有限公司": "国家电网", "中国南方电网有限责任公司": "南方电网",
        "中国移动通信集团有限公司": "中国移动",
        "宝钢集团有限公司": "宝钢", "中国宝武钢铁集团有限公司": "中国宝武",
        "中国工商银行股份有限公司": "工行", "中国建设银行股份有限公司": "建行",
        "招商银行股份有限公司": "招行",
    }

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    def _get(self, *keys, default=""):
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val]
        return []

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_heading(self, text):
        return add_section_heading(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, caption=""):
        return add_table_from_data(self.doc, headers, rows, col_widths, caption)

    def _enterprise_short_name(self) -> str:
        full = self._get("enterprise_name", default="")
        if not full:
            return ""
        if full in self.ENTERPRISE_SHORT:
            return self.ENTERPRISE_SHORT[full]
        # 兜底：取"有限公司"前的核心字号
        for kw in ["有限公司", "股份有限公司", "集团有限公司", "集团"]:
            if kw in full:
                return full.split(kw)[0].rstrip("（）()")
        return full

    def _add_title(self):
        """标题：黑体二号居中，动态填充"XX 奖学金申请书"
        优先级：scholarship_name 已含"申请书" → 直接用；不含"申请书"但含"奖学金" → 补后缀；
                无 scholarship_name → 取企业简称 + "奖学金申请书"；无企业名 → 用占位。"""
        scholarship = self._get("scholarship_name", default="")
        short = self._enterprise_short_name()
        if scholarship:
            if "申请书" in scholarship:
                title = scholarship
            elif "奖学金" in scholarship:
                title = f"{scholarship}申请书"
            else:
                title = f"{scholarship}奖学金申请书"
        elif short:
            title = f"{short}奖学金申请书"
        else:
            title = "企业专项奖学金申请书"
        add_title(self.doc, title)
    def _add_salutation(self):
        salutation = self._get("salutation",
                               default="尊敬的学院领导、评审委员会：")
        add_salutation_paragraph(self.doc, salutation)
    def _add_opening(self):
        opening = self._get("opening", default="")
        if opening:
            self.add_para(opening)
            return
        name = self._get("name")
        college = self._get("college")
        major = self._get("major")
        grade = self._get("grade")
        apply_year = self._get("apply_year", default="2024-2025 学年")
        scholarship = self._get("scholarship_name",
                                default=self._get("enterprise_name",
                                                  default="企业专项奖学金"))
        gpa = self._get("gpa")
        rank = self._get("rank")
        internship_short = ""
        internship = self._get_list("internship")
        if internship and isinstance(internship[0], dict):
            it = internship[0]
            ent = it.get("company", "")
            dur = it.get("duration", "")
            if ent and dur:
                internship_short = f"曾在{ent}实习{dur}"
        parts = []
        if name and college and major and grade:
            parts.append(f"我是{college}{major}{grade}学生{name}，"
                         f"特申请{apply_year}{scholarship}。")
        else:
            parts.append(f"特申请{apply_year}{scholarship}。")
        data_parts = []
        if gpa:
            data_parts.append(f"GPA {gpa}")
        if rank:
            data_parts.append(f"专业排名第 {rank}")
        if internship_short:
            data_parts.append(internship_short)
        if data_parts:
            parts.append("；".join(data_parts) + "。")
        parts.append("现将本人情况汇报如下：")
        self.add_para("".join(parts))
    def _add_ideology(self):
        self.add_heading("一、思想方面")
        ideology = self._get("ideology", default="")
        if ideology:
            if isinstance(ideology, list):
                for p in ideology:
                    self.add_para(p)
            else:
                self.add_para(ideology)
            return
        political = self._get("political_status", default="共青团员")
        party_history = self._get("party_history", default="")
        party_activities = self._get_list("party_activities")
        parts = []
        if political:
            if "党员" in political:
                parts.append(f"作为一名{political}，我始终坚持学习党的理论，"
                             "认真学习习近平新时代中国特色社会主义思想与党的二十大和二十届三中全会精神。")
            elif "积极分子" in political:
                parts.append("作为一名入党积极分子，我认真学习习近平新时代中国特色社会主义思想，"
                             "时刻以党员标准要求自己。")
            else:
                parts.append(f"作为一名{political}，我拥护中国共产党的领导，"
                             "认真学习党的创新理论。")
        if party_history:
            parts.append(party_history)
        if party_activities:
            parts.append("；".join(party_activities))
        parts.append("在日常生活中，我注重理论学习与实践结合，关注时政热点，提升思想觉悟。")
        self.add_para("".join(parts))
    def _add_academics(self):
        self.add_heading("二、学习方面")
        academics = self._get("academics", default="")
        if academics and isinstance(academics, str):
            self.add_para(academics)
            self._add_core_courses_table()
            self._add_academics_summary()
            return
        gpa = self._get("gpa")
        weighted = self._get("weighted_avg")
        rank = self._get("rank")
        rank_total = self._get("rank_total")
        course_count = self._get("course_count")
        high_score_count = self._get("high_score_count")
        parts = []
        if gpa:
            gpa_str = f"本学年 GPA {gpa}"
            if weighted:
                gpa_str += f"，加权平均分 {weighted}"
            if rank:
                gpa_str += f"，专业排名第 {rank}"
                if rank_total:
                    try:
                        r_num = int(str(rank).split("/")[0])
                        total_num = int(rank_total)
                        if total_num > 0:
                            pct = round(r_num / total_num * 100, 1)
                            gpa_str += f"（前 {pct}%）"
                    except (ValueError, IndexError):
                        pass
            parts.append(gpa_str + "。")
        if course_count and high_score_count:
            parts.append(f"修读 {course_count} 门课程，{high_score_count} 门 80 分以上。")
        if parts:
            self.add_para("".join(parts))
        self._add_core_courses_table()
        self._add_academics_summary()
    def _add_core_courses_table(self):
        courses = self._get_list("core_courses")
        if not courses:
            return
        rows = []
        for c in courses:
            if not isinstance(c, dict):
                continue
            rows.append([str(c.get("name", "")), str(c.get("credit", "")),
                         str(c.get("score", ""))])
        if rows:
            self.add_table(["课程名称", "学分", "成绩"], rows,
                           col_widths=[8.0, 3.0, 3.0], caption="主干课程成绩：")
    def _add_academics_summary(self):
        summary = self._get("academics_summary", default="")
        if summary:
            self.add_para(summary)
            return
        cet4 = self._get("cet4")
        cet6 = self._get("cet6")
        computer_level = self._get("computer_level")
        course_highlight = self._get("course_highlight", default="")
        study_method = self._get("study_method",
                                  default="学习上注重课前预习与课后总结，"
                                          "建立知识体系；遇到问题主动与老师、同学讨论。")
        parts = []
        if course_highlight:
            parts.append(course_highlight + "。")
        lang_parts = []
        if cet4:
            lang_parts.append(f"CET-4 {cet4} 分")
        if cet6:
            lang_parts.append(f"CET-6 {cet6} 分")
        if lang_parts:
            parts.append("、".join(lang_parts) + "，能熟练阅读英文文献；")
        if computer_level:
            parts.append(f"计算机{computer_level}。")
        if study_method:
            parts.append(study_method)
        if parts:
            self.add_para("".join(parts))
    def _add_industry_understanding(self):
        """行业认识（250~350 字，企业专项专属重点）：
        行业现状 + 行业痛点 + 行业趋势 + 该企业在行业中的地位 + 与企业价值观契合度
        """
        self.add_heading("三、行业认识方面")
        industry_text = self._get("industry_understanding", default="")
        if industry_text:
            if isinstance(industry_text, list):
                for p in industry_text:
                    self.add_para(p)
            else:
                self.add_para(industry_text)
            return

        enterprise_name = self._enterprise_short_name() or "该企业"
        enterprise_industry = self._get("enterprise_industry", default="")
        enterprise_position = self._get("enterprise_position", default="")
        industry_status = self._get("industry_status", default="")
        industry_pain = self._get("industry_pain", default="")
        industry_trend = self._get("industry_trend", default="")
        enterprise_value = self._get("enterprise_value", default="")

        # 段 1：行业现状 + 痛点 + 趋势
        parts1 = []
        if industry_status:
            parts1.append(f"我了解到，{industry_status}")
            if industry_pain:
                parts1.append(f"但同时行业面临以下痛点：{industry_pain}")
            if industry_trend:
                parts1.append(f"展望未来 3~5 年，{industry_trend}")
        elif enterprise_industry:
            parts1.append(f"我了解到{enterprise_industry}行业的发展现状。")
        if parts1:
            self.add_para("".join(parts1))

        # 段 2：该企业在行业中的地位 + 价值观契合度
        parts2 = []
        if enterprise_position:
            parts2.append(f"在{enterprise_industry or '相关'}行业中，"
                          f"{enterprise_name}{enterprise_position}。")
        if enterprise_value:
            parts2.append(f"我认同{enterprise_name}'{enterprise_value}'的核心价值观，"
                          "这与我专注深耕的个人追求高度契合。")
        if parts2:
            self.add_para("".join(parts2))
    def _add_career_plan(self):
        """职业规划（250~350 字，企业专项专属重点）：
        三阶段表格（短期 1~3 年/中期 3~5 年/长期 5~10 年）+ 与该企业契合度说明
        """
        self.add_heading("四、职业规划方面")
        career_text = self._get("career_plan", default="")
        if career_text:
            if isinstance(career_text, list):
                for p in career_text:
                    self.add_para(p)
            else:
                self.add_para(career_text)
            self._add_career_table()
            self._add_enterprise_match()
            return

        enterprise_name = self._enterprise_short_name() or "目标企业"
        intro = self._get("career_intro", default="")
        if intro:
            self.add_para(intro)
        else:
            self.add_para(f"我制定了三阶段职业规划，与{enterprise_name}业务方向高度契合：")
        self._add_career_table()
        self._add_enterprise_match()
    def _add_career_table(self):
        career_short = self._get("career_short", default="")
        career_mid = self._get("career_mid", default="")
        career_long = self._get("career_long", default="")
        if not (career_short or career_mid or career_long):
            return

        def split_goal_path(text):
            if not text:
                return ("", "")
            for sep in ["；", ";", "。路径：", "。路径:", "，路径：", "，路径:"]:
                if sep in text:
                    goal, path = text.split(sep, 1)
                    return (goal.strip(), path.strip())
            return (text.strip(), "")

        short_goal, short_path = split_goal_path(career_short)
        mid_goal, mid_path = split_goal_path(career_mid)
        long_goal, long_path = split_goal_path(career_long)
        rows = [
            ["短期", "1~3 年", short_goal, short_path],
            ["中期", "3~5 年", mid_goal, mid_path],
            ["长期", "5~10 年", long_goal, long_path],
        ]
        self.add_table(["阶段", "时间", "目标", "路径"], rows,
                       col_widths=[1.8, 2.0, 5.5, 5.7],
                       caption="三阶段职业规划：")
    def _add_enterprise_match(self):
        enterprise_match = self._get("enterprise_match", default="")
        if enterprise_match:
            self.add_para(enterprise_match)
            return

        enterprise_name = self._enterprise_short_name() or "该企业"
        parts = []
        course_match = self._get("course_match", default="")
        project_match = self._get("project_match", default="")
        internship_match = self._get("internship_match", default="")
        if course_match:
            parts.append(f"在校已学习{course_match}等核心课程；")
        if project_match:
            parts.append(f"主持校级大创项目《{project_match}》，"
                         f"与{enterprise_name}业务方向一致；")
        if internship_match:
            parts.append(internship_match + "；")
        if parts:
            self.add_para("上述规划与" + enterprise_name +
                          "业务方向高度契合：" + "".join(parts))
        else:
            self.add_para(f"上述规划与{enterprise_name}业务方向一致，"
                          "我将持续关注该企业技术动态，为未来入职做好充分准备。")
    def _add_research_practice(self):
        """科研与实践方面（150~250 字）：
        大创/竞赛（弱化要求）+ 该企业实习/项目合作（重点）+ 学生工作 + 志愿服务
        """
        self.add_heading("五、科研与实践方面")
        rp_text = self._get("research_practice", default="")
        if rp_text and isinstance(rp_text, str):
            self.add_para(rp_text)
            self._add_internship_items()
            self._add_competition_items()
            self._add_social_practice()
            return
        self._add_internship_items()
        self._add_competition_items()
        self._add_research_items()
        self._add_social_practice()
    def _add_internship_items(self):
        internships = self._get_list("internship")
        cooperations = self._get_list("project_cooperation")
        enterprise_name = self._enterprise_short_name() or "该企业"

        if internships:
            parts = []
            for it in internships:
                if not isinstance(it, dict):
                    continue
                company = it.get("company", "")
                position = it.get("position", "")
                duration = it.get("duration", "")
                output = it.get("output", "")
                evaluation = it.get("evaluation", "")
                seg = (f"{duration} " if duration else "") + \
                      (f"在{company}" if company else "实习")
                if position:
                    seg += f"{position}岗"
                seg += "实习"
                if output:
                    seg += f"，{output}"
                if evaluation:
                    seg += f"，获{evaluation}评价"
                if seg:
                    parts.append(seg + "。")
            if parts:
                self.add_para("".join(parts))

        if cooperations:
            parts = []
            for c in cooperations:
                if not isinstance(c, dict):
                    continue
                project = c.get("project", "")
                role = c.get("role", "")
                duration = c.get("duration", "")
                output = c.get("output", "")
                seg = (f"{duration} " if duration else "") + f"参与{enterprise_name}"
                if project:
                    seg += f"《{project}》"
                seg += "项目合作"
                if role:
                    seg += f"，担任{role}"
                if output:
                    seg += f"，{output}"
                if seg:
                    parts.append(seg + "。")
            if parts:
                self.add_para("".join(parts))
    def _add_competition_items(self):
        competitions = self._get_list("competitions")
        if not competitions:
            return
        comp_parts = []
        for c in competitions:
            if not isinstance(c, dict):
                continue
            time = c.get("time", "")
            name = c.get("name", "")
            award = c.get("award", "")
            role = c.get("role", "")
            seg = (f"{time} " if time else "") + name
            if award:
                seg += f" {award}"
            if role:
                seg += f"（{role}）"
            if seg:
                comp_parts.append(seg + "；")
        if comp_parts:
            comp_parts[-1] = comp_parts[-1].rstrip("；") + "。"
            self.add_para("学科竞赛方面，" + "".join(comp_parts))
    def _add_research_items(self):
        projects = self._get_list("research_projects")
        if not projects:
            return
        proj_parts = []
        for p in projects:
            if not isinstance(p, dict):
                continue
            name = p.get("name", "")
            level = p.get("level", "")
            role = p.get("role", "")
            duration = p.get("duration", "")
            output = p.get("output", "")
            seg = f"{role}《{name}》" if role and name else f"《{name}》" if name else ""
            if level:
                seg += f"（{level}）"
            if duration:
                seg += f"，{duration}"
            if output:
                seg += f"，{output}"
            if seg:
                proj_parts.append(seg + "。")
        if proj_parts:
            self.add_para("".join(proj_parts))
    def _add_social_practice(self):
        position = self._get("position", default="")
        position_work = self._get("position_work", default="")
        volunteer_hours = self._get("volunteer_hours", default="")
        volunteer_detail = self._get("volunteer_detail", default="")
        parts = []
        if position:
            seg = f"担任{position}"
            if position_work:
                seg += f"，主要工作：{position_work}"
            parts.append(seg + "。")
        if volunteer_hours or volunteer_detail:
            seg = f"累计志愿服务时长 {volunteer_hours} 小时" if volunteer_hours else ""
            if volunteer_detail:
                seg = (seg + "：" if seg else "") + volunteer_detail
            parts.append(seg + "。")
        if parts:
            self.add_para("".join(parts))
    def _add_life(self):
        self.add_heading("六、生活方面")
        life = self._get("life", default="")
        if life:
            if isinstance(life, list):
                for p in life:
                    self.add_para(p)
            else:
                self.add_para(life)
            return
        dorm_role = self._get("dorm_role", default="")
        dorm_activity = self._get("dorm_activity", default="")
        dorm_honor = self._get("dorm_honor", default="")
        interpersonal = self._get("interpersonal", default="")
        lifestyle = self._get("lifestyle",
                               default="生活中我注重勤俭节约，作息规律。")
        parts = [lifestyle]
        if dorm_role and dorm_activity:
            seg = f"担任{dorm_role}期间，{dorm_activity}"
            if dorm_honor:
                seg += f"，{dorm_honor}"
            parts.append(seg + "。")
        if interpersonal:
            parts.append(interpersonal + "。")
        self.add_para("".join(parts))
    def _add_ending(self):
        ending = self._get("ending", default="")
        if ending:
            self.add_para(ending)
        else:
            enterprise_name = self._enterprise_short_name() or "目标企业"
            self.add_para(
                "以上是我本学年的基本情况。无论结果如何，"
                "我都将继续努力学习、深入行业，"
                f"为未来加入{enterprise_name}做好充分准备。"
                "恳请评审委员会予以考虑。"
            )
        add_cizhi_paragraph(self.doc, "此致")   # "此致"另起一行，空两格
        add_jingli_paragraph(self.doc, "敬礼！")  # "敬礼！"另起一行，顶格
    def _add_signature(self):
        self.doc.add_paragraph()  # 空一行
        name = self._get("name", default="")
        apply_date = self._get("apply_date", default="")
        if name:
            add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if apply_date:
            add_right_aligned_paragraph(self.doc, apply_date)
    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排标题/称呼/开头/思想/学习/行业认识/职业规划/
        科研实践/生活/结尾/落款，生成 docx"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._add_title()
            self._add_salutation()
            self._add_opening()
            self._add_ideology()
            self._add_academics()
            self._add_industry_understanding()
            self._add_career_plan()
            self._add_research_practice()
            self._add_life()
            self._add_ending()
            self._add_signature()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise
    def _save(self, output_path: str) -> str:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申请书已生成：{output_path}")
        return str(output_path)
    def _validate_data(self) -> List[str]:
        warnings = []

        # P0 必采字段
        p0_fields = [
            ("name", "申请人姓名"), ("college", "学院"), ("major", "专业"),
            ("grade", "年级"), ("gpa", "GPA"), ("rank", "专业排名"),
            ("enterprise_name", "申请企业全称"), ("scholarship_name", "奖学金全称"),
        ]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        # 企业专属字段校验
        if not self._get("industry_status") and not self._get("industry_understanding"):
            warnings.append("缺少 行业现状数据（industry_status），行业认识段将内容单薄")
        if not self._get("enterprise_position"):
            warnings.append("缺少 该企业在行业中的地位（enterprise_position）")
        for f, label in [("career_short", "短期职业规划"),
                         ("career_mid", "中期职业规划"),
                         ("career_long", "长期职业规划")]:
            if not self._get(f):
                warnings.append(f"缺少 {label}（{f}），三阶段表格将不完整")

        # 主干课程（弱化要求：4~6 门）
        courses = self._get_list("core_courses")
        if not courses:
            warnings.append("缺少 主干课程（core_courses），将省略主干课程表")
        elif len(courses) < 4:
            warnings.append(f"主干课程仅 {len(courses)} 门，建议 4~6 门")

        # 排名校验
        rank_str = str(self._get("rank", default=""))
        if rank_str and "/" not in rank_str:
            if not self._get("rank_total"):
                warnings.append(
                    f"排名 '{rank_str}' 缺基数，应为 'X/N' 格式或补填 rank_total")

        # 企业专项硬门槛校验：GPA 前 30%
        try:
            rank_num = int(rank_str.split("/")[0]) if "/" in rank_str else 0
            rank_total_num = 0
            if "/" in rank_str:
                rank_total_num = int(rank_str.split("/")[1])
            elif self._get("rank_total"):
                rank_total_num = int(self._get("rank_total"))
            if rank_num and rank_total_num:
                pct = rank_num / rank_total_num * 100
                if pct > 30:
                    warnings.append(
                        f"排名前 {pct:.1f}% 不满足企业专项前 30% 硬门槛，建议改申其他奖学金")
        except (ValueError, IndexError):
            pass

        # 行业认识套话检测
        industry_text = str(self._get("industry_status", default="")) + \
                        str(self._get("industry_understanding", default=""))
        cliche_phrases = ["领军企业", "发展前景广阔", "深受用户喜爱",
                          "卓越企业", "伟大企业", "贵公司"]
        for phrase in cliche_phrases:
            if phrase in industry_text:
                warnings.append(f"行业认识段含套话'{phrase}'，建议替换为具体数据")
                break

        # 职业规划与该企业相关性校验
        career_text = (str(self._get("career_short", default="")) +
                       str(self._get("career_mid", default="")) +
                       str(self._get("career_long", default="")))
        enterprise_name = self._enterprise_short_name()
        if enterprise_name and career_text:
            if enterprise_name not in career_text and \
               "公司" not in career_text and "企业" not in career_text:
                warnings.append(
                    f"职业规划中未提及申请企业'{enterprise_name}'，"
                    "建议明确短期规划加入该企业")

        self.warnings = warnings
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings

DEFAULT_DATA = {
    "name": "李伟", "student_id": "2022112345", "gender": "男",
    "college": "计算机科学与技术学院", "major": "计算机科学与技术",
    "grade": "2022 级 大三", "class_name": "计科 2202 班",
    "political_status": "共青团员", "phone": "138XXXXXXXX",
    "apply_year": "2024-2025 学年", "apply_date": "2025 年 6 月 15 日",
    "salutation": "尊敬的学院领导、评审委员会：",
    "enterprise_name": "华为技术有限公司", "scholarship_name": "华为奖学金",
    "enterprise_industry": "信息通信技术（ICT）",
    "enterprise_position": "作为全球领先 ICT 基础设施和智能终端提供商，5G 设备全球市场份额第一（30.8%），自研麒麟芯片突破先进制程封锁，鸿蒙 OS 自主生态装机量超 9 亿台，业务覆盖 170 余个国家与地区",
    "industry_status": "ICT 行业 2024 年全球市场规模超 5 万亿美元，5G 标准必要专利数中国企业占比超 38%（工信部 2024 年数据）。",
    "industry_pain": "一是先进制程受限，国产 EDA 工具市场占有率不足 10%；二是 AI 算力供需失衡，2024 年国内智算中心缺口约 60%；三是开源生态尚不成熟，openEuler 等国产 OS 服务器市场份额仅 35%。",
    "industry_trend": "AI 算力需求年增 5 倍（Gartner 预测），国产替代深化推进，预计 2027 年国产 EDA 工具份额突破 30%。",
    "enterprise_value": "以客户为中心、长期艰苦奋斗",
    "career_short": "加入华为云分布式存储研发岗；通过校招入职，参与对象存储服务 OBS 核心模块开发，3 年内主导 1 个核心特性交付",
    "career_mid": "成为分布式存储领域技术骨干；主导 2~3 个核心模块研发，申请发明专利 2 项，在国内核心期刊或 top 会议发表论文 1 篇",
    "career_long": "助力国家关键基础软件自主可控战略；担任分布式存储技术负责人，带领团队突破关键基础软件'卡脖子'问题",
    "enterprise_match": "上述规划与华为云业务方向高度契合：在校已学习分布式系统、操作系统、计算机网络等核心课程；主持校级大创项目《基于 openEuler 的分布式存储性能优化》，与华为云 OBS 业务方向一致；2024 年暑期在华为云实习 3 个月，熟悉公司研发流程与文化。",
    "gpa": "3.78/4.0", "weighted_avg": "87.5", "rank": "5/87", "rank_total": "87",
    "course_count": "12", "high_score_count": "8",
    "core_courses": [
        {"name": "数据结构", "credit": "4", "score": "92"},
        {"name": "操作系统", "credit": "4", "score": "90"},
        {"name": "计算机网络", "credit": "3", "score": "91"},
        {"name": "数据库原理", "credit": "3", "score": "89"},
        {"name": "计算机组成原理", "credit": "4", "score": "88"},
        {"name": "分布式系统", "credit": "3", "score": "90"}],
    "course_highlight": "数据结构 92、操作系统 90、计算机网络 91、分布式系统 90，专业核心课平均 90 分",
    "cet4": "542", "cet6": "498",
    "computer_level": "二级 C 语言（优秀）、三级数据库技术",
    "study_method": "学习上注重课前预习与课后总结，建立知识体系；遇到问题主动与老师、同学讨论。",
    "party_history": "2023.09 提交入党申请书，2024.03 列为入党积极分子。",
    "party_activities": [
        "参加学院分党校第 7 期培训班（2024.09-2024.12）结业",
        "2024.10 参与主题党日活动'红色教育基地走访'，撰写调研报告 1 份（约 3000 字）",
        "提交思想汇报 3 篇",
    ],
    "ideology": "", "academics": "", "industry_understanding": "",
    "career_plan": "", "research_practice": "",
    "research_projects": [
        {"name": "基于 openEuler 的分布式存储性能优化", "level": "校级大创项目",
         "role": "主持", "duration": "2024.03-2025.03",
         "output": "性能提升 15%，项目结题评估优秀"},
    ],
    "competitions": [
        {"name": "全国大学生数学建模竞赛", "award": "省级一等奖",
         "time": "2024.11", "role": "队长，负责整体建模与论文主笔"},
        {"name": "中国大学生计算机设计大赛", "award": "省级二等奖",
         "time": "2024.05", "role": "核心成员，负责后端开发"},
    ],
    "internship": [
        {"company": "华为云", "position": "对象存储服务 OBS 研发",
         "duration": "2024.07-2024.09",
         "output": "负责分布式块存储性能优化模块开发，性能提升 15%",
         "evaluation": "优秀实习生"},
    ],
    "project_cooperation": [],
    "position": "班级学习委员（2024.09-2025.06）",
    "position_work": "组织班级学习经验交流会 4 次，建立'一对一'帮扶机制，服务同学 20 余人次，班级平均 GPA 由 3.42 提升至 3.55",
    "volunteer_hours": "80",
    "volunteer_detail": "参与三下乡 1 次（2024.07），累计支教 8 课时；担任图书馆管理员（2024.09-2025.06，每月 6 小时）",
    "dorm_role": "宿舍长", "dorm_activity": "组织宿舍 5 次集体活动",
    "dorm_honor": "宿舍连续两学期获评'文明宿舍'",
    "interpersonal": "与同学相处融洽，曾帮助室友完成 1 次重要实验调试",
    "lifestyle": "生活中我注重勤俭节约，作息规律。", "life": "", "ending": "",
}

def main():
    parser = argparse.ArgumentParser(
        description="企业专项奖学金申请书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n"
            "JSON 字段定义详见 SKILL.md 第三章信息采集清单。\n"
            "企业专项专属字段：enterprise_name/scholarship_name/"
            "industry_status/career_short/career_mid/career_long。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据（华为奖学金）生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据（华为奖学金）生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)

if __name__ == "__main__":
    main()
