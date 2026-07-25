#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
优秀毕业设计/论文申报书 docx 生成器（v1.0）

4 段结构输出：①课题选择与意义 ②研究方法与技术路线（含研究阶段表）
③创新点（含创新点表，按 3 类：理论/方法/应用分类）④成果与应用价值（含应用价值表）

字数三档：brief 短档 ~2000 字（院级） / standard 中档 ~2700 字（校级，默认） / enhanced 长档 ~3200 字
申报规格：school 校级 / college 院级
毕业设计类型：experimental 实验研究 / engineering 工程设计 / theoretical 理论研究 / software 软件开发 / art_design 艺术设计
申请人类型：undergraduate / master / doctor / junior_college

硬门槛 4 项：答辩通过 + 成绩排名（前 20% 校级 / 前 30% 院级）+ 查重合格 + 导师推荐

格式：A4 / 页边距上下 2.54cm 左右 2.5cm / 黑体二号标题 / 宋体小四 1.5 倍行距首行缩进 2 字符
研究阶段表（4 列）+ 创新点表（3 列）+ 应用价值表（3 列）/ 此致敬礼 / 落款右对齐（3 行）/ 摘要末尾或独立页

使用方式：
    python build.py --data data.json --out output.docx
    python build.py --demo --out demo.docx
    python build.py --demo --word-count-version brief --out demo_brief.docx
    python build.py --demo --apply-level college --out demo_college.docx

JSON 字段详见 SKILL.md §8 信息采集清单与 §14 JSON Schema 完整字段定义。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ===== 字体与格式常量 =====
FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号
SIZE_XIAO_WU = Pt(9)        # 小五
SIZE_LIU = Pt(7.5)          # 六号

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

# ===== 字数档位与申报规格常量 =====
WORD_COUNT_VERSIONS = ("brief", "standard", "enhanced")
WORD_COUNT_TARGETS = {"brief": 2000, "standard": 2700, "enhanced": 3200}
APPLY_LEVELS = ("school", "college")
THESIS_TYPES = ("experimental", "engineering", "theoretical", "software", "art_design")
GRADUATE_TYPES = ("undergraduate", "master", "doctor", "junior_college")
INNOVATION_TYPES = ("theoretical", "method", "application")


# ===== 字体/段落/表格辅助函数 =====

def set_run_font(run, font_name: str = FONT_SONG, font_size=SIZE_XIAO_SI,
                 bold: bool = False, color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG, font_size=SIZE_WU,
                  bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  line_spacing: float = 1.25) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(doc, text: str, font_name: str = FONT_SONG,
    font_size=SIZE_XIAO_SI, bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True, line_spacing: float = 1.5,
    space_before: float = 0, space_after: float = 0):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_title(doc, text: str):
    """标题：黑体二号，居中，段前段后 12pt"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
                                      alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False, space_before=12, space_after=12)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
                                     alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=indent, line_spacing=1.5)


def add_salutation_paragraph(doc, text: str):
    """称呼：顶格（不缩进），宋体小四"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
                                     alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_section_heading(doc, text: str):
    """正文小节标题（一、二、三…）：黑体小四加粗，首行缩进 2 字符"""
    return add_paragraph_with_format(doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SI, bold=True,
                                     alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
                                     line_spacing=1.5, space_before=6, space_after=3)


def add_cizhi_paragraph(doc, text: str = "此致"):
    """'此致'另起一行，空两格（首行缩进 2 字符）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
                                     alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True, line_spacing=1.5)


def add_jingli_paragraph(doc, text: str = "敬礼！"):
    """'敬礼！'另起一行，顶格（不缩进）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
                                     alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False, line_spacing=1.5)


def add_right_aligned_paragraph(doc, text: str):
    """右对齐段落（落款用）"""
    return add_paragraph_with_format(doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
                                     alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=False, line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None, caption: str = "",
                        header_bold: bool = True):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）"""
    if caption:
        add_paragraph_with_format(
            doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)
    if not rows:
        return None
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG, font_size=SIZE_WU, bold=header_bold)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            set_cell_text(cells[j], val, font_name=FONT_SONG, font_size=SIZE_WU, bold=False)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ===== 主构建类 =====
class OutstandingThesisDocBuilder:
    """优秀毕业设计/论文申报书 docx 构建器（4 段结构 + 摘要）"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []

    # ---------- 字段安全访问 ----------
    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def _get_list(self, key: str) -> List[Any]:
        """安全取列表字段，非列表返回空列表"""
        val = self._get(key, default=[])
        if isinstance(val, list):
            return val
        if isinstance(val, str):
            return [val] if val else []
        return []

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_heading(self, text):
        return add_section_heading(self.doc, text)

    def add_table(self, headers, rows, col_widths=None, caption=""):
        return add_table_from_data(self.doc, headers, rows, col_widths, caption)

    # ---------- 路由解析 ----------
    def _resolve_apply_level(self) -> str:
        """解析 apply_level 字段（默认 school）"""
        level = str(self._get("apply_level", default="school")).lower().strip()
        if level in APPLY_LEVELS:
            return level
        cn_map = {"校级": "school", "校": "school", "院级": "college", "院": "college"}
        return cn_map.get(level, "school")

    def _resolve_word_count_version(self) -> str:
        """解析 word_count_version 字段（默认 standard）"""
        version = str(self._get("word_count_version", default="standard")).lower().strip()
        if version in WORD_COUNT_VERSIONS:
            return version
        cn_map = {"短": "brief", "精简": "brief", "中": "standard", "标准": "standard",
                  "长": "enhanced", "加强": "enhanced", "增强": "enhanced"}
        return cn_map.get(version, "standard")

    def _resolve_thesis_type(self) -> str:
        """解析 thesis_type 字段（默认 experimental）"""
        tt = str(self._get("thesis_type", default="experimental")).lower().strip()
        if tt in THESIS_TYPES:
            return tt
        cn_map = {"实验研究": "experimental", "实验": "experimental", "工程设计": "engineering",
                  "工程": "engineering", "理论研究": "theoretical", "理论": "theoretical",
                  "软件开发": "software", "软件": "software", "艺术设计": "art_design", "艺术": "art_design"}
        return cn_map.get(tt, "experimental")

    def _resolve_graduate_type(self) -> str:
        """解析 graduate_type 字段（默认 undergraduate）"""
        gt = str(self._get("graduate_type", default="undergraduate")).lower().strip()
        if gt in GRADUATE_TYPES:
            return gt
        cn_map = {"本科": "undergraduate", "本科生": "undergraduate", "硕士": "master",
                  "硕士研究生": "master", "博士": "doctor", "博士研究生": "doctor",
                  "专科": "junior_college", "专科生": "junior_college", "高职": "junior_college"}
        return cn_map.get(gt, "undergraduate")

    def _resolve_level_text(self) -> str:
        """规格文本（用于标题与开头）：校级 / 院级"""
        level = self._resolve_apply_level()
        return {"school": "校级", "college": "院级"}.get(level, "校级")

    # ---------- 标题与称呼 ----------
    def _resolve_title_text(self) -> str:
        """根据 apply_level + thesis_type 动态生成标题"""
        level_text = self._resolve_level_text()
        tt = self._resolve_thesis_type()
        gt = self._resolve_graduate_type()
        if gt in ("master", "doctor"):
            return "优秀学位论文申报书"
        if tt == "theoretical":
            return f"{level_text}优秀毕业论文申报书"
        return f"{level_text}优秀毕业设计申报书"

    def _add_title(self):
        """标题：黑体二号居中"""
        add_title(self.doc, self._resolve_title_text())

    def _add_salutation(self):
        """称呼：顶格，宋体小四，全角冒号"""
        salutation = self._get("salutation", default="尊敬的评审委员会：")
        add_salutation_paragraph(self.doc, salutation)

    # ---------- 开头段落 ----------
    def _add_opening(self):
        """开头段落：身份 + 申报规格 + 毕设题目 + 答辩成绩 + 导师 + 进入正文"""
        opening = self._get("opening", default="")
        if opening:
            self.add_para(opening)
            return
        name = self._get("name")
        college = self._get("college")
        major = self._get("major")
        grade = self._get("grade")
        apply_year = self._get("apply_year", default="2025 届")
        level_text = self._resolve_level_text()
        thesis_title = self._get("thesis_title", default="")
        advisor = self._get("advisor", default="")
        defense_score = self._get("defense_score", default="")
        defense_rank = self._get("defense_rank", default="")
        plagiarism = self._get("plagiarism_rate", default="")
        version = self._resolve_word_count_version()

        parts = []
        if name and college and major and grade:
            parts.append(f"我是{college}{major}{grade}学生{name}，毕业设计题目《{thesis_title}》，"
                         f"指导教师为{advisor}，特此申报{apply_year}{level_text}优秀毕业设计/论文。")
        else:
            parts.append(f"特此申报{apply_year}{level_text}优秀毕业设计/论文。")

        data_parts = []
        if defense_score:
            score_seg = f"毕设答辩成绩 {defense_score} 分"
            try:
                if int(float(defense_score)) >= 90:
                    score_seg += "（优秀）"
                elif int(float(defense_score)) >= 80:
                    score_seg += "（良好）"
            except (ValueError, TypeError):
                pass
            data_parts.append(score_seg)
        if defense_rank:
            rank_seg = f"专业排名 {defense_rank}"
            pct = self._calc_rank_percent(defense_rank)
            if pct is not None:
                rank_seg += f"（前 {pct}%）"
            data_parts.append(rank_seg)
        if plagiarism and version != "brief":
            data_parts.append(f"查重率 {plagiarism}")
        sig = self._build_signature_achievements()
        if sig:
            data_parts.append(sig)
        if data_parts:
            parts.append("；".join(data_parts) + "。")
        parts.append("现将毕业设计情况汇报如下：")
        self.add_para("".join(parts))

    def _calc_rank_percent(self, rank_str: str) -> Optional[float]:
        """根据 'X/N' 计算前 X%"""
        try:
            if "/" in str(rank_str):
                r_num = int(str(rank_str).split("/")[0])
                total_num = int(str(rank_str).split("/")[1])
            else:
                return None
            if total_num > 0:
                return round(r_num / total_num * 100, 1)
        except (ValueError, IndexError):
            pass
        return None

    def _build_signature_achievements(self) -> str:
        """构建标志成果摘要（SCI 一作 / 发明专利第一发明人 任一）"""
        parts = []
        achievements = self._get_list("achievements")
        # 论文标志
        for a in achievements:
            if not isinstance(a, dict):
                continue
            cat, level, order, name = (str(a.get("category", "")), str(a.get("level", "")),
                                        str(a.get("author_order", "")), str(a.get("name", "")))
            if cat == "论文" and ("SCI" in level or "SSCI" in level or "CSSCI" in level) and "第一" in order:
                parts.append(f"以第一作者发表{level}论文 1 篇《{name}》")
                break
        # 专利标志
        for a in achievements:
            if not isinstance(a, dict):
                continue
            cat, level, order, name = (str(a.get("category", "")), str(a.get("level", "")),
                                        str(a.get("author_order", "")), str(a.get("name", "")))
            if cat == "专利" and "发明" in level and "第一" in order:
                parts.append(f"作为第一发明人申请发明专利 1 项《{name}》")
                break
        return "；".join(parts) if parts else ""

    # ---------- 一、课题选择与意义段 ----------
    def _add_topic_selection(self):
        """课题选择与意义（400~600 字）：课题来源 + 科学/工程问题 + 学术/应用意义 + 学科前沿对接 + 研究目标"""
        self.add_heading("一、课题选择与意义")
        topic = self._get("topic_selection", default="")
        if topic and isinstance(topic, str):
            self.add_para(topic)
            return
        # 第 1 段：课题来源 + 科学/工程问题
        project_source = self._get("project_source", default="")
        research_problem = self._get("research_problem", default="")
        parts = []
        if project_source:
            parts.append(f"本课题来源于{project_source}，")
        if research_problem:
            parts.append(f"针对{research_problem}，")
        thesis_title = self._get("thesis_title", default="")
        if thesis_title:
            parts.append(f"研究{thesis_title}。")
        if parts:
            self.add_para("".join(parts))

        # 第 2 段：学术/应用意义 + 学科前沿对接
        significance = self._get("research_significance", default="")
        if significance:
            self.add_para(significance + "。")

        # 第 3 段：研究目标
        goals = self._get_list("research_goals")
        if goals:
            goal_parts = []
            for i, g in enumerate(goals, start=1):
                if isinstance(g, str):
                    goal_parts.append(f"（{i}）{g}")
                elif isinstance(g, dict):
                    goal_parts.append(f"（{i}）{g.get('goal', '') or g.get('content', '')}")
            if goal_parts:
                self.add_para("研究目标包括：" + "；".join(goal_parts) + "。")

    # ---------- 二、研究方法与技术路线段 ----------
    def _add_research_method(self):
        """研究方法与技术路线（500~700 字）：研究方法主线 + 技术路线 + 关键技术 + 数据/材料 + 研究阶段表"""
        self.add_heading("二、研究方法与技术路线")
        research = self._get("research_method_text", default="")
        if research and isinstance(research, str):
            self.add_para(research)
            self._add_research_stage_table()
            return

        # 第 1 段：研究方法主线（按 thesis_type 分流）
        method = self._get("research_method", default="")
        if method:
            self.add_para(f"研究方法采用{method}路线。")
        else:
            tt = self._resolve_thesis_type()
            default_method = self._default_method_by_thesis_type(tt)
            self.add_para(f"研究方法采用{default_method}路线。")

        # 第 2 段：技术路线 + 关键技术
        technical_route = self._get("technical_route", default="")
        if technical_route:
            self.add_para(f"技术路线：{technical_route}。")

        key_techs = self._get_list("key_technologies")
        if key_techs:
            tech_parts = []
            for i, t in enumerate(key_techs, start=1):
                if not isinstance(t, dict):
                    continue
                name = t.get("name", "")
                step = t.get("application_step", "") or t.get("step", "")
                problem = t.get("problem_solved", "") or t.get("problem", "")
                seg = f"{'①②③④⑤⑥'[i-1] if i <= 6 else f'（{i}）'}{name}"
                if step:
                    seg += f"（{step}）"
                if problem:
                    seg += f"，解决{problem}"
                tech_parts.append(seg)
            if tech_parts:
                self.add_para("关键技术包括：" + "；".join(tech_parts) + "。")

        # 第 3 段：数据/材料/装置
        data_material = self._get("data_material", default="")
        if data_material:
            self.add_para(f"实验数据与材料：{data_material}。")

        # 研究阶段表
        self._add_research_stage_table()

    def _default_method_by_thesis_type(self, tt: str) -> str:
        """按 thesis_type 返回默认研究方法主线"""
        return {
            "experimental": "文献调研 → 实验设计 → 数据采集 → 模型训练 → 性能评估",
            "engineering": "需求分析 → 方案设计 → 详细设计 → 仿真验证 → 优化迭代",
            "theoretical": "文献综述 → 命题提出 → 推导/证明 → 验证 → 结论",
            "software": "需求分析 → 系统设计 → 编码实现 → 测试验证 → 部署运行",
            "art_design": "主题确立 → 调研分析 → 创意构思 → 作品制作 → 展示答辩",
        }.get(tt, "文献调研 → 实验设计 → 数据采集 → 模型训练 → 性能评估")

    def _add_research_stage_table(self):
        """研究阶段表（4 列：阶段/时间/工作内容/产出）"""
        stages = self._get_list("research_stages")
        rows = []
        for s in stages:
            if isinstance(s, dict):
                rows.append([str(s.get("stage", "") or s.get("name", "")),
                             str(s.get("time", "") or s.get("date", "")),
                             str(s.get("content", "") or s.get("work", "")),
                             str(s.get("output", "") or s.get("result", ""))])
        if rows:
            self.add_table(["阶段", "时间", "工作内容", "产出"], rows,
                           col_widths=[3.0, 3.5, 5.5, 3.5], caption="研究阶段：")

    # ---------- 三、创新点段 ----------
    def _add_innovation_points(self):
        """创新点（400~600 字）：按 3 类（理论/方法/应用）分别阐述 + 创新点表"""
        self.add_heading("三、创新点")
        innovation_text = self._get("innovation_text", default="")
        if innovation_text and isinstance(innovation_text, str):
            self.add_para(innovation_text)
            self._add_innovation_table()
            return

        innovations = self._get_list("innovation_points")
        if not innovations:
            self.add_para("本课题创新点详见下表。")
            self._add_innovation_table()
            return

        # 按 3 类分组
        type_groups = {"theoretical": [], "method": [], "application": []}
        for inv in innovations:
            if isinstance(inv, dict):
                t = inv.get("type", "")
                if t in type_groups:
                    type_groups[t].append(inv)

        self.add_para(f"本课题创新点 {len(innovations)} 项：")
        type_labels = {"theoretical": "理论创新", "method": "方法创新", "application": "应用创新"}
        for t, label in type_labels.items():
            for i, inv in enumerate(type_groups[t], start=1):
                seg = self._build_innovation_paragraph(inv, label, i)
                if seg:
                    self.add_para(seg)

        self._add_innovation_table()

    def _build_innovation_paragraph(self, inv: Dict, category_label: str, idx: int) -> str:
        """构建单条创新点段落（80-120 字）"""
        name = inv.get("name", "")
        description = inv.get("description", "")
        comparison = inv.get("comparison", "")
        support = inv.get("support_material", "")
        seg = f"{'①②③④⑤'[idx-1] if idx <= 5 else f'（{idx}）'}**{category_label}**："
        if name:
            seg += f"{name}，"
        if description:
            seg += description
        if comparison:
            seg += f"，{comparison}"
        if support:
            seg += f"。支撑材料：{support}"
        if not seg.endswith("。"):
            seg += "。"
        return seg

    def _add_innovation_table(self):
        """创新点表（3 列：类别/创新点描述/支撑材料）"""
        innovations = self._get_list("innovation_points")
        rows = []
        type_label = {"theoretical": "理论创新", "method": "方法创新", "application": "应用创新"}
        for inv in innovations:
            if not isinstance(inv, dict):
                continue
            type_str = type_label.get(inv.get("type", ""), "")
            name = inv.get("name", "")
            description = inv.get("description", "")
            comparison = inv.get("comparison", "")
            desc = name + ("——" + description if description else "")
            if comparison:
                desc += f"（{comparison}）"
            support = inv.get("support_material", "")
            rows.append([type_str, desc, support])
        if rows:
            self.add_table(["类别", "创新点描述", "支撑材料"], rows,
                           col_widths=[2.0, 9.5, 4.0], caption="创新点汇总：")

    # ---------- 四、成果与应用价值段 ----------
    def _add_achievements_value(self):
        """成果与应用价值（300~500 字）：成果清单 + 应用场景 + 后续研究方向 + 应用价值表"""
        self.add_heading("四、成果与应用价值")
        value_text = self._get("value_text", default="")
        if value_text and isinstance(value_text, str):
            self.add_para(value_text)
            self._add_application_value_table()
            return

        # 第 1 段：成果清单（按类别分组）
        achievements = self._get_list("achievements")
        if achievements:
            cat_groups = {}
            for a in achievements:
                if isinstance(a, dict):
                    cat_groups.setdefault(a.get("category", ""), []).append(a)
            ach_parts = []
            for cat, items in cat_groups.items():
                item_parts = []
                for it in items:
                    seg = it.get("name", "")
                    if it.get("level"):
                        seg += f"（{it.get('level')}）"
                    if it.get("author_order"):
                        seg += f"，{it.get('author_order')}"
                    if it.get("time"):
                        seg += f"，{it.get('time')}"
                    item_parts.append(seg)
                if item_parts:
                    ach_parts.append(f"{cat}：{'；'.join(item_parts)}")
            if ach_parts:
                self.add_para("成果清单：" + "；".join(ach_parts) + "。")

        # 第 2 段：应用场景（standard/enhanced 显示）
        version = self._resolve_word_count_version()
        scenarios = self._get_list("application_scenarios")
        if scenarios and version != "brief":
            scene_parts = []
            for i, s in enumerate(scenarios, start=1):
                if not isinstance(s, dict):
                    continue
                scene = s.get("scene", "")
                value = s.get("value", "")
                evidence = s.get("evidence", "")
                seg = f"{'①②③④⑤'[i-1] if i <= 5 else f'（{i}）'}{scene}"
                if value:
                    seg += f"（{value}）"
                if evidence:
                    seg += f"，采纳证明：{evidence}"
                scene_parts.append(seg)
            if scene_parts:
                self.add_para("应用场景：" + "；".join(scene_parts) + "。")

        # 第 3 段：后续研究方向（enhanced 才显示）
        future = self._get_list("future_research")
        if future and version == "enhanced":
            future_parts = []
            for i, f in enumerate(future, start=1):
                if not isinstance(f, dict):
                    continue
                problem = f.get("problem", "")
                breakthrough = f.get("breakthrough", "")
                prospect = f.get("prospect", "")
                seg = f"{'①②③④⑤'[i-1] if i <= 5 else f'（{i}）'}{problem}"
                if breakthrough:
                    seg += f"，预期突破：{breakthrough}"
                if prospect:
                    seg += f"，应用前景：{prospect}"
                future_parts.append(seg)
            if future_parts:
                self.add_para("后续研究方向：" + "；".join(future_parts) + "。")

        # 应用价值表
        self._add_application_value_table()

    def _add_application_value_table(self):
        """应用价值表（3 列：场景/价值描述/采纳证明）"""
        scenarios = self._get_list("application_scenarios")
        rows = []
        for s in scenarios:
            if isinstance(s, dict):
                rows.append([str(s.get("scene", "")),
                             str(s.get("value", "")),
                             str(s.get("evidence", ""))])
        if rows:
            self.add_table(["应用场景", "价值描述", "采纳证明"], rows,
                           col_widths=[4.5, 5.5, 5.5], caption="应用价值汇总：")

    # ---------- 结尾段 ----------
    def _add_ending(self):
        """结尾（80~150 字）：毕设总结 + 致谢导师 + 朴素表态 + 此致/敬礼！"""
        ending = self._get("ending", default="")
        if ending:
            self.add_para(ending)
        else:
            advisor = self._get("advisor", default="")
            thesis_title = self._get("thesis_title", default="")
            summary_seg = (f"本毕业设计在课题方向取得了一定突破，"
                           f"完成《{thesis_title}》的研究目标。")
            if advisor:
                thanks_seg = f"感谢导师{advisor}一年来的悉心指导。"
            else:
                thanks_seg = "感谢导师一年来的悉心指导。"
            attitude_seg = ("无论评选结果如何，我都将以优秀毕业设计的标准严格要求自己，"
                            "在研究生阶段继续深入研究。")
            self.add_para(summary_seg + thanks_seg + attitude_seg)
        add_cizhi_paragraph(self.doc, "此致")
        add_jingli_paragraph(self.doc, "敬礼！")

    def _add_signature(self):
        """落款：右对齐，3 行（申请人 + 指导教师 + 日期）"""
        self.doc.add_paragraph()
        name = self._get("name", default="")
        advisor = self._get("advisor", default="")
        apply_date = self._get("apply_date", default="")
        if name:
            add_right_aligned_paragraph(self.doc, f"申请人：{name}")
        if advisor:
            add_right_aligned_paragraph(self.doc, f"指导教师：{advisor}")
        if apply_date:
            add_right_aligned_paragraph(self.doc, apply_date)

    # ---------- 摘要 ----------
    def _add_abstract(self):
        """摘要（300~500 字，4 句结构）：独立段落，置于申报书末尾或独立页"""
        abstract_pos = str(self._get("abstract_position", default="end")).lower().strip()
        if abstract_pos == "page":
            add_page_break(self.doc)
        # 摘要标题
        add_paragraph_with_format(
            self.doc, "摘要", font_name=FONT_HEI, font_size=SIZE_XIAO_SI, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            line_spacing=1.5, space_before=6, space_after=3)

        abstract = self._get("abstract", default="")
        if abstract:
            self.add_para(abstract)
            return

        # 自动生成摘要
        parts = self._build_abstract()
        if parts:
            self.add_para("".join(parts))

    def _build_abstract(self) -> str:
        """构建摘要 4 句结构（300~500 字）"""
        thesis_title = self._get("thesis_title", default="")
        research_problem = self._get("research_problem", default="")
        research_method = self._get("research_method", default="")
        technical_route = self._get("technical_route", default="")
        innovations = self._get_list("innovation_points")
        achievements = self._get_list("achievements")
        scenarios = self._get_list("application_scenarios")

        parts = []
        # 第 1 句：课题背景一句话
        bg_seg = f"本课题研究《{thesis_title}》" if thesis_title else "本课题"
        if research_problem:
            bg_seg += f"，针对{research_problem}"
        bg_seg += "。"
        parts.append(bg_seg)

        # 第 2 句：研究方法一句话
        if research_method:
            method_seg = f"采用{research_method}路线"
            if technical_route:
                method_seg += f"，基于{technical_route}"
            method_seg += "开展研究。"
            parts.append(method_seg)

        # 第 3 句：主要创新点 1-3 条
        if innovations:
            type_label = {"theoretical": "理论创新", "method": "方法创新", "application": "应用创新"}
            innov_parts = []
            for i, inv in enumerate(innovations[:3], start=1):
                if not isinstance(inv, dict):
                    continue
                type_str = type_label.get(inv.get("type", ""), "")
                name = inv.get("name", "")
                comparison = inv.get("comparison", "")
                seg = f"{'①②③'[i-1]}{type_str}——{name}"
                if comparison:
                    seg += f"（{comparison}）"
                innov_parts.append(seg)
            if innov_parts:
                parts.append("主要创新点：" + "；".join(innov_parts) + "。")

        # 第 4 句：主要成果与应用价值
        result_parts = []
        if achievements:
            cat_count = {}
            for a in achievements:
                if isinstance(a, dict):
                    cat = a.get("category", "")
                    cat_count[cat] = cat_count.get(cat, 0) + 1
            ach_seg = "、".join([f"{cat} {cnt} 项" for cat, cnt in cat_count.items()])
            if ach_seg:
                result_parts.append(f"研究成果包括{ach_seg}")
        if scenarios:
            result_parts.append(f"已应用于 {len(scenarios)} 个场景")
        if result_parts:
            parts.append("，".join(result_parts) + "。")

        return "".join(parts) if parts else ""

    # ===== 主构建方法 =====
    def _build_application(self):
        """主构建：标题 → 称呼 → 开头 → 课题选择 → 研究方法 → 创新点 → 成果应用 → 结尾 → 落款 → 摘要"""
        self._add_title()
        self._add_salutation()
        self._add_opening()
        self._add_topic_selection()
        self._add_research_method()
        self._add_innovation_points()
        self._add_achievements_value()
        self._add_ending()
        self._add_signature()
        self._add_abstract()

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：校验数据 + 构建 docx + 保存"""
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()
            self._build_application()
            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 优秀毕业设计/论文申报书已生成：{output_path}")
        return str(output_path)

    # ===== 数据校验 =====
    def _validate_data(self) -> List[str]:
        """校验数据完整性，返回缺失字段列表（不阻断生成）。含硬门槛 4 项校验。"""
        warnings = []
        version = self._resolve_word_count_version()
        level = self._resolve_apply_level()

        # P0 必采字段
        p0_fields = [("name", "申请人姓名"), ("college", "学院"), ("major", "专业"),
                     ("grade", "年级"), ("thesis_title", "毕设题目"),
                     ("advisor", "指导教师"), ("defense_score", "答辩成绩")]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        # 创新点校验
        innovations = self._get_list("innovation_points")
        if not innovations:
            warnings.append("缺少 创新点（innovation_points），建议至少 1 项")
        elif level == "school" and len(innovations) < 2:
            warnings.append(f"校级申报建议至少 2 项创新点，当前 {len(innovations)} 项")

        # 答辩成绩校验
        defense_score = str(self._get("defense_score", default=""))
        if defense_score:
            try:
                score_val = float(defense_score)
                if level == "school" and score_val < 90:
                    warnings.append(f"校级建议答辩成绩 ≥90 分（优秀），当前 {score_val} 分")
                elif level == "college" and score_val < 80:
                    warnings.append(f"院级建议答辩成绩 ≥80 分（良好），当前 {score_val} 分")
            except ValueError:
                pass

        # 排名校验
        rank_str = str(self._get("defense_rank", default=""))
        if rank_str and "/" in rank_str:
            try:
                rank_num = int(rank_str.split("/")[0])
                rank_total = int(rank_str.split("/")[1])
                if rank_total > 0:
                    pct = rank_num / rank_total * 100
                    if level == "school" and pct > 20:
                        warnings.append(f"校级建议答辩排名前 20%，当前前 {pct:.1f}%")
                    elif level == "college" and pct > 30:
                        warnings.append(f"院级建议答辩排名前 30%，当前前 {pct:.1f}%")
            except (ValueError, IndexError):
                pass

        # 查重率校验
        plagiarism = str(self._get("plagiarism_rate", default=""))
        if plagiarism:
            try:
                plag_num = float("".join(c for c in plagiarism if c.isdigit() or c == "."))
                gt = self._resolve_graduate_type()
                limits = {"undergraduate": 30, "master": 10, "doctor": 5}
                limit = limits.get(gt, 30)
                if plag_num > limit:
                    warnings.append(f"{gt} 查重率建议 ≤{limit}%，当前 {plag_num}%")
            except ValueError:
                pass

        # 标志成果校验
        if not self._get_list("achievements"):
            warnings.append("缺少 成果清单（achievements），建议至少 1 项论文/专利/软著")
        # 应用场景校验
        scenarios = self._get_list("application_scenarios")
        if not scenarios:
            warnings.append("缺少 应用场景（application_scenarios），建议至少 1 个")
        elif level == "school" and len(scenarios) < 3:
            warnings.append(f"校级建议至少 3 个应用场景，当前 {len(scenarios)} 个")
        # 研究阶段校验
        stages = self._get_list("research_stages")
        if not stages:
            warnings.append("缺少 研究阶段表（research_stages），建议 4-6 行")
        elif len(stages) < 4:
            warnings.append(f"研究阶段仅 {len(stages)} 行，建议 4-6 行")
        # 字数版本提示
        target = WORD_COUNT_TARGETS.get(version, 2700)
        warnings.append(f"ℹ 字数版本 {version}，目标 ~{target} 字")

        self.warnings = warnings
        if warnings:
            print("⚠️ 数据校验警告：", file=sys.stderr)
            for w in warnings:
                print(f"  - {w}", file=sys.stderr)
        return warnings


# ===== 默认示例数据 =====
DEFAULT_DATA = {
    # ----- 通用字段（11）+ 撰写控制（4）-----
    "name": "张明", "student_id": "2021123456", "gender": "男",
    "college": "计算机科学与技术学院", "major": "计算机科学与技术",
    "grade": "2021 级 大四", "class_name": "计科 2101 班", "phone": "138XXXXXXXX",
    "apply_level": "school", "apply_year": "2025 届", "apply_date": "2025 年 6 月 10 日",
    "salutation": "尊敬的评审委员会：",
    "word_count_version": "standard", "abstract_position": "end", "opening": "", "ending": "",
    # ----- 毕设字段（10）-----
    "thesis_title": "基于对比学习的小样本医学影像分割方法研究",
    "thesis_type": "experimental", "advisor": "张伟 教授", "advisor_title": "教授",
    "defense_score": "92", "defense_rank": "1/87", "plagiarism_rate": "8.5%",
    "defense_date": "2025 年 5 月 28 日",
    "research_duration": "2024.09-2025.05（9 个月）", "graduate_type": "undergraduate",
    # ----- 课题字段（4）-----
    "project_source": "导师国家自然科学基金项目子课题（项目编号 62XXXXXXXX）",
    "research_problem": "现有监督学习方法在小样本场景下精度不足的问题",
    "research_significance": "学术意义在于改进现有对比学习框架在医学影像领域的迁移性能，应用意义在于为基层医院提供低成本 AI 辅助诊断方案，与'健康中国 2030'战略对接",
    "research_goals": ["构建小样本医学影像数据集", "设计对比学习分割网络", "在公开数据集上验证方法有效性"],
    # ----- 研究方法字段（6）-----
    "research_method": "文献调研 → 实验设计 → 数据采集 → 模型训练 → 性能评估",
    "technical_route": "基于对比学习预训练 + Transformer 分割网络 + 迁移学习策略",
    "key_technologies": [
        {"name": "对比学习预训练", "application_step": "特征学习阶段", "problem_solved": "小样本场景特征学习不足"},
        {"name": "Transformer 分割网络", "application_step": "分割阶段", "problem_solved": "医学影像全局依赖建模"},
        {"name": "迁移学习策略", "application_step": "跨设备适配阶段", "problem_solved": "跨设备数据分布差异"},
    ],
    "data_material": "ISIC 2018 数据集（2 595 张皮肤镜图像）+ 自建数据集（300 张基层医院皮肤镜图像，已脱敏）",
    "research_stages": [
        {"stage": "1 文献调研", "time": "2024.09-2024.10", "content": "调研 50+ 篇文献", "output": "文献综述报告"},
        {"stage": "2 实验设计", "time": "2024.11-2024.12", "content": "设计实验方案", "output": "实验方案文档"},
        {"stage": "3 数据采集", "time": "2025.01-2025.02", "content": "采集 300 张图像", "output": "数据集"},
        {"stage": "4 模型训练", "time": "2025.03-2025.04", "content": "训练对比学习模型", "output": "训练日志"},
        {"stage": "5 性能评估", "time": "2025.04-2025.05", "content": "对比实验 + 消融实验", "output": "实验报告"},
    ],
    "experiments": [{"name": "小样本分割对比实验", "dataset": "ISIC 2018", "metric": "Dice",
                     "baseline": "0.785", "ours": "0.870", "improvement": "+8.5%"}],
    # ----- 创新点字段（1 列表）-----
    "innovation_points": [
        {"type": "theoretical", "name": "对比学习小样本分割理论框架",
         "description": "首次将对比学习引入皮肤镜图像分割任务，构建小样本场景下的分割理论框架",
         "comparison": "与 U-Net、DeepLab v3+ 相比，Dice 系数提升 8.5%",
         "support_material": "SCI 二区论文 1 篇（已发表）"},
        {"type": "method", "name": "多尺度特征融合模块 MSFF",
         "description": "设计跨尺度注意力机制，实现病灶边界精修",
         "comparison": "与基线 U-Net 相比，边界 IoU 提升 5.2%",
         "support_material": "发明专利 1 项（实质审查中）"},
        {"type": "application", "name": "基层医院皮肤镜辅助诊断部署",
         "description": "将所提方法部署于基层医院皮肤镜设备，实现 0.3 秒/张实时分割",
         "comparison": "已与 XX 县人民医院合作验证，诊断准确率提升 12%",
         "support_material": "合作验证报告 1 份"},
    ],
    # ----- 成果应用字段（4）-----
    "achievements": [
        {"category": "论文", "name": "基于对比学习的法律问答系统", "level": "SCI 二区",
         "author_order": "第一作者", "time": "2025.03"},
        {"category": "专利", "name": "一种基于对比学习的法律问答方法", "level": "发明专利",
         "author_order": "第一发明人", "time": "实质审查中"},
        {"category": "软著", "name": "皮肤镜图像分割软件 V1.0", "level": "软件著作权",
         "author_order": "第一完成人", "time": "2025.02"},
        {"category": "代码", "name": "MSFF 开源代码", "level": "GitHub 开源",
         "author_order": "维护者", "time": "2025.04"},
    ],
    "application_scenarios": [
        {"scene": "基层医院皮肤镜辅助诊断", "value": "诊断准确率提升 12%", "evidence": "XX 县人民医院合作报告"},
        {"scene": "在线医疗 AI 接口", "value": "已接入 XX 互联网医院", "evidence": "接入证明"},
        {"scene": "医学院教学辅助", "value": "已用于 XX 医学院皮肤科教学", "evidence": "教学应用证明"},
    ],
    "adoption_proofs": [{"unit": "XX 县人民医院", "content": "合作验证皮肤镜辅助诊断系统",
                          "date": "2025.04", "proof_no": "XX-2025-001"}],
    "future_research": [
        {"problem": "多模态融合诊断（CT + 皮肤镜）", "breakthrough": "提升诊断准确率至 95%+",
         "prospect": "应用于综合医院"},
        {"problem": "联邦学习框架下的隐私保护分割", "breakthrough": "实现跨医院数据协作",
         "prospect": "推广至全国医院"},
    ],
    "abstract": "",
}


# ===== CLI 入口 =====
def main():
    parser = argparse.ArgumentParser(
        description="优秀毕业设计/论文申报书 docx 生成器（4 段结构 + 摘要）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=("示例：\n"
                "  python build.py --data data.json --out output.docx\n"
                "  python build.py --demo --out demo.docx\n"
                "  python build.py --demo --word-count-version brief --out demo_brief.docx\n"
                "  python build.py --demo --apply-level college --out demo_college.docx\n"
                "\n"
                "JSON 字段定义详见 SKILL.md §8 信息采集清单与 §14 JSON Schema 完整字段定义。"),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True, help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true", help="使用内置示例数据生成演示文档")
    parser.add_argument("--word-count-version", type=str, default=None,
                        choices=list(WORD_COUNT_VERSIONS),
                        help="覆盖 word_count_version 字段（brief 短档 / standard 中档 / enhanced 长档）")
    parser.add_argument("--apply-level", type=str, default=None,
                        choices=list(APPLY_LEVELS),
                        help="覆盖 apply_level 字段（school 校级 / college 院级）")
    args = parser.parse_args()

    if args.demo:
        data = dict(DEFAULT_DATA)
        print("ℹ️ 使用内置示例数据生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    # CLI 覆盖字段
    if args.word_count_version:
        data["word_count_version"] = args.word_count_version
    if args.apply_level:
        data["apply_level"] = args.apply_level

    builder = OutstandingThesisDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
