#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
保研推免申请书 docx 生成器

格式标准：A4 / 页边距上下 2.54cm 左右 2.5cm / 正文宋体小四 1.5 倍行距首行缩进 2 字符 / 一级标题黑体三号居中 / 二级标题黑体小三左对齐 / 表格宋体五号居中。

栏目：封面 / 基本信息 / 学业情况 / 科研经历 / 个人陈述 / 推荐信情况 / 附件清单 / 签字栏。

使用：python build.py --data data.json --out output.docx
      python build.py --demo --out demo.docx

JSON 字段详见 SKILL.md 第三节。
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ===== 字体与格式常量 =====

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

SIZE_ER = Pt(22)
SIZE_XIAO_ER = Pt(18)
SIZE_SAN = Pt(16)
SIZE_XIAO_SAN = Pt(15)
SIZE_SI = Pt(14)
SIZE_XIAO_SI = Pt(12)
SIZE_WU = Pt(10.5)
SIZE_XIAO_WU = Pt(9)

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

# 禁用句检测列表（个人陈述中不得出现）
FORBIDDEN_PHRASES = [
    "梦寐以求", "从小就对", "倍感荣幸", "倍加珍惜",
    "师资雄厚", "学风严谨", "历史悠久", "学习氛围浓厚",
]


# ===== 工具函数 =====

def set_run_font(run, font_name: str = FONT_SONG,
                 font_size=SIZE_XIAO_SI, bold: bool = False,
                 color: Optional[RGBColor] = None) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_cell_text(cell, text: str, font_name: str = FONT_SONG,
                  font_size=SIZE_WU, bold: bool = False,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(doc, text: str, font_name: str = FONT_SONG,
                              font_size=SIZE_XIAO_SI, bold: bool = False,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT,
                              first_line_indent: bool = True,
                              line_spacing: float = 1.5,
                              space_before: float = 0,
                              space_after: float = 0):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return p


def add_heading_level1(doc, text: str):
    """一级标题：黑体三号，居中，段前段后 12pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
        space_before=12, space_after=12)


def add_heading_level2(doc, text: str):
    """二级标题：黑体小三，左对齐，段前段后 6pt"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
        space_before=6, space_after=6)


def add_body_paragraph(doc, text: str, indent: bool = True):
    """正文段落：宋体小四，1.5 倍行距，首行缩进 2 字符"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=indent,
        line_spacing=1.5)


def add_table_from_data(doc, headers: List[str], rows: List[List[str]],
                        col_widths: Optional[List[float]] = None,
                        first_col_bold: bool = False):
    """从数据创建表格，自动应用规范格式（表头加粗居中、数据居中）

    Args:
        doc: Document 对象
        headers: 表头列表
        rows: 数据行列表
        col_widths: 各列宽度（cm），可选
        first_col_bold: 首列是否加粗（用于"字段名-内容"型表格）
    """
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=SIZE_WU, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            bold = first_col_bold and j == 0
            set_cell_text(cells[j], val, font_name=FONT_SONG,
                          font_size=SIZE_WU, bold=bold)
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def dicts_to_rows(dicts: List[Dict[str, Any]], keys: List[str]) -> List[List[str]]:
    """将字典列表转为表格行（按 keys 顺序取值，自动转 str）"""
    rows = []
    for d in dicts:
        if not isinstance(d, dict):
            continue
        rows.append([safe_str(d.get(k, "")) for k in keys])
    return rows


def setup_page(doc):
    """设置 A4 页面与页边距（上下 2.54cm，左右 2.5cm）"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.bottom_margin = Cm(MARGIN_TOP_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_RIGHT_CM)
    section.right_margin = Cm(MARGIN_LEFT_RIGHT_CM)


def add_page_number(doc):
    """页脚添加居中页码（宋体五号）"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_WU)


def add_page_break(doc):
    """插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def count_chinese_chars(text: str) -> int:
    """统计中文字符数（用于个人陈述字数自检）"""
    if not text:
        return 0
    return sum(1 for c in text if '\u4e00' <= c <= '\u9fff')


def check_forbidden_phrases(text: str) -> List[str]:
    """检查个人陈述中是否含禁用句，返回命中的禁用句列表"""
    if not text:
        return []
    return [p for p in FORBIDDEN_PHRASES if p in text]


def parse_rank(rank_str: str):
    """解析专业排名字符串，返回 (分子, 分母) 或 None。支持 "1/87" "第1/87" 等格式"""
    if not rank_str:
        return None
    s = str(rank_str).replace("第", "").replace("名", "").replace(" ", "")
    if "/" not in s:
        return None
    parts = s.split("/")
    if len(parts) != 2:
        return None
    try:
        return int(parts[0]), int(parts[1])
    except ValueError:
        return None


def safe_str(val, default: str = "") -> str:
    """安全转字符串，None 返回默认值"""
    if val is None:
        return default
    if isinstance(val, float) and val.is_integer():
        return str(int(val))
    return str(val)


def join_english_scores(cet4, cet6, toefl, ielts) -> str:
    """合并英语成绩为单一字符串"""
    parts = []
    if cet4:
        parts.append(f"CET-4 {safe_str(cet4)}")
    if cet6:
        parts.append(f"CET-6 {safe_str(cet6)}")
    if toefl:
        parts.append(f"TOEFL {safe_str(toefl)}")
    if ielts:
        parts.append(f"IELTS {safe_str(ielts)}")
    return " / ".join(parts) if parts else ""


# ===== ApplicationDocBuilder 主类 =====

class ApplicationDocBuilder:
    """保研推免申请书 docx 构建器"""

    def __init__(self):
        self.doc = Document()
        setup_page(self.doc)
        add_page_number(self.doc)
        style = self.doc.styles["Normal"]
        style.font.name = FONT_SONG
        style.font.size = SIZE_XIAO_SI
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_SONG)
        rFonts.set(qn("w:ascii"), FONT_SONG)
        rFonts.set(qn("w:hAnsi"), FONT_SONG)
        self.data: Dict[str, Any] = {}
        self.warnings: List[str] = []
        self.statement_text: str = ""

    def _get(self, *keys, default=""):
        """安全取嵌套字段，缺字段返回默认值"""
        cur = self.data
        for k in keys:
            if not isinstance(cur, dict):
                return default
            cur = cur.get(k, default)
        return cur if cur is not None else default

    def add_h1(self, text):
        return add_heading_level1(self.doc, text)

    def add_h2(self, text):
        return add_heading_level2(self.doc, text)

    def add_para(self, text, indent=True):
        return add_body_paragraph(self.doc, text, indent=indent)

    def add_table(self, headers, rows, col_widths=None, first_col_bold=False):
        return add_table_from_data(self.doc, headers, rows, col_widths, first_col_bold)

    def add_page_break(self):
        add_page_break(self.doc)

    def _table_or_placeholder(self, headers, items, keys, col_widths,
                              placeholder, first_col_bold=False):
        """统一处理"有数据则建表，无数据则占位"逻辑"""
        if items and isinstance(items, list):
            rows = dicts_to_rows(items, keys)
            self.add_table(headers, rows, col_widths=col_widths,
                          first_col_bold=first_col_bold)
        else:
            self.add_para(placeholder)

    # ---- 封面 ----

    def _add_cover(self):
        """封面：黑体二号标题 + 副标题 + 5 行下划线信息"""
        for _ in range(2):
            self.doc.add_paragraph()

        add_paragraph_with_format(
            self.doc, "保研推免申请书",
            font_name=FONT_HEI, font_size=SIZE_ER, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
            space_before=12, space_after=12)

        channel = self._get("apply_channel", default="")
        target_school = self._get("target_school", default="")
        subtitle_parts = []
        if target_school:
            subtitle_parts.append(target_school)
        if channel:
            subtitle_parts.append(channel)
        if subtitle_parts:
            subtitle = f"（{' '.join(subtitle_parts)}）"
            add_paragraph_with_format(
                self.doc, subtitle,
                font_name=FONT_HEI, font_size=SIZE_SAN, bold=False,
                alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
                space_after=24)

        for _ in range(2):
            self.doc.add_paragraph()

        info_items = [
            ("申请人", self._get("applicant_name")),
            ("专    业", self._get("applicant_major")),
            ("申请院校", self._get("target_school")),
            ("申请方向", self._get("target_direction")),
            ("申请日期", self._get("apply_date")),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = 2.0
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run_label = p.add_run(f"{label}：")
            set_run_font(run_label, font_name=FONT_HEI,
                         font_size=SIZE_SI, bold=True)
            run_value = p.add_run(value if value else "                    ")
            set_run_font(run_value, font_name=FONT_SONG, font_size=SIZE_SI)
            run_value.underline = True

        self.add_page_break()

    # ---- 基本信息 ----

    def _add_basic_info(self):
        """一、基本信息表（10 行 2 列）"""
        self.add_h1("一、基本信息")
        basic = self._get("basic_info", default={})
        if not isinstance(basic, dict):
            basic = {}

        english_str = join_english_scores(
            basic.get("cet4") or self._get("cet4"),
            basic.get("cet6") or self._get("cet6"),
            basic.get("toefl") or self._get("toefl"),
            basic.get("ielts") or self._get("ielts"))

        rows = [
            ["姓名", basic.get("name", self._get("applicant_name"))],
            ["性别", basic.get("gender", self._get("applicant_gender"))],
            ["出生年月", basic.get("birth", self._get("applicant_birth"))],
            ["政治面貌", basic.get("political", self._get("applicant_political"))],
            ["学院", basic.get("college", self._get("applicant_college"))],
            ["专业", basic.get("major", self._get("applicant_major"))],
            ["班级", basic.get("class", self._get("applicant_class"))],
            ["联系电话", basic.get("phone", self._get("applicant_phone"))],
            ["邮箱", basic.get("email", self._get("applicant_email"))],
            ["外语水平", english_str or "（请填写 CET-4/6 + TOEFL/IELTS 成绩）"],
        ]
        rows = [[k, safe_str(v)] for k, v in rows]
        self.add_table(["字段", "内容"], rows,
                       col_widths=[4.0, 12.0], first_col_bold=True)

    # ---- 学业情况 ----

    def _add_academics(self):
        """二、学业情况（含学业硬指标段落+主修课程表+奖学金表）"""
        self.add_h1("二、学业情况")
        academics = self._get("academics", default={})
        if not isinstance(academics, dict):
            academics = {}

        # 2.1 学业硬指标段落
        self.add_h2("（一）学业硬指标")
        gpa = academics.get("gpa_6sem", "")
        weighted = academics.get("weighted_avg", "")
        rank = academics.get("rank", "")
        rank_percent = academics.get("rank_percent", "")
        recommend_qualified = academics.get("recommend_qualified", "")

        parts = []
        if gpa:
            parts.append(f"前 6 学期 GPA {gpa}")
        if weighted:
            parts.append(f"加权平均分 {weighted}")
        if rank:
            parts.append(f"专业排名第 {rank}")
        if rank_percent:
            parts.append(f"前 {rank_percent}")
        if recommend_qualified:
            obtained = recommend_qualified in ["是", True, "已取得"]
            parts.append(f"已{'取得' if obtained else '未取得'}本校推免资格")
        if parts:
            self.add_para("（" + "，".join(parts) + "。）")
        else:
            self.add_para("（请填写学业硬指标：GPA（含分母）/ 加权平均分 / 专业排名（X/N）/ 排名百分比 / 是否已取得推免资格。）")

        # 2.2 主修课程成绩表
        self.add_h2("（二）主修课程成绩")
        courses = academics.get("main_courses", [])
        self._table_or_placeholder(
            ["学期", "课程名", "学分", "成绩"], courses,
            ["semester", "course", "credit", "score"],
            col_widths=[3.0, 7.5, 2.5, 3.0],
            placeholder="（请填写主修课程成绩表，10~15 门专业核心课与高分课程，按学期正序排列。）")

        # 2.3 奖学金获得情况表
        self.add_h2("（三）奖学金获得情况")
        scholarships = academics.get("scholarships", [])
        self._table_or_placeholder(
            ["时间", "奖项全称", "级别"], scholarships,
            ["time", "name", "level"],
            col_widths=[3.0, 9.0, 4.0],
            placeholder="（请填写奖学金获得情况表，按时间倒序，注明级别（国家/省/校/院）。）")

        # 2.4 推免资格确认
        self.add_h2("（四）推免资格确认")
        if recommend_qualified:
            confirm = (f"申请人已取得本校 {self._get('applicant_major', 'XX 专业')} "
                       f"2026 届推免资格（{academics.get('recommend_date', '2025.09')} 公布）。")
            self.add_para(confirm)
        else:
            self.add_para("（请填写推免资格确认：是否已取得本校推免资格，公布日期。）")

    # ---- 科研经历【重点】 ----

    def _add_research(self):
        """三、科研经历（大创表+论文表+专利表+竞赛表+实验室叙述）"""
        self.add_h1("三、科研经历")
        research = self._get("research", default={})
        if not isinstance(research, dict):
            research = {}

        # 3.1 大创/科研立项表
        self.add_h2("（一）大创/科研立项")
        self._table_or_placeholder(
            ["项目名称", "级别", "角色", "起止时间", "项目产出"],
            research.get("projects", []),
            ["name", "level", "role", "duration", "output"],
            col_widths=[4.5, 2.0, 1.8, 2.7, 5.0],
            placeholder="（请填写大创/科研立项表，每项含项目名/级别/角色/起止/产出。）")

        # 3.2 发表论文表
        self.add_h2("（二）发表论文")
        self._table_or_placeholder(
            ["题目", "期刊或会议", "级别", "排序", "状态", "时间"],
            research.get("papers", []),
            ["title", "journal", "level", "author_order", "status", "time"],
            col_widths=[4.5, 3.0, 1.8, 1.2, 1.5, 2.0],
            placeholder="（请填写发表论文表，每篇含题目/期刊/级别/作者排序/状态/时间。）")

        # 3.3 专利表
        self.add_h2("（三）专利")
        self._table_or_placeholder(
            ["名称", "类型", "申请号", "状态", "发明人排序"],
            research.get("patents", []),
            ["name", "type", "apply_no", "status", "inventor_order"],
            col_widths=[5.0, 2.0, 3.5, 2.5, 3.0],
            placeholder="（请填写专利表，每项含名称/类型/申请号/状态/发明人排序。）")

        # 3.4 学科竞赛获奖表
        self.add_h2("（四）学科竞赛获奖")
        self._table_or_placeholder(
            ["竞赛全称", "奖项", "级别", "获奖时间"],
            research.get("contests", []),
            ["name", "award", "level", "time"],
            col_widths=[6.5, 2.5, 2.5, 4.5],
            placeholder="（请填写学科竞赛获奖表，每项含竞赛名/奖项/级别/时间。）")

        # 3.5 实验室/导师科研参与叙述
        self.add_h2("（五）实验室/导师科研参与")
        lab = research.get("lab_participation", "")
        if lab:
            self.add_para(lab)
        else:
            self.add_para("（请填写实验室/导师科研参与叙述，100~300 字，含进入时间/承担工作/产出收获/与申请方向关联。）")

    # ---- 个人陈述【重点】 ----

    def _add_personal_statement(self):
        """四、个人陈述（800~1500 字，三段论）"""
        self.add_h1("四、个人陈述")
        statement = self._get("personal_statement", default={})
        if isinstance(statement, str):
            para1, para2, para3, full_text = statement, "", "", statement
        elif isinstance(statement, dict):
            para1 = statement.get("paragraph_1", "")
            para2 = statement.get("paragraph_2", "")
            para3 = statement.get("paragraph_3", "")
            full_text = para1 + para2 + para3
        else:
            para1 = para2 = para3 = full_text = ""

        self.statement_text = full_text
        cn_count = count_chinese_chars(full_text)

        # 字数提示
        if cn_count > 0:
            self.add_para(f"（个人陈述共 {cn_count} 字，建议 800~1500 字。）", indent=False)

        # 三段标题与内容
        self.add_h2("（一）为什么选这个方向")
        if para1:
            self.add_para(para1)
        else:
            self.add_para("（请填写第一段，300~500 字。结构：具体事件 → 兴趣触发 → 深入探索 → 方向确认。必含本科科研具体经历、对方向的核心认识、下一步主动行动。）")

        self.add_h2("（二）为什么选这个学校")
        if para2:
            self.add_para(para2)
        else:
            self.add_para("（请填写第二段，300~500 字。结构：导师匹配 + 平台资源 + 已建立联系。必含目标院校具体优势数据、与拟申请导师的具体联系、目标院校具体资源。）")

        self.add_h2("（三）研究生阶段计划")
        if para3:
            self.add_para(para3)
        else:
            self.add_para("（请填写第三段，200~500 字。结构：课程学习 + 科研计划 + 长期目标。必含具体课程、具体方向与产出预期、长期职业规划。）")

        # 禁用句检测与字数检查已在 _validate_data 中完成，避免重复警告

    # ---- 推荐信情况 ----

    def _add_recommendations(self):
        """五、推荐信情况（2 位推荐人表格）"""
        self.add_h1("五、推荐信情况")
        recs = self._get("recommendations", default=[])
        if not isinstance(recs, list):
            recs = []

        if recs:
            self._table_or_placeholder(
                ["推荐人", "职称", "关系", "研究方向", "联系方式"], recs,
                ["name", "title", "relation", "research_area", "contact"],
                col_widths=[2.5, 1.8, 3.5, 4.2, 4.0], placeholder="")
            send_note = self._get("recommendations_send_note", default="")
            if send_note:
                self.add_para(send_note)
            else:
                self.add_para("（推荐信将由推荐人签字后直接寄送至申请院校招生办公室。）")
        else:
            self.add_para("（请填写推荐信情况表，2 位推荐人，每人含姓名/职称/关系/研究方向/联系方式。一般推荐人 1 为本校科研导师，推荐人 2 为任课教师或班主任。）")

    # ---- 附件清单 ----

    def _add_attachments(self):
        """六、附件清单（6~8 项，编号列表）"""
        self.add_h1("六、附件清单")
        attachments = self._get("attachments", default=[])
        if not isinstance(attachments, list):
            attachments = []

        if attachments:
            for i, a in enumerate(attachments, 1):
                if isinstance(a, dict):
                    text = f"{i}. {safe_str(a.get('name', ''))}"
                    if a.get("count"):
                        text += f" {a.get('count')} 份"
                    if a.get("note"):
                        text += f"（{a.get('note')}）"
                else:
                    text = f"{i}. {safe_str(a)}"
                self.add_para(text, indent=False)
        else:
            default_attachments = [
                "本科成绩单 1 份（教务处盖章）",
                "CET-6 与 TOEFL 成绩单复印件各 1 份",
                "论文录用通知 / 投稿确认邮件 1 份",
                "专利受理通知书 1 份",
                "获奖证书复印件各 1 份",
                "大创立项书与结题证明 1 份",
                "推荐信 2 封（由推荐人签字后直接寄送）",
            ]
            for i, a in enumerate(default_attachments, 1):
                self.add_para(f"{i}. {a}", indent=False)

        # 签字栏
        self.add_h1("七、签字栏")
        for _ in range(4):
            self.doc.add_paragraph()
        self.add_para(
            "申请人签字：____________________    "
            "日期：______年____月____日", indent=False)
        for _ in range(3):
            self.doc.add_paragraph()
        self.add_para(
            "学院推免工作小组盖章：____________________    "
            "日期：______年____月____日", indent=False)

    # ---- 主构建方法 ----

    def build(self, data: Dict[str, Any], output_path: str) -> str:
        """主构建方法：编排 7 大栏目，生成 docx

        Args:
            data: 申报书字段字典
            output_path: 输出 docx 路径
        Returns: 实际保存路径
        """
        try:
            self.data = data if isinstance(data, dict) else {}
            self._validate_data()

            self._add_cover()
            self._add_basic_info()
            self._add_academics()
            self._add_research()
            self._add_personal_statement()
            self._add_recommendations()
            self._add_attachments()

            # 打印累积警告
            if self.warnings:
                print("⚠️ 数据校验警告：", file=sys.stderr)
                for w in self.warnings:
                    print(f"  - {w}", file=sys.stderr)

            return self._save(output_path)
        except Exception as e:
            sys.stderr.write(f"❌ 生成失败：{e}\n")
            raise

    def _save(self, output_path: str) -> str:
        """保存文档，自动创建目录

        Args:
            output_path: 输出 docx 路径
        Returns: 实际保存路径
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        print(f"✅ 申报书已生成：{output_path}")
        return str(output_path)

    # ---- 数据校验 ----

    def _validate_data(self) -> List[str]:
        """校验数据完整性，生成警告（不阻断生成）

        校验维度：P0 必填 / GPA / 排名（前 5%）/ CET-6（≥425）/ 科研项目必填 / 个人陈述字数（800~1500）与禁用句 / 推荐人数量。
        警告累积到 self.warnings，build() 末尾统一打印到 stderr。
        """
        warnings = []

        # P0 必填字段
        p0_fields = [
            ("applicant_name", "申请人姓名"), ("applicant_major", "专业"),
            ("target_school", "申请学校"), ("target_direction", "申请方向"),
            ("apply_date", "申请日期"),
        ]
        for key, name in p0_fields:
            if not self._get(key):
                warnings.append(f"缺少 {name}（{key}）")

        # 学业硬指标校验
        academics = self._get("academics", default={})
        if not isinstance(academics, dict):
            academics = {}

        gpa = academics.get("gpa_6sem", "")
        if not gpa:
            warnings.append("缺少 GPA（academics.gpa_6sem），保研硬指标")

        rank = academics.get("rank", "")
        if rank:
            parsed = parse_rank(rank)
            if parsed:
                rank_num, rank_den = parsed
                if rank_den > 0:
                    percent = rank_num / rank_den * 100
                    if percent > 5:
                        warnings.append(
                            f"专业排名 {rank}，前 {percent:.1f}%，超过多数 985 院校门槛线（前 5%）")
        else:
            warnings.append("缺少专业排名（academics.rank），保研硬指标")

        # CET-6 校验（≥425）
        basic = self._get("basic_info", default={})
        if not isinstance(basic, dict):
            basic = {}
        cet6 = academics.get("cet6") or self._get("cet6") or basic.get("cet6")
        if cet6:
            try:
                cet6_str = safe_str(cet6).split()[0] if " " in safe_str(cet6) else safe_str(cet6)
                cet6_num = int(cet6_str)
                if cet6_num < 425:
                    warnings.append(f"CET-6 成绩 {cet6_num} 低于 425 分门槛")
            except (ValueError, TypeError):
                pass
        else:
            warnings.append("缺少 CET-6 成绩（多数院校要求 ≥425）")

        # 科研项目必填（至少 1 项大创或论文）
        research = self._get("research", default={})
        if not isinstance(research, dict):
            research = {}
        projects = research.get("projects", []) or []
        papers = research.get("papers", []) or []
        if not projects and not papers:
            warnings.append("缺少科研项目/论文，保研软实力几乎为空，建议补充")

        # 个人陈述字数与禁用句
        statement = self._get("personal_statement", default={})
        if isinstance(statement, dict):
            full = (statement.get("paragraph_1", "") +
                    statement.get("paragraph_2", "") +
                    statement.get("paragraph_3", ""))
        elif isinstance(statement, str):
            full = statement
        else:
            full = ""

        cn = count_chinese_chars(full)
        if full and cn < 800:
            warnings.append(f"个人陈述字数 {cn} 不足 800 字")
        elif full and cn > 1500:
            warnings.append(f"个人陈述字数 {cn} 超过 1500 字")

        forbidden = check_forbidden_phrases(full)
        if forbidden:
            warnings.append(f"个人陈述含禁用句：{', '.join(forbidden)}（一票否决项）")

        # 推荐人数量
        recs = self._get("recommendations", default=[])
        if isinstance(recs, list) and len(recs) < 2:
            warnings.append(f"推荐人仅 {len(recs)} 位，建议 2 位")

        self.warnings.extend(warnings)
        return warnings


# ===== 默认示例数据 =====

DEFAULT_DATA = {
    "applicant_name": "张三", "applicant_gender": "男",
    "applicant_birth": "2003 年 5 月", "applicant_political": "中共党员",
    "applicant_college": "计算机科学与技术学院",
    "applicant_major": "计算机科学与技术", "applicant_class": "计科 2201 班",
    "applicant_phone": "138XXXXXXXX",
    "applicant_email": "zhangsan@xxx.edu.cn",
    "target_school": "XX 大学", "target_college": "计算机学院",
    "target_direction": "自然语言处理", "target_advisor": "XX 教授",
    "apply_channel": "夏令营", "apply_date": "2025 年 6 月 10 日",
    "basic_info": {
        "name": "张三", "gender": "男", "birth": "2003 年 5 月",
        "political": "中共党员", "college": "计算机科学与技术学院",
        "major": "计算机科学与技术", "class": "计科 2201 班",
        "phone": "138XXXXXXXX", "email": "zhangsan@xxx.edu.cn",
        "cet4": "580", "cet6": "580", "toefl": "105",
    },
    "academics": {
        "gpa_6sem": "3.92/4.0", "weighted_avg": "91.5/100",
        "rank": "1/87", "rank_percent": "1.1%",
        "recommend_qualified": "是", "recommend_date": "2025.09",
        "main_courses": [
            {"semester": "第 1 学期", "course": "高等数学（上）", "credit": "5", "score": "95"},
            {"semester": "第 1 学期", "course": "C 语言程序设计", "credit": "3", "score": "96"},
            {"semester": "第 2 学期", "course": "高等数学（下）", "credit": "5", "score": "93"},
            {"semester": "第 2 学期", "course": "数据结构", "credit": "4", "score": "94"},
            {"semester": "第 3 学期", "course": "离散数学", "credit": "3", "score": "92"},
            {"semester": "第 3 学期", "course": "计算机组成原理", "credit": "4", "score": "90"},
            {"semester": "第 4 学期", "course": "操作系统", "credit": "4", "score": "93"},
            {"semester": "第 4 学期", "course": "计算机网络", "credit": "4", "score": "91"},
            {"semester": "第 5 学期", "course": "机器学习", "credit": "3", "score": "95"},
            {"semester": "第 5 学期", "course": "自然语言处理", "credit": "3", "score": "96"},
            {"semester": "第 6 学期", "course": "深度学习", "credit": "3", "score": "94"},
            {"semester": "第 6 学期", "course": "数据挖掘", "credit": "3", "score": "92"},
        ],
        "scholarships": [
            {"time": "2025.05", "name": "国家奖学金（专业前 1%）", "level": "国家级"},
            {"time": "2024.11", "name": "校级一等奖学金（专业前 5%）", "level": "校级"},
            {"time": "2024.05", "name": "全国大学生英语竞赛三等奖", "level": "国家级"},
            {"time": "2023.11", "name": "校级一等奖学金", "level": "校级"},
        ],
    },
    "research": {
        "projects": [
            {"name": "基于对比学习的法律问答系统", "level": "国家级",
             "role": "负责人", "duration": "2024.03-2025.03",
             "output": "准确率提升 7pp，论文投 EMNLP 2024"},
            {"name": "XX 老师课题组低资源对话项目", "level": "校级",
             "role": "参与", "duration": "2023.09-2024.06",
             "output": "中文数据集 5000 条，1 份规范文档"},
        ],
        "papers": [
            {"title": "XX 法律问答中的对比学习", "journal": "EMNLP 2024",
             "level": "CCF-B", "author_order": "1", "status": "在审",
             "time": "2024.06"},
            {"title": "基于预训练模型的中文分词", "journal": "计算机应用",
             "level": "中文核心", "author_order": "1", "status": "已发表",
             "time": "2023.05"},
        ],
        "patents": [
            {"name": "一种基于对比学习的文本分类方法", "type": "发明专利",
             "apply_no": "CN2024XXXXXXX", "status": "实质审查",
             "inventor_order": "1（学生一作）"},
        ],
        "contests": [
            {"name": "全国大学生数学建模竞赛", "award": "国家二等奖",
             "level": "国家级", "time": "2024.11"},
            {"name": "ACM-ICPC 亚洲区域赛", "award": "银奖",
             "level": "国际级", "time": "2024.10"},
            {"name": "中国大学生计算机设计大赛", "award": "二等奖",
             "level": "国家级", "time": "2024.06"},
        ],
        "lab_participation": ("2023.09 起进入 XX 教授课题组，参与《低资源场景下的对话理解》项目。"
            "主要负责：构建含 5000 条多轮对话的中文数据集；复现 3 个基线模型并优化对比学习模块；"
            "独立完成数据标注规范文档 1 份。期间掌握了 PyTorch、HuggingFace Transformers 等工具，"
            "对低资源 NLP 有了系统认识，为研究生阶段在该方向深入研究奠定基础。"),
    },
    "personal_statement": {
        "paragraph_1": ("我对自然语言处理的兴趣，源于大二参加校级大创项目《基于对比学习的法律问答系统》时的经历，该经历决定了我本科后续三年的学习与科研走向。"
            "当时我们用 BERT 微调做意图识别，准确率卡在 82% 上不去，后来尝试对比学习才突破到 89%，这件事让我深受启发。"
            "这次经历让我意识到：表征学习是 NLP 的核心问题之一，也是限制低资源场景性能的关键瓶颈。"
            "此后我阅读了 XX 老师 2024 年 ACL 论文《XX》，复现了其中的对比学习框架，并在自己项目上做了改进实验，"
            "进一步确认了在该方向深入研究的决心，并立志在该方向做出有影响力的工作。"
            "本科三年级上学期，我选修了《自然语言处理》与《深度学习》课程，取得了 96 分与 94 分的成绩，"
            "系统补齐了理论基础。这段经历也让我系统掌握了 PyTorch、HuggingFace Transformers 等工具，"
            "具备了独立开展 NLP 研究的初步能力。基于上述背景，我希望在研究生阶段继续深入这一方向，特别是在低资源场景下的鲁棒表征问题上。"),
        "paragraph_2": ("选择 XX 大学基于三点考虑：一是贵校计算机学院在自然语言处理方向拥有 3 位 IEEE Fellow，"
            "近 3 年在 ACL、EMNLP、NAACL 发表论文 20+ 篇，研究实力全国领先；"
            "二是已与拟申请导师 XX 教授邮件沟通 3 次，老师的研究方向（对比学习与知识蒸馏融合）"
            "与本人在对比学习上的兴趣高度契合，老师建议我提前阅读其课题组 2024 年 ACL 论文并复现核心代码；"
            "三是贵校拥有 XX 国家重点实验室，配备 A100 GPU 集群 8 卡节点，"
            "能提供本方向所需的大规模预训练计算资源，这是多数院校难以提供的硬件条件。"
            "以上三点让我确信贵校是开展研究生阶段 NLP 研究的最佳选择，也是我能将本科阶段积累延伸为系统化研究成果的理想平台。"),
        "paragraph_3": ("研究生阶段我的计划分三步：研一系统修读《深度学习》《自然语言处理》《统计学习方法》等核心课程，"
            "重点补强数学基础，同时跟随导师完成文献调研 50 篇；研二在对比学习与知识蒸馏融合方向深入研究，"
            "重点突破低资源场景下的鲁棒表征问题，预期在 CCF-B 类会议发表论文 1 篇，并申请发明专利 1 项；"
            "研三总结成果申请博士，长期目标是成为高校 NLP 方向研究者，推动低资源场景下的语言理解技术落地。"
            "为此我计划每季度阅读 20 篇顶会论文，每年参与 1 次学术会议交流，与导师保持紧密讨论，逐步形成独立的研究品味与能力。"
            "我相信，贵校提供的优质学术资源与导师悉心指导，能够帮助我实现上述规划，我也将以最大热情投入研究生学习与研究。"),
    },
    "recommendations": [
        {"name": "李教授", "title": "教授", "relation": "科研导师",
         "research_area": "自然语言处理", "contact": "li@xxx.edu.cn"},
        {"name": "王副教授", "title": "副教授", "relation": "任课教师（机器学习）",
         "research_area": "机器学习", "contact": "wang@xxx.edu.cn"},
    ],
    "recommendations_send_note": ("推荐信将由推荐人签字后直接寄送至 XX 大学计算机学院招生办公室，"
        "预计 2025 年 6 月 25 日前寄出。"),
    "attachments": [
        {"name": "本科成绩单", "count": "1", "note": "教务处盖章"},
        {"name": "CET-6 与 TOEFL 成绩单复印件", "count": "各 1", "note": ""},
        {"name": "EMNLP 2024 投稿确认邮件", "count": "1", "note": "含在审证明"},
        {"name": "《计算机应用》论文录用通知与见刊页面", "count": "1", "note": ""},
        {"name": "发明专利受理通知书", "count": "1", "note": ""},
        {"name": "国家奖学金证书、数学建模国二、ACM-ICPC 银奖证书复印件", "count": "各 1", "note": ""},
        {"name": "国家级大创立项书与结题证明", "count": "1", "note": ""},
        {"name": "推荐信", "count": "2", "note": "由推荐人签字后直接寄送"},
    ],
}


# ===== CLI 入口 =====

def main():
    parser = argparse.ArgumentParser(
        description="保研推免申请书 docx 生成器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例：\n"
            "  python build.py --data data.json --out output.docx\n"
            "  python build.py --demo --out demo.docx\n"
            "\n"
            "JSON 字段定义详见 SKILL.md 第三节。"
        ),
    )
    parser.add_argument("--data", type=str, default=None,
                        help="数据 JSON 文件路径（与 --demo 二选一）")
    parser.add_argument("--out", type=str, required=True,
                        help="输出 docx 文件路径")
    parser.add_argument("--demo", action="store_true",
                        help="使用内置示例数据生成演示文档")
    args = parser.parse_args()

    if args.demo:
        data = DEFAULT_DATA
        print("ℹ️ 使用内置示例数据生成演示文档")
    elif args.data:
        if not os.path.exists(args.data):
            sys.stderr.write(f"❌ 数据文件不存在：{args.data}\n")
            sys.exit(1)
        try:
            with open(args.data, "r", encoding="utf-8") as f:
                data = json.load(f)
        except json.JSONDecodeError as e:
            sys.stderr.write(f"❌ JSON 解析失败：{e}\n")
            sys.exit(1)
    else:
        sys.stderr.write("❌ 必须提供 --data 或 --demo 参数\n")
        parser.print_help()
        sys.exit(1)

    builder = ApplicationDocBuilder()
    try:
        builder.build(data, args.out)
    except Exception as e:
        sys.stderr.write(f"❌ 生成失败：{e}\n")
        sys.exit(1)


if __name__ == "__main__":
    main()
