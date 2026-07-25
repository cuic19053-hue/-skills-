#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
review_simulator.py — 大学生申报书评审模拟器

本模块基于 review_criteria.json 中 30 个子 skill 的评审标准，
对学生填写的申报书进行多维度模拟评审，输出评分、评语、等级与改进建议。

主要功能：
    1. ReviewSimulator(skill_name, application_text, applicant_data) — 主类
    2. simulate_review() — 模拟评审，返回评分 + 评语
    3. generate_review_report(out_path) — 生成评审报告 docx
    4. compare_applicants(applicant_list) — 多人对比排序
    5. CLI 接口：python review_simulator.py --skill national_scholarship --data data.json

依赖：
    - review_criteria.json（评审标准配置）
    - plagiarism_checker.py（查重预检）
    - python-docx（生成报告；如未安装则降级输出 txt/md）

Author: T38 子智能体
Version: 1.0.0
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from dataclasses import dataclass, field, asdict
from typing import Any, Dict, List, Optional, Tuple

# ---------------------------------------------------------------------------
# 路径常量
# ---------------------------------------------------------------------------
THIS_DIR = os.path.dirname(os.path.abspath(__file__))
CRITERIA_PATH = os.path.join(THIS_DIR, "review_criteria.json")

# 同目录的查重模块
sys.path.insert(0, THIS_DIR)
try:
    from plagiarism_checker import check_plagiarism, PlagiarismReport  # type: ignore
except ImportError:  # pragma: no cover
    check_plagiarism = None  # type: ignore
    PlagiarismReport = None  # type: ignore

# ---------------------------------------------------------------------------
# 数据类
# ---------------------------------------------------------------------------
@dataclass
class DimensionScore:
    """单维度评分。"""
    name: str
    weight: float
    max_score: float
    raw_score: float          # 0-100 原始分
    weighted_score: float     # 加权后分（raw * weight / 100）
    comment: str = ""
    evidence: List[str] = field(default_factory=list)
    grade: str = ""           # A/B/C/D


@dataclass
class ReviewResult:
    """完整评审结果。"""
    skill_name: str
    skill_label: str
    applicant_name: str = ""
    total_score: float = 0.0
    max_score: float = 100.0
    grade: str = ""           # A/B/C/D
    passed: bool = False
    dimensions: List[DimensionScore] = field(default_factory=list)
    overall_comment: str = ""
    strengths: List[str] = field(default_factory=list)
    weaknesses: List[str] = field(default_factory=list)
    improvements: List[str] = field(default_factory=list)
    veto_triggered: List[str] = field(default_factory=list)
    plagiarism: Optional[Dict] = None

    def to_dict(self) -> Dict:
        d = asdict(self)
        return d

    def to_json(self, indent: int = 2) -> str:
        return json.dumps(self.to_dict(), ensure_ascii=False, indent=indent)


# ---------------------------------------------------------------------------
# 评审模拟器主类
# ---------------------------------------------------------------------------
class ReviewSimulator:
    """
    评审模拟器。

    用法：
        sim = ReviewSimulator(
            skill_name="national_scholarship",
            application_text="本人自入学以来……",
            applicant_data={
                "name": "张三",
                "gpa_rank_percent": 5,
                "comprehensive_rank_percent": 8,
                "papers": [{"level": "SCI", "author_order": 1}],
                "competitions": [{"level": "国家级", "award": "一等奖"}],
                "volunteer_hours": 80,
            },
        )
        result = sim.simulate_review()
        sim.generate_review_report("/tmp/report.docx")
    """

    GRADE_THRESHOLDS = {"A": 90, "B": 80, "C": 70, "D": 0}
    GRADE_LABELS = {
        "A": "优（强烈推荐通过）",
        "B": "良（建议通过）",
        "C": "中（待定/有保留通过）",
        "D": "差（不建议通过）",
    }

    def __init__(
        self,
        skill_name: str,
        application_text: str = "",
        applicant_data: Optional[Dict[str, Any]] = None,
    ) -> None:
        self.skill_name = skill_name
        self.application_text = application_text or ""
        self.applicant_data = applicant_data or {}
        self.criteria: Dict = self._load_criteria()
        self.skill_config: Dict = self.criteria.get(skill_name)
        if not self.skill_config:
            raise ValueError(
                f"未找到 skill '{skill_name}' 的评审标准，"
                f"可用: {', '.join(k for k in self.criteria.keys() if k != '_meta')}"
            )
        self.applicant_name: str = self.applicant_data.get("name", "未署名申请人")

    # -------------------- 准入判断 --------------------
    def _load_criteria(self) -> Dict:
        if not os.path.exists(CRITERIA_PATH):
            raise FileNotFoundError(f"评审标准配置文件不存在: {CRITERIA_PATH}")
        with open(CRITERIA_PATH, "r", encoding="utf-8") as f:
            return json.load(f)

    def check_hard_gates(self) -> Tuple[bool, List[str]]:
        """
        检查硬门槛（一票否决项的客观部分）。
        返回 (是否通过, 失败原因列表)。
        """
        gates = self.skill_config.get("hard_gates", {})
        failures: List[str] = []
        data = self.applicant_data

        # GPA 排名
        if "gpa_rank_top" in gates:
            rank = data.get("gpa_rank_percent")
            if rank is None:
                failures.append(f"未提供 GPA 排名百分比（硬门槛要求 {gates['gpa_rank_top']}）")
            else:
                limit = self._parse_percent(gates["gpa_rank_top"])
                if rank > limit:
                    failures.append(
                        f"GPA 排名 {rank}% 超过硬门槛 {gates['gpa_rank_top']}"
                    )

        # 综合排名
        if "comprehensive_rank_top" in gates:
            rank = data.get("comprehensive_rank_percent")
            if rank is None:
                failures.append(
                    f"未提供综合排名百分比（硬门槛要求 {gates['comprehensive_rank_top']}）"
                )
            else:
                limit = self._parse_percent(gates["comprehensive_rank_top"])
                if rank > limit:
                    failures.append(
                        f"综合排名 {rank}% 超过硬门槛 {gates['comprehensive_rank_top']}"
                    )

        # 必修课不及格
        if gates.get("required_courses_fail") is False:
            if data.get("has_failed_required"):
                failures.append("必修课有不及格（硬门槛禁止）")

        # 违纪记录
        if gates.get("no_discipline_record"):
            if data.get("has_discipline_record"):
                failures.append("存在违纪处分记录（硬门槛禁止）")

        # 家庭经济困难
        if gates.get("poverty_registered"):
            if not data.get("poverty_registered"):
                failures.append("未通过家庭经济困难认定（硬门槛要求）")

        # 公司注册（创业实践）
        if gates.get("company_registered"):
            if not data.get("company_registered"):
                failures.append("公司未实际注册（创业实践硬门槛）")

        # 党员身份
        if gates.get("party_member_required"):
            if not data.get("is_party_member"):
                failures.append("非中共党员（硬门槛要求）")

        # 学生干部
        if gates.get("student_cadre_required"):
            if not data.get("is_student_cadre"):
                failures.append("无学生干部任职经历（硬门槛要求）")

        # 应届毕业
        if gates.get("graduate_required"):
            if not data.get("is_graduate"):
                failures.append("非应届毕业生（硬门槛要求）")

        # 英语成绩
        if gates.get("english_required"):
            if not data.get("english_qualified"):
                failures.append("英语成绩未达标（硬门槛要求）")

        # 指导教师
        if gates.get("advisor_required"):
            if not data.get("advisor"):
                failures.append("无指导教师（硬门槛要求）")

        return (len(failures) == 0, failures)

    @staticmethod
    def _parse_percent(s: str) -> float:
        """'10%' -> 10.0"""
        return float(str(s).replace("%", "").strip())

    # -------------------- 模拟评审 --------------------
    def simulate_review(self) -> ReviewResult:
        """执行完整评审流程。"""
        result = ReviewResult(
            skill_name=self.skill_name,
            skill_label=self.skill_config.get("name", self.skill_name),
            applicant_name=self.applicant_name,
            max_score=self.skill_config.get("max_score", 100),
        )

        # 1. 查重预检
        plagiarism_dict: Optional[Dict] = None
        if self.application_text and check_plagiarism is not None:
            try:
                pr: PlagiarismReport = check_plagiarism(
                    text=self.application_text,
                    skill_name=self.skill_name,
                )
                plagiarism_dict = {
                    "overall_similarity": pr.overall_similarity,
                    "max_similarity": pr.max_similarity,
                    "hit_count": pr.hit_count,
                    "grade": pr.grade,
                    "passed": pr.passed,
                }
                if not pr.passed:
                    result.veto_triggered.append(
                        f"查重不通过：整体相似度 {pr.overall_similarity*100:.1f}%，"
                        f"等级 {pr.grade}（一票否决）"
                    )
            except Exception as e:  # pragma: no cover
                plagiarism_dict = {"error": str(e)}
        result.plagiarism = plagiarism_dict

        # 2. 硬门槛
        gate_ok, gate_failures = self.check_hard_gates()
        if not gate_ok:
            result.veto_triggered.extend(gate_failures)

        # 3. 一票否决项（配置文件中显式声明）
        veto_items = self._check_veto_items()
        result.veto_triggered.extend(veto_items)

        # 4. 多维度评分
        result.dimensions = self._score_dimensions()

        # 5. 汇总
        total = sum(d.weighted_score for d in result.dimensions)
        # 归一化到 max_score
        if result.max_score and result.max_score != 100:
            total = total * result.max_score / 100.0
        result.total_score = round(total, 2)

        # 6. 等级判定
        if result.veto_triggered:
            result.grade = "D"
            result.passed = False
            result.total_score = min(result.total_score, 59.99)
        else:
            for grade, threshold in self.GRADE_THRESHOLDS.items():
                if result.total_score >= threshold:
                    result.grade = grade
                    break
            result.passed = result.grade in ("A", "B")

        # 7. 评语
        result.overall_comment = self._generate_overall_comment(result)
        result.strengths = self._extract_strengths(result)
        result.weaknesses = self._extract_weaknesses(result)
        result.improvements = self._generate_improvements(result)

        return result

    # -------------------- 一票否决项检查 --------------------
    def _check_veto_items(self) -> List[str]:
        triggered: List[str] = []
        veto_list = self.skill_config.get("veto_items", [])
        data = self.applicant_data
        text = self.application_text

        # 字数检查（适用于 party_application / thought_report 等有字数硬性要求的）
        for item in veto_list:
            # 字数下限
            m = re.search(r"字数\s*<\s*(\d+)", item)
            if m:
                min_words = int(m.group(1))
                if len(text) < min_words:
                    triggered.append(f"字数 {len(text)} < {min_words}（一票否决）")
                continue
            # 字数超限
            m = re.search(r"超字数", item)
            if m and text:
                # 各 skill 字数上限不同，简化处理：检查是否给出 word_count_max
                wmax = self.skill_config.get("hard_gates", {}).get("word_count_max")
                if wmax and len(text) > wmax:
                    triggered.append(f"字数 {len(text)} > 上限 {wmax}（一票否决）")
                continue
            # 必引理论（政治类）
            if "必引" in item or "指导思想" in item:
                required_terms = self._extract_required_terms(item)
                missing = [t for t in required_terms if t and t not in text]
                if missing:
                    triggered.append(f"必引要点缺失：{', '.join(missing)}（一票否决）")
                continue
            # 查重率
            if "查重率" in item:
                pl_rate = self.applicant_data.get("plagiarism_rate")
                if pl_rate and pl_rate > 30:
                    triggered.append(f"查重率 {pl_rate}% > 30%（一票否决）")
                continue
            # 其他 veto 项标记为"需人工核查"
            # （不在此自动判定，仅作记录）
        return triggered

    @staticmethod
    def _extract_required_terms(item: str) -> List[str]:
        """从 veto 项描述中抽取必引术语（简化）。"""
        # 党/习近平新时代/二十大/两个确立/两个维护/全心全意为人民服务
        terms = []
        if "党" in item:
            terms.append("中国共产党")
        if "习近平" in item or "新时代" in item:
            terms.append("习近平新时代中国特色社会主义思想")
        if "二十大" in item:
            terms.append("党的二十大")
        if "两个确立" in item:
            terms.append("两个确立")
        if "两个维护" in item:
            terms.append("两个维护")
        if "为人民服务" in item:
            terms.append("全心全意为人民服务")
        return terms

    # -------------------- 维度评分 --------------------
    def _score_dimensions(self) -> List[DimensionScore]:
        """对每个维度评分。"""
        out: List[DimensionScore] = []
        for dim_cfg in self.skill_config.get("dimensions", []):
            name = dim_cfg["name"]
            weight = dim_cfg["weight"]
            max_s = dim_cfg.get("max", weight)
            raw, comment, evidence = self._score_single_dimension(name, dim_cfg)
            weighted = raw * weight / 100.0
            grade = self._grade_from_score(raw)
            out.append(DimensionScore(
                name=name,
                weight=weight,
                max_score=max_s,
                raw_score=round(raw, 2),
                weighted_score=round(weighted, 2),
                comment=comment,
                evidence=evidence,
                grade=grade,
            ))
        return out

    def _score_single_dimension(
        self, dim_name: str, dim_cfg: Dict
    ) -> Tuple[float, str, List[str]]:
        """对单维度评分，返回 (0-100 原始分, 评语, 证据列表)。"""
        data = self.applicant_data
        evidence: List[str] = []
        score = 60.0  # 基础分

        # —— 学业表现 ——
        if "学业" in dim_name or "学习" in dim_name:
            rank = data.get("gpa_rank_percent")
            if rank is not None:
                if rank <= 3:
                    score = 98
                    evidence.append(f"GPA 排名 {rank}%（前 3%）")
                elif rank <= 5:
                    score = 95
                    evidence.append(f"GPA 排名 {rank}%（前 5%）")
                elif rank <= 10:
                    score = 90
                    evidence.append(f"GPA 排名 {rank}%（前 10%）")
                elif rank <= 20:
                    score = 82
                    evidence.append(f"GPA 排名 {rank}%（前 20%）")
                elif rank <= 30:
                    score = 75
                    evidence.append(f"GPA 排名 {rank}%（前 30%）")
                else:
                    score = 65
                    evidence.append(f"GPA 排名 {rank}%（>30%，较弱）")
            else:
                score = 70
                evidence.append("未提供 GPA 排名，按基础分计")
            # 英语等级加分
            eng = data.get("english_level", "")
            if eng in ("CET-6", "CET6", "六级"):
                score = min(score + 3, 100)
                evidence.append("通过英语六级")
            elif eng in ("CET-4", "CET4", "四级"):
                evidence.append("通过英语四级")
            elif eng in ("IELTS-7", "TOEFL-100"):
                score = min(score + 5, 100)
                evidence.append(f"高水平英语：{eng}")

        # —— 创新能力 / 科研 ——
        elif "创新" in dim_name or "科研" in dim_name or "学术" in dim_name:
            papers = data.get("papers", []) or []
            competitions = data.get("competitions", []) or []
            patents = data.get("patents", []) or []
            projects = data.get("innovation_projects", []) or []
            score = 60
            for p in papers:
                lvl = p.get("level", "")
                order = p.get("author_order", 99)
                if lvl in ("SCI", "EI", "CSSCI") and order == 1:
                    score = min(score + 20, 100)
                    evidence.append(f"第一作者 {lvl} 论文 1 篇")
                elif lvl in ("核心", "中文核心") and order == 1:
                    score = min(score + 12, 100)
                    evidence.append(f"第一作者核心期刊论文 1 篇")
                elif lvl in ("SCI", "EI", "CSSCI"):
                    score = min(score + 8, 100)
                    evidence.append(f"合作 {lvl} 论文 1 篇")
                else:
                    score = min(score + 3, 100)
                    evidence.append("普通论文 1 篇")
            for c in competitions:
                lvl = c.get("level", "")
                award = c.get("award", "")
                if lvl in ("国际", "国家级") and "一等" in award:
                    score = min(score + 15, 100)
                    evidence.append(f"{lvl} 一等奖")
                elif lvl in ("国际", "国家级"):
                    score = min(score + 10, 100)
                    evidence.append(f"{lvl} {award}")
                elif lvl == "省级" and "一等" in award:
                    score = min(score + 8, 100)
                    evidence.append("省级一等奖")
                elif lvl == "省级":
                    score = min(score + 5, 100)
                    evidence.append(f"省级 {award}")
            for pt in patents:
                if pt.get("inventor_order") == 1 and pt.get("type") == "发明":
                    score = min(score + 12, 100)
                    evidence.append("第一发明人发明专利 1 项")
                else:
                    score = min(score + 4, 100)
                    evidence.append("专利 1 项")
            for proj in projects:
                lvl = proj.get("level", "")
                if lvl == "国家级":
                    score = min(score + 8, 100)
                    evidence.append("国家级大创/科研项目")
                elif lvl == "省级":
                    score = min(score + 5, 100)
                    evidence.append("省级项目")
                else:
                    score = min(score + 2, 100)
                    evidence.append("校级项目")
            if not evidence:
                evidence.append("未提供科研/创新成果，按基础分计")

        # —— 实践 / 社会工作 ——
        elif "实践" in dim_name or "社会工作" in dim_name or "工作实绩" in dim_name:
            hours = data.get("volunteer_hours", 0) or 0
            cadre = data.get("cadre_positions", []) or []
            social = data.get("social_practice", []) or []
            score = 60
            if hours >= 100:
                score = min(score + 15, 100)
                evidence.append(f"志愿服务 {hours} 小时")
            elif hours >= 50:
                score = min(score + 10, 100)
                evidence.append(f"志愿服务 {hours} 小时")
            elif hours >= 20:
                score = min(score + 5, 100)
                evidence.append(f"志愿服务 {hours} 小时")
            for c in cadre:
                pos = c.get("position", "")
                tenure = c.get("tenure_months", 0) or 0
                lvl = c.get("level", "院级")
                if lvl == "校级" and tenure >= 12:
                    score = min(score + 12, 100)
                    evidence.append(f"校级 {pos} 任职 {tenure} 月")
                elif lvl == "院级" and tenure >= 12:
                    score = min(score + 8, 100)
                    evidence.append(f"院级 {pos} 任职 {tenure} 月")
                elif tenure >= 6:
                    score = min(score + 4, 100)
                    evidence.append(f"{pos} 任职 {tenure} 月")
            for s in social:
                award = s.get("award", "")
                if "省级" in award or "国家级" in award:
                    score = min(score + 8, 100)
                    evidence.append(f"社会实践 {award}")
                else:
                    score = min(score + 3, 100)
                    evidence.append("参与社会实践")
            if not evidence:
                evidence.append("未提供实践经历，按基础分计")

        # —— 思想品德 / 政治理论 ——
        elif "思想" in dim_name or "政治" in dim_name or "品德" in dim_name:
            score = 75
            if data.get("is_party_member"):
                score = min(score + 12, 100)
                evidence.append("中共党员")
            elif data.get("is_probationary_member"):
                score = min(score + 8, 100)
                evidence.append("预备党员")
            elif data.get("is_activist"):
                score = min(score + 5, 100)
                evidence.append("入党积极分子")
            if data.get("is_league_member"):
                evidence.append("共青团员")
            # 文本中政治理论引用
            text = self.application_text
            if "党的二十大" in text:
                evidence.append("引用党的二十大")
            if "习近平新时代中国特色社会主义思想" in text:
                evidence.append("引用习近平新时代中国特色社会主义思想")
            if not evidence:
                evidence.append("未提供思想政治表现佐证，按基础分计")

        # —— 材料规范 / 规范性 ——
        elif "规范" in dim_name or "材料" in dim_name:
            score = 80
            text = self.application_text
            wmin = self.skill_config.get("hard_gates", {}).get("word_count_min")
            wmax = self.skill_config.get("hard_gates", {}).get("word_count_max")
            if wmin and len(text) < wmin:
                score -= 15
                evidence.append(f"字数 {len(text)} < 下限 {wmin}")
            elif wmax and len(text) > wmax:
                score -= 10
                evidence.append(f"字数 {len(text)} > 上限 {wmax}")
            else:
                evidence.append(f"字数 {len(text)} 符合规范")
            # 检查必填字段
            required_fields = data.get("_required_fields_missing", [])
            if required_fields:
                score -= 5 * len(required_fields)
                evidence.append(f"缺失字段: {', '.join(required_fields)}")
            # 错别字粗检（仅基础启发式）
            typos = self._count_likely_typos(text)
            if typos > 0:
                score -= min(typos * 2, 10)
                evidence.append(f"疑似错别字 {typos} 处")
            score = max(score, 50)
            if not evidence:
                evidence.append("材料规范")

        # —— 团队 ——
        elif "团队" in dim_name:
            team = data.get("team_members", []) or []
            score = 60
            if len(team) >= 3:
                score += 10
                evidence.append(f"团队 {len(team)} 人，规模合理")
            disciplines = set(m.get("discipline", "") for m in team if m.get("discipline"))
            if len(disciplines) >= 2:
                score += 8
                evidence.append(f"跨 {len(disciplines)} 个专业")
            if data.get("advisor"):
                score += 7
                evidence.append("有指导教师")
            score = min(score, 100)
            if not evidence:
                evidence.append("未提供团队信息")

        # —— 其他维度（家庭经济/行业兴趣/职业规划/身体素质等）——
        else:
            score = 75
            evidence.append(f"{dim_name} 按基础分评估，需人工复核")
            # 家庭经济困难
            if "经济" in dim_name:
                if data.get("poverty_registered"):
                    score = 90
                    evidence.append("已通过家庭经济困难认定")
                if data.get("poverty_level") == "特困":
                    score = min(score + 5, 100)
                    evidence.append("特困等级")
            # 行业兴趣
            if "行业" in dim_name or "职业" in dim_name:
                if data.get("industry_match"):
                    score = 88
                    evidence.append("职业规划与目标行业匹配")
                if data.get("internship_relevant"):
                    score = min(score + 5, 100)
                    evidence.append("有相关实习经历")

        # 兜底
        score = max(0, min(100, score))
        comment = self._build_dimension_comment(dim_name, score, evidence)
        return score, comment, evidence

    @staticmethod
    def _count_likely_typos(text: str) -> int:
        """粗略统计疑似错别字（仅作启发式，统计非常用叠字/重复字）。"""
        if not text:
            return 0
        # 简化：检测连续 3 个以上相同字符
        return len(re.findall(r"(.)\1{2,}", text))

    def _build_dimension_comment(
        self, dim_name: str, score: float, evidence: List[str]
    ) -> str:
        grade = self._grade_from_score(score)
        label = self.GRADE_LABELS.get(grade, "")
        ev_text = "；".join(evidence[:3]) if evidence else "无具体佐证"
        return f"{dim_name}维度评分 {score:.1f}/100（{label}）。主要依据：{ev_text}。"

    @staticmethod
    def _grade_from_score(score: float) -> str:
        if score >= 90:
            return "A"
        if score >= 80:
            return "B"
        if score >= 70:
            return "C"
        return "D"

    # -------------------- 评语生成 --------------------
    def _generate_overall_comment(self, result: ReviewResult) -> str:
        parts: List[str] = []
        parts.append(
            f"申请人 {result.applicant_name} 申报 {result.skill_label}，"
            f"模拟评审总分 {result.total_score}/{result.max_score}，"
            f"等级 {result.grade}（{self.GRADE_LABELS.get(result.grade, '')}）。"
        )
        if result.veto_triggered:
            parts.append(
                f"触发一票否决项 {len(result.veto_triggered)} 项："
                + "；".join(result.veto_triggered)
            )
        if result.plagiarism:
            p = result.plagiarism
            if "overall_similarity" in p:
                parts.append(
                    f"查重预检：整体相似度 {p['overall_similarity']*100:.1f}%，"
                    f"等级 {p.get('grade', '-')}。"
                )
        # 最高/最低维度
        if result.dimensions:
            best = max(result.dimensions, key=lambda d: d.raw_score)
            worst = min(result.dimensions, key=lambda d: d.raw_score)
            parts.append(
                f"最强维度：{best.name}（{best.raw_score:.1f}）；"
                f"最弱维度：{worst.name}（{worst.raw_score:.1f}）。"
            )
        if result.passed:
            parts.append("综合建议：通过预审，可正式提交。")
        else:
            parts.append("综合建议：未通过预审，需根据改进建议修改后重审。")
        return " ".join(parts)

    def _extract_strengths(self, result: ReviewResult) -> List[str]:
        out: List[str] = []
        for d in result.dimensions:
            if d.raw_score >= 85:
                out.append(f"【{d.name}】{d.comment}")
        if result.plagiarism and result.plagiarism.get("grade") == "A":
            out.append("【原创性】查重等级 A，原创性良好。")
        return out

    def _extract_weaknesses(self, result: ReviewResult) -> List[str]:
        out: List[str] = []
        for d in result.dimensions:
            if d.raw_score < 75:
                out.append(f"【{d.name}】{d.comment}")
        if result.veto_triggered:
            out.append(f"【一票否决】共 {len(result.veto_triggered)} 项")
        if result.plagiarism and not result.plagiarism.get("passed", True):
            out.append("【原创性】查重未通过。")
        return out

    def _generate_improvements(self, result: ReviewResult) -> List[str]:
        out: List[str] = []
        for d in result.dimensions:
            if d.raw_score < 75:
                if "学业" in d.name:
                    out.append(f"提升 {d.name}：补强 GPA/排名，必要时申请破格通道。")
                elif "创新" in d.name or "科研" in d.name:
                    out.append(f"提升 {d.name}：补充论文/竞赛/专利等具体成果佐证。")
                elif "实践" in d.name:
                    out.append(f"提升 {d.name}：增加志愿服务时长、补全学生干部任职证明。")
                elif "思想" in d.name or "政治" in d.name:
                    out.append(f"提升 {d.name}：补充入党/积极分子证明，引用最新时政。")
                elif "规范" in d.name or "材料" in d.name:
                    out.append(f"提升 {d.name}：核对字数与必填字段，纠正错别字。")
                else:
                    out.append(f"提升 {d.name}：补充具体佐证材料。")
        if result.veto_triggered:
            out.append("【紧急】必须解决一票否决项后才能通过预审。")
        if result.plagiarism and not result.plagiarism.get("passed", True):
            out.append("【紧急】查重未通过，需重写命中段落（参考 plagiarism_checker 报告）。")
        return out

    # -------------------- 报告生成 --------------------
    def generate_review_report(self, out_path: str) -> str:
        """
        生成评审报告 docx（如 python-docx 不可用则降级为 txt）。

        :param out_path: 输出路径（.docx 或 .txt）
        :return: 实际输出路径
        """
        if not hasattr(self, "_last_result"):
            self._last_result = self.simulate_review()
        result = self._last_result

        if out_path.endswith(".docx"):
            try:
                return self._write_docx(result, out_path)
            except ImportError:
                # 降级
                txt_path = out_path.replace(".docx", ".txt")
                return self._write_txt(result, txt_path)
        return self._write_txt(result, out_path)

    def _write_docx(self, result: ReviewResult, out_path: str) -> str:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        doc = Document()
        # 页面设置
        for section in doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # 标题
        h = doc.add_heading(f"{result.skill_label} 评审模拟报告", level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本信息
        p = doc.add_paragraph()
        p.add_run(f"申请人：{result.applicant_name}\n").bold = True
        p.add_run(f"申报项目：{result.skill_label}（{result.skill_name}）\n")
        p.add_run(f"模拟评审总分：{result.total_score}/{result.max_score}\n")
        p.add_run(f"评审等级：{result.grade} - {self.GRADE_LABELS.get(result.grade, '')}\n")
        p.add_run(f"是否通过预审：{'✅ 通过' if result.passed else '❌ 不通过'}\n")
        if result.plagiarism and "overall_similarity" in result.plagiarism:
            p.add_run(
                f"查重相似度：{result.plagiarism['overall_similarity']*100:.1f}%"
                f"（等级 {result.plagiarism.get('grade', '-')}）\n"
            )

        # 一票否决
        if result.veto_triggered:
            doc.add_heading("⚠️ 一票否决项", level=1)
            for v in result.veto_triggered:
                doc.add_paragraph(v, style="List Bullet")

        # 维度评分表
        doc.add_heading("📊 多维度评分", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "维度"
        hdr[1].text = "权重"
        hdr[2].text = "原始分"
        hdr[3].text = "加权分"
        hdr[4].text = "等级"
        for d in result.dimensions:
            row = table.add_row().cells
            row[0].text = d.name
            row[1].text = f"{d.weight}%"
            row[2].text = f"{d.raw_score:.1f}"
            row[3].text = f"{d.weighted_score:.2f}"
            row[4].text = d.grade
        # 评语
        doc.add_heading("📝 维度评语", level=1)
        for d in result.dimensions:
            doc.add_paragraph(f"【{d.name}】{d.comment}")

        # 优势与不足
        if result.strengths:
            doc.add_heading("✅ 优势", level=1)
            for s in result.strengths:
                doc.add_paragraph(s, style="List Bullet")
        if result.weaknesses:
            doc.add_heading("❌ 不足", level=1)
            for w in result.weaknesses:
                doc.add_paragraph(w, style="List Bullet")
        if result.improvements:
            doc.add_heading("🔧 改进建议", level=1)
            for i in result.improvements:
                doc.add_paragraph(i, style="List Bullet")

        # 总评
        doc.add_heading("📋 综合评语", level=1)
        doc.add_paragraph(result.overall_comment)

        # 落款
        doc.add_paragraph("")
        sig = doc.add_paragraph("评审模拟器 v1.0 自动生成")
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.save(out_path)
        return out_path

    def _write_txt(self, result: ReviewResult, out_path: str) -> str:
        lines: List[str] = []
        lines.append(f"{'='*60}")
        lines.append(f"  {result.skill_label} 评审模拟报告")
        lines.append(f"{'='*60}")
        lines.append(f"申请人        : {result.applicant_name}")
        lines.append(f"申报项目      : {result.skill_label}（{result.skill_name}）")
        lines.append(f"模拟评审总分  : {result.total_score}/{result.max_score}")
        lines.append(f"评审等级      : {result.grade} - {self.GRADE_LABELS.get(result.grade, '')}")
        lines.append(f"是否通过预审  : {'✅ 通过' if result.passed else '❌ 不通过'}")
        if result.plagiarism and "overall_similarity" in result.plagiarism:
            lines.append(
                f"查重相似度    : {result.plagiarism['overall_similarity']*100:.1f}%"
                f"（等级 {result.plagiarism.get('grade', '-')}）"
            )
        lines.append("")
        if result.veto_triggered:
            lines.append("⚠️ 一票否决项：")
            for v in result.veto_triggered:
                lines.append(f"  - {v}")
            lines.append("")
        lines.append("📊 多维度评分：")
        lines.append(f"  {'维度':<12}{'权重':>6}{'原始分':>8}{'加权分':>8}{'等级':>6}")
        for d in result.dimensions:
            lines.append(
                f"  {d.name:<12}{d.weight:>5}%{d.raw_score:>8.1f}{d.weighted_score:>8.2f}{d.grade:>6}"
            )
        lines.append("")
        lines.append("📝 维度评语：")
        for d in result.dimensions:
            lines.append(f"  【{d.name}】{d.comment}")
        lines.append("")
        if result.strengths:
            lines.append("✅ 优势：")
            for s in result.strengths:
                lines.append(f"  - {s}")
        if result.weaknesses:
            lines.append("❌ 不足：")
            for w in result.weaknesses:
                lines.append(f"  - {w}")
        if result.improvements:
            lines.append("🔧 改进建议：")
            for i in result.improvements:
                lines.append(f"  - {i}")
        lines.append("")
        lines.append("📋 综合评语：")
        lines.append(result.overall_comment)
        lines.append("")
        lines.append("评审模拟器 v1.0 自动生成")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return out_path


# ---------------------------------------------------------------------------
# 多人对比
# ---------------------------------------------------------------------------
def compare_applicants(applicant_list: List[Dict]) -> Dict:
    """
    多人对比排序。

    :param applicant_list: [{"skill_name": "...", "application_text": "...",
                             "applicant_data": {...}}, ...]
    :return: {"ranking": [...], "summary": "..."}
    """
    results: List[Tuple[int, ReviewResult]] = []
    for idx, app in enumerate(applicant_list):
        try:
            sim = ReviewSimulator(
                skill_name=app["skill_name"],
                application_text=app.get("application_text", ""),
                applicant_data=app.get("applicant_data", {}),
            )
            res = sim.simulate_review()
            results.append((idx, res))
        except Exception as e:
            results.append((idx, ReviewResult(
                skill_name=app.get("skill_name", "?"),
                skill_label=app.get("skill_name", "?"),
                applicant_name=app.get("applicant_data", {}).get("name", f"#{idx}"),
                total_score=0,
                overall_comment=f"评审失败：{e}",
            )))

    # 排序：一票否决的排最后；其余按 total_score 降序
    def sort_key(t: Tuple[int, ReviewResult]) -> Tuple[int, float]:
        _, r = t
        return (0 if r.passed else 1, -r.total_score)

    results.sort(key=sort_key)

    ranking = []
    for rank, (idx, r) in enumerate(results, 1):
        ranking.append({
            "rank": rank,
            "original_index": idx,
            "applicant_name": r.applicant_name,
            "skill_name": r.skill_name,
            "total_score": r.total_score,
            "grade": r.grade,
            "passed": r.passed,
            "veto_count": len(r.veto_triggered),
            "comment": r.overall_comment,
        })

    passed_n = sum(1 for r in ranking if r["passed"])
    summary = (
        f"共 {len(ranking)} 名申请人，通过预审 {passed_n} 名，"
        f"未通过 {len(ranking) - passed_n} 名。"
        f"最高分 {max((r['total_score'] for r in ranking), default=0):.2f}，"
        f"最低分 {min((r['total_score'] for r in ranking), default=0):.2f}。"
    )
    return {"ranking": ranking, "summary": summary, "results": [r.to_dict() for _, r in results]}


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def _build_arg_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        description="大学生申报书评审模拟器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""\
示例：
  python review_simulator.py --skill national_scholarship --data data.json
  python review_simulator.py --skill national_scholarship --data data.json --text-file text.txt --out report.docx
  python review_simulator.py --compare applicants.json --out ranking.json
  python review_simulator.py --list-skills
""",
    )
    p.add_argument("--skill", help="子 skill 名称")
    p.add_argument("--data", help="申请人数据 JSON 文件路径")
    p.add_argument("--text", help="申报书正文文本")
    p.add_argument("--text-file", help="申报书正文文件路径")
    p.add_argument("--out", help="输出报告路径（.docx 或 .json）")
    p.add_argument("--json", action="store_true", help="输出 JSON 格式评审结果")
    p.add_argument("--compare", help="多人对比模式，输入 JSON 文件（applicant_list）")
    p.add_argument("--list-skills", action="store_true", help="列出所有支持的 skill")
    return p


def _list_skills() -> None:
    if not os.path.exists(CRITERIA_PATH):
        print(f"❌ 找不到评审标准: {CRITERIA_PATH}", file=sys.stderr)
        return
    with open(CRITERIA_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)
    skills = [k for k in data.keys() if k != "_meta"]
    print(f"支持的子 skill（共 {len(skills)} 个）：")
    for k in sorted(skills):
        cfg = data[k]
        dims = " / ".join(d["name"] for d in cfg.get("dimensions", []))
        print(f"  - {k:<28} {cfg.get('name', '')}  [{dims}]")


def main(argv: Optional[List[str]] = None) -> int:
    parser = _build_arg_parser()
    args = parser.parse_args(argv)

    if args.list_skills:
        _list_skills()
        return 0

    # 多人对比模式
    if args.compare:
        if not os.path.exists(args.compare):
            print(f"❌ 文件不存在: {args.compare}", file=sys.stderr)
            return 2
        with open(args.compare, "r", encoding="utf-8") as f:
            applicant_list = json.load(f)
        if not isinstance(applicant_list, list):
            print("❌ --compare 文件需为 JSON 数组", file=sys.stderr)
            return 2
        result = compare_applicants(applicant_list)
        out = json.dumps(result, ensure_ascii=False, indent=2)
        if args.out:
            with open(args.out, "w", encoding="utf-8") as f:
                f.write(out)
            print(f"✅ 对比结果已写入: {args.out}")
        else:
            print(out)
        return 0

    # 单人评审
    if not args.skill:
        parser.print_help()
        return 2

    applicant_data: Dict[str, Any] = {}
    if args.data:
        if not os.path.exists(args.data):
            print(f"❌ 文件不存在: {args.data}", file=sys.stderr)
            return 2
        with open(args.data, "r", encoding="utf-8") as f:
            applicant_data = json.load(f)

    text = args.text or ""
    if args.text_file:
        if not os.path.exists(args.text_file):
            print(f"❌ 文件不存在: {args.text_file}", file=sys.stderr)
            return 2
        with open(args.text_file, "r", encoding="utf-8") as f:
            text = f.read()

    try:
        sim = ReviewSimulator(
            skill_name=args.skill,
            application_text=text,
            applicant_data=applicant_data,
        )
    except (ValueError, FileNotFoundError) as e:
        print(f"❌ 初始化失败：{e}", file=sys.stderr)
        return 2

    result = sim.simulate_review()

    if args.out and (args.out.endswith(".docx") or args.out.endswith(".txt")):
        path = sim.generate_review_report(args.out)
        print(f"✅ 评审报告已写入: {path}")
        print(f"   总分 {result.total_score}/{result.max_score}，等级 {result.grade}，"
              f"{'通过' if result.passed else '未通过'}")
        return 0 if result.passed else 1

    if args.json or (args.out and args.out.endswith(".json")):
        out = result.to_json()
        if args.out:
            with open(args.out, "w", encoding="utf-8") as f:
                f.write(out)
            print(f"✅ 评审结果已写入: {args.out}")
        else:
            print(out)
    else:
        # 控制台简报
        print(f"申请人      : {result.applicant_name}")
        print(f"申报项目    : {result.skill_label}")
        print(f"总分        : {result.total_score}/{result.max_score}")
        print(f"等级        : {result.grade} - {ReviewSimulator.GRADE_LABELS.get(result.grade, '')}")
        print(f"是否通过    : {'✅ 通过' if result.passed else '❌ 未通过'}")
        if result.veto_triggered:
            print(f"一票否决    : {len(result.veto_triggered)} 项")
            for v in result.veto_triggered:
                print(f"  - {v}")
        print("\n维度评分：")
        for d in result.dimensions:
            print(f"  {d.name:<12}  {d.raw_score:>6.1f}/100  权重 {d.weight}%  加权 {d.weighted_score:>6.2f}  等级 {d.grade}")
        if result.plagiarism and "overall_similarity" in result.plagiarism:
            print(f"\n查重        : {result.plagiarism['overall_similarity']*100:.1f}%（等级 {result.plagiarism.get('grade', '-')}）")
        print(f"\n综合评语    : {result.overall_comment}")
    return 0 if result.passed else 1


if __name__ == "__main__":
    sys.exit(main())
