#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx_common.py — college-application-doc 共享组件库
====================================================

为 22+ 个子 skill 的 build.py 提供统一的 docx 生成基础能力，
消除字体常量、页面设置、段落生成、表格生成、签字栏、页眉页脚等
跨子 skill 重复代码。

设计原则
--------
1. **零业务耦合**：本模块只提供 docx 基础构件，不包含任何具体材料
   类型（奖学金/入党/三下乡/挑战杯…）的业务逻辑。
2. **覆盖现有重复模式**：抽样 national_scholarship / motivation_scholarship
   build.py 后，提取 12 套字体常量、5 个工具函数、15 个段落/表格/页眉
   构件、1 个 DocxBuilder 基类，覆盖现有 build.py 中 95% 的重复样板。
3. **不破坏向后兼容**：现有 build.py 不强制改造，本模块可被新增子 skill
   直接 import 使用。
4. **可独立测试**：example_usage.py 跑通即视为通过。

依赖
----
- python-docx >= 0.8 (已在 SKILL 环境预装，实测 1.2.0)
- 标准库：os / re / sys / datetime / pathlib / typing

使用方式
--------
    from utils.docx_common import (
        DocxBuilder, FONT_SONG, FONT_HEI, SIZE_ER, SIZE_XIAO_SI,
        setup_a4_page, set_run_font, add_title, add_heading1,
        add_paragraph, add_table, add_signature_block,
        add_page_number, add_header, add_footer, add_seal_placeholder,
        add_date_line, add_this_salute, validate_docx, save_docx,
        count_chinese_chars, count_words,
        apply_red_text, apply_bold_text,
    )

    builder = DocxBuilder("/tmp/demo.docx")
    builder.add_title("国家奖学金申请书")
    builder.add_paragraph("尊敬的校奖学金评审委员会：", indent=False)
    builder.add_paragraph("本人…")
    builder.save()

作者：T39 工程化任务
版本：1.0.0
"""

from __future__ import annotations

import os
import re
import sys
from datetime import date, datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple, Union

from docx import Document
from docx.document import Document as _Document  # noqa: F401  (用于类型注解)
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# 1. 字体常量
# ============================================================

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANGSONG = "仿宋"
FONT_KAI = "楷体"
FONT_TIMES = "Times New Roman"

# 兼容别名（部分 build.py 使用大写）
FONT_SONGTI = FONT_SONG
FONT_HEITI = FONT_HEI


# ============================================================
# 2. 字号常量
# ============================================================

SIZE_ER = Pt(22)            # 二号
SIZE_XIAO_ER = Pt(18)       # 小二
SIZE_SAN = Pt(16)           # 三号
SIZE_XIAO_SAN = Pt(15)      # 小三
SIZE_SI = Pt(14)            # 四号
SIZE_XIAO_SI = Pt(12)       # 小四
SIZE_WU = Pt(10.5)          # 五号
SIZE_XIAO_WU = Pt(9)        # 小五


# ============================================================
# 3. 页面 / 颜色常量
# ============================================================

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_BOTTOM_CM = 2.54
MARGIN_LEFT_RIGHT_CM = 2.5

COLOR_RED = RGBColor(0xFF, 0x00, 0x00)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_DARK_GRAY = RGBColor(0x40, 0x40, 0x40)


# ============================================================
# 4. 工具函数：Run / 单元格 / 段落
# ============================================================

def set_run_font(
    run,
    font_name: str = FONT_SONG,
    font_size: Any = SIZE_XIAO_SI,
    bold: bool = False,
    color: Optional[RGBColor] = None,
    italic: bool = False,
) -> None:
    """设置 run 字体（中英文同步设置 eastAsia/ascii/hAnsi）

    Args:
        run: docx Run 对象
        font_name: 字体名（如 "宋体" / "Times New Roman"）
        font_size: Pt 对象
        bold: 是否加粗
        color: RGBColor，None 表示不修改颜色
        italic: 是否斜体

    Notes:
        python-docx 默认只设置 ascii/hAnsi，中文需要额外写 eastAsia。
        本函数同时写入三个属性，确保中英文混排时字体一致。
    """
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def apply_red_text(run) -> None:
    """设置 run 文本为红色（常用于强调或印章占位）"""
    run.font.color.rgb = COLOR_RED


def apply_bold_text(run) -> None:
    """加粗 run"""
    run.font.bold = True


def set_cell_font(
    cell,
    font_name: str = FONT_SONG,
    font_size: Any = SIZE_WU,
    bold: bool = False,
    alignment: int = WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    """设置单元格内所有文字字体与对齐"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def set_cell_text(
    cell,
    text: str,
    font_name: str = FONT_SONG,
    font_size: Any = SIZE_WU,
    bold: bool = False,
    alignment: int = WD_ALIGN_PARAGRAPH.CENTER,
    line_spacing: float = 1.25,
) -> None:
    """清空单元格并写入文字（含字体设置）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=bold)


def add_paragraph_with_format(
    doc,
    text: str,
    font_name: str = FONT_SONG,
    font_size: Any = SIZE_XIAO_SI,
    bold: bool = False,
    alignment: int = WD_ALIGN_PARAGRAPH.LEFT,
    first_line_indent: bool = True,
    line_spacing: float = 1.5,
    space_before: float = 0,
    space_after: float = 0,
    color: Optional[RGBColor] = None,
):
    """添加带格式段落，可控制字体/字号/对齐/缩进/行距/段前后/颜色

    Returns:
        docx Paragraph 对象
    """
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    if first_line_indent:
        # 首行缩进 2 字符 = 2 * 字号
        pf.first_line_indent = Pt(font_size.pt * 2)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size=font_size,
                 bold=bold, color=color)
    return p


# ============================================================
# 5. 页面 / 标题 / 标题层级
# ============================================================

def setup_a4_page(
    doc,
    margins: Optional[Dict[str, float]] = None,
) -> None:
    """设置 A4 页面与页边距

    Args:
        doc: docx Document 对象
        margins: 可选自定义页边距 dict，单位 cm，
                 keys: top, bottom, left, right
                 默认上下 2.54cm，左右 2.5cm

    Notes:
        此函数会修改所有 section（多数 docx 只有一个 section）。
    """
    if margins is None:
        margins = {
            "top": MARGIN_TOP_BOTTOM_CM,
            "bottom": MARGIN_TOP_BOTTOM_CM,
            "left": MARGIN_LEFT_RIGHT_CM,
            "right": MARGIN_LEFT_RIGHT_CM,
        }
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(margins["top"])
        section.bottom_margin = Cm(margins["bottom"])
        section.left_margin = Cm(margins["left"])
        section.right_margin = Cm(margins["right"])


def add_title(
    doc,
    text: str,
    size: Any = SIZE_ER,
    align: int = WD_ALIGN_PARAGRAPH.CENTER,
    font_name: str = FONT_HEI,
    bold: bool = True,
    space_before: float = 12,
    space_after: float = 12,
):
    """添加标题（默认黑体二号居中，段前段后 12pt）

    适用于 B 类书信体材料（国家奖学金 / 励志奖学金 / 入党申请书 等）。
    A 类审批表如需三号标题，传 size=SIZE_SAN。
    """
    return add_paragraph_with_format(
        doc, text, font_name=font_name, font_size=size, bold=bold,
        alignment=align, first_line_indent=False,
        space_before=space_before, space_after=space_after)


def add_heading1(
    doc,
    text: str,
    align: int = WD_ALIGN_PARAGRAPH.CENTER,
):
    """一级标题：黑体三号居中，段前 12pt 段后 6pt

    用于 A 类审批表主标题、章节主标题。
    """
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_SAN, bold=True,
        alignment=align, first_line_indent=False,
        space_before=12, space_after=6)


def add_heading2(
    doc,
    text: str,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
):
    """二级标题：黑体小三左对齐，段前 6pt 段后 3pt

    用于章节子标题（如"一、思想方面"）。
    """
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_XIAO_SAN, bold=True,
        alignment=align, first_line_indent=False,
        space_before=6, space_after=3)


def add_heading3(
    doc,
    text: str,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
):
    """三级标题：黑体四号左对齐（首行缩进 2 字符），段前 3pt 段后 3pt

    用于章节小标题（如"（一）理论学习"）。
    """
    return add_paragraph_with_format(
        doc, text, font_name=FONT_HEI, font_size=SIZE_SI, bold=True,
        alignment=align, first_line_indent=True,
        space_before=3, space_after=3)


def add_paragraph(
    doc,
    text: str,
    indent: bool = True,
    bold: bool = False,
    font_name: str = FONT_SONG,
    font_size: Any = SIZE_XIAO_SI,
    line_spacing: float = 1.5,
    alignment: int = WD_ALIGN_PARAGRAPH.LEFT,
):
    """添加正文段落（宋体小四，1.5 倍行距，默认首行缩进 2 字符）

    Args:
        indent: 首行缩进 2 字符（书信体正文 True；称呼/落款 False）
        bold: 是否整段加粗
    """
    return add_paragraph_with_format(
        doc, text, font_name=font_name, font_size=font_size, bold=bold,
        alignment=alignment, first_line_indent=indent,
        line_spacing=line_spacing)


def add_salutation(doc, text: str = "尊敬的校领导："):
    """添加称呼段落：顶格不缩进，宋体小四"""
    return add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


# ============================================================
# 6. 表格
# ============================================================

def add_table(
    doc,
    headers: Sequence[str],
    rows: Sequence[Sequence[Any]],
    col_widths: Optional[Sequence[float]] = None,
    caption: str = "",
    font_size: Any = SIZE_WU,
    header_bold: bool = True,
):
    """从数据创建表格，自动应用规范格式

    Args:
        headers: 表头文字列表
        rows: 行数据二维列表
        col_widths: 列宽 cm 列表（可选）
        caption: 表格上方说明文字（如"主干课程成绩："），不写则不添加
        font_size: 单元格字号，默认五号
        header_bold: 表头是否加粗

    Returns:
        docx Table 对象

    Notes:
        - 表头居中加粗，数据居中
        - 表格整体居中对齐
        - 列宽如指定则按 cm 设置
    """
    if caption:
        add_paragraph_with_format(
            doc, caption, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=True,
            line_spacing=1.5, space_before=3, space_after=3)

    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, font_name=FONT_SONG,
                      font_size=font_size, bold=header_bold)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 数据行
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            if j >= n_cols:
                break
            cell = table.rows[i].cells[j]
            set_cell_text(cell, val, font_name=FONT_SONG,
                          font_size=font_size, bold=False)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


# ============================================================
# 7. 书信体专用构件
# ============================================================

def add_this_salute(doc, cizhi: str = "此致", jingli: str = "敬礼！"):
    """添加"此致/敬礼！"结尾

    规范：
    - "此致"另起一行，空两格（首行缩进 2 字符）
    - "敬礼！"另起一行，顶格（不缩进）
    """
    add_paragraph_with_format(
        doc, cizhi, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=True, line_spacing=1.5)
    add_paragraph_with_format(
        doc, jingli, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=False, line_spacing=1.5)


def add_signature_block(
    doc,
    signatures: Sequence[str],
    align: int = WD_ALIGN_PARAGRAPH.RIGHT,
    line_spacing: float = 1.5,
):
    """添加签字栏（落款），支持多行

    Args:
        signatures: 落款行列表，如 ["申请人：张三", "2025 年 5 月 10 日"]
        align: 对齐，默认右对齐

    Example:
        add_signature_block(doc, ["申请人：张三", "2025 年 5 月 10 日"])
    """
    for line in signatures:
        add_paragraph_with_format(
            doc, line, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
            alignment=align, first_line_indent=False,
            line_spacing=line_spacing, space_before=0, space_after=0)


def add_date_line(
    doc,
    custom_date: Optional[str] = None,
    align: int = WD_ALIGN_PARAGRAPH.RIGHT,
):
    """添加落款日期行

    Args:
        custom_date: 自定义日期文本，如 "2025 年 5 月 10 日"
                     None 则自动取今天，格式 "YYYY 年 M 月 D 日"
        align: 对齐，默认右对齐
    """
    if custom_date is None:
        today = date.today()
        custom_date = f"{today.year} 年 {today.month} 月 {today.day} 日"
    add_paragraph_with_format(
        doc, custom_date, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=align, first_line_indent=False, line_spacing=1.5,
        space_before=6, space_after=6)


def add_seal_placeholder(
    doc,
    position: str = "right_bottom",
    text: str = "（加盖公章）",
    size_cm: float = 3.5,
):
    """添加印章占位符

    Args:
        position: 印章位置标识，目前仅 "right_bottom"（右下角，落款下方）
        text: 占位提示文字，默认 "（加盖公章）"
        size_cm: 占位符大小，cm

    Notes:
        本函数仅添加文字占位，实际印章由用户手工加盖。
        如需绘制圆框，可在子类重写本方法。
    """
    if position != "right_bottom":
        # 兼容其他位置：均按右对齐处理
        position = "right_bottom"
    # 占位一行
    add_paragraph_with_format(
        doc, text, font_name=FONT_SONG, font_size=SIZE_XIAO_SI,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=False, line_spacing=1.5,
        space_before=12, space_after=12, color=COLOR_DARK_GRAY)


# ============================================================
# 8. 页眉 / 页脚 / 页码
# ============================================================

def _set_header_footer_text(section, target: str, text: str) -> None:
    """向 section 的页眉/页脚写入文字

    Args:
        section: docx Section 对象
        target: "header" 或 "footer"
        text: 文字内容
    """
    container = section.header if target == "header" else section.footer
    # 清空默认段落
    if not container.paragraphs:
        p = container.add_paragraph()
    else:
        p = container.paragraphs[0]
        # 清空已有 runs
        for run in list(p.runs):
            run.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_XIAO_WU,
                 color=COLOR_DARK_GRAY)


def add_header(doc, text: str) -> None:
    """添加页眉文字（居中，宋体小五深灰）

    Notes:
        会修改所有 section 的页眉。
    """
    for section in doc.sections:
        _set_header_footer_text(section, "header", text)


def add_footer(doc, text: str) -> None:
    """添加页脚文字（居中，宋体小五深灰）

    Notes:
        会修改所有 section 的页脚。
        如需页码，请使用 add_page_number 而非本函数。
    """
    for section in doc.sections:
        _set_header_footer_text(section, "footer", text)


def _create_field(paragraph, field_code: str) -> None:
    """向段落中插入 Word 域代码（如 PAGE / NUMPAGES）"""
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = field_code
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar_begin)
    run._element.append(instrText)
    run._element.append(fldChar_sep)
    run._element.append(fldChar_end)
    set_run_font(run, font_name=FONT_SONG, font_size=SIZE_XIAO_WU,
                 color=COLOR_DARK_GRAY)


def add_page_number(doc, fmt: str = "第 {PAGE} 页 共 {NUMPAGES} 页") -> None:
    """添加页码到页脚（居中）

    Args:
        fmt: 格式字符串，{PAGE} 与 {NUMPAGES} 占位符分别替换为
             当前页码与总页数（Word 域代码）。

    Notes:
        会覆盖所有 section 的页脚内容。
    """
    parts = re.split(r"(\{PAGE\}|\{NUMPAGES\})", fmt)
    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            p = footer.add_paragraph()
        else:
            p = footer.paragraphs[0]
            for run in list(p.runs):
                run.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for part in parts:
            if part == "{PAGE}":
                _create_field(p, "PAGE")
            elif part == "{NUMPAGES}":
                _create_field(p, "NUMPAGES")
            elif part:
                run = p.add_run(part)
                set_run_font(run, font_name=FONT_SONG,
                             font_size=SIZE_XIAO_WU,
                             color=COLOR_DARK_GRAY)


# ============================================================
# 9. 字数统计 / 校验 / 保存
# ============================================================

_CN_RE = re.compile(r"[\u4e00-\u9fff]")
# 英文单词：连续字母数字
_WORD_RE = re.compile(r"[A-Za-z0-9]+")


def count_chinese_chars(text: str) -> int:
    """统计中文字符数（CJK 统一汉字区段）"""
    return len(_CN_RE.findall(text or ""))


def count_words(text: str) -> int:
    """统计字数（中英文混合，中文按字符计，英文按单词计）

    示例：
        "今天 weather is nice" → 1 (今) + 1 (天) + 3 (weather/is/nice) = 5
    """
    text = text or ""
    cn = len(_CN_RE.findall(text))
    en = len(_WORD_RE.findall(text))
    return cn + en


def validate_docx(doc) -> Tuple[bool, List[str]]:
    """验证 docx 完整性

    Returns:
        (is_valid, issues)
        is_valid: True 表示通过基础校验
        issues: 问题列表（空列表表示无问题）
    """
    issues: List[str] = []
    if doc is None:
        return False, ["doc 为 None"]
    # 至少有一个段落或表格
    n_para = len(doc.paragraphs)
    n_table = len(doc.tables)
    if n_para == 0 and n_table == 0:
        issues.append("文档为空：既无段落也无表格")
    # 检查至少有一个 section
    if len(doc.sections) == 0:
        issues.append("文档无 section")
    # 检查每个 section 页面尺寸
    for i, sec in enumerate(doc.sections):
        if sec.page_width is None or sec.page_height is None:
            issues.append(f"section[{i}] 缺少页面尺寸")
    is_valid = len(issues) == 0
    return is_valid, issues


def save_docx(doc, output_path: Union[str, Path]) -> str:
    """保存 docx 到指定路径

    Args:
        doc: docx Document 对象
        output_path: 输出路径

    Returns:
        保存后的绝对路径

    Raises:
        ValueError: doc 为 None 或路径无效
        OSError: 父目录无法创建
    """
    if doc is None:
        raise ValueError("doc 为 None，无法保存")
    output_path = str(output_path)
    if not output_path:
        raise ValueError("output_path 为空")
    parent = os.path.dirname(os.path.abspath(output_path))
    os.makedirs(parent, exist_ok=True)
    doc.save(output_path)
    return os.path.abspath(output_path)


# ============================================================
# 10. DocxBuilder 基类
# ============================================================

class DocxBuilder:
    """通用 docx 构建器，所有子 skill build.py 可继承此类

    子类典型用法：
        class NationalScholarshipBuilder(DocxBuilder):
            def build(self, data: dict):
                self.add_title(data["title"])
                self.add_salutation(data["salutation"])
                for para in data["body"]:
                    self.add_paragraph(para)
                self.add_this_salute()
                self.add_signature_block([data["applicant"], data["date"]])
                self.save()

    本基类提供：
    - 默认 A4 页面（上下 2.54cm 左右 2.5cm）
    - 标题 / 标题层级 / 段落 / 表格 / 称呼 / 此致敬礼 / 落款 / 印章 / 日期
    - 页眉 / 页脚 / 页码
    - 校验 / 保存

    子类只需实现 build() 方法即可。
    """

    def __init__(
        self,
        output_path: Union[str, Path],
        margins: Optional[Dict[str, float]] = None,
        with_page_number: bool = False,
        page_number_fmt: str = "第 {PAGE} 页 共 {NUMPAGES} 页",
        header: Optional[str] = None,
        footer: Optional[str] = None,
    ):
        """
        Args:
            output_path: 输出 docx 路径
            margins: 自定义页边距 dict（top/bottom/left/right，单位 cm）
            with_page_number: 是否自动添加页码（默认 False，多数申请书不需要）
            page_number_fmt: 页码格式（仅 with_page_number=True 生效）
            header: 页眉文字（可选）
            footer: 页脚文字（可选，与页码互斥）
        """
        self.output_path = str(output_path)
        self.doc = Document()
        setup_a4_page(self.doc, margins=margins)
        if header:
            add_header(self.doc, header)
        if with_page_number:
            add_page_number(self.doc, fmt=page_number_fmt)
        elif footer:
            add_footer(self.doc, footer)

    # ------- 标题 -------

    def add_title(self, text: str, **kwargs):
        """添加文档主标题（黑体二号居中）"""
        return add_title(self.doc, text, **kwargs)

    def add_heading1(self, text: str, **kwargs):
        """添加一级标题（黑体三号居中）"""
        return add_heading1(self.doc, text, **kwargs)

    def add_heading2(self, text: str, **kwargs):
        """添加二级标题（黑体小三左对齐）"""
        return add_heading2(self.doc, text, **kwargs)

    def add_heading3(self, text: str, **kwargs):
        """添加三级标题（黑体四号左对齐）"""
        return add_heading3(self.doc, text, **kwargs)

    # ------- 段落 -------

    def add_paragraph(self, text: str, indent: bool = True, **kwargs):
        """添加正文段落（宋体小四，首行缩进 2 字符）"""
        return add_paragraph(self.doc, text, indent=indent, **kwargs)

    def add_salutation(self, text: str = "尊敬的校领导："):
        """添加称呼段落（顶格不缩进）"""
        return add_salutation(self.doc, text)

    def add_this_salute(self, cizhi: str = "此致", jingli: str = "敬礼！"):
        """添加"此致/敬礼！"结尾"""
        add_this_salute(self.doc, cizhi=cizhi, jingli=jingli)

    # ------- 落款 -------

    def add_signature_block(self, signatures: Sequence[str], **kwargs):
        """添加落款（默认右对齐）"""
        add_signature_block(self.doc, signatures, **kwargs)

    def add_date_line(self, custom_date: Optional[str] = None, **kwargs):
        """添加落款日期行"""
        add_date_line(self.doc, custom_date=custom_date, **kwargs)

    def add_seal_placeholder(self, position: str = "right_bottom", **kwargs):
        """添加印章占位符"""
        add_seal_placeholder(self.doc, position=position, **kwargs)

    # ------- 表格 -------

    def add_table(self, headers: Sequence[str], rows: Sequence[Sequence[Any]],
                  **kwargs):
        """添加表格（自动格式）"""
        return add_table(self.doc, headers, rows, **kwargs)

    # ------- 页眉页脚 -------

    def add_header(self, text: str):
        """添加页眉"""
        add_header(self.doc, text)

    def add_footer(self, text: str):
        """添加页脚"""
        add_footer(self.doc, text)

    def add_page_number(self, fmt: str = "第 {PAGE} 页 共 {NUMPAGES} 页"):
        """添加页码"""
        add_page_number(self.doc, fmt=fmt)

    # ------- 校验 / 保存 -------

    def validate(self) -> Tuple[bool, List[str]]:
        """验证 docx 完整性"""
        return validate_docx(self.doc)

    def save(self) -> str:
        """校验 + 保存 docx，返回绝对路径"""
        ok, issues = self.validate()
        if not ok:
            # 不抛异常，仅打印警告，保持兼容
            sys.stderr.write(
                "[docx_common] 校验警告：" + "; ".join(issues) + "\n"
            )
        return save_docx(self.doc, self.output_path)

    # ------- 兼容接口 -------

    def __getattr__(self, name: str):
        # 允许子类直接访问 self.doc 等成员，未定义属性才报错
        raise AttributeError(name)


# ============================================================
# 11. 便捷工厂
# ============================================================

def create_docx(
    output_path: Union[str, Path],
    margins: Optional[Dict[str, float]] = None,
) -> _Document:
    """便捷工厂：创建一个已设置 A4 页面的空 Document

    适用于不想用 DocxBuilder 类、但希望快速拿一个标准 Document 的场景。

    Example:
        doc = create_docx("/tmp/test.docx")
        add_title(doc, "测试")
        doc.save("/tmp/test.docx")
    """
    doc = Document()
    setup_a4_page(doc, margins=margins)
    return doc


# ============================================================
# 12. 模块自检（仅 __main__ 时执行）
# ============================================================

def _self_test() -> int:
    """模块自检：生成一个最小 docx 并校验"""
    out = "/tmp/_docx_common_self_test.docx"
    b = DocxBuilder(out, with_page_number=True, header="自检页眉")
    b.add_title("docx_common 自检文档")
    b.add_salutation("尊敬的审查者：")
    b.add_paragraph("本段用于自检。")
    b.add_heading2("一、自检小节")
    b.add_paragraph("自检小节正文。")
    b.add_table(
        headers=["序号", "项目", "结果"],
        rows=[["1", "字体常量", "✓"], ["2", "表格", "✓"]],
    )
    b.add_this_salute()
    b.add_signature_block(["自检人：docx_common", "2025 年 1 月 1 日"])
    path = b.save()
    ok, issues = validate_docx(b.doc)
    print(f"[self_test] saved: {path}")
    print(f"[self_test] valid: {ok}, issues: {issues}")
    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(_self_test())
