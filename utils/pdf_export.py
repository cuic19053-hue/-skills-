#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF 输出共享模块（college-application-doc / utils / pdf_export.py）
====================================================================

为 22+ 个子 skill（national_scholarship / motivation_scholarship /
university_scholarship / enterprise_scholarship / single_scholarship /
challenge_cup / internet_plus / innovation_research / college_research /
university_research / entrepreneurship_training / entrepreneurship_practice /
party_application / party_full_member / thought_report / youth_league_application
/ outstanding_student / outstanding_graduate / outstanding_cadre /
civilized_student / grant_application / graduate_recommendation /
selected_graduate / social_survey / class_collective / major_transfer /
policy_lecture / tech_service / volunteer_teaching / western_plan 等）
提供统一的 docx → PDF 转换、PDF 合并、页码、水印、元数据、完整性校验能力。

设计目标
--------
1. **统一接口**：所有子 skill 通过 `from utils.pdf_export import docx_to_pdf`
   即可获得 PDF 输出能力，无需各自实现。
2. **多引擎兜底**：按 LibreOffice headless → docx2pdf → python-docx + reportlab
   顺序尝试，保证不同环境下均可用。
3. **高校场景**：内置 A+B 两套材料合并、添加"仅供提交"水印、设置元数据等
   高校申报常用功能。
4. **清晰报错**：缺 LibreOffice / docx2pdf 时给出明确安装指引，不让用户
   看到模糊的 ModuleNotFoundError。

环境要求
--------
- 优先：LibreOffice >= 7.0（命令 `libreoffice` 或 `soffice`，已预装）
- 可选：docx2pdf（pip install docx2pdf；Windows 端依赖 Word，Linux 端依赖 LibreOffice）
- 必装：pypdf >= 4.0（页码/水印/元数据/合并，已预装）
- 兜底：reportlab + python-docx（重建简单 docx，已预装）

CLI 用法
--------
    # 基本转换（自动同名 .pdf）
    python pdf_export.py --input xxx.docx

    # 指定输出路径
    python pdf_export.py --input xxx.docx --output xxx.pdf

    # 合并多个 PDF（A+B 两套材料）
    python pdf_export.py --merge a.pdf b.pdf --output merged.pdf

    # 添加页码
    python pdf_export.py --input xxx.pdf --page-number

    # 添加水印
    python pdf_export.py --input xxx.pdf --watermark "仅供提交"

    # 设置元数据
    python pdf_export.py --input xxx.pdf --metadata title="国家奖学金申请" author="张三"

    # 校验 PDF 完整性
    python pdf_export.py --input xxx.pdf --validate

退出码
------
    0 = 成功
    1 = 输入文件不存在
    2 = 转换失败（所有引擎均不可用）
    3 = PDF 完整性校验失败
    4 = 参数错误
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Iterable, List, Optional, Sequence, Tuple

# ------------------------------------------------------------------
# 引擎可用性探测（在 import 时执行一次，避免每次调用都重复探测）
# ------------------------------------------------------------------

_LIBREOFFICE_BIN: Optional[str] = None
for _candidate in ("libreoffice", "libreoffice7.6", "libreoffice7.4",
                   "soffice", "/usr/bin/libreoffice", "/usr/bin/soffice"):
    if shutil.which(_candidate):
        _LIBREOFFICE_BIN = _candidate
        break

try:
    import docx2pdf as _docx2pdf  # type: ignore
    _HAS_DOCX2PDF = True
except ImportError:
    _HAS_DOCX2PDF = False

try:
    from pypdf import PdfReader, PdfWriter  # type: ignore
    from pypdf.generic import NameObject, TextStringObject  # type: ignore
    _HAS_PYPDF = True
    _PDF_LIB = "pypdf"
except ImportError:
    try:
        from PyPDF2 import PdfReader, PdfWriter  # type: ignore
        from PyPDF2.generic import NameObject, TextStringObject  # type: ignore
        _HAS_PYPDF = True
        _PDF_LIB = "PyPDF2"
    except ImportError:
        _HAS_PYPDF = False
        _PDF_LIB = None

try:
    from reportlab.pdfgen import canvas  # type: ignore
    from reportlab.lib.pagesizes import A4  # type: ignore
    from reportlab.pdfbase import pdfmetrics  # type: ignore
    from reportlab.pdfbase.ttfonts import TTFont  # type: ignore
    _HAS_REPORTLAB = True
except ImportError:
    _HAS_REPORTLAB = False

try:
    from docx import Document as _DocxDocument  # type: ignore
    _HAS_PYTHON_DOCX = True
except ImportError:
    _HAS_PYTHON_DOCX = False


# ============================================================
# 自定义异常
# ============================================================

class PDFExportError(Exception):
    """PDF 导出基础异常"""


class LibreOfficeNotAvailableError(PDFExportError):
    """LibreOffice 未安装"""


class ConversionError(PDFExportError):
    """所有转换引擎均失败"""


class PDFValidationError(PDFExportError):
    """PDF 完整性校验失败"""


# ============================================================
# 工具函数：日志输出
# ============================================================

def _log(msg: str, level: str = "INFO") -> None:
    """统一日志格式输出到 stderr"""
    prefix = {"INFO": "ℹ️", "OK": "✅", "WARN": "⚠️", "ERR": "❌"}.get(level, "•")
    print(f"{prefix} {msg}", file=sys.stderr)


# ============================================================
# 引擎 1：LibreOffice headless（首选）
# ============================================================

def _convert_with_libreoffice(
    docx_path: Path,
    output_dir: Path,
    timeout: int = 120,
) -> Optional[Path]:
    """
    使用 LibreOffice headless 模式将 docx 转换为 PDF。

    LibreOffice 是当前最可靠的 docx → PDF 转换方案，能正确处理
    中文字体（宋体/黑体/仿宋/楷体）、表格、页眉页脚、首行缩进、
    段落对齐等所有 python-docx 写出的格式特性。

    Args:
        docx_path: 输入 docx 文件路径
        output_dir: 输出目录（PDF 将生成在此目录下，文件名与 docx 同名）
        timeout: 超时秒数（默认 120s，足够处理 50 页文档）

    Returns:
        成功时返回 PDF 文件路径，失败返回 None
    """
    if not _LIBREOFFICE_BIN:
        return None

    output_dir.mkdir(parents=True, exist_ok=True)
    # LibreOffice headless 必须使用独立的 user profile 防止并发冲突
    profile_dir = output_dir / ".lo_profile"
    cmd = [
        _LIBREOFFICE_BIN,
        f"-env:UserInstallation=file://{profile_dir}",
        "--headless",
        "--norestore",
        "--nolockcheck",
        "--convert-to", "pdf",
        "--outdir", str(output_dir),
        str(docx_path),
    ]
    try:
        proc = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=timeout,
            check=False,
        )
    except subprocess.TimeoutExpired:
        _log(f"LibreOffice 转换超时（{timeout}s）", "WARN")
        return None
    except Exception as exc:  # pragma: no cover - 防御性
        _log(f"LibreOffice 调用异常：{exc}", "WARN")
        return None

    expected_pdf = output_dir / (docx_path.stem + ".pdf")
    if proc.returncode != 0 or not expected_pdf.exists():
        stderr_text = proc.stderr.decode("utf-8", errors="replace")[:200]
        _log(f"LibreOffice 转换失败（rc={proc.returncode}）{stderr_text}", "WARN")
        return None
    return expected_pdf


# ============================================================
# 引擎 2：docx2pdf（次选，Windows 端依赖 MS Word）
# ============================================================

def _convert_with_docx2pdf(docx_path: Path, pdf_path: Path) -> Optional[Path]:
    """
    使用 docx2pdf 库转换 docx → PDF。

    注意：
    - Windows 上 docx2pdf 调用 MS Word COM 接口，转换质量最高
    - Linux/macOS 上 docx2pdf 内部仍是调用 LibreOffice
    - 多数高校申报场景在 Linux 服务器上，本引擎仅作补充

    Args:
        docx_path: 输入 docx 文件路径
        pdf_path: 输出 PDF 文件路径

    Returns:
        成功时返回 PDF 路径，失败返回 None
    """
    if not _HAS_DOCX2PDF:
        return None
    try:
        _docx2pdf.convert(str(docx_path), str(pdf_path))  # type: ignore[attr-defined]
    except Exception as exc:
        _log(f"docx2pdf 转换失败：{exc}", "WARN")
        return None
    if pdf_path.exists() and pdf_path.stat().st_size > 0:
        return pdf_path
    return None


# ============================================================
# 引擎 3：python-docx + reportlab 重建（兜底，仅简单 docx）
# ============================================================

def _convert_with_reportlab(docx_path: Path, pdf_path: Path) -> Optional[Path]:
    """
    兜底方案：解析 docx 文本并用 reportlab 重建 PDF。

    警告：此方案仅处理段落文本、加粗、表格，不处理页眉页脚、复杂样式、
    中文字体嵌入可能缺失。仅在前两种引擎均不可用时使用。

    Args:
        docx_path: 输入 docx 文件路径
        pdf_path: 输出 PDF 文件路径

    Returns:
        成功时返回 PDF 路径，失败返回 None
    """
    if not (_HAS_REPORTLAB and _HAS_PYTHON_DOCX):
        return None
    try:
        from reportlab.lib.units import cm
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import (SimpleDocTemplate, Paragraph,
                                         Spacer, Table, TableStyle)
        from reportlab.lib import colors
    except ImportError:
        return None

    # 尝试注册系统中文字体
    _register_chinese_fonts()

    doc = _DocxDocument(str(docx_path))

    pdf_doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        topMargin=2.54 * cm,
        bottomMargin=2.54 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "CNBody", parent=styles["Normal"],
        fontName="CN" if "CN" in pdfmetrics.getRegisteredFontNames() else "Helvetica",
        fontSize=12, leading=18, firstLineIndent=24,
        spaceBefore=0, spaceAfter=6,
    )
    heading_style = ParagraphStyle(
        "CNHeading", parent=styles["Normal"],
        fontName="CN-Bold" if "CN-Bold" in pdfmetrics.getRegisteredFontNames()
        else "Helvetica-Bold",
        fontSize=16, leading=22, alignment=1,
        spaceBefore=12, spaceAfter=12,
    )

    story: List = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 6))
            continue
        # 识别标题（简化规则：居中且字号大）
        if (para.alignment is not None
                and para.alignment == 1  # WD_ALIGN_PARAGRAPH.CENTER
                and len(text) < 40):
            story.append(Paragraph(_xml_escape(text), heading_style))
        else:
            story.append(Paragraph(_xml_escape(text), body_style))

    # 简单处理表格
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([_xml_escape(cell.text.strip()) for cell in row.cells])
        if data:
            try:
                tbl = Table(data)
                tbl.setStyle(TableStyle([
                    ("FONTNAME", (0, 0), (-1, -1),
                     "CN" if "CN" in pdfmetrics.getRegisteredFontNames() else "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 6))
            except Exception:
                pass

    try:
        pdf_doc.build(story)
    except Exception as exc:
        _log(f"reportlab 重建失败：{exc}", "WARN")
        return None
    if pdf_path.exists() and pdf_path.stat().st_size > 0:
        return pdf_path
    return None


def _register_chinese_fonts() -> None:
    """在 reportlab 中注册系统中文字体（如 WenQuanYi / Noto CJK）"""
    if "CN" in pdfmetrics.getRegisteredFontNames():
        return
    candidates = [
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    ]
    for path in candidates:
        if Path(path).exists():
            try:
                pdfmetrics.registerFont(TTFont("CN", path))
                # 同时注册粗体（简化处理：复用同字体）
                pdfmetrics.registerFont(TTFont("CN-Bold", path))
                return
            except Exception:
                continue


def _xml_escape(text: str) -> str:
    """转义 XML 特殊字符供 reportlab Paragraph 使用"""
    return (text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;"))


# ============================================================
# 主函数：docx_to_pdf
# ============================================================

def docx_to_pdf(
    docx_path: str | Path,
    pdf_path: Optional[str | Path] = None,
    *,
    engine: str = "auto",
    timeout: int = 120,
) -> Path:
    """
    将 docx 文件转换为 PDF（多引擎兜底）。

    转换优先级（engine="auto" 时）：
        1. LibreOffice headless  —— 最可靠，能完整保留 docx 格式
        2. docx2pdf              —— Windows 上调用 Word，质量最高
        3. reportlab + python-docx —— 兜底，仅简单文档，中文字体可能丢失

    Args:
        docx_path: 输入 docx 文件路径
        pdf_path:  输出 PDF 路径；None 时自动同名 .pdf（与 docx 同目录）
        engine:    引擎选择：auto / libreoffice / docx2pdf / reportlab
        timeout:   LibreOffice 超时秒数

    Returns:
        生成的 PDF 文件绝对路径

    Raises:
        FileNotFoundError: docx 文件不存在
        ConversionError:   所有引擎均失败
        LibreOfficeNotAvailableError: 指定 libreoffice 但系统未安装
    """
    docx_path = Path(docx_path).resolve()
    if not docx_path.exists():
        raise FileNotFoundError(f"docx 文件不存在：{docx_path}")
    if docx_path.suffix.lower() != ".docx":
        _log(f"输入文件扩展名非 .docx：{docx_path.name}", "WARN")

    if pdf_path is None:
        pdf_path = docx_path.with_suffix(".pdf")
    pdf_path = Path(pdf_path).resolve()
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    # 引擎顺序
    if engine == "auto":
        engines: Tuple[str, ...] = ("libreoffice", "docx2pdf", "reportlab")
    else:
        engines = (engine,)

    last_error: Optional[str] = None
    for eng in engines:
        if eng == "libreoffice":
            if not _LIBREOFFICE_BIN:
                last_error = (
                    "LibreOffice 未安装。安装方式：\n"
                    "  Ubuntu/Debian: sudo apt install libreoffice\n"
                    "  CentOS/RHEL:   sudo yum install libreoffice\n"
                    "  macOS:         brew install --cask libreoffice\n"
                    "  Windows:       https://www.libreoffice.org/download/"
                )
                _log("LibreOffice 未安装，跳过", "WARN")
                continue
            _log(f"尝试 LibreOffice 转换：{docx_path.name}")
            # LibreOffice 输出到临时目录，再 move 到目标路径（避免 outdir 污染）
            with tempfile.TemporaryDirectory(prefix="lo_pdf_") as tmpdir:
                result = _convert_with_libreoffice(
                    docx_path, Path(tmpdir), timeout=timeout)
                if result and result.exists():
                    shutil.move(str(result), str(pdf_path))
                    _log(f"LibreOffice 转换成功：{pdf_path.name}", "OK")
                    return pdf_path
        elif eng == "docx2pdf":
            if not _HAS_DOCX2PDF:
                last_error = "docx2pdf 未安装：pip install docx2pdf"
                _log("docx2pdf 未安装，跳过", "WARN")
                continue
            _log(f"尝试 docx2pdf 转换：{docx_path.name}")
            result = _convert_with_docx2pdf(docx_path, pdf_path)
            if result:
                _log(f"docx2pdf 转换成功：{pdf_path.name}", "OK")
                return pdf_path
        elif eng == "reportlab":
            if not (_HAS_REPORTLAB and _HAS_PYTHON_DOCX):
                last_error = "reportlab/python-docx 未安装"
                _log("reportlab/python-docx 未安装，跳过", "WARN")
                continue
            _log(f"尝试 reportlab 重建：{docx_path.name}")
            _log("注意：reportlab 兜底方案可能丢失复杂格式，建议安装 LibreOffice",
                 "WARN")
            result = _convert_with_reportlab(docx_path, pdf_path)
            if result:
                _log(f"reportlab 重建成功：{pdf_path.name}", "OK")
                return pdf_path
        else:
            raise ValueError(f"未知引擎：{eng}（可选：auto/libreoffice/docx2pdf/reportlab）")

    # 所有引擎都失败
    hint = last_error or "未知原因"
    raise ConversionError(
        f"docx → PDF 转换失败（所有引擎均不可用）。\n"
        f"输入：{docx_path}\n"
        f"原因：{hint}\n"
        f"建议：优先安装 LibreOffice 以获得最佳转换质量。"
    )


# ============================================================
# PDF 合并
# ============================================================

def merge_pdfs(
    pdf_list: Sequence[str | Path],
    output_path: str | Path,
) -> Path:
    """
    合并多个 PDF 文件（典型场景：A 类审批表 + B 类申请书合并为一份提交材料）。

    Args:
        pdf_list:    要合并的 PDF 路径列表（按顺序合并）
        output_path: 合并后输出路径

    Returns:
        合并后的 PDF 绝对路径

    Raises:
        FileNotFoundError: 任一输入 PDF 不存在
        PDFExportError:    pypdf/PyPDF2 未安装或合并失败
    """
    if not _HAS_PYPDF:
        raise PDFExportError(
            "合并 PDF 需要 pypdf 或 PyPDF2：pip install pypdf"
        )
    if not pdf_list:
        raise ValueError("pdf_list 不能为空")

    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    writer = PdfWriter()
    for idx, pdf in enumerate(pdf_list, 1):
        pdf = Path(pdf).resolve()
        if not pdf.exists():
            raise FileNotFoundError(f"第 {idx} 个 PDF 不存在：{pdf}")
        _log(f"读取：{pdf.name}")
        reader = PdfReader(str(pdf))
        for page in reader.pages:
            writer.add_page(page)
        # 保留源文档书签作为 outline（若存在）
        try:
            if reader.outline:
                _log(f"  → 包含 {len(reader.pages)} 页")
        except Exception:
            pass

    with open(output_path, "wb") as f:
        writer.write(f)

    size_kb = output_path.stat().st_size / 1024
    _log(f"合并完成：{output_path.name}（{size_kb:.1f} KB，{len(pdf_list)} 份）", "OK")
    return output_path


# ============================================================
# 添加页码
# ============================================================

def add_page_number(
    pdf_path: str | Path,
    *,
    output_path: Optional[str | Path] = None,
    position: str = "bottom-center",
    font_size: int = 9,
    format: str = "{page}/{total}",
) -> Path:
    """
    为 PDF 每页添加页码。

    Args:
        pdf_path:     输入 PDF 路径
        output_path:  输出路径；None 时覆盖原文件
        position:     位置：bottom-center / bottom-right / bottom-left /
                      top-center / top-right / top-left
        font_size:    页码字号
        format:       页码格式，{page} 与 {total} 为占位符

    Returns:
        处理后的 PDF 绝对路径
    """
    if not _HAS_PYPDF:
        raise PDFExportError("添加页码需要 pypdf：pip install pypdf")
    if not _HAS_REPORTLAB:
        raise PDFExportError("添加页码需要 reportlab：pip install reportlab")

    pdf_path = Path(pdf_path).resolve()
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF 不存在：{pdf_path}")

    if output_path is None:
        output_path = pdf_path
    output_path = Path(output_path).resolve()

    reader = PdfReader(str(pdf_path))
    total = len(reader.pages)
    writer = PdfWriter()

    # 位置映射
    pos_map = {
        "bottom-center": (lambda w, h, tw: (w / 2 - tw / 2, 20)),
        "bottom-right":  (lambda w, h, tw: (w - tw - 30, 20)),
        "bottom-left":   (lambda w, h, tw: (30, 20)),
        "top-center":    (lambda w, h, tw: (w / 2 - tw / 2, h - 30)),
        "top-right":     (lambda w, h, tw: (w - tw - 30, h - 30)),
        "top-left":      (lambda w, h, tw: (30, h - 30)),
    }
    if position not in pos_map:
        raise ValueError(f"未知位置：{position}（可选：{list(pos_map.keys())}）")

    _register_chinese_fonts()
    font_name = "CN" if "CN" in pdfmetrics.getRegisteredFontNames() else "Helvetica"

    with tempfile.TemporaryDirectory(prefix="pgnum_") as tmpdir:
        for idx, page in enumerate(reader.pages, 1):
            box = page.mediabox
            w = float(box.width)
            h = float(box.height)
            text = format.format(page=idx, total=total)
            tw = len(text) * font_size * 0.5  # 估算文本宽度

            # 生成单页水印 PDF
            wm_pdf = Path(tmpdir) / f"pg_{idx}.pdf"
            c = canvas.Canvas(str(wm_pdf), pagesize=(w, h))
            c.setFont(font_name, font_size)
            x, y = pos_map[position](w, h, tw)
            c.drawCentredString(x + tw / 2, y, text)
            c.save()

            wm_reader = PdfReader(str(wm_pdf))
            page.merge_page(wm_reader.pages[0])
            writer.add_page(page)

        with open(output_path, "wb") as f:
            writer.write(f)

    _log(f"已添加页码（{total} 页，{position}）：{output_path.name}", "OK")
    return output_path


# ============================================================
# 添加水印
# ============================================================

def add_watermark(
    pdf_path: str | Path,
    watermark_text: str,
    *,
    output_path: Optional[str | Path] = None,
    opacity: float = 0.15,
    font_size: int = 60,
    rotation: int = 45,
    color: Tuple[float, float, float] = (0.5, 0.5, 0.5),
) -> Path:
    """
    为 PDF 每页添加斜向水印（典型：高校申报材料加"仅供提交""草稿"等）。

    Args:
        pdf_path:        输入 PDF 路径
        watermark_text:  水印文本
        output_path:     输出路径；None 时覆盖原文件
        opacity:         透明度 0~1（默认 0.15 较淡）
        font_size:       水印字号
        rotation:        旋转角度（默认 45° 斜向）
        color:           RGB 颜色三元组（0~1）

    Returns:
        处理后的 PDF 绝对路径
    """
    if not _HAS_PYPDF:
        raise PDFExportError("添加水印需要 pypdf：pip install pypdf")
    if not _HAS_REPORTLAB:
        raise PDFExportError("添加水印需要 reportlab：pip install reportlab")

    pdf_path = Path(pdf_path).resolve()
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF 不存在：{pdf_path}")
    if not watermark_text:
        raise ValueError("watermark_text 不能为空")

    if output_path is None:
        output_path = pdf_path
    output_path = Path(output_path).resolve()

    reader = PdfReader(str(pdf_path))
    total = len(reader.pages)
    writer = PdfWriter()

    _register_chinese_fonts()
    font_name = "CN" if "CN" in pdfmetrics.getRegisteredFontNames() else "Helvetica"

    with tempfile.TemporaryDirectory(prefix="wm_") as tmpdir:
        # 水印 PDF 按页生成（每页尺寸可能不同）
        for idx, page in enumerate(reader.pages):
            box = page.mediabox
            w = float(box.width)
            h = float(box.height)

            wm_pdf = Path(tmpdir) / f"wm_{idx}.pdf"
            c = canvas.Canvas(str(wm_pdf), pagesize=(w, h))
            c.setFont(font_name, font_size)
            c.setFillColorRGB(*color)
            try:
                c.setFillAlpha(opacity)
            except Exception:
                # 旧版 reportlab 无 setFillAlpha
                pass
            c.saveState()
            c.translate(w / 2, h / 2)
            c.rotate(rotation)
            c.drawCentredString(0, 0, watermark_text)
            c.restoreState()
            c.save()

            wm_reader = PdfReader(str(wm_pdf))
            page.merge_page(wm_reader.pages[0])
            writer.add_page(page)

        with open(output_path, "wb") as f:
            writer.write(f)

    _log(f"已添加水印「{watermark_text}」（{total} 页）：{output_path.name}", "OK")
    return output_path


# ============================================================
# 设置 PDF 元数据
# ============================================================

def set_pdf_metadata(
    pdf_path: str | Path,
    title: Optional[str] = None,
    author: Optional[str] = None,
    subject: Optional[str] = None,
    keywords: Optional[str] = None,
    creator: Optional[str] = None,
) -> Path:
    """
    设置 PDF 元数据（标题/作者/主题/关键词/创建者）。

    高校电子申报系统常根据 PDF 元数据归档，建议设置：
        - title:   文档标题（如"国家奖学金申请书"）
        - author:  申请人姓名
        - subject: 学号 + 院系
        - keywords: 申报类型 + 年度

    Args:
        pdf_path: 输入 PDF 路径（原地覆盖）
        title/author/subject/keywords/creator: 元数据字段

    Returns:
        处理后的 PDF 绝对路径
    """
    if not _HAS_PYPDF:
        raise PDFExportError("设置元数据需要 pypdf：pip install pypdf")

    pdf_path = Path(pdf_path).resolve()
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF 不存在：{pdf_path}")

    reader = PdfReader(str(pdf_path))
    writer = PdfWriter(clone_from=reader)

    # pypdf 元数据字段映射
    meta_map = {
        "/Title": title,
        "/Author": author,
        "/Subject": subject,
        "/Keywords": keywords,
        "/Creator": creator,
        "/Producer": "college-application-doc/utils/pdf_export.py",
    }
    for key, value in meta_map.items():
        if value is not None:
            writer.add_metadata({key: str(value)})

    tmp = pdf_path.with_suffix(".meta.tmp.pdf")
    with open(tmp, "wb") as f:
        writer.write(f)
    shutil.move(str(tmp), str(pdf_path))

    _log(f"已设置元数据：{pdf_path.name}", "OK")
    return pdf_path


# ============================================================
# PDF 完整性校验
# ============================================================

def validate_pdf(pdf_path: str | Path) -> bool:
    """
    验证 PDF 文件完整性（结构、页数、可解析性）。

    Args:
        pdf_path: 输入 PDF 路径

    Returns:
        True = 完整可用

    Raises:
        FileNotFoundError: 文件不存在
        PDFValidationError: 校验失败（详细原因在异常 message）
    """
    if not _HAS_PYPDF:
        raise PDFExportError("校验 PDF 需要 pypdf：pip install pypdf")

    pdf_path = Path(pdf_path).resolve()
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF 不存在：{pdf_path}")

    size = pdf_path.stat().st_size
    if size == 0:
        raise PDFValidationError(f"PDF 文件为空：{pdf_path}")

    # 检查文件头
    with open(pdf_path, "rb") as f:
        header = f.read(8)
    if not header.startswith(b"%PDF-"):
        raise PDFValidationError(
            f"文件头非 %PDF-，不是有效 PDF：{pdf_path}（header={header!r}）"
        )

    # 检查 EOF 标记
    with open(pdf_path, "rb") as f:
        f.seek(-1024 if size > 1024 else 0, 2)
        tail = f.read()
    if b"%%EOF" not in tail:
        raise PDFValidationError(
            f"未找到 %%EOF 标记，PDF 可能被截断：{pdf_path}"
        )

    # 尝试用 pypdf 解析
    try:
        reader = PdfReader(str(pdf_path))
        n_pages = len(reader.pages)
        if n_pages == 0:
            raise PDFValidationError(f"PDF 页数为 0：{pdf_path}")
        # 尝试读取第一页内容流（确保不是损坏的对象）
        _ = reader.pages[0].mediabox
    except PDFValidationError:
        raise
    except Exception as exc:
        raise PDFValidationError(
            f"PDF 解析失败：{pdf_path}（{type(exc).__name__}: {exc}）"
        ) from exc

    _log(f"PDF 完整性校验通过：{pdf_path.name}（{n_pages} 页，{size/1024:.1f} KB）",
         "OK")
    return True


# ============================================================
# CLI 接口
# ============================================================

def _build_argparser() -> argparse.ArgumentParser:
    """构建 CLI 参数解析器"""
    p = argparse.ArgumentParser(
        prog="pdf_export.py",
        description="docx → PDF 转换 / PDF 合并 / 页码 / 水印 / 元数据 / 校验（college-application-doc 共享工具）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例：
  # docx 转 PDF（自动同名 .pdf）
  python pdf_export.py --input xxx.docx

  # 指定输出路径
  python pdf_export.py --input xxx.docx --output xxx.pdf

  # 合并多个 PDF（A+B 两套材料）
  python pdf_export.py --merge a.pdf b.pdf --output merged.pdf

  # 添加页码
  python pdf_export.py --input xxx.pdf --page-number

  # 添加水印
  python pdf_export.py --input xxx.pdf --watermark "仅供提交"

  # 设置元数据
  python pdf_export.py --input xxx.pdf --metadata title="国家奖学金申请" author="张三"

  # 校验 PDF 完整性
  python pdf_export.py --input xxx.pdf --validate
""",
    )
    p.add_argument("--input", "-i", help="输入文件（docx 或 pdf）")
    p.add_argument("--output", "-o", help="输出文件路径")
    p.add_argument("--engine", default="auto",
                   choices=["auto", "libreoffice", "docx2pdf", "reportlab"],
                   help="docx → PDF 转换引擎（默认 auto）")
    p.add_argument("--merge", nargs="+", metavar="PDF",
                   help="合并多个 PDF（按顺序），配合 --output 使用")
    p.add_argument("--page-number", action="store_true",
                   help="为 PDF 添加页码")
    p.add_argument("--watermark", metavar="TEXT",
                   help='添加水印文本（如 "仅供提交"）')
    p.add_argument("--metadata", nargs="+", metavar="KEY=VALUE",
                   help='设置元数据，如 title="国家奖学金" author="张三"')
    p.add_argument("--validate", action="store_true",
                   help="校验 PDF 完整性")
    p.add_argument("--timeout", type=int, default=120,
                   help="LibreOffice 转换超时秒数（默认 120）")
    return p


def _parse_metadata(items: List[str]) -> dict:
    """解析 --metadata key=value 列表为字典"""
    result: dict = {}
    for item in items:
        if "=" not in item:
            raise ValueError(f"--metadata 参数格式错误，应为 key=value：{item}")
        key, _, value = item.partition("=")
        key = key.strip().lower()
        if key not in {"title", "author", "subject", "keywords", "creator"}:
            raise ValueError(f"未知元数据字段：{key}（可选：title/author/subject/keywords/creator）")
        result[key] = value
    return result


def main(argv: Optional[Sequence[str]] = None) -> int:
    """CLI 主入口"""
    parser = _build_argparser()
    args = parser.parse_args(argv)

    # 合并模式
    if args.merge:
        if not args.output:
            _log("--merge 模式必须指定 --output", "ERR")
            return 4
        try:
            merge_pdfs(args.merge, args.output)
            return 0
        except Exception as exc:
            _log(f"合并失败：{exc}", "ERR")
            return 2

    # 单文件操作模式
    if not args.input:
        parser.print_help()
        return 4

    input_path = Path(args.input)
    if not input_path.exists():
        _log(f"输入文件不存在：{input_path}", "ERR")
        return 1

    suffix = input_path.suffix.lower()

    try:
        # docx → PDF 转换
        if suffix == ".docx":
            pdf_path = docx_to_pdf(
                input_path,
                args.output,
                engine=args.engine,
                timeout=args.timeout,
            )
            _log(f"✓ 转换完成：{pdf_path}", "OK")
            # 转换后追加操作
            _post_process(pdf_path, args)
            return 0

        # PDF 后处理
        if suffix == ".pdf":
            if args.validate:
                validate_pdf(input_path)
                return 0
            # 复制到输出路径（避免原地操作丢失原件）
            target = Path(args.output) if args.output else input_path
            if args.output and args.output != str(input_path):
                shutil.copy2(input_path, target)
            _post_process(target, args)
            return 0

        _log(f"不支持的文件类型：{suffix}（仅支持 .docx / .pdf）", "ERR")
        return 4

    except LibreOfficeNotAvailableError as exc:
        _log(str(exc), "ERR")
        return 2
    except ConversionError as exc:
        _log(str(exc), "ERR")
        return 2
    except PDFValidationError as exc:
        _log(str(exc), "ERR")
        return 3
    except FileNotFoundError as exc:
        _log(str(exc), "ERR")
        return 1
    except Exception as exc:
        _log(f"未预期错误：{type(exc).__name__}: {exc}", "ERR")
        return 2


def _post_process(pdf_path: Path, args: argparse.Namespace) -> None:
    """对生成的 PDF 执行后处理（页码/水印/元数据/校验）"""
    if args.page_number:
        add_page_number(pdf_path)
    if args.watermark:
        add_watermark(pdf_path, args.watermark)
    if args.metadata:
        meta = _parse_metadata(args.metadata)
        set_pdf_metadata(pdf_path, **meta)
    if args.validate:
        validate_pdf(pdf_path)


# ============================================================
# 模块自检
# ============================================================

def selftest() -> dict:
    """
    运行模块自检，返回各引擎/库的可用性。

    Returns:
        dict 报告各组件状态，便于排查环境问题
    """
    report = {
        "libreoffice": {
            "available": _LIBREOFFICE_BIN is not None,
            "path": _LIBREOFFICE_BIN,
        },
        "docx2pdf": {"available": _HAS_DOCX2PDF},
        "pypdf_lib": {"available": _HAS_PYPDF, "name": _PDF_LIB},
        "reportlab": {"available": _HAS_REPORTLAB},
        "python_docx": {"available": _HAS_PYTHON_DOCX},
        "recommended_engine": (
            "libreoffice" if _LIBREOFFICE_BIN
            else ("docx2pdf" if _HAS_DOCX2PDF
                  else ("reportlab" if _HAS_REPORTLAB and _HAS_PYTHON_DOCX
                        else None))
        ),
    }
    return report


if __name__ == "__main__":
    # 支持 --selftest 快速环境检查
    if "--selftest" in sys.argv:
        report = selftest()
        for k, v in report.items():
            print(f"{k:20s}: {v}")
        sys.exit(0)
    sys.exit(main())
