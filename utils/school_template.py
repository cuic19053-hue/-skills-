#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
学校模板适配层（School Template Adapter）
==========================================

为 22+ 个子 skill（national_scholarship / motivation_scholarship /
innovation_research / party_application / outstanding_graduate 等）提供
统一的学校差异适配能力，让 build.py 不再写死学校特定的：

  - 页眉/页脚（含学校名 / 校徽 / 页码）
  - 印章（红章位置 / 大小 / 透明度）
  - 字段名映射（如"学习成绩" vs "GPA" vs "绩点"）
  - 签字栏（学生 / 辅导员 / 院系 / 学校 4 方）
  - 字体偏好（部分学校要求仿宋而非宋体）
  - 页边距偏好
  - 申请理由字数限制（如清华 200 字 vs 北大 250 字）

设计原则
--------
1. **零侵入**：现有 build.py 无需改动即可使用，仅在需要时调用本模块
2. **配置驱动**：所有学校差异通过 JSON 配置文件描述，新增学校只需新增
   一个 JSON 文件 + 一行注册
3. **渐进降级**：未配置的字段自动回退到 default 模板
4. **可热加载**：runtime 可通过 `register_template()` 注入新学校配置

典型用法
--------
    from utils.school_template import load_template, apply_template

    tpl = load_template("tsinghua")
    apply_template(doc, tpl)            # 应用页边距/页眉/页脚/印章
    gpa_label = tpl.get_field("gpa")    # -> "GPA" 而非"学习成绩"
    body_font = tpl.fonts["body"]       # -> "仿宋" 而非"宋体"

目录结构
--------
    utils/
    ├── school_template.py        # 本文件
    └── schools/
        ├── template_default.json # 默认配置（所有字段的兜底值）
        ├── template_tsinghua.json
        ├── template_pku.json
        ├── template_zju.json
        └── template_whu.json

Author: T37 sub-agent
"""

from __future__ import annotations

import json
import os
import re
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

# ============================================================
# 路径常量
# ============================================================

_THIS_DIR = Path(__file__).resolve().parent
SCHOOLS_DIR = _THIS_DIR / "schools"
DEFAULT_TEMPLATE_NAME = "default"

# 字段名规范化正则：用于把用户传入的学校别名（如 "THU" / "清华" /
# "清华大学"）映射到统一的配置文件名 "tsinghua"
_ALIAS_NORMALIZE_RULES = [
    (re.compile(r"清华(大学)?"), "tsinghua"),
    (re.compile(r"北京大学|北大"), "pku"),
    (re.compile(r"浙江(大学)?|浙大"), "zju"),
    (re.compile(r"武汉(大学)?|武大"), "whu"),
    (re.compile(r"复旦(大学)?"), "fudan"),
    (re.compile(r"上海交通(大学)?|上交"), "sjtu"),
    (re.compile(r"南京(大学)?|南大"), "nju"),
    (re.compile(r"中山(大学)?|中大"), "sysu"),
    (re.compile(r"华中科技(大学)?|华科"), "hust"),
    (re.compile(r"中国人民(大学)?|人大"), "ruc"),
]


# ============================================================
# SchoolTemplate 数据类
# ============================================================


class SchoolTemplate:
    """封装一所学校的全部排版/字段/签字差异。

    所有字段都设有默认值，构造时缺失的字段会从 `template_default.json`
    自动补齐（参见 `load_template`）。
    """

    # --- 必填字段 ---
    school_name: str

    # --- 页眉页脚 ---
    header: Dict[str, Any]
    footer: Dict[str, Any]

    # --- 印章 ---
    seal: Dict[str, Any]

    # --- 字段名映射 ---
    field_mapping: Dict[str, str]

    # --- 字体偏好 ---
    fonts: Dict[str, Any]

    # --- 页边距 ---
    margins: Dict[str, float]

    # --- 签字栏顺序与角色 ---
    signature_blocks: List[str]

    # --- 申请理由字数限制 ---
    apply_reason_chars: Dict[str, int]

    def __init__(self, config: Dict[str, Any]) -> None:
        if "school_name" not in config:
            raise ValueError("学校配置必须包含 school_name 字段")
        self.school_name = config["school_name"]
        self.header = config.get("header", {"enabled": False, "text": "", "logo": ""})
        self.footer = config.get(
            "footer", {"enabled": True, "text": "", "page_number": True}
        )
        self.seal = config.get(
            "seal", {"enabled": False, "position": "right_bottom", "size_cm": 4}
        )
        self.field_mapping = config.get(
            "field_mapping",
            {
                "gpa": "学习成绩",
                "rank": "排名",
                "major": "专业",
                "class": "班级",
            },
        )
        self.fonts = config.get(
            "fonts",
            {"body": "宋体", "heading": "黑体", "title_size": "22"},
        )
        self.margins = config.get(
            "margins",
            {
                "top_cm": 2.54,
                "bottom_cm": 2.54,
                "left_cm": 2.5,
                "right_cm": 2.5,
            },
        )
        self.signature_blocks = config.get(
            "signature_blocks",
            ["申请人", "辅导员", "院系负责人", "学校负责人"],
        )
        self.apply_reason_chars = config.get(
            "apply_reason_chars", {"min": 180, "max": 200}
        )
        # 保留原始 config 以便扩展字段透传
        self._raw: Dict[str, Any] = config

    # ------------------------------------------------------------
    # 字段映射 API
    # ------------------------------------------------------------

    def get_field(self, field: str, default: Optional[str] = None) -> str:
        """获取字段名映射。

        >>> tpl = SchoolTemplate({"school_name": "x", "field_mapping": {"gpa": "GPA"}})
        >>> tpl.get_field("gpa")
        'GPA'
        >>> tpl.get_field("non_exists", "fallback")
        'fallback'
        """
        return self.field_mapping.get(field, default or field)

    def remap_record(self, record: Dict[str, Any]) -> Dict[str, Any]:
        """把字典的 key 按字段映射重命名（用于把内部统一 key 翻译成
        学校特定展示名）。"""
        out: Dict[str, Any] = {}
        for k, v in record.items():
            out[self.get_field(k)] = v
        return out

    # ------------------------------------------------------------
    # 字体 API
    # ------------------------------------------------------------

    def body_font(self) -> str:
        return str(self.fonts.get("body", "宋体"))

    def heading_font(self) -> str:
        return str(self.fonts.get("heading", "黑体"))

    def title_size_pt(self) -> float:
        """title_size 字段支持两种格式："22"（磅值字符串）或 22（数字）"""
        size = self.fonts.get("title_size", 22)
        try:
            return float(size)
        except (TypeError, ValueError):
            return 22.0

    # ------------------------------------------------------------
    # 字数限制 API
    # ------------------------------------------------------------

    def apply_reason_min(self) -> int:
        return int(self.apply_reason_chars.get("min", 180))

    def apply_reason_max(self) -> int:
        return int(self.apply_reason_chars.get("max", 200))

    def validate_apply_reason(self, text: str) -> Dict[str, Any]:
        """检查申请理由字数是否在区间内，返回 {ok, length, min, max, message}"""
        length = len(text)
        lo, hi = self.apply_reason_min(), self.apply_reason_max()
        if length < lo:
            ok, msg = False, f"字数 {length} 少于下限 {lo}"
        elif length > hi:
            ok, msg = False, f"字数 {length} 超过上限 {hi}"
        else:
            ok, msg = True, f"字数 {length} 在区间 [{lo}, {hi}] 内"
        return {"ok": ok, "length": length, "min": lo, "max": hi, "message": msg}

    # ------------------------------------------------------------
    # 字典/JSON 序列化
    # ------------------------------------------------------------

    def to_dict(self) -> Dict[str, Any]:
        return deepcopy(self._raw)

    def to_json(self, indent: int = 2) -> str:
        return json.dumps(self._raw, ensure_ascii=False, indent=indent)

    def __repr__(self) -> str:
        return f"SchoolTemplate(school_name={self.school_name!r})"


# ============================================================
# 配置加载 / 注册中心
# ============================================================


def _normalize_school_name(school_name: str) -> str:
    """把各种别名规范化为配置文件名片段。

    - "清华大学" / "清华" / "THU" -> "tsinghua"
    - "北京大学" / "北大" / "PKU" -> "pku"
    - "浙江" / "浙大" -> "zju"
    - "武汉" / "武大" -> "whu"
    - 其他保持原样（小写化）
    """
    if not school_name:
        return DEFAULT_TEMPLATE_NAME
    name = str(school_name).strip()
    # 英文缩写直接小写
    if re.fullmatch(r"[A-Za-z]{2,10}", name):
        return name.lower()
    for pattern, target in _ALIAS_NORMALIZE_RULES:
        if pattern.search(name):
            return target
    return name.lower()


def _template_file_path(school_name: str) -> Path:
    """构造 template_<name>.json 路径"""
    normalized = _normalize_school_name(school_name)
    return SCHOOLS_DIR / f"template_{normalized}.json"


# 运行时注册中心：非文件来源的学校配置（通过 register_template 注入）
_RUNTIME_REGISTRY: Dict[str, Dict[str, Any]] = {}


def register_template(school_name: str, config_dict: Dict[str, Any]) -> SchoolTemplate:
    """运行时注册一所新学校，无需写 JSON 文件。

    示例：

        register_template("fudan", {
            "school_name": "复旦大学",
            "fonts": {"body": "仿宋"},
            "margins": {"top_cm": 2.8, "bottom_cm": 2.8,
                        "left_cm": 2.5, "right_cm": 2.5},
        })

    之后再 `load_template("fudan")` 即可拿到。
    """
    if "school_name" not in config_dict:
        config_dict = {**config_dict, "school_name": school_name}
    normalized = _normalize_school_name(school_name)
    _RUNTIME_REGISTRY[normalized] = deepcopy(config_dict)
    return SchoolTemplate(config_dict)


def _merge_with_default(config: Dict[str, Any]) -> Dict[str, Any]:
    """以 template_default.json 为底，逐字段深合并用户配置。

    缺失字段回退到默认；列表/标量直接覆盖；字典递归合并。
    """
    default_path = SCHOOLS_DIR / f"template_{DEFAULT_TEMPLATE_NAME}.json"
    if not default_path.exists():
        return deepcopy(config)
    with open(default_path, "r", encoding="utf-8") as f:
        default_cfg = json.load(f)

    def _deep_merge(base: Dict[str, Any], override: Dict[str, Any]) -> Dict[str, Any]:
        merged = deepcopy(base)
        for k, v in override.items():
            if k in merged and isinstance(merged[k], dict) and isinstance(v, dict):
                merged[k] = _deep_merge(merged[k], v)
            else:
                merged[k] = deepcopy(v)
        return merged

    return _deep_merge(default_cfg, config)


def load_template(school_name: str = DEFAULT_TEMPLATE_NAME) -> SchoolTemplate:
    """加载一所学校的模板配置。

    查找顺序：
      1. 运行时注册表（`register_template` 注入的）
      2. `schools/template_<name>.json` 文件
      3. 若都不存在，回退到 `template_default.json`，并在配置中
         用传入的 school_name 替换默认值（方便后续打印输出）

    参数：
        school_name: 学校名（支持中英文别名，如 "清华"/"清华大学"/"THU"）

    返回：
        SchoolTemplate 实例

    异常：
        FileNotFoundError: 当 schools/ 目录完全不存在时
    """
    if not SCHOOLS_DIR.exists():
        raise FileNotFoundError(
            f"学校配置目录不存在: {SCHOOLS_DIR}\n"
            "请确认 school_template.py 与 schools/ 在同一父目录下。"
        )

    normalized = _normalize_school_name(school_name)

    # 1. 运行时注册表
    if normalized in _RUNTIME_REGISTRY:
        cfg = _merge_with_default(_RUNTIME_REGISTRY[normalized])
        return SchoolTemplate(cfg)

    # 2. 文件
    cfg_path = _template_file_path(normalized)
    if cfg_path.exists():
        with open(cfg_path, "r", encoding="utf-8") as f:
            cfg = json.load(f)
        cfg = _merge_with_default(cfg)
        return SchoolTemplate(cfg)

    # 3. 默认模板兜底
    default_path = SCHOOLS_DIR / f"template_{DEFAULT_TEMPLATE_NAME}.json"
    if not default_path.exists():
        raise FileNotFoundError(
            f"未找到学校配置: {cfg_path}，且默认模板 {default_path} 也不存在"
        )
    with open(default_path, "r", encoding="utf-8") as f:
        cfg = json.load(f)
    # 用用户传入的 school_name 替换默认显示名（保留功能等价）
    cfg["school_name"] = school_name if school_name else cfg.get("school_name", "默认")
    return SchoolTemplate(cfg)


def list_supported_schools() -> List[Dict[str, str]]:
    """列出所有当前可用学校配置。

    返回：[{"id": "tsinghua", "school_name": "清华大学", "source": "file"}, ...]
    """
    result: List[Dict[str, str]] = []
    seen: set = set()

    # 文件来源
    if SCHOOLS_DIR.exists():
        for p in sorted(SCHOOLS_DIR.glob("template_*.json")):
            stem = p.stem  # template_tsinghua
            sid = stem.replace("template_", "", 1)
            if sid == DEFAULT_TEMPLATE_NAME:
                continue
            try:
                with open(p, "r", encoding="utf-8") as f:
                    cfg = json.load(f)
                display_name = cfg.get("school_name", sid)
            except (json.JSONDecodeError, OSError):
                display_name = sid
            result.append({"id": sid, "school_name": display_name, "source": "file"})
            seen.add(sid)

    # 运行时来源
    for sid, cfg in _RUNTIME_REGISTRY.items():
        if sid in seen or sid == DEFAULT_TEMPLATE_NAME:
            continue
        result.append(
            {
                "id": sid,
                "school_name": cfg.get("school_name", sid),
                "source": "runtime",
            }
        )

    # 默认模板放第一个
    default_path = SCHOOLS_DIR / f"template_{DEFAULT_TEMPLATE_NAME}.json"
    if default_path.exists():
        try:
            with open(default_path, "r", encoding="utf-8") as f:
                cfg = json.load(f)
            display_name = cfg.get("school_name", "默认")
        except (json.JSONDecodeError, OSError):
            display_name = "默认"
        result.insert(0, {"id": "default", "school_name": display_name, "source": "file"})

    return result


def get_field_mapping(school_name: str, field: str) -> str:
    """便捷函数：直接拿一所学校的字段映射。

    等价于 `load_template(school_name).get_field(field)`。
    """
    return load_template(school_name).get_field(field)


# ============================================================
# apply_template - 将模板应用到 docx Document
# ============================================================


def _set_section_margins(section, margins: Dict[str, float]) -> None:
    """应用页边距配置到 section 对象"""
    from docx.shared import Cm

    section.top_margin = Cm(float(margins.get("top_cm", 2.54)))
    section.bottom_margin = Cm(float(margins.get("bottom_cm", 2.54)))
    section.left_margin = Cm(float(margins.get("left_cm", 2.5)))
    section.right_margin = Cm(float(margins.get("right_cm", 2.5)))


def _set_page_header(section, header_cfg: Dict[str, Any], school_name: str) -> None:
    """配置页眉。

    header_cfg 字段：
      - enabled: bool 是否启用
      - text: str 自定义页眉文字（为空时用 school_name）
      - logo: str 校徽图片路径（绝对路径或相对于 schools/ 目录）
      - align: str 对齐方式 left/center/right（默认 center）
    """
    if not header_cfg.get("enabled", False):
        return
    header = section.header
    header.is_linked_to_previous = False
    # 取已有段落或新增
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    # 清空已有 runs
    for r in list(p.runs):
        r.text = ""
    text = header_cfg.get("text") or school_name
    align = header_cfg.get("align", "center")
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)

    # 插入校徽图片（可选）
    logo = header_cfg.get("logo", "")
    if logo:
        logo_path = Path(logo)
        if not logo_path.is_absolute():
            logo_path = SCHOOLS_DIR / logo
        if logo_path.exists():
            run = p.add_run()
            run.add_picture(str(logo_path), width=None)
            # 文字 run
            text_run = p.add_run("  " + text)
        else:
            text_run = p.add_run(text)
    else:
        text_run = p.add_run(text)

    # 字体设置（页眉一般用宋体小五）
    from docx.oxml.ns import qn
    from docx.shared import Pt

    text_run.font.size = Pt(9)
    text_run.font.name = "宋体"
    rPr = text_run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:ascii"), "宋体")
    rFonts.set(qn("w:hAnsi"), "宋体")


def _set_page_footer(section, footer_cfg: Dict[str, Any]) -> None:
    """配置页脚，支持页码字段。

    footer_cfg 字段：
      - enabled: bool
      - text: str 自定义页脚文字
      - page_number: bool 是否显示页码
    """
    if not footer_cfg.get("enabled", False):
        return
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        p = footer.paragraphs[0]
    else:
        p = footer.add_paragraph()
    for r in list(p.runs):
        r.text = ""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    text = footer_cfg.get("text", "")
    if text:
        run = p.add_run(text + "    ")

    if footer_cfg.get("page_number", False):
        # 插入 PAGE 域
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        run = p.add_run()
        run.font.size = Pt(9)
        run.font.name = "宋体"
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run._element.append(fldChar_begin)
        run._element.append(instrText)
        run._element.append(fldChar_end)


def _apply_seal_placeholder(doc, seal_cfg: Dict[str, Any]) -> None:
    """在文档末尾添加一个印章占位段落（实际盖章由学校人工完成，
    这里只标注位置与大小，便于打印校对）。

    seal_cfg 字段：
      - enabled: bool
      - position: "right_bottom" / "left_bottom" / "center_bottom"
      - size_cm: float 印章直径
      - transparency: 0-100 透明度（仅文档说明用，实际不渲染）
    """
    if not seal_cfg.get("enabled", False):
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    align_map = {
        "right_bottom": WD_ALIGN_PARAGRAPH.RIGHT,
        "left_bottom": WD_ALIGN_PARAGRAPH.LEFT,
        "center_bottom": WD_ALIGN_PARAGRAPH.CENTER,
    }
    align = align_map.get(seal_cfg.get("position", "right_bottom"), WD_ALIGN_PARAGRAPH.RIGHT)
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(0)
    run = p.add_run(f"〔此处加盖学校公章 Ø{seal_cfg.get('size_cm', 4)}cm〕")
    run.font.size = Pt(9)
    run.font.name = "宋体"
    from docx.oxml.ns import qn

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:ascii"), "宋体")
    rFonts.set(qn("w:hAnsi"), "宋体")


def apply_template(
    doc,
    template: Union[SchoolTemplate, str, Dict[str, Any]],
    *,
    apply_seal: bool = True,
) -> None:
    """将学校模板配置应用到 docx Document 对象。

    会执行：
      1. 设置页边距（所有 section）
      2. 设置页眉（学校名 + 校徽）
      3. 设置页脚（文字 + 页码）
      4. （可选）添加印章占位段落

    参数：
        doc: docx.Document 对象
        template: SchoolTemplate 实例 / 学校名 / 配置字典
        apply_seal: 是否在文档末尾添加印章占位（默认 True）

    示例：

        from docx import Document
        from utils.school_template import load_template, apply_template

        doc = Document()
        apply_template(doc, "tsinghua")
    """
    if isinstance(template, str):
        tpl = load_template(template)
    elif isinstance(template, dict):
        tpl = SchoolTemplate(_merge_with_default(template))
    elif isinstance(template, SchoolTemplate):
        tpl = template
    else:
        raise TypeError(f"不支持的 template 类型: {type(template)}")

    # 1. 页边距 + 页眉 + 页脚（针对每个 section）
    for section in doc.sections:
        _set_section_margins(section, tpl.margins)
        _set_page_header(section, tpl.header, tpl.school_name)
        _set_page_footer(section, tpl.footer)

    # 2. 印章占位
    if apply_seal:
        _apply_seal_placeholder(doc, tpl.seal)


# ============================================================
# 签字栏生成辅助
# ============================================================


def build_signature_table(
    doc,
    template: Union[SchoolTemplate, str],
    *,
    date_str: str = "    年    月    日",
) -> None:
    """按学校模板生成 4 方签字栏表格。

    默认 4 行：申请人 / 辅导员 / 院系负责人 / 学校负责人。
    清华版会插入"导师"行；北大版会改为"班主任"。
    """
    if isinstance(template, str):
        tpl = load_template(template)
    else:
        tpl = template

    blocks = tpl.signature_blocks
    rows = len(blocks)
    table = doc.add_table(rows=rows, cols=2)
    table.alignment = 1  # 居中

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for i, role in enumerate(blocks):
        # 左列：角色名
        left_cell = table.cell(i, 0)
        left_cell.text = ""
        p_left = left_cell.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_left.add_run(role + "签字：")
        run.font.size = Pt(11)
        run.font.name = "宋体"
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), "宋体")

        # 右列：签字留空 + 日期
        right_cell = table.cell(i, 1)
        right_cell.text = ""
        p_right = right_cell.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run2 = p_right.add_run("                    " + date_str)
        run2.font.size = Pt(11)
        run2.font.name = "宋体"
        rPr2 = run2._element.get_or_add_rPr()
        rFonts2 = rPr2.find(qn("w:rFonts"))
        if rFonts2 is None:
            rFonts2 = OxmlElement("w:rFonts")
            rPr2.insert(0, rFonts2)
        rFonts2.set(qn("w:eastAsia"), "宋体")


# ============================================================
# CLI / 自检入口
# ============================================================


def _self_test() -> int:
    """简易自检：列出学校、加载每个学校、检查字段映射"""
    print("=" * 60)
    print("SchoolTemplate 自检")
    print("=" * 60)
    schools = list_supported_schools()
    print(f"支持学校数量: {len(schools)}")
    for s in schools:
        print(f"  - [{s['source']:8s}] {s['id']:12s} -> {s['school_name']}")

    print()
    print("字段映射对照表（gpa / rank / major / class）：")
    print(f"  {'学校':<12} {'gpa':<10} {'rank':<10} {'major':<10} {'class':<10}")
    for s in schools:
        if s["id"] == "default":
            continue
        tpl = load_template(s["id"])
        print(
            f"  {tpl.school_name:<10} {tpl.get_field('gpa'):<10} "
            f"{tpl.get_field('rank'):<10} {tpl.get_field('major'):<10} "
            f"{tpl.get_field('class'):<10}"
        )

    print()
    print("字体偏好对照表：")
    for s in schools:
        tpl = load_template(s["id"])
        print(
            f"  {tpl.school_name:<10} body={tpl.body_font():<6} "
            f"heading={tpl.heading_font():<6} title={tpl.title_size_pt()}pt"
        )

    print()
    print("申请理由字数区间：")
    for s in schools:
        tpl = load_template(s["id"])
        print(
            f"  {tpl.school_name:<10} [{tpl.apply_reason_min()}, {tpl.apply_reason_max()}]"
        )
    return 0


if __name__ == "__main__":
    import sys

    sys.exit(_self_test())
